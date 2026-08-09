from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parent
OUT = DOCS_DIR / "Bao_cao_khoa_luan_CinemaBooking_Full.docx"


BLUE = "1F4E79"
DARK = "0B2545"
MUTED = "52657A"
LIGHT_BLUE = "E8F1FA"
LIGHT_GRAY = "F4F6F8"
LIGHT_YELLOW = "FFF8E8"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 14, BLUE, 12, 6),
        ("Heading 3", 13, DARK, 10, 4),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def para(doc: Document, text: str = "", bold: bool = False, italic: bool = False, align=None, size: float | None = None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code(doc: Document, code: str, caption: str | None = None) -> None:
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(MUTED)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    add_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
    p = cell.paragraphs[0]
    for i, line in enumerate(code.strip("\n").splitlines()):
        if i:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("111827")
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    add_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)
    p.add_run("\n" + body)
    doc.add_paragraph()


def add_placeholder(doc: Document, title: str, description: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    add_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFFDF5")
    set_cell_margins(cell, top=180, bottom=180, start=180, end=180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("9A5B00")
    p.add_run("\n" + description)
    doc.add_paragraph()


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    add_table_borders(table)
    set_repeat_table_header(table.rows[0])

    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_BLUE)
        set_cell_margins(hdr[i])
        if widths:
            set_cell_width(hdr[i], widths[i])
        p = hdr[i].paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(DARK)

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_margins(cells[i])
            if widths:
                set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[i].paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
    doc.add_paragraph()


def add_figure_plan(doc: Document, no: str, name: str, purpose: str, placement: str) -> None:
    add_placeholder(
        doc,
        f"Hình {no}. {name}",
        f"Mục đích: {purpose}\nVị trí chèn đề xuất: {placement}\nGợi ý: có thể vẽ bằng draw.io, Mermaid, PlantUML hoặc chèn ảnh giao diện thật từ project.",
    )


def chapter(doc: Document, title: str) -> None:
    doc.add_page_break()
    doc.add_heading(title, level=1)


def build_cover(doc: Document) -> None:
    para(doc, "TRƯỜNG ĐẠI HỌC/CƠ SỞ ĐÀO TẠO: ........................................", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "KHOA: ......................................................................................", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("\n\n")
    para(doc, "BÁO CÁO KHÓA LUẬN TỐT NGHIỆP", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=20)
    doc.add_paragraph()
    para(doc, "ĐỀ TÀI", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    para(doc, "XÂY DỰNG HỆ THỐNG ĐẶT VÉ XEM PHIM TRỰC TUYẾN CINEMABOOKING.VN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18)
    doc.add_paragraph("\n")
    add_matrix(
        doc,
        ["Thông tin", "Nội dung"],
        [
            ["Sinh viên thực hiện", "........................................................"],
            ["Mã sinh viên", "........................................................"],
            ["Lớp", "........................................................"],
            ["Giảng viên hướng dẫn", "........................................................"],
            ["Chuyên ngành", "Công nghệ thông tin / Kỹ thuật phần mềm"],
            ["Niên khóa", "........................................................"],
        ],
        [2600, 6760],
    )
    doc.add_paragraph("\n\n\n")
    para(doc, "TP. ................, năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)


def build_intro_pages(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("LỜI CAM ĐOAN", level=1)
    para(doc, "Em xin cam đoan báo cáo khóa luận với đề tài “Xây dựng hệ thống đặt vé xem phim trực tuyến CinemaBooking.vn” là kết quả tìm hiểu, phân tích, thiết kế và xây dựng dựa trên chính dự án đã thực hiện. Nội dung trong báo cáo được trình bày theo hiểu biết của cá nhân, bám sát mã nguồn, cơ sở dữ liệu, giao diện và các chức năng đã triển khai trong hệ thống.")
    para(doc, "Các phần mô tả giải pháp, phân tích nghiệp vụ, thiết kế cơ sở dữ liệu, thiết kế luồng xử lý và đánh giá hệ thống được viết lại bằng ngôn ngữ riêng, không sao chép nguyên văn từ tài liệu có sẵn. Những công nghệ, thư viện và nền tảng được sử dụng trong dự án đều được trình bày với mục đích giải thích vai trò trong hệ thống.")
    para(doc, "Sinh viên thực hiện\n\n................................")

    doc.add_page_break()
    doc.add_heading("LỜI CẢM ƠN", level=1)
    para(doc, "Trong quá trình thực hiện khóa luận, em đã có cơ hội hệ thống hóa kiến thức về lập trình web, thiết kế cơ sở dữ liệu, bảo mật ứng dụng, thanh toán trực tuyến và xây dựng trải nghiệm người dùng. Đề tài CinemaBooking.vn không chỉ là một ứng dụng đặt vé xem phim đơn thuần mà còn là một bài thực hành tổng hợp nhiều vấn đề gần với sản phẩm thực tế như giữ ghế đồng thời, phân quyền theo vai trò, xác thực bằng JWT, thanh toán qua cổng, gửi email vé điện tử, soát vé QR và cập nhật trạng thái ghế theo thời gian thực.")
    para(doc, "Em xin gửi lời cảm ơn đến giảng viên hướng dẫn đã định hướng, góp ý và hỗ trợ em trong quá trình hoàn thiện đề tài. Em cũng cảm ơn gia đình, bạn bè và các nguồn tài liệu kỹ thuật chính thống đã giúp em có thêm kiến thức để hoàn thiện sản phẩm.")

    doc.add_page_break()
    doc.add_heading("TÓM TẮT ĐỀ TÀI", level=1)
    para(doc, "Đề tài tập trung xây dựng hệ thống đặt vé xem phim trực tuyến CinemaBooking.vn với hai phần chính: backend Spring Boot và frontend React. Hệ thống cho phép khách hàng xem phim, xem rạp, chọn suất chiếu, giữ ghế, áp dụng mã giảm giá, thanh toán qua VNPay hoặc SePay/VietQR, nhận vé điện tử qua email và sử dụng mã QR để vào rạp. Bên cạnh đó, hệ thống cung cấp cổng quản trị cho admin và nhân viên rạp, hỗ trợ quản lý phim, rạp, phòng, ghế, suất chiếu, đơn đặt vé, thanh toán, khuyến mãi, người dùng, nhật ký thao tác và soát vé QR.")
    para(doc, "Điểm nổi bật của hệ thống là luồng giữ ghế có timeout, cập nhật trạng thái ghế realtime bằng WebSocket, phân quyền RBAC kết hợp giới hạn phạm vi nhân viên theo rạp phụ trách, xác thực JWT với refresh token rotation, gửi email xác thực tài khoản và vé điện tử, tích hợp thanh toán VNPay và SePay/VietQR có cơ chế callback/webhook, cùng các scheduler xử lý ghế/booking hết hạn. Cơ sở dữ liệu được quản lý bằng PostgreSQL và Flyway, có index phục vụ các truy vấn thường dùng, đồng thời có audit log để truy vết thao tác quan trọng.")
    add_matrix(
        doc,
        ["Từ khóa", "Nội dung"],
        [
            ["Cinema booking", "Đặt vé xem phim trực tuyến, chọn ghế, mua vé, soát vé."],
            ["RBAC", "Phân quyền theo vai trò ADMIN, STAFF, USER và permission chi tiết."],
            ["JWT", "Xác thực API bằng access token, refresh token, logout và revoke session."],
            ["Realtime", "Cập nhật seat map qua WebSocket/STOMP."],
            ["Payment", "Thanh toán VNPay, SePay/VietQR, callback/webhook, payment event."],
            ["QR ticket", "Sinh QR cho từng vé và check-in bằng camera hoặc ảnh."],
        ],
        [2000, 7360],
    )

    doc.add_page_break()
    doc.add_heading("MỤC LỤC", level=1)
    para(doc, "Ghi chú: Đây là bản Word có cấu trúc heading đầy đủ. Trong Microsoft Word, có thể vào References -> Table of Contents -> Automatic Table để tạo mục lục tự động.")
    add_bullets(
        doc,
        [
            "Chương 1. Tổng quan đề tài",
            "Chương 2. Cơ sở lý thuyết và công nghệ sử dụng",
            "Chương 3. Phân tích yêu cầu hệ thống",
            "Chương 4. Thiết kế hệ thống",
            "Chương 5. Xây dựng và triển khai hệ thống",
            "Chương 6. Kiểm thử và đánh giá",
            "Chương 7. Kết luận và hướng phát triển",
            "Phụ lục. Danh sách sơ đồ, giao diện, API, tài khoản test và checklist bảo vệ",
        ],
    )

    doc.add_page_break()
    doc.add_heading("DANH MỤC HÌNH VÀ SƠ ĐỒ ĐỀ XUẤT", level=1)
    add_matrix(
        doc,
        ["Mã hình", "Tên sơ đồ/giao diện", "Mục đích"],
        [
            ["Hình 1", "Sơ đồ ngữ cảnh hệ thống", "Xác định actor và hệ thống ngoài tương tác với CinemaBooking.vn."],
            ["Hình 2", "Sơ đồ BFD", "Phân rã chức năng theo nhóm nghiệp vụ."],
            ["Hình 3", "Use Case tổng quát", "Mô tả chức năng theo ADMIN, STAFF, USER."],
            ["Hình 4", "DFD Level 0", "Mô tả luồng dữ liệu tổng thể."],
            ["Hình 5", "DFD Level 1 - Đặt vé và thanh toán", "Mô tả xử lý chi tiết từ giữ ghế đến payment."],
            ["Hình 6", "ERD", "Mô hình dữ liệu và quan hệ bảng."],
            ["Hình 7", "Sơ đồ lớp nghiệp vụ booking", "Mô tả class/entity chính."],
            ["Hình 8", "State diagram booking/payment/seat/ticket", "Mô tả vòng đời trạng thái."],
            ["Hình 9", "Sequence login JWT", "Mô tả đăng nhập và refresh token."],
            ["Hình 10", "Sequence đặt vé realtime", "Mô tả giữ ghế và WebSocket."],
            ["Hình 11", "Sequence VNPay", "Mô tả redirect, callback và xác nhận thanh toán."],
            ["Hình 12", "Sequence SePay/VietQR", "Mô tả QR, webhook và xác nhận tự động."],
            ["Hình 13", "Sequence soát vé QR", "Mô tả staff check-in đúng rạp/suất."],
            ["Hình 14", "Deployment diagram", "Mô tả frontend, backend, database và dịch vụ ngoài."],
            ["Hình 15+", "Ảnh giao diện", "Homepage, chọn ghế, checkout, vé của tôi, admin, staff scanner."],
        ],
        [1300, 3400, 4660],
    )


def chapter_1(doc: Document) -> None:
    chapter(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI")
    doc.add_heading("1.1. Lý do chọn đề tài", level=2)
    para(doc, "Trong bối cảnh các dịch vụ giải trí ngày càng được số hóa, nhu cầu đặt vé xem phim trực tuyến trở nên phổ biến. Người dùng mong muốn có thể xem thông tin phim, chọn rạp gần mình, xem lịch chiếu, chọn ghế và thanh toán ngay trên điện thoại hoặc máy tính mà không cần đến quầy. Đối với rạp phim, việc vận hành bán vé trực tuyến giúp giảm tải cho nhân viên, hạn chế sai sót khi bán vé thủ công, tăng khả năng quản lý doanh thu và hỗ trợ chăm sóc khách hàng tốt hơn.")
    para(doc, "Tuy nhiên, một hệ thống đặt vé xem phim không chỉ dừng ở các chức năng CRUD cơ bản. Bài toán thực tế có nhiều điểm phức tạp: nhiều người có thể cùng xem một sơ đồ ghế; ghế cần được giữ tạm trong một khoảng thời gian; thanh toán có thể thành công, thất bại hoặc bị bỏ dở; vé cần được sinh QR để soát tại rạp; nhân viên chỉ nên thao tác trong phạm vi rạp được phân công; admin cần theo dõi doanh thu, booking, thanh toán và nhật ký hệ thống. Vì vậy, đề tài CinemaBooking.vn được lựa chọn nhằm xây dựng một hệ thống gần với sản phẩm thực tế, có đầy đủ các luồng nghiệp vụ quan trọng của đặt vé xem phim trực tuyến.")

    doc.add_heading("1.2. Mục tiêu của đề tài", level=2)
    add_bullets(
        doc,
        [
            "Xây dựng hệ thống đặt vé xem phim trực tuyến với giao diện thân thiện cho khách hàng, nhân viên và quản trị viên.",
            "Cho phép người dùng xem phim, rạp, lịch chiếu, chọn ghế theo sơ đồ và thanh toán trực tuyến.",
            "Thiết kế luồng giữ ghế có thời hạn để hạn chế tình trạng đặt trùng ghế khi nhiều người thao tác đồng thời.",
            "Tích hợp thanh toán VNPay và SePay/VietQR, hỗ trợ callback/webhook để cập nhật trạng thái giao dịch.",
            "Sinh vé điện tử kèm QR cho từng ghế sau khi thanh toán thành công.",
            "Xây dựng chức năng soát vé QR cho nhân viên, có kiểm tra đúng rạp, đúng suất chiếu và cửa sổ check-in.",
            "Áp dụng phân quyền RBAC, đồng thời giới hạn nhân viên theo rạp phụ trách.",
            "Cung cấp trang quản trị để quản lý phim, rạp, phòng, ghế, suất chiếu, booking, thanh toán, khuyến mãi, người dùng và audit log.",
            "Tối ưu hệ thống bằng cache, index, phân trang, tránh N+1 query, scheduler xử lý dữ liệu hết hạn và WebSocket cập nhật realtime.",
        ],
    )

    doc.add_heading("1.3. Phạm vi đề tài", level=2)
    add_matrix(
        doc,
        ["Nhóm phạm vi", "Nội dung thực hiện"],
        [
            ["Khách hàng", "Đăng ký, xác thực email, đăng nhập, Google login, xem phim/rạp/lịch chiếu, chọn ghế, áp mã giảm giá, thanh toán, nhận vé, xem vé của tôi, cập nhật hồ sơ."],
            ["Nhân viên", "Xem rạp phụ trách, xem dữ liệu vận hành theo scope, tạo/sửa/hủy suất chiếu trong rạp được gán, soát vé QR bằng camera hoặc ảnh."],
            ["Quản trị viên", "Quản lý dữ liệu master, quản lý user/role, gán staff theo rạp, theo dõi dashboard, payment, audit log."],
            ["Thanh toán", "VNPay sandbox và SePay/VietQR có webhook xác nhận tự động; MoMo có gateway scaffold nhưng đang tắt khi chưa có credential phù hợp."],
            ["Realtime", "Cập nhật trạng thái ghế HOLD/BOOKED/AVAILABLE qua WebSocket cho các client đang xem cùng suất chiếu."],
            ["Báo cáo", "Tập trung vào phân tích, thiết kế, xây dựng, kiểm thử và đánh giá hệ thống đặt vé xem phim."],
        ],
        [2200, 7160],
    )

    doc.add_heading("1.4. Phương pháp thực hiện", level=2)
    para(doc, "Đề tài được thực hiện theo hướng phân tích nghiệp vụ trước, sau đó thiết kế cơ sở dữ liệu, xây dựng backend, xây dựng frontend, tích hợp các dịch vụ ngoài và cuối cùng kiểm thử luồng nghiệp vụ. Trong quá trình phát triển, hệ thống được cải tiến dần theo các vấn đề thực tế phát sinh: lỗi giữ ghế hết hạn, cập nhật realtime chưa mượt, phân quyền staff theo rạp, giao diện thanh toán QR, email vé, lỗi callback thanh toán khi tắt ngrok, tối ưu query và index.")
    add_numbered(
        doc,
        [
            "Khảo sát nghiệp vụ đặt vé xem phim trực tuyến và xác định actor chính.",
            "Phân tích chức năng, phi chức năng và các trạng thái nghiệp vụ quan trọng.",
            "Thiết kế kiến trúc backend/frontend/database theo hướng tách lớp.",
            "Xây dựng API REST, WebSocket, scheduler, thanh toán, email và admin portal.",
            "Kiểm thử bằng unit test, integration test, test thủ công và các tình huống demo bảo vệ.",
            "Tối ưu trải nghiệm người dùng, phân quyền, query, cache, index và luồng lỗi.",
        ],
    )

    doc.add_heading("1.5. Đóng góp chính của hệ thống", level=2)
    add_callout(
        doc,
        "Điểm nhấn của khóa luận",
        "Hệ thống không chỉ là CRUD phim/rạp/suất chiếu. Phần có giá trị nhất nằm ở các luồng gần product thật: giữ ghế chống đặt trùng, realtime seat map, payment callback/webhook, QR ticket, staff scope, refresh token rotation, audit log và scheduler.",
        LIGHT_YELLOW,
    )
    add_bullets(
        doc,
        [
            "Mô hình dữ liệu tương đối đầy đủ cho một hệ thống cinema booking.",
            "Luồng booking có kiểm soát trạng thái ghế, booking, payment và ticket.",
            "Cơ chế phân quyền kết hợp permission-based access control và staff cinema assignment.",
            "Thanh toán hỗ trợ cả cổng redirect và QR ngân hàng có webhook.",
            "Hệ thống có khả năng phục hồi trạng thái khi người dùng bỏ ngang thanh toán.",
            "Frontend có các màn hình cho khách hàng, nhân viên và admin, bám sát nghiệp vụ thực tế.",
        ],
    )

    doc.add_heading("1.6. Bố cục báo cáo", level=2)
    para(doc, "Báo cáo gồm bảy chương. Chương 1 trình bày tổng quan đề tài. Chương 2 giới thiệu cơ sở lý thuyết và công nghệ. Chương 3 phân tích yêu cầu hệ thống. Chương 4 trình bày thiết kế hệ thống. Chương 5 mô tả quá trình xây dựng và triển khai. Chương 6 trình bày kiểm thử và đánh giá. Chương 7 kết luận và hướng phát triển.")


def chapter_2(doc: Document) -> None:
    chapter(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG")
    doc.add_heading("2.1. Tổng quan hệ thống đặt vé xem phim trực tuyến", level=2)
    para(doc, "Một hệ thống đặt vé xem phim trực tuyến gồm nhiều thành phần phối hợp: quản lý phim, quản lý rạp/phòng/ghế, quản lý suất chiếu, đặt vé, giữ ghế, thanh toán, phát hành vé điện tử, soát vé và quản trị vận hành. Đặc điểm nổi bật của loại hệ thống này là dữ liệu thay đổi theo thời gian thực, đặc biệt ở sơ đồ ghế. Một ghế có thể đang trống ở thời điểm này nhưng bị giữ hoặc bán ở thời điểm tiếp theo.")
    para(doc, "Vì vậy, hệ thống cần xử lý tốt tính nhất quán dữ liệu. Khi nhiều người cùng chọn một ghế, backend phải đảm bảo chỉ một người giữ được ghế. Khi người dùng bỏ ngang thanh toán, ghế phải được trả lại. Khi thanh toán thành công, ghế phải chuyển sang đã đặt và vé QR phải được sinh ra. Những yêu cầu này khiến bài toán đặt vé xem phim trở thành ví dụ tốt để học transaction, locking, scheduler, realtime và tích hợp payment gateway.")

    doc.add_heading("2.2. Kiến trúc client-server", level=2)
    para(doc, "CinemaBooking.vn sử dụng kiến trúc client-server. Frontend React chạy trên trình duyệt, chịu trách nhiệm hiển thị giao diện và tương tác với người dùng. Backend Spring Boot cung cấp REST API, xử lý nghiệp vụ, bảo mật, truy vấn cơ sở dữ liệu và tích hợp dịch vụ ngoài. PostgreSQL lưu trữ dữ liệu bền vững. Các dịch vụ như VNPay, SePay, Google OAuth và SMTP Mail được xem là hệ thống ngoài.")
    add_figure_plan(doc, "2.1", "Sơ đồ kiến trúc tổng thể", "Minh họa React Client, Spring Boot API, PostgreSQL, VNPay, SePay, Google OAuth, SMTP Mail và WebSocket.", "Chương 2 hoặc Chương 4, ngay sau phần kiến trúc.")
    add_code(
        doc,
        """
React/Vite Client
  -> REST API /api/v1/...
  -> WebSocket /ws hoặc /ws-native
  -> Spring Boot Backend
       -> PostgreSQL
       -> VNPay / SePay / Google / SMTP
        """,
        "Mô tả ngắn kiến trúc client-server của hệ thống",
    )

    doc.add_heading("2.3. Spring Boot và mô hình phân lớp", level=2)
    para(doc, "Backend được xây dựng bằng Spring Boot 3.5.14 và Java 21. Spring Boot giúp giảm cấu hình thủ công, tích hợp sẵn web server, Spring MVC, Spring Security, Spring Data JPA, validation, mail, WebSocket, cache và testing. Mã nguồn backend được tổ chức theo mô hình phân lớp: Controller nhận request, Service xử lý nghiệp vụ, Repository truy vấn dữ liệu, Mapper chuyển đổi Entity/DTO, Entity ánh xạ bảng database.")
    add_matrix(
        doc,
        ["Lớp", "Vai trò", "Ví dụ trong hệ thống"],
        [
            ["Controller", "Nhận HTTP request, validate đầu vào cơ bản, gọi service và trả ApiResponse.", "BookingController, PaymentController, ShowtimeController, AuthenticationController."],
            ["Service", "Xử lý nghiệp vụ chính, transaction, phân quyền phạm vi, gọi repository/gateway/email/websocket.", "BookingServiceImpl, PaymentServiceImpl, ShowtimeServiceImpl, UserServiceImpl."],
            ["Repository", "Truy vấn dữ liệu bằng Spring Data JPA, JPQL/native query, projection.", "BookingRepository, SeatStatusRepository, PaymentRepository."],
            ["Mapper", "Chuyển Entity sang DTO response và request sang Entity.", "BookingMapper, PaymentMapper, ShowtimeMapper."],
            ["Entity", "Đại diện bảng dữ liệu trong JPA.", "Booking, Payment, Ticket, Showtime, SeatStatus, User."],
            ["DTO", "Dữ liệu vào/ra API, tránh expose trực tiếp entity.", "CreateBookingRequest, BookingResponse, TicketCheckInRequest."],
        ],
        [1800, 4700, 2860],
    )

    doc.add_heading("2.4. React, TypeScript và Vite", level=2)
    para(doc, "Frontend sử dụng React 19, TypeScript và Vite. React phù hợp để xây dựng giao diện dạng component, TypeScript giúp phát hiện lỗi kiểu dữ liệu sớm hơn, còn Vite hỗ trợ môi trường dev nhanh và build production gọn. Frontend được chia theo nhóm trang: public, user, admin, staff; các API được tách theo domain trong thư mục src/api; trạng thái đăng nhập được quản lý bằng Zustand.")
    add_matrix(
        doc,
        ["Công nghệ frontend", "Vai trò trong hệ thống"],
        [
            ["React", "Xây dựng giao diện component cho homepage, chọn ghế, checkout, admin portal, scanner."],
            ["TypeScript", "Tăng độ an toàn kiểu dữ liệu khi gọi API và xử lý state."],
            ["Vite", "Dev server, proxy API local, build frontend."],
            ["React Router", "Khai báo route public/user/admin/staff và ProtectedRoute."],
            ["Zustand", "Lưu access token, user info, permissions, xử lý login/logout."],
            ["Axios", "Gọi API, gắn Bearer token, tự refresh token khi 401."],
            ["Leaflet", "Hiển thị bản đồ rạp và tọa độ rạp."],
            ["Recharts", "Hiển thị biểu đồ dashboard."],
            ["html5-qrcode", "Quét QR bằng camera hoặc file ảnh cho staff."],
        ],
        [2400, 6960],
    )

    doc.add_heading("2.5. PostgreSQL, JPA và Flyway", level=2)
    para(doc, "PostgreSQL được sử dụng làm hệ quản trị cơ sở dữ liệu chính. Hệ thống dùng UUID cho khóa chính, các khóa ngoại để đảm bảo quan hệ dữ liệu, CHECK constraint để giới hạn trạng thái hợp lệ và index để tối ưu truy vấn. Spring Data JPA/Hibernate giúp ánh xạ entity với bảng, nhưng hệ thống vẫn dùng các query tùy chỉnh khi cần tối ưu danh sách admin, dashboard, seat map và scheduler.")
    para(doc, "Flyway được dùng để quản lý phiên bản cơ sở dữ liệu. Các file migration nằm trong src/main/resources/db/migration. Khi ứng dụng khởi động, Flyway kiểm tra những migration đã chạy và tự chạy migration mới. Cách này an toàn hơn việc để Hibernate tự tạo bảng trong production, vì mỗi thay đổi schema đều có lịch sử rõ ràng.")

    doc.add_heading("2.6. JWT, refresh token và OAuth Google", level=2)
    para(doc, "JWT được sử dụng để xác thực các API protected. Sau khi đăng nhập, backend phát access token ngắn hạn chứa subject, userId, token_use và scope permission. Refresh token dùng để xin access token mới khi access token hết hạn. Refresh token được lưu trong HttpOnly cookie và lưu dạng hash ở database, có cơ chế rotation để phát hiện reuse. Ngoài đăng nhập bằng tài khoản nội bộ, hệ thống còn hỗ trợ Google login bằng Google ID token; backend xác minh issuer, audience và email_verified trước khi cấp token nội bộ.")

    doc.add_heading("2.7. RBAC và staff scope", level=2)
    para(doc, "RBAC là mô hình phân quyền dựa trên vai trò. Người dùng có role, role có permission, permission quyết định API nào được phép truy cập. CinemaBooking.vn sử dụng các role ADMIN, STAFF và USER. Điểm nâng cao là STAFF không chỉ bị giới hạn bởi permission mà còn bị giới hạn theo rạp được phân công thông qua bảng staff_cinemas và StaffCinemaScopeService.")

    doc.add_heading("2.8. WebSocket realtime", level=2)
    para(doc, "WebSocket được dùng để cập nhật sơ đồ ghế theo thời gian thực. Khi một người giữ ghế, thanh toán thành công, hủy booking hoặc ghế hết hạn giữ, backend publish SeatStatusEvent tới topic của suất chiếu. Các client đang xem cùng seat map sẽ nhận event và cập nhật màu ghế mà không cần refresh trang.")

    doc.add_heading("2.9. Thanh toán trực tuyến", level=2)
    para(doc, "Hệ thống hỗ trợ hai luồng thanh toán chính. VNPay là cổng thanh toán dạng redirect: backend tạo payment URL, người dùng chuyển sang trang VNPay và VNPay gọi callback về backend. SePay/VietQR là luồng QR ngân hàng: backend tạo QR có nội dung chuyển khoản, người dùng chuyển khoản bằng app ngân hàng, SePay gửi webhook khi tiền vào tài khoản để hệ thống tự xác nhận booking.")

    doc.add_heading("2.10. QR ticket và soát vé", level=2)
    para(doc, "Sau khi booking thành công, mỗi ghế có một ticket riêng kèm QR code. Khi staff quét QR, backend kiểm tra QR hợp lệ, vé còn active, booking đã thanh toán thành công, đúng rạp, đúng suất chiếu và nằm trong cửa sổ check-in. Nếu hợp lệ, ticket chuyển sang USED và lưu check_in_time cùng nhân viên quét.")

    doc.add_heading("2.11. Scheduler, cache, index và audit", level=2)
    para(doc, "Scheduler giúp hệ thống tự xử lý những việc không phụ thuộc thao tác trực tiếp của người dùng: trả ghế hết hạn, expire booking chờ thanh toán, đồng bộ trạng thái suất chiếu, dọn token hết hạn. Cache bằng Caffeine được dùng cho dữ liệu ít thay đổi như phim, rạp, phòng, ghế và khuyến mãi. Index trong PostgreSQL được thiết kế theo các truy vấn phổ biến để giảm scan bảng. Audit log và payment event giúp truy vết thao tác quan trọng, hỗ trợ debug và đối soát.")


def chapter_3(doc: Document) -> None:
    chapter(doc, "CHƯƠNG 3. PHÂN TÍCH YÊU CẦU HỆ THỐNG")
    doc.add_heading("3.1. Mô tả bài toán", level=2)
    para(doc, "Bài toán đặt ra là xây dựng một nền tảng đặt vé xem phim trực tuyến cho phép khách hàng chủ động tìm phim, chọn rạp, chọn suất chiếu, chọn ghế, thanh toán và nhận vé điện tử. Hệ thống cũng cần hỗ trợ nhân viên rạp soát vé QR và hỗ trợ quản trị viên vận hành dữ liệu phim/rạp/suất chiếu, theo dõi booking, thanh toán, người dùng và báo cáo doanh thu.")
    para(doc, "Trong môi trường thực tế, nhiều người dùng có thể cùng mua vé cho một suất chiếu. Nếu hệ thống không kiểm soát đồng thời, hai người có thể cùng đặt một ghế. Nếu người dùng giữ ghế rồi không thanh toán, ghế sẽ bị khóa vô ích và ảnh hưởng doanh thu. Nếu nhân viên quét nhầm vé ở rạp khác hoặc suất khác, vé có thể bị dùng sai. Vì vậy, hệ thống cần phân tích rõ các trạng thái và quy tắc nghiệp vụ.")

    doc.add_heading("3.2. Tác nhân hệ thống", level=2)
    add_matrix(
        doc,
        ["Tác nhân", "Mô tả", "Mục tiêu sử dụng"],
        [
            ["Khách hàng (USER)", "Người dùng đặt vé xem phim trực tuyến.", "Xem phim/rạp, chọn ghế, thanh toán, nhận vé, xem lịch sử vé."],
            ["Nhân viên rạp (STAFF)", "Người vận hành tại một hoặc nhiều rạp được admin phân công.", "Soát vé QR, quản lý suất chiếu trong phạm vi rạp, hỗ trợ booking/payment theo rạp."],
            ["Quản trị viên (ADMIN)", "Người quản trị toàn hệ thống.", "Quản lý dữ liệu, phân quyền, xem dashboard, audit, xử lý vận hành."],
            ["Google OAuth", "Dịch vụ xác thực ngoài.", "Xác minh danh tính khi người dùng đăng nhập bằng Google."],
            ["VNPay", "Cổng thanh toán redirect.", "Xử lý thanh toán và trả callback."],
            ["SePay/VietQR", "Dịch vụ xác nhận chuyển khoản ngân hàng.", "Tạo QR chuyển khoản và gửi webhook khi nhận tiền."],
            ["SMTP Mail", "Dịch vụ gửi email.", "Gửi xác thực tài khoản, reset mật khẩu, vé điện tử, thông báo hủy suất."],
        ],
        [1800, 3900, 3660],
    )
    add_figure_plan(doc, "3.1", "Sơ đồ ngữ cảnh hệ thống", "Thể hiện CinemaBooking.vn nằm giữa người dùng, nhân viên, admin và các hệ thống ngoài.", "Đầu Chương 3 sau phần tác nhân.")

    doc.add_heading("3.3. Sơ đồ phân rã chức năng BFD", level=2)
    para(doc, "Sơ đồ BFD nên chia hệ thống thành các nhóm chức năng lớn để người đọc thấy phạm vi nghiệp vụ. Với CinemaBooking.vn, các chức năng chính gồm quản lý tài khoản, quản lý dữ liệu rạp phim, đặt vé, thanh toán, phát hành vé, soát vé, quản trị và báo cáo.")
    add_code(
        doc,
        """
CinemaBooking.vn
├─ Quản lý tài khoản
│  ├─ Đăng ký, xác thực email
│  ├─ Đăng nhập password/Google
│  ├─ Refresh token, logout, quản lý phiên
│  └─ Hồ sơ cá nhân, đổi mật khẩu
├─ Quản lý phim - rạp - phòng - ghế
│  ├─ Phim chiếu rạp
│  ├─ Rạp và tọa độ bản đồ
│  ├─ Phòng chiếu
│  └─ Sinh/quản lý ghế
├─ Quản lý suất chiếu
│  ├─ Tạo/sửa/hủy suất
│  ├─ Đồng bộ UPCOMING/ONGOING/ENDED
│  └─ Kiểm tra trùng lịch phòng
├─ Đặt vé
│  ├─ Chọn suất
│  ├─ Giữ ghế
│  ├─ Áp mã giảm giá
│  └─ Tạo booking
├─ Thanh toán
│  ├─ VNPay
│  ├─ SePay/VietQR
│  ├─ Webhook/callback
│  └─ Payment event/đối soát
├─ Vé điện tử và soát vé
│  ├─ Sinh QR
│  ├─ Gửi email vé
│  └─ Staff check-in đúng rạp/suất
└─ Quản trị
   ├─ User/role/permission/staff scope
   ├─ Dashboard
   ├─ Audit log
   └─ Báo cáo vận hành
        """,
        "BFD dạng cây để chuyển thành hình trong báo cáo",
    )
    add_figure_plan(doc, "3.2", "Sơ đồ BFD", "Phân rã chức năng hệ thống thành các nhóm chính và chức năng con.", "Chương 3.")

    doc.add_heading("3.4. Yêu cầu chức năng", level=2)
    add_matrix(
        doc,
        ["Mã", "Chức năng", "Mô tả"],
        [
            ["FR-01", "Đăng ký tài khoản", "Người dùng đăng ký bằng username, email, mật khẩu và xác thực email trước khi đăng nhập."],
            ["FR-02", "Đăng nhập", "Hỗ trợ đăng nhập bằng username/password và Google ID token."],
            ["FR-03", "Xem phim", "Khách có thể xem danh sách phim, chi tiết phim, trailer, độ tuổi, thể loại và lịch chiếu."],
            ["FR-04", "Xem rạp", "Khách có thể xem rạp theo thành phố, xem bản đồ, địa chỉ và lịch chiếu của rạp."],
            ["FR-05", "Chọn ghế", "Người dùng đã đăng nhập có thể chọn ghế theo sơ đồ phòng."],
            ["FR-06", "Giữ ghế", "Ghế được chuyển sang HOLD trong thời gian cấu hình, ngăn người khác đặt trùng."],
            ["FR-07", "Tạo booking", "Sau khi giữ ghế, hệ thống tạo booking PENDING với chi tiết ghế và tổng tiền."],
            ["FR-08", "Áp mã giảm giá", "Người dùng có thể áp hoặc gỡ mã giảm giá trước khi tạo mã thanh toán cố định."],
            ["FR-09", "Thanh toán VNPay", "Hệ thống tạo URL VNPay, xử lý callback và cập nhật booking/payment."],
            ["FR-10", "Thanh toán SePay/VietQR", "Hệ thống tạo QR chuyển khoản, nhận webhook và xác nhận tự động."],
            ["FR-11", "Sinh vé QR", "Sau khi thanh toán thành công, hệ thống sinh ticket QR cho từng ghế."],
            ["FR-12", "Gửi email", "Gửi email xác thực, reset mật khẩu, vé điện tử và thông báo hủy suất."],
            ["FR-13", "Vé của tôi", "Người dùng xem vé hợp lệ, đơn chờ thanh toán, thất bại, hết hạn hoặc đã hủy."],
            ["FR-14", "Soát vé QR", "Staff quét QR bằng camera/file, kiểm tra đúng rạp/suất và đổi ticket sang USED."],
            ["FR-15", "Quản lý phim", "Admin tạo/sửa/xóa mềm phim, lọc và phân trang."],
            ["FR-16", "Quản lý rạp", "Admin tạo/sửa/xóa mềm rạp, quản lý tọa độ map và thành phố."],
            ["FR-17", "Quản lý phòng/ghế", "Admin/staff có quyền quản lý phòng, sinh ghế theo layout và cập nhật loại ghế."],
            ["FR-18", "Quản lý suất chiếu", "Admin/staff tạo/sửa/hủy suất, kiểm tra overlap, trạng thái và scope rạp."],
            ["FR-19", "Quản lý booking/payment", "Admin/staff xem booking/payment theo quyền và scope rạp."],
            ["FR-20", "Quản lý người dùng", "Admin tạo/sửa/khóa user, set role, gán rạp cho staff."],
            ["FR-21", "Dashboard/audit", "Admin/staff xem thống kê và admin xem nhật ký thao tác."],
        ],
        [900, 2400, 6060],
    )

    doc.add_heading("3.5. Yêu cầu phi chức năng", level=2)
    add_matrix(
        doc,
        ["Nhóm", "Yêu cầu", "Cách hệ thống đáp ứng"],
        [
            ["Bảo mật", "API protected phải có token và permission.", "Spring Security, JWT Resource Server, @PreAuthorize, RBAC."],
            ["Bảo mật phiên", "Refresh token không nên lưu raw và phải có khả năng revoke.", "Lưu token_hash, refresh token rotation, invalidated_token, revoke session."],
            ["Nhất quán dữ liệu", "Không được đặt trùng ghế.", "Transaction, lock khi giữ ghế, unique seat_showtime, status check."],
            ["Realtime", "Nhiều client xem cùng suất phải thấy trạng thái ghế mới.", "WebSocket/STOMP topic theo showtime."],
            ["Khả dụng", "User bỏ thanh toán thì ghế tự trả.", "HoldExpireScheduler và PendingBookingExpireScheduler."],
            ["Hiệu năng", "Danh sách admin phải phân trang, query lớn có index.", "Pageable, index, projection, cache master data."],
            ["Mở rộng", "Có thể thêm cổng thanh toán mới.", "PaymentGateway interface và enum PaymentMethod."],
            ["Dễ bảo trì", "Tách controller/service/repository/mapper/DTO.", "Kiến trúc phân lớp rõ ràng."],
            ["Truy vết", "Biết ai thao tác gì khi xảy ra sự cố.", "AdminAuditLog, AuthAuditLog, PaymentEvent."],
            ["Trải nghiệm", "Giao diện rõ ràng, responsive, lỗi tiếng Việt.", "React, Tailwind, toast, form validation, filter/search."],
        ],
        [1700, 3500, 4160],
    )

    doc.add_heading("3.6. Use Case tổng quát", level=2)
    add_figure_plan(doc, "3.3", "Use Case tổng quát", "Mô tả các chức năng của USER, STAFF, ADMIN và actor ngoài.", "Chương 3 sau bảng yêu cầu chức năng.")
    add_matrix(
        doc,
        ["Actor", "Use case chính"],
        [
            ["USER", "Đăng ký, xác thực email, đăng nhập, xem phim/rạp, chọn ghế, đặt vé, áp mã giảm giá, thanh toán, xem vé, cập nhật hồ sơ, đổi mật khẩu."],
            ["STAFF", "Xem rạp phụ trách, xem suất/booking/payment trong scope, tạo/sửa/hủy suất, soát QR, xem dashboard vận hành."],
            ["ADMIN", "Quản lý phim/rạp/phòng/ghế/suất/booking/payment/promotion/user, gán rạp cho staff, xem audit/dashboard."],
            ["Payment Gateway", "Nhận yêu cầu thanh toán, xử lý giao dịch, trả callback/webhook."],
            ["Mail Service", "Gửi email xác thực, reset mật khẩu, vé điện tử, thông báo hủy suất."],
        ],
        [1800, 7560],
    )

    doc.add_heading("3.7. DFD Level 0 và Level 1", level=2)
    para(doc, "DFD Level 0 nên thể hiện các luồng dữ liệu chính giữa người dùng, hệ thống và kho dữ liệu. DFD Level 1 nên tách sâu các xử lý quan trọng như đặt vé/thanh toán và soát vé.")
    add_figure_plan(doc, "3.4", "DFD Level 0", "Thể hiện dữ liệu đi từ USER/STAFF/ADMIN/payment/mail vào hệ thống và các kho dữ liệu chính.", "Chương 3.")
    add_figure_plan(doc, "3.5", "DFD Level 1 - Luồng đặt vé và thanh toán", "Thể hiện chọn ghế, giữ ghế, tạo booking, tạo payment, callback/webhook, sinh ticket.", "Chương 3.")

    doc.add_heading("3.8. Quy tắc nghiệp vụ quan trọng", level=2)
    add_bullets(
        doc,
        [
            "Chỉ được đặt vé cho suất UPCOMING và còn nằm ngoài thời gian cutoff trước giờ chiếu.",
            "Mỗi showtime phải có đủ seat_status tương ứng với toàn bộ ghế active của phòng.",
            "Ghế chỉ có thể được giữ nếu trạng thái hiện tại là AVAILABLE.",
            "Khi giữ ghế, seat_status chuyển sang HOLD, có hold_by và hold_until.",
            "Booking chỉ được tạo nếu các ghế đang HOLD bởi chính user hiện tại và chưa hết hạn.",
            "Nếu thanh toán thành công, booking SUCCESS, payment SUCCESS, seat_status BOOKED và ticket ACTIVE được sinh.",
            "Nếu thanh toán thất bại/hết hạn/hủy, ghế được trả AVAILABLE nếu vẫn đang thuộc hold của booking đó.",
            "Mỗi booking detail tương ứng một ghế và có thể có một ticket QR riêng.",
            "Staff check-in phải chọn đúng rạp/suất; vé sai rạp hoặc sai suất không được chuyển USED.",
            "Staff chỉ thao tác trên rạp được admin phân công.",
            "Suất chiếu đã có vé used không được hủy bằng policy thông thường.",
            "Dữ liệu giao dịch không xóa cứng để phục vụ lịch sử, audit và xử lý khiếu nại.",
        ],
    )


def chapter_4(doc: Document) -> None:
    chapter(doc, "CHƯƠNG 4. THIẾT KẾ HỆ THỐNG")
    doc.add_heading("4.1. Kiến trúc tổng thể", level=2)
    para(doc, "Hệ thống được thiết kế theo mô hình tách biệt frontend và backend. Frontend React giao tiếp với backend qua REST API và WebSocket. Backend Spring Boot chịu trách nhiệm xác thực, phân quyền, nghiệp vụ, thanh toán, email và truy cập database. PostgreSQL là nơi lưu trữ dữ liệu chính. Các dịch vụ bên ngoài gồm Google OAuth, VNPay, SePay/VietQR và SMTP Mail.")
    add_figure_plan(doc, "4.1", "Kiến trúc tổng thể hệ thống", "Mô tả frontend, backend, database và các dịch vụ ngoài.", "Đầu Chương 4.")

    doc.add_heading("4.2. Thiết kế module backend", level=2)
    add_matrix(
        doc,
        ["Module", "Class chính", "Nhiệm vụ"],
        [
            ["Authentication", "AuthenticationController, AuthenticationService, SecurityConfig, CustomJwtDecoder", "Đăng nhập, Google login, refresh token, logout, quản lý session."],
            ["User/Profile", "UserController, UserServiceImpl, UserRepository", "Đăng ký, xác thực email, reset mật khẩu, cập nhật hồ sơ, quản lý user."],
            ["Movie/Cinema", "MovieController, CinemaController, MovieServiceImpl, CinemaServiceImpl", "Quản lý phim, rạp, map, dữ liệu public."],
            ["Room/Seat", "RoomController, SeatController, RoomServiceImpl, SeatServiceImpl", "Quản lý phòng chiếu, sinh ghế, loại ghế."],
            ["Showtime", "ShowtimeController, ShowtimeServiceImpl, ShowtimeStatusSyncService", "Tạo/sửa/hủy suất, lọc lịch chiếu, đồng bộ trạng thái."],
            ["Booking", "BookingController, BookingServiceImpl", "Giữ ghế, tạo booking, áp mã giảm giá, expire/cancel booking."],
            ["Payment", "PaymentController, PaymentServiceImpl, PaymentGateway implementations", "Khởi tạo VNPay/SePay, callback/webhook, payment events."],
            ["Ticket", "TicketController, TicketQrCodeService, QrCodeImageService", "Sinh QR, xem vé, check-in QR."],
            ["WebSocket", "WebSocketConfig, SeatStatusPublisher, SeatStatusEvent", "Đẩy trạng thái ghế realtime."],
            ["Scheduler", "HoldExpireScheduler, PendingBookingExpireScheduler, ShowtimeStatusSyncScheduler, TokenCleanupTask", "Xử lý nền theo thời gian."],
            ["Audit", "AdminAuditLogInterceptor, AuthAuditService, PaymentEventService", "Ghi nhật ký thao tác và event thanh toán."],
        ],
        [1800, 3900, 3660],
    )
    add_figure_plan(doc, "4.2", "Component diagram backend", "Mô tả các module backend và quan hệ Controller-Service-Repository-Gateway.", "Chương 4.")

    doc.add_heading("4.3. Thiết kế module frontend", level=2)
    add_matrix(
        doc,
        ["Nhóm frontend", "File/thư mục", "Nhiệm vụ"],
        [
            ["Routing", "src/router/AppRouter.tsx", "Định nghĩa route public, user, admin, staff và lazy loading."],
            ["Layout", "PublicLayout, AuthLayout, AdminLayout", "Tổ chức khung giao diện theo nhóm người dùng."],
            ["Auth state", "src/stores/authStore.ts", "Lưu access token, user info, permissions, xử lý login/logout."],
            ["API client", "src/api/axiosClient.ts", "Gắn token, refresh token, retry request sau 401."],
            ["Public pages", "HomePage, MovieDetailPage, CinemaMapPage, CinemaDetailPage", "Xem phim/rạp/lịch chiếu không bắt buộc đăng nhập."],
            ["User pages", "SeatSelectionPage, CheckoutPage, MyBookingsPage, TicketDetailPage, ProfilePage", "Đặt vé, thanh toán, xem vé và hồ sơ."],
            ["Admin pages", "AdminDashboardPage, AdminMoviePage, AdminShowtimePage, AdminUserPage...", "Quản trị và vận hành dữ liệu."],
            ["Staff pages", "StaffAssignedCinemasPage, StaffTicketScannerPage", "Rạp phụ trách và soát vé QR."],
            ["Hooks", "useSeatWebSocket, useDebounce", "Logic tái sử dụng cho realtime và tìm kiếm."],
        ],
        [2000, 3500, 3860],
    )

    doc.add_heading("4.4. Thiết kế cơ sở dữ liệu", level=2)
    para(doc, "Cơ sở dữ liệu được thiết kế xoay quanh các nhóm bảng: người dùng/phân quyền, dữ liệu rạp phim, đặt vé/thanh toán/vé, audit và token. Các bảng nghiệp vụ quan trọng sử dụng UUID làm khóa chính, các enum trạng thái được ràng buộc bằng CHECK constraint, các quan hệ chính được bảo vệ bằng khóa ngoại và các cột truy vấn thường xuyên được đánh index.")
    add_figure_plan(doc, "4.3", "ERD hệ thống CinemaBooking.vn", "Thể hiện đầy đủ users, roles, permissions, movies, cinemas, rooms, seats, showtimes, seat_status, bookings, booking_details, payments, payment_events, tickets, staff_cinemas, audit logs, refresh_tokens.", "Chương 4.")
    add_matrix(
        doc,
        ["Nhóm bảng", "Bảng", "Ý nghĩa"],
        [
            ["RBAC", "users, roles, permissions, users_roles, roles_permissions", "Quản lý tài khoản, vai trò và quyền."],
            ["Staff scope", "staff_cinemas", "Gán nhân viên với rạp phụ trách."],
            ["Rạp phim", "movies, cinemas, rooms, seats, showtimes", "Dữ liệu phim, rạp, phòng, ghế vật lý và suất chiếu."],
            ["Ghế theo suất", "seat_status", "Trạng thái AVAILABLE/HOLD/BOOKED cho từng ghế trong từng suất."],
            ["Đặt vé", "bookings, booking_details", "Đơn đặt vé và chi tiết ghế."],
            ["Thanh toán", "payments, payment_events", "Giao dịch thanh toán và nhật ký event payment."],
            ["Vé", "tickets", "QR ticket từng ghế, trạng thái ACTIVE/USED/CANCELLED."],
            ["Bảo mật phiên", "refresh_tokens, invalidated_token", "Quản lý refresh token và blacklist access token đã logout."],
            ["Audit", "admin_audit_logs, auth_audit_logs", "Truy vết thao tác quản trị và đăng nhập."],
        ],
        [1700, 3000, 4660],
    )

    doc.add_heading("4.5. Thiết kế trạng thái nghiệp vụ", level=2)
    add_figure_plan(doc, "4.4", "State diagram Booking/Payment/Seat/Ticket/Showtime", "Mô tả vòng đời trạng thái quan trọng trong hệ thống.", "Chương 4.")
    add_matrix(
        doc,
        ["Đối tượng", "Trạng thái", "Luồng chuyển chính"],
        [
            ["SeatStatus", "AVAILABLE, HOLD, BOOKED", "AVAILABLE -> HOLD khi user giữ ghế; HOLD -> BOOKED khi thanh toán thành công; HOLD -> AVAILABLE khi hủy/thất bại/hết hạn."],
            ["Booking", "PENDING, SUCCESS, FAILED, CANCELLED, EXPIRED", "PENDING -> SUCCESS khi payment success; PENDING -> FAILED khi payment fail; PENDING -> CANCELLED khi user/admin hủy; PENDING -> EXPIRED khi quá hạn."],
            ["Payment", "PENDING, SUCCESS, FAILED, EXPIRED", "PENDING -> SUCCESS khi gateway xác nhận; PENDING -> FAILED/EXPIRED khi lỗi hoặc quá hạn."],
            ["Ticket", "ACTIVE, USED, CANCELLED", "ACTIVE -> USED khi check-in; ACTIVE -> CANCELLED khi booking/suất bị hủy."],
            ["Showtime", "UPCOMING, ONGOING, ENDED, CANCELLED", "UPCOMING -> ONGOING -> ENDED theo thời gian; hoặc CANCELLED khi admin/staff hủy."],
        ],
        [1800, 2700, 4860],
    )

    doc.add_heading("4.6. Thiết kế luồng đăng nhập và bảo mật", level=2)
    add_figure_plan(doc, "4.5", "Sequence diagram Login JWT và Refresh Token", "Mô tả LoginPage, Auth API, AuthenticationService, UserRepository, RefreshTokenRepository và axios refresh.", "Chương 4.")
    add_code(
        doc,
        """
LoginPage -> authApi.login -> POST /auth/token
AuthenticationController -> AuthenticationService.authenticate
AuthenticationService -> UserRepository.findByUsername
AuthenticationService -> BCryptPasswordEncoder.matches
AuthenticationService -> generateAccessToken + generateRefreshToken
AuthenticationController -> Set-Cookie cinema_refresh_token HttpOnly
Frontend -> authStore.login
        """,
        "Luồng tuần tự login password",
    )

    doc.add_heading("4.7. Thiết kế luồng đặt vé và giữ ghế", level=2)
    add_figure_plan(doc, "4.6", "Activity diagram đặt vé", "Mô tả từ xem lịch chiếu, chọn ghế, giữ ghế, tạo booking đến thanh toán.", "Chương 4.")
    add_figure_plan(doc, "4.7", "Sequence diagram giữ ghế realtime", "Mô tả SeatSelectionPage, BookingServiceImpl.holdSeats, SeatStatusRepository lock và SeatStatusPublisher.", "Chương 4.")
    add_code(
        doc,
        """
User chọn ghế
  -> holdSeats(showtimeId, seatIds)
  -> findForUpdateByShowtimeAndSeats
  -> kiểm tra AVAILABLE
  -> set HOLD + hold_by + hold_until
  -> afterCommit publish HOLD qua WebSocket
        """,
        "Luồng giữ ghế chống đặt trùng",
    )

    doc.add_heading("4.8. Thiết kế luồng thanh toán", level=2)
    add_figure_plan(doc, "4.8", "Sequence diagram VNPay", "Mô tả tạo payment URL, redirect, callback, verify hash, update booking/payment, sinh ticket.", "Chương 4.")
    add_figure_plan(doc, "4.9", "Sequence diagram SePay/VietQR", "Mô tả tạo QR, user chuyển khoản, SePay webhook, verify HMAC/API key, xác nhận payment.", "Chương 4.")

    doc.add_heading("4.9. Thiết kế luồng soát vé QR", level=2)
    add_figure_plan(doc, "4.10", "Sequence diagram QR check-in", "Mô tả StaffTicketScannerPage, TicketController, BookingServiceImpl.checkInTicket và TicketRepository.", "Chương 4.")
    add_code(
        doc,
        """
Staff chọn rạp/suất
  -> quét QR bằng camera/file
  -> POST /api/v1/tickets/check-in
  -> normalizeAndValidate QR
  -> tìm ticket
  -> kiểm tra ACTIVE, booking SUCCESS
  -> kiểm tra đúng cinemaId, showtimeId
  -> kiểm tra staff scope và cửa sổ check-in
  -> set USED + check_in_time + checked_in_by
        """,
        "Luồng check-in QR chuẩn product",
    )

    doc.add_heading("4.10. Thiết kế giao diện người dùng", level=2)
    para(doc, "Giao diện được thiết kế theo ba nhóm chính: giao diện khách hàng, giao diện nhân viên và giao diện admin. Giao diện khách hàng ưu tiên thao tác nhanh, rõ thông tin phim/rạp/suất/ghế/thanh toán. Giao diện admin/staff ưu tiên khả năng quét dữ liệu, lọc, phân trang, thao tác vận hành và trạng thái rõ ràng.")
    for no, name, purpose in [
        ("4.11", "Giao diện trang chủ", "Hiển thị danh sách phim, mua vé theo phim/rạp, tìm kiếm thành phố/rạp."),
        ("4.12", "Giao diện chi tiết phim", "Hiển thị thông tin phim và lịch chiếu."),
        ("4.13", "Giao diện chọn ghế", "Hiển thị sơ đồ ghế, loại ghế, ghế đã đặt/đang giữ, tổng tiền."),
        ("4.14", "Giao diện thanh toán", "Hiển thị tóm tắt booking, mã giảm giá, VNPay và Quét QR ngân hàng."),
        ("4.15", "Giao diện vé của tôi", "Hiển thị vé hợp lệ, đơn chờ thanh toán, đơn đã hủy/hết hạn."),
        ("4.16", "Giao diện admin dashboard", "Hiển thị KPI, biểu đồ doanh thu và top phim."),
        ("4.17", "Giao diện quản lý suất chiếu", "Hiển thị lọc ngày/rạp/phim/trạng thái và thao tác tạo/hủy suất."),
        ("4.18", "Giao diện staff scanner", "Chọn rạp/suất, quét QR bằng camera hoặc file."),
    ]:
        add_figure_plan(doc, no, name, purpose, "Chương 4 hoặc Chương 5, phần giao diện.")

    doc.add_heading("4.11. Thiết kế bảo mật", level=2)
    add_bullets(
        doc,
        [
            "API public được khai báo rõ trong SecurityConfig, các API còn lại yêu cầu authentication.",
            "Permission được kiểm tra bằng @PreAuthorize tại controller.",
            "Staff scope được kiểm tra ở service để tránh thao tác sai rạp.",
            "Refresh token lưu HttpOnly cookie và DB chỉ lưu hash.",
            "Access token logout được đưa vào invalidated_token tới khi hết hạn.",
            "Google ID token được verify issuer/audience/email_verified.",
            "Webhook thanh toán cần API key/HMAC để tránh request giả.",
            "Mật khẩu lưu bằng BCrypt, không lưu plain text.",
            "Validation DTO và GlobalExceptionHandler giúp tránh dữ liệu sai và lỗi thô rò ra client.",
        ],
    )

    doc.add_heading("4.12. Thiết kế triển khai", level=2)
    add_figure_plan(doc, "4.19", "Deployment diagram", "Mô tả React build, Spring Boot API, PostgreSQL, SMTP, gateway thanh toán, domain/ngrok khi demo.", "Cuối Chương 4.")
    add_matrix(
        doc,
        ["Môi trường", "Thành phần", "Ghi chú"],
        [
            ["Local dev", "Vite dev server, Spring Boot, PostgreSQL Docker", "Frontend dùng proxy BACKEND_PROXY_TARGET tới backend local."],
            ["Demo mobile", "Ngrok HTTPS -> Vite dev server -> proxy backend", "Camera mobile cần HTTPS; callback payment cần URL public."],
            ["Production đề xuất", "Static frontend, backend API, PostgreSQL managed, SMTP thật, payment keys thật", "Không dùng mock-data.sql, không commit .env."],
        ],
        [2000, 4200, 3160],
    )


def chapter_5(doc: Document) -> None:
    chapter(doc, "CHƯƠNG 5. XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG")
    doc.add_heading("5.1. Môi trường phát triển", level=2)
    add_matrix(
        doc,
        ["Nhóm", "Công nghệ", "Phiên bản/ghi chú"],
        [
            ["Backend", "Java, Spring Boot", "Java 21, Spring Boot 3.5.14."],
            ["Database", "PostgreSQL, Flyway", "Schema version bằng db/migration, Docker Compose cho local."],
            ["Security", "Spring Security, OAuth2 Resource Server, Nimbus JOSE JWT", "JWT HS512, BCrypt, refresh token rotation."],
            ["Realtime", "Spring WebSocket/STOMP", "Seat status events theo topic showtime."],
            ["Email", "Spring Mail, Thymeleaf", "Mailtrap cho dev, SMTP thật cho production."],
            ["Frontend", "React, TypeScript, Vite", "React 19, Vite 8, TypeScript 6."],
            ["UI", "Tailwind CSS, lucide-react, Recharts, Leaflet", "Giao diện responsive, bản đồ và biểu đồ."],
            ["QR", "ZXing backend, html5-qrcode frontend", "Sinh QR ticket và quét QR bằng camera/file."],
            ["Testing", "JUnit, Spring Boot Test, Testcontainers, Vitest", "Test backend có DB ảo PostgreSQL bằng Docker."],
        ],
        [1800, 3500, 4060],
    )

    doc.add_heading("5.2. Cấu trúc mã nguồn backend", level=2)
    add_code(
        doc,
        """
src/main/java/com/cinema/booking
├─ configuration
├─ controller
├─ dto/request, dto/response
├─ entity
├─ enums
├─ exception
├─ mapper
├─ payment
├─ repository
├─ security
├─ service / service/impl
├─ specification
├─ util
└─ websocket
        """,
        "Cấu trúc backend Spring Boot",
    )
    para(doc, "Cấu trúc này giúp mỗi nhóm trách nhiệm nằm ở đúng vị trí. Controller không chứa logic phức tạp; Service là nơi xử lý nghiệp vụ; Repository tập trung query; DTO giúp định nghĩa dữ liệu vào/ra; Mapper giảm lặp code chuyển đổi; Exception và ErrorCode chuẩn hóa lỗi; Payment package hỗ trợ mở rộng cổng thanh toán.")

    doc.add_heading("5.3. Cấu trúc mã nguồn frontend", level=2)
    add_code(
        doc,
        """
src
├─ api
├─ components
├─ constants
├─ hooks
├─ pages
│  ├─ admin
│  ├─ public
│  ├─ staff
│  └─ user
├─ router
├─ stores
├─ types
└─ utils
        """,
        "Cấu trúc frontend React",
    )
    para(doc, "Frontend được chia theo domain màn hình và chức năng dùng chung. Các trang public không bắt buộc đăng nhập, các trang user cần token, các trang admin/staff cần permission. API được tách thành từng file để page gọi rõ ràng và dễ bảo trì.")

    doc.add_heading("5.4. Xây dựng luồng đăng ký, xác thực email và đăng nhập", level=2)
    para(doc, "Luồng đăng ký bắt đầu từ RegisterPage, gửi request tới UserController. Backend tạo user, mã hóa mật khẩu bằng BCrypt, sinh token xác thực email, lưu hash token và gửi email xác thực. Khi người dùng bấm link, VerifyEmailPage gọi API xác thực email. Chỉ user đã emailVerified mới được AuthenticationService cho đăng nhập.")
    para(doc, "Luồng đăng nhập sử dụng AuthenticationController và AuthenticationService. Sau khi kiểm tra mật khẩu, trạng thái user và rate limit, backend sinh access token và refresh token. Refresh token được ghi vào HttpOnly cookie, access token trả về body cho frontend lưu trong authStore. Axios interceptor tự gắn access token vào các request tiếp theo.")
    add_code(
        doc,
        """
AuthenticationService.authenticate()
  -> authRateLimitService.check()
  -> userRepository.findByUsername()
  -> passwordEncoder.matches()
  -> validateUserCanAuthenticate()
  -> issueTokenPair()
  -> authAuditService.record()
        """,
        "Các bước chính trong login password",
    )

    doc.add_heading("5.5. Xây dựng RBAC và staff scope", level=2)
    para(doc, "Danh sách permission được định nghĩa trong PermissionName.java và được seed bởi ApplicationInitConfig. Role ADMIN được gán toàn bộ permission; STAFF được gán các quyền vận hành; USER được gán các quyền đặt vé và quản lý profile. Controller dùng @PreAuthorize để kiểm tra permission. Với STAFF, các service như ShowtimeServiceImpl, BookingServiceImpl và Ticket check-in gọi StaffCinemaScopeService để kiểm tra staff có được phân công rạp tương ứng không.")

    doc.add_heading("5.6. Xây dựng quản lý phim, rạp, phòng, ghế và suất chiếu", level=2)
    para(doc, "Admin có thể quản lý phim, rạp, phòng, ghế và suất chiếu. Khi tạo suất chiếu, hệ thống kiểm tra end_time phải sau start_time, kiểm tra phòng không bị overlap trong khoảng thời gian có buffer dọn phòng và tự sinh seat_status cho toàn bộ ghế active của phòng. Đây là bước quan trọng vì mỗi suất chiếu cần một sơ đồ trạng thái ghế riêng.")
    add_code(
        doc,
        """
createShowtime()
  -> validate endTime > startTime
  -> load movie, room
  -> validate staff scope by cinema
  -> check overlap with cleaning buffer
  -> save showtime
  -> clone room seats to seat_status AVAILABLE
        """,
        "Luồng tạo suất chiếu",
    )

    doc.add_heading("5.7. Xây dựng luồng giữ ghế và đặt vé", level=2)
    para(doc, "Khi người dùng chọn ghế, frontend gọi API giữ ghế. Backend dùng transaction và lock khi đọc SeatStatus để chống race condition. Nếu ghế AVAILABLE, hệ thống chuyển sang HOLD, lưu hold_by và hold_until. Sau khi transaction commit, SeatStatusPublisher gửi event WebSocket để các client khác thấy ghế đang được giữ.")
    para(doc, "Khi tạo booking, backend kiểm tra lại các ghế vẫn đang HOLD bởi đúng user hiện tại và chưa hết hạn. Sau đó hệ thống tính tổng tiền dựa trên base_price của showtime và price_multiplier của từng seat. Nếu có promotion code, hệ thống validate mã và tính discount. Booking được tạo với trạng thái PENDING và payment_expires_at.")

    doc.add_heading("5.8. Xây dựng thanh toán VNPay và SePay/VietQR", level=2)
    para(doc, "PaymentServiceImpl đóng vai trò trung tâm điều phối thanh toán. Các cổng thanh toán được đóng gói qua PaymentGateway interface, giúp hệ thống dễ mở rộng thêm gateway mới. VNPayPaymentGateway tạo URL redirect kèm chữ ký. SePayPaymentGateway tạo payload/QR chuyển khoản theo cấu hình ngân hàng. Khi callback/webhook hợp lệ, backend cập nhật payment SUCCESS, booking SUCCESS, ghế BOOKED và sinh ticket QR.")
    add_matrix(
        doc,
        ["Cổng", "Kiểu tương tác", "Điểm cần xử lý"],
        [
            ["VNPay", "Redirect sang trang thanh toán, callback về backend.", "Tạo vnp_TxnRef, ký hash, verify callback, xử lý success/fail."],
            ["SePay/VietQR", "QR chuyển khoản, webhook khi tiền vào tài khoản.", "Tạo đúng amount/content, khóa thay đổi promotion sau khi QR tạo, verify webhook API key/HMAC."],
            ["MoMo", "Gateway scaffold, mặc định tắt.", "Chỉ bật khi có credential hợp lệ."],
        ],
        [1800, 3900, 3660],
    )

    doc.add_heading("5.9. Xây dựng vé điện tử và QR check-in", level=2)
    para(doc, "Sau thanh toán thành công, hệ thống sinh một ticket cho mỗi booking detail. QR code không chỉ là mã ngẫu nhiên mà được ký bằng secret để tránh giả mạo. Staff khi quét QR phải chọn rạp và suất chiếu đang soát. Backend kiểm tra QR hợp lệ, vé ACTIVE, booking SUCCESS, đúng rạp, đúng suất, đúng cửa sổ check-in và staff có scope rạp. Nếu hợp lệ, ticket chuyển sang USED.")

    doc.add_heading("5.10. Xây dựng WebSocket realtime seat map", level=2)
    para(doc, "WebSocketConfig mở endpoint cho client kết nối. Frontend dùng hook useSeatWebSocket để subscribe topic theo showtimeId. BookingServiceImpl, HoldExpireScheduler và các luồng thanh toán gọi SeatStatusPublisher sau khi DB commit. Nhờ vậy, khi một người giữ ghế hoặc thanh toán thành công, những người khác đang xem cùng suất thấy màu ghế đổi gần như ngay lập tức.")

    doc.add_heading("5.11. Xây dựng scheduler", level=2)
    para(doc, "Các scheduler giúp hệ thống tự động xử lý những trường hợp người dùng bỏ ngang hoặc thời gian hệ thống thay đổi. HoldExpireScheduler trả ghế HOLD quá hạn. PendingBookingExpireScheduler chuyển booking PENDING quá hạn sang EXPIRED và trả ghế. ShowtimeStatusSyncScheduler cập nhật trạng thái suất chiếu theo thời gian. TokenCleanupTask dọn token hết hạn.")

    doc.add_heading("5.12. Xây dựng cache, query optimization và audit", level=2)
    para(doc, "Cache Caffeine được cấu hình trong application.yaml và CacheConfig. Các dữ liệu ít thay đổi như movies, cinemas, cinema map, rooms by cinema, seats by room và promotions được cache để giảm query. Khi admin tạo/sửa/xóa dữ liệu liên quan, @CacheEvict được dùng để xóa cache cũ. Query admin và dashboard được tối ưu bằng pagination, index, projection và fetch query để hạn chế N+1.")
    para(doc, "Audit được chia thành auth audit, admin audit và payment event. Auth audit ghi các sự kiện login/refresh/logout. Admin audit ghi thao tác quản trị. Payment event ghi toàn bộ biến động quan trọng của payment, callback, webhook, expire và refund requested.")

    doc.add_heading("5.13. Cấu hình chạy dự án", level=2)
    add_matrix(
        doc,
        ["File", "Vai trò"],
        [
            ["application.yaml", "Cấu hình datasource, Flyway, JPA, cache, mail, payment, booking timeout, showtime, JWT, ticket, logging."],
            [".env backend", "Chứa biến môi trường cho database, secret, mail, VNPay, SePay, Google, timeout."],
            [".env frontend", "Chứa VITE_API_BASE_URL, BACKEND_PROXY_TARGET, Google client id và cấu hình dev server."],
            ["docker-compose.yml", "Chạy PostgreSQL local."],
            ["database/database.sql", "Tạo database từ đầu khi reset thủ công."],
            ["database/mock-data.sql", "Seed dữ liệu demo/test nhanh."],
            ["src/main/resources/db/migration", "Nguồn migration chính của Flyway."],
        ],
        [2400, 6960],
    )


def chapter_6(doc: Document) -> None:
    chapter(doc, "CHƯƠNG 6. KIỂM THỬ VÀ ĐÁNH GIÁ")
    doc.add_heading("6.1. Chiến lược kiểm thử", level=2)
    para(doc, "Hệ thống được kiểm thử bằng cả automated test và kiểm thử thủ công. Backend có unit test cho QR code và integration test cho auth, booking, payment callback, security và user management. Integration test sử dụng Testcontainers để tạo PostgreSQL tạm thời, giúp test không ảnh hưởng database thật. Frontend có test cho các form login/register và các command lint/build để kiểm tra TypeScript, cấu trúc code và khả năng build production.")
    add_matrix(
        doc,
        ["Nhóm test", "Công cụ", "Mục tiêu"],
        [
            ["Backend unit test", "JUnit", "Kiểm tra service nhỏ như ticket QR, QR image."],
            ["Backend integration test", "Spring Boot Test, Testcontainers PostgreSQL", "Kiểm tra API và luồng nghiệp vụ có database thật tạm thời."],
            ["Security test", "Spring Security Test", "Kiểm tra quyền truy cập, user chỉ xem dữ liệu của mình, staff scope."],
            ["Payment test", "PaymentCallbackIntegrationTest, SePayPaymentGatewayTest", "Kiểm tra callback/webhook/gateway."],
            ["Frontend test", "Vitest, Testing Library", "Kiểm tra form login/register và validate UI."],
            ["Manual test", "Browser, Mailtrap, VNPay sandbox, SePay webhook, QR scanner", "Kiểm tra trải nghiệm end-to-end như người dùng thật."],
        ],
        [2200, 3400, 3760],
    )

    doc.add_heading("6.2. Danh sách test case tiêu biểu", level=2)
    add_matrix(
        doc,
        ["Mã", "Tình huống", "Kết quả mong đợi"],
        [
            ["TC-01", "Đăng ký email mới hợp lệ.", "Tạo user, gửi email xác thực, user chưa verify không đăng nhập được."],
            ["TC-02", "Xác thực email bằng token hợp lệ.", "emailVerified=true, user đăng nhập được."],
            ["TC-03", "Đăng nhập sai mật khẩu.", "Trả lỗi thân thiện, không reload trang, ghi auth audit thất bại."],
            ["TC-04", "Access token hết hạn.", "Axios gọi refresh, nhận access token mới và retry request."],
            ["TC-05", "Staff truy cập rạp không được gán.", "Backend trả 403, không cho tạo/sửa/soát vé."],
            ["TC-06", "Hai user cùng giữ một ghế.", "Chỉ một user giữ thành công, user còn lại nhận lỗi ghế không khả dụng."],
            ["TC-07", "Ghế HOLD hết hạn.", "Scheduler trả ghế AVAILABLE và WebSocket cập nhật UI."],
            ["TC-08", "Áp mã giảm giá hợp lệ.", "Tổng tiền booking giảm đúng, pending payment cũ hết hạn nếu số tiền đổi."],
            ["TC-09", "Tạo QR SePay rồi đổi mã giảm giá.", "UI khóa mã hoặc yêu cầu tạo lại QR mới, tránh thanh toán sai số tiền."],
            ["TC-10", "VNPay thanh toán thành công.", "Payment SUCCESS, Booking SUCCESS, Seat BOOKED, Ticket ACTIVE, gửi email vé."],
            ["TC-11", "VNPay thanh toán thất bại.", "Payment FAILED, Booking FAILED, ghế AVAILABLE."],
            ["TC-12", "SePay webhook hợp lệ.", "Hệ thống xác nhận payment/booking tự động và chuyển trang kết quả."],
            ["TC-13", "Webhook SePay sai secret/API key.", "Không cập nhật payment, ghi event thất bại nếu có."],
            ["TC-14", "Staff quét vé đúng rạp/suất.", "Ticket chuyển USED, lưu check_in_time và checked_in_by."],
            ["TC-15", "Staff quét vé sai rạp.", "Không check-in, báo vé không thuộc rạp này."],
            ["TC-16", "Quét lại vé đã dùng.", "Không đổi trạng thái, trả thông tin đã dùng."],
            ["TC-17", "Admin hủy suất đã có booking SUCCESS chưa check-in.", "Booking CANCELLED, ticket CANCELLED, ghi refund requested, gửi email."],
            ["TC-18", "Admin xem dashboard.", "KPI, biểu đồ và top phim hiển thị đúng theo dữ liệu."],
            ["TC-19", "Admin lọc payment theo ngày/trạng thái/cổng.", "Danh sách phân trang, không lỗi query, dữ liệu đúng filter."],
            ["TC-20", "User xem vé của chính mình.", "Chỉ thấy booking/ticket của user đang đăng nhập."],
        ],
        [900, 4300, 4160],
    )

    doc.add_heading("6.3. Kiểm thử hiệu năng và tối ưu truy vấn", level=2)
    para(doc, "Hệ thống đã được chú ý ở các điểm có nguy cơ chậm khi dữ liệu tăng: danh sách admin có phân trang, truy vấn dashboard dùng aggregation, các bảng giao dịch có index theo status/created_at/showtime/user, seat_status có index theo showtime/status/hold_until, booking/payment có index phục vụ lọc theo ngày và trạng thái. Ngoài ra, application.yaml có profile sql-profile để bật log SQL, log slow query và thống kê Hibernate khi cần profiling.")
    add_code(
        doc,
        """
spring.profiles.active=sql-profile
HIBERNATE_SLOW_QUERY_MS=50

Mục tiêu:
- xem SQL thật Hibernate tạo ra
- phát hiện query chậm
- dùng EXPLAIN ANALYZE để kiểm tra index
- điều chỉnh repository/fetch/projection nếu cần
        """,
        "Cấu hình profiling SQL đề xuất",
    )

    doc.add_heading("6.4. Đánh giá mức độ hoàn thành", level=2)
    add_matrix(
        doc,
        ["Tiêu chí", "Đánh giá"],
        [
            ["Chức năng khách hàng", "Đáp ứng đầy đủ luồng xem phim/rạp, chọn ghế, áp mã, thanh toán, xem vé và profile."],
            ["Chức năng nhân viên", "Có rạp phụ trách, quản lý vận hành theo scope và soát QR đúng rạp/suất."],
            ["Chức năng admin", "Có dashboard, CRUD dữ liệu chính, quản lý user/role/scope, booking/payment/audit."],
            ["Bảo mật", "Có JWT, refresh token rotation, RBAC, staff scope, password hash, email verification."],
            ["Thanh toán", "Có VNPay và SePay/VietQR, payment event, timeout và xử lý trạng thái."],
            ["Realtime", "Có WebSocket cập nhật seat map."],
            ["Cơ sở dữ liệu", "Schema đầy đủ, có migration Flyway, index, constraint, audit table."],
            ["UX/UI", "Giao diện public/user/admin/staff khá đầy đủ, có filter, phân trang, trạng thái rõ."],
            ["Mở rộng", "Có thể thêm gateway mới, cache provider mới, báo cáo nâng cao, refund tự động."],
        ],
        [2600, 6760],
    )

    doc.add_heading("6.5. Hạn chế hiện tại", level=2)
    add_bullets(
        doc,
        [
            "Luồng hoàn tiền hiện ở mức ghi nhận refund requested và xử lý thủ công, chưa tích hợp refund tự động với từng cổng.",
            "Chưa triển khai push notification/mobile app; thông báo chủ yếu qua email và UI.",
            "Chưa có hệ thống recommendation phim cá nhân hóa theo lịch sử xem.",
            "Chưa có module quản lý combo bắp nước, voucher phức tạp hoặc loyalty/member tier.",
            "Chưa triển khai CI/CD và observability production đầy đủ như metrics, tracing, alerting.",
            "Bản đồ hiện dùng dữ liệu rạp mẫu, chưa tích hợp tìm đường chi tiết bằng dịch vụ bản đồ production.",
            "Một số sơ đồ trong báo cáo cần được vẽ/chèn ảnh chính thức trước khi nộp bản cuối.",
        ],
    )


def chapter_7(doc: Document) -> None:
    chapter(doc, "CHƯƠNG 7. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    doc.add_heading("7.1. Kết luận", level=2)
    para(doc, "Khóa luận đã xây dựng được hệ thống CinemaBooking.vn với đầy đủ các chức năng cốt lõi của một nền tảng đặt vé xem phim trực tuyến. Hệ thống hỗ trợ khách hàng tìm phim/rạp/suất chiếu, chọn ghế, giữ ghế, thanh toán, nhận vé QR và xem lịch sử vé. Đồng thời, hệ thống cung cấp cổng quản trị cho admin và nhân viên, bao gồm quản lý dữ liệu rạp phim, quản lý suất chiếu, đơn đặt vé, thanh toán, khuyến mãi, người dùng, audit log và soát vé QR.")
    para(doc, "Về mặt kỹ thuật, đề tài đã áp dụng nhiều nội dung quan trọng của phát triển ứng dụng web hiện đại: Spring Boot phân lớp, PostgreSQL/Flyway, Spring Security/JWT, RBAC, refresh token rotation, WebSocket realtime, scheduler, cache, index, payment gateway, webhook, email template, QR code, React frontend, route guard, Axios interceptor và quản lý state bằng Zustand. Các thành phần này giúp hệ thống không chỉ chạy được mà còn có cấu trúc dễ mở rộng và gần với yêu cầu vận hành thực tế.")

    doc.add_heading("7.2. Kết quả đạt được", level=2)
    add_bullets(
        doc,
        [
            "Hoàn thiện backend API cho các nghiệp vụ chính.",
            "Hoàn thiện frontend cho khách hàng, admin và staff.",
            "Thiết kế database đầy đủ và có migration/index/constraint.",
            "Xây dựng luồng giữ ghế có timeout và realtime WebSocket.",
            "Tích hợp thanh toán VNPay và SePay/VietQR.",
            "Sinh vé QR và soát vé đúng rạp/suất.",
            "Bổ sung xác thực email, reset mật khẩu, gửi vé qua email.",
            "Bổ sung RBAC, staff scope, audit log, payment event và scheduler.",
            "Có tài liệu chạy dự án, tài liệu luồng hệ thống và checklist demo.",
        ],
    )

    doc.add_heading("7.3. Hướng phát triển", level=2)
    add_matrix(
        doc,
        ["Hướng phát triển", "Mô tả"],
        [
            ["Refund tự động", "Tích hợp API hoàn tiền của VNPay/SePay hoặc quy trình duyệt hoàn tiền nhiều bước."],
            ["Combo và sản phẩm phụ", "Bán bắp nước, combo, voucher sản phẩm kèm vé."],
            ["Loyalty", "Tích điểm, hạng thành viên, ưu đãi theo lịch sử đặt vé."],
            ["Notification", "Push notification/email/SMS nhắc lịch chiếu, thông báo hủy suất."],
            ["Mobile app", "Xây dựng app mobile hoặc PWA để tăng trải nghiệm người dùng."],
            ["Recommendation", "Gợi ý phim/rạp/suất theo vị trí và hành vi người dùng."],
            ["Monitoring", "Bổ sung metrics, tracing, alerting, dashboard kỹ thuật."],
            ["CI/CD", "Tự động test/build/deploy bằng GitHub Actions hoặc pipeline tương đương."],
            ["Redis", "Dùng Redis cho cache phân tán, rate limit phân tán hoặc pub/sub khi chạy nhiều backend instance."],
            ["Microservice", "Tách payment/ticket/notification service khi hệ thống lớn."],
        ],
        [2500, 6860],
    )

    doc.add_heading("7.4. Bài học rút ra", level=2)
    para(doc, "Thông qua đề tài, em hiểu rõ hơn rằng một hệ thống web thực tế không chỉ là giao diện đẹp và các API CRUD. Phần khó nằm ở tính nhất quán dữ liệu, bảo mật, xử lý lỗi, timeout, trạng thái nghiệp vụ, trải nghiệm người dùng và khả năng vận hành. Việc xây dựng CinemaBooking.vn giúp em rèn luyện tư duy thiết kế hệ thống, tổ chức mã nguồn, phân quyền, tích hợp dịch vụ ngoài và kiểm thử các luồng nghiệp vụ phức tạp.")


def appendix(doc: Document) -> None:
    chapter(doc, "PHỤ LỤC")
    doc.add_heading("A. Danh sách sơ đồ nên hoàn thiện trước khi nộp", level=2)
    add_matrix(
        doc,
        ["Sơ đồ", "Có nên đưa vào bản cuối?", "Lý do"],
        [
            ["Sơ đồ ngữ cảnh", "Bắt buộc", "Giúp hội đồng hiểu hệ thống tương tác với ai."],
            ["BFD", "Bắt buộc", "Phù hợp chương phân tích chức năng."],
            ["Use Case", "Bắt buộc", "Thể hiện actor và chức năng."],
            ["DFD Level 0/1", "Rất nên", "Thể hiện luồng dữ liệu."],
            ["ERD", "Bắt buộc", "Database là phần mạnh của đề tài."],
            ["Activity diagram", "Rất nên", "Giải thích quy trình đặt vé/thanh toán/check-in."],
            ["Sequence diagram", "Rất nên", "Thể hiện rõ backend/frontend/payment/websocket tương tác."],
            ["State diagram", "Rất nên", "Rất hợp booking/payment/seat/ticket/showtime."],
            ["Class diagram", "Nên có", "Chỉ nên vẽ class nghiệp vụ chính, tránh vẽ toàn bộ quá rối."],
            ["Deployment diagram", "Nên có", "Giải thích triển khai local/demo/product."],
        ],
        [3000, 1800, 4560],
    )

    doc.add_heading("B. Danh sách ảnh giao diện nên chèn", level=2)
    add_bullets(
        doc,
        [
            "Trang chủ: danh sách phim và mua vé theo rạp.",
            "Trang chi tiết phim và lịch chiếu.",
            "Trang bản đồ rạp.",
            "Trang chọn ghế.",
            "Trang thanh toán VNPay/Quét QR ngân hàng.",
            "Trang kết quả giao dịch.",
            "Trang Vé của tôi.",
            "Trang hồ sơ cá nhân.",
            "Admin dashboard.",
            "Admin quản lý phim/rạp/phòng ghế/suất chiếu.",
            "Admin quản lý booking/payment/user/audit.",
            "Staff rạp phụ trách.",
            "Staff soát vé QR.",
            "Email vé điện tử trong Mailtrap.",
        ],
    )

    doc.add_heading("C. Tài khoản demo", level=2)
    add_matrix(
        doc,
        ["Role", "Username", "Password", "Ghi chú"],
        [
            ["ADMIN", "admin", "admin123 hoặc cấu hình trong app.admin.default-password", "Quản trị toàn hệ thống."],
            ["STAFF", "staff1", "123456", "Soát vé, vận hành rạp được phân công."],
            ["USER", "user1", "123456", "Đặt vé, thanh toán, xem vé."],
            ["USER", "user2", "123456", "Test multi-user."],
        ],
        [1500, 2300, 2800, 2760],
    )

    doc.add_heading("D. Checklist demo bảo vệ", level=2)
    add_numbered(
        doc,
        [
            "Khởi động PostgreSQL bằng Docker Compose.",
            "Chạy backend Spring Boot và kiểm tra Swagger.",
            "Chạy frontend React/Vite.",
            "Đăng nhập admin, xem dashboard.",
            "Admin tạo/sửa rạp hoặc suất chiếu.",
            "Admin gán rạp cho staff.",
            "Đăng nhập user, chọn phim/rạp/suất và giữ ghế.",
            "Mở trình duyệt thứ hai để thấy ghế HOLD realtime.",
            "Áp mã giảm giá và thanh toán bằng VNPay hoặc SePay.",
            "Xem email vé điện tử.",
            "Đăng nhập staff, chọn rạp/suất đang mở check-in.",
            "Quét QR đúng và thử quét lại vé đã dùng.",
            "Admin hủy suất chưa check-in và xem trạng thái hoàn tiền/audit/payment event.",
        ],
    )

    doc.add_heading("E. Gợi ý câu hỏi hội đồng có thể hỏi", level=2)
    add_matrix(
        doc,
        ["Câu hỏi", "Ý trả lời nên nắm"],
        [
            ["Vì sao cần seat_status?", "Vì ghế vật lý thuộc phòng, còn trạng thái ghế thay đổi theo từng suất chiếu."],
            ["Làm sao chống đặt trùng ghế?", "Dùng transaction, lock/update có điều kiện trạng thái, unique seat_showtime và kiểm tra HOLD."],
            ["Vì sao WebSocket publish sau commit?", "Để client chỉ nhận trạng thái chắc chắn đã lưu DB."],
            ["JWT và refresh token khác nhau thế nào?", "Access token ngắn hạn để gọi API; refresh token dài hơn để cấp access token mới, lưu hash và rotate."],
            ["Staff scope là gì?", "Nhân viên chỉ được thao tác trên rạp được admin phân công, ngoài permission chung."],
            ["SePay khác VNPay thế nào?", "VNPay redirect/callback; SePay dùng QR ngân hàng và webhook xác nhận tiền vào."],
            ["Nếu user bỏ thanh toán thì sao?", "Scheduler expire booking và trả ghế AVAILABLE."],
            ["Tại sao dùng Flyway?", "Quản lý version schema rõ ràng, triển khai nhất quán giữa môi trường."],
            ["Tại sao không cache seat_status?", "Seat status thay đổi realtime, cache lâu dễ stale và gây sai dữ liệu."],
            ["Hạn chế hiện tại là gì?", "Chưa refund tự động, chưa CI/CD/monitoring production đầy đủ, chưa combo/loyalty/mobile app."],
        ],
        [3600, 5760],
    )


def main() -> None:
    doc = Document()
    apply_styles(doc)
    add_page_number(doc.sections[0])

    build_cover(doc)
    build_intro_pages(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc)
    appendix(doc)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

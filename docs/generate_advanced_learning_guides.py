from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parent


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("52657A")
    p.paragraph_format.space_after = Pt(12)

    meta = doc.add_paragraph()
    meta_run = meta.add_run("CinemaBooking.vn - Tài liệu học luồng hệ thống, phục vụ bảo vệ khóa luận và phỏng vấn")
    meta_run.bold = True
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = RGBColor.from_string("9A5B00")


def add_callout(doc: Document, title: str, body: str, fill: str = "F4F8FC") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    add_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")
    p.add_run("\n" + body)
    doc.add_paragraph()


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    add_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FC")
    p = cell.paragraphs[0]
    for line_no, line in enumerate(code.strip("\n").splitlines()):
        if line_no:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("172033")
    doc.add_paragraph()


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    add_table_borders(table)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_shading(hdr[i], "E8EEF5")
        if widths:
            set_cell_width(hdr[i], widths[i])
        p = hdr[i].paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string("0B2545")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            if widths:
                set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[i].paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9.2)
    doc.add_paragraph()


def add_common_tail(doc: Document, interview_questions: list[tuple[str, str]], checklist: list[str]) -> None:
    doc.add_heading("Câu hỏi bảo vệ/phỏng vấn nên tự trả lời được", level=1)
    add_matrix(
        doc,
        ["Câu hỏi", "Ý chính cần trả lời"],
        [[q, a] for q, a in interview_questions],
        [4200, 5160],
    )
    doc.add_heading("Checklist tự học", level=1)
    add_bullets(doc, checklist)


def save_doc(doc: Document, filename: str) -> Path:
    path = DOCS_DIR / filename
    doc.save(path)
    return path


def build_auth_doc() -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(
        doc,
        "Luồng Auth, JWT, Refresh Token và Google Login",
        "Giải thích cách người dùng đăng nhập, server phát token, client tự refresh token, logout và chống reuse refresh token.",
    )
    add_callout(
        doc,
        "Hiểu nhanh",
        "Access token giống thẻ ra vào ngắn hạn. Refresh token giống chìa khóa gia hạn phiên, được lưu HttpOnly cookie và lưu dạng hash trong DB. Khi access token hết hạn, frontend gọi /auth/refresh để lấy access token mới.",
    )
    doc.add_heading("Sơ đồ luồng tổng quát", level=1)
    add_code(
        doc,
        """
User nhập username/password
        |
        v
LoginPage.tsx -> authApi.login()
        |
        v
POST /auth/token -> AuthenticationController.authenticate()
        |
        v
AuthenticationService.authenticate()
  - kiểm tra rate limit
  - tìm user
  - so BCrypt password
  - kiểm tra active/deleted/emailVerified
  - sinh access token + refresh token
        |
        v
Response:
  - access token trả về body
  - refresh token ghi vào HttpOnly cookie
        |
        v
authStore lưu access token, user info, permissions
        |
        v
axiosClient tự gắn Authorization: Bearer <access token>
        """,
    )
    doc.add_heading("Các class chính và vai trò", level=1)
    add_matrix(
        doc,
        ["Class/File", "Vai trò", "Điểm cần nhớ"],
        [
            ["AuthenticationController.java", "Mở API /auth/token, /auth/google, /auth/refresh, /auth/logout, /auth/sessions.", "Controller không tự xử lý nghiệp vụ sâu, chỉ nhận request, gọi service, ghi/clear cookie refresh token."],
            ["AuthenticationService.java", "Trung tâm auth: login, Google login, refresh, logout, revoke session, verify token.", "Nơi sinh JWT, hash refresh token, chống refresh token reuse và ghi audit."],
            ["SecurityConfig.java", "Cấu hình Spring Security.", "Permit public endpoint, bật Resource Server JWT, dùng CustomJwtDecoder, map scope thành authority."],
            ["CustomJwtDecoder.java", "Giải mã JWT trong request protected.", "Khi request có Bearer token, Spring gọi decoder để xác thực token."],
            ["RefreshToken.java", "Entity lưu refresh token đã hash.", "Không lưu refresh token raw trong DB, giảm rủi ro nếu DB lộ."],
            ["InvalidatedToken.java", "Blacklist access token đã logout.", "Access token vốn stateless, nên khi logout cần lưu jti vào bảng invalidated_token tới khi token hết hạn."],
            ["authStore.ts", "Zustand store lưu trạng thái đăng nhập trên frontend.", "Lưu access token, user, permissions; logout dọn localStorage."],
            ["axiosClient.ts", "Axios interceptor gắn token và tự refresh.", "Khi API trả 401, client gọi /auth/refresh một lần, retry request cũ."],
            ["ProtectedRoute.tsx", "Chặn route cần đăng nhập/quyền.", "Nếu chưa login hoặc token hết hạn thì chuyển /login; nếu thiếu permission thì về trang chủ."],
        ],
        [2600, 3300, 3460],
    )
    doc.add_heading("Giải thích backend từng đoạn quan trọng", level=1)
    doc.add_heading("@RestController và /auth", level=2)
    add_code(
        doc,
        """
@RestController
@RequestMapping("/auth")
class AuthenticationController {
    @PostMapping("/token")
    ApiResponse<AuthenticationResponse> authenticate(...)
}
        """,
    )
    add_bullets(
        doc,
        [
            "@RestController: class này trả JSON trực tiếp, không trả view HTML.",
            "@RequestMapping(\"/auth\"): mọi endpoint trong class bắt đầu bằng /auth.",
            "@PostMapping(\"/token\"): API đăng nhập bằng username/password là POST /auth/token.",
            "ApiResponse<T>: format response thống nhất cho frontend, có code/message/result/timestamp/path.",
        ],
    )
    doc.add_heading("Login password trong AuthenticationService", level=2)
    add_code(
        doc,
        """
authRateLimitService.check(rateLimitKey, LOGIN_MAX_ATTEMPTS, LOGIN_WINDOW);
user = userRepository.findByUsername(username).orElseThrow(...);
validateUserCanAuthenticate(user);
if (!passwordEncoder.matches(request.getPassword(), user.getPassword())) {
    throw new AppException(ErrorCode.UNAUTHENTICATED);
}
AuthenticationResponse response = issueTokenPair(user, servletRequest);
authAuditService.record(EVENT_LOGIN_PASSWORD, user, username, true, null, servletRequest);
        """,
    )
    add_bullets(
        doc,
        [
            "Rate limit: chống brute-force, ví dụ nhập sai liên tục trong 15 phút.",
            "findByUsername: lấy user từ DB. Nếu không có thì không cho login.",
            "validateUserCanAuthenticate: user bị khóa, bị xóa mềm, hoặc email chưa xác thực thì không cho login.",
            "BCrypt matches: so mật khẩu plain người dùng nhập với hash BCrypt trong DB.",
            "issueTokenPair: sinh access token và refresh token cùng lúc.",
            "authAuditService.record: ghi nhật ký login thành công/thất bại để admin điều tra bảo mật.",
        ],
    )
    doc.add_heading("Access token chứa gì?", level=2)
    add_code(
        doc,
        """
JWTClaimsSet jwtClaimsSet = new JWTClaimsSet.Builder()
    .subject(user.getUsername())
    .issuer("cinema-booking")
    .expirationTime(...)
    .jwtID(UUID.randomUUID().toString())
    .claim("token_use", "access")
    .claim("scope", buildScope(user))
    .claim("userId", user.getId().toString())
    .build();
        """,
    )
    add_bullets(
        doc,
        [
            "subject: username của người dùng.",
            "issuer: hệ thống phát hành token, dùng để biết token thuộc app nào.",
            "expirationTime: access token ngắn hạn, hết hạn thì không dùng được nữa.",
            "jwtID: mã duy nhất của token, dùng để blacklist khi logout.",
            "token_use=access: phân biệt access token với refresh token.",
            "scope: chuỗi quyền như ROLE_USER MOVIE_VIEW BOOKING_CREATE; Spring Security biến scope thành authority.",
            "userId: frontend/backend có thể biết id user mà không cần query thêm trong một số trường hợp.",
        ],
    )
    doc.add_heading("Refresh token rotation và chống reuse", level=2)
    add_code(
        doc,
        """
verifyToken(refreshToken, true);
RefreshToken current = refreshTokenRepository.findByTokenHash(hashToken(refreshToken)).orElseThrow(...);
String newRefreshToken = generateRefreshToken(user, servletRequest);
current.setRevokedAt(LocalDateTime.now());
current.setRevokedReason("ROTATED");
current.setReplacedByTokenId(newRefreshTokenId);
        """,
    )
    add_bullets(
        doc,
        [
            "Mỗi lần refresh thành công, refresh token cũ bị thu hồi và token mới được cấp.",
            "Nếu refresh token cũ bị dùng lại sau khi đã rotate, hệ thống coi là reuse detected.",
            "Khi reuse detected, hệ thống revoke toàn bộ refresh token active của user để bảo vệ tài khoản.",
            "Đây là cách làm chuẩn hơn kiểu refresh token cố định dùng mãi tới khi hết hạn.",
        ],
    )
    doc.add_heading("Frontend tự refresh token như thế nào?", level=1)
    add_code(
        doc,
        """
axiosClient.interceptors.response.use(
  response => response,
  async error => {
    if (error.response?.status === 401 && !originalRequest._retry) {
      const accessToken = await refreshAccessToken();
      originalRequest.headers.Authorization = `Bearer ${accessToken}`;
      return axiosClient(originalRequest);
    }
  }
);
        """,
    )
    add_bullets(
        doc,
        [
            "Interceptor giống người gác cổng của frontend: request đi ra thì gắn token, response lỗi thì xử lý tập trung.",
            "Nếu access token hết hạn, API trả 401. Frontend không bắt user login lại ngay mà gọi /auth/refresh.",
            "refreshPromise giúp nhiều request 401 cùng lúc không gọi refresh trùng nhiều lần.",
            "Nếu refresh cũng lỗi, authStore.logout() dọn dữ liệu và đưa user về /login.",
        ],
    )
    doc.add_heading("Google Login khác password login ở đâu?", level=1)
    add_bullets(
        doc,
        [
            "Frontend lấy Google ID token từ Google.",
            "Backend decode ID token bằng Google JWK Set URI.",
            "Backend kiểm tra issuer là https://accounts.google.com và audience khớp Google client id của app.",
            "Nếu email_verified=true, backend tìm user theo email. Nếu chưa có, tạo user mới role USER.",
            "Sau đó backend vẫn phát access token/refresh token nội bộ của hệ thống, nên các API còn lại không phụ thuộc trực tiếp Google.",
        ],
    )
    add_common_tail(
        doc,
        [
            ("Vì sao không lưu refresh token raw trong DB?", "Vì nếu DB bị lộ, attacker không lấy được token thật; hệ thống chỉ lưu SHA-256 hash để đối chiếu."),
            ("Vì sao refresh token cần rotate?", "Để phát hiện token bị đánh cắp. Token cũ dùng lại sau khi đã rotate là dấu hiệu bất thường."),
            ("Vì sao access token vẫn lưu localStorage?", "Dễ dùng cho SPA, nhưng cần access token ngắn hạn. Refresh token bảo mật hơn vì đặt HttpOnly cookie."),
            ("Logout xử lý gì?", "Blacklist access token bằng jti và revoke refresh token/cookie."),
            ("Google login có thay thế JWT nội bộ không?", "Không. Google chỉ xác minh danh tính ban đầu, sau đó hệ thống vẫn dùng JWT nội bộ."),
        ],
        [
            "Tự vẽ lại luồng /auth/token từ frontend tới database.",
            "Giải thích được khác nhau giữa access token và refresh token.",
            "Mở AuthenticationService.java và tìm được generateAccessToken, generateRefreshToken, refreshToken, logout.",
            "Mở axiosClient.ts và giải thích được vì sao có refreshPromise.",
            "Trả lời được reuse refresh token nguy hiểm như thế nào.",
        ],
    )
    return save_doc(doc, "Luong_Auth_JWT_Session_GoogleLogin_CinemaBooking.docx")


def build_rbac_doc() -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(
        doc,
        "Luồng RBAC, Permission và Staff Scope theo rạp",
        "Giải thích cách hệ thống phân quyền ADMIN/STAFF/USER, kiểm tra permission ở backend/frontend và giới hạn dữ liệu nhân viên theo rạp phụ trách.",
    )
    add_callout(
        doc,
        "Hiểu nhanh",
        "RBAC trả lời câu hỏi: user có quyền làm API này không? Staff scope trả lời câu hỏi sâu hơn: staff có quyền thao tác trên đúng rạp được phân công không?",
    )
    doc.add_heading("Mô hình dữ liệu phân quyền", level=1)
    add_code(
        doc,
        """
users --< users_roles >-- roles --< roles_permissions >-- permissions

Ví dụ:
user staff1
  -> ROLE_STAFF
  -> permissions: SHOWTIME_CREATE, TICKET_CHECKIN, BOOKING_VIEW_ALL, ...

staff_cinemas
  -> staff_id
  -> cinema_id
  -> active
        """,
    )
    add_matrix(
        doc,
        ["Bảng/Class", "Ý nghĩa", "Tại sao cần"],
        [
            ["users", "Tài khoản người dùng.", "Lưu username, email, password hash, trạng thái active/deleted, avatar."],
            ["roles", "Vai trò như ADMIN, STAFF, USER.", "Gom nhóm quyền theo chức năng thực tế."],
            ["permissions", "Quyền cụ thể như MOVIE_VIEW, SHOWTIME_CREATE.", "Dùng để bảo vệ từng API nhỏ."],
            ["users_roles", "Bảng nối user-role.", "Một user có thể có nhiều role."],
            ["roles_permissions", "Bảng nối role-permission.", "Một role có nhiều permission."],
            ["staff_cinemas", "Bảng gán staff phụ trách rạp.", "Không cho staff xem/thao tác toàn bộ rạp."],
            ["PermissionName.java", "Enum danh sách quyền chuẩn trong code.", "Tránh gõ chuỗi quyền lung tung."],
            ["ApplicationInitConfig.java", "Seed role/permission/admin/user/staff mặc định.", "Đảm bảo app chạy lần đầu có RBAC core."],
        ],
        [2200, 3500, 3660],
    )
    doc.add_heading("Backend chặn quyền bằng @PreAuthorize", level=1)
    add_code(
        doc,
        """
@PreAuthorize("hasAuthority('SHOWTIME_CREATE')")
@PostMapping
ApiResponse<ShowtimeResponse> createShowtime(@RequestBody ShowtimeCreationRequest request)
        """,
    )
    add_bullets(
        doc,
        [
            "@EnableMethodSecurity trong SecurityConfig bật kiểm tra @PreAuthorize.",
            "hasAuthority('SHOWTIME_CREATE') nghĩa là JWT của user phải có permission SHOWTIME_CREATE trong scope.",
            "Nếu thiếu quyền, Spring Security ném AccessDeniedException và GlobalExceptionHandler/JwtAccessDeniedHandler trả lỗi 403.",
            "Cách này tốt hơn check role cứng, vì sau này role thay đổi nhưng permission vẫn giữ đúng nghĩa nghiệp vụ.",
        ],
    )
    doc.add_heading("Scope trong JWT đến từ đâu?", level=1)
    add_code(
        doc,
        """
private String buildScope(User user) {
    StringJoiner stringJoiner = new StringJoiner(" ");
    user.getRoles().forEach(role -> {
        stringJoiner.add("ROLE_" + role.getName());
        role.getPermissions().forEach(p -> stringJoiner.add(p.getName()));
    });
    return stringJoiner.toString();
}
        """,
    )
    add_bullets(
        doc,
        [
            "Khi login, backend đọc roles và permissions của user.",
            "Token sẽ có claim scope chứa cả ROLE_ADMIN/ROLE_STAFF/ROLE_USER và permission chi tiết.",
            "SecurityConfig dùng JwtGrantedAuthoritiesConverter với authorityPrefix rỗng, nên permission trong scope trở thành authority trực tiếp.",
        ],
    )
    doc.add_heading("Staff scope: chặn đúng rạp", level=1)
    add_code(
        doc,
        """
public void validateCurrentStaffCanAccessCinema(UUID cinemaId) {
    if (cinemaId == null || !isStaffButNotAdmin()) {
        return;
    }
    boolean assigned = staffCinemaRepository.existsActiveAssignment(
            SecurityUtils.getCurrentUserId(),
            cinemaId);
    if (!assigned) {
        throw new AppException(ErrorCode.UNAUTHORIZED);
    }
}
        """,
    )
    add_bullets(
        doc,
        [
            "ADMIN bỏ qua scope vì admin quản lý toàn hệ thống.",
            "STAFF thì phải có record active trong staff_cinemas cho rạp đó.",
            "Service nghiệp vụ phải gọi validateCurrentStaffCanAccessCinema trước khi tạo/sửa/hủy suất hoặc soát vé theo rạp.",
            "Đây là lớp bảo vệ dữ liệu theo phạm vi, nằm sau lớp permission.",
        ],
    )
    doc.add_heading("Frontend dùng permission thế nào?", level=1)
    add_code(
        doc,
        """
<ProtectedRoute permission="PAYMENT_VIEW_ALL">
  <AdminPaymentPage />
</ProtectedRoute>

hasPermission: (perm) => {
  return permissions.includes(perm) || permissions.includes('ALL');
}
        """,
    )
    add_bullets(
        doc,
        [
            "Frontend route guard giúp ẩn/chặn màn hình không phù hợp.",
            "Sidebar admin/staff dựa vào permission để hiện menu tương ứng.",
            "Quan trọng: frontend chỉ cải thiện UX, bảo mật thật vẫn phải nằm ở backend bằng @PreAuthorize và service scope.",
        ],
    )
    doc.add_heading("Quy tắc phân quyền hiện tại nên nhớ", level=1)
    add_matrix(
        doc,
        ["Role", "Nên được làm", "Không nên được làm"],
        [
            ["ADMIN", "Quản trị phim, rạp, phòng, ghế, suất chiếu, user, role, audit, dashboard, payment.", "Không có giới hạn scope rạp."],
            ["STAFF", "Xem dashboard vận hành, tạo/sửa/hủy suất ở rạp phụ trách, xem booking/payment thuộc rạp, soát QR.", "Không quản lý role/permission, không tạo admin, không refund thật nếu chưa có quy trình duyệt."],
            ["USER", "Xem phim/rạp/suất, giữ ghế, đặt vé, thanh toán, xem vé của mình, cập nhật profile.", "Không xem dữ liệu người khác, không truy cập admin/staff."],
        ],
        [1700, 4300, 3360],
    )
    add_common_tail(
        doc,
        [
            ("RBAC khác staff scope thế nào?", "RBAC kiểm tra quyền chức năng; staff scope kiểm tra phạm vi dữ liệu cụ thể."),
            ("Vì sao không chỉ dùng ROLE_STAFF?", "Vì role quá rộng. Permission giúp chia nhỏ quyền như SHOWTIME_CREATE, TICKET_CHECKIN."),
            ("Frontend ProtectedRoute có đủ bảo mật không?", "Không. Frontend chỉ là UX. Backend mới là lớp bảo vệ thật."),
            ("Nếu staff đổi rạp phụ trách thì sao?", "Cập nhật staff_cinemas, các request sau đó sẽ bị scope service kiểm tra theo assignment mới."),
            ("Vì sao ADMIN không bị scope?", "Admin cần vận hành toàn hệ thống, scope sẽ gây cản trở quản trị."),
        ],
        [
            "Mở PermissionName.java và đọc toàn bộ permission.",
            "Mở SecurityConfig.java và tìm @EnableMethodSecurity.",
            "Mở một controller bất kỳ và chỉ ra @PreAuthorize bảo vệ API nào.",
            "Giải thích StaffCinemaScopeService bằng lời của bạn.",
            "Phân biệt 403 thiếu quyền và 401 chưa đăng nhập/token sai.",
        ],
    )
    return save_doc(doc, "Luong_RBAC_Permission_StaffScope_CinemaBooking.docx")


def build_scheduler_doc() -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(
        doc,
        "Luồng Scheduler, giữ ghế hết hạn và đồng bộ trạng thái suất chiếu",
        "Giải thích các job nền chạy định kỳ: trả ghế HOLD hết hạn, hết hạn booking PENDING, dọn token, cập nhật UPCOMING/ONGOING/ENDED.",
    )
    add_callout(
        doc,
        "Hiểu nhanh",
        "Scheduler là những nhân viên chạy nền. Người dùng không bấm gì nhưng hệ thống vẫn tự dọn ghế giữ quá hạn, tự hủy booking chờ thanh toán, tự cập nhật suất chiếu đang chiếu/kết thúc.",
    )
    doc.add_heading("Các scheduler hiện có", level=1)
    add_matrix(
        doc,
        ["Scheduler", "Tần suất/cấu hình", "Nhiệm vụ"],
        [
            ["HoldExpireScheduler", "booking.expired-hold-scan-delay-ms", "Tìm seat_status HOLD quá hold_until, trả về AVAILABLE, push WebSocket."],
            ["PendingBookingExpireScheduler", "booking.expired-booking-scan-delay-ms", "Tìm booking PENDING quá payment_expires_at hoặc legacy timeout, expire booking."],
            ["ShowtimeStatusSyncScheduler", "showtime.status-sync-delay-ms", "Đồng bộ trạng thái showtime theo thời gian hiện tại."],
            ["TokenCleanupTask", "cron 2 giờ sáng", "Dọn invalidated token và refresh token hết hạn."],
        ],
        [2500, 2500, 4360],
    )
    doc.add_heading("Luồng giữ ghế và tự trả ghế", level=1)
    add_code(
        doc,
        """
User chọn ghế
  -> BookingServiceImpl.holdSeats()
  -> seat_status.status = HOLD
  -> hold_by = userId
  -> hold_until = now + BOOKING_SEAT_HOLD_MINUTES
  -> WebSocket báo ghế HOLD cho mọi client

Nếu user không thanh toán
  -> HoldExpireScheduler chạy nền
  -> tìm HOLD có hold_until <= now
  -> update AVAILABLE + clear hold_by/hold_until
  -> sau commit mới publish WebSocket AVAILABLE
        """,
    )
    doc.add_heading("Code HoldExpireScheduler cần hiểu", level=1)
    add_code(
        doc,
        """
@Scheduled(fixedDelayString = "${booking.expired-hold-scan-delay-ms:30000}")
@Transactional
public void releaseExpiredHolds() {
    List<ExpiredSeatHoldProjection> expired =
        seatStatusRepository.findExpiredHoldRows(LocalDateTime.now(), expiredHoldScanLimit);

    int releasedCount = seatStatusRepository.releaseExpiredHoldsByIds(expiredIds);

    TransactionSynchronizationManager.registerSynchronization(new TransactionSynchronization() {
        @Override
        public void afterCommit() {
            seatIdsByShowtime.forEach(seatStatusPublisher::publishAvailable);
        }
    });
}
        """,
    )
    add_bullets(
        doc,
        [
            "@Scheduled: Spring tự gọi method theo chu kỳ.",
            "fixedDelay: đợi method chạy xong rồi mới đếm delay cho lần tiếp theo.",
            "@Transactional: tất cả update DB trong method nằm trong một transaction.",
            "findExpiredHoldRows dùng projection để lấy đúng dữ liệu cần thiết, nhẹ hơn load entity lớn.",
            "releaseExpiredHoldsByIds update theo batch, tránh update từng ghế gây chậm.",
            "afterCommit: chỉ publish WebSocket sau khi DB commit thành công. Nếu DB rollback thì client không nhận trạng thái sai.",
        ],
    )
    doc.add_heading("Luồng booking PENDING hết hạn", level=1)
    add_code(
        doc,
        """
User tạo booking PENDING
  -> booking.payment_expires_at = now + pending timeout
  -> payment PENDING nếu đã tạo thanh toán

Nếu hết giờ
  -> PendingBookingExpireScheduler.findExpiredPendingBookingIds()
  -> BookingServiceImpl.expirePendingBooking(id)
  -> booking.status = EXPIRED
  -> payment.status = EXPIRED/FAILED tùy luồng
  -> seat_status trả AVAILABLE
  -> WebSocket refresh ghế
        """,
    )
    add_bullets(
        doc,
        [
            "HOLD và PENDING là hai tầng khác nhau: HOLD giữ ghế tạm; PENDING là đơn đã tạo nhưng chưa thanh toán xong.",
            "payment_expires_at giúp mỗi booking có hạn thanh toán rõ ràng.",
            "Scheduler xử lý trường hợp user tắt tab, back browser, mạng lag hoặc cổng thanh toán không callback.",
        ],
    )
    doc.add_heading("Đồng bộ trạng thái suất chiếu", level=1)
    add_bullets(
        doc,
        [
            "UPCOMING: chưa tới giờ bắt đầu.",
            "ONGOING: đang trong khoảng start_time đến end_time.",
            "ENDED: đã qua end_time.",
            "CANCELLED: admin/staff hủy suất.",
            "ShowtimeStatusSyncScheduler chạy định kỳ để status không bị cũ khi thời gian trôi qua.",
            "ShowtimeServiceImpl.getAllShowtimes cũng gọi synchronizeCurrentStatuses() trước khi trả danh sách admin, giúp màn hình quản lý không hiển thị sai trạng thái quá lâu.",
        ],
    )
    doc.add_heading("Cấu hình liên quan trong .env/application.yaml", level=1)
    add_matrix(
        doc,
        ["Config", "Ý nghĩa", "Gợi ý product"],
        [
            ["BOOKING_SEAT_HOLD_MINUTES", "Số phút giữ ghế khi user chọn ghế.", "Thường 5 phút là hợp lý cho flow bán vé online."],
            ["BOOKING_PENDING_TIMEOUT_MINUTES", "Số phút booking chờ thanh toán.", "Nên đồng bộ với UX countdown hoặc ngắn hơn tùy cổng thanh toán."],
            ["BOOKING_EXPIRED_HOLD_SCAN_DELAY_MS", "Chu kỳ quét ghế HOLD hết hạn.", "15-30 giây đủ mượt, không quá tải DB."],
            ["BOOKING_EXPIRED_BOOKING_SCAN_DELAY_MS", "Chu kỳ quét booking PENDING hết hạn.", "15-30 giây cho demo/product nhỏ."],
            ["showtime.status-sync-delay-ms", "Chu kỳ đồng bộ trạng thái showtime.", "60 giây là hợp lý."],
        ],
        [2900, 3500, 2960],
    )
    add_common_tail(
        doc,
        [
            ("Vì sao scheduler cần afterCommit?", "Để client chỉ thấy trạng thái mới khi DB chắc chắn đã lưu thành công."),
            ("HOLD khác PENDING thế nào?", "HOLD là trạng thái ghế; PENDING là trạng thái booking/thanh toán."),
            ("Nếu scheduler chết thì sao?", "Ghế/booking có thể chậm được trả. Khi app chạy lại scheduler tiếp tục quét dựa vào thời gian trong DB."),
            ("Vì sao dùng scan limit?", "Để mỗi lần quét không xử lý quá nhiều row gây spike DB."),
            ("Có nên dùng queue không?", "Khi hệ thống lớn có thể dùng queue/delayed job, nhưng scheduler DB scan hiện đủ tốt cho monolith/demo."),
        ],
        [
            "Mở HoldExpireScheduler.java và giải thích từng annotation.",
            "Mở PendingBookingExpireScheduler.java và vẽ lại luồng expire booking.",
            "Giải thích vì sao không cần refresh trang ghế khi có WebSocket.",
            "Nêu được rủi ro nếu không có scheduler.",
            "Biết config timeout nằm ở .env/application.yaml.",
        ],
    )
    return save_doc(doc, "Luong_Scheduler_GiuGhe_HetHan_TrangThaiSuatChieu_CinemaBooking.docx")


def build_exception_audit_doc() -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(
        doc,
        "Luồng Exception Handling, Audit Log và Soft Delete",
        "Giải thích cách backend chuẩn hóa lỗi trả về client, ghi nhật ký thao tác admin/auth/payment và chiến lược xóa mềm/xóa cứng.",
    )
    add_callout(
        doc,
        "Hiểu nhanh",
        "Product thật không nên để lỗi Java thô rơi ra frontend. Service ném AppException(ErrorCode.X), GlobalExceptionHandler đổi thành JSON dễ hiểu. Với thao tác quan trọng, hệ thống ghi audit để truy vết ai làm gì, lúc nào.",
    )
    doc.add_heading("Luồng lỗi chuẩn", level=1)
    add_code(
        doc,
        """
Service phát hiện lỗi nghiệp vụ
  -> throw new AppException(ErrorCode.SEAT_ALREADY_BOOKED)
        |
        v
GlobalExceptionHandler.handlingAppException()
        |
        v
Response JSON:
{
  "code": ...,
  "message": "...",
  "timestamp": "...",
  "path": "/api/v1/..."
}
        |
        v
Frontend axios catch -> toast/message tiếng Việt thân thiện
        """,
    )
    add_matrix(
        doc,
        ["Class/File", "Vai trò", "Điểm học"],
        [
            ["AppException.java", "Exception nghiệp vụ có ErrorCode.", "Thay vì throw RuntimeException lung tung, service ném lỗi có mã rõ ràng."],
            ["ErrorCode.java", "Danh sách mã lỗi, message, HTTP status.", "Một nguồn chuẩn cho backend và frontend hiểu lỗi."],
            ["GlobalExceptionHandler.java", "Bắt exception toàn cục.", "Dùng @ControllerAdvice và @ExceptionHandler để trả response thống nhất."],
            ["ApiResponse.java", "Wrapper response chuẩn.", "API thành công/thất bại đều cùng shape."],
            ["JwtAuthenticationEntryPoint.java", "Trả lỗi 401 khi token thiếu/sai/hết hạn.", "Dành riêng cho Spring Security authentication."],
            ["JwtAccessDeniedHandler.java", "Trả lỗi 403 khi thiếu quyền.", "Dành riêng cho authorization."],
        ],
        [2600, 3300, 3460],
    )
    doc.add_heading("GlobalExceptionHandler xử lý những nhóm lỗi nào?", level=1)
    add_matrix(
        doc,
        ["Nhóm lỗi", "Handler", "Ý nghĩa product"],
        [
            ["AppException", "handlingAppException", "Lỗi nghiệp vụ đã biết: hết ghế, thiếu quyền, booking hết hạn..."],
            ["Validation", "MethodArgumentNotValidException, ConstraintViolationException", "Form sai dữ liệu: thiếu field, sai min/max, format không hợp lệ."],
            ["Authentication/AccessDenied", "AuthenticationCredentialsNotFoundException, AccessDeniedException", "401/403 rõ ràng cho client."],
            ["Bad body/type/method", "HttpMessageNotReadable, MethodArgumentTypeMismatch, MethodNotSupported", "Client gọi API sai vẫn nhận lỗi đẹp."],
            ["Optimistic locking", "ObjectOptimisticLockingFailureException", "Xung đột cập nhật đồng thời, quan trọng với ghế/seat_status."],
            ["Data integrity", "DataIntegrityViolationException", "Trùng unique key, FK lỗi, constraint DB."],
            ["Exception chung", "Exception.class", "Fallback 500, có log stacktrace để dev điều tra."],
        ],
        [2200, 3400, 3760],
    )
    doc.add_heading("Audit log trong hệ thống", level=1)
    add_matrix(
        doc,
        ["Audit", "Class/Bảng", "Dùng cho"],
        [
            ["Auth audit", "AuthAuditLog, AuthAuditService", "Ghi login password, login Google, refresh token, logout, revoke session."],
            ["Admin audit", "AdminAuditLog, AdminAuditLogInterceptor", "Ghi thao tác quản trị như tạo/sửa/xóa user, movie, showtime..."],
            ["Payment event", "PaymentEvent, PaymentEventService", "Ghi event thanh toán: initiate, callback, webhook, refund requested, reconciliation."],
        ],
        [1800, 3200, 4360],
    )
    add_bullets(
        doc,
        [
            "@Transactional(propagation = REQUIRES_NEW) trong audit/payment event giúp log vẫn có transaction riêng.",
            "Audit giúp trả lời câu hỏi khi có sự cố: ai thao tác, IP nào, user-agent nào, thành công hay thất bại.",
            "Payment event không chỉ để debug mà còn phục vụ đối soát thanh toán với cổng.",
        ],
    )
    doc.add_heading("Soft delete và hard delete", level=1)
    add_matrix(
        doc,
        ["Kiểu xóa", "Dùng khi nào", "Ví dụ trong hệ thống"],
        [
            ["Soft delete", "Dữ liệu từng liên quan giao dịch, cần giữ lịch sử.", "movies, cinemas, rooms, seats, showtimes thường set is_deleted=true hoặc status=CANCELLED."],
            ["Hard delete", "Dữ liệu tạm/kỹ thuật, không cần lịch sử lâu dài hoặc cascade hợp lý.", "invalidated_token hết hạn, refresh token hết hạn, bảng nối khi gỡ role/assignment."],
            ["Cancel with policy", "Không xóa suất có booking mà chuyển trạng thái và xử lý vé/thanh toán.", "ShowtimeServiceImpl.cancelShowtimeWithPolicy."],
        ],
        [1900, 3800, 3660],
    )
    doc.add_heading("Vì sao không xóa cứng booking/payment/ticket?", level=2)
    add_bullets(
        doc,
        [
            "Đơn hàng và thanh toán là chứng từ. Product thật cần giữ để hỗ trợ khách, khiếu nại, báo cáo doanh thu.",
            "Ticket đã dùng cần giữ check_in_time để đối soát soát vé.",
            "Nếu xóa cứng, dashboard, audit, email vé, lịch sử giao dịch dễ bị mất dữ liệu.",
        ],
    )
    doc.add_heading("Frontend nên hiển thị lỗi thế nào?", level=1)
    add_bullets(
        doc,
        [
            "Không show stacktrace/JDBC message cho user.",
            "Map message backend sang tiếng Việt thân thiện nếu cần.",
            "Với lỗi validation, đọc result.errors để highlight đúng field.",
            "Với 401, redirect login hoặc refresh token tự động.",
            "Với 403, thông báo không đủ quyền thay vì reload trắng trang.",
        ],
    )
    add_common_tail(
        doc,
        [
            ("Vì sao cần GlobalExceptionHandler?", "Để lỗi API thống nhất, không rò stacktrace và frontend dễ xử lý."),
            ("AppException khác Exception thường?", "AppException là lỗi nghiệp vụ có ErrorCode; Exception thường là lỗi bất ngờ/fallback."),
            ("Soft delete có nhược điểm gì?", "Query phải luôn lọc is_deleted=false; dữ liệu lớn cần index phù hợp."),
            ("Audit log có nên rollback cùng nghiệp vụ không?", "Tùy loại. Hệ thống dùng REQUIRES_NEW cho nhiều audit/event để tăng khả năng ghi dấu vết."),
            ("Khi hủy suất có vé đã thanh toán thì làm gì?", "Không xóa booking; chuyển CANCELLED, hủy ticket active, ghi refund requested/payment event, gửi email."),
        ],
        [
            "Mở GlobalExceptionHandler.java và đọc các @ExceptionHandler.",
            "Mở ErrorCode.java và hiểu code/message/status.",
            "Giải thích vì sao dashboard không nên tính dữ liệu is_deleted=true.",
            "Nêu được khác nhau giữa CANCELLED booking và deleted booking.",
            "Tự mô tả luồng lỗi từ service tới toast frontend.",
        ],
    )
    return save_doc(doc, "Luong_Exception_Audit_SoftDelete_CinemaBooking.docx")


def build_cache_query_doc() -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(
        doc,
        "Luồng Cache, Query Optimization, Index và Flyway Migration",
        "Giải thích cách hệ thống giảm query thừa, tránh N+1, cache dữ liệu ít đổi, dùng index và Flyway để quản lý schema.",
    )
    add_callout(
        doc,
        "Hiểu nhanh",
        "Cache là nhớ kết quả đọc ít thay đổi. Index là mục lục của database. Flyway là lịch sử version database. Ba thứ này giúp hệ thống nhanh và dễ triển khai nhất quán.",
    )
    doc.add_heading("Cache hiện tại dùng gì?", level=1)
    add_code(
        doc,
        """
@Configuration
@EnableCaching
public class CacheConfig {
    public static final String MOVIES = "movies";
    public static final String CINEMAS = "cinemas";
    public static final String CINEMA_MAP = "cinema-map";
    public static final String ROOMS_BY_CINEMA = "rooms-by-cinema";
    public static final String SEATS_BY_ROOM = "seats-by-room";
    public static final String PROMOTIONS = "promotions";
}
        """,
    )
    add_bullets(
        doc,
        [
            "@EnableCaching bật Spring Cache abstraction.",
            "@Cacheable lưu kết quả method đọc. Lần sau cùng key thì trả từ cache, không query DB.",
            "@CacheEvict xóa cache khi dữ liệu thay đổi.",
            "Hiện chưa thấy Redis trong code. Nếu sau này thêm Redis, service vẫn dùng @Cacheable/@CacheEvict, chỉ đổi provider cache.",
        ],
    )
    doc.add_heading("Dữ liệu nào nên cache?", level=1)
    add_matrix(
        doc,
        ["Dữ liệu", "Có nên cache?", "Lý do"],
        [
            ["Movies", "Có", "Danh sách phim thay đổi ít, đọc nhiều ở homepage."],
            ["Cinemas / cinema map", "Có", "Rạp/tọa độ ít đổi, nhiều user xem map/danh sách rạp."],
            ["Rooms by cinema", "Có", "Phòng chiếu gần như master data."],
            ["Seats by room", "Có", "Ghế vật lý của phòng ít đổi."],
            ["Promotions active", "Có nhưng phải evict khi dùng/cập nhật", "Mã giảm giá đọc nhiều, nhưng used_count/active thay đổi cần cẩn thận."],
            ["Seat status", "Không cache lâu", "Trạng thái ghế thay đổi realtime, cache dễ stale."],
            ["Bookings/payments", "Không cache phổ thông", "Dữ liệu giao dịch riêng user, thay đổi liên tục."],
        ],
        [2200, 1900, 5260],
    )
    doc.add_heading("Tránh N+1 query", level=1)
    add_bullets(
        doc,
        [
            "N+1 xảy ra khi lấy danh sách N entity, rồi mỗi entity lại lazy load thêm quan hệ bằng 1 query riêng.",
            "Cách xử lý trong hệ thống: repository dùng fetch join/entity graph/projection/DTO query cho màn danh sách cần dữ liệu nested.",
            "Ví dụ payment admin cần payment + booking + user + showtime + movie + room + cinema thì query nên join sẵn, không map từng payment rồi lazy load từng phần.",
            "Projection như ExpiredSeatHoldProjection giúp scheduler chỉ lấy id/showtimeId/seatId thay vì load cả SeatStatus entity graph.",
        ],
    )
    doc.add_heading("Index là gì?", level=1)
    add_callout(
        doc,
        "Ví dụ dễ hiểu",
        "Database không có index giống cuốn sách không có mục lục: muốn tìm tên phim phải lật từng trang. Có index giống có mục lục: nhảy thẳng tới vùng dữ liệu cần tìm.",
        "FFF8E8",
    )
    add_matrix(
        doc,
        ["Index", "Tối ưu cho query nào", "Lưu ý"],
        [
            ["idx_showtimes_start_time", "Lọc suất chiếu theo ngày/giờ.", "Rất quan trọng cho lịch chiếu."],
            ["idx_seat_status_lookup(showtime_id,status)", "Lấy ghế theo suất và trạng thái.", "Quan trọng cho seat map/check ghế."],
            ["idx_bookings_user", "Vé/đơn của tôi.", "Tối ưu trang MyBookings."],
            ["idx_payments_created_at/status/method", "Admin lọc thanh toán theo ngày/trạng thái/cổng.", "Giảm scan bảng payment lớn."],
            ["idx_staff_cinemas_staff/cinema", "Staff scope.", "Kiểm tra assignment nhanh."],
            ["unique email/username", "Đăng ký/login.", "Vừa chống trùng vừa tăng tốc tìm kiếm."],
        ],
        [3000, 3800, 2560],
    )
    doc.add_heading("Flyway Migration", level=1)
    add_bullets(
        doc,
        [
            "Flyway đọc các file trong src/main/resources/db/migration.",
            "File đặt tên V1__create_cinema_schema.sql, V2__..., V10__...",
            "Khi app khởi động, Flyway kiểm tra bảng flyway_schema_history để biết migration nào đã chạy.",
            "Không sửa migration cũ đã chạy ở môi trường khác; thêm file V tiếp theo để thay đổi schema.",
            "database/database.sql là file tổng hợp để tạo mới nhanh; db/migration là nguồn chuẩn cho triển khai app tự migrate.",
        ],
    )
    doc.add_heading("SQL profiling / EXPLAIN ANALYZE nên dùng khi nào?", level=1)
    add_bullets(
        doc,
        [
            "Khi một API chậm hoặc dashboard tính lâu.",
            "Khi nghi ngờ query scan toàn bảng dù có điều kiện lọc.",
            "Khi thêm index mới và muốn chứng minh index có tác dụng.",
            "Khi số lượng booking/payment/seat_status tăng lớn.",
        ],
    )
    add_code(
        doc,
        """
EXPLAIN ANALYZE
SELECT *
FROM showtimes
WHERE start_time >= '2026-08-01'
  AND start_time < '2026-08-08'
  AND is_deleted = false;
        """,
    )
    add_bullets(
        doc,
        [
            "EXPLAIN cho biết kế hoạch chạy query.",
            "ANALYZE chạy thật query rồi đo thời gian thật.",
            "Nếu thấy Seq Scan trên bảng lớn, có thể cần index phù hợp.",
            "Không nên thêm index vô tội vạ vì index làm insert/update chậm hơn và tốn dung lượng.",
        ],
    )
    add_common_tail(
        doc,
        [
            ("Cache khác index thế nào?", "Cache nằm ở tầng app/cache provider, index nằm trong DB."),
            ("Vì sao không cache seat_status?", "Vì ghế thay đổi realtime, cache stale sẽ gây đặt trùng hoặc UI sai."),
            ("Flyway để làm gì?", "Quản lý version schema để môi trường dev/test/prod đồng bộ."),
            ("N+1 query là gì?", "1 query lấy danh sách, rồi N query phụ lấy quan hệ từng item."),
            ("Khi nào thêm Redis?", "Khi chạy nhiều instance backend hoặc cần cache phân tán, session/rate limit chia sẻ."),
        ],
        [
            "Mở CacheConfig.java và đọc cache name.",
            "Tìm @Cacheable/@CacheEvict trong service/repository.",
            "Mở db/migration và hiểu thứ tự V1, V2, ...",
            "Giải thích được tại sao index giúp lọc nhanh.",
            "Nêu được 2 cách tránh N+1 trong JPA.",
        ],
    )
    return save_doc(doc, "Luong_Cache_Query_Index_Flyway_CinemaBooking.docx")


def build_frontend_doc() -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(
        doc,
        "Kiến trúc Frontend React, Router, State, API và UX",
        "Giải thích cách React client tổ chức route, layout, protected route, Zustand auth store, Axios interceptor, page admin/staff/user và luồng dữ liệu từ API lên giao diện.",
    )
    add_callout(
        doc,
        "Hiểu nhanh",
        "Frontend React là lớp trải nghiệm người dùng. Nó không quyết định bảo mật cuối cùng, nhưng nó giúp user thấy đúng màn hình, gọi đúng API, hiển thị đúng loading/error và tự refresh token mượt.",
    )
    doc.add_heading("Cấu trúc thư mục frontend", level=1)
    add_matrix(
        doc,
        ["Thư mục/File", "Vai trò", "Cách học"],
        [
            ["src/pages/public", "Trang public: home, phim, rạp, login/register.", "Học luồng khách chưa đăng nhập."],
            ["src/pages/user", "Trang user đã login: chọn ghế, checkout, vé của tôi, profile.", "Học luồng đặt vé end-to-end."],
            ["src/pages/admin", "Trang quản trị phim/rạp/suất/booking/payment/user/audit.", "Học CRUD, filter, pagination."],
            ["src/pages/staff", "Trang nhân viên: rạp phụ trách, soát QR.", "Học staff scope và scanner UX."],
            ["src/api", "Từng module gọi backend.", "Mỗi api file tương ứng domain backend."],
            ["src/stores", "Zustand state global.", "Hiện quan trọng nhất là authStore."],
            ["src/router/AppRouter.tsx", "Khai báo route và permission.", "Đọc để hiểu app có những màn nào."],
            ["src/components/layout", "Layout public/admin/auth.", "Hiểu shell giao diện."],
            ["src/hooks", "Hook tái sử dụng như WebSocket/debounce.", "Học logic không gắn chặt page."],
        ],
        [2300, 3400, 3660],
    )
    doc.add_heading("Router và lazy loading", level=1)
    add_code(
        doc,
        """
const HomePage = lazy(() => import('../pages/public/HomePage'));

<Suspense fallback={<PageLoader />}>
  <Routes>
    <Route element={<AuthLayout />}>...</Route>
    <Route element={<AdminLayout />}>...</Route>
    <Route element={<PublicLayout />}>...</Route>
  </Routes>
</Suspense>
        """,
    )
    add_bullets(
        doc,
        [
            "lazy(): tách code theo page, không tải toàn bộ app ngay lần đầu.",
            "Suspense: hiển thị PageLoader trong lúc bundle page đang tải.",
            "AuthLayout: trang login/register không có navbar/footer.",
            "AdminLayout: sidebar quản trị cho admin/staff.",
            "PublicLayout: navbar/footer cho trang khách hàng.",
        ],
    )
    doc.add_heading("ProtectedRoute bảo vệ màn hình", level=1)
    add_code(
        doc,
        """
if (!token || isTokenExpired(token)) {
  return <Navigate to="/login" state={{ from: location }} replace />;
}

if (permission && !hasPermission(permission)) {
  return <Navigate to="/" replace />;
}
        """,
    )
    add_bullets(
        doc,
        [
            "Nếu chưa có token hoặc token hết hạn, user bị chuyển về login.",
            "Nếu route yêu cầu permission mà user không có, chuyển về trang chủ.",
            "State from có thể dùng để sau login quay lại trang trước.",
            "Đây là UX guard. Backend vẫn phải kiểm tra thật bằng JWT và @PreAuthorize.",
        ],
    )
    doc.add_heading("authStore quản lý trạng thái đăng nhập", level=1)
    add_bullets(
        doc,
        [
            "token: access token dùng gắn Authorization header.",
            "refreshToken: fallback trong localStorage nếu backend trả, nhưng luồng mới ưu tiên HttpOnly cookie.",
            "user: thông tin user để hiển thị avatar/tên.",
            "permissions: danh sách quyền để ẩn/hiện menu/nút.",
            "login(): lưu token/user/permissions vào localStorage và state.",
            "updateTokens(): cập nhật access token mới sau refresh.",
            "logout(): dọn localStorage và chuyển /login.",
            "hasPermission(): kiểm tra quyền trên frontend.",
        ],
    )
    doc.add_heading("axiosClient là cầu nối API", level=1)
    add_code(
        doc,
        """
axiosClient.interceptors.request.use(config => {
  const token = useAuthStore.getState().token;
  if (token && !config.headers.Authorization) {
    config.headers.Authorization = `Bearer ${token}`;
  }
  return config;
});
        """,
    )
    add_bullets(
        doc,
        [
            "Request interceptor tự gắn token, page không cần tự set header.",
            "Response interceptor xử lý 401 tập trung, tự refresh token và retry.",
            "isAuthAttempt tránh vòng lặp refresh khi chính /auth/refresh bị lỗi.",
            "withCredentials=true để browser gửi HttpOnly refresh cookie lên backend.",
        ],
    )
    doc.add_heading("Luồng dữ liệu của một page điển hình", level=1)
    add_code(
        doc,
        """
User mở trang
  -> React component mount
  -> useEffect gọi api module
  -> axiosClient gửi request
  -> backend trả ApiResponse.result
  -> page setState
  -> render loading / empty / list / error
        """,
    )
    add_bullets(
        doc,
        [
            "Page nên giữ state UI: filters, selected item, modal open, loading.",
            "API module chỉ nên biết endpoint và request/response type.",
            "Utils format tiền/ngày giúp UI đồng nhất.",
            "Toast dùng cho feedback nhanh: áp mã thành công, lỗi thanh toán, check-in thất bại...",
        ],
    )
    doc.add_heading("Các luồng frontend quan trọng", level=1)
    add_matrix(
        doc,
        ["Luồng", "File chính", "Điểm học"],
        [
            ["Đặt vé", "MovieDetailPage/CinemaDetailPage -> SeatSelectionPage -> CheckoutPage -> PaymentResultPage", "Đi từ public showtime tới giữ ghế, booking, payment."],
            ["Vé của tôi", "MyBookingsPage, TicketDetailPage", "Phân loại vé hợp lệ/đơn đã đặt, QR, trạng thái hoàn tiền/hủy."],
            ["Admin", "Admin*Page.tsx", "Filter, pagination, modal CRUD, permission theo nút."],
            ["Staff scanner", "StaffTicketScannerPage.tsx", "Chọn city/cinema/showtime, camera/file QR, gọi API check-in có scope."],
            ["Profile", "ProfilePage.tsx", "Cập nhật thông tin cá nhân, đổi mật khẩu, session."],
            ["Map/rạp", "CinemaMapPage, HomePage", "Tìm rạp, map, thành phố không dấu, horizontal/vertical scroll UX."],
        ],
        [1900, 4100, 3360],
    )
    doc.add_heading("Nguyên tắc UX/product đã dùng", level=1)
    add_bullets(
        doc,
        [
            "Không bắt login khi chỉ xem phim/rạp/lịch chiếu.",
            "Chỉ bắt login khi chọn ghế/checkout/xem vé.",
            "Payment QR khóa mã giảm giá sau khi QR đã tạo để tránh chuyển sai số tiền.",
            "Staff scanner bắt chọn đúng rạp/suất trước khi check-in để tránh quét nhầm vé.",
            "Admin/staff menu chỉ hiện những mục user có quyền.",
            "Form admin có filter ngày/trạng thái để vận hành dữ liệu lớn.",
        ],
    )
    add_common_tail(
        doc,
        [
            ("Vì sao dùng lazy loading?", "Giảm bundle ban đầu, trang mở nhanh hơn."),
            ("ProtectedRoute có thay backend security không?", "Không, chỉ là UX guard."),
            ("Vì sao dùng axios interceptor?", "Tập trung logic gắn token/refresh token, tránh lặp ở từng API."),
            ("Zustand khác Context?", "Zustand nhẹ, dễ đọc state ngoài React component như trong interceptor."),
            ("Khi API trả lỗi validation thì page nên làm gì?", "Hiển thị message đúng field hoặc toast thân thiện, không reload trang."),
        ],
        [
            "Mở AppRouter.tsx và đọc toàn bộ route.",
            "Mở authStore.ts và giải thích login/logout/updateTokens.",
            "Mở axiosClient.ts và giải thích request/response interceptor.",
            "Chọn một page admin và mô tả state/filter/API/render.",
            "Tự demo luồng user từ homepage tới thanh toán thành công.",
        ],
    )
    return save_doc(doc, "Luong_Frontend_React_Router_State_API_UX_CinemaBooking.docx")


def build_email_doc() -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(
        doc,
        "Luồng Email: xác thực tài khoản, reset mật khẩu, gửi vé và thông báo hủy suất",
        "Giải thích cách backend tạo token email, render template HTML, gửi mail bất đồng bộ và những điểm cần bảo mật.",
    )
    add_callout(
        doc,
        "Hiểu nhanh",
        "Email là kênh xác nhận ngoài app. Đăng ký cần xác thực email, quên mật khẩu cần reset link, mua vé thành công cần gửi vé/QR, hủy suất cần thông báo và hướng dẫn hoàn tiền.",
    )
    doc.add_heading("Các file/class chính", level=1)
    add_matrix(
        doc,
        ["Class/File", "Vai trò", "Điểm học"],
        [
            ["EmailService.java", "Interface gửi email.", "Controller/service chỉ phụ thuộc interface, dễ đổi provider."],
            ["EmailServiceImpl.java", "Render template và gửi mail.", "Dùng @Async để gửi nền, @Transactional(readOnly=true) khi cần load booking/ticket."],
            ["templates/email-verification.html", "Template xác thực email.", "Có link verify token."],
            ["templates/password-reset.html", "Template reset mật khẩu.", "Có link reset token hết hạn."],
            ["templates/ticket-email.html", "Template vé điện tử.", "Chứa thông tin phim/rạp/phòng/ghế/QR."],
            ["templates/showtime-cancellation-email.html", "Template hủy suất.", "Thông báo lý do và trạng thái xử lý hoàn tiền."],
            ["UserServiceImpl.java", "Tạo user, token verification/reset.", "Token nên lưu hash, không lưu raw."],
        ],
        [2700, 3200, 3460],
    )
    doc.add_heading("Luồng đăng ký và xác thực email", level=1)
    add_code(
        doc,
        """
User đăng ký
  -> RegisterPage.tsx gọi /api/v1/users/register
  -> UserServiceImpl tạo user emailVerified=false
  -> sinh verification token raw
  -> lưu hash token + expiresAt vào user
  -> EmailServiceImpl gửi link verify qua email

User bấm link
  -> VerifyEmailPage.tsx đọc token từ URL
  -> gọi /api/v1/users/verify-email
  -> backend hash token nhập vào, so với hash trong DB
  -> nếu đúng và chưa hết hạn: emailVerified=true
        """,
    )
    add_bullets(
        doc,
        [
            "Không nên cho user login nếu email chưa xác thực, AuthenticationService.validateUserCanAuthenticate sẽ chặn.",
            "Token gửi qua email là raw token, DB chỉ nên lưu hash để nếu DB lộ cũng không dùng được link.",
            "expiresAt giới hạn thời gian token sống, giảm rủi ro link cũ bị dùng lại.",
        ],
    )
    doc.add_heading("Luồng quên mật khẩu", level=1)
    add_code(
        doc,
        """
ForgotPasswordPage
  -> nhập email
  -> backend nếu email tồn tại thì tạo password_reset_token
  -> gửi link reset

ResetPasswordPage
  -> nhập mật khẩu mới
  -> backend kiểm tra token hash + expiresAt
  -> BCrypt hash password mới
  -> xóa reset token khỏi user
        """,
    )
    add_bullets(
        doc,
        [
            "Không nên trả lời quá rõ email có tồn tại hay không để tránh enumeration.",
            "Mật khẩu mới phải validate đủ mạnh ở frontend và backend.",
            "Sau reset password thực tế nên revoke các refresh token cũ để bảo mật hơn.",
        ],
    )
    doc.add_heading("Luồng gửi vé qua email", level=1)
    add_bullets(
        doc,
        [
            "Khi payment SUCCESS, backend chuyển booking SUCCESS, seat_status BOOKED, sinh ticket QR.",
            "EmailServiceImpl load booking kèm showtime/movie/room/cinema/booking details/tickets.",
            "Template ticket-email.html hiển thị thông tin vé: phim, rạp, địa chỉ, giờ chiếu, ghế, tổng tiền, QR từng vé.",
            "Email nên gửi sau commit transaction thanh toán, tránh gửi vé khi DB rollback.",
            "Nếu gửi email fail, booking vẫn thành công; hệ thống có thể log và cho user xem vé trong app.",
        ],
    )
    doc.add_heading("Luồng hủy suất và email hoàn tiền", level=1)
    add_bullets(
        doc,
        [
            "Admin/staff hủy suất bằng policy.",
            "Booking SUCCESS bị chuyển CANCELLED, ticket ACTIVE bị chuyển CANCELLED.",
            "Payment SUCCESS được ghi PaymentEventType.REFUND_REQUESTED.",
            "Email hủy suất gửi lý do, thông tin suất chiếu và trạng thái đang xử lý hoàn tiền.",
        ],
    )
    doc.add_heading("Async email cần hiểu", level=1)
    add_bullets(
        doc,
        [
            "@Async giúp request chính không phải chờ SMTP gửi xong.",
            "@EnableAsync nằm trong AsyncConfig, cho phép Spring chạy method async bằng thread pool.",
            "Vì async chạy ở thread khác, nếu cần đọc lazy relation thì method nên mở transaction read-only riêng.",
            "Không nên để lỗi gửi email làm rollback payment đã thành công.",
        ],
    )
    add_common_tail(
        doc,
        [
            ("Vì sao email token nên lưu hash?", "Nếu DB lộ thì attacker không có raw token để xác thực/reset."),
            ("Nếu gửi vé qua email lỗi thì booking có hủy không?", "Không. Vé vẫn nằm trong app; email là kênh thông báo bổ sung."),
            ("Vì sao gửi email sau commit?", "Để tránh gửi thông tin không tồn tại nếu transaction rollback."),
            ("Email verification bảo vệ gì?", "Giảm tài khoản rác và đảm bảo user sở hữu email."),
            ("Có nên dùng queue cho email?", "Khi scale lớn nên dùng queue/retry, hiện @Async đủ cho monolith/demo."),
        ],
        [
            "Mở EmailServiceImpl.java và tìm các method gửi mail.",
            "Mở template ticket-email.html để xem dữ liệu vé render thế nào.",
            "Giải thích vì sao login chặn emailVerified=false.",
            "Tự mô tả luồng reset password từ trang frontend tới DB.",
            "Nêu được xử lý khi SMTP lỗi.",
        ],
    )
    return save_doc(doc, "Luong_Email_Verification_Reset_Ticket_Notification_CinemaBooking.docx")


def build_admin_staff_doc() -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(
        doc,
        "Luồng Admin/Staff vận hành rạp: suất chiếu, phòng ghế, booking, payment và soát vé",
        "Giải thích cách các màn quản trị/vận hành liên kết với backend, permission và staff cinema assignment.",
    )
    add_callout(
        doc,
        "Hiểu nhanh",
        "Admin là người quản trị toàn hệ thống. Staff là người vận hành theo rạp được phân công. Cùng vào admin portal nhưng menu, dữ liệu và hành động được giới hạn theo permission + scope.",
    )
    doc.add_heading("Admin portal có những nhóm màn hình nào?", level=1)
    add_matrix(
        doc,
        ["Màn hình", "Người dùng chính", "Chức năng"],
        [
            ["Tổng quan", "Admin/Staff có DASHBOARD_VIEW", "KPI doanh thu, vé bán, suất chiếu, top phim."],
            ["Phim chiếu rạp", "Admin", "CRUD phim, trạng thái NOW_SHOWING/COMING_SOON/ENDED."],
            ["Rạp chiếu", "Admin", "CRUD rạp, tọa độ map, thành phố, địa chỉ."],
            ["Phòng & ghế", "Admin/Staff scope", "Quản lý phòng, sinh ghế theo layout, xem ghế."],
            ["Suất chiếu", "Admin/Staff scope", "Tạo/sửa/hủy suất, lọc ngày/rạp/phim/status."],
            ["Đơn đặt vé", "Admin/Staff scope", "Xem booking, trạng thái, chi tiết ghế/vé."],
            ["Thanh toán", "Admin/Staff scope", "Xem payment, event, đối soát."],
            ["Người dùng", "Admin", "Tạo/sửa/khóa user, set role, gán rạp cho staff."],
            ["Nhật ký", "Admin", "Audit log admin/auth/payment."],
            ["Rạp phụ trách", "Staff", "Staff xem mình được phân công rạp nào."],
            ["Soát vé QR", "Staff", "Chọn rạp/suất hôm nay, quét QR bằng camera/file."],
        ],
        [2200, 2500, 4660],
    )
    doc.add_heading("Tạo suất chiếu chuẩn product", level=1)
    add_code(
        doc,
        """
Admin/Staff mở form tạo suất
  -> chọn phim
  -> chọn thành phố
  -> chọn rạp theo thành phố/scope
  -> chọn phòng thuộc rạp
  -> chọn giờ bắt đầu/kết thúc
  -> nhập giá vé cơ bản
  -> backend kiểm tra:
       endTime > startTime
       phòng không overlap + cleaning buffer 15 phút
       staff có quyền ở rạp đó
  -> tạo showtime
  -> clone toàn bộ seats của room thành seat_status AVAILABLE
        """,
    )
    add_bullets(
        doc,
        [
            "Tách seats vật lý và seat_status theo suất là thiết kế đúng: ghế A1 tồn tại trong phòng, nhưng trạng thái A1 khác nhau theo từng suất.",
            "Cleaning buffer 15 phút tránh xếp hai suất sát nhau làm không có thời gian dọn phòng.",
            "Staff chỉ thấy rạp/phòng được phân công để tránh tạo suất sai rạp.",
        ],
    )
    doc.add_heading("Hủy suất có policy", level=1)
    add_bullets(
        doc,
        [
            "Không hủy nếu suất đã có ticket USED.",
            "Booking PENDING/SUCCESS bị chuyển CANCELLED.",
            "Ghế trả AVAILABLE để không còn giữ/bán cho suất bị hủy.",
            "Ticket ACTIVE chuyển CANCELLED.",
            "Payment SUCCESS ghi refund requested để admin xử lý hoàn tiền.",
            "User ở trang Vé của tôi nên thấy trạng thái hủy/đang xử lý hoàn tiền, không hiện nút chọn lại ghế cho suất đã hủy/kết thúc.",
        ],
    )
    doc.add_heading("Soát vé QR chuẩn", level=1)
    add_code(
        doc,
        """
Staff chọn:
  - thành phố
  - rạp được phân công
  - suất chiếu hôm nay đang mở check-in

Quét QR:
  -> QR hợp lệ?
  -> ticket ACTIVE?
  -> booking SUCCESS?
  -> đúng rạp?
  -> đúng suất chiếu?
  -> trong cửa sổ check-in?
  -> set ticket USED + check_in_time
        """,
    )
    add_bullets(
        doc,
        [
            "Chọn rạp/suất trước khi quét giúp tránh quét nhầm vé ở rạp khác hoặc suất khác.",
            "Vé đã USED không cho dùng lại, trả thông tin giờ đã quét nếu có.",
            "Camera và file upload đều là nguồn QR, nhưng backend check-in mới quyết định hợp lệ.",
        ],
    )
    doc.add_heading("Vì sao staff không nên có toàn quyền?", level=1)
    add_bullets(
        doc,
        [
            "Staff vận hành tại rạp, không quản trị toàn hệ thống.",
            "Staff có thể xem/tạo/sửa/hủy suất trong phạm vi rạp phụ trách.",
            "Staff có thể xem booking/payment thuộc rạp mình để hỗ trợ khách.",
            "Staff không nên quản lý user/role, xóa dữ liệu master toàn hệ thống hoặc refund thật nếu chưa có quy trình duyệt.",
        ],
    )
    add_common_tail(
        doc,
        [
            ("Vì sao tạo showtime phải sinh seat_status?", "Vì mỗi suất có trạng thái ghế riêng."),
            ("Vì sao staff cần chọn rạp trước khi soát QR?", "Để kiểm tra vé thuộc đúng rạp/suất, tránh check-in nhầm."),
            ("Vì sao hủy suất không xóa booking?", "Booking/payment là lịch sử giao dịch và phục vụ hoàn tiền/audit."),
            ("Cleaning buffer dùng để làm gì?", "Tránh overlap thực tế giữa hai suất vì cần dọn phòng/chuyển khách."),
            ("Staff xem rạp phụ trách ở đâu?", "Trang StaffAssignedCinemasPage trong admin portal."),
        ],
        [
            "Demo được staff chỉ thấy rạp được gán.",
            "Tạo thử suất bằng staff và giải thích vì sao chỉ chọn được rạp scope.",
            "Hủy suất có booking và mô tả trạng thái booking/payment/ticket.",
            "Soát QR đúng/sai rạp/sai suất/đã dùng.",
            "Giải thích cho hội đồng sự khác nhau giữa Admin và Staff.",
        ],
    )
    return save_doc(doc, "Luong_Admin_Staff_Operation_CinemaBooking.docx")


def build_study_roadmap_doc() -> Path:
    doc = Document()
    apply_styles(doc)
    add_title(
        doc,
        "Lộ trình học toàn bộ hệ thống CinemaBooking.vn",
        "Bản đồ đọc tài liệu và đọc code theo thứ tự để hiểu sâu backend Spring Boot, frontend React, bảo mật, thanh toán, realtime và vận hành product.",
    )
    add_callout(
        doc,
        "Cách dùng bộ tài liệu",
        "Đừng đọc ngẫu nhiên. Hãy đi theo thứ tự: kiến trúc tổng quan -> Auth/JWT -> RBAC -> Booking/Payment/WebSocket -> Scheduler -> Admin/Staff -> Exception/Audit -> Cache/Query -> Frontend. Mỗi vòng đọc hãy chạy app và tự demo lại một luồng.",
    )
    doc.add_heading("Bộ file nên đọc theo thứ tự", level=1)
    add_matrix(
        doc,
        ["Thứ tự", "File Word", "Bạn sẽ hiểu được gì"],
        [
            ["1", "Huong_dan_kien_truc_va_luong_hoat_dong_CinemaBookingSystem.docx", "Bức tranh tổng thể: database, backend, frontend, luồng user/admin/staff."],
            ["2", "Luong_Auth_JWT_Session_GoogleLogin_CinemaBooking.docx", "Đăng nhập, JWT, refresh token, logout, Google login, bảo mật phiên."],
            ["3", "Luong_RBAC_Permission_StaffScope_CinemaBooking.docx", "Role/permission, @PreAuthorize, quyền frontend, staff chỉ thao tác rạp được gán."],
            ["4", "Luong_thanh_toan_VNPay_SePay_CinemaBooking.docx", "Thanh toán VNPay và SePay/VietQR từ frontend đến backend/callback/webhook."],
            ["5", "Luong_WebSocket_Realtime_SeatMap_CinemaBooking.docx hoặc bản _hoc_nhanh", "Realtime ghế HOLD/BOOKED/AVAILABLE qua WebSocket."],
            ["6", "Luong_Scheduler_GiuGhe_HetHan_TrangThaiSuatChieu_CinemaBooking.docx", "Job nền trả ghế, expire booking, sync trạng thái suất chiếu, dọn token."],
            ["7", "Luong_Admin_Staff_Operation_CinemaBooking.docx", "Cách admin/staff vận hành rạp, tạo suất, soát vé, hủy suất, hoàn tiền thủ công."],
            ["8", "Luong_Exception_Audit_SoftDelete_CinemaBooking.docx", "Chuẩn hóa lỗi, audit log, payment event, xóa mềm/xóa cứng."],
            ["9", "Luong_Cache_Query_Index_Flyway_CinemaBooking.docx", "Cache, index, tránh N+1, Flyway migration, profiling query."],
            ["10", "Luong_Email_Verification_Reset_Ticket_Notification_CinemaBooking.docx", "Email xác thực, reset password, gửi vé, thông báo hủy suất."],
            ["11", "Luong_Frontend_React_Router_State_API_UX_CinemaBooking.docx", "Router, layout, ProtectedRoute, Zustand, axios interceptor, UX page."],
        ],
        [900, 4300, 4160],
    )
    doc.add_heading("Lộ trình học 7 ngày", level=1)
    add_matrix(
        doc,
        ["Ngày", "Mục tiêu", "Việc cần làm"],
        [
            ["Ngày 1", "Hiểu tổng quan hệ thống", "Đọc file kiến trúc tổng quan, mở database.sql, vẽ lại ERD đơn giản."],
            ["Ngày 2", "Hiểu Auth/JWT/RBAC", "Đọc Auth + RBAC, mở SecurityConfig, AuthenticationService, PermissionName, AppRouter."],
            ["Ngày 3", "Hiểu booking/seat/payment", "Demo chọn ghế -> checkout -> VNPay/SePay; đọc file payment và BookingServiceImpl."],
            ["Ngày 4", "Hiểu realtime/scheduler", "Mở hai browser đặt cùng suất, xem WebSocket đổi ghế; đọc scheduler docs."],
            ["Ngày 5", "Hiểu admin/staff", "Login admin/staff, tạo suất, gán rạp, soát QR; giải thích staff scope."],
            ["Ngày 6", "Hiểu exception/audit/cache/query", "Cố tình tạo lỗi, xem response; mở audit logs; xem migration/index."],
            ["Ngày 7", "Tập bảo vệ/phỏng vấn", "Tự trả lời checklist cuối mỗi file, quay demo 5-7 phút hoặc tập thuyết trình."],
        ],
        [1100, 3000, 5260],
    )
    doc.add_heading("Thứ tự đọc code backend", level=1)
    add_numbered(
        doc,
        [
            "Bắt đầu từ CinemaBookingSystemApplication.java để biết app bật @EnableScheduling, @EnableJpaAuditing và config properties.",
            "Đọc entity trong src/main/java/com/cinema/booking/entity để hiểu dữ liệu.",
            "Đọc repository để biết query lấy dữ liệu và chỗ tối ưu fetch/projection/index.",
            "Đọc service impl vì đây là nơi nghiệp vụ thật nằm: BookingServiceImpl, PaymentServiceImpl, ShowtimeServiceImpl, UserServiceImpl.",
            "Đọc controller để biết API public/protected và @PreAuthorize.",
            "Đọc configuration/security để hiểu JWT, cache, CORS/proxy, WebSocket.",
            "Đọc scheduler/task để hiểu các việc tự chạy nền.",
            "Đọc exception/audit/payment event để hiểu product thật xử lý lỗi và truy vết.",
            "Cuối cùng đọc test để biết hệ thống đang được kiểm chứng bằng tình huống nào.",
        ],
    )
    doc.add_heading("Thứ tự đọc code frontend", level=1)
    add_numbered(
        doc,
        [
            "Đọc main.tsx và App.tsx để biết React app được mount ra sao.",
            "Đọc router/AppRouter.tsx để thấy toàn bộ đường đi của app.",
            "Đọc stores/authStore.ts và api/axiosClient.ts để hiểu login/token/refresh.",
            "Đọc api/*.ts theo domain: movieApi, bookingApi, paymentApi, userApi.",
            "Đọc pages/public để hiểu user chưa đăng nhập xem phim/rạp.",
            "Đọc pages/user để hiểu đặt vé, checkout, vé của tôi, profile.",
            "Đọc pages/admin và pages/staff để hiểu vận hành product.",
            "Đọc hooks/useSeatWebSocket.ts để nối với tài liệu WebSocket.",
            "Đọc utils/format.ts để hiểu format tiền/ngày thống nhất.",
        ],
    )
    doc.add_heading("Checklist bảo vệ nhanh", level=1)
    add_bullets(
        doc,
        [
            "Giải thích được vì sao hệ thống tách seats và seat_status.",
            "Giải thích được vì sao refresh token lưu hash và rotate.",
            "Giải thích được vì sao staff phải có staff_cinemas scope.",
            "Giải thích được vì sao SePay cần webhook còn VNPay dùng callback/return URL.",
            "Giải thích được vì sao WebSocket publish sau commit.",
            "Giải thích được scheduler làm gì khi user bỏ thanh toán.",
            "Giải thích được soft delete dùng cho dữ liệu nghiệp vụ quan trọng.",
            "Giải thích được @PreAuthorize khác ProtectedRoute.",
            "Giải thích được index/cache/Flyway giúp product triển khai và scale thế nào.",
            "Demo được 3 luồng: user đặt vé, staff soát vé, admin hủy suất/xem hoàn tiền.",
        ],
    )
    doc.add_heading("Câu trả lời mẫu khi phỏng vấn", level=1)
    add_matrix(
        doc,
        ["Câu hỏi", "Cách trả lời ngắn gọn"],
        [
            ["Dự án này có gì giống product thật?", "Có RBAC, staff scope theo rạp, giữ ghế timeout, thanh toán gateway, webhook, ticket QR, audit log, scheduler, cache/index và realtime seat map."],
            ["Bạn tối ưu N+1 thế nào?", "Dùng fetch join/projection/DTO query cho màn danh sách cần nested data, cache master data, index theo filter phổ biến."],
            ["Nếu payment callback bị mất thì sao?", "Booking có timeout/scheduler; với SePay webhook tự xác nhận, payment event giúp đối soát. Product lớn có thêm reconciliation job."],
            ["Nếu hai user chọn cùng ghế?", "DB update có điều kiện status/version/transaction; chỉ một người giữ được, WebSocket báo realtime cho người còn lại."],
            ["Vì sao không cho staff toàn quyền?", "Nguyên tắc least privilege: staff chỉ vận hành rạp được giao, admin mới quản trị toàn hệ thống."],
        ],
        [3400, 5960],
    )
    return save_doc(doc, "Lo_trinh_hoc_toan_bo_he_thong_CinemaBooking.docx")


def main() -> None:
    generated = [
        build_study_roadmap_doc(),
        build_auth_doc(),
        build_rbac_doc(),
        build_scheduler_doc(),
        build_exception_audit_doc(),
        build_cache_query_doc(),
        build_frontend_doc(),
        build_email_doc(),
        build_admin_staff_doc(),
    ]
    print("Generated documents:")
    for path in generated:
        print(f"- {path}")


if __name__ == "__main__":
    main()

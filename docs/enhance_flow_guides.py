from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PAYMENT_DOC = "docs/Luong_thanh_toan_VNPay_SePay_CinemaBooking.docx"
WEBSOCKET_DOC = "docs/Luong_WebSocket_Realtime_SeatMap_CinemaBooking.docx"
WEBSOCKET_FALLBACK_DOC = "docs/Luong_WebSocket_Realtime_SeatMap_CinemaBooking_hoc_nhanh.docx"
LEARNING_MARKER = "PHẦN HỌC NHANH VÀ ÔN BẢO VỆ"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="111827", size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def setup_learning_styles(doc):
    styles = doc.styles
    if "LearningCode" not in styles:
        style = styles.add_style("LearningCode", 1)
        style.font.name = "Consolas"
        style.font.size = Pt(8.5)
        style.font.color.rgb = RGBColor.from_string("111827")
        style.paragraph_format.left_indent = Inches(0.18)
        style.paragraph_format.right_indent = Inches(0.18)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(8)


def add_code(doc, text):
    paragraph = doc.add_paragraph(style="LearningCode")
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    return paragraph


def add_callout(doc, title, body, fill="EFF6FF", title_color="1D4ED8"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""

    title_paragraph = cell.paragraphs[0]
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor.from_string(title_color)

    body_paragraph = cell.add_paragraph()
    body_paragraph.paragraph_format.space_before = Pt(3)
    body_run = body_paragraph.add_run(body)
    body_run.font.name = "Arial"
    body_run.font.size = Pt(9.5)
    body_run.font.color.rgb = RGBColor.from_string("1F2937")
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, True, "FFFFFF", 9)
        shade_cell(header_cells[index], "111827")
        if widths:
            header_cells[index].width = widths[index]

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, index == 0, "111827", 8.8)
            if widths:
                cells[index].width = widths[index]
    doc.add_paragraph()
    return table


def add_section_marker(doc):
    doc.add_page_break()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("PHẦN HỌC NHANH VÀ ÔN BẢO VỆ")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("0F172A")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Phần này tóm tắt lại tài liệu theo cách dễ học: nhìn sơ đồ trước, sau đó nối sang class và code."
    ).italic = True


def has_learning_section(doc):
    return any(LEARNING_MARKER in paragraph.text for paragraph in doc.paragraphs)


def save_with_fallback(doc, primary_path, fallback_path=None):
    try:
        doc.save(primary_path)
        return primary_path
    except PermissionError:
        if not fallback_path:
            raise
        doc.save(fallback_path)
        return fallback_path


def enhance_payment_doc():
    doc = Document(PAYMENT_DOC)
    if has_learning_section(doc):
        print(f"Skipped existing learning section: {PAYMENT_DOC}")
        return PAYMENT_DOC
    setup_learning_styles(doc)
    add_section_marker(doc)

    doc.add_heading("A. Cách đọc luồng thanh toán trong 15 phút", level=1)
    add_numbered(
        doc,
        [
            "Đọc sơ đồ tổng quan để biết booking, payment, ghế và vé liên quan với nhau như thế nào.",
            "Đọc riêng luồng VNPay để hiểu kiểu redirect: user đi sang cổng thanh toán rồi callback quay về backend.",
            "Đọc riêng luồng Quét QR ngân hàng để hiểu kiểu webhook: user chuyển khoản, SePay báo tiền vào cho backend.",
            "Đọc bảng câu hỏi bảo vệ để tập trả lời ngắn gọn: vì sao kiểm tra amount, vì sao cần idempotency, vì sao QR phải tạo lại khi đổi mã giảm giá.",
        ],
    )

    add_callout(
        doc,
        "Ý chính cần nhớ",
        "Frontend chỉ hiển thị và gửi yêu cầu. Backend mới là nơi quyết định booking có được thành công hay không. Backend luôn kiểm tra user, trạng thái booking, hạn giữ vé, số tiền, chữ ký callback/webhook và trạng thái payment trước khi sinh vé.",
        "ECFDF5",
        "047857",
    )

    doc.add_heading("B. Sơ đồ dữ liệu: từ giữ ghế đến có vé", level=1)
    add_code(
        doc,
        """User chọn ghế
  -> seat_status = HOLD
  -> booking = PENDING
  -> payment = PENDING
  -> thanh toán thành công
  -> payment = SUCCESS
  -> booking = SUCCESS
  -> seat_status = BOOKED
  -> sinh ticket QR
  -> gửi email vé cho khách"""
    )
    doc.add_paragraph(
        "Nếu thanh toán thất bại hoặc hết hạn, hướng đi sẽ ngược lại: payment FAILED/EXPIRED, booking FAILED/EXPIRED, ghế được trả về AVAILABLE để người khác có thể đặt."
    )

    doc.add_heading("C. Sơ đồ luồng VNPay dễ nhớ", level=1)
    add_code(
        doc,
        """CheckoutPage
  -> paymentApi.initiatePayment(method = VNPAY)
  -> PaymentController.initiatePayment()
  -> PaymentServiceImpl.initiatePayment()
  -> VnPayPaymentGateway.createPaymentUrl()
  -> frontend redirect sang VNPay
  -> user thanh toán trên VNPay
  -> VNPay callback về /payments/vnpay-callback
  -> PaymentServiceImpl.handleVNPayCallback()
  -> kiểm tra chữ ký + amount + hạn giữ vé
  -> BookingServiceImpl.handlePaymentSuccess()
  -> booking SUCCESS, ghế BOOKED, sinh ticket"""
    )

    doc.add_heading("D. Sơ đồ luồng Quét QR ngân hàng dễ nhớ", level=1)
    add_code(
        doc,
        """CheckoutPage
  -> paymentApi.initiatePayment(method = SEPAY)
  -> PaymentServiceImpl.initiatePayment()
  -> SePayPaymentGateway.createPaymentUrl()
  -> backend trả sepay://pay?... chứa qrUrl, amount, transferContent
  -> frontend hiển thị QR
  -> user chuyển khoản bằng app ngân hàng
  -> SePay gửi webhook về /payments/sepay-webhook
  -> PaymentServiceImpl.handleSePayWebhook()
  -> kiểm tra API key/HMAC + amount + nội dung chuyển khoản
  -> BookingServiceImpl.handlePaymentSuccess()
  -> booking SUCCESS, ghế BOOKED, sinh ticket"""
    )

    doc.add_heading("E. Bảng thuật ngữ thanh toán", level=1)
    add_table(
        doc,
        ["Thuật ngữ", "Hiểu đơn giản", "Trong code/hệ thống"],
        [
            ("Booking", "Đơn đặt vé của user.", "Booking entity, BookingServiceImpl."),
            ("Payment", "Một lần tạo thanh toán cho booking.", "Payment entity, PaymentServiceImpl."),
            ("PaymentEvent", "Nhật ký từng bước của payment.", "PaymentEventService, admin payment events."),
            ("Callback", "Cổng thanh toán redirect/gọi lại sau giao dịch.", "VNPay callback."),
            ("Webhook", "Server bên ngoài tự báo sự kiện về backend.", "SePay webhook."),
            ("Idempotency", "Callback/webhook gửi nhiều lần vẫn không xử lý trùng.", "Kiểm tra payment status trước khi update."),
            ("Amount check", "So tiền callback/webhook với tiền trong DB.", "Chống thanh toán thiếu/sai số tiền."),
            ("Transaction lock", "Khóa payment khi xử lý callback/webhook.", "findLockedByTransactionNo."),
        ],
        [Inches(1.45), Inches(2.35), Inches(2.7)],
    )

    doc.add_heading("F. Đọc code theo thứ tự nào?", level=1)
    add_numbered(
        doc,
        [
            "Mở CheckoutPage.tsx để xem người dùng bấm thanh toán và frontend gọi API.",
            "Mở paymentApi.ts để xem endpoint frontend gọi xuống backend.",
            "Mở PaymentController.java để xem request được nhận ở đâu.",
            "Mở PaymentServiceImpl.java để đọc nghiệp vụ chính: tạo payment, callback VNPay, webhook SePay.",
            "Mở VnPayPaymentGateway.java hoặc SePayPaymentGateway.java để xem từng cổng tạo URL/QR như thế nào.",
            "Mở BookingServiceImpl.java để xem lúc payment SUCCESS thì ghế và ticket được cập nhật ra sao.",
            "Mở PaymentEventServiceImpl.java để hiểu admin đối soát payment bằng nhật ký nào.",
        ],
    )

    doc.add_heading("G. Câu hỏi bảo vệ và cách trả lời ngắn", level=1)
    add_table(
        doc,
        ["Câu hỏi", "Trả lời trọng tâm"],
        [
            (
                "Vì sao backend không tin amount từ frontend?",
                "Vì request frontend có thể bị sửa hoặc stale khi đổi mã giảm giá. Backend lấy booking.totalPrice trong DB làm chuẩn.",
            ),
            (
                "Vì sao QR phải tạo lại khi đổi mã giảm giá?",
                "QR đã cố định số tiền và nội dung chuyển khoản. Đổi mã làm tổng tiền đổi, nên QR cũ phải bỏ để tránh chuyển sai tiền.",
            ),
            (
                "Nếu VNPay callback hai lần thì sao?",
                "Backend kiểm tra payment status. Nếu đã SUCCESS/FAILED/EXPIRED thì không xử lý lại, chỉ trả kết quả hiện tại.",
            ),
            (
                "Webhook SePay khác nút Tôi đã chuyển khoản như thế nào?",
                "Webhook là xác nhận thật từ server SePay. Nút Tôi đã chuyển khoản chỉ refetch booking để kiểm tra webhook đã đến chưa.",
            ),
            (
                "Khi payment thành công thì ai sinh vé?",
                "PaymentService xác nhận payment, sau đó gọi BookingServiceImpl.handlePaymentSuccess để đổi ghế BOOKED và sinh ticket.",
            ),
            (
                "Nếu callback đến sau khi hết hạn giữ vé?",
                "Hệ thống không xác nhận vé nữa. Payment/booking chuyển EXPIRED để tránh giữ ghế quá lâu hoặc bán sai.",
            ),
        ],
        [Inches(2.45), Inches(4.05)],
    )

    doc.add_heading("H. Checklist tự học nhanh", level=1)
    add_bullets(
        doc,
        [
            "Bạn giải thích được khác nhau giữa VNPay redirect và SePay webhook.",
            "Bạn biết payment SUCCESS chưa đủ; BookingService còn phải đổi booking, ghế và sinh ticket.",
            "Bạn biết vì sao cần kiểm tra amount ở backend.",
            "Bạn biết vì sao callback/webhook cần idempotent.",
            "Bạn biết đọc log/payment event để debug giao dịch lỗi.",
            "Bạn biết frontend không tự xác nhận vé, chỉ hiển thị trạng thái backend trả về.",
        ],
    )

    return save_with_fallback(doc, PAYMENT_DOC)


def enhance_websocket_doc():
    doc = Document(WEBSOCKET_DOC)
    if has_learning_section(doc):
        print(f"Skipped existing learning section: {WEBSOCKET_DOC}")
        return WEBSOCKET_DOC
    setup_learning_styles(doc)
    add_section_marker(doc)

    doc.add_heading("A. Cách đọc luồng WebSocket trong 15 phút", level=1)
    add_numbered(
        doc,
        [
            "Đầu tiên nhớ công thức: HTTP lấy snapshot, WebSocket nhận delta, DB là nguồn sự thật.",
            "Đọc WebSocketConfig để hiểu backend mở endpoint và topic như thế nào.",
            "Đọc SeatStatusPublisher để hiểu backend phát event xuống frontend.",
            "Đọc useSeatWebSocket và SeatSelectionPage để hiểu frontend subscribe và đổi màu ghế.",
            "Đọc scheduler để hiểu vì sao user tắt tab thì ghế vẫn tự nhả khi hết hạn.",
        ],
    )

    add_callout(
        doc,
        "Ý chính cần nhớ",
        "WebSocket không phải nơi quyết định ghế có được giữ hay không. Backend vẫn dùng database transaction và lock để quyết định. WebSocket chỉ là kênh thông báo nhanh cho các màn hình đang mở.",
        "ECFDF5",
        "047857",
    )

    doc.add_heading("B. Sơ đồ realtime tổng quát", level=1)
    add_code(
        doc,
        """Client mở trang chọn ghế
  -> HTTP GET seat map để lấy toàn bộ ghế ban đầu
  -> WebSocket subscribe /topic/seatmap/{showtimeId}

User A giữ ghế
  -> REST API holdSeats()
  -> DB lock seat_status
  -> AVAILABLE -> HOLD
  -> transaction commit
  -> SeatStatusPublisher gửi event
  -> User B đang mở cùng suất chiếu nhận event
  -> UI User B đổi màu ghế"""
    )

    doc.add_heading("C. Snapshot và delta là gì?", level=1)
    add_table(
        doc,
        ["Khái niệm", "Hiểu đơn giản", "Trong hệ thống"],
        [
            ("Snapshot", "Ảnh chụp toàn bộ trạng thái ghế lúc mở trang.", "bookingApi.getSeatMap(showtimeId)."),
            ("Delta", "Một thay đổi nhỏ vừa xảy ra.", "SeatStatusEvent cho 1 seatId."),
            ("Fallback refetch", "Cơ chế tự đồng bộ lại nếu lỡ mất WebSocket event.", "React Query refetchInterval/focus refetch."),
            ("Topic", "Kênh realtime riêng cho một suất chiếu.", "/topic/seatmap/{showtimeId}."),
            ("Publisher", "Class phát tin realtime.", "SeatStatusPublisher."),
            ("Subscriber", "Frontend đang nghe topic.", "useSeatWebSocket."),
        ],
        [Inches(1.45), Inches(2.4), Inches(2.65)],
    )

    doc.add_heading("D. Sơ đồ các trạng thái ghế", level=1)
    add_code(
        doc,
        """AVAILABLE
  -> user giữ ghế thành công
  -> HOLD
      -> thanh toán thành công
      -> BOOKED

HOLD
  -> user hủy / thanh toán thất bại / hết hạn giữ vé
  -> AVAILABLE"""
    )
    doc.add_paragraph(
        "SELECTED và MY_HOLD là trạng thái giao diện để người dùng dễ nhìn. Trạng thái DB chính vẫn là AVAILABLE, HOLD và BOOKED."
    )

    doc.add_heading("E. Đọc code theo thứ tự nào?", level=1)
    add_numbered(
        doc,
        [
            "Mở WebSocketConfig.java để xem endpoint /ws-native và broker /topic.",
            "Mở SecurityConfig.java để xem vì sao WebSocket endpoint public nhưng API giữ ghế vẫn cần JWT.",
            "Mở SeatStatusEvent.java để xem payload event gồm những field nào.",
            "Mở SeatStatusPublisher.java để xem backend publish event ra topic.",
            "Mở BookingServiceImpl.holdSeats để xem DB lock và đổi AVAILABLE -> HOLD.",
            "Mở BookingServiceImpl.handlePaymentSuccess để xem HOLD -> BOOKED.",
            "Mở HoldExpireScheduler/PendingBookingExpireScheduler để xem HOLD -> AVAILABLE khi hết hạn.",
            "Mở useSeatWebSocket.ts để xem frontend subscribe topic.",
            "Mở SeatSelectionPage.tsx để xem event làm đổi màu ghế như thế nào.",
        ],
    )

    doc.add_heading("F. Vì sao publish sau commit?", level=1)
    add_code(
        doc,
        """Sai cách:
  update DB
  publish WebSocket ngay
  transaction rollback
  -> frontend đã đổi màu nhưng DB không đổi

Đúng cách:
  update DB
  transaction commit thành công
  publish WebSocket
  -> frontend nhận event đúng với DB"""
    )
    doc.add_paragraph(
        "Đây là điểm rất quan trọng khi trình bày. Realtime chỉ đáng tin khi event phản ánh trạng thái đã commit trong database."
    )

    doc.add_heading("G. Bảng câu hỏi bảo vệ và cách trả lời ngắn", level=1)
    add_table(
        doc,
        ["Câu hỏi", "Trả lời trọng tâm"],
        [
            (
                "WebSocket có chống bán trùng ghế không?",
                "Không. Chống bán trùng ghế nằm ở backend transaction và DB lock. WebSocket chỉ thông báo trạng thái mới.",
            ),
            (
                "Vì sao cần HTTP snapshot nếu đã có WebSocket?",
                "WebSocket chỉ gửi thay đổi mới. Khi mở trang cần HTTP để lấy toàn bộ trạng thái ban đầu.",
            ),
            (
                "Nếu mất WebSocket thì UI có lệch không?",
                "Có thể lệch tạm thời, nên frontend có React Query refetch để tự đồng bộ lại.",
            ),
            (
                "Vì sao topic có showtimeId?",
                "Để client chỉ nhận event của suất chiếu đang xem, giảm nhiễu và giảm dữ liệu gửi xuống.",
            ),
            (
                "Vì sao WebSocket public vẫn an toàn?",
                "Client chỉ nhận trạng thái ghế. Muốn giữ ghế/thanh toán vẫn phải gọi API có JWT và permission.",
            ),
            (
                "Khi thanh toán thành công, realtime diễn ra thế nào?",
                "BookingService đổi ghế sang BOOKED, sau commit SeatStatusPublisher gửi BOOKED event xuống các client.",
            ),
        ],
        [Inches(2.55), Inches(3.95)],
    )

    doc.add_heading("H. Checklist tự học nhanh", level=1)
    add_bullets(
        doc,
        [
            "Bạn giải thích được snapshot khác delta.",
            "Bạn biết DB lock mới là phần chống trùng ghế.",
            "Bạn biết WebSocket topic theo showtimeId.",
            "Bạn biết backend publish sau commit để tránh UI lệch DB.",
            "Bạn biết frontend cleanup subscription khi rời trang hoặc đổi suất chiếu.",
            "Bạn biết scheduler nhả ghế hết hạn dù user đã tắt tab.",
            "Bạn biết test realtime bằng cách mở hai trình duyệt cùng một suất chiếu.",
        ],
    )

    return save_with_fallback(doc, WEBSOCKET_DOC, WEBSOCKET_FALLBACK_DOC)


if __name__ == "__main__":
    print(enhance_payment_doc())
    print(enhance_websocket_doc())

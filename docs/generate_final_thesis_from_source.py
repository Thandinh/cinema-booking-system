from __future__ import annotations

"""Generate thesis deliverables from the live CinemaBooking source tree.

This script deliberately does not open or depend on any existing thesis report.
It reads only backend/frontend implementation, migrations, configuration and tests.
"""

import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Inches, Mm, Pt, RGBColor

from thesis_academic_expansion import (
    add_formal_abbreviations,
    add_formal_pre_toc_pages,
    add_formal_summary_and_preface,
    expand_chapter_one,
    expand_chapter_three,
    expand_chapter_two,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FRONTEND = ROOT.parent / "frontend" / "cinema-client"
BACKEND_JAVA = ROOT / "src" / "main" / "java" / "com" / "cinema" / "booking"
MIGRATIONS = ROOT / "src" / "main" / "resources" / "db" / "migration"
TESTS = ROOT / "src" / "test" / "java"
REPORT = DOCS / "Bao_cao_khoa_luan_CinemaBooking_Final_Updated_Showtime_Checkin.docx"
PLANTUML = DOCS / "CinemaBooking_PlantUML_Final.txt"
SOURCE_ANALYSIS = DOCS / "CinemaBooking_Source_Analysis.md"
API_INVENTORY = DOCS / "CinemaBooking_API_Inventory.md"

NAVY = "07172E"
BLUE = "1D4ED8"
SLATE = "475569"
LIGHT_BLUE = "EAF2FF"
LIGHT_GRAY = "F8FAFC"
LIGHT_GOLD = "FFF8E1"
BORDER = "CBD5E1"


@dataclass
class Endpoint:
    module: str
    http_method: str
    endpoint: str
    controller: str
    method: str
    request: str
    response: str
    authorization: str


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def safe_relative(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT)).replace("\\", "/")
    except ValueError:
        return str(path).replace("\\", "/")


def collapse_ws(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def join_path(base: str, tail: str) -> str:
    base = base.rstrip("/")
    tail = tail.strip()
    if not tail:
        return base or "/"
    if not tail.startswith("/"):
        tail = "/" + tail
    return (base + tail) if base else tail


def extract_quoted(annotation_line: str) -> str:
    match = re.search(r'"([^"]*)"', annotation_line)
    return match.group(1) if match else ""


def parse_controllers() -> list[Endpoint]:
    endpoints: list[Endpoint] = []
    controller_dir = BACKEND_JAVA / "controller"
    mapping_names = {
        "Get": "GET",
        "Post": "POST",
        "Put": "PUT",
        "Patch": "PATCH",
        "Delete": "DELETE",
    }

    for path in sorted(controller_dir.glob("*.java")):
        text = read(path)
        class_match = re.search(r"public\s+class\s+(\w+)", text)
        if not class_match:
            continue
        controller = class_match.group(1)
        before_class = text[: class_match.start()]
        roots = re.findall(r"@RequestMapping\s*\(\s*\"([^\"]*)\"", before_class)
        root = roots[-1] if roots else ""
        module = controller.removesuffix("Controller")
        lines = text.splitlines()
        pending: list[str] = []
        index = 0

        while index < len(lines):
            line = lines[index].strip()
            if line.startswith("@"):
                pending.append(line)
                index += 1
                continue

            if pending and "(" in line:
                signature = line
                paren_balance = line.count("(") - line.count(")")
                while paren_balance > 0 and index + 1 < len(lines):
                    index += 1
                    addition = lines[index].strip()
                    signature += " " + addition
                    paren_balance += addition.count("(") - addition.count(")")

                mapping = next((entry for entry in pending if re.match(r"@(Get|Post|Put|Patch|Delete)Mapping", entry)), None)
                if mapping:
                    kind = re.match(r"@(Get|Post|Put|Patch|Delete)Mapping", mapping).group(1)
                    http_method = mapping_names[kind]
                    endpoint = join_path(root, extract_quoted(mapping))
                    pre_auth = next((entry for entry in pending if entry.startswith("@PreAuthorize")), None)
                    if pre_auth:
                        authorization = extract_quoted(pre_auth) or pre_auth
                    else:
                        authorization = infer_filter_authorization(http_method, endpoint)

                    method_match = re.search(r"(?:public\s+)?(.+?)\s+(\w+)\s*\((.*)\)", collapse_ws(signature))
                    if method_match:
                        response = method_match.group(1)
                        method_name = method_match.group(2)
                        params = method_match.group(3)
                    else:
                        response = "[CHƯA XÁC MINH ĐƯỢC TỪ SOURCE CODE]"
                        method_name = "[không phân tích được chữ ký]"
                        params = ""
                    request = describe_parameters(params)
                    endpoints.append(Endpoint(module, http_method, endpoint, controller, method_name, request, response, authorization))
                pending = []
            elif line and not line.startswith("//") and not line.startswith("*"):
                pending = []
            index += 1
    return endpoints


def infer_filter_authorization(http_method: str, endpoint: str) -> str:
    public_post = {
        "/auth/token", "/auth/google", "/auth/introspect", "/auth/logout", "/auth/refresh",
        "/api/v1/users/register", "/api/v1/users/verify-email", "/api/v1/users/resend-verification",
        "/api/v1/users/forgot-password", "/api/v1/users/reset-password",
        "/api/v1/payments/momo-ipn", "/api/v1/payments/sepay-webhook",
    }
    public_get_prefixes = (
        "/api/v1/movies", "/api/v1/showtimes", "/api/v1/bookings/showtimes",
        "/api/v1/cinemas", "/api/v1/payments/vnpay-callback", "/api/v1/payments/momo-return",
    )
    if http_method == "POST" and endpoint in public_post:
        return "PUBLIC (SecurityConfig)"
    if http_method == "GET" and any(endpoint.startswith(prefix) for prefix in public_get_prefixes):
        return "PUBLIC (SecurityConfig)"
    return "Authenticated by SecurityConfig; scope may be enforced in service"


def describe_parameters(params: str) -> str:
    parts: list[str] = []
    if "@RequestBody" in params:
        body_match = re.search(r"@RequestBody(?:\([^)]*\))?\s+(?:@\w+(?:\([^)]*\))?\s+)*(\w+)", params)
        parts.append("Body: " + (body_match.group(1) if body_match else "RequestBody"))
    for name, label in (("@PathVariable", "Path"), ("@RequestParam", "Query"), ("@ModelAttribute", "Query model")):
        if name in params:
            count = params.count(name)
            parts.append(f"{label}: {count} tham số")
    if "Pageable" in params:
        parts.append("Pageable")
    if "Authentication" in params:
        parts.append("Authentication context")
    if "HttpServletRequest" in params:
        parts.append("HTTP request context")
    return "; ".join(parts) if parts else "Không có body; xem tham số method"


def frontend_calls() -> dict[str, list[str]]:
    result: dict[str, list[str]] = defaultdict(list)
    if not FRONTEND.exists():
        return result
    # TypeScript response generics can be nested, e.g. get<ApiResponse<Page<T>>>(...).
    # Scan up to the invocation's opening parenthesis instead of parsing nested angle brackets.
    pattern = re.compile(r"axiosClient\.(?:get|post|put|patch|delete)\s*[^()]*\(\s*([`'\"])(/[^`'\"]+)\1")
    for path in sorted((FRONTEND / "src" / "api").glob("*.ts")):
        text = read(path)
        for match in pattern.finditer(text):
            endpoint = re.sub(r"\$\{[^}]+\}", "{id}", match.group(2))
            result[endpoint].append(path.name)
    return result


def frontend_usage_for(endpoint: str, calls: dict[str, list[str]]) -> str:
    if endpoint in calls:
        return ", ".join(sorted(set(calls[endpoint])))
    normalized = re.sub(r"/\{[^}]+\}", "/{id}", endpoint)
    if normalized in calls:
        return ", ".join(sorted(set(calls[normalized])))
    return "Không tìm thấy lời gọi trực tiếp từ frontend/api"


def entity_names() -> list[str]:
    return [path.stem for path in sorted((BACKEND_JAVA / "entity").glob("*.java"))]


def test_names() -> list[str]:
    return [safe_relative(path) for path in sorted(TESTS.rglob("*Test.java"))]


def source_counts() -> dict[str, int]:
    return {
        "controllers": len(list((BACKEND_JAVA / "controller").glob("*.java"))),
        "services": len(list((BACKEND_JAVA / "service").glob("*.java"))) + len(list((BACKEND_JAVA / "service" / "impl").glob("*.java"))),
        "repositories": len(list((BACKEND_JAVA / "repository").rglob("*.java"))),
        "entities": len(entity_names()),
        "migrations": len(list(MIGRATIONS.glob("*.sql"))),
        "tests": len(test_names()),
        "endpoints": len(parse_controllers()),
        "frontend_pages": len(list((FRONTEND / "src" / "pages").rglob("*.tsx"))) if FRONTEND.exists() else 0,
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        side_node = margins.find(qn(f"w:{side}"))
        if side_node is None:
            side_node = OxmlElement(f"w:{side}")
            margins.append(side_node)
        side_node.set(qn("w:w"), str(value))
        side_node.set(qn("w:type"), "dxa")


def add_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_table_width(table, width_dxa=9360) -> None:
    tbl_pr = table._tbl.tblPr
    node = tbl_pr.find(qn("w:tblW"))
    if node is None:
        node = OxmlElement("w:tblW")
        tbl_pr.append(node)
    node.set(qn("w:w"), str(width_dxa))
    node.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    configure_section_layout(doc.sections[0])
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Inches(0.3)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size in (("Title", 18), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Inches(0)
    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styles["Heading 3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(11)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Inches(0)
    for style_name in ("Figure Caption", "Table Caption"):
        try:
            custom = styles[style_name]
        except KeyError:
            custom = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        custom.base_style = caption
        custom.font.name = "Times New Roman"
        custom._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        custom.font.size = Pt(11)
        custom.font.italic = True
        custom.font.color.rgb = RGBColor(0, 0, 0)
        custom.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        custom.paragraph_format.first_line_indent = Inches(0)


def configure_section_layout(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.0)


def set_page_number_format(section, fmt: str, start: int = 1) -> None:
    sect_pr = section._sectPr
    page_number_type = sect_pr.find(qn("w:pgNumType"))
    if page_number_type is None:
        page_number_type = OxmlElement("w:pgNumType")
        sect_pr.append(page_number_type)
    page_number_type.set(qn("w:fmt"), fmt)
    page_number_type.set(qn("w:start"), str(start))


def start_numbered_section(doc: Document, fmt: str, start: int = 1):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section_layout(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for paragraph_node in section.footer.paragraphs:
        paragraph_node.clear()
    set_page_number_format(section, fmt, start)
    add_page_number(section)
    return section


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def enable_field_updates(doc: Document) -> None:
    """Ask Microsoft Word to recalculate TOC and list fields when opening."""
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_toc(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Mở file trong Microsoft Word, bấm Update Field để cập nhật mục lục tự động."
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def add_style_toc(doc: Document, style_name: str, placeholder_text: str) -> None:
    paragraph_node = doc.add_paragraph()
    paragraph_node.paragraph_format.first_line_indent = Inches(0)
    run = paragraph_node.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f'TOC \\h \\z \\t "{style_name},1"'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = placeholder_text
    separate.append(text_node)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def add_inline_text(paragraph_node, text: str, *, bold: bool = False, italic: bool = False, color: str | None = None) -> None:
    """Render Markdown-style inline code as Word character formatting, not literal backticks."""
    for index, part in enumerate(text.split("`")):
        if not part:
            continue
        run = paragraph_node.add_run(part)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if index % 2 == 1:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10.5)


def paragraph(doc: Document, text: str = "", bold: bool = False, italic: bool = False, color: str | None = None) -> None:
    p = doc.add_paragraph()
    add_inline_text(p, text, bold=bold, italic=italic, color=color)


def add_hyperlink(paragraph_node, text: str, url: str) -> None:
    relationship_id = paragraph_node.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph_node._p.append(hyperlink)


def hyperlink_paragraph(doc: Document, label: str, url: str, suffix: str = "") -> None:
    p = doc.add_paragraph()
    add_inline_text(p, f"{label}: ", bold=True)
    add_hyperlink(p, url, url)
    if suffix:
        add_inline_text(p, suffix)


def bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_inline_text(p, item)


def table(doc: Document, headers: list[str], rows: list[list[str]], font_size=9.5, caption: str | None = None) -> None:
    if caption:
        caption_paragraph = doc.add_paragraph(caption, style="Table Caption")
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t)
    add_table_borders(t)
    set_repeat_table_header(t.rows[0])
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            run = cell.paragraphs[0].add_run(value.replace("`", ""))
            run.font.size = Pt(font_size)
    doc.add_paragraph()


def callout(doc: Document, title: str, body: str, fill=LIGHT_BLUE) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t)
    add_table_borders(t)
    cell = t.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 130, 150, 130, 150)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run("\n" + body)
    doc.add_paragraph()


def placeholder(doc: Document, marker: str, explanation: str, caption: str | None = None) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(t)
    add_table_borders(t)
    cell = t.cell(0, 0)
    set_cell_shading(cell, LIGHT_GOLD)
    set_cell_margins(cell, 170, 160, 170, 160)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(marker)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("9A5B00")
    p.add_run("\n" + explanation)
    if caption:
        caption_paragraph = doc.add_paragraph(caption, style="Figure Caption")
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def code_block(doc: Document, code: str, caption: str, font_size: float = 8.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(caption)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(SLATE)
    t = doc.add_table(rows=1, cols=1)
    set_table_width(t)
    add_table_borders(t)
    row_properties = t.rows[0]._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    row_properties.append(cannot_split)
    cell = t.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, 120, 130, 120, 130)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    for number, line in enumerate(code.strip().splitlines()):
        if number:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(font_size)
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    def cover_line(text: str, size: float, *, bold: bool = False, after: float = 0) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0, 0, 0)

    def cover_information() -> None:
        info = doc.add_table(rows=4, cols=2)
        info.alignment = WD_TABLE_ALIGNMENT.CENTER
        info.autofit = False
        labels = ["Sinh viên thực hiện:", "Mã số sinh viên:", "Lớp:", "Giảng viên hướng dẫn:"]
        values = ["[HỌ VÀ TÊN SINH VIÊN]", "[MÃ SỐ SINH VIÊN]", "[LỚP]", "[HỌ VÀ TÊN GIẢNG VIÊN]"]
        for index, label in enumerate(labels):
            info.cell(index, 0).text = label
            info.cell(index, 1).text = values[index]
            for cell in info.rows[index].cells:
                cell.paragraphs[0].paragraph_format.first_line_indent = Inches(0)
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)
                for run in cell.paragraphs[0].runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(13)
            info.cell(index, 0).paragraphs[0].runs[0].bold = True

    # Bìa chính: trang trọng, tối giản và để trống thông tin đơn vị chưa được cung cấp.
    cover_line("[TÊN CƠ QUAN CHỦ QUẢN]", 13, bold=True)
    cover_line("[TÊN TRƯỜNG ĐẠI HỌC]", 14, bold=True)
    cover_line("[TÊN KHOA/BỘ MÔN]", 14, bold=True, after=30)
    cover_line("KHÓA LUẬN TỐT NGHIỆP", 20, bold=True)
    cover_line("NGÀNH CÔNG NGHỆ THÔNG TIN", 14, bold=True, after=34)
    cover_line("ĐỀ TÀI", 13, bold=True, after=8)
    cover_line("XÂY DỰNG HỆ THỐNG ĐẶT VÉ XEM PHIM", 20, bold=True)
    cover_line("TRỰC TUYẾN CINEMABOOKING", 20, bold=True, after=38)
    cover_information()
    for _ in range(3):
        doc.add_paragraph()
    cover_line("[ĐỊA DANH], NĂM 2026", 13, bold=True)
    doc.add_page_break()

    # Bìa phụ: dùng cho trang tên đề tài bên trong cuốn báo cáo.
    cover_line("[TÊN CƠ QUAN CHỦ QUẢN]", 13, bold=True)
    cover_line("[TÊN TRƯỜNG ĐẠI HỌC]", 14, bold=True)
    cover_line("[TÊN KHOA/BỘ MÔN]", 14, bold=True, after=26)
    cover_line("[HỌ VÀ TÊN SINH VIÊN]", 14, bold=True, after=22)
    cover_line("XÂY DỰNG HỆ THỐNG ĐẶT VÉ XEM PHIM", 20, bold=True)
    cover_line("TRỰC TUYẾN CINEMABOOKING", 20, bold=True, after=20)
    cover_line("KHÓA LUẬN TỐT NGHIỆP", 16, bold=True)
    cover_line("NGÀNH CÔNG NGHỆ THÔNG TIN", 13, bold=True, after=32)
    cover_information()
    for _ in range(4):
        doc.add_paragraph()
    cover_line("[ĐỊA DANH], NĂM 2026", 13, bold=True)


def add_toc_page(doc: Document) -> None:
    doc.add_heading("MỤC LỤC", level=1)
    add_toc(doc)
    doc.add_page_break()


def add_chapter_one(doc: Document, counts: dict[str, int]) -> None:
    doc.add_heading("CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", level=1)
    doc.add_heading("1.1. Lý do chọn đề tài", level=2)
    paragraph(doc, "Trong quy trình bán vé tại rạp, số ghế của mỗi suất chiếu là hữu hạn nhưng được nhiều khách hàng cùng quan tâm. Khi quy trình này được chuyển lên môi trường web, khoảng thời gian từ lúc chọn ghế đến lúc nhận kết quả thanh toán làm phát sinh trạng thái giữ chỗ tạm thời. Nếu hai yêu cầu đến gần như đồng thời hoặc kết quả thanh toán bị gửi lặp, việc chỉ kiểm tra trạng thái trên giao diện không đủ để ngăn một ghế được bán nhiều lần.")
    paragraph(doc, "Bài toán còn liên quan đến hoạt động sau bán vé. Vé điện tử phải được kiểm tra đúng rạp và đúng suất chiếu; nhân viên chỉ được thao tác trên những rạp được phân công; đơn của suất chiếu bị hủy cần được đưa vào quy trình quản lý yêu cầu hoàn tiền có thể theo dõi. Những yêu cầu này tạo ra mối liên hệ chặt chẽ giữa thiết kế cơ sở dữ liệu, giao dịch, phân quyền và trải nghiệm người dùng. Vì vậy, đề tài CinemaBooking được chọn để xây dựng và đánh giá một hệ thống đặt vé hoàn chỉnh trên nền ReactJS, Spring Boot và PostgreSQL.")
    doc.add_heading("1.2. Mục tiêu đề tài", level=2)
    doc.add_heading("1.2.1. Mục tiêu tổng quát", level=3)
    paragraph(doc, "Xây dựng hệ thống web cho phép khách hàng xem phim, xem rạp/lịch chiếu, chọn ghế, giữ ghế có thời hạn, tạo đơn, thanh toán và nhận vé QR; đồng thời cung cấp khu vực vận hành cho nhân viên rạp và quản trị viên.")
    doc.add_heading("1.2.2. Mục tiêu cụ thể", level=3)
    bullets(doc, [
        "Xây dựng REST API phân lớp Controller → Service → Repository → PostgreSQL.",
        "Xác thực bằng JWT access token và refresh token có lưu hash, rotation, thu hồi phiên và audit.",
        "Phân quyền RBAC theo permission; với STAFF còn kiểm soát phạm vi rạp được phân công.",
        "Xử lý giữ ghế, hết hạn giữ ghế, đặt vé và thanh toán với transaction, khóa dữ liệu và trạng thái rõ ràng.",
        "Phát đồng bộ trạng thái ghế theo thời gian thực bằng WebSocket/STOMP.",
        "Tích hợp các luồng thanh toán VNPay và SePay/VietQR; hỗ trợ tạo, theo dõi và ghi nhận kết quả yêu cầu hoàn tiền theo quy trình vận hành.",
    ])
    doc.add_heading("1.3. Đối tượng nghiên cứu", level=2)
    paragraph(doc, "Đối tượng nghiên cứu của đề tài là các quy trình nghiệp vụ và giải pháp kỹ thuật trong hệ thống đặt vé xem phim trực tuyến, bao gồm quản lý phim, rạp, phòng chiếu, ghế và suất chiếu; xác thực và phân quyền người dùng; giữ ghế trong môi trường có nhiều yêu cầu đồng thời; tạo đơn đặt vé, áp dụng khuyến mãi, thanh toán, phát hành vé điện tử và soát vé bằng mã QR. Bên cạnh đó, đề tài nghiên cứu cách tổ chức kiến trúc client-server, thiết kế cơ sở dữ liệu quan hệ, xây dựng REST API và đồng bộ trạng thái ghế theo thời gian thực.")
    doc.add_heading("1.4. Phạm vi đề tài", level=2)
    table(doc, ["Nhóm", "Phạm vi chức năng"], [
        ["Khách hàng", "Tra cứu phim, rạp và lịch chiếu; lựa chọn khu vực, giữ ghế, tạo đơn đặt vé, áp dụng khuyến mãi và thanh toán; nhận, xem và quản lý vé điện tử; cập nhật hồ sơ và quản lý các phiên đăng nhập cá nhân."],
        ["Nhân viên", "Thực hiện nghiệp vụ vận hành tại các rạp được phân công: xem dữ liệu phim và rạp; xem, cập nhật phòng chiếu và ghế; tạo, cập nhật hoặc hủy suất chiếu theo chính sách; theo dõi đơn đặt vé, thanh toán và vé; xem số liệu vận hành; soát vé QR đúng rạp và suất chiếu."],
        ["Quản trị viên", "Có toàn bộ quyền của hệ thống, bao gồm các nghiệp vụ của nhân viên trên phạm vi tất cả các rạp. Ngoài ra, quản trị viên thực hiện quản lý phim, rạp, phòng chiếu, ghế, suất chiếu, người dùng, phân quyền, khuyến mãi, đơn đặt vé, thanh toán, yêu cầu hoàn tiền, thống kê và nhật ký vận hành."],
    ], caption="Bảng 1.1. Phạm vi chức năng của đề tài")
    doc.add_heading("1.5. Phương pháp và quy trình thực hiện", level=2)
    doc.add_heading("1.5.1. Phương pháp thực hiện", level=3)
    paragraph(doc, "Đề tài được thực hiện thông qua các bước khảo sát yêu cầu, mô hình hóa nghiệp vụ, thiết kế kiến trúc, cơ sở dữ liệu và giao diện, hiện thực các phân hệ, sau đó kiểm thử bằng các mã kiểm thử và tình huống sử dụng thực tế. Nội dung báo cáo được đối chiếu với các thành phần hiện thực của hệ thống, bao gồm giao diện React, các lớp Controller, Service, Repository và các migration cơ sở dữ liệu. Cách tiếp cận này giúp bảo đảm các nhận định về giữ ghế, giao dịch, thanh toán và phân quyền có căn cứ rõ ràng.")
    doc.add_heading("1.5.2. Quy trình thực hiện đề tài", level=3)
    paragraph(doc, "Quá trình thực hiện đề tài được tổ chức thành sáu giai đoạn: khảo sát yêu cầu, phân tích nghiệp vụ, thiết kế hệ thống, hiện thực chương trình, kiểm thử và đánh giá kết quả. Mặc dù được trình bày theo trình tự, các giai đoạn không hoàn toàn tách biệt mà được rà soát và điều chỉnh trong quá trình phát triển. Kết quả kiểm thử có thể dẫn đến việc cập nhật quy tắc nghiệp vụ, thiết kế dữ liệu hoặc cách tổ chức giao diện. Cách thực hiện này giúp kiểm soát phạm vi và duy trì sự nhất quán giữa yêu cầu, thiết kế và sản phẩm được xây dựng.")
    table(doc, ["Giai đoạn", "Nội dung thực hiện", "Kết quả"], [
        ["Khảo sát yêu cầu", "Xác định nhóm người sử dụng, nhu cầu đặt vé và quy trình vận hành tại rạp.", "Yêu cầu và phạm vi của đề tài."],
        ["Phân tích nghiệp vụ", "Phân tích quy trình xác thực, chọn ghế, đặt vé, thanh toán, phát hành và soát vé.", "Mô hình trường hợp sử dụng và các quy tắc nghiệp vụ."],
        ["Thiết kế hệ thống", "Thiết kế kiến trúc, cơ sở dữ liệu, REST API, giao diện và các luồng xử lý.", "Sơ đồ thiết kế, mô hình dữ liệu và đặc tả giao tiếp."],
        ["Hiện thực", "Xây dựng front-end, back-end, cơ sở dữ liệu và các thành phần tích hợp.", "Phiên bản CinemaBooking đáp ứng phạm vi chức năng của đề tài."],
        ["Kiểm thử", "Kiểm tra chức năng, phân quyền và thanh toán; xây dựng kịch bản đánh giá xử lý đồng thời.", "Kết quả kiểm thử và các vấn đề cần điều chỉnh."],
        ["Đánh giá", "Đối chiếu kết quả với mục tiêu và xác định những hạn chế.", "Kết luận và hướng phát triển của đề tài."],
    ], caption="Bảng 1.2. Các giai đoạn thực hiện đề tài")
    doc.add_heading("1.6. Cấu trúc báo cáo", level=2)
    paragraph(doc, "Ngoài phần tài liệu tham khảo và phụ lục, báo cáo được tổ chức thành năm chương. Chương 1 trình bày tổng quan về đề tài; Chương 2 trình bày cơ sở lý thuyết và các công nghệ được sử dụng; Chương 3 tập trung phân tích yêu cầu và thiết kế hệ thống; Chương 4 mô tả quá trình hiện thực và kiểm thử; Chương 5 tổng kết các kết quả đạt được, phân tích những hạn chế và đề xuất hướng phát triển.")


def add_chapter_two(doc: Document) -> None:
    doc.add_heading("CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ", level=1)
    doc.add_heading("2.1. Kiến trúc ứng dụng web", level=2)
    paragraph(doc, "Kiến trúc client-server phân tách ứng dụng thành phía client đảm nhiệm giao diện và phía server quản lý dữ liệu, bảo mật cùng các quy tắc nghiệp vụ. Hai phía giao tiếp qua một hợp đồng xác định trước, nhờ đó giao diện có thể thay đổi mà không phải đưa logic nghiệp vụ quan trọng xuống trình duyệt [1]. Đối với hệ thống đặt vé, cách phân tách này còn giúp máy chủ giữ vai trò quyết định đối với giá vé, quyền sở hữu lượt giữ ghế, kết quả thanh toán và hiệu lực của vé điện tử.")
    paragraph(doc, "CinemaBooking áp dụng mô hình trên với front-end React chạy trên trình duyệt, back-end Spring Boot cung cấp RESTful API và PostgreSQL lưu trữ dữ liệu bền vững. Dữ liệu thông thường được trao đổi dưới dạng JSON qua HTTP. Riêng trạng thái ghế cần phản ánh nhanh đến nhiều người dùng được truyền bổ sung qua WebSocket. RESTful API vẫn là kênh thực hiện lệnh và lấy dữ liệu ban đầu; WebSocket chỉ phân phối những thay đổi phát sinh sau đó.")

    doc.add_heading("2.2. ReactJS và TypeScript", level=2)
    paragraph(doc, "ReactJS xây dựng giao diện từ các component có thể tái sử dụng. Mỗi component nhận dữ liệu đầu vào qua props và quản lý trạng thái cần thiết bằng Hook. Khi trạng thái thay đổi, React cập nhật lại phần giao diện liên quan thay vì yêu cầu tải lại toàn bộ trang [2]. TypeScript bổ sung hệ thống kiểu tĩnh cho JavaScript, hỗ trợ phát hiện sớm sai khác giữa dữ liệu API, thuộc tính component và trạng thái của biểu mẫu [3].")
    paragraph(doc, "Front-end của CinemaBooking sử dụng React 19, TypeScript, Vite và React Router. Trạng thái được phân chia theo phạm vi sử dụng thay vì tập trung toàn bộ vào một kho dùng chung. Trạng thái chỉ phục vụ một màn hình được giữ tại component; dữ liệu biểu mẫu được quản lý cùng cơ chế validation; thông tin xác thực, giao diện và khu vực ưu tiên được chia sẻ bằng Zustand; dữ liệu nhận từ máy chủ có vòng đời tải, lưu tạm và làm mới được quản lý bằng React Query. Sự phân chia này làm rõ nguồn gốc dữ liệu và hạn chế việc render hoặc gọi API không cần thiết.")
    paragraph(doc, "React Router tổ chức các màn hình theo địa chỉ truy cập và hỗ trợ bảo vệ những khu vực dành cho người dùng đã đăng nhập, nhân viên hoặc quản trị viên. Các mô-đun API theo từng nghiệp vụ sử dụng chung một HTTP client để thống nhất địa chỉ máy chủ, access token, cookie xác thực và cách xử lý lỗi. Cơ chế bảo vệ route ở client nhằm điều hướng trải nghiệm người dùng; quyết định cấp quyền cuối cùng vẫn thuộc về Spring Security ở phía server.")
    table(doc, ["Loại trạng thái", "Công cụ sử dụng", "Ví dụ", "Mục đích"], [
        ["Trạng thái giao diện cục bộ", "React Hook", "Tab, hộp thoại, bộ lọc, trạng thái camera", "Gắn trạng thái với vòng đời của một component."],
        ["Dữ liệu biểu mẫu", "React Hook Form và Zod", "Đăng nhập, đăng ký, hồ sơ và biểu mẫu quản trị", "Quản lý giá trị, validation và lỗi nhập liệu."],
        ["Trạng thái dùng chung", "Zustand", "Phiên người dùng, giao diện và khu vực ưu tiên", "Chia sẻ dữ liệu giữa nhiều route và component."],
        ["Dữ liệu máy chủ", "React Query", "Phim, rạp và lịch chiếu theo khu vực", "Quản lý loading, cache, stale state và refetch."],
        ["Dữ liệu thời gian thực", "STOMP Hook", "Thay đổi trạng thái ghế theo suất chiếu", "Tiếp nhận sự kiện do server chủ động gửi."],
    ], caption="Bảng 2.1. Phân loại trạng thái được sử dụng ở phía front-end")

    doc.add_heading("2.3. Java và Spring Boot", level=2)
    paragraph(doc, "Java là ngôn ngữ hướng đối tượng có hệ thống kiểu chặt chẽ và hệ sinh thái phù hợp với các ứng dụng nghiệp vụ. Spring Boot đơn giản hóa quá trình khởi tạo ứng dụng Spring thông qua cơ chế tự động cấu hình, quản lý bean và các starter theo chức năng [4]. Dependency Injection cho phép một thành phần khai báo các phụ thuộc cần thiết thay vì tự khởi tạo chúng, qua đó giảm liên kết cứng và thuận lợi hơn cho kiểm thử.")
    paragraph(doc, "Phía server của CinemaBooking sử dụng Java 21 và Spring Boot 3.5.14. Spring Web đảm nhiệm REST API; Spring Validation kiểm tra dữ liệu đầu vào; Spring Data JPA hỗ trợ truy cập PostgreSQL; Spring Security thực hiện xác thực và phân quyền; các mô-đun cache, mail và WebSocket phục vụ những yêu cầu bổ sung. Ứng dụng được triển khai theo kiến trúc nguyên khối có phân chia mô-đun, phù hợp với phạm vi đề tài và cho phép các transaction đặt vé, thanh toán được kiểm soát trong cùng một cơ sở dữ liệu.")

    doc.add_heading("2.4. PostgreSQL, JPA và xử lý giao dịch", level=2)
    paragraph(doc, "PostgreSQL là hệ quản trị cơ sở dữ liệu quan hệ hỗ trợ khóa ngoại, ràng buộc kiểm tra, chỉ mục, transaction và khóa ở mức bản ghi. Các chỉ mục được lựa chọn theo những đường truy vấn thường xuyên để hỗ trợ truy xuất dữ liệu hiệu quả hơn [5]. Trong CinemaBooking, cơ sở dữ liệu không chỉ lưu thông tin mà còn tham gia bảo vệ các bất biến nghiệp vụ, chẳng hạn một ghế chỉ có một trạng thái trong mỗi suất chiếu, thời gian kết thúc phải sau thời gian bắt đầu và một số trạng thái chỉ được nhận các giá trị hợp lệ.")
    paragraph(doc, "JPA quy định cách ánh xạ đối tượng Java sang mô hình quan hệ [6], còn Hibernate là implementation được sử dụng trong dự án. Spring Data JPA cung cấp repository, truy vấn dẫn xuất, JPQL và native SQL [7]. Một rủi ro thường gặp của ánh xạ quan hệ là N+1 query: sau một truy vấn lấy danh sách, mỗi phần tử tiếp tục tạo thêm truy vấn để tải quan hệ [8]. Hệ thống hạn chế rủi ro này bằng fetch query, projection, tải theo lô và tắt Open Session in View; hiệu quả thực tế cần tiếp tục được kiểm tra qua log SQL và EXPLAIN ANALYZE khi quy mô dữ liệu tăng.")
    paragraph(doc, "Transaction bảo đảm một nhóm thay đổi hoặc cùng được ghi nhận hoặc cùng bị hoàn tác [9]. Đối với ghế và thanh toán, chỉ kiểm tra trạng thái ở tầng ứng dụng là chưa đủ vì nhiều request có thể diễn ra đồng thời. Do đó, hệ thống kết hợp khóa bi quan, trường phiên bản cho khóa lạc quan, cập nhật có điều kiện và ràng buộc duy nhất tại cơ sở dữ liệu. Khóa bản ghi làm cho các transaction cạnh tranh phải chờ nhau tại tài nguyên cần thay đổi, phù hợp với cơ chế khóa tường minh của PostgreSQL [10]. Mỗi cơ chế giải quyết một dạng xung đột khác nhau và tạo thành nhiều lớp bảo vệ thay vì phụ thuộc vào một biện pháp duy nhất.")
    table(doc, ["Cơ chế", "Nguyên lý", "Ưu điểm", "Phạm vi vận dụng"], [
        ["Khóa bi quan", "Khóa bản ghi trong thời gian transaction xử lý", "Ngăn transaction khác đồng thời sửa cùng tài nguyên", "Giữ ghế, xử lý payment và check-in."],
        ["Khóa lạc quan", "So sánh giá trị phiên bản khi cập nhật", "Phát hiện dữ liệu đã bị thay đổi mà không giữ khóa dài", "Trạng thái ghế có trường version."],
        ["Cập nhật có điều kiện", "Chỉ cập nhật khi trạng thái và chủ sở hữu còn phù hợp", "Gộp bước kiểm tra và thay đổi vào một thao tác nguyên tử", "Xác nhận hoặc giải phóng lượt giữ ghế."],
        ["Ràng buộc duy nhất", "Cơ sở dữ liệu từ chối dữ liệu trùng điều kiện", "Tạo lớp bảo vệ cuối cùng trước request chạy song song", "Trạng thái ghế, booking và payment đang chờ."],
    ], caption="Bảng 2.2. Các cơ chế kiểm soát đồng thời được vận dụng trong hệ thống")

    doc.add_heading("2.5. RESTful API", level=2)
    paragraph(doc, "REST là phong cách kiến trúc tổ chức giao tiếp quanh tài nguyên và sử dụng ngữ nghĩa của HTTP [1]. URI xác định tài nguyên, HTTP method biểu diễn thao tác, còn status code phản ánh kết quả xử lý theo ngữ nghĩa của giao thức [11]. Cách thiết kế này tạo giao diện giao tiếp rõ ràng giữa client và server, đồng thời cho phép các màn hình khác nhau tái sử dụng cùng một nghiệp vụ.")
    paragraph(doc, "CinemaBooking trao đổi dữ liệu JSON và dùng DTO để xác định contract của request, response. Bean Validation loại bỏ dữ liệu đầu vào không hợp lệ trước khi vào logic nghiệp vụ; bộ xử lý ngoại lệ tập trung chuyển lỗi thành status code và cấu trúc phản hồi thống nhất. Những endpoint callback hoặc webhook được mở cho payment gateway không đồng nghĩa với việc bỏ qua bảo mật: dữ liệu vẫn phải được xác minh chữ ký, số tiền, trạng thái và tính lũy đẳng trước khi cập nhật giao dịch.")

    doc.add_heading("2.6. MVC và kiến trúc phân lớp", level=2)
    paragraph(doc, "Trong ứng dụng web truyền thống, mô hình MVC phân tách Model, View và Controller [1]. CinemaBooking vận dụng nguyên tắc này trong bối cảnh client-server: React đảm nhiệm View độc lập; Controller phía server tiếp nhận HTTP request; mô hình nghiệp vụ và dữ liệu được tổ chức qua Service, Entity và Repository. Vì vậy, kiến trúc thực tế phù hợp hơn với cách gọi kiến trúc REST phân lớp thay vì MVC server-rendered thuần túy.")
    paragraph(doc, "Sự phân lớp giúp giới hạn trách nhiệm và hướng phụ thuộc. Controller không trực tiếp điều phối transaction phức tạp; Repository không quyết định chính sách quyền; Entity không được trả nguyên trạng ra giao diện. Service trở thành nơi tập trung quy tắc nghiệp vụ và ranh giới transaction, còn DTO và mapper tách contract API khỏi mô hình lưu trữ.")
    table(doc, ["Thành phần", "Trách nhiệm chính", "Giới hạn trách nhiệm"], [
        ["Controller", "Ánh xạ HTTP, tiếp nhận tham số, kích hoạt validation và tạo response", "Không chứa quy trình nghiệp vụ hoặc transaction dài."],
        ["Service", "Xử lý quy tắc nghiệp vụ, quyền sở hữu, chuyển trạng thái và transaction", "Không phụ thuộc chi tiết trình bày của React."],
        ["Repository", "Thực hiện truy vấn, khóa và lưu dữ liệu", "Không điều phối giao diện hoặc payment gateway."],
        ["Entity", "Ánh xạ cấu trúc và quan hệ dữ liệu", "Không được dùng thay cho contract API."],
        ["DTO và Mapper", "Xác định dữ liệu vào, ra và chuyển đổi mô hình", "Không chứa quy tắc thanh toán hoặc phân quyền."],
        ["Gateway/Adapter", "Đóng gói khác biệt của hệ thống thanh toán bên ngoài", "Không tự quyết định trạng thái booking."],
    ], caption="Bảng 2.3. Trách nhiệm của các thành phần trong kiến trúc phân lớp")

    doc.add_heading("2.7. Authentication và authorization", level=2)
    paragraph(doc, "Authentication xác định danh tính của chủ thể, trong khi authorization quyết định chủ thể đó được phép thực hiện hành động nào. CinemaBooking hỗ trợ đăng nhập bằng tên đăng nhập, mật khẩu và Google ID Token; token do Google cung cấp được xác minh tại phía máy chủ trước khi thông tin định danh được chấp nhận [12]. Sau khi xác thực thành công, server phát access token dạng JWT có thời hạn ngắn. JWT chứa các claim đã được ký, cho phép Resource Server kiểm tra tính hợp lệ và quyền truy cập mà không cần lưu một HTTP session cho từng request [13].")
    paragraph(doc, "Trong CinemaBooking, refresh token được quản lý như một phiên đăng nhập ở phía server: giá trị lưu trong cơ sở dữ liệu được băm, token được luân chuyển sau mỗi lần làm mới [14] và có thể bị thu hồi khi đăng xuất. Front-end không lưu refresh token trong localStorage và gửi yêu cầu làm mới với cookie. Các thuộc tính bảo vệ của cookie như HttpOnly, SameSite và Secure cần được đối chiếu với cấu hình triển khai thực tế; vì vậy, báo cáo không xem đây là đặc tính đã được xác minh chỉ từ mã nguồn client. Access token được giữ trong trạng thái ứng dụng và gửi bằng Bearer header. Cách kết hợp hai loại token nhằm cân bằng khả năng xác thực request với yêu cầu quản lý và thu hồi phiên đăng nhập.")
    paragraph(doc, "Authorization được thực hiện theo hai chiều. RBAC ánh xạ người dùng với vai trò và permission để kiểm soát loại hành động; cinema scope giới hạn dữ liệu mà nhân viên được phép thao tác theo rạp được phân công. Kiểm tra route ở front-end chỉ hỗ trợ trải nghiệm, còn Spring Security và kiểm tra tại Service mới là lớp kiểm soát có thẩm quyền [15]. Các endpoint đăng nhập và giữ ghế còn được giới hạn tần suất để giảm thử mật khẩu hoặc spam tài nguyên. Bộ đếm hiện đặt trong bộ nhớ của từng tiến trình, phù hợp với triển khai một instance; giới hạn phân tán dùng chung nhiều instance được xác định là hướng phát triển.")

    doc.add_heading("2.8. WebSocket/STOMP", level=2)
    paragraph(doc, "WebSocket duy trì kết nối hai chiều giữa client và server, phù hợp với dữ liệu cần được đẩy đến trình duyệt ngay khi thay đổi. STOMP bổ sung mô hình destination, publish và subscribe trên kết nối WebSocket, giúp các bên trao đổi thông điệp theo kênh thay vì tự xây dựng một giao thức riêng [16].")
    paragraph(doc, "Trong CinemaBooking, mỗi suất chiếu có một topic trạng thái ghế. Khi mở trang chọn ghế, client lấy ảnh chụp ban đầu bằng REST API rồi đăng ký topic tương ứng để nhận các thay đổi HOLD, AVAILABLE hoặc BOOKED. Sự kiện chỉ được phát sau khi transaction đã commit, nhờ đó giao diện không nhận trạng thái bị hoàn tác. PostgreSQL vẫn là nguồn dữ liệu có thẩm quyền; WebSocket đảm nhiệm đồng bộ thay đổi giữa các client. Simple broker hiện tại phù hợp với một instance và cần được thay thế bằng broker dùng chung nếu hệ thống mở rộng theo chiều ngang.")

    doc.add_heading("2.9. Các công nghệ hỗ trợ", level=2)
    table(doc, ["Thành phần", "Ứng dụng trong hệ thống", "Vai trò"], [
        ["Flyway", "Quản lý các migration của PostgreSQL", "Theo dõi phiên bản lược đồ và kiểm soát thay đổi cấu trúc dữ liệu."],
        ["Caffeine", "Lưu đệm phim, rạp, phòng, ghế và khuyến mãi", "Giảm số lần truy vấn đối với dữ liệu có tần suất đọc cao và ít thay đổi."],
        ["ZXing", "Tạo hình ảnh QR cho từng vé điện tử", "Mã hóa định danh vé phục vụ hiển thị và soát vé."],
        ["Thymeleaf và Spring Mail", "Gửi thư xác thực, đặt lại mật khẩu, vé và thông báo hủy suất", "Cung cấp kênh thông báo qua thư điện tử."],
        ["Google Identity", "Xác minh Google ID Token tại phía máy chủ", "Hỗ trợ đăng nhập bằng tài khoản Google."],
        ["Leaflet", "Hiển thị bản đồ và vị trí rạp", "Hỗ trợ tra cứu rạp theo khu vực và khoảng cách."],
        ["Testcontainers", "Khởi tạo PostgreSQL container trong kiểm thử tích hợp", "Tạo môi trường cơ sở dữ liệu gần với hệ thống thực tế."],
    ], caption="Bảng 2.4. Các công nghệ hỗ trợ được sử dụng trong hệ thống")
    paragraph(doc, "Flyway được dùng cùng chế độ kiểm tra ánh xạ của Hibernate để cấu trúc cơ sở dữ liệu thay đổi theo các migration có thứ tự, thay vì để ORM tự sửa lược đồ khi khởi động. Caffeine chỉ lưu đệm dữ liệu có tần suất đọc cao và ít thay đổi như phim, rạp, phòng hoặc khuyến mãi; trạng thái ghế, booking và payment không được đưa vào cache vì gắn trực tiếp với transaction và thay đổi thường xuyên.")
    paragraph(doc, "Các scheduler xử lý dữ liệu phụ thuộc thời gian như lượt giữ ghế, booking chờ thanh toán, trạng thái suất chiếu và token hết hạn. ZXing, Spring Mail và Thymeleaf phục vụ phát hành QR cùng thông báo qua thư điện tử. Testcontainers hỗ trợ integration test với PostgreSQL gần môi trường thực tế, trong khi các kiểm thử đơn vị tập trung vào thành phần có quy tắc độc lập. Việc lựa chọn công cụ kiểm thử được trình bày chi tiết hơn tại Chương 4.")


def add_chapter_three(doc: Document, endpoints: list[Endpoint]) -> None:
    doc.add_heading("CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", level=1)

    doc.add_heading("3.1. Phân tích bài toán", level=2)
    paragraph(doc, "Quy trình mua vé xem phim trực tuyến bắt đầu từ việc tra cứu phim, rạp và lịch chiếu, tiếp tục qua các bước chọn ghế, giữ chỗ, tạo đơn, thanh toán và nhận vé điện tử. Bên cạnh khách hàng, hệ thống còn phục vụ nhân viên vận hành tại rạp và quản trị viên quản lý dữ liệu trên phạm vi toàn hệ thống. Mỗi nhóm người dùng có mục tiêu, phạm vi dữ liệu và quyền thao tác khác nhau, do đó bài toán không thể được giải quyết bằng một tập chức năng quản lý dữ liệu đơn thuần.")
    paragraph(doc, "Đặc điểm quan trọng nhất của nghiệp vụ đặt vé là ghế vừa được hiển thị cho nhiều người cùng xem, vừa là tài nguyên chỉ có thể bán một lần trong một suất chiếu. Trạng thái trên giao diện có thể thay đổi trong khoảng thời gian rất ngắn; hai khách hàng có thể chọn cùng một ghế trước khi màn hình kịp đồng bộ. Hệ thống vì vậy phải kiểm soát lượt giữ ghế tại phía máy chủ, giới hạn thời gian giữ và sử dụng transaction cùng ràng buộc cơ sở dữ liệu để ngăn bán trùng.")
    paragraph(doc, "Thanh toán tạo thêm một ranh giới với hệ thống bên ngoài. Kết quả có thể được gửi qua callback hoặc webhook, có thể đến muộn hay lặp lại, trong khi người dùng cũng có thể tải lại trang hoặc khởi tạo thao tác nhiều lần. Đơn đặt vé chỉ được xác nhận sau khi chữ ký, số tiền, mã giao dịch và trạng thái hiện tại đều hợp lệ. Việc phát hành vé phải có tính lũy đẳng để cùng một kết quả thanh toán không tạo nhiều bộ vé.")
    paragraph(doc, "Ở khâu vận hành, vé điện tử phải được sử dụng đúng rạp, đúng suất chiếu và trong khoảng thời gian cho phép. Nhân viên chỉ được thao tác trong các rạp đã được phân công, còn quản trị viên chịu trách nhiệm quản lý toàn hệ thống và xử lý các nghiệp vụ nhạy cảm như xác nhận kết quả của yêu cầu hoàn tiền. Từ những yêu cầu này, thiết kế CinemaBooking tập trung vào bốn vấn đề: phân quyền đúng phạm vi, bảo toàn trạng thái giao dịch, xử lý yêu cầu đồng thời và cung cấp trải nghiệm cập nhật kịp thời cho người sử dụng.")

    doc.add_heading("3.2. Tác nhân và yêu cầu chức năng", level=2)
    paragraph(doc, "Hệ thống có ba tác nhân con người và hai hệ thống thanh toán bên ngoài. Khách hàng sử dụng các chức năng mua vé; nhân viên thực hiện công việc vận hành theo rạp được phân công; quản trị viên quản lý dữ liệu và nghiệp vụ trên toàn hệ thống. VNPay và SePay được xem là tác nhân phụ vì chủ động trao đổi kết quả giao dịch với back-end.")
    table(doc, ["Tác nhân", "Mục tiêu sử dụng", "Chức năng và phạm vi chính"], [
        ["Khách hàng (USER)", "Tìm rạp, suất chiếu phù hợp và hoàn tất quá trình mua vé.", "Đăng ký, xác thực email, đăng nhập, tra cứu phim/rạp/lịch chiếu, chủ động dùng vị trí hiện tại để ưu tiên rạp gần, giữ ghế, tạo đơn, áp dụng khuyến mãi, thanh toán, xem vé và quản lý hồ sơ, phiên đăng nhập của chính mình."],
        ["Nhân viên rạp (STAFF)", "Hỗ trợ vận hành tại những rạp được phân công.", "Xem dữ liệu rạp, phòng và ghế; quản lý suất chiếu theo quyền; theo dõi đơn và giao dịch trong phạm vi rạp; soát vé đúng rạp và suất chiếu. Không có quyền xác nhận kết quả yêu cầu hoàn tiền."],
        ["Quản trị viên (ADMIN)", "Quản trị và giám sát toàn bộ hệ thống.", "Có đầy đủ quyền quản lý danh mục, người dùng, phân quyền, suất chiếu, đơn đặt vé, thanh toán, yêu cầu hoàn tiền, thống kê và nhật ký vận hành trên tất cả các rạp."],
        ["VNPay", "Xử lý giao dịch qua cổng thanh toán.", "Tiếp nhận yêu cầu có chữ ký, cung cấp trang thanh toán và trả kết quả qua callback để hệ thống đối chiếu."],
        ["SePay", "Xác nhận giao dịch chuyển khoản ngân hàng.", "Gửi webhook sau khi ghi nhận tiền vào tài khoản; dữ liệu được kiểm tra khóa xác thực hoặc HMAC, nội dung và số tiền trước khi chấp nhận."],
    ], caption="Bảng 3.1. Tác nhân và yêu cầu chức năng chính của hệ thống")

    doc.add_heading("3.3. Yêu cầu phi chức năng", level=2)
    paragraph(doc, "Yêu cầu phi chức năng được xác định từ những rủi ro có thể ảnh hưởng đến độ tin cậy của quá trình bán vé. Do đề tài chưa thực hiện kiểm thử tải trong môi trường vận hành chính thức, các tiêu chí dưới đây tập trung vào hành vi có thể kiểm chứng bằng kiểm thử chức năng, kiểm thử tích hợp và phân tích dữ liệu, thay vì đưa ra thông lượng hoặc thời gian phản hồi chưa được đo đạc.")
    table(doc, ["Nhóm yêu cầu", "Tiêu chí cần đạt", "Định hướng thiết kế"], [
        ["Bảo mật", "Chỉ chủ thể hợp lệ được truy cập chức năng và dữ liệu thuộc phạm vi cho phép; token và thông tin nhạy cảm không được lưu, truyền dưới dạng dễ khai thác.", "JWT có chữ ký; refresh token được băm, luân chuyển và thu hồi; RBAC kết hợp phạm vi rạp; giới hạn tần suất đăng nhập và giữ ghế; ghi nhật ký bảo mật."],
        ["Toàn vẹn dữ liệu", "Một ghế không được bán cho nhiều khách trong cùng suất; callback lặp không tạo thêm vé; trạng thái chỉ chuyển khi thỏa điều kiện nghiệp vụ.", "Khóa ngoại, CHECK constraint, chỉ mục duy nhất, transaction, khóa bản ghi, cập nhật có điều kiện và kiểm tra idempotency."],
        ["Hiệu năng", "Danh sách lớn được phân trang; dữ liệu đọc thường xuyên không tạo số lượng truy vấn tăng tuyến tính theo quan hệ; dữ liệu ít thay đổi có thể được tái sử dụng.", "Projection và fetch query, batch fetching, chỉ mục theo truy vấn, tắt Open Session in View và cache chọn lọc."],
        ["Tính kịp thời", "Các trình duyệt đang xem cùng suất chiếu nhận được thay đổi ghế sau khi dữ liệu đã được ghi nhận thành công.", "REST cung cấp trạng thái ban đầu; STOMP/WebSocket phân phối sự kiện sau khi transaction commit."],
        ["Độ tin cậy", "Ghế và đơn chờ không bị giữ vô thời hạn; lỗi bên ngoài không làm mất khả năng đối soát; thao tác lặp có kết quả nhất quán.", "Thời hạn giữ ghế và thanh toán, scheduler xử lý dữ liệu hết hạn, payment event, reconciliation và quy tắc lũy đẳng."],
        ["Quyền riêng tư và khả dụng", "Chức năng định vị chỉ hoạt động sau khi người dùng chủ động yêu cầu; việc từ chối quyền, hết thời gian chờ, trình duyệt không hỗ trợ hoặc dữ liệu tọa độ rạp không hợp lệ không được chặn luồng tra cứu thủ công.", "Tọa độ hiện tại chỉ được giữ trong state của màn hình; CinemaBooking không lưu tọa độ chính xác vào localStorage hoặc cơ sở dữ liệu. Thành phố ưu tiên có thể được lưu để duy trì trải nghiệm giữa các lần truy cập."],
        ["Khả năng bảo trì", "Các thay đổi ở giao diện, nghiệp vụ và dữ liệu được giới hạn trong đúng thành phần; lược đồ có thể nâng cấp có kiểm soát.", "Kiến trúc phân lớp, DTO/mapper, adapter cổng thanh toán, xử lý ngoại lệ tập trung, Flyway migration và kiểm thử tự động."],
    ], font_size=8.5, caption="Bảng 3.2. Yêu cầu phi chức năng và định hướng thiết kế")

    doc.add_heading("3.4. Sơ đồ và đặc tả trường hợp sử dụng", level=2)
    paragraph(doc, "Sơ đồ trường hợp sử dụng mô tả ranh giới chức năng giữa khách hàng, nhân viên, quản trị viên và các hệ thống thanh toán. Những chức năng công khai như xem phim, rạp và lịch chiếu không yêu cầu đăng nhập; các nghiệp vụ tạo đơn, thanh toán, vận hành và quản trị chỉ được thực hiện sau khi xác định danh tính và quyền truy cập.")
    placeholder(doc, "[CHÈN FILE: usecase.png]", "Sơ đồ thể hiện ba tác nhân USER, STAFF, ADMIN; hai hệ thống bên ngoài VNPay, SePay; cùng ranh giới chức năng và quan hệ include/extend phù hợp.", caption="Hình 3.1. Sơ đồ trường hợp sử dụng tổng quát của hệ thống CinemaBooking")

    doc.add_heading("3.4.1. Trường hợp sử dụng của khách hàng", level=3)
    paragraph(doc, "Khách hàng có thể tra cứu nội dung công khai trước khi đăng nhập. Các chức năng giữ ghế, tạo đơn, thanh toán và xem vé yêu cầu tài khoản hợp lệ; dữ liệu đơn, giao dịch và vé còn được kiểm tra quyền sở hữu để ngăn truy cập chéo giữa các khách hàng.")
    table(doc, ["Mã", "Trường hợp sử dụng", "Tiền điều kiện", "Luồng chính", "Ngoại lệ chính"], [
        ["UC-U01", "Đăng ký và xác thực email", "Tên đăng nhập và email chưa được sử dụng.", "Tạo tài khoản, gửi thư xác thực và kích hoạt bằng token còn hiệu lực.", "Dữ liệu trùng, token sai hoặc hết hạn."],
        ["UC-U02", "Đăng nhập", "Tài khoản hợp lệ và đang hoạt động.", "Xác minh mật khẩu hoặc Google ID Token, sau đó phát phiên đăng nhập.", "Sai thông tin, vượt giới hạn tần suất, tài khoản bị khóa hoặc email chưa xác thực."],
        ["UC-U03", "Tra cứu rạp gần và lịch chiếu", "Không bắt buộc đăng nhập; chức năng Gần tôi cần trình duyệt hỗ trợ định vị và người dùng cấp quyền.", "Chọn khu vực thủ công hoặc yêu cầu vị trí hiện tại; hệ thống ưu tiên thành phố, rạp theo khoảng cách rồi hiển thị các suất đang mở bán.", "Từ chối quyền, định vị hết thời gian chờ, trình duyệt không hỗ trợ, tọa độ rạp không hợp lệ hoặc không có rạp trong bán kính 10 km; người dùng vẫn có thể chọn rạp thủ công."],
        ["UC-U04", "Giữ ghế và tạo đơn", "Đã đăng nhập; suất còn mở bán; ghế hợp lệ.", "Giữ ghế có thời hạn, xác nhận quyền giữ và tạo đơn chờ thanh toán.", "Ghế không còn khả dụng, lượt giữ hết hạn hoặc đơn chờ đã tồn tại."],
        ["UC-U05", "Áp dụng khuyến mãi", "Đơn đang chờ và chưa có mã thanh toán cố định số tiền.", "Kiểm tra điều kiện mã, cập nhật mức giảm và tổng thanh toán.", "Mã hết hạn, hết lượt, không đạt giá trị tối thiểu hoặc payment đang hoạt động."],
        ["UC-U06", "Thanh toán", "Đơn còn hiệu lực và thuộc khách hàng.", "Chọn phương thức, tạo giao dịch và chờ provider xác nhận.", "Sai chữ ký, lệch số tiền, giao dịch thất bại hoặc hết hạn."],
        ["UC-U07", "Xem và sử dụng vé", "Đơn đã thanh toán thành công.", "Tải vé theo từng ghế, hiển thị hoặc lưu QR và thông tin suất chiếu.", "Vé không thuộc khách hàng, đã hủy hoặc không còn hiệu lực."],
        ["UC-U08", "Quản lý hồ sơ và phiên", "Đã đăng nhập.", "Cập nhật hồ sơ, đổi mật khẩu, xem và thu hồi phiên đăng nhập.", "Mật khẩu hiện tại sai hoặc phiên không thuộc tài khoản."],
    ], font_size=8.0, caption="Bảng 3.3. Đặc tả các trường hợp sử dụng chính của khách hàng")

    doc.add_heading("3.4.2. Trường hợp sử dụng của nhân viên rạp", level=3)
    paragraph(doc, "Quyền của nhân viên được kiểm soát đồng thời theo permission và phạm vi rạp. Việc một lựa chọn không xuất hiện trên giao diện không được xem là biện pháp bảo mật; back-end tiếp tục giới hạn dữ liệu và xác minh rạp được phân công trước mỗi nghiệp vụ liên quan.")
    table(doc, ["Mã", "Trường hợp sử dụng", "Quyền yêu cầu", "Giới hạn dữ liệu"], [
        ["UC-S01", "Xem rạp phụ trách", "Quyền vận hành tương ứng.", "Chỉ các rạp được phân công cho nhân viên hiện tại."],
        ["UC-S02", "Tạo hoặc cập nhật suất chiếu", "SHOWTIME_CREATE hoặc SHOWTIME_UPDATE.", "Phòng chiếu phải thuộc rạp được phân công."],
        ["UC-S03", "Xem phòng và sơ đồ ghế", "ROOM_VIEW hoặc SEAT_VIEW.", "Chỉ phòng và ghế thuộc phạm vi rạp."],
        ["UC-S04", "Theo dõi đơn và giao dịch", "BOOKING_VIEW_ALL hoặc PAYMENT_VIEW_ALL.", "Kết quả truy vấn được giới hạn theo danh sách rạp được phép."],
        ["UC-S05", "Soát vé QR", "TICKET_CHECKIN.", "Vé phải đúng rạp, đúng suất chiếu và trong cửa sổ check-in."],
        ["UC-S06", "Xem số liệu vận hành", "DASHBOARD_VIEW hoặc ANALYTICS_VIEW.", "Số liệu chỉ thuộc các rạp được phân công."],
    ], font_size=8.0, caption="Bảng 3.4. Đặc tả các trường hợp sử dụng của nhân viên rạp")

    doc.add_heading("3.4.3. Trường hợp sử dụng của quản trị viên", level=3)
    paragraph(doc, "Quản trị viên có toàn bộ permission và không bị giới hạn theo rạp. Tuy nhiên, quyền rộng không thay thế các điều kiện nghiệp vụ: dữ liệu đang được tham chiếu phải tuân theo chính sách xóa, suất chiếu không hợp lệ không được tạo, vé sai trạng thái không được soát và kết quả yêu cầu hoàn tiền chỉ được ghi nhận khi yêu cầu đang ở trạng thái phù hợp.")
    table(doc, ["Nhóm nghiệp vụ", "Chức năng quản trị", "Kết quả hoặc dữ liệu liên quan"], [
        ["Danh mục", "Quản lý phim, rạp, phòng, ghế và sinh sơ đồ ghế theo lô.", "Dữ liệu nền phục vụ lập lịch và bán vé."],
        ["Lịch chiếu", "Tạo, cập nhật, hủy và theo dõi trạng thái suất chiếu.", "Suất chiếu, trạng thái ghế, đơn và yêu cầu hoàn tiền được tạo khi hủy."],
        ["Người dùng và phân quyền", "Tạo, cập nhật, khóa, xóa mềm, gán vai trò và phân công rạp.", "Tài khoản, vai trò, permission và phạm vi làm việc của nhân viên."],
        ["Khuyến mãi", "Tạo, cập nhật, xóa mềm và theo dõi khả năng sử dụng.", "Điều kiện giảm giá, giới hạn và số lượt đã giữ hoặc sử dụng."],
        ["Thanh toán và yêu cầu hoàn tiền", "Theo dõi giao dịch, sự kiện, đối soát; tạo yêu cầu và ghi nhận kết quả xử lý hoàn tiền.", "Payment, PaymentEvent, Refund và trạng thái đơn đặt vé."],
        ["Giám sát", "Xem dashboard, analytics và nhật ký quản trị, xác thực.", "Số liệu tổng hợp và thông tin truy vết hoạt động."],
    ], font_size=8.0, caption="Bảng 3.5. Các nhóm trường hợp sử dụng của quản trị viên")

    doc.add_heading("3.5. Kiến trúc tổng thể và phân rã mô-đun", level=2)
    paragraph(doc, "CinemaBooking được tổ chức theo kiến trúc nguyên khối mô-đun. Front-end React và back-end Spring Boot là hai ứng dụng độc lập giao tiếp qua RESTful API; PostgreSQL là nơi lưu trữ dữ liệu bền vững; WebSocket cung cấp kênh sự kiện; các adapter tách biệt việc tích hợp cổng thanh toán, thư điện tử và dịch vụ bên ngoài. Trong back-end, các mô-đun cùng được triển khai trong một tiến trình và dùng chung cơ sở dữ liệu, nhưng được phân chia theo trách nhiệm nghiệp vụ.")
    paragraph(doc, "Kiến trúc này phù hợp với phạm vi đề tài vì các quy trình giữ ghế, tạo đơn và xác nhận thanh toán có thể sử dụng transaction trong cùng một cơ sở dữ liệu. So với việc tách microservice sớm, cách tổ chức nguyên khối mô-đun giảm chi phí vận hành và giao tiếp phân tán, đồng thời vẫn duy trì ranh giới đủ rõ để mở rộng từng phân hệ về sau.")
    placeholder(doc, "[CHÈN FILE: Component Diagram.svg]", "Sơ đồ thể hiện React client, REST API, Spring Security, các mô-đun nghiệp vụ, Repository, PostgreSQL, WebSocket và adapter tích hợp bên ngoài.", caption="Hình 3.2. Sơ đồ thành phần của hệ thống CinemaBooking")
    table(doc, ["Mô-đun", "Đối tượng nghiệp vụ", "Trách nhiệm chính"], [
        ["Xác thực và người dùng", "Tài khoản, vai trò, permission, refresh token, phân công rạp", "Xác định danh tính, quản lý phiên, hồ sơ, quyền truy cập và phạm vi làm việc."],
        ["Danh mục rạp chiếu", "Phim, rạp, phòng, ghế và khuyến mãi", "Cung cấp dữ liệu nền cho tra cứu, lập lịch và tính giá."],
        ["Suất chiếu", "Suất chiếu và trạng thái ghế", "Lập lịch, mở bán, cập nhật trạng thái và xử lý hủy suất."],
        ["Đặt vé", "Lượt giữ ghế, đơn đặt vé và chi tiết đơn", "Giữ ghế, tính tiền, tạo đơn và quản lý vòng đời chờ thanh toán."],
        ["Thanh toán", "Giao dịch, sự kiện thanh toán và yêu cầu hoàn tiền", "Khởi tạo giao dịch, xác nhận kết quả, đối soát và quản lý yêu cầu hoàn tiền."],
        ["Vé điện tử", "Vé và mã QR", "Phát hành, hiển thị và soát vé theo đúng ngữ cảnh."],
        ["Thời gian thực và tác vụ nền", "Sự kiện ghế và dữ liệu phụ thuộc thời gian", "Đồng bộ trạng thái sau commit và xử lý dữ liệu hết hạn."],
        ["Thống kê và nhật ký", "Dữ liệu tổng hợp và nhật ký hoạt động", "Cung cấp dashboard, báo cáo và khả năng truy vết."],
    ], caption="Bảng 3.6. Phân rã các mô-đun nghiệp vụ của hệ thống")

    doc.add_heading("3.6. Thiết kế lớp và trạng thái nghiệp vụ", level=2)
    doc.add_heading("3.6.1. Mô hình lớp miền nghiệp vụ", level=3)
    paragraph(doc, "Mô hình lớp biểu diễn các đối tượng cốt lõi và quan hệ giữa chúng. Nhóm người dùng liên kết với vai trò, permission, phiên đăng nhập và rạp được phân công. Nhóm rạp chiếu tổ chức theo chuỗi rạp–phòng–ghế; phim được lập lịch thành suất chiếu trong một phòng; mỗi ghế của suất chiếu có một trạng thái riêng. Đơn đặt vé gồm nhiều chi tiết ghế, liên kết với giao dịch, vé điện tử, sự kiện thanh toán và yêu cầu hoàn tiền.")
    placeholder(doc, "[CHÈN FILE: Class Diagram.svg]", "Sơ đồ trình bày các thuộc tính chính và mối quan hệ giữa các lớp miền; không hiển thị getter, setter hoặc những chi tiết triển khai không ảnh hưởng đến thiết kế.", caption="Hình 3.3. Sơ đồ lớp miền nghiệp vụ của hệ thống CinemaBooking")

    doc.add_heading("3.6.2. Trạng thái và bất biến nghiệp vụ", level=3)
    paragraph(doc, "Các đối tượng giao dịch được quản lý bằng trạng thái thay vì xóa hoặc ghi đè lịch sử. Điều kiện chuyển trạng thái được kiểm tra ở tầng nghiệp vụ, trong khi CHECK constraint và chỉ mục duy nhất tạo lớp bảo vệ bổ sung tại cơ sở dữ liệu. Bảng 3.7 tổng hợp những trạng thái có ảnh hưởng trực tiếp đến quy trình bán vé.")
    table(doc, ["Đối tượng", "Trạng thái", "Chuyển đổi tiêu biểu", "Bất biến cần bảo đảm"], [
        ["SeatStatus", "AVAILABLE, HOLD, BOOKED", "AVAILABLE→HOLD; HOLD→BOOKED hoặc AVAILABLE.", "HOLD phải có người giữ và thời hạn; chỉ một trạng thái cho mỗi cặp ghế–suất."],
        ["Booking", "PENDING, SUCCESS, FAILED, CANCELLED, EXPIRED, REFUND_PENDING, REFUNDED", "PENDING→trạng thái kết thúc; SUCCESS→REFUND_PENDING→REFUNDED.", "REFUNDED là trạng thái nội bộ sau khi người vận hành ghi nhận yêu cầu đã được xử lý; tối đa một đơn PENDING cho mỗi khách và suất."],
        ["Payment", "PENDING, SUCCESS, FAILED, EXPIRED, REFUND_PENDING, REFUNDED, REFUND_FAILED", "PENDING→SUCCESS/FAILED/EXPIRED; SUCCESS→nhánh quản lý yêu cầu hoàn tiền.", "Tối đa một payment PENDING cho mỗi booking; mã giao dịch không trùng khi có giá trị."],
        ["Ticket", "ACTIVE, USED, CANCELLED", "ACTIVE→USED hoặc CANCELLED.", "Mỗi chi tiết đơn có tối đa một vé; vé USED lưu thời gian và nhân viên soát."],
        ["Showtime", "UPCOMING, ONGOING, ENDED, CANCELLED", "UPCOMING→ONGOING→ENDED hoặc CANCELLED.", "Thời gian kết thúc sau thời gian bắt đầu; không trùng phòng trong khoảng dọn phòng 15 phút; suất hủy không được mở bán."],
        ["Refund", "PENDING, PROCESSING, SUCCESS, FAILED, CANCELLED", "PENDING→PROCESSING→SUCCESS hoặc FAILED.", "SUCCESS ghi nhận kết quả do người vận hành xác nhận; hệ thống chưa tự gọi API hoàn tiền của nhà cung cấp."],
    ], font_size=8.0, caption="Bảng 3.7. Trạng thái chính và các bất biến nghiệp vụ")
    paragraph(doc, "Các trạng thái REFUND_PENDING, REFUNDED và REFUND_FAILED phản ánh tiến trình quản lý yêu cầu trong cơ sở dữ liệu. Chúng không đồng nghĩa Back-end đã tự động thực hiện giao dịch chuyển tiền qua API của VNPay hoặc SePay. Trong phiên bản hiện tại, quản trị viên xử lý việc hoàn trả theo quy trình bên ngoài hệ thống, sau đó ghi nhận kết quả thành công hoặc thất bại trên màn hình vận hành.")
    placeholder(doc, "[CHÈN FILE: [WORD Hình 3.9] Sơ đồ trạng thái đơn đặt vé.png]", "Sơ đồ làm rõ các nhánh từ PENDING đến SUCCESS, FAILED, CANCELLED, EXPIRED và quá trình quản lý yêu cầu từ REFUND_PENDING đến trạng thái được người vận hành ghi nhận.", caption="Hình 3.4. Sơ đồ trạng thái của đơn đặt vé")
    placeholder(doc, "[CHÈN FILE: Hình 3.10] Sơ đồ trạng thái giao dịch thanh toán.png]", "Sơ đồ thể hiện vòng đời Payment và nhánh quản lý yêu cầu hoàn tiền; REFUNDED là trạng thái được ghi nhận trong hệ thống, không phải kết quả từ API hoàn tiền tự động.", caption="Hình 3.5. Sơ đồ trạng thái của giao dịch thanh toán")

    doc.add_heading("3.7. Thiết kế cơ sở dữ liệu", level=2)
    paragraph(doc, "Mô hình dữ liệu được xây dựng theo hướng chuẩn hóa các thực thể nghiệp vụ và sử dụng khóa ngoại để duy trì mối quan hệ. Những dữ liệu có lịch sử như đơn đặt vé, giao dịch, vé, sự kiện thanh toán và yêu cầu hoàn tiền được lưu thành các bảng riêng, cho phép truy vết một giao dịch từ thời điểm khởi tạo đến kết quả được ghi nhận.")
    placeholder(doc, "[CHÈN FILE: erd.svg]", "Sử dụng ERD tổng quan với tên bảng, khóa chính, khóa ngoại và quan hệ chính có thể đọc được khi in trên khổ A4.", caption="Hình 3.6. Sơ đồ quan hệ dữ liệu tổng quan của hệ thống CinemaBooking")
    table(doc, ["Nhóm dữ liệu", "Các bảng chính", "Quan hệ và ràng buộc nổi bật"], [
        ["Người dùng và phân quyền", "users, roles, permissions, users_roles, roles_permissions, staff_cinemas", "Tên đăng nhập và email duy nhất; quan hệ nhiều–nhiều giữa người dùng, vai trò và permission; phân công nhân viên theo rạp."],
        ["Danh mục và lịch chiếu", "movies, cinemas, rooms, seats, showtimes, promotions", "Rạp lưu latitude và longitude phục vụ marker, tính khoảng cách; ghế duy nhất trong phòng; thời gian suất hợp lệ; trạng thái được giới hạn; dữ liệu danh mục hỗ trợ xóa mềm."],
        ["Đặt vé và vé điện tử", "seat_status, bookings, booking_details, tickets", "Một trạng thái cho mỗi ghế–suất; dữ liệu HOLD phải nhất quán; một vé cho mỗi chi tiết đơn; giới hạn đơn PENDING."],
        ["Thanh toán và yêu cầu hoàn tiền", "payments, payment_events, refunds", "Mã giao dịch duy nhất; giới hạn payment PENDING và refund đang hoạt động; lưu phản hồi sự kiện dưới dạng JSONB."],
        ["Phiên và nhật ký", "refresh_tokens, invalidated_token, auth_audit_logs, admin_audit_logs", "Refresh token được băm; có chỉ mục theo phiên, thời hạn, tác nhân, hành động và thời gian."],
    ], font_size=8.0, caption="Bảng 3.8. Phân nhóm dữ liệu và các ràng buộc quan trọng")
    paragraph(doc, "Chỉ mục được thiết kế theo các đường truy vấn có tần suất cao hoặc ảnh hưởng trực tiếp đến tác vụ nền, gồm tra cứu suất chiếu theo phim, rạp và thời gian; tải sơ đồ ghế; tìm đơn hoặc payment đang chờ; quét dữ liệu hết hạn; lọc sự kiện thanh toán và nhật ký. Bên cạnh chỉ mục thông thường, chỉ mục duy nhất một phần được sử dụng để giới hạn các bản ghi đang hoạt động mà vẫn giữ được lịch sử trạng thái kết thúc.")
    paragraph(doc, "Bảng cinemas lưu vĩ độ và kinh độ của rạp dưới dạng số thực. API bản đồ chỉ chọn các rạp đang hoạt động, chưa bị xóa và có đủ hai giá trị tọa độ; phía client tiếp tục kiểm tra miền hợp lệ của vĩ độ [-90, 90] và kinh độ [-180, 180] trước khi tạo marker hoặc tính khoảng cách. Cách kiểm tra hai lớp giúp một bản ghi thiếu hoặc sai tọa độ không làm hỏng toàn bộ danh sách bản đồ.")
    paragraph(doc, "Phim, rạp, phòng, ghế, suất chiếu, khuyến mãi và người dùng áp dụng xóa mềm để không làm mất quan hệ với dữ liệu đã phát sinh. Đơn đặt vé, chi tiết đơn, payment, ticket, payment event và yêu cầu refund được giữ lại phục vụ đối soát, khiếu nại và kiểm toán. Việc hủy hoặc ngừng khai thác được biểu diễn bằng trạng thái nghiệp vụ thay vì xóa vật lý dữ liệu giao dịch.")

    doc.add_heading("3.8. Thiết kế RESTful API", level=2)
    paragraph(doc, "RESTful API là ranh giới giao tiếp giữa front-end và back-end. URI được tổ chức theo tài nguyên; HTTP method thể hiện thao tác; DTO xác định dữ liệu đầu vào và đầu ra; status code cùng cấu trúc ApiResponse biểu diễn kết quả thống nhất. Controller tiếp nhận yêu cầu và validation, Service thực hiện quyền truy cập, quy tắc nghiệp vụ và transaction, còn Repository làm việc với PostgreSQL.")
    paragraph(doc, "Các API tra cứu phim, rạp, lịch chiếu và sơ đồ ghế được công khai nhưng chỉ trả dữ liệu có thể hiển thị. API của khách hàng yêu cầu đăng nhập và tiếp tục kiểm tra quyền sở hữu. API vận hành yêu cầu permission và phạm vi rạp. Callback hoặc webhook được mở cho hệ thống thanh toán bên ngoài, nhưng kết quả chỉ được chấp nhận sau khi xác minh chữ ký hoặc khóa xác thực, số tiền, mã giao dịch và trạng thái hiện tại.")
    representative_specs = [
        ("Xác thực", "POST", "/auth/token", "Tên đăng nhập và mật khẩu", "Công khai", "Tạo phiên đăng nhập và cấp access token."),
        ("Bản đồ rạp", "GET", "/api/v1/cinemas/map", "Không có body; trả rạp có tọa độ", "Công khai", "Cung cấp id, tên, địa chỉ, thành phố, latitude, longitude và trạng thái để client hiển thị marker, tính khoảng cách."),
        ("Tra cứu sơ đồ ghế", "GET", "/api/v1/showtimes/{id}/seats", "Mã suất chiếu", "Công khai", "Trả trạng thái hiện tại của các ghế trong suất chiếu."),
        ("Giữ ghế", "POST", "/api/v1/bookings/hold", "Mã suất chiếu và danh sách mã ghế", "Khách hàng có quyền đặt vé", "Giữ các ghế hợp lệ trong một khoảng thời gian xác định."),
        ("Tạo đơn", "POST", "/api/v1/bookings", "Mã suất chiếu, danh sách mã ghế và mã khuyến mãi nếu có", "Khách hàng có quyền đặt vé", "Tạo đơn chờ thanh toán từ lượt giữ ghế còn hiệu lực."),
        ("Áp dụng khuyến mãi", "PATCH", "/api/v1/bookings/{id}/promotion", "Mã đơn đặt vé và mã khuyến mãi", "Chủ sở hữu đơn đặt vé", "Kiểm tra điều kiện và tính lại số tiền tại server."),
        ("Khởi tạo thanh toán", "POST", "/api/v1/payments/initiate", "Mã đơn, phương thức và số tiền cần thanh toán", "Chủ sở hữu có quyền thanh toán", "Tạo hoặc sử dụng lại giao dịch chờ phù hợp."),
        ("Nhận kết quả VNPay", "GET", "/api/v1/payments/vnpay-callback", "Các tham số giao dịch và chữ ký do VNPay gửi", "Công khai cho VNPay", "Xác minh kết quả và chuyển người dùng về trang kết quả."),
        ("Nhận thông báo SePay", "POST", "/api/v1/payments/sepay-webhook", "Thông tin giao dịch chuyển khoản và dữ liệu xác thực", "Công khai cho SePay", "Đối chiếu giao dịch trước khi xác nhận thanh toán."),
        ("Soát vé", "POST", "/api/v1/tickets/check-in", "Mã QR, mã rạp và mã suất chiếu", "Nhân viên có quyền soát vé", "Xác minh ngữ cảnh và chuyển vé hợp lệ sang trạng thái đã sử dụng."),
        ("Hủy suất chiếu", "POST", "/api/v1/showtimes/{id}/cancel", "Mã suất chiếu và lý do hủy", "Người có quyền cập nhật suất chiếu", "Hủy suất theo chính sách và tạo bản ghi yêu cầu hoàn tiền khi cần."),
    ]
    endpoint_index = {(endpoint.http_method, endpoint.endpoint): endpoint for endpoint in endpoints}
    representative_rows = []
    for module, method, path, request_data, access, purpose in representative_specs:
        endpoint = endpoint_index.get((method, path))
        if endpoint is None:
            raise ValueError(f"Representative API is missing from controller inventory: {method} {path}")
        representative_rows.append([
            module,
            f"{method} {path}",
            request_data,
            access,
            purpose,
        ])
    table(doc, ["Nghiệp vụ", "API tiêu biểu", "Dữ liệu vào", "Quyền truy cập", "Kết quả"], representative_rows, font_size=7.5, caption="Bảng 3.9. Một số API tiêu biểu của các luồng nghiệp vụ chính")


def add_chapter_four(doc: Document, counts: dict[str, int]) -> None:
    doc.add_heading("CHƯƠNG 4. HIỆN THỰC HỆ THỐNG VÀ KIỂM THỬ", level=1)
    doc.add_heading("4.1. Môi trường phát triển", level=2)
    paragraph(doc, "CinemaBooking được phát triển theo mô hình client-server. Phần Back-end cung cấp REST API và các kênh WebSocket, phần Front-end đảm nhiệm giao diện và tương tác với người dùng, còn PostgreSQL lưu trữ dữ liệu nghiệp vụ. Các công cụ và phiên bản chính được sử dụng trong quá trình hiện thực được trình bày tại Bảng 4.1.")
    table(doc, ["Thành phần", "Phiên bản và cấu hình sử dụng"], [
        ["Back-end", "Java 21, Maven và Spring Boot 3.5.14."],
        ["Front-end", "React 19, TypeScript và Vite."],
        ["Cơ sở dữ liệu", "PostgreSQL 15; Flyway quản lý các migration từ V1 đến V14; Hibernate sử dụng chế độ ddl-auto=validate."],
        ["Môi trường kiểm thử", "JUnit, Spring Boot Test, Testcontainers, PostgreSQL container và Vitest."],
        ["Tài liệu API", "Springdoc OpenAPI và giao diện Swagger UI."],
    ], caption="Bảng 4.1. Môi trường và công cụ phát triển")
    doc.add_heading("4.2. Hiện thực Back-end", level=2)
    paragraph(doc, "Back-end được tổ chức theo kiến trúc phân lớp. Controller tiếp nhận request và kiểm tra dữ liệu đầu vào; Service thực hiện các quy tắc nghiệp vụ trong phạm vi transaction; Repository truy cập PostgreSQL thông qua Spring Data JPA. Entity chỉ được sử dụng bên trong tầng nghiệp vụ, trong khi dữ liệu trao đổi với Front-end được biểu diễn bằng DTO. Cách tổ chức này giúp tách trách nhiệm giữa các tầng và hạn chế việc thay đổi giao diện API ảnh hưởng trực tiếp đến mô hình dữ liệu.")
    table(doc, ["Tầng", "Trách nhiệm", "Thành phần tiêu biểu"], [
        ["Controller", "Tiếp nhận REST request, validation dữ liệu và kiểm tra permission trước khi chuyển sang tầng nghiệp vụ.", "BookingController, PaymentController, TicketController."],
        ["Service", "Xử lý quy tắc nghiệp vụ, transaction, trạng thái booking/payment và phạm vi rạp của nhân viên.", "BookingServiceImpl, PaymentServiceImpl, RefundServiceImpl."],
        ["Repository", "Thực hiện truy vấn, phân trang, projection và khóa dữ liệu khi xử lý đồng thời.", "SeatStatusRepository, BookingRepository, PaymentRepository."],
        ["Entity, DTO và Mapper", "Biểu diễn dữ liệu lưu trữ và chuyển đổi dữ liệu trao đổi giữa API với tầng nghiệp vụ.", "Booking, Payment, SeatStatus, HoldSeatRequest và các response DTO."],
        ["Security", "Xác thực JWT, quản lý phiên đăng nhập, giới hạn tần suất và kiểm soát truy cập theo permission.", "SecurityConfig, AuthenticationService, CustomJwtDecoder."],
        ["Hạ tầng nghiệp vụ", "Kết nối cổng thanh toán, phát sự kiện realtime, gửi email và thực thi tác vụ định kỳ.", "PaymentGateway, SeatStatusPublisher, EmailService và các Scheduler."],
    ], caption="Bảng 4.2. Trách nhiệm các lớp trong Back-end")
    doc.add_heading("4.2.1. Hiện thực xác thực và quản lý phiên", level=3)
    paragraph(doc, "Hệ thống hỗ trợ đăng nhập bằng tên đăng nhập và mật khẩu hoặc bằng Google ID token. Sau khi thông tin đăng nhập được xác minh, Back-end phát access token có thời hạn ngắn để gọi API và refresh token dùng để duy trì phiên. Refresh token không được lưu ở dạng nguyên bản trong cơ sở dữ liệu mà được băm trước khi lưu; bản ghi phiên còn chứa thời hạn, thông tin thiết bị và trạng thái thu hồi.")
    paragraph(doc, "Khi access token hết hạn, client gửi refresh token để yêu cầu cấp cặp token mới. AuthenticationService khóa bản ghi phiên trong thời gian xử lý, kiểm tra thời hạn và trạng thái token, sau đó thu hồi token cũ trước khi phát token mới. Cơ chế refresh token rotation này làm giảm khả năng một token đã bị sử dụng tiếp tục được tái sử dụng [14]. Khi người dùng đăng xuất, đổi mật khẩu hoặc thu hồi phiên, hệ thống vô hiệu hóa token tương ứng; trường authVersion trên tài khoản hỗ trợ loại bỏ các access token được phát trước thời điểm thay đổi bảo mật.")
    paragraph(doc, "Các request đăng nhập được giới hạn tần suất theo cửa sổ thời gian nhằm giảm nguy cơ dò mật khẩu tự động. Tuy nhiên, cơ chế hiện tại lưu bộ đếm trong bộ nhớ của một instance ứng dụng. Vì vậy, distributed rate limiting dùng chung giữa nhiều máy chủ được xác định là hướng phát triển khi hệ thống triển khai theo mô hình phân tán.")

    doc.add_heading("4.2.2. Hiện thực giữ ghế và tạo đơn đặt vé", level=3)
    paragraph(doc, "Giữ ghế là nghiệp vụ có nguy cơ xảy ra race condition cao nhất vì nhiều người dùng có thể chọn cùng một ghế gần như đồng thời. BookingServiceImpl xử lý yêu cầu trong transaction và truy vấn các bản ghi SeatStatus cần thay đổi bằng khóa ghi. Danh sách ghế được khóa theo thứ tự ổn định, sau đó hệ thống chỉ chuyển trạng thái từ AVAILABLE sang HOLD khi toàn bộ ghế vẫn còn hợp lệ. Thông tin người giữ và thời điểm hết hạn được lưu cùng trạng thái ghế; trường version hỗ trợ phát hiện xung đột cập nhật ở tầng JPA.")
    paragraph(doc, "Sau khi giữ ghế thành công, người dùng có thể tạo booking ở trạng thái PENDING. Giá vé được tính lại tại Back-end từ giá cơ bản của suất chiếu và hệ số loại ghế, thay vì tin vào số tiền do client gửi lên. BookingDetail được tạo cho từng ghế và thời hạn thanh toán được lưu tại paymentExpiresAt. Ràng buộc duy nhất trong cơ sở dữ liệu ngăn một người dùng tạo nhiều booking PENDING cho cùng một suất chiếu, qua đó bổ sung một lớp bảo vệ cho trường hợp double click hoặc request bị gửi lặp.")
    paragraph(doc, "Sự kiện thay đổi trạng thái ghế chỉ được phát qua WebSocket sau khi transaction đã hoàn tất. Các client đang theo dõi cùng suất chiếu vì vậy nhận được trạng thái đã được commit, tránh trường hợp giao diện hiển thị một thay đổi sau đó bị rollback trong cơ sở dữ liệu. Yêu cầu giữ ghế cũng được giới hạn tần suất để hạn chế hành vi spam và giảm tải cho các truy vấn khóa.")

    doc.add_heading("4.2.3. Hiện thực thanh toán, ghi nhận sự kiện và quản lý yêu cầu hoàn tiền", level=3)
    paragraph(doc, "PaymentServiceImpl là thành phần điều phối vòng đời của giao dịch thanh toán, còn PaymentGateway định nghĩa giao diện xử lý chung cho các cổng thanh toán. Cách tổ chức này giúp tách các quy tắc nghiệp vụ của hệ thống khỏi những yêu cầu kỹ thuật riêng của từng nhà cung cấp. Đối với VNPay, hệ thống tạo địa chỉ chuyển hướng để người dùng thực hiện thanh toán và tiếp nhận kết quả thông qua callback. Đối với SePay, hệ thống tạo mã QR chuyển khoản và tiếp nhận webhook khi giao dịch ngân hàng được ghi nhận.")
    paragraph(doc, "Trước khi xác nhận một giao dịch thành công, hệ thống đối chiếu mã giao dịch, số tiền, trạng thái đơn đặt vé và thông tin xác thực tương ứng với từng cổng thanh toán. Kết quả chỉ được cập nhật khi các điều kiện kiểm tra đều hợp lệ. Quá trình này được thiết kế theo nguyên tắc idempotency, nghĩa là cùng một kết quả thanh toán chỉ được ghi nhận một lần, kể cả khi callback hoặc webhook được nhà cung cấp gửi lại nhiều lần.")
    paragraph(doc, "Các sự kiện phát sinh trong quá trình thanh toán được lưu trong PaymentEvent để phục vụ truy vết, đối soát và phân tích sự cố. Phương thức ghi nhận sự kiện của PaymentEventServiceImpl được cấu hình với cơ chế truyền giao dịch REQUIRES_NEW. Vì vậy, thao tác này được thực hiện trong một transaction độc lập và có thể lưu lại dấu vết xử lý ngay cả khi transaction nghiệp vụ chính phát sinh lỗi và bị hoàn tác. Cơ chế này giúp quản trị viên xác định thời điểm tiếp nhận dữ liệu, kết quả kiểm tra và nguyên nhân thất bại của từng giao dịch.")
    paragraph(doc, "Khi một suất chiếu có đơn đã thanh toán bị hủy, RefundServiceImpl tạo yêu cầu hoàn tiền và quản lý mối liên hệ giữa yêu cầu này với đơn đặt vé và giao dịch thanh toán. Trong phạm vi đề tài, phân hệ này hỗ trợ tạo yêu cầu, theo dõi trạng thái và ghi nhận kết quả xử lý. Hệ thống chưa tích hợp API hoàn tiền trực tiếp của VNPay hoặc SePay. Việc hoàn trả được thực hiện theo quy trình bên ngoài hệ thống, sau đó quản trị viên có quyền cập nhật kết quả thành công hoặc thất bại. Do đó, trạng thái REFUNDED được hiểu là kết quả hoàn tiền đã được người vận hành ghi nhận trong CinemaBooking, thay vì một giao dịch chuyển tiền được Back-end tự động thực hiện.")
    table(doc, ["Nhóm thông tin", "Dữ liệu theo dõi", "Ý nghĩa vận hành"], [
        ["Phạm vi thống kê", "Khoảng thời gian, tổng số sự kiện và số đơn đặt vé bị ảnh hưởng", "Xác định phạm vi dữ liệu được theo dõi trong từng lần kiểm tra."],
        ["Kết quả xử lý", "Số sự kiện thành công, thất bại và tỷ lệ lỗi", "Đánh giá mức độ ổn định của quá trình xác nhận thanh toán."],
        ["Nguyên nhân lỗi", "Lỗi nhà cung cấp, chữ ký không hợp lệ, sai lệch số tiền, giao dịch thất bại hoặc hết hạn", "Hỗ trợ phân loại sự cố và xác định hướng kiểm tra phù hợp."],
        ["Phân bố theo cổng", "Số lượng lỗi theo từng phương thức thanh toán", "Khoanh vùng cổng thanh toán phát sinh nhiều lỗi trong khoảng thời gian được khảo sát."],
        ["Chỉ báo giám sát", "Tỷ lệ lỗi và giá trị healthStatus", "Cung cấp thông tin tổng hợp cho màn hình vận hành thanh toán."],
    ], font_size=8.5, caption="Bảng 4.3. Các nhóm thông tin giám sát sự kiện thanh toán")
    paragraph(doc, "Các nhóm thông tin trong Bảng 4.3 được tổng hợp từ nhật ký PaymentEvent thông qua API quản trị được bảo vệ bởi quyền truy cập thanh toán. Chức năng này hỗ trợ quản trị viên phát hiện các trường hợp sai chữ ký, sai lệch số tiền và lỗi phát sinh theo từng cổng. Đây là cơ chế giám sát nghiệp vụ ở cấp ứng dụng; việc thu thập log, metric và cảnh báo tập trung trên nhiều máy chủ được xác định là hướng phát triển của hệ thống.")
    doc.add_heading("4.2.4. Validation, xử lý ngoại lệ và audit", level=3)
    paragraph(doc, "Dữ liệu đầu vào được kiểm tra bằng Jakarta Bean Validation tại các request DTO. GlobalExceptionHandler chuyển các lỗi nghiệp vụ, lỗi validation, lỗi xác thực, lỗi phân quyền và một số lỗi persistence thành cấu trúc ApiResponse thống nhất gồm mã lỗi, thông báo, đường dẫn và thời điểm xảy ra lỗi. Nhờ đó, Front-end có thể hiển thị thông báo phù hợp mà không phụ thuộc vào stack trace hoặc thông điệp kỹ thuật của cơ sở dữ liệu.")
    paragraph(doc, "Các thao tác quản trị và sự kiện xác thực quan trọng được ghi vào nhật ký audit. Đối với callback hoặc webhook public, việc không yêu cầu JWT chỉ nhằm cho phép hệ thống bên ngoài gọi vào endpoint; payload vẫn phải vượt qua bước xác minh chữ ký hoặc API key và được đối chiếu với dữ liệu thanh toán đang lưu trước khi làm thay đổi trạng thái nghiệp vụ.")

    doc.add_heading("4.2.5. Tối ưu truy vấn và sử dụng cache", level=3)
    paragraph(doc, "Caffeine cache được áp dụng cho những dữ liệu có tần suất đọc cao nhưng ít thay đổi như phim, rạp, bản đồ rạp, phòng, ghế cấu hình và khuyến mãi. Khi dữ liệu tương ứng được cập nhật, cache được xóa để lần đọc tiếp theo lấy lại dữ liệu mới. Trạng thái ghế theo suất chiếu không được cache vì thay đổi liên tục và yêu cầu tính nhất quán cao trong quá trình giữ ghế, thanh toán và trả ghế.")
    paragraph(doc, "Cấu hình JPA tắt Open Session in View, đặt default batch fetch size bằng 50 và sử dụng projection hoặc fetch query cho các màn hình danh sách. Các API phân trang được chuyển sang PageResponse ổn định trước khi trả về client. Các biện pháp này giúp giảm nguy cơ N+1 query và tránh phát sinh truy vấn ngoài phạm vi Service; tuy nhiên, mỗi endpoint mới vẫn cần được theo dõi bằng SQL profile và EXPLAIN ANALYZE với dữ liệu thực tế thay vì mặc định cho rằng N+1 đã được loại bỏ hoàn toàn.")

    doc.add_heading("4.2.6. Tác vụ định kỳ, email và thông báo realtime", level=3)
    paragraph(doc, "Các Scheduler xử lý những công việc không phụ thuộc vào thao tác trực tiếp của người dùng, bao gồm trả ghế HOLD đã hết hạn, đóng booking PENDING quá thời gian thanh toán, cập nhật trạng thái suất chiếu và dọn token không còn hiệu lực. Tác vụ đồng bộ lịch chiếu chuyển suất UPCOMING sang ONGOING khi đến giờ bắt đầu và sang ENDED khi hết giờ. Các suất CANCELLED được giữ nguyên để bảo toàn lịch sử vận hành. Khi trạng thái ghế thay đổi, SeatStatusPublisher gửi sự kiện STOMP đến đúng topic của suất chiếu để các client đang mở sơ đồ ghế cập nhật ngay.")
    paragraph(doc, "Khi tạo hoặc điều chỉnh một suất chiếu, dịch vụ kiểm tra xung đột trong cùng phòng và dành thêm 15 phút cho việc dọn, ổn định và chuẩn bị phòng ở hai đầu khoảng thời gian. Danh sách lịch công khai áp dụng chính sách khác với màn hình quản trị: chỉ các suất UPCOMING còn khả năng đặt vé, bắt đầu sau thời điểm ngừng bán 15 phút và nằm trong bảy ngày được công bố mới được gửi đến khách hàng. Nhờ đó, lịch đã qua hoặc đã hủy vẫn được lưu để quản trị nhưng không tiếp tục xuất hiện trong hành trình mua vé.")
    paragraph(doc, "EmailService đảm nhiệm thư xác thực tài khoản, khôi phục mật khẩu, gửi vé điện tử và thông báo hủy suất chiếu. Việc tách nội dung email thành template giúp phần trình bày độc lập với logic nghiệp vụ và thuận tiện hơn khi cần thay đổi nhận diện hoặc bổ sung loại thông báo mới.")
    doc.add_heading("4.3. Hiện thực Front-end", level=2)
    paragraph(doc, "Front-end được xây dựng dưới dạng single-page application bằng React và TypeScript. AppRouter phân chia các route công khai, route yêu cầu đăng nhập và route dành cho quản trị viên hoặc nhân viên. Những page có kích thước lớn được tải theo nhu cầu để giảm lượng JavaScript cần tải ở lần truy cập đầu tiên. ProtectedRoute kiểm tra trạng thái đăng nhập và permission trước khi hiển thị nội dung được bảo vệ.")
    paragraph(doc, "Mọi request REST được gửi qua một Axios client dùng chung. Access token được gắn vào header Authorization, còn refresh token được trình duyệt gửi bằng cookie. Khi nhiều request đồng thời nhận lỗi 401 do access token hết hạn, client chỉ tạo một request refresh rồi cho các request còn lại chờ kết quả; cách xử lý single-flight này tránh gửi nhiều yêu cầu làm mới phiên cùng lúc. React Query quản lý dữ liệu từ server, trong khi auth store lưu trạng thái phiên và quyền phục vụ điều hướng giao diện.")
    table(doc, ["Khu vực", "Thành phần hiện thực", "Vai trò"], [
        ["Điều hướng và bố cục", "AppRouter, ProtectedRoute và các layout", "Tổ chức route, lazy loading và kiểm soát truy cập theo trạng thái đăng nhập hoặc permission."],
        ["Giao tiếp API", "axiosClient, các module API và auth store", "Gửi request, gắn access token, làm mới phiên và chuẩn hóa lỗi trả về."],
        ["Tra cứu rạp theo vị trí", "CinemaMapPage, RegionalShowtimeBrowser và utils/location", "Xin quyền định vị theo thao tác người dùng, kiểm tra tọa độ, tính khoảng cách Haversine và ưu tiên rạp gần."],
        ["Đặt vé và realtime", "SeatSelectionPage, bookingApi và useSeatWebSocket", "Tải sơ đồ ghế ban đầu, giữ ghế và nhận thay đổi trạng thái theo suất chiếu."],
        ["Thanh toán và vé", "CheckoutPage, PaymentResultPage, MyBookingsPage và TicketDetailPage", "Khởi tạo thanh toán, theo dõi kết quả và hiển thị đơn đặt vé hoặc vé điện tử."],
        ["Quản trị và nhân viên", "Các page trong khu vực admin và staff", "Vận hành dữ liệu hệ thống, thanh toán, yêu cầu hoàn tiền, rạp phụ trách và soát vé."],
        ["Khả năng phục hồi giao diện", "ErrorBoundary, React Query, toast và utility", "Xử lý lỗi runtime, cache dữ liệu đọc, thông báo và định dạng dữ liệu nhất quán."],
    ], caption="Bảng 4.4. Cấu trúc hiện thực Front-end")
    doc.add_heading("4.3.1. Giao diện xác thực và tra cứu thông tin", level=3)
    paragraph(doc, "Người dùng có thể tra cứu phim, rạp và lịch chiếu mà chưa cần đăng nhập. Lịch công khai chỉ hiển thị các suất còn mở bán trong khoảng thời gian hệ thống công bố; suất đã qua, đang chiếu, đã kết thúc hoặc bị hủy không được đưa vào lựa chọn đặt vé. Chỉ khi bắt đầu giữ ghế hoặc sử dụng chức năng tài khoản, hệ thống mới điều hướng đến trang đăng nhập. Các biểu mẫu đăng nhập, đăng ký và khôi phục mật khẩu kiểm tra dữ liệu ở phía client để phản hồi sớm, nhưng Back-end vẫn thực hiện validation độc lập nhằm bảo đảm quy tắc không bị bỏ qua khi API được gọi từ một client khác.")
    paragraph(doc, "Chức năng Gần tôi xuất hiện tại khu vực mua vé theo rạp trên trang chủ và trang bản đồ rạp. Khi người dùng nhấn nút, client mới gọi Browser Geolocation API; trong thời gian chờ, nút chuyển sang trạng thái đang xử lý để hạn chế thao tác lặp. Cấu hình dùng thời gian chờ 10 giây, cho phép tái sử dụng kết quả vị trí tối đa 60 giây; trang bản đồ yêu cầu độ chính xác cao hơn, còn khu vực lịch chiếu ưu tiên tốc độ và khả năng đáp ứng [17].")
    paragraph(doc, "Sau khi nhận vị trí, utility `location.ts` kiểm tra miền tọa độ và áp dụng công thức Haversine [18] với hằng số bán kính 6.371 km. Danh sách rạp được sắp từ gần đến xa; nếu hai khoảng cách bằng nhau, tên rạp được dùng làm tiêu chí ổn định. Thành phố được ưu tiên theo rạp gần nhất và người dùng vẫn có thể đổi thành phố, tìm kiếm hoặc chọn rạp khác. Trên bản đồ, bán kính 10 km dùng để đếm và gắn nhãn rạp gần; nếu không có rạp nào trong bán kính này, hệ thống vẫn giữ danh sách đã sắp theo khoảng cách để người dùng có phương án tiếp tục, thay vì trả về màn hình rỗng.")
    paragraph(doc, "Các nhánh lỗi được xử lý theo hướng phục hồi: khi người dùng từ chối quyền, định vị hết thời gian chờ, không xác định được vị trí hoặc trình duyệt không hỗ trợ, giao diện hiển thị thông báo tiếng Việt và giữ nguyên khả năng chọn thành phố, rạp thủ công. Rạp có tọa độ thiếu hoặc ngoài miền hợp lệ không được dùng để tạo marker hay xếp hạng gần. Tọa độ chính xác của người dùng chỉ tồn tại trong state của component ở phiên hiển thị hiện tại; dữ liệu này không được ghi vào localStorage, không được lưu trong cơ sở dữ liệu và trong luồng client hiện tại cũng không được gửi đến endpoint `/api/v1/cinemas/nearest`. LocalStorage chỉ giữ tên thành phố ưu tiên cùng nguồn lựa chọn `manual` hoặc `location`.")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN ĐĂNG NHẬP]", "Chụp riêng trang đăng nhập, thể hiện tên đăng nhập, mật khẩu, Google Login và liên kết khôi phục tài khoản.", caption="Hình 4.1. Giao diện đăng nhập")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN ĐĂNG KÝ]", "Chụp riêng trang đăng ký, thể hiện các trường dữ liệu và thông báo validation.", caption="Hình 4.2. Giao diện đăng ký tài khoản")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN TRANG CHỦ]", "Thể hiện thanh điều hướng, danh sách phim và khu vực mua vé theo rạp.", caption="Hình 4.3. Giao diện trang chủ")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN CHI TIẾT PHIM]", "Thể hiện thông tin phim, trailer, ngôn ngữ, phụ đề và lịch chiếu theo khu vực.", caption="Hình 4.4. Giao diện chi tiết phim và lịch chiếu")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN DANH SÁCH VÀ BẢN ĐỒ RẠP]", "Chụp sau khi đã cho phép truy cập vị trí, thể hiện rõ nút Gần tôi, dấu vị trí người dùng, danh sách rạp được sắp theo khoảng cách, nhãn khoảng cách hoặc số rạp trong bán kính 10 km và marker rạp trên bản đồ. Chỉ dùng dữ liệu vị trí thử nghiệm, không để lộ địa chỉ sinh hoạt thực tế trong ảnh nộp báo cáo.", caption="Hình 4.5. Giao diện tra cứu rạp chiếu trên bản đồ và chức năng Gần tôi")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN CHI TIẾT RẠP]", "Thể hiện địa chỉ, thông tin rạp, ngày chiếu, phim và các giờ chiếu đang mở bán.", caption="Hình 4.6. Giao diện chi tiết rạp và lịch chiếu")

    doc.add_heading("4.3.2. Giao diện chọn ghế và thanh toán", level=3)
    paragraph(doc, "Hành trình mua vé được tổ chức theo trình tự chọn suất chiếu, chọn ghế, xác nhận đơn và thanh toán. SeatSelectionPage lấy snapshot sơ đồ ghế bằng REST API, sau đó đăng ký kênh WebSocket của suất chiếu để nhận những thay đổi mới. Trạng thái do server gửi về được hợp nhất vào dữ liệu đang hiển thị mà không tải lại toàn trang. Khi yêu cầu giữ ghế bị từ chối vì ghế đã thuộc người khác, giao diện cập nhật lại sơ đồ và hiển thị thông báo nghiệp vụ thay vì coi đây là lỗi hệ thống.")
    paragraph(doc, "CheckoutPage hiển thị thông tin booking do Back-end trả về, cho phép áp dụng khuyến mãi và lựa chọn VNPay hoặc quét QR ngân hàng. Với SePay, mã QR chứa số tiền và nội dung chuyển khoản cố định; nếu người dùng đổi mã giảm giá sau khi QR đã được tạo, giao diện yêu cầu tạo lại mã thanh toán để tránh sai lệch số tiền. PaymentResultPage chỉ hiển thị thành công sau khi trạng thái đã được Back-end xác nhận, không dựa vào thao tác bấm nút của người dùng.")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN CHỌN GHẾ]", "Thể hiện thông tin suất chiếu, màn chiếu, chú giải, các loại ghế và phần tóm tắt đặt vé.", caption="Hình 4.7. Giao diện lựa chọn ghế")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN XÁC NHẬN ĐƠN VÀ PHƯƠNG THỨC THANH TOÁN]", "Thể hiện thông tin đơn, mã giảm giá, tổng tiền và các phương thức thanh toán.", caption="Hình 4.8. Giao diện xác nhận đơn và chọn phương thức thanh toán")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN THANH TOÁN QR NGÂN HÀNG]", "Thể hiện mã QR, số tiền, tài khoản nhận và nội dung chuyển khoản.", caption="Hình 4.9. Giao diện thanh toán bằng mã QR ngân hàng")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN KẾT QUẢ GIAO DỊCH]", "Thể hiện trạng thái giao dịch, thông tin đơn và điều hướng sau thanh toán.", caption="Hình 4.10. Giao diện kết quả giao dịch")

    doc.add_heading("4.3.3. Giao diện vé và tài khoản khách hàng", level=3)
    paragraph(doc, "Sau thanh toán, người dùng theo dõi đơn và vé trong khu vực tài khoản. Các trạng thái đang chờ, thất bại, hết hạn, đã hủy, đang chờ xử lý yêu cầu hoàn tiền và đã được ghi nhận hoàn tiền được trình bày tách biệt với vé còn hiệu lực. Trang chi tiết vé tập trung mã QR cùng thông tin phim, rạp, phòng, ghế và thời gian chiếu; trang hồ sơ cho phép cập nhật thông tin cá nhân, ảnh đại diện và quản lý các phiên đăng nhập đang hoạt động.")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN ĐƠN ĐÃ ĐẶT]", "Thể hiện danh sách đơn theo trạng thái, thông tin phim, rạp, suất chiếu và thao tác phù hợp.", caption="Hình 4.11. Giao diện quản lý các đơn đã đặt")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN CHI TIẾT VÉ ĐIỆN TỬ]", "Thể hiện mã QR và đầy đủ thông tin phim, rạp, phòng, ghế, thời gian của vé.", caption="Hình 4.12. Giao diện chi tiết vé điện tử")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN HỒ SƠ VÀ PHIÊN ĐĂNG NHẬP]", "Thể hiện thông tin hồ sơ và khu vực theo dõi, thu hồi các phiên đăng nhập.", caption="Hình 4.13. Giao diện hồ sơ cá nhân và quản lý phiên")

    doc.add_heading("4.3.4. Giao diện quản trị hệ thống", level=3)
    paragraph(doc, "Khu vực quản trị được thiết kế cho các thao tác vận hành lặp lại. Danh sách sử dụng bộ lọc, phân trang và trạng thái rõ ràng; các biểu mẫu tạo hoặc cập nhật được đặt trong modal có vùng cuộn riêng để phù hợp với màn hình có chiều cao khác nhau. Thành phần giao diện chỉ hiển thị thao tác mà tài khoản có permission tương ứng, trong khi Back-end tiếp tục kiểm tra quyền tại API để tránh phụ thuộc vào việc ẩn nút ở phía client.")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN DASHBOARD QUẢN TRỊ]", "Thể hiện các chỉ số tổng quan, biểu đồ doanh thu và nhóm phim có doanh thu cao.", caption="Hình 4.14. Giao diện dashboard quản trị")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN QUẢN LÝ PHIM]", "Thể hiện danh sách phim, bộ lọc, trạng thái và thao tác thêm hoặc cập nhật phim.", caption="Hình 4.15. Giao diện quản lý phim")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN QUẢN LÝ RẠP CHIẾU]", "Thể hiện bộ lọc thành phố, danh sách rạp và các thao tác quản trị.", caption="Hình 4.16. Giao diện quản lý rạp chiếu")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN QUẢN LÝ PHÒNG VÀ GHẾ]", "Thể hiện việc chọn thành phố, rạp, phòng và cấu hình sơ đồ ghế.", caption="Hình 4.17. Giao diện quản lý phòng chiếu và ghế")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN QUẢN LÝ SUẤT CHIẾU]", "Thể hiện danh sách suất, thời gian, trạng thái và các thao tác tạo hoặc hủy suất.", caption="Hình 4.18. Giao diện quản lý suất chiếu")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN QUẢN LÝ NGƯỜI DÙNG VÀ NHÂN VIÊN]", "Thể hiện các tab khách hàng, nhân viên và thao tác phân công rạp.", caption="Hình 4.19. Giao diện quản lý người dùng và nhân viên")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN QUẢN LÝ ĐƠN ĐẶT VÉ]", "Thể hiện bộ lọc thành phố, rạp, trạng thái và danh sách đơn.", caption="Hình 4.20. Giao diện quản lý đơn đặt vé")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN QUẢN LÝ GIAO DỊCH THANH TOÁN]", "Thể hiện danh sách giao dịch, cổng thanh toán, số tiền, trạng thái và bộ lọc.", caption="Hình 4.21. Giao diện quản lý giao dịch thanh toán")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN QUẢN LÝ HOÀN TIỀN]", "Thể hiện các yêu cầu hoàn tiền, trạng thái xử lý và thao tác dành cho quản trị viên.", caption="Hình 4.22. Giao diện quản lý yêu cầu hoàn tiền")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN ĐỐI SOÁT THANH TOÁN]", "Thể hiện các bất nhất giữa booking, payment và ticket cần được kiểm tra.", caption="Hình 4.23. Giao diện đối soát thanh toán")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN NHẬT KÝ VẬN HÀNH]", "Thể hiện nhật ký callback, webhook, thay đổi trạng thái và các thông tin truy vết.", caption="Hình 4.24. Giao diện nhật ký vận hành")

    doc.add_heading("4.3.5. Giao diện vận hành dành cho nhân viên", level=3)
    paragraph(doc, "Giao diện nhân viên chỉ hiển thị các rạp thuộc phạm vi được quản trị viên phân công. Nhân viên có thể theo dõi phòng, ghế, suất chiếu, booking và thanh toán trong phạm vi được phép. Trên màn hình soát vé, danh sách suất được lọc tiếp theo rạp và chỉ gồm những suất đang mở check-in, từ 60 phút trước giờ bắt đầu đến 30 phút sau giờ bắt đầu theo cấu hình hiện tại. Camera và chức năng đọc tệp ảnh chỉ được kích hoạt sau khi người dùng đã chọn đủ rạp và suất chiếu.")
    paragraph(doc, "Kết quả quét được trình bày theo từng nguyên nhân để nhân viên có thể xử lý ngay tại cửa rạp. Vé hợp lệ hiển thị xác nhận thành công cùng thông tin khách, phim, phòng, ghế và thời điểm soát. Vé đã dùng cho biết trạng thái đã sử dụng; mã không hợp lệ, vé sai rạp, sai suất, chưa đến giờ hoặc quá thời gian check-in đều có thông báo riêng. Trong mọi nhánh bị từ chối, vé không bị chuyển sang USED. Việc kiểm tra tương tự cũng được áp dụng với tài khoản quản trị khi sử dụng chức năng soát vé, ngoại trừ phạm vi rạp không bị giới hạn bởi phân công nhân viên.")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN RẠP NHÂN VIÊN PHỤ TRÁCH]", "Thể hiện danh sách rạp được phân công và thông tin vận hành liên quan.", caption="Hình 4.25. Giao diện danh sách rạp nhân viên phụ trách")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN NHÂN VIÊN QUẢN LÝ SUẤT CHIẾU]", "Thể hiện danh sách và biểu mẫu suất chiếu trong phạm vi rạp được phân công.", caption="Hình 4.26. Giao diện nhân viên quản lý suất chiếu")
    placeholder(doc, "[CHÈN ẢNH GIAO DIỆN CHỌN NGỮ CẢNH SOÁT VÉ]", "Thể hiện danh sách rạp thuộc phạm vi phụ trách, các suất đang mở check-in và khu vực camera hoặc tải tệp sau khi đã chọn đủ ngữ cảnh.", caption="Hình 4.27. Giao diện lựa chọn rạp và suất chiếu để soát vé")
    placeholder(doc, "[CHÈN ẢNH TRẠNG THÁI KHÔNG CÓ SUẤT ĐANG MỞ CHECK-IN]", "Thể hiện thông báo khi rạp được chọn chưa có suất chiếu nằm trong cửa sổ check-in.", caption="Hình 4.28. Trạng thái không có suất chiếu đang mở check-in")
    placeholder(doc, "[CHÈN ẢNH KẾT QUẢ SOÁT VÉ THÀNH CÔNG]", "Thể hiện thông tin vé hợp lệ sau khi được xác nhận và chuyển sang trạng thái đã sử dụng.", caption="Hình 4.29. Kết quả soát vé hợp lệ")
    placeholder(doc, "[CHÈN ẢNH KẾT QUẢ VÉ ĐÃ ĐƯỢC SỬ DỤNG]", "Thể hiện cảnh báo vé đã được soát trước đó cùng thông tin thời điểm hoặc nhân viên thực hiện nếu có.", caption="Hình 4.30. Kết quả soát lại vé đã được sử dụng")
    placeholder(doc, "[CHÈN ẢNH KẾT QUẢ MÃ QR KHÔNG HỢP LỆ]", "Thể hiện phản hồi khi mã bị sửa, không đúng định dạng hoặc ảnh không thể đọc được.", caption="Hình 4.31. Kết quả xác thực mã QR không hợp lệ")
    placeholder(doc, "[CHÈN ẢNH KẾT QUẢ VÉ SAI RẠP HOẶC SAI SUẤT]", "Thể hiện thông báo theo đúng nguyên nhân khi vé không thuộc rạp hoặc suất chiếu đang được lựa chọn.", caption="Hình 4.32. Kết quả soát vé sai rạp hoặc sai suất chiếu")
    placeholder(doc, "[CHÈN ẢNH KẾT QUẢ NGOÀI CỬA SỔ CHECK-IN]", "Thể hiện thông báo khi quét quá sớm hoặc sau thời điểm hệ thống cho phép soát vé.", caption="Hình 4.33. Kết quả soát vé ngoài cửa sổ check-in")
    doc.add_heading("4.4. Kế hoạch triển khai hệ thống", level=2)
    paragraph(doc, "[KHÔNG GIAN TRỐNG - SINH VIÊN SẼ TỰ BỔ SUNG CHI TIẾT QUÁ TRÌNH DEPLOY HỆ THỐNG THỰC TẾ VÀO ĐÂY SAU]")
    doc.add_heading("4.5. Quản lý mã nguồn bằng Git và GitHub", level=2)
    paragraph(doc, "Mã nguồn của hệ thống được tổ chức thành hai kho Git độc lập, tương ứng với phần máy chủ (Back-end) và ứng dụng khách (Front-end), sau đó được lưu trữ trên GitHub. Việc phân tách này phù hợp với kiến trúc của hệ thống, đồng thời hỗ trợ quản lý phiên bản, theo dõi lịch sử thay đổi và phát triển từng thành phần tương đối độc lập. Tại thời điểm hoàn thiện báo cáo, nhánh main của mỗi kho lưu trữ phiên bản mã nguồn được sử dụng để đối chiếu với nội dung trình bày trong khóa luận.")
    paragraph(doc, "Thông tin các kho mã nguồn của đề tài:")
    hyperlink_paragraph(doc, "Kho mã nguồn Back-end", "https://github.com/Thandinh/cinema-booking-system")
    hyperlink_paragraph(doc, "Kho mã nguồn Front-end", "https://github.com/Thandinh/cinema-booking-client")
    paragraph(doc, "Các kho mã nguồn trên là minh chứng cho quá trình hiện thực hệ thống, cho phép đối chiếu cấu trúc chương trình, lịch sử thay đổi và các tệp kiểm thử với nội dung phân tích trong báo cáo. Để bảo đảm khả năng truy vết và tái lập kết quả, phiên bản mã nguồn sử dụng khi nghiệm thu cần được xác định bằng mã commit hoặc thẻ phiên bản tương ứng.")
    paragraph(doc, "Trong quá trình phát triển, các thay đổi được ghi nhận theo từng phạm vi cụ thể như bổ sung chức năng, sửa lỗi hoặc hoàn thiện mã nguồn. Trước khi được tích hợp vào nhánh main, các thay đổi cần được kiểm tra bằng những ca kiểm thử liên quan và quy trình biên dịch ở chế độ sản xuất. Phương thức quản lý này hỗ trợ xác định nguyên nhân khi phát sinh lỗi hồi quy, khôi phục phiên bản ổn định và duy trì tính nhất quán giữa hai thành phần của hệ thống [19].")

    doc.add_heading("4.6. Kiểm thử chức năng", level=2)
    doc.add_heading("4.6.1. Phương pháp và phạm vi kiểm thử", level=3)
    paragraph(doc, "Hoạt động kiểm thử được thực hiện ở cả Back-end và Front-end. Đối với Back-end, unit test kiểm tra các thành phần có đầu vào, đầu ra tương đối độc lập; integration test khởi tạo Spring context và PostgreSQL bằng Testcontainers để kiểm tra đồng thời Controller, Service, Repository, JPA mapping, Flyway migration và ràng buộc dữ liệu. Cách kiểm thử này phù hợp với các nghiệp vụ phụ thuộc vào transaction, khóa bản ghi và unique constraint, vốn khó đánh giá đầy đủ bằng cơ sở dữ liệu mô phỏng trong bộ nhớ.")
    paragraph(doc, "Đối với Front-end, Vitest kiểm tra validation, xử lý đăng nhập, component và các utility quan trọng. Production build được chạy sau test để phát hiện lỗi TypeScript, import không hợp lệ hoặc lỗi đóng gói mà unit test có thể không đi qua. Phạm vi của các nhóm test hiện có được tổng hợp tại Bảng 4.5.")
    table(doc, ["Nhóm", "Kiểm thử hiện có", "Mục tiêu"], [
        ["Khởi động ứng dụng", "CinemaBookingSystemApplicationTests", "Kiểm tra Spring context và cấu hình cơ bản của ứng dụng."],
        ["Xác thực và bảo mật", "AuthenticationServiceIntegrationTest; BookingPaymentSecurityIntegrationTest", "Đăng nhập, refresh token rotation, quản lý phiên, ownership và permission boundary."],
        ["Đặt vé và xử lý đồng thời", "BookingWorkflowIntegrationTest", "Giữ ghế, tạo booking PENDING, timeout, tranh chấp ghế và request gửi lặp."],
        ["Thanh toán", "PaymentCallbackIntegrationTest; SePayPaymentGatewayTest", "Callback/webhook, idempotency, đối chiếu số tiền và dữ liệu cổng thanh toán."],
        ["API và quản lý dữ liệu", "GlobalExceptionHandlerIntegrationTest; UserManagementIntegrationTest; HomeShowtimeFeedIntegrationTest", "Cấu trúc lỗi, quản lý người dùng và dữ liệu lịch chiếu công khai."],
        ["Vé điện tử", "TicketQrCodeServiceTest; QrCodeImageServiceTest", "Chữ ký QR, phát hiện dữ liệu bị sửa và tạo ảnh QR."],
        ["Giới hạn tần suất", "FixedWindowRateLimitServiceTest", "Số request được phép và thời gian thử lại sau khi vượt giới hạn."],
        ["Front-end", "Bảy tệp kiểm thử bằng Vitest; trong đó location.test.ts và RegionalShowtimeBrowser.test.tsx kiểm tra chức năng vị trí", "Validation, đăng nhập, carousel, lịch chiếu theo khu vực; Haversine, rạp gần nhất, sắp xếp tọa độ không hợp lệ về cuối và ưu tiên rạp/thành phố sau định vị."],
    ], caption="Bảng 4.5. Các nhóm kiểm thử hiện có trong mã nguồn")

    doc.add_heading("4.6.2. Kết quả kiểm thử Back-end", level=3)
    paragraph(doc, "Bộ kiểm thử Back-end được chạy bằng Maven trong môi trường có Docker. Testcontainers tạo một PostgreSQL container độc lập, áp dụng đầy đủ các Flyway migration trước khi thực hiện test và xóa môi trường sau khi hoàn tất. Kết quả Surefire ghi nhận 12 test suite với 53 test case; toàn bộ test case đều đạt, không có trường hợp thất bại, lỗi hoặc bị bỏ qua.")
    paragraph(doc, "Trong quá trình chạy có xuất hiện log cảnh báo về việc tái sử dụng refresh token và vi phạm unique constraint của booking hoặc payment đang chờ. Đây là dữ liệu lỗi được các test case chủ động tạo ra để kiểm tra cơ chế bảo mật, idempotency và xử lý request đồng thời. Các test tương ứng đã xác nhận hệ thống trả về kết quả nghiệp vụ mong đợi, vì vậy những log này không được tính là lỗi của lần kiểm thử.")

    doc.add_heading("4.6.3. Kết quả kiểm thử và biên dịch Front-end", level=3)
    paragraph(doc, "Vitest được chạy ở chế độ một lần và ghi nhận 20 test case thuộc 7 tệp đều thành công. Sau đó, lệnh production build thực hiện kiểm tra TypeScript và đóng gói ứng dụng bằng Vite. Quá trình build xử lý 2.666 module và hoàn tất mà không phát sinh lỗi. Kết quả tổng hợp của cả hai phần hệ thống được trình bày tại Bảng 4.6.")
    table(doc, ["Hạng mục", "Kết quả", "Thời điểm thực hiện"], [
        ["Back-end Maven/JUnit", "12 test suite; 53 test; 53 đạt; 0 thất bại; 0 lỗi; 0 bỏ qua", "17/08/2026"],
        ["Front-end Vitest", "7 tệp; 20 test; 20 đạt; 0 thất bại", "17/08/2026"],
        ["Front-end production build", "TypeScript và Vite build thành công; 2.666 module được xử lý", "17/08/2026"],
    ], caption="Bảng 4.6. Kết quả thực thi kiểm thử và biên dịch hệ thống")
    paragraph(doc, "Báo cáo chi tiết của Back-end được Maven lưu trong thư mục target/surefire-reports; kết quả Front-end được ghi nhận từ đầu ra của Vitest và Vite. Những số liệu này phản ánh đúng lần thực thi ngày 17/08/2026 và cho thấy các trường hợp đã được tự động hóa không phát hiện lỗi hồi quy. Kết quả không được dùng để suy rộng thành giới hạn tải hoặc mức độ an toàn tuyệt đối của hệ thống.")

    doc.add_heading("4.6.4. Các kịch bản kiểm thử nghiệp vụ tiêu biểu", level=3)
    paragraph(doc, "Bên cạnh việc thống kê theo lớp kiểm thử, các test case được đối chiếu với những luồng người dùng có ảnh hưởng trực tiếp đến dữ liệu. Bảng 4.7 trình bày một số kịch bản tiêu biểu, kết quả mong đợi và bằng chứng kiểm chứng hiện có. Những thao tác thuần giao diện như camera, khả năng hiển thị trên nhiều kích thước màn hình và hành vi điều hướng vẫn cần được quan sát thủ công trên trình duyệt hoặc thiết bị thật.")
    table(doc, ["Mã", "Tình huống", "Kết quả mong đợi", "Bằng chứng kiểm chứng"], [
        ["TC-01", "Đăng nhập sai mật khẩu hoặc tài khoản không hợp lệ", "Không phát token; API trả thông báo xác thực phù hợp.", "AuthenticationServiceIntegrationTest."],
        ["TC-02", "Refresh token cũ bị sử dụng lại sau rotation", "Phát hiện reuse và thu hồi phiên liên quan.", "AuthenticationServiceIntegrationTest."],
        ["TC-03", "Hai request gần như đồng thời giữ cùng một ghế", "Chỉ một request giữ ghế thành công; request còn lại nhận lỗi nghiệp vụ.", "BookingWorkflowIntegrationTest."],
        ["TC-04", "Tạo lặp booking hoặc payment PENDING", "Không tạo hai bản ghi cạnh tranh cho cùng phạm vi đã ràng buộc.", "BookingWorkflowIntegrationTest và BookingPaymentSecurityIntegrationTest."],
        ["TC-05", "Callback thanh toán hợp lệ được gửi nhiều lần", "Booking và ticket chỉ được xác nhận một lần; lần gọi sau đi qua nhánh idempotent.", "PaymentCallbackIntegrationTest."],
        ["TC-06", "QR vé bị sửa, đã sử dụng, sai rạp hoặc sai suất đang soát", "Không check-in và không làm thay đổi trạng thái ticket; kết quả phản ánh đúng nguyên nhân.", "TicketQrCodeServiceTest và BookingWorkflowIntegrationTest."],
        ["TC-07", "Tạo suất mới cách suất hiện có 10 phút trong cùng phòng", "Từ chối vì chưa đáp ứng khoảng dọn phòng 15 phút.", "BookingWorkflowIntegrationTest."],
        ["TC-08", "Tra cứu lịch công khai có suất quá gần giờ chiếu, đang chiếu, đã hủy hoặc ngoài bảy ngày", "Chỉ suất UPCOMING còn trong khoảng mở bán được trả về.", "HomeShowtimeFeedIntegrationTest."],
        ["TC-09", "Nhân viên tải danh sách suất để soát vé", "Chỉ suất thuộc rạp được phân công và nằm trong cửa sổ check-in được hiển thị.", "BookingWorkflowIntegrationTest."],
        ["TC-10", "Dữ liệu form đăng nhập, đăng ký hoặc tìm kiếm không hợp lệ", "Giao diện hiển thị lỗi và không gửi request sai cấu trúc.", "Vitest."],
        ["TC-11", "Quét QR bằng camera/tệp ảnh và kiểm tra responsive", "Luồng thao tác hoạt động ổn định trên trình duyệt và thiết bị mục tiêu.", "Kiểm thử thủ công; bổ sung ảnh hoặc biên bản khi nộp."],
        ["TC-12", "Cho phép quyền vị trí tại khu vực lịch chiếu", "Thành phố và rạp gần nhất được chọn; các thành phố, rạp được sắp từ gần đến xa và khoảng cách được hiển thị.", "RegionalShowtimeBrowser.test.tsx và location.test.ts."],
        ["TC-13", "Từ chối quyền, định vị hết thời gian chờ hoặc trình duyệt không hỗ trợ", "Hiển thị thông báo phù hợp; không mất danh sách và vẫn chọn được thành phố, rạp thủ công.", "Kiểm thử thủ công trên trình duyệt; thông báo lỗi được ánh xạ trong location.ts."],
        ["TC-14", "Vị trí hợp lệ nhưng không có rạp trong bán kính 10 km", "Hiển thị số rạp gần bằng 0 nhưng vẫn sắp và hiển thị các rạp gần nhất để người dùng tiếp tục.", "Kiểm thử component/thủ công trên CinemaMapPage."],
        ["TC-15", "Danh sách có rạp thiếu hoặc sai tọa độ", "Không tạo marker sai và không phát sinh lỗi toàn trang; rạp không hợp lệ nằm sau các rạp tính được khoảng cách.", "location.test.ts; kiểm tra hasValidCoordinates tại CinemaMapPage."],
        ["TC-16", "Kết thúc phiên tra cứu sau khi dùng Gần tôi", "Tọa độ chính xác không được ghi vào localStorage hoặc cơ sở dữ liệu; chỉ thành phố ưu tiên và nguồn lựa chọn được lưu.", "Rà soát cityPreferenceStore.ts, CinemaMapPage và RegionalShowtimeBrowser."],
    ], font_size=8.5, caption="Bảng 4.7. Các kịch bản kiểm thử nghiệp vụ tiêu biểu")

    doc.add_heading("4.7. Kiểm thử xử lý đồng thời bằng Apache JMeter", level=2)
    paragraph(doc, "Apache JMeter được sử dụng để tạo nhiều yêu cầu HTTP trong cùng một khoảng thời gian và quan sát cách hệ thống phản hồi khi các yêu cầu cùng cạnh tranh một tài nguyên [20]. Trong phạm vi đề tài, một phép thử có kiểm soát đã được thực hiện đối với nghiệp vụ giữ ghế. Mục đích của phép thử không phải xác định công suất tối đa của hệ thống, mà kiểm chứng quy tắc quan trọng: tại một thời điểm, một ghế của một suất chiếu chỉ được phép thuộc về một phiên giữ hợp lệ.")
    paragraph(doc, "Kết quả trong mục này là số liệu ghi nhận từ một lần chạy trên môi trường phát triển cục bộ ngày 17/08/2026. Do số lượng người dùng ảo nhỏ, thời gian chạy ngắn và JMeter được vận hành ở chế độ giao diện, các số liệu không được ngoại suy thành năng lực tải của môi trường triển khai thực tế.")

    doc.add_heading("4.7.1. Mục tiêu và nguyên tắc kiểm thử", level=3)
    paragraph(doc, "Kịch bản sử dụng tám tài khoản thử nghiệm độc lập. Mỗi luồng đăng nhập để nhận access token, sau đó cùng gửi yêu cầu giữ một ghế đang ở trạng thái AVAILABLE của cùng một suất chiếu. Dữ liệu gửi đến API giữ ghế gồm showtimeId và danh sách chứa cùng một seatId. Synchronizing Timer được đặt trước sampler giữ ghế để tám luồng chờ tại cùng một điểm và được giải phóng gần như đồng thời.")
    paragraph(doc, "Kịch bản được xem là đạt khi tám tài khoản đều đăng nhập thành công, chỉ một yêu cầu giữ ghế được chấp nhận và bảy yêu cầu còn lại bị từ chối bằng phản hồi nghiệp vụ nhất quán. Các phản hồi bị từ chối không được là lỗi xác thực, lỗi giới hạn tần suất hoặc lỗi máy chủ. Trước mỗi lần chạy, kết quả cũ được xóa và ghế đầu vào được kiểm tra còn khả dụng để tránh làm sai lệch phép thử.")

    doc.add_heading("4.7.2. Thiết lập kịch bản", level=3)
    table(doc, ["Thành phần", "Giá trị thiết lập", "Mục đích"], [
        ["Thread Group", "8 virtual user; Ramp-up 1 giây; Loop Count 1", "Tạo tám luồng độc lập và chỉ thực hiện một vòng kiểm thử."],
        ["Dữ liệu tài khoản", "8 tài khoản thử nghiệm khác nhau được đọc từ tệp CSV", "Bảo đảm mỗi yêu cầu giữ ghế đại diện cho một người dùng riêng."],
        ["Đăng nhập", "POST /auth/token; JSON Extractor lấy access token", "Tạo ngữ cảnh xác thực hợp lệ cho từng luồng."],
        ["Yêu cầu giữ ghế", "POST /api/v1/bookings/hold; cùng showtimeId và seatId", "Tạo tranh chấp có chủ đích trên cùng một tài nguyên ghế."],
        ["Synchronizing Timer", "Number of Simulated Users to Group by: 8", "Phát tám yêu cầu giữ ghế gần như đồng thời."],
        ["Tiêu chí chấp nhận", "1 phản hồi thành công; 7 phản hồi HTTP 409; không có HTTP 401, 429 hoặc 5xx", "Phân biệt từ chối đúng nghiệp vụ với lỗi của hệ thống."],
    ], font_size=8.8, caption="Bảng 4.8. Cấu hình kiểm thử đồng thời khi giữ cùng một ghế")
    placeholder(doc, "[CHÈN ẢNH CẤU HÌNH THREAD GROUP]", "Chèn ảnh thể hiện Number of Threads, Ramp-up Period và Loop Count của kịch bản tám người dùng.", caption="Hình 4.34. Cấu hình Thread Group cho kiểm thử giữ ghế đồng thời")
    placeholder(doc, "[CHÈN ẢNH CẤU HÌNH SYNCHRONIZING TIMER]", "Chèn ảnh thể hiện tám người dùng được gom nhóm trước khi gửi yêu cầu giữ ghế.", caption="Hình 4.35. Cấu hình Synchronizing Timer cho yêu cầu giữ ghế")

    doc.add_heading("4.7.3. Kết quả thực nghiệm", level=3)
    paragraph(doc, "Lần chạy tạo tổng cộng 16 mẫu, gồm tám yêu cầu đăng nhập và tám yêu cầu giữ ghế. Tất cả yêu cầu đăng nhập đều thành công. Đối với thao tác giữ ghế, một yêu cầu được chấp nhận và bảy yêu cầu nhận HTTP 409 do ghế đã được phiên khác giữ. Thời gian phản hồi trung bình của yêu cầu đăng nhập là 98 ms; thời gian phản hồi trung bình của yêu cầu giữ ghế là 121 ms trong điều kiện của lần thử này.")
    placeholder(doc, "[CHÈN ẢNH SUMMARY REPORT]", "Chèn ảnh Summary Report thể hiện 8 mẫu đăng nhập, 8 mẫu giữ ghế, thời gian phản hồi và tỷ lệ JMeter đánh dấu lỗi.", caption="Hình 4.36. Kết quả kiểm thử đồng thời trong Summary Report")
    table(doc, ["Nhóm yêu cầu", "Số mẫu", "Kết quả nghiệp vụ", "Trung bình", "Nhỏ nhất", "Lớn nhất", "Error % của JMeter"], [
        ["Đăng nhập", "8", "8 thành công", "98 ms", "83 ms", "181 ms", "0,00%"],
        ["Giữ cùng một ghế", "8", "1 thành công; 7 bị từ chối đúng nghiệp vụ", "121 ms", "119 ms", "125 ms", "87,50%"],
    ], font_size=8.5, caption="Bảng 4.9. Kết quả kiểm thử tám người dùng đồng thời giữ cùng một ghế")
    paragraph(doc, "Tỷ lệ lỗi 87,50% ở sampler giữ ghế là cách JMeter thống kê bảy phản hồi HTTP 409. Trong ngữ cảnh kiểm thử này, HTTP 409 là kết quả mong đợi vì nó biểu thị xung đột trạng thái tài nguyên, không phải sự cố xử lý. Vì vậy, tỷ lệ lỗi tổng hợp của JMeter không được diễn giải trực tiếp thành tỷ lệ lỗi chất lượng của hệ thống.")

    doc.add_heading("4.7.4. Đối chiếu nội dung phản hồi", level=3)
    paragraph(doc, "Yêu cầu được chấp nhận trả về HTTP 200 và nội dung xác nhận thao tác giữ ghế thành công. Response header được lưu cùng response body nhằm chứng minh cả trạng thái HTTP lẫn dữ liệu nghiệp vụ của kết quả này.")
    placeholder(doc, "[CHÈN ẢNH RESPONSE BODY CỦA YÊU CẦU GIỮ GHẾ THÀNH CÔNG]", "Chèn ảnh phần Response data của yêu cầu duy nhất được hệ thống chấp nhận.", caption="Hình 4.37. Nội dung phản hồi của yêu cầu giữ ghế thành công")
    placeholder(doc, "[CHÈN ẢNH RESPONSE HEADER CỦA YÊU CẦU GIỮ GHẾ THÀNH CÔNG]", "Chèn ảnh response header thể hiện trạng thái HTTP 200 của yêu cầu thành công.", caption="Hình 4.38. Header phản hồi của yêu cầu giữ ghế thành công")
    paragraph(doc, "Bảy yêu cầu còn lại trả về HTTP 409 cùng thông báo cho biết ghế không còn khả dụng do đã được một người dùng khác giữ. Phản hồi này cho phép client hướng dẫn khách hàng chọn ghế khác mà không làm lộ lỗi kỹ thuật nội bộ. Không ghi nhận HTTP 401, HTTP 429 hoặc HTTP 5xx trong nhóm yêu cầu của lần chạy.")
    placeholder(doc, "[CHÈN ẢNH RESPONSE BODY CỦA YÊU CẦU BỊ TỪ CHỐI]", "Chèn ảnh phần Response data của một yêu cầu nhận lỗi nghiệp vụ do ghế đã được giữ.", caption="Hình 4.39. Nội dung phản hồi khi ghế đã được người dùng khác giữ")
    placeholder(doc, "[CHÈN ẢNH RESPONSE HEADER HTTP 409]", "Chèn ảnh response header thể hiện HTTP 409 Conflict của yêu cầu bị từ chối.", caption="Hình 4.40. Header phản hồi HTTP 409 khi xảy ra tranh chấp ghế")
    paragraph(doc, "Kết quả cho thấy trong phạm vi tám yêu cầu đồng thời, hệ thống duy trì được điều kiện một ghế chỉ có một người giữ hợp lệ. Các yêu cầu đến sau được từ chối có kiểm soát thay vì ghi đè trạng thái hoặc phát sinh lỗi máy chủ. Kết quả này bổ sung bằng chứng thực nghiệm cho các kiểm thử tích hợp về transaction và khóa dữ liệu, nhưng chưa thay thế một đợt kiểm thử tải kéo dài.")

    doc.add_heading("4.7.5. Giới hạn của phép thử", level=3)
    paragraph(doc, "Phép thử chỉ đánh giá xung đột trên một ghế với tám người dùng trong một vòng chạy. Chỉ số throughput hiển thị trong Summary Report không được sử dụng để công bố công suất hệ thống vì thời lượng mẫu ngắn, JMeter chạy kèm các listener giao diện và tài nguyên máy thử chưa được chuẩn hóa. Các đợt đánh giá tiếp theo cần chạy JMeter ở chế độ không giao diện, tăng thời lượng và số người dùng, theo dõi percentile của thời gian phản hồi, connection pool, CPU, bộ nhớ và PostgreSQL, đồng thời tách riêng kịch bản tải API tra cứu với kịch bản callback hoặc webhook gửi lặp.")

    doc.add_heading("4.8. Đánh giá kiểm thử", level=2)
    paragraph(doc, "Kết quả 53 test Back-end và 20 test Front-end đều đạt cho thấy các luồng đã được tự động hóa hoạt động ổn định trên môi trường kiểm thử. Việc sử dụng PostgreSQL Testcontainers giúp kết quả của các test liên quan đến transaction, migration và constraint có giá trị thực tế hơn so với khi thay thế bằng một cơ sở dữ liệu trong bộ nhớ. Production build thành công cũng xác nhận phiên bản Front-end không còn lỗi TypeScript hoặc lỗi đóng gói tại thời điểm kiểm tra.")
    paragraph(doc, "Phép thử JMeter bổ sung góc nhìn thực nghiệm cho tình huống nhiều người cùng giữ một ghế: tám yêu cầu đồng thời chỉ tạo một kết quả giữ ghế thành công, còn bảy yêu cầu được từ chối bằng HTTP 409. Kết quả phù hợp với quy tắc nghiệp vụ và không xuất hiện lỗi 5xx. Sự kết hợp giữa kiểm thử tích hợp và phép thử ở tầng HTTP giúp kiểm tra cơ chế xử lý đồng thời từ hai cấp độ khác nhau.")
    paragraph(doc, "Mặc dù vậy, số lượng test đạt không đồng nghĩa mọi tình huống vận hành đều đã được bao phủ. Camera trên nhiều loại thiết bị, khả năng tiếp cận, độ ổn định mạng, lỗi thực tế từ ngân hàng và tải dài hạn vẫn cần được kiểm thử thêm trong môi trường gần với production. Đề tài chưa công bố ngưỡng throughput, TPS, percentile thời gian phản hồi hoặc số người dùng đồng thời tối đa vì phép thử JMeter hiện tại chỉ nhằm xác minh tính đúng đắn của tranh chấp ghế.")

    doc.add_heading("4.9. Tiểu kết chương 4", level=2)
    paragraph(doc, "Chương 4 đã trình bày cách hiện thực các thành phần chính của CinemaBooking trên Spring Boot, React và PostgreSQL. Các cơ chế xác thực, giữ ghế, thanh toán, phát hành vé, realtime, quản lý yêu cầu hoàn tiền, audit và quản trị được tổ chức thành các lớp có trách nhiệm riêng, đồng thời sử dụng transaction và ràng buộc cơ sở dữ liệu để bảo vệ tính nhất quán của các nghiệp vụ quan trọng.")
    paragraph(doc, "Bộ kiểm thử hiện tại đã chạy thành công trên cả Back-end và Front-end, trong đó các integration test sử dụng PostgreSQL thật thông qua Testcontainers. Phép thử JMeter với tám người dùng đồng thời cũng cho thấy chỉ một yêu cầu giữ cùng một ghế được chấp nhận. Các kết quả này cung cấp bằng chứng cho những trường hợp đã được kiểm tra, nhưng chưa thay thế cho đo tải dài hạn, kiểm thử bảo mật chuyên sâu và đánh giá trên môi trường triển khai chính thức. Những giới hạn đó là cơ sở để xác định các nội dung cần tiếp tục hoàn thiện ở Chương 5.")


def add_chapter_five(doc: Document) -> None:
    doc.add_heading("CHƯƠNG 5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
    doc.add_heading("5.1. Kết quả đạt được", level=2)
    paragraph(doc, "Kết quả chính của đề tài là một ứng dụng web có thể vận hành trọn vẹn quy trình từ công bố lịch chiếu đến soát vé tại rạp. Phim, rạp, phòng, sơ đồ ghế và suất chiếu được quản lý trên cùng một cơ sở dữ liệu. Dữ liệu này được sử dụng trực tiếp khi khách hàng tìm lịch, khi nhân viên kiểm tra vé và khi quản trị viên theo dõi hoạt động bán vé, vì vậy các màn hình không phải duy trì những bản sao nghiệp vụ tách rời.")
    paragraph(doc, "Ở hành trình mua vé, người dùng có thể chọn khu vực, phim, rạp, ngày chiếu và ghế ngồi trước khi tạo đơn. Ghế đã chọn chỉ được giữ đến một thời điểm cụ thể; hết thời hạn, máy chủ thu hồi ghế và gửi trạng thái mới đến những trình duyệt đang xem cùng suất chiếu. Mã khuyến mãi được tính tại phía máy chủ. Hai lựa chọn thanh toán đang hoạt động là VNPay và chuyển khoản qua mã QR được SePay xác nhận. Khi số tiền và kết quả giao dịch hợp lệ, đơn được chốt, ghế chuyển sang trạng thái đã bán, vé QR được tạo và thư thông báo được gửi cho khách hàng.")
    paragraph(doc, "Khu vực vận hành phản ánh sự khác biệt giữa quyền quản trị toàn hệ thống và trách nhiệm tại từng rạp. Nhân viên nhìn thấy các rạp được phân công, chỉ thao tác trên dữ liệu thuộc phạm vi đó và phải chọn đúng rạp, đúng suất trước khi soát vé. Quản trị viên quản lý dữ liệu nền, tài khoản, khuyến mãi, giao dịch, yêu cầu hoàn tiền và nhật ký. Đối với yêu cầu hoàn tiền, quản trị viên chỉ ghi nhận kết quả đã được xử lý theo quy trình bên ngoài; hệ thống chưa tự động chuyển tiền qua API nhà cung cấp. Cách kiểm soát này được thực hiện ở dịch vụ phía máy chủ, do đó việc thay đổi tham số trên trình duyệt không đủ để mở rộng phạm vi truy cập.")
    paragraph(doc, "Về kỹ thuật, ReactJS đảm nhiệm giao diện; Spring Boot xử lý bảo mật và quy tắc nghiệp vụ; PostgreSQL bảo vệ quan hệ dữ liệu. Tranh chấp ghế được xử lý bằng giao dịch, khóa bản ghi, trường phiên bản và ràng buộc duy nhất. Phiên đăng nhập sử dụng access token và refresh token có cơ chế thay thế, thu hồi. Flyway quản lý thay đổi lược đồ, còn các kiểm thử tích hợp tập trung vào xác thực, đặt vé, thanh toán, ngoại lệ và vé QR. Phép thử JMeter với tám người dùng đồng thời đã bổ sung bằng chứng rằng chỉ một yêu cầu giữ cùng một ghế được chấp nhận trong điều kiện thử nghiệm. Báo cáo không xem kết quả này là phép đo công suất hoặc bằng chứng cho khả năng triển khai quy mô lớn.")
    doc.add_heading("5.2. Hạn chế", level=2)
    paragraph(doc, "Phiên bản hiện tại phù hợp hơn với mô hình một máy chủ ứng dụng. Bộ môi giới STOMP và bộ nhớ đệm Caffeine đều nằm trong tiến trình đang chạy. Nếu khởi động nhiều máy chủ, một sự kiện ghế hoặc một lần xóa bộ nhớ đệm tại máy này chưa tự động được chuyển đến các máy còn lại. Vì vậy, khả năng mở rộng theo chiều ngang mới được phân tích ở mức thiết kế, chưa được kiểm chứng bằng triển khai phân tán.")
    paragraph(doc, "Cơ chế giới hạn tần suất đối với đăng nhập và giữ ghế hiện sử dụng bộ đếm cửa sổ thời gian cố định trong bộ nhớ của từng tiến trình. Cơ chế này đáp ứng mục tiêu hạn chế thao tác lặp trên một máy chủ nhưng chưa đồng bộ hạn mức khi ứng dụng chạy trên nhiều máy. Tương tự, hệ thống đã có PaymentEvent, báo cáo giám sát và chức năng đối soát ở cấp ứng dụng, song chưa có hàng đợi chuyên biệt để tự động thử lại theo khoảng chờ tăng dần và chưa tích hợp nền tảng thu thập log, metric, cảnh báo tập trung.")
    paragraph(doc, "Phạm vi tích hợp thanh toán vẫn còn giới hạn bởi môi trường của nhà cung cấp. MoMo có lớp kết nối nhưng đang tắt trong cấu hình mặc định. Quy trình quản lý yêu cầu hoàn tiền đã lưu được yêu cầu, trạng thái và sự kiện xử lý, song chưa gọi API hoàn tiền chính thức của từng cổng. Trạng thái thành công hoặc thất bại hiện được quản trị viên ghi nhận sau khi xử lý bên ngoài hệ thống. Việc tự động hóa phần này cần tài khoản đối tác, quy định đối soát và kịch bản thử nghiệm riêng của nhà cung cấp.")
    paragraph(doc, "Phép thử đồng thời bằng JMeter mới tập trung vào tính đúng đắn khi tám người dùng giữ cùng một ghế, chưa phải một đợt đánh giá công suất trên môi trường gần với triển khai thật. Đề tài chưa công bố ngưỡng thông lượng, thời gian phản hồi ở các phân vị, số người dùng đồng thời tối đa hoặc kết quả phục hồi sau sự cố. Quy trình tích hợp liên tục, quản lý bí mật, sao lưu và giám sát tập trung cũng chưa có minh chứng vận hành thực tế.")
    doc.add_heading("5.3. Hướng phát triển", level=2)
    paragraph(doc, "Ưu tiên gần nhất là hoàn thiện một môi trường triển khai có tên miền, TLS, cơ chế lưu bí mật ngoài mã nguồn và lịch sao lưu PostgreSQL. Log, số liệu sử dụng tài nguyên và lỗi thanh toán cần được tập trung để có thể phát hiện sự cố thay vì chỉ kiểm tra thủ công trên từng máy chủ. Tệp cấu hình, kết quả Maven, Vitest và JMeter cần tiếp tục được lưu theo từng phiên bản phát hành để duy trì bằng chứng kiểm thử có thể đối chiếu.")
    paragraph(doc, "Khi nhu cầu vượt khả năng của một máy chủ, kênh sự kiện ghế có thể chuyển sang RabbitMQ hoặc ActiveMQ; Redis có thể hỗ trợ đồng bộ bộ nhớ đệm và giới hạn tần suất giữa nhiều tiến trình. Các tác vụ đối soát hoặc xử lý lại sự kiện thanh toán có thể được đưa vào hàng đợi với chính sách thử lại và khoảng chờ tăng dần; log, metric và cảnh báo có thể được tập trung trên một nền tảng quan sát chuyên dụng. Đây là các hướng phát triển chưa được hiện thực trong phiên bản hiện tại, khác với bộ giới hạn tần suất cục bộ, PaymentEvent, báo cáo giám sát và chức năng đối soát đã có trong hệ thống. Việc triển khai chỉ nên được thực hiện sau khi kết quả đo tải và yêu cầu vận hành chứng minh sự cần thiết.")
    paragraph(doc, "Về nghiệp vụ, hệ thống có thể tích hợp API hoàn tiền trực tiếp của VNPay và các nhà cung cấp tiếp theo, bổ sung bước phê duyệt và báo cáo đối soát để thay thế việc ghi nhận kết quả thủ công hiện nay. Về chất lượng, cần mở rộng kiểm thử hợp đồng API, kiểm thử cạnh tranh ghế, kiểm thử xuyên suốt trên trình duyệt và thiết bị di động, đồng thời thực hiện rà soát bảo mật và khả năng tiếp cận trước khi đưa sản phẩm đến người dùng thực tế.")


def add_formal_conclusion(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("KẾT LUẬN", level=1)
    paragraph(doc, "CinemaBooking đã hoàn thành mục tiêu xây dựng một quy trình đặt vé có thể theo dõi từ lúc khách hàng chọn suất đến lúc vé được kiểm tra tại rạp. Phần khó nhất của đề tài không nằm ở việc tạo một bản ghi đặt vé, mà ở chỗ giữ trạng thái ghế nhất quán trong thời gian chờ thanh toán và xử lý kết quả trả về từ hệ thống bên ngoài. Việc đặt các kiểm tra quyết định tại máy chủ và cơ sở dữ liệu giúp giao diện phản ánh nghiệp vụ thay vì tự quyết định trạng thái bán vé.")
    paragraph(doc, "Đề tài cũng cho thấy một hệ thống web hoàn chỉnh cần kết hợp nhiều lớp kiến thức: ReactJS tổ chức trải nghiệm sử dụng, Spring Boot điều phối quy tắc và bảo mật, PostgreSQL duy trì tính toàn vẹn, còn WebSocket nối thay đổi đã xác nhận với các trình duyệt đang hoạt động. Phân quyền theo vai trò được bổ sung bằng phạm vi rạp của nhân viên; kết quả thanh toán được lưu cùng sự kiện để hỗ trợ truy tìm; vé QR chỉ được chuyển sang đã sử dụng sau khi qua đủ điều kiện kiểm tra.")
    paragraph(doc, "Sản phẩm vẫn còn khoảng cách so với một dịch vụ thương mại triển khai ở quy mô lớn, đặc biệt ở hạ tầng nhiều máy chủ, kiểm thử tải, giám sát và hoàn tiền tự động. Việc nhận diện rõ các giới hạn này giúp xác định hướng phát triển có thứ tự ưu tiên. Đồng thời, quá trình đối chiếu giao diện, dịch vụ, truy vấn và migration đã giúp sinh viên củng cố năng lực phân tích hệ thống, xử lý giao dịch và kiểm chứng một nhận định kỹ thuật bằng mã nguồn cụ thể.")


def add_formal_appendices(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("PHỤ LỤC A. MA TRẬN TRUY VẾT YÊU CẦU", level=1)
    paragraph(doc, "Bảng PL.A.1 chỉ ra nơi mỗi yêu cầu quan trọng được hiện thực ở giao diện và phía máy chủ. Các đường dẫn tiêu biểu giúp đối chiếu nội dung của Chương 3 và Chương 4 với chương trình, nhưng không thay thế tài liệu kiểm kê toàn bộ lớp và phương thức của dự án.")
    table(doc, ["Mã", "Yêu cầu", "Giao diện/API", "Back-end và dữ liệu"], [
        ["FR-01", "Đăng ký và xác thực email", "RegisterPage; /api/v1/auth/register; verify-email", "AuthenticationService; EmailService; users và token xác thực"],
        ["FR-02", "Đăng nhập và làm mới phiên", "LoginPage; authApi; axiosClient", "AuthenticationController; AuthenticationService; refresh_tokens; invalidated_token"],
        ["FR-03", "Tra cứu phim, rạp và lịch chiếu còn mở bán", "HomePage; MovieDetailPage; CinemaMapPage", "Movie/Cinema/Showtime Controller-Service-Repository; cửa sổ lịch công khai; cache dữ liệu ít đổi"],
        ["FR-04", "Giữ ghế có thời hạn", "SeatSelectionPage; bookingApi.holdSeats", "BookingServiceImpl; SeatStatusRepository; seat_status; lock và version"],
        ["FR-05", "Tạo đơn đặt vé", "CheckoutPage; POST /api/v1/bookings", "BookingServiceImpl; bookings; booking_details; partial unique index"],
        ["FR-06", "Áp dụng khuyến mãi", "CheckoutPage; promotionApi", "PromotionService; promotion_id; total_price và discount_amount"],
        ["FR-07", "Thanh toán VNPay", "CheckoutPage; PaymentResultPage", "PaymentController; PaymentServiceImpl; VnPayPaymentGateway; payments/events"],
        ["FR-08", "Thanh toán QR ngân hàng", "SePay QR form và polling kết quả", "SePayPaymentGateway; webhook; xác minh API key/HMAC; amount/content"],
        ["FR-09", "Phát hành và xem vé QR", "MyBookingsPage; TicketDetailPage", "TicketQrCodeService; tickets; BookingService payment-success"],
        ["FR-10", "Soát vé theo rạp và suất", "StaffTicketScannerPage", "TicketController; BookingServiceImpl.checkInTicket; staff_cinemas"],
        ["FR-11", "Đồng bộ trạng thái ghế", "useSeatWebSocket", "WebSocketConfig; SeatStatusPublisher; SeatStatusEvent"],
        ["FR-12", "Thu hồi giữ ghế hết hạn", "Client nhận AVAILABLE qua topic", "HoldExpireScheduler; BookingExpirationService; seat_status/bookings"],
        ["FR-13", "Hủy suất và quản lý yêu cầu hoàn tiền", "Admin showtime/payment pages", "ShowtimeService; RefundServiceImpl; refunds; payment_events"],
        ["FR-14", "Quản trị và phân quyền", "ProtectedRoute; admin/staff layout", "SecurityConfig; @PreAuthorize; ApplicationInitConfig; RBAC tables"],
        ["FR-15", "Theo dõi và kiểm toán", "Admin audit/payment event pages", "AdminAuditLogInterceptor; AuthAuditService; audit/payment event tables"],
        ["FR-16", "Quản lý vòng đời và xung đột suất chiếu", "AdminShowtimePage; lịch công khai và bộ chọn suất check-in", "ShowtimeServiceImpl; ShowtimeRepository; ShowtimeStatusSyncScheduler; khoảng dọn phòng và cửa sổ thời gian"],
        ["FR-17", "Tìm rạp gần vị trí hiện tại", "CinemaMapPage; RegionalShowtimeBrowser; cinemaApi.getMapData; utils/location", "GET /api/v1/cinemas/map; CinemaMapResponse; cinemas.latitude/longitude; endpoint /nearest sẵn có nhưng không được luồng client hiện tại sử dụng"],
    ], caption="Bảng PL.A.1. Ma trận truy vết yêu cầu và thành phần hiện thực")

    doc.add_page_break()
    doc.add_heading("PHỤ LỤC B. KỊCH BẢN KIỂM THỬ NGHIỆP VỤ", level=1)
    paragraph(doc, "Các kịch bản sau kết hợp kiểm thử chức năng, bảo mật và xử lý đồng thời. Những trường hợp chưa có kết quả đo được ghi dưới dạng kết quả kỳ vọng; sinh viên bổ sung minh chứng thực thi trước khi nộp bản in cuối cùng.")
    table(doc, ["Mã", "Tiền điều kiện và thao tác", "Kết quả mong đợi", "Loại"], [
        ["TC-01", "Đăng ký bằng username/email chưa tồn tại và password hợp lệ; mở liên kết email.", "Tài khoản được tạo, email được xác minh và có thể đăng nhập.", "Chức năng"],
        ["TC-02", "Đăng nhập sai nhiều lần trong cửa sổ rate limit.", "Trả thông báo tiếng Việt phù hợp; chặn tạm thời khi vượt ngưỡng; ghi auth audit.", "Bảo mật"],
        ["TC-03", "Dùng refresh token hợp lệ hai lần liên tiếp.", "Lần đầu rotation thành công; token cũ bị revoke và không được tái sử dụng.", "Bảo mật"],
        ["TC-04", "Hai user đồng thời giữ cùng showtimeId và seatId.", "Tối đa một user giữ ghế; user còn lại nhận lỗi ghế vừa được người khác chọn.", "Đồng thời"],
        ["TC-05", "User rời trang đến khi holdUntil hết hạn.", "Scheduler trả ghế AVAILABLE; client đang subscribe nhận sự kiện realtime.", "Lifecycle"],
        ["TC-06", "Tạo hai booking PENDING cho cùng user, suất và bộ ghế.", "Transaction/index từ chối trạng thái pending trùng hoặc service tái sử dụng luồng hợp lệ.", "Toàn vẹn"],
        ["TC-07", "VNPay callback có checksum sai hoặc amount khác payment.", "Không cập nhật SUCCESS, không tạo ticket; payment event ghi kết quả từ chối.", "Thanh toán"],
        ["TC-08", "Gửi callback thành công hợp lệ lặp lại.", "Xử lý idempotent; không tạo ticket hoặc email nghiệp vụ trùng không kiểm soát.", "Thanh toán"],
        ["TC-09", "SePay webhook đúng nội dung nhưng thiếu/sai API key hoặc HMAC.", "Webhook bị từ chối; booking và payment không đổi trạng thái.", "Thanh toán"],
        ["TC-10", "Áp dụng mã giảm giá sau khi đã tạo QR cố định số tiền.", "UI yêu cầu hủy/tạo lại mã thanh toán; không chấp nhận số tiền không khớp.", "UX/Thanh toán"],
        ["TC-11", "Staff quét vé ACTIVE nhưng chọn sai rạp hoặc sai suất.", "Không set USED; trả thông báo đúng nguyên nhân và giữ vé có thể sử dụng tại đúng điểm.", "Phân quyền"],
        ["TC-12", "Quét lại vé đã USED.", "Từ chối và hiển thị thời điểm/người soát nếu response cho phép.", "Nghiệp vụ"],
        ["TC-13", "Staff truy cập booking/payment của rạp không được phân công.", "Bị từ chối ở service scope dù có permission chức năng.", "Phân quyền"],
        ["TC-14", "Hủy suất đã có booking thanh toán thành công.", "Suất không còn mở bán; booking chuyển sang REFUND_PENDING, hệ thống tạo yêu cầu và audit/event để quản trị viên tiếp tục xử lý.", "Vận hành"],
        ["TC-15", "Tạo hoặc cập nhật suất trong cùng phòng khi khoảng cách với suất hiện có nhỏ hơn 15 phút.", "Hệ thống từ chối và không tạo dữ liệu lịch bị xung đột.", "Toàn vẹn"],
        ["TC-16", "Tra cứu lịch công khai với suất quá gần giờ bắt đầu, đang chiếu, đã kết thúc, đã hủy hoặc ngoài bảy ngày.", "Chỉ suất UPCOMING còn trong khoảng mở bán được trả về cho khách hàng.", "Lifecycle"],
        ["TC-17", "Nhân viên chọn một rạp được phân công để lấy danh sách suất đang mở soát vé.", "Chỉ suất bắt đầu trong khoảng từ 30 phút trước đến 60 phút sau thời điểm hiện tại được trả về; rạp ngoài phạm vi bị từ chối.", "Phân quyền"],
        ["TC-18", "Tám virtual user dùng tám tài khoản và cùng giữ một ghế qua Synchronizing Timer.", "Một request thành công; bảy request nhận HTTP 409 đúng nghiệp vụ; trung bình của sampler giữ ghế là 121 ms trong lần chạy.", "JMeter thực nghiệm"],
        ["TC-19", "Cấp quyền vị trí và đặt tọa độ thử nghiệm gần một rạp có dữ liệu hợp lệ.", "Client chọn rạp gần nhất, sắp xếp thành phố/rạp theo Haversine và hiển thị khoảng cách.", "Vitest: location.test.ts; RegionalShowtimeBrowser.test.tsx"],
        ["TC-20", "Từ chối quyền vị trí, mô phỏng lỗi timeout hoặc môi trường không có navigator.geolocation.", "Thông báo hướng dẫn chọn thủ công; danh sách rạp và các bộ chọn vẫn sử dụng được.", "Kiểm thử giao diện thủ công; đối chiếu location.ts"],
        ["TC-21", "Dùng vị trí cách mọi rạp hơn 10 km.", "Thông báo không có rạp trong bán kính 10 km nhưng vẫn hiển thị các rạp gần nhất theo thứ tự khoảng cách.", "Kiểm thử component/thủ công"],
        ["TC-22", "Dữ liệu bản đồ có tọa độ NaN, thiếu hoặc ngoài miền vĩ độ/kinh độ.", "Bỏ qua marker không hợp lệ; phép tính trả vô cực và đưa mục này về cuối thay vì làm hỏng màn hình.", "Vitest: location.test.ts; rà soát CinemaMapPage"],
        ["TC-23", "Dùng Gần tôi rồi tải lại hoặc đóng trang.", "Không có latitude/longitude chính xác trong localStorage hay dữ liệu gửi về CinemaBooking; chỉ lưu thành phố ưu tiên và nguồn location.", "Rà soát source và Application Storage của trình duyệt"],
    ], caption="Bảng PL.B.1. Kịch bản kiểm thử trọng yếu")

    doc.add_page_break()
    doc.add_heading("PHỤ LỤC C. TỪ ĐIỂN DỮ LIỆU", level=1)
    paragraph(doc, "Phụ lục này tổng hợp vai trò và các ràng buộc chính của những nhóm bảng trong cơ sở dữ liệu. Sơ đồ ERD tại Hình 3.8 được sử dụng làm mô hình quan hệ thống nhất; báo cáo không bổ sung một ERD chi tiết thứ hai.")
    doc.add_heading("C.1. Từ điển dữ liệu theo nhóm bảng", level=2)
    table(doc, ["Nhóm bảng", "Các bảng", "Vai trò và ràng buộc chính"], [
        ["Nhận dạng và phân quyền", "roles, permissions, users, users_roles, roles_permissions, staff_cinemas", "Quản lý tài khoản, RBAC và phạm vi rạp; tên đăng nhập, email, tên vai trò và tên quyền có ràng buộc duy nhất."],
        ["Danh mục", "movies, cinemas, rooms, seats, showtimes, promotions", "Lưu dữ liệu phim và cơ sở vật chất; cinemas có latitude/longitude cho bản đồ và tìm gần; ghế duy nhất trong phòng; thời gian kết thúc suất phải sau thời gian bắt đầu; trạng thái được giới hạn bằng CHECK."],
        ["Đặt vé", "seat_status, bookings, booking_details, tickets", "Một trạng thái cho mỗi ghế–suất; HOLD phải có người giữ và thời hạn; một đơn PENDING cho mỗi khách–suất; một vé cho mỗi chi tiết đặt vé."],
        ["Thanh toán", "payments, payment_events, refunds", "Một giao dịch PENDING trên mỗi đơn; lưu lịch sử callback/webhook; tối đa một yêu cầu hoàn tiền đang hoạt động cho mỗi giao dịch."],
        ["Bảo mật và nhật ký", "refresh_tokens, invalidated_token, auth_audit_logs, admin_audit_logs", "Lưu hash refresh token, JWT bị thu hồi và dấu vết xác thực hoặc thao tác quản trị; có chỉ mục theo người dùng, loại sự kiện và thời gian."],
    ], font_size=8.5, caption="Bảng PL.C.1. Từ điển dữ liệu theo nhóm bảng")


def add_advanced_appendices_d_to_g(doc: Document) -> None:
    """Add the source-traceable technical appendices selected for the final thesis."""
    doc.add_page_break()
    doc.add_heading("PHỤ LỤC D. GIỮ GHẾ, BOOKING VÀ RACE CONDITION", level=1)
    doc.add_heading("D.1. Mô hình dữ liệu trạng thái ghế", level=2)
    paragraph(doc, "Ghế vật lý nằm ở `seats`, còn ghế của từng suất nằm ở `seat_status`. Một seat có nhiều SeatStatus theo nhiều showtime; unique constraint `(seat_id, showtime_id)` bảo đảm một ghế không có hai state map cho cùng suất. SeatStatus lưu `status`, `hold_by`, `hold_until` và `version`. CHECK constraint ở migration buộc HOLD phải đi kèm thông tin hold phù hợp, tránh dữ liệu nửa vời. Vì vậy availability không được suy ra từ BookingDetail một cách chậm và dễ sai, mà được lấy trực tiếp từ seat_status.")
    doc.add_heading("D.2. Quy trình giữ ghế (hold seat)", level=2)
    bullets(doc, [
        "1. SeatSelectionPage tải initial map bằng GET `/api/v1/showtimes/{id}/seats` và subscribe WebSocket topic cùng showtime.",
        "2. Khi khách chọn ghế, frontend gửi `HoldSeatRequest` có showtimeId và danh sách seatIds đến POST `/api/v1/bookings/hold`.",
        "3. BookingServiceImpl kiểm tra rate limit `SeatHoldRateLimitService` theo user/IP để giảm spam thao tác giữ ghế.",
        "4. Service kiểm tra suất và thực hiện cleanup các expiry liên quan theo logic hiện hành; repository lấy/khóa SeatStatus mục tiêu trong transaction.",
        "5. Chỉ những row AVAILABLE hợp lệ được đổi thành HOLD, gán holdBy là user hiện tại và holdUntil theo `booking.seat-hold-minutes`.",
        "6. Transaction commit thành công mới đăng ký publish event HOLD qua SeatStatusPublisher. Nếu transaction rollback, không có broadcast trạng thái giả.",
        "7. Response trả HoldSeatResponse với holdUntil để frontend hiển thị countdown; UI không phải nguồn quyết định cuối cùng, backend mới là authority.",
    ])
    doc.add_heading("D.3. Tại sao hai khách không mua cùng ghế", level=2)
    paragraph(doc, "Giả sử A và B gần như đồng thời giữ A1 cùng showtime. Cả hai request đi vào transaction khác nhau. Repository lock/atomic condition khiến chỉ transaction lấy quyền thay đổi state AVAILABLE trước mới có thể update HOLD. Transaction sau đọc được row đã HOLD hoặc update count không đủ, service trả lỗi nghiệp vụ 'one or more seats are not available'. Không được biến lỗi đó thành 500 hoặc retry mù quáng trên frontend vì seat đã thuộc phiên giữ của khách khác. `@Version` của SeatStatus là lớp phát hiện conflict khi ORM update cạnh tranh; unique seat-showtime là ràng buộc nền dữ liệu.")
    paragraph(doc, "Sau hold, booking tạo `PENDING` và payment expiry. V14 thêm unique index pending theo user/showtime và pending payment theo booking để giảm double click/refresh tạo bản ghi pending song song. PaymentService còn khóa/kiểm tra booking/payment trước khi khởi tạo gateway. Điều này không có nghĩa mọi request trùng sẽ trả cùng response; tùy thời điểm có thể nhận response tái sử dụng payment hoặc business error. Ý nghĩa đúng là database không cho trạng thái nguy hiểm như hai payment pending cạnh tranh không kiểm soát tồn tại cho cùng đơn/gateway.")
    doc.add_heading("D.4. Hết hạn và nhả ghế", level=2)
    paragraph(doc, "`HoldExpireScheduler` chạy fixed delay theo `BOOKING_EXPIRED_HOLD_SCAN_DELAY_MS`, lấy batch projection expired, update theo id, rồi nhóm seat theo showtime để publish AVAILABLE sau commit. `PendingBookingExpireScheduler` quét Booking PENDING theo payment_expires_at/legacy cutoff, gọi `BookingService.expirePendingBooking` từng booking để một record lỗi không làm block cả batch. Khi payment thất bại, user hủy hoặc booking hết hạn, business service trả ghế về AVAILABLE nếu state/owner phù hợp. Đây là lý do refresh UI không phải điều kiện cần để ghế được giải phóng.")
    doc.add_page_break()
    doc.add_heading("D.5. Dữ liệu mẫu để giải thích khi bảo vệ", level=2)
    paragraph(doc, "Ví dụ dưới đây minh họa đúng cấu trúc của `HoldSeatRequest` và `HoldSeatResponse`. Các UUID chỉ là ký hiệu thay thế; khi chạy thực tế, chúng được lấy từ dữ liệu của suất chiếu và ghế tương ứng.")
    code_block(doc, "POST /api/v1/bookings/hold\nAuthorization: Bearer <access-token>\nContent-Type: application/json\n\n{\n  \"showtimeId\": \"<showtime-uuid>\",\n  \"seatIds\": [\"<seat-uuid-A1>\", \"<seat-uuid-A2>\"]\n}", "Mẫu request giữ ghế", font_size=9.5)
    paragraph(doc, "Khi xử lý thành công, endpoint trả `ApiResponse<HoldSeatResponse>`. Phần `result` chứa danh sách ghế đã giữ, thời điểm hết hạn, số tiền tạm tính và thông báo để giao diện hiển thị bộ đếm thời gian.")
    code_block(doc, "{\n  \"showtimeId\": \"<showtime-uuid>\",\n  \"heldSeatIds\": [\"<seat-uuid-A1>\", \"<seat-uuid-A2>\"],\n  \"holdUntil\": \"2026-08-18T20:15:00\",\n  \"estimatedTotalPrice\": 120000.00,\n  \"message\": \"Ghế đã được giữ trong <n> phút\"\n}", "Mẫu result của HoldSeatResponse", font_size=9.5)

    doc.add_page_break()
    doc.add_heading("PHỤ LỤC E. PAYMENT, CALLBACK, WEBHOOK VÀ QUẢN LÝ YÊU CẦU REFUND", level=1)
    doc.add_heading("E.1. Trừu tượng hóa luồng thanh toán", level=2)
    paragraph(doc, "`PaymentGateway` là abstraction của provider; `VnPayPaymentGateway`, `SePayPaymentGateway` và `MomoPaymentGateway` là adapter. PaymentServiceImpl chọn gateway theo PaymentMethod nhưng không để controller phải biết chi tiết checksum/QR/HMAC của từng nhà cung cấp. Tách adapter giúp thêm gateway mới mà không trộn string ký/hash vào code booking. Method/status/event types được enum và migration CHECK constraint để tránh giá trị tự do trong database.")
    doc.add_heading("E.2. VNPay", level=2)
    paragraph(doc, "Khách chọn VNPay, CheckoutPage gọi `paymentApi.initiatePayment(bookingId, method, amount)`. PaymentController nhận bookingId/method/amount, Service lock/validate booking PENDING/amount/hết hạn và tạo hoặc tái sử dụng payment phù hợp. VnPayPaymentGateway tạo URL. Browser redirect đến payment page của VNPay. Sau giao dịch, VNPay gọi callback endpoint công khai. Công khai ở đây chỉ bỏ yêu cầu Bearer JWT để VNPay có thể gọi; PaymentService vẫn verify checksum, response code, transaction, amount và trạng thái trước khi cập nhật. Callback bị lặp lại phải đi qua logic idempotent để không sinh ticket hai lần.")
    doc.add_heading("E.3. SePay/VietQR", level=2)
    paragraph(doc, "SePay là flow chuyển khoản. Gateway dựng ảnh/URL QR từ bank code, account number, account name, amount và content prefix. Khách quét bằng app ngân hàng; SePay quan sát giao dịch tiền vào rồi gọi webhook. Backend không chấp nhận chỉ vì client bấm 'đã chuyển khoản': thông tin xác nhận đến từ webhook đã xác minh API key/HMAC và payment được đối chiếu amount/content. UI có thể polling trạng thái nhưng polling chỉ đọc trạng thái server xác nhận. Nếu customer đổi promotion sau khi QR được sinh, QR có amount cũ; do đó mã nguồn và giao diện yêu cầu bỏ QR và sinh QR mới thay vì âm thầm cập nhật số tiền trên UI.")
    doc.add_heading("E.4. Sự kiện thanh toán, đối soát và yêu cầu hoàn tiền", level=2)
    paragraph(doc, "`payment_events` lưu `eventType`, trạng thái payment/booking trước-sau, success, message và JSONB payload. Nó không thay payment bảng chính mà là audit/event stream cho màn hình vận hành, retry diagnosis, signature failure, amount mismatch và reconciliation. PaymentEventServiceImpl dùng transaction mới trong các điểm ghi event để trace sự kiện có thể tồn tại độc lập theo thiết kế service. AdminPaymentPage đọc danh sách payment, refund, reconciliation và event qua PaymentController.")
    paragraph(doc, "Khi hủy suất chiếu theo policy, service có thể chuyển booking thành REFUND_PENDING và tạo Refund PENDING cho payment thành công. RefundServiceImpl cung cấp các thao tác `requestRefund`, `markRefunded` và `markRefundFailed` để quản lý yêu cầu. Các thao tác `markRefunded` và `markRefundFailed` chỉ ghi nhận kết quả do người vận hành cung cấp; chúng không gọi API nhà cung cấp để chuyển tiền. Khi quản trị viên xác nhận kết quả thành công, Payment chuyển thành REFUNDED và Booking chuyển thành REFUNDED nếu đang ở REFUND_PENDING; nếu thất bại, Payment chuyển thành REFUND_FAILED để tiếp tục theo dõi. Cách xử lý này vừa bảo toàn lịch sử vừa tránh làm mất hiệu lực booking gốc một cách vô điều kiện trong trường hợp phát sinh thanh toán đến muộn.")
    table(doc, ["Provider/flow", "Điểm vào", "Kiểm chứng trước success", "Tác động success"], [
        ["VNPay", "POST initiate; GET vnpay-callback", "Checksum, response/status, amount, payment/booking state, idempotency.", "Payment SUCCESS; Booking SUCCESS; seat BOOKED; ticket/email/WS."],
        ["SePay", "POST initiate; POST sepay-webhook", "API key/HMAC, amount, content, payment lookup, idempotency.", "Như VNPay, sau khi webhook được xác nhận."],
        ["MoMo", "Adapter và endpoint return/IPN tồn tại", "Cần cấu hình `MOMO_ENABLED` và credential hợp lệ.", "Không coi là active mặc định vì config false."],
        ["Yêu cầu refund", "Admin/staff scope payment APIs", "Refund pending/processing, cinema scope và kết quả do operator ghi nhận.", "Payment/Booking/PaymentEvent cập nhật theo kết quả được ghi nhận; chưa gọi provider refund API."],
    ], caption="Bảng PL.E.1. Đối chiếu các luồng thanh toán và hoàn tiền")

    doc.add_page_break()
    doc.add_heading("PHỤ LỤC F. WEBSOCKET, SCHEDULER VÀ TÍNH NHẤT QUÁN UI", level=1)
    doc.add_heading("F.1. Kết nối realtime", level=2)
    paragraph(doc, "`WebSocketConfig` bật `@EnableWebSocketMessageBroker`, cấu hình simple broker prefix `/topic`, application destination prefix `/app`, đăng ký `/ws` SockJS fallback và `/ws-native` Native WebSocket. Frontend useSeatWebSocket ưu tiên native endpoint; URL được xây theo `VITE_WS_BASE_URL` hoặc protocol/host hiện tại, chuyển https sang wss. Client STOMP subscribe `/topic/seatmap/{showtimeId}` sau onConnect, reconnect sau 3 giây và deactivate khi component unmount hoặc showtime đổi.")
    paragraph(doc, "`SeatStatusPublisher` dùng SimpMessagingTemplate publish `SeatStatusEvent` gồm showtimeId, seatId, status, heldByUserId, holdUntil và event time. Hold gửi event từng ghế để truyền holdUntil/user; booked/available có `publishBulk`. React giữ callback trong ref để mỗi re-render không làm reconnect/re-subscribe. Khi event đến, SeatSelectionPage cập nhật phần tử seat map tương ứng. REST vẫn cần thiết cho initial load/recovery: WebSocket không phải database và client mới kết nối phải GET map ban đầu.")
    doc.add_heading("F.2. Phát sự kiện sau khi transaction commit", level=2)
    paragraph(doc, "Nếu publisher gửi event trong khi transaction chưa commit, database có thể rollback còn client đã đổi màu ghế. HoldExpireScheduler thể hiện rõ `TransactionSynchronizationManager.registerSynchronization` và `afterCommit` trước `publishAvailable`. Booking/payment service cũng tổ chức event sau khi dữ liệu thành công theo flow. Đây là nguyên tắc nhất quán giữa state bền vững và state hiển thị; không phải WebSocket tự đảm bảo transaction.")
    doc.add_heading("F.3. Giới hạn khi mở rộng nhiều máy chủ", level=2)
    paragraph(doc, "Mã nguồn hiện dùng Spring simple broker trong memory, phù hợp ứng dụng một backend instance hoặc demo/khóa luận. Nếu nhiều instance backend, subscriber kết nối instance A sẽ không tự nhận message simple broker ở instance B. Hướng phát triển là STOMP broker relay RabbitMQ/ActiveMQ hoặc publish event qua Redis/Kafka, rồi mỗi instance broadcast local. Đây là đề xuất phát triển, không được mô tả là đã có trong hệ thống hiện tại.")

    doc.add_page_break()
    doc.add_heading("PHỤ LỤC G. CƠ SỞ DỮ LIỆU, CHỈ MỤC, BỘ NHỚ ĐỆM VÀ HIỆU NĂNG", level=1)
    doc.add_heading("G.1. Flyway là nguồn schema", level=2)
    paragraph(doc, "`application.yaml` cấu hình Flyway locations `classpath:db/migration`, validate-on-migrate=true và Hibernate `ddl-auto=validate`. Khi startup, Flyway áp migration theo version; Hibernate kiểm tra mapping entity khớp schema nhưng không tự tạo/biến đổi bảng. V1 tạo schema nền; V2 tối ưu pending payment lookup; V3 payment events; V4 refresh tokens; V5 auth audit; V6 staff cinemas; V7 SePay method; V8 email lower unique; V9 showtime status index; V10 admin filter index; V11 refund; V12 normalize event type; V13/V14 harden pending payment/booking concurrency.")
    doc.add_heading("G.2. Các index cần giải thích", level=2)
    table(doc, ["Index/constraint", "Query hoặc invariant phục vụ", "Lý do"], [
        ["unique_seat_showtime", "Lấy/cập nhật trạng thái ghế theo suất.", "Không thể có hai seat state cho cùng ghế/suất."],
        ["idx_seat_status_hold_release / hold_until", "HoldExpireScheduler scan hold hết hạn.", "Tránh full scan toàn bộ seat_status."],
        ["idx_bookings_pending_expires_*", "PendingBookingExpireScheduler tìm booking hết hạn.", "Batch cleanup theo expiry hiệu quả."],
        ["uq_bookings_pending_user_showtime", "Double click/create pending lại cùng user/suất.", "Chặn duplicate pending state ở tầng DB."],
        ["uq_payments_pending_booking(_method)", "Tạo/retry payment khi mạng chập chờn.", "Chặn nhiều pending payment cạnh tranh."],
        ["uq_payments_transaction_no", "Callback/webhook provider theo transaction.", "Định danh giao dịch ngoài không bị trùng."],
        ["payment_events/audit time indexes", "Admin filter/reconciliation/trace theo thời gian.", "Hỗ trợ vận hành dữ liệu lớn dần."],
    ], caption="Bảng PL.G.1. Index và constraint phục vụ nghiệp vụ")
    doc.add_heading("G.3. Cache đúng nơi", level=2)
    paragraph(doc, "Caffeine cache names được khai báo trong application.yaml và CacheConfig. Movie/cinema/map/room/seat definition/promotion là dữ liệu ít đổi và nhiều người đọc; service đánh dấu @Cacheable, write methods @CacheEvict. SeatStatus, Booking và Payment không được cache bởi chúng thay đổi theo transaction/callback/timeouts. Cache giảm số query read lặp lại nhưng không thay thế index, không phải distributed cache và có TTL 10 phút mặc định qua `CACHE_CAFFEINE_SPEC`.")
    doc.add_heading("G.4. Tránh N+1", level=2)
    paragraph(doc, "N+1 xảy ra khi lấy N entity rồi mỗi entity lazy load relation bằng một query riêng. Mã nguồn giảm rủi ro bằng DTO mapper, repository projection/fetch query cho các màn hình list, pagination, `hibernate.default_batch_fetch_size=50`, cache ở catalog và `open-in-view=false`. Việc tắt open-in-view buộc code trả response trong service transaction nên lỗi lazy load xuất hiện sớm thay vì âm thầm phát query lúc serialize JSON. Đây không phải tuyên bố không bao giờ có N+1; khi thêm endpoint mới phải theo dõi sql-profile, Hibernate statistics và EXPLAIN ANALYZE query thực tế.")


def add_technical_appendices(doc: Document, endpoints: list[Endpoint], counts: dict[str, int]) -> None:
    """Add the source-traceable depth needed for learning and thesis defense."""
    doc.add_page_break()
    doc.add_heading("PHỤ LỤC A. BẢN ĐỒ SOURCE CODE VÀ TRACEABILITY", level=1)
    paragraph(doc, "Phụ lục này không đưa thêm chức năng mới vào mô tả. Mỗi mục chỉ diễn giải lại vai trò của class, annotation và configuration đang hiện diện trong source. Mục tiêu là giúp người đọc mở IDE theo đúng đường dẫn, nhìn thấy quan hệ giữa code và báo cáo, đồng thời có cơ sở trả lời khi hội đồng hỏi sâu.")
    table(doc, ["Vùng source", "Vai trò", "Điểm bắt đầu nên đọc"], [
        ["src/main/java/com/cinema/booking/controller", "Lớp giao tiếp REST: route, HTTP method, validation, permission.", "AuthenticationController, BookingController, PaymentController, TicketController."],
        ["src/main/java/com/cinema/booking/service/impl", "Nghiệp vụ và transaction. Không nên đọc controller xong mà bỏ qua phần này.", "AuthenticationService, BookingServiceImpl, PaymentServiceImpl, RefundServiceImpl."],
        ["src/main/java/com/cinema/booking/repository", "Query JPA/native, projection, @Lock/@EntityGraph khi có; là bằng chứng cho data access và concurrency.", "SeatStatusRepository, BookingRepository, PaymentRepository, RefreshTokenRepository."],
        ["src/main/java/com/cinema/booking/entity", "Mô hình object/JPA; đối chiếu với Flyway để biết bảng và quan hệ.", "User, SeatStatus, Booking, Payment, Ticket, Refund."],
        ["src/main/resources/db/migration", "Nguồn schema ưu tiên: V1 tạo nền, V2–V14 bổ sung index/lifecycle/refund/concurrency.", "V1__create_cinema_schema.sql; V11, V13, V14."],
        ["cinema-client/src", "Router/layout/page/component/API/store/hook của React.", "router/AppRouter.tsx; api/axiosClient.ts; stores/authStore.ts; hooks/useSeatWebSocket.ts."],
    ])
    callout(doc, "Cách đọc trong 15 phút", "Chọn một hành động UI, ví dụ bấm Giữ ghế. Mở SeatSelectionPage để xem payload, bookingApi để thấy endpoint, BookingController để thấy HTTP boundary, BookingServiceImpl để hiểu transaction, SeatStatusRepository để thấy điều kiện/lock, migration V1/V14 để thấy unique index và cuối cùng useSeatWebSocket/SeatStatusPublisher để hiểu giao diện khác đổi màu ghế như thế nào.")

    doc.add_heading("PHỤ LỤC B. LUỒNG AUTHENTICATION, JWT VÀ SESSION", level=1)
    doc.add_heading("B.1. Thành phần tham gia", level=2)
    table(doc, ["Tệp/class", "Trách nhiệm xác minh từ source"], [
        ["AuthenticationController.java", "Expose token, Google login, introspect, refresh, logout và API quản lý session."],
        ["AuthenticationService.java", "Xử lý password/Google, tạo/verify signed JWT, refresh rotation, revoke session, audit, rate limit."],
        ["SecurityConfig.java", "Khai báo public endpoint và catch-all authenticated; tạo JwtAuthenticationConverter không thêm prefix cho authority."],
        ["CustomJwtDecoder.java", "Decoder được Resource Server dùng để xác minh access token trong request protected."],
        ["RefreshToken.java / RefreshTokenRepository.java", "Persist hash token và lock record trong refresh để chống dùng đồng thời."],
        ["authStore.ts / axiosClient.ts", "Giữ access token, quyền và thông tin user; đính Bearer, refresh một lần cho nhóm request 401."],
    ])
    doc.add_heading("B.2. Login bằng username/password", level=2)
    paragraph(doc, "Tại React, LoginPage gọi `authApi.login`, qua `axiosClient`. API client gửi POST `/auth/token` với AuthenticationRequest. AuthenticationController nhận request và chuyển `AuthenticationService.authenticate`. Service chuẩn hóa username, tạo rate-limit key từ username/request, sau đó gọi AuthRateLimitService trước khi truy vấn user. UserRepository tìm user theo username; `validateUserCanAuthenticate` kiểm tra trạng thái active/deleted/email verified theo logic service; PasswordEncoder so sánh mật khẩu thô với hash BCrypt lưu trong database.")
    paragraph(doc, "Khi xác thực hợp lệ, `issueTokenPair` tạo access token và refresh token. Token chứa issuer `cinema-booking`, token use (`access` hoặc `refresh`), subject, JWT ID, expiry, auth version và authority. Refresh token không được lưu plaintext trong database: service băm token, lưu tokenHash vào `refresh_tokens`, đồng thời ghi metadata cần thiết để quản lý session. AuthenticationResponse trả lại token/session response. Đây là lý do bản thân database bị lộ không trực tiếp làm lộ refresh token đã cấp.")
    paragraph(doc, "Nếu lỗi password, user bị khóa, rate limit vượt ngưỡng hoặc user không tồn tại, service vẫn ghi AuthAuditLog thành công/thất bại theo `AuthAuditService`. GlobalExceptionHandler chuyển exception sang mã/lời nhắn API. Frontend hiển thị thông báo tiếng Việt ở layer UI, không để browser reload trang login. Bản thân controller không tự ghi SQL hoặc tự kiểm tra password; đây là phân tách trách nhiệm hợp lý giữa transport và business layer.")
    code_block(doc, "LoginPage\n  → authApi.login(credentials)\n  → POST /auth/token\n  → AuthenticationController.authenticate(...)\n  → AuthenticationService.authenticate(...)\n  → UserRepository + PasswordEncoder\n  → issueTokenPair + RefreshTokenRepository\n  → AuthenticationResponse\n  → authStore.login(accessToken, user, permissions)", "Trace source của login password")
    doc.add_heading("B.3. Access token, refresh token và rotation", level=2)
    paragraph(doc, "Access token có thời hạn ngắn theo `JWT_ACCESS_TOKEN_VALID_DURATION` (mặc định config là 900 giây). Khi API protected trả 401, interceptor trong `axiosClient.ts` không gửi refresh song song cho từng request. Biến module `refreshPromise` là single-flight: request đầu khởi tạo POST `/auth/refresh`; request còn lại chờ cùng Promise. Nếu nhận access token mới, mỗi request retry đúng một lần bằng cờ `_retry`; nếu không thể refresh, client clear session và điều hướng login. Cách này tránh vòng lặp 401 vô hạn và tránh bão request refresh khi nhiều API cùng hết hạn.")
    paragraph(doc, "Backend verify refresh token ở `AuthenticationService.refreshToken`. Trước hết signature, token_use, expiry và auth_version được kiểm tra. Service băm token input, lấy `RefreshToken` bằng repository lock, kiểm tra token record chưa revoked/chưa expired/đúng JWT ID, tạo refresh token mới, đánh dấu token cũ `ROTATED` và lưu `replacedByTokenId`. Nếu attacker dùng lại refresh token cũ đã rotation thì bản ghi không còn active và request bị từ chối. Đây là điểm khác quan trọng giữa refresh token rotation và việc chỉ ký một JWT dài hạn không quản lý server-side.")
    paragraph(doc, "`authStore.ts` hiện chủ động loại refresh token khỏi localStorage; `axiosClient` đặt `withCredentials: true`. Điều này cho thấy client được tổ chức theo hướng refresh token không lưu trong JavaScript persistent storage. Tuy nhiên, thuộc tính cookie cụ thể như HttpOnly/SameSite/Secure cần đối chiếu đúng với response/cấu hình chạy thực tế; phần đó không nên suy diễn khi source không cho thấy header cookie tại đây. [CHƯA XÁC MINH ĐƯỢC TỪ SOURCE CODE] toàn bộ thuộc tính reverse proxy/cookie ở môi trường deployment.")
    doc.add_heading("B.4. Logout, revoke session và auth version", level=2)
    paragraph(doc, "Logout nhận access token và refresh token theo LogoutRequest. Service tìm user từ refresh token, invalidate access token bằng JWT ID với expiry vào `invalidated_token`, revoke refresh token với lý do `LOGOUT`, rồi ghi audit. API session cho phép lấy danh sách session, revoke một session hoặc revoke session khác. Khi đổi mật khẩu hoặc security action liên quan, `auth_version` trên User là tín hiệu để token cũ có claim version khác bị loại khi verify. Điều này giảm nguy cơ access token cũ tiếp tục dùng sau sự kiện bảo mật nhạy cảm.")
    doc.add_heading("B.5. Google login", level=2)
    paragraph(doc, "Frontend có Google ID token request trong `authApi.ts`; backend nhận GoogleLoginRequest. AuthenticationService decode/validate token theo Google issuer/JWK set, kiểm tra claim `email_verified`, tìm user theo email không phân biệt hoa thường hoặc tạo user Google mới, cập nhật avatar còn thiếu nếu source token có dữ liệu phù hợp và phát token pair của hệ thống. Đây là mô hình frontend lấy Google Identity credential rồi gửi token sang backend để backend xác minh. Nó không phải Spring OAuth2 redirect login vì SecurityConfig/pom không được dùng theo flow `oauth2Login()` trong source khảo sát.")

    doc.add_heading("PHỤ LỤC C. RBAC VÀ STAFF CINEMA SCOPE", level=1)
    doc.add_heading("C.1. RBAC theo permission", level=2)
    paragraph(doc, "RBAC trong source không chỉ kiểm tra role string ở route. `ApplicationInitConfig` tạo Permission từ `PermissionName.values()`, xây Role USER/STAFF từ tập permission riêng và cho ADMIN toàn bộ permission. `SecurityConfig` bật `@EnableMethodSecurity`; các method controller dùng `@PreAuthorize(\"hasAuthority('...')\")`. JwtAuthenticationConverter bỏ authority prefix để permission trong claim có thể khớp trực tiếp với các chuỗi như `SHOWTIME_CREATE`, `TICKET_CHECKIN` hoặc `PAYMENT_VIEW_ALL`.")
    table(doc, ["Role", "Ví dụ permission/capability từ source", "Giới hạn"], [
        ["USER", "MOVIE_VIEW, CINEMA_VIEW, SHOWTIME_VIEW, SEAT_VIEW, BOOKING_CREATE, PAYMENT_CREATE, TICKET_VIEW_OWN, PROFILE_UPDATE.", "Không có permission quản trị/đối soát/refund toàn hệ thống."],
        ["STAFF", "SHOWTIME_CREATE/UPDATE, TICKET_CHECKIN, BOOKING_VIEW_ALL, PAYMENT_VIEW_ALL, DASHBOARD/REPORT/ANALYTICS.", "Phải qua StaffCinemaScopeService; không tự nhiên có quyền trên mọi rạp."],
        ["ADMIN", "Tập permission đầy đủ được gán trong ApplicationInitConfig.", "Vẫn chịu validation/business policy, ví dụ trạng thái/refund workflow."],
    ])
    doc.add_heading("C.2. Scope theo rạp", level=2)
    paragraph(doc, "Bảng `staff_cinemas` (Flyway V6) liên kết `staff_id` và `cinema_id`. `StaffCinemaScopeService` là lớp tập trung kiểm tra user hiện tại có phải staff nhưng không phải admin, lấy danh sách cinema id được phân công và validate quyền truy cập cinema. Các service điều hành như showtime, room/seat, payment/refund/search và scanner gọi service scope thay vì tin vào cinemaId do frontend gửi. Đây là lớp bảo vệ chống IDOR: staff thay query string cinemaId vẫn không có quyền đọc/sửa rạp ngoài assignment.")
    paragraph(doc, "Ví dụ check-in không chỉ yêu cầu `TICKET_CHECKIN`, mà TicketController buộc request có `cinemaId` và `showtimeId`; BookingService kiểm tra ticket thuộc đúng cinema/suất, staff được gán cinema đó, và nằm trong check-in window. Sự kết hợp permission + business context + data scope mới là lý do phân quyền đúng với vận hành rạp.")

    doc.add_heading("PHỤ LỤC D. GIỮ GHẾ, BOOKING VÀ RACE CONDITION", level=1)
    doc.add_heading("D.1. Mô hình dữ liệu trạng thái ghế", level=2)
    paragraph(doc, "Ghế vật lý nằm ở `seats`, còn ghế của từng suất nằm ở `seat_status`. Một seat có nhiều SeatStatus theo nhiều showtime; unique constraint `(seat_id, showtime_id)` bảo đảm một ghế không có hai state map cho cùng suất. SeatStatus lưu `status`, `hold_by`, `hold_until` và `version`. CHECK constraint ở migration buộc HOLD phải đi kèm thông tin hold phù hợp, tránh dữ liệu nửa vời. Vì vậy availability không được suy ra từ BookingDetail một cách chậm và dễ sai, mà được lấy trực tiếp từ seat_status.")
    doc.add_heading("D.2. Hold seat step-by-step", level=2)
    bullets(doc, [
        "1. SeatSelectionPage tải initial map bằng GET `/api/v1/showtimes/{id}/seats` và subscribe WebSocket topic cùng showtime.",
        "2. Khi khách chọn ghế, frontend gửi `HoldSeatRequest` có showtimeId và danh sách seatIds đến POST `/api/v1/bookings/hold`.",
        "3. BookingServiceImpl kiểm tra rate limit `SeatHoldRateLimitService` theo user/IP để giảm spam thao tác giữ ghế.",
        "4. Service kiểm tra suất và thực hiện cleanup các expiry liên quan theo logic hiện hành; repository lấy/khóa SeatStatus mục tiêu trong transaction.",
        "5. Chỉ những row AVAILABLE hợp lệ được đổi thành HOLD, gán holdBy là user hiện tại và holdUntil theo `booking.seat-hold-minutes`.",
        "6. Transaction commit thành công mới đăng ký publish event HOLD qua SeatStatusPublisher. Nếu transaction rollback, không có broadcast trạng thái giả.",
        "7. Response trả HoldSeatResponse với holdUntil để frontend hiển thị countdown; UI không phải nguồn quyết định cuối cùng, backend mới là authority.",
    ])
    doc.add_heading("D.3. Tại sao hai khách không mua cùng ghế", level=2)
    paragraph(doc, "Giả sử A và B gần như đồng thời giữ A1 cùng showtime. Cả hai request đi vào transaction khác nhau. Repository lock/atomic condition khiến chỉ transaction lấy quyền thay đổi state AVAILABLE trước mới có thể update HOLD. Transaction sau đọc được row đã HOLD hoặc update count không đủ, service trả lỗi nghiệp vụ 'one or more seats are not available'. Không được biến lỗi đó thành 500 hoặc retry mù quáng trên frontend vì seat đã thuộc phiên giữ của khách khác. `@Version` của SeatStatus là lớp phát hiện conflict khi ORM update cạnh tranh; unique seat-showtime là ràng buộc nền dữ liệu.")
    paragraph(doc, "Sau hold, booking tạo `PENDING` và payment expiry. V14 thêm unique index pending theo user/showtime và pending payment theo booking để giảm double click/refresh tạo bản ghi pending song song. PaymentService còn khóa/kiểm tra booking/payment trước khi khởi tạo gateway. Điều này không có nghĩa mọi request trùng sẽ trả cùng response; tùy thời điểm có thể nhận response tái sử dụng payment hoặc business error. Ý nghĩa đúng là database không cho trạng thái nguy hiểm như hai payment pending cạnh tranh không kiểm soát tồn tại cho cùng đơn/gateway.")
    doc.add_heading("D.4. Hết hạn và nhả ghế", level=2)
    paragraph(doc, "`HoldExpireScheduler` chạy fixed delay theo `BOOKING_EXPIRED_HOLD_SCAN_DELAY_MS`, lấy batch projection expired, update theo id, rồi nhóm seat theo showtime để publish AVAILABLE sau commit. `PendingBookingExpireScheduler` quét Booking PENDING theo payment_expires_at/legacy cutoff, gọi `BookingService.expirePendingBooking` từng booking để một record lỗi không làm block cả batch. Khi payment thất bại, user hủy hoặc booking hết hạn, business service trả ghế về AVAILABLE nếu state/owner phù hợp. Đây là lý do refresh UI không phải điều kiện cần để ghế được giải phóng.")
    doc.add_heading("D.5. Payload mẫu để giải thích khi bảo vệ", level=2)
    code_block(doc, "POST /api/v1/bookings/hold\nAuthorization: Bearer <access-token>\n{\n  \"showtimeId\": \"<uuid>\",\n  \"seatIds\": [\"<uuid-A1>\", \"<uuid-A2>\"]\n}\n\nKết quả thành công: ApiResponse<HoldSeatResponse>\n- danh sách ghế được giữ\n- holdUntil (ISO date-time)", "Payload mô tả theo HoldSeatRequest/Response; UUID thực tế thay đổi theo dữ liệu")

    doc.add_heading("PHỤ LỤC E. PAYMENT, CALLBACK, WEBHOOK VÀ QUẢN LÝ YÊU CẦU REFUND", level=1)
    doc.add_heading("E.1. Payment abstraction", level=2)
    paragraph(doc, "`PaymentGateway` là abstraction của provider; `VnPayPaymentGateway`, `SePayPaymentGateway` và `MomoPaymentGateway` là adapter. PaymentServiceImpl chọn gateway theo PaymentMethod nhưng không để controller phải biết chi tiết checksum/QR/HMAC của từng nhà cung cấp. Tách adapter giúp thêm gateway mới mà không trộn string ký/hash vào code booking. Method/status/event types được enum và migration CHECK constraint để tránh giá trị tự do trong database.")
    doc.add_heading("E.2. VNPay", level=2)
    paragraph(doc, "Khách chọn VNPay, CheckoutPage gọi `paymentApi.initiatePayment(bookingId, method, amount)`. PaymentController nhận bookingId/method/amount, Service lock/validate booking PENDING/amount/hết hạn và tạo hoặc tái sử dụng payment phù hợp. VnPayPaymentGateway tạo URL. Browser redirect đến payment page của VNPay. Sau giao dịch, VNPay gọi callback endpoint công khai. Công khai ở đây chỉ bỏ yêu cầu Bearer JWT để VNPay có thể gọi; PaymentService vẫn verify checksum, response code, transaction, amount và trạng thái trước khi cập nhật. Callback bị lặp lại phải đi qua logic idempotent để không sinh ticket hai lần.")
    doc.add_heading("E.3. SePay/VietQR", level=2)
    paragraph(doc, "SePay là flow chuyển khoản. Gateway dựng ảnh/URL QR từ bank code, account number, account name, amount và content prefix. Khách quét bằng app ngân hàng; SePay quan sát giao dịch tiền vào rồi gọi webhook. Backend không chấp nhận chỉ vì client bấm 'đã chuyển khoản': thông tin xác nhận đến từ webhook đã xác minh API key/HMAC và payment được đối chiếu amount/content. UI có thể polling trạng thái nhưng polling chỉ đọc trạng thái server xác nhận. Nếu customer đổi promotion sau khi QR được sinh, QR có amount cũ; do đó source/UX yêu cầu bỏ QR và sinh QR mới thay vì âm thầm cập nhật số tiền trên UI.")
    doc.add_heading("E.4. Payment event, reconciliation và quản lý yêu cầu refund", level=2)
    paragraph(doc, "`payment_events` lưu `eventType`, trạng thái payment/booking trước-sau, success, message và JSONB payload. Nó không thay payment bảng chính mà là audit/event stream cho màn hình vận hành, retry diagnosis, signature failure, amount mismatch và reconciliation. PaymentEventServiceImpl dùng transaction mới trong các điểm ghi event để trace sự kiện có thể tồn tại độc lập theo thiết kế service. AdminPaymentPage đọc danh sách payment, refund, reconciliation và event qua PaymentController.")
    paragraph(doc, "Khi hủy suất chiếu theo policy, service có thể chuyển booking thành REFUND_PENDING và tạo Refund PENDING cho payment thành công. RefundServiceImpl cung cấp các thao tác `requestRefund`, `markRefunded` và `markRefundFailed` để quản lý yêu cầu. Các thao tác `markRefunded` và `markRefundFailed` chỉ ghi nhận kết quả do người vận hành cung cấp; chúng không gọi API nhà cung cấp để chuyển tiền. Khi quản trị viên xác nhận kết quả thành công, Payment chuyển thành REFUNDED và Booking chuyển thành REFUNDED nếu đang ở REFUND_PENDING; nếu thất bại, Payment chuyển thành REFUND_FAILED để tiếp tục theo dõi. Cách xử lý này vừa bảo toàn lịch sử vừa tránh làm mất hiệu lực booking gốc một cách vô điều kiện trong trường hợp phát sinh thanh toán đến muộn.")
    table(doc, ["Provider/flow", "Điểm vào", "Kiểm chứng trước success", "Tác động success"], [
        ["VNPay", "POST initiate; GET vnpay-callback", "Checksum, response/status, amount, payment/booking state, idempotency.", "Payment SUCCESS; Booking SUCCESS; seat BOOKED; ticket/email/WS."],
        ["SePay", "POST initiate; POST sepay-webhook", "API key/HMAC, amount, content, payment lookup, idempotency.", "Như VNPay, sau khi webhook được xác nhận."],
        ["MoMo", "Adapter và endpoint return/IPN tồn tại", "Cần cấu hình `MOMO_ENABLED` và credential hợp lệ.", "Không coi là active mặc định vì config false."],
        ["Yêu cầu refund", "Admin/staff scope payment APIs", "Refund pending/processing, cinema scope và kết quả do operator ghi nhận.", "Payment/Booking/PaymentEvent cập nhật theo kết quả được ghi nhận; chưa gọi provider refund API."],
    ])

    doc.add_heading("PHỤ LỤC F. WEBSOCKET, SCHEDULER VÀ TÍNH NHẤT QUÁN UI", level=1)
    doc.add_heading("F.1. Kết nối realtime", level=2)
    paragraph(doc, "`WebSocketConfig` bật `@EnableWebSocketMessageBroker`, cấu hình simple broker prefix `/topic`, application destination prefix `/app`, đăng ký `/ws` SockJS fallback và `/ws-native` Native WebSocket. Frontend useSeatWebSocket ưu tiên native endpoint; URL được xây theo `VITE_WS_BASE_URL` hoặc protocol/host hiện tại, chuyển https sang wss. Client STOMP subscribe `/topic/seatmap/{showtimeId}` sau onConnect, reconnect sau 3 giây và deactivate khi component unmount hoặc showtime đổi.")
    paragraph(doc, "`SeatStatusPublisher` dùng SimpMessagingTemplate publish `SeatStatusEvent` gồm showtimeId, seatId, status, heldByUserId, holdUntil và event time. Hold gửi event từng ghế để truyền holdUntil/user; booked/available có `publishBulk`. React giữ callback trong ref để mỗi re-render không làm reconnect/re-subscribe. Khi event đến, SeatSelectionPage cập nhật phần tử seat map tương ứng. REST vẫn cần thiết cho initial load/recovery: WebSocket không phải database và client mới kết nối phải GET map ban đầu.")
    doc.add_heading("F.2. Publish after commit", level=2)
    paragraph(doc, "Nếu publisher gửi event trong khi transaction chưa commit, database có thể rollback còn client đã đổi màu ghế. HoldExpireScheduler thể hiện rõ `TransactionSynchronizationManager.registerSynchronization` và `afterCommit` trước `publishAvailable`. Booking/payment service cũng tổ chức event sau khi dữ liệu thành công theo flow. Đây là nguyên tắc nhất quán giữa state bền vững và state hiển thị; không phải WebSocket tự đảm bảo transaction.")
    doc.add_heading("F.3. Giới hạn scale", level=2)
    paragraph(doc, "Source dùng Spring simple broker trong memory, phù hợp ứng dụng một backend instance hoặc demo/khóa luận. Nếu nhiều instance backend, subscriber kết nối instance A sẽ không tự nhận message simple broker ở instance B. Hướng phát triển là STOMP broker relay RabbitMQ/ActiveMQ hoặc publish event qua Redis/Kafka, rồi mỗi instance broadcast local. Đây là đề xuất phát triển, không được mô tả là đã có trong source.")

    doc.add_heading("PHỤ LỤC G. DATABASE, INDEX, CACHE VÀ HIỆU NĂNG", level=1)
    doc.add_heading("G.1. Flyway là nguồn schema", level=2)
    paragraph(doc, "`application.yaml` cấu hình Flyway locations `classpath:db/migration`, validate-on-migrate=true và Hibernate `ddl-auto=validate`. Khi startup, Flyway áp migration theo version; Hibernate kiểm tra mapping entity khớp schema nhưng không tự tạo/biến đổi bảng. V1 tạo schema nền; V2 tối ưu pending payment lookup; V3 payment events; V4 refresh tokens; V5 auth audit; V6 staff cinemas; V7 SePay method; V8 email lower unique; V9 showtime status index; V10 admin filter index; V11 refund; V12 normalize event type; V13/V14 harden pending payment/booking concurrency.")
    doc.add_heading("G.2. Các index cần giải thích", level=2)
    table(doc, ["Index/constraint", "Query hoặc invariant phục vụ", "Lý do"], [
        ["unique_seat_showtime", "Lấy/cập nhật trạng thái ghế theo suất.", "Không thể có hai seat state cho cùng ghế/suất."],
        ["idx_seat_status_hold_release / hold_until", "HoldExpireScheduler scan hold hết hạn.", "Tránh full scan toàn bộ seat_status."],
        ["idx_bookings_pending_expires_*", "PendingBookingExpireScheduler tìm booking hết hạn.", "Batch cleanup theo expiry hiệu quả."],
        ["uq_bookings_pending_user_showtime", "Double click/create pending lại cùng user/suất.", "Chặn duplicate pending state ở tầng DB."],
        ["uq_payments_pending_booking(_method)", "Tạo/retry payment khi mạng chập chờn.", "Chặn nhiều pending payment cạnh tranh."],
        ["uq_payments_transaction_no", "Callback/webhook provider theo transaction.", "Định danh giao dịch ngoài không bị trùng."],
        ["payment_events/audit time indexes", "Admin filter/reconciliation/trace theo thời gian.", "Hỗ trợ vận hành dữ liệu lớn dần."],
    ])
    doc.add_heading("G.3. Cache đúng nơi", level=2)
    paragraph(doc, "Caffeine cache names được khai báo trong application.yaml và CacheConfig. Movie/cinema/map/room/seat definition/promotion là dữ liệu ít đổi và nhiều người đọc; service đánh dấu @Cacheable, write methods @CacheEvict. SeatStatus, Booking và Payment không được cache bởi chúng thay đổi theo transaction/callback/timeouts. Cache giảm số query read lặp lại nhưng không thay thế index, không phải distributed cache và có TTL 10 phút mặc định qua `CACHE_CAFFEINE_SPEC`.")
    doc.add_heading("G.4. Tránh N+1", level=2)
    paragraph(doc, "N+1 xảy ra khi lấy N entity rồi mỗi entity lazy load relation bằng một query riêng. Source giảm rủi ro bằng DTO mapper, repository projection/fetch query cho các màn hình list, pagination, `hibernate.default_batch_fetch_size=50`, cache ở catalog và `open-in-view=false`. Việc tắt open-in-view buộc code trả response trong service transaction nên lỗi lazy load xuất hiện sớm thay vì âm thầm phát query lúc serialize JSON. Đây không phải tuyên bố không bao giờ có N+1; khi thêm endpoint mới phải theo dõi sql-profile, Hibernate statistics và EXPLAIN ANALYZE query thực tế.")

    doc.add_heading("PHỤ LỤC H. FRONTEND UI, API VÀ ERROR HANDLING", level=1)
    doc.add_heading("H.1. Điều hướng và route protection", level=2)
    paragraph(doc, "AppRouter lazy import các page để giảm tải initial bundle. Auth routes nằm trong AuthLayout; admin/staff trong AdminLayout; public/customer trong PublicLayout. ProtectedRoute nhận permission cho admin/staff routes, còn customer routes yêu cầu đăng nhập. Đây là guard UX: route không hiển thị màn hình khi state client không có quyền. Server vẫn là lớp bảo vệ cuối, vì người dùng có thể gọi API trực tiếp bên ngoài browser.")
    doc.add_heading("H.2. API layer", level=2)
    paragraph(doc, "Các module `movieApi`, `cinemaApi`, `showtimeApi`, `bookingApi`, `paymentApi`, `ticketApi`, `roomSeatApi`, `promotionApi`, `userApi`, `analyticsApi`, `auditLogApi`, `authApi` gom URL và type. Page không nên tự copy chuỗi endpoint. `axiosClient` có base URL theo `VITE_API_BASE_URL`; local dev để rỗng để Vite proxy `/api`, `/auth`, `/ws` và `/ws-native` tới `BACKEND_PROXY_TARGET`. Deploy tách domain có thể đặt URL API/WS qua env.")
    doc.add_heading("H.3. Validation và trải nghiệm thanh toán", level=2)
    paragraph(doc, "Frontend có React Hook Form/Zod trong dependency và form pages; backend vẫn đặt jakarta validation để mọi caller, kể cả Postman hay callback, chịu cùng quy tắc. Ở checkout, sự thay đổi promotion làm cập nhật Booking phía server; nếu SePay QR đã được tạo, UX yêu cầu tạo lại QR để amount/content khớp. Result page đọc trạng thái đã xác nhận thay vì tin vào click của client. Đây là cách tránh case customer vừa tạo QR 275.000đ rồi giảm đơn xuống 245.000đ nhưng vẫn chuyển khoản QR cũ.")
    doc.add_heading("H.4. Error boundary và thông báo", level=2)
    paragraph(doc, "`App.tsx` có React ErrorBoundary toàn ứng dụng; Toast và API error handler ở page/API layer trả thông tin dễ hiểu. Ở backend, ApiResponse + GlobalExceptionHandler chuẩn hóa code/message/path/timestamp. Hai phía có nhiệm vụ khác nhau: client dịch lỗi nghiệp vụ thành UI; backend không được dựa vào UI để giữ nguyên tắc dữ liệu.")

    doc.add_heading("PHỤ LỤC I. CHECKLIST DEMO VÀ CÂU HỎI BẢO VỆ", level=1)
    doc.add_heading("I.1. Checklist demo theo thứ tự an toàn", level=2)
    bullets(doc, [
        "1. Chạy PostgreSQL/Docker hoặc database đã migrate; chạy backend, frontend; kiểm tra Swagger và trang chủ public.",
        "2. Đăng ký user, nhận email xác thực (nếu SMTP cấu hình), login password và xem profile/session.",
        "3. Mở cùng một suất ở hai browser/user; A giữ ghế, B quan sát ghế chuyển HOLD realtime; B thử giữ cùng ghế và nhận business error.",
        "4. A tạo booking, áp khuyến mãi, demo VNPay hoặc SePay. Với SePay, chỉ xác nhận sau webhook/provider event, không chỉ click client.",
        "5. Mở vé QR, login staff có assignment, chọn cinema/showtime phù hợp, quét ticket một lần rồi quét lại để thấy USED; thử context sai rạp/suất.",
        "6. Demo admin: tạo/sửa catalog, xem filter booking/payment/audit, hủy suất theo policy, xem REFUND_PENDING/refund event nếu dữ liệu test có sẵn.",
        "7. Chạy test backend/frontend đã có; chỉ báo cáo pass/fail thực tế, không tự tạo TPS/coverage." ,
    ])
    doc.add_heading("I.2. Ba câu hỏi xoáy và cách trả lời", level=2)
    table(doc, ["Câu hỏi", "Ý trả lời có căn cứ source"], [
        ["Nếu hai khách bấm cùng ghế cùng lúc, hệ thống bảo đảm thế nào?", "Không tin UI. SeatStatus theo seat+showtime có unique constraint, service transaction và repository lock/atomic update; version phát hiện conflict; pending booking/payment có partial unique index; thất bại trả business error, không bán trùng."],
        ["Tại sao callback VNPay/SePay lại public, có nguy hiểm không?", "Gateway ngoài không có JWT của user nên endpoint cần public ở filter. Nhưng PaymentService/Gateway vẫn verify checksum hoặc API key/HMAC, đối chiếu transaction/amount/content và idempotency. Public route không đồng nghĩa bỏ xác thực nghiệp vụ."],
        ["Tại sao dùng cả RBAC và staff scope?", "Permission trả lời 'có được làm loại thao tác này không'; scope trả lời 'được làm trên rạp nào'. Nếu chỉ RBAC, staff có SHOWTIME_UPDATE có thể sửa rạp khác bằng ID đoán được."],
    ])
    doc.add_heading("I.3. Những điều cần nói trung thực", level=2)
    paragraph(doc, "Không khẳng định đã benchmark 10.000 user, đã deploy production, đã dùng Redis/RabbitMQ/Kafka/CI-CD hay hoàn tiền tự động qua provider nếu chưa có bằng chứng/source/config/kết quả chạy. Cách trả lời tốt là nêu đúng hiện trạng: hệ thống hiện có transaction/lock/index/cache local/scheduler/payment event; khi scale nhiều instance sẽ chuyển sang broker/cache/queue phân tán như hướng phát triển.")
def add_references(doc: Document) -> None:
    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    refs = [
        ("[1] R. T. Fielding, Architectural Styles and the Design of Network-based Software Architectures, Doctoral dissertation, University of California, Irvine, 2000", "https://ics.uci.edu/~fielding/pubs/dissertation/top.htm", "18/08/2026"),
        ("[2] React Team, React Documentation – Components and Hooks", "https://react.dev/reference/react", "12/08/2026"),
        ("[3] Microsoft, The TypeScript Handbook – The Basics", "https://www.typescriptlang.org/docs/handbook/2/basic-types.html", "18/08/2026"),
        ("[4] Spring Team, Spring Boot Reference Documentation", "https://docs.spring.io/spring-boot/reference/", "12/08/2026"),
        ("[5] PostgreSQL Global Development Group, PostgreSQL Documentation – Indexes", "https://www.postgresql.org/docs/current/indexes.html", "12/08/2026"),
        ("[6] Eclipse Foundation, Jakarta Persistence 3.2 Specification", "https://jakarta.ee/specifications/persistence/3.2/", "18/08/2026"),
        ("[7] Spring Team, Spring Data JPA Reference Documentation", "https://docs.spring.io/spring-data/jpa/reference/", "18/08/2026"),
        ("[8] Hibernate Team, Hibernate ORM User Guide", "https://docs.hibernate.org/orm/current/userguide/html_single/Hibernate_User_Guide.html", "18/08/2026"),
        ("[9] PostgreSQL Global Development Group, PostgreSQL Documentation – Transactions", "https://www.postgresql.org/docs/current/tutorial-transactions.html", "18/08/2026"),
        ("[10] PostgreSQL Global Development Group, PostgreSQL Documentation – Explicit Locking", "https://www.postgresql.org/docs/current/explicit-locking.html", "18/08/2026"),
        ("[11] R. Fielding, M. Nottingham và J. Reschke, HTTP Semantics, RFC 9110", "https://www.rfc-editor.org/rfc/rfc9110.html", "18/08/2026"),
        ("[12] Google, Verify the Google ID token on your server side", "https://developers.google.com/identity/gsi/web/guides/verify-google-id-token", "18/08/2026"),
        ("[13] M. Jones, J. Bradley và N. Sakimura, JSON Web Token (JWT), RFC 7519", "https://www.rfc-editor.org/rfc/rfc7519.html", "12/08/2026"),
        ("[14] T. Lodderstedt, J. Bradley, A. Labunets và D. Fett, Best Current Practice for OAuth 2.0 Security, RFC 9700", "https://www.rfc-editor.org/rfc/rfc9700.html", "18/08/2026"),
        ("[15] Spring Team, Spring Security Method Security", "https://docs.spring.io/spring-security/reference/servlet/authorization/method-security.html", "12/08/2026"),
        ("[16] Spring Team, STOMP over WebSocket", "https://docs.spring.io/spring-framework/reference/web/websocket/stomp.html", "12/08/2026"),
        ("[17] World Wide Web Consortium (W3C), Geolocation", "https://www.w3.org/TR/geolocation/", "18/08/2026"),
        ("[18] R. W. Sinnott, Virtues of the Haversine, Sky & Telescope, vol. 68, no. 2, p. 159, 1984", "https://ui.adsabs.harvard.edu/abs/1984S%26T....68..159S/abstract", "18/08/2026"),
        ("[19] Git Project, Git Reference Manual", "https://git-scm.com/docs", "12/08/2026"),
        ("[20] Apache Software Foundation, Apache JMeter User Manual", "https://jmeter.apache.org/usermanual/index.html", "12/08/2026"),
    ]
    for title, url, accessed_at in refs:
        p = doc.add_paragraph()
        add_inline_text(p, f"{title}, ")
        add_hyperlink(p, url, url)
        add_inline_text(p, f" (truy cập ngày {accessed_at}).")


def create_report(endpoints: list[Endpoint], counts: dict[str, int]) -> None:
    doc = Document()
    configure_document(doc)
    enable_field_updates(doc)
    add_cover(doc)
    start_numbered_section(doc, "lowerRoman", 1)
    add_formal_pre_toc_pages(doc)
    add_toc_page(doc)
    add_formal_abbreviations(doc)
    doc.add_heading("DANH MỤC HÌNH ẢNH", level=1)
    add_style_toc(doc, "Figure Caption", "Mở file trong Microsoft Word, bấm Update Field để cập nhật danh mục hình.")
    doc.add_page_break()
    doc.add_heading("DANH MỤC BẢNG", level=1)
    add_style_toc(doc, "Table Caption", "Mở file trong Microsoft Word, bấm Update Field để cập nhật danh mục bảng.")
    doc.add_page_break()
    add_formal_summary_and_preface(doc)
    start_numbered_section(doc, "decimal", 1)
    add_chapter_one(doc, counts)
    expand_chapter_one(doc)
    doc.add_page_break()
    add_chapter_two(doc)
    expand_chapter_two(doc)
    doc.add_page_break()
    add_chapter_three(doc, endpoints)
    expand_chapter_three(doc)
    doc.add_page_break()
    add_chapter_four(doc, counts)
    doc.add_page_break()
    add_chapter_five(doc)
    doc.add_page_break()
    add_references(doc)
    add_formal_appendices(doc)
    add_advanced_appendices_d_to_g(doc)
    doc.save(REPORT)


def write_api_inventory(endpoints: list[Endpoint], calls: dict[str, list[str]]) -> None:
    groups: dict[str, list[Endpoint]] = defaultdict(list)
    for endpoint in endpoints:
        groups[endpoint.module].append(endpoint)
    output = [
        "# CinemaBooking API Inventory",
        "",
        "Tài liệu này được trích từ annotation trong `src/main/java/com/cinema/booking/controller`. Source code là căn cứ ưu tiên. Cột frontend được dò từ `cinema-client/src/api/*.ts`; nếu không có kết quả thì không suy diễn rằng API không thể được gọi gián tiếp.",
        "",
        f"Số endpoint trích được: **{len(endpoints)}**.",
        "",
        "## API CONSISTENCY FINDINGS",
        "",
        "- API client frontend tập trung trong `src/api`, giúp đối chiếu dễ hơn. Các URL có nội suy `${...}` được chuẩn hóa thành `{id}` trong quá trình dò.",
        "- Một số endpoint backend là callback/webhook hoặc scheduler support; không có lời gọi trực tiếp từ frontend là có chủ đích, ví dụ VNPay callback, SePay webhook, MoMo IPN.",
        "- `MomoPaymentGateway` tồn tại ở backend nhưng `momo.enabled` mặc định là `false`; frontend không được xem là đang cung cấp cổng MoMo mặc định nếu chưa có cấu hình hợp lệ.",
        "- Có hai route đọc sơ đồ ghế cùng gọi `bookingService.getSeatMap`: `/api/v1/showtimes/{id}/seats` và `/api/v1/bookings/showtimes/{showtimeId}/seats`. React API client gọi route thứ nhất. Source chưa thể hiện annotation deprecation cho route thứ hai; nên chọn một canonical route và có giai đoạn deprecate trước khi xóa để tránh tăng bề mặt API không cần thiết.",
        "- Tương tự, `BookingController` và `TicketController` đều có POST check-in, cùng xác thực QR/cinema/showtime rồi gọi `bookingService.checkInTicket`; React gọi `/api/v1/tickets/check-in`. Nên xem route ở `TicketController` là canonical vì domain rõ ràng hơn và quản lý alias còn lại theo chính sách tương thích API.",
        "- Không phát hiện endpoint React gọi sai HTTP method từ các API client tĩnh đã dò. Các lời gọi động ngoài thư mục `src/api` cần được kiểm tra lại nếu về sau thêm code mới.",
        "",
    ]
    names = {
        "Authentication": "Authentication APIs", "User": "User APIs", "Movie": "Movie APIs", "Cinema": "Cinema APIs",
        "Room": "Room APIs", "Seat": "Seat APIs", "Showtime": "Showtime APIs", "Booking": "Booking APIs",
        "Payment": "Payment APIs", "Ticket": "Ticket APIs", "Promotion": "Promotion APIs", "Analytics": "Analytics APIs",
        "AdminAuditLog": "Admin Audit APIs", "AuthAuditLog": "Authentication Audit APIs",
    }
    for module in sorted(groups):
        output.extend([f"## {names.get(module, module + ' APIs')}", ""])
        output.extend([
            "| Module | HTTP Method | Endpoint | Controller | Method | Request DTO/Params | Response | Authorization | Front-end sử dụng |",
            "| --- | --- | --- | --- | --- | --- | --- | --- | --- |",
        ])
        for item in groups[module]:
            frontend = frontend_usage_for(item.endpoint, calls).replace("|", "\\|")
            fields = [item.module, item.http_method, f"`{item.endpoint}`", item.controller, item.method, item.request, item.response, item.authorization, frontend]
            output.append("| " + " | ".join(field.replace("\n", " ") for field in fields) + " |")
        output.append("")
    API_INVENTORY.write_text("\n".join(output), encoding="utf-8")


def plantuml_text() -> str:
    return r'''# CINEMABOOKING - PLANTUML FINAL
# Căn cứ: source front-end/back-end, Flyway V1-V14 và configuration hiện hành.
# File diagrams_plantuml.txt chỉ được tham khảo về ý tưởng bố cục; nội dung dưới đây đã được đối chiếu lại với implementation hiện tại.
#
# HƯỚNG DẪN DÙNG TRONG BÁO CÁO
# 1. Sao chép trọn vẹn MỘT khối PlantUML, bao gồm dòng mở đầu và dòng kết thúc khối.
# 2. Xuất SVG để chèn Word (ưu tiên) hoặc PNG độ phân giải cao.
# 3. Chèn ảnh tại marker [CHÈN FILE: ...] tương ứng trong Chương 3.
# 4. Không dán toàn bộ file vào PlantUML trong một lần nếu chỉ cần vẽ một sơ đồ.
#
# DANH MỤC SƠ ĐỒ DÙNG TRONG WORD
# [WORD: Hình 3.1] Use Case Diagram - CinemaBooking
# [WORD: Hình 3.2] Class Diagram - Domain Model
# [WORD: Hình 3.3] Sequence - Temporary Seat Hold
# [WORD: Hình 3.4] Activity - Seat Hold and Booking
# [WORD: Hình 3.5] Sequence - VNPay Payment
# [WORD: Hình 3.6] Sequence - SePay VietQR Payment
# [WORD: Hình 3.7] Sequence - Ticket Check-in
# [WORD: Hình 3.8] ERD - PostgreSQL CinemaBooking
# [WORD: Hình 3.9] State Diagram - Booking
# [WORD: Hình 3.10] State Diagram - Payment
# [WORD: Hình 3.11] Sequence - Access Token Expiry and Refresh Token Rotation
# [WORD: Hình 3.12] Sequence - WebSocket Seat Map Synchronization
# [WORD: Hình 3.13] Activity - Showtime Cancellation and Refund Request
# [WORD: Hình 3.14] Activity - Refund Processing
# [WORD: Hình 3.15] Component Diagram - CinemaBooking

' [WORD: Hình 3.1] Use Case tổng quát
@startuml
title Use Case Diagram - Hệ thống CinemaBooking
left to right direction
skinparam shadowing false
skinparam packageStyle rectangle
actor "Khách hàng\n(USER)" as USER
actor "Nhân viên rạp\n(STAFF)" as STAFF
actor "Quản trị viên\n(ADMIN)" as ADMIN
actor "VNPay" as VNPAY
actor "SePay / VietQR" as SEPAY

rectangle "CinemaBooking" {
  usecase "Xem phim, rạp, lịch chiếu" as UC_BROWSE
  usecase "Đăng ký / xác thực email\nreset mật khẩu / đăng nhập" as UC_AUTH
  usecase "Chọn ghế, giữ ghế, tạo đơn" as UC_BOOK
  usecase "Áp dụng / bỏ khuyến mãi" as UC_PROMO
  usecase "Thanh toán VNPay" as UC_VNPAY
  usecase "Thanh toán QR ngân hàng" as UC_SEPAY
  usecase "Xem đơn, vé QR, hồ sơ, phiên" as UC_ACCOUNT
  usecase "Soát vé QR đúng rạp/suất" as UC_CHECKIN
  usecase "Quản lý suất chiếu, phòng/ghế\ntrong rạp được phân công" as UC_STAFF_OPS
  usecase "Quản trị catalog, người dùng, RBAC\nbooking/payment/refund request/audit/analytics" as UC_ADMIN
}
USER --> UC_BROWSE
USER --> UC_AUTH
USER --> UC_BOOK
USER --> UC_PROMO
USER --> UC_VNPAY
USER --> UC_SEPAY
USER --> UC_ACCOUNT
STAFF --> UC_CHECKIN
STAFF --> UC_STAFF_OPS
ADMIN --> UC_ADMIN
VNPAY --> UC_VNPAY
SEPAY --> UC_SEPAY
@enduml

' [WORD: Hình 3.15] Sơ đồ thành phần
@startuml
title Component Diagram - CinemaBooking
top to bottom direction
skinparam shadowing false
skinparam linetype ortho
actor USER
actor STAFF
actor ADMIN
component "React + TypeScript + Vite\nPages / Components / Zustand / Axios" as FE
component "Spring Security\nJWT + RBAC + Staff Scope" as SECURITY
component "REST Controllers\nDTO Validation / ApiResponse" as API
component "Business Services\nAuth / Booking / Payment / Ticket / Admin" as SERVICE
component "WebSocket STOMP\nSeatStatusPublisher" as WS
component "JPA Repositories\nProjection / Lock Query" as REPO
database "PostgreSQL\nFlyway schema" as DB
component "Caffeine Cache" as CACHE
component "Scheduler\nHold / Pending booking / Showtime" as SCHEDULER
cloud "VNPay" as VNPAY
cloud "SePay / VietQR" as SEPAY
cloud "SMTP / Mail" as MAIL
cloud "Google ID Token" as GOOGLE
USER --> FE
STAFF --> FE
ADMIN --> FE
FE --> SECURITY : Bearer JWT
FE --> API : HTTPS JSON
FE <--> WS : /ws-native, STOMP
SECURITY --> API
API --> SERVICE
SERVICE --> REPO
SERVICE --> CACHE
REPO --> DB
SCHEDULER --> SERVICE
SERVICE --> WS : publish after commit
SERVICE --> VNPAY
SERVICE --> SEPAY
SERVICE --> MAIL
SERVICE --> GOOGLE
@enduml

' [WORD: Hình 3.2] Sơ đồ lớp miền nghiệp vụ
@startuml
title Class Diagram - Mô hình miền nghiệp vụ CinemaBooking
hide methods
skinparam shadowing false
skinparam classAttributeIconSize 0
class User {
  UUID id
  String username
  String email
  Boolean emailVerified
  Integer authVersion
  Boolean isActive
  Boolean isDeleted
}
class Role {
  UUID id
  String name
}
class Permission {
  UUID id
  String name
}
class StaffCinema {
  UUID staffId
  UUID cinemaId
}
class Movie {
  UUID id
  String title
  Integer duration
  MovieStatus status
  Boolean isDeleted
}
class Cinema {
  UUID id
  String name
  String city
  Double latitude
  Double longitude
  Boolean isActive
}
class Room {
  UUID id
  String name
  Boolean isDeleted
}
class Seat {
  UUID id
  String rowLabel
  Integer seatNumber
  SeatType seatType
  BigDecimal priceMultiplier
}
class Showtime {
  UUID id
  LocalDateTime startTime
  LocalDateTime endTime
  BigDecimal basePrice
  ShowtimeStatus status
}
class SeatStatus {
  UUID id
  SeatStatusType status
  LocalDateTime holdUntil
  Integer version
}
class Booking {
  UUID id
  BigDecimal totalPrice
  BigDecimal discountAmount
  BookingStatus status
  String secureToken
  LocalDateTime paymentExpiresAt
}
class BookingDetail {
  UUID id
  BigDecimal priceAtBooking
}
class Payment {
  UUID id
  BigDecimal amount
  PaymentMethod method
  PaymentStatus status
  String transactionNo
}
class PaymentEvent {
  UUID id
  PaymentEventType eventType
  Boolean success
  LocalDateTime createdAt
}
class Refund {
  UUID id
  BigDecimal amount
  RefundStatus status
  String reason
  LocalDateTime requestedAt
}
class Ticket {
  UUID id
  String qrCode
  TicketStatus status
  LocalDateTime checkInTime
}
class RefreshToken {
  UUID id
  String tokenHash
  LocalDateTime expiresAt
  LocalDateTime revokedAt
}
class InvalidatedToken {
  String id
  LocalDateTime expiryTime
}
class AuthAuditLog {
  UUID id
  String eventType
  Boolean success
  LocalDateTime createdAt
}
class AdminAuditLog {
  UUID id
  String action
  String resourceType
  Boolean success
  LocalDateTime createdAt
}
User "*" -- "*" Role : users_roles
Role "*" -- "*" Permission : roles_permissions
User "1" -- "*" StaffCinema
Cinema "1" -- "*" StaffCinema
Cinema "1" -- "*" Room
Room "1" -- "*" Seat
Movie "1" -- "*" Showtime
Room "1" -- "*" Showtime
Showtime "1" -- "*" SeatStatus
Seat "1" -- "*" SeatStatus
User "1" -- "*" Booking
Showtime "1" -- "*" Booking
Booking "1" -- "*" BookingDetail
Seat "1" -- "*" BookingDetail
BookingDetail "1" -- "0..1" Ticket
Booking "1" -- "*" Payment
Payment "1" -- "*" PaymentEvent
Booking "1" -- "*" PaymentEvent
Payment "1" -- "*" Refund
Booking "1" -- "*" Refund
User "1" -- "*" RefreshToken
User "1" -- "*" AuthAuditLog
User "0..1" -- "*" AdminAuditLog : actor
@enduml

' [WORD: Hình 3.8] Sơ đồ quan hệ cơ sở dữ liệu
@startuml
title ERD - Cơ sở dữ liệu PostgreSQL CinemaBooking
hide circle
skinparam shadowing false
entity users {
  * id : uuid <<PK>>
  --
  username : varchar <<UQ>>
  email : varchar <<UQ>>
  auth_version : integer
  is_active : boolean
  is_deleted : boolean
}
entity roles {
  * id : uuid <<PK>>
  --
  name : varchar <<UQ>>
}
entity permissions {
  * id : uuid <<PK>>
  --
  name : varchar <<UQ>>
}
entity users_roles {
  * user_id : uuid <<FK>>
  * role_id : uuid <<FK>>
}
entity roles_permissions {
  * role_id : uuid <<FK>>
  * permission_id : uuid <<FK>>
}
entity staff_cinemas {
  * staff_id : uuid <<FK>>
  * cinema_id : uuid <<FK>>
}
entity movies {
  * id : uuid <<PK>>
  title : varchar
  status : varchar
}
entity cinemas {
  * id : uuid <<PK>>
  name : varchar
  city : varchar
  latitude : double
  longitude : double
}
entity rooms {
  * id : uuid <<PK>>
  cinema_id : uuid <<FK>>
  name : varchar
}
entity seats {
  * id : uuid <<PK>>
  room_id : uuid <<FK>>
  row_label : varchar
  seat_number : int
}
entity showtimes {
  * id : uuid <<PK>>
  movie_id : uuid <<FK>>
  room_id : uuid <<FK>>
  start_time : timestamp
  end_time : timestamp
  status : varchar
}
entity seat_status {
  * id : uuid <<PK>>
  seat_id : uuid <<FK>>
  showtime_id : uuid <<FK>>
  status : varchar
  hold_by : uuid <<FK>>
  hold_until : timestamp
  version : int
}
entity bookings {
  * id : uuid <<PK>>
  user_id : uuid <<FK>>
  showtime_id : uuid <<FK>>
  promotion_id : uuid <<FK>>
  status : varchar
  payment_expires_at : timestamp
}
entity booking_details {
  * id : uuid <<PK>>
  booking_id : uuid <<FK>>
  seat_id : uuid <<FK>>
}
entity payments {
  * id : uuid <<PK>>
  booking_id : uuid <<FK>>
  method : varchar
  status : varchar
  transaction_no : varchar <<UQ>>
}
entity payment_events {
  * id : uuid <<PK>>
  payment_id : uuid <<FK>>
  booking_id : uuid <<FK>>
  event_type : varchar
  payload : jsonb
}
entity refunds {
  * id : uuid <<PK>>
  booking_id : uuid <<FK>>
  payment_id : uuid <<FK>>
  status : varchar
}
entity tickets {
  * id : uuid <<PK>>
  booking_detail_id : uuid <<FK,UQ>>
  qr_code : text <<UQ>>
  status : varchar
}
entity refresh_tokens {
  * id : uuid <<PK>>
  user_id : uuid <<FK>>
  token_hash : varchar
}
users ||--o{ users_roles
roles ||--o{ users_roles
roles ||--o{ roles_permissions
permissions ||--o{ roles_permissions
users ||--o{ staff_cinemas
cinemas ||--o{ staff_cinemas
cinemas ||--o{ rooms
rooms ||--o{ seats
movies ||--o{ showtimes
rooms ||--o{ showtimes
showtimes ||--o{ seat_status
seats ||--o{ seat_status
users ||--o{ bookings
showtimes ||--o{ bookings
bookings ||--o{ booking_details
seats ||--o{ booking_details
booking_details ||--o| tickets
bookings ||--o{ payments
payments ||--o{ payment_events
bookings ||--o{ payment_events
payments ||--o{ refunds
bookings ||--o{ refunds
users ||--o{ refresh_tokens
@enduml

' [WORD: Hình 3.11] Access token hết hạn và luân chuyển refresh token
@startuml
title Sequence Diagram - Access Token hết hạn và Refresh Token Rotation
autonumber
actor User
participant "React Client\naxiosClient" as FE
participant AuthenticationController as CTRL
participant AuthenticationService as AUTH
database PostgreSQL as DB
User -> FE : thao tác trên phiên đã đăng nhập
FE -> CTRL : API request kèm access token đã hết hạn
CTRL --> FE : 401 Unauthorized
FE -> FE : gom các request 401 vào một refresh promise
FE -> CTRL : POST /auth/refresh\nHttpOnly refresh cookie
CTRL -> AUTH : refreshToken(request metadata)
AUTH -> DB : khóa bản ghi refresh token theo hash
AUTH -> AUTH : xác minh token_use, expiry, auth_version
alt refresh token hợp lệ
  AUTH -> DB : thu hồi token cũ; lưu token mới đã băm
  AUTH --> CTRL : cặp token mới
  CTRL --> FE : access token mới + refresh cookie mới
  FE -> FE : phát lại các request đang chờ
else token không hợp lệ hoặc đã bị thu hồi
  AUTH --> CTRL : lỗi xác thực
  CTRL --> FE : 401 Unauthorized
  FE -> FE : xóa phiên cục bộ và chuyển đến trang đăng nhập
end
@enduml

' [WORD: Hình 3.3] Luồng giữ ghế tạm thời
@startuml
title Sequence Diagram - Giữ ghế tạm thời
autonumber
actor Customer
participant "React SeatSelectionPage" as FE
participant BookingController as CTRL
participant BookingServiceImpl as SERVICE
participant SeatStatusRepository as REPO
database PostgreSQL as DB
participant SeatStatusPublisher as WS
Customer -> FE : choose seats
FE -> CTRL : POST /api/v1/bookings/hold\nHoldSeatRequest(showtimeId, seatIds)
CTRL -> SERVICE : holdSeats(request)
SERVICE -> SERVICE : rate limit + validate showtime
SERVICE -> REPO : lock target SeatStatus rows
REPO -> DB : SELECT/UPDATE in transaction
alt all target seats AVAILABLE
  SERVICE -> DB : AVAILABLE -> HOLD, holdBy, holdUntil
  SERVICE -> WS : register publish after commit
  SERVICE --> CTRL : HoldSeatResponse
  WS -> FE : /topic/seatmap/{showtimeId} HOLD event
else a seat is not available
  SERVICE --> CTRL : business error; no partial success
end
@enduml

' [WORD: Hình 3.4] Quy trình hoạt động giữ ghế và tạo đơn
@startuml
title Activity Diagram - Quy trình giữ ghế và tạo đơn đặt vé
skinparam shadowing false
skinparam linetype ortho
start
:Khách hàng lựa chọn ghế;
:POST /api/v1/bookings/hold\nHoldSeatRequest(showtimeId, seatIds);
:Kiểm tra rate limit và suất chiếu;
:Khóa các dòng SeatStatus trong transaction;
if (Tất cả ghế ở trạng thái AVAILABLE?) then (Có)
  :Chuyển AVAILABLE -> HOLD;\nLưu holdBy và holdUntil;
  :Đăng ký phát sự kiện HOLD sau commit;
  :Trả HoldSeatResponse;
  :Khách hàng xác nhận tạo đơn;
  :POST /api/v1/bookings;
  :Kiểm tra hold còn hiệu lực\nvà thuộc khách hàng hiện tại;
  if (Điều kiện giữ ghế hợp lệ?) then (Có)
    :Tạo Booking PENDING;\nTạo BookingDetail;\nLưu paymentExpiresAt;
    :Trả BookingResponse;
  else (Không)
    :Trả lỗi nghiệp vụ;\nKhông tạo đơn đặt vé;
  endif
else (Không)
  :Trả lỗi ghế không khả dụng;\nKhông cập nhật một phần;
endif
stop
@enduml

' [WORD: Hình 3.5] Luồng thanh toán VNPay
@startuml
title Sequence Diagram - Luồng thanh toán VNPay
autonumber
actor Customer
participant React as FE
participant PaymentController as CTRL
participant PaymentServiceImpl as SERVICE
participant VnPayPaymentGateway as GATEWAY
database PostgreSQL as DB
participant VNPay
participant "Ticket / Email / WS" as POST
Customer -> FE : choose VNPay
FE -> CTRL : POST /api/v1/payments/initiate
CTRL -> SERVICE : initiatePayment(bookingId, VNPAY, amount)
SERVICE -> DB : lock/validate booking and pending payment
SERVICE -> GATEWAY : create payment URL + secure hash
GATEWAY --> SERVICE : redirect URL
SERVICE -> DB : save/reuse PENDING payment + PAYMENT_EVENT
SERVICE --> FE : payment URL
FE -> VNPay : redirect browser
VNPay -> CTRL : GET /vnpay-callback
CTRL -> SERVICE : process VNPay callback
SERVICE -> GATEWAY : verify checksum / response / amount
alt verified success and first valid processing
  SERVICE -> DB : payment SUCCESS; booking SUCCESS; seats BOOKED
  SERVICE -> POST : create tickets, async email, WS BOOKED
else invalid/failed/expired/duplicate
  SERVICE -> DB : record event; apply safe status transition only
end
CTRL --> FE : redirect/result view
@enduml

' [WORD: Hình 3.6] Luồng thanh toán SePay/VietQR
@startuml
title Sequence Diagram - Luồng thanh toán SePay/VietQR
autonumber
actor Customer
participant React as FE
participant PaymentController as CTRL
participant PaymentServiceImpl as SERVICE
participant SePayPaymentGateway as GATEWAY
database PostgreSQL as DB
participant SePay
participant "Ticket / Email / WS" as POST
Customer -> FE : select bank QR payment
FE -> CTRL : POST /api/v1/payments/initiate
CTRL -> SERVICE : initiatePayment(bookingId, SEPAY, amount)
SERVICE -> DB : validate PENDING booking and amount
SERVICE -> GATEWAY : build QR account/amount/content
SERVICE -> DB : save PENDING payment + event
SERVICE --> FE : QR payment data
Customer -> SePay : bank transfer using QR
SePay -> CTRL : POST /sepay-webhook
CTRL -> SERVICE : process webhook
SERVICE -> GATEWAY : verify API key/HMAC
SERVICE -> DB : locate payment; compare content and amount
alt verified first success
  SERVICE -> DB : payment SUCCESS; booking SUCCESS; seats BOOKED
  SERVICE -> POST : tickets, email, WS
else mismatch/invalid/duplicate
  SERVICE -> DB : store PaymentEvent; do not issue ticket
end
FE -> CTRL : GET payment/booking result
CTRL --> FE : current confirmed status
@enduml

' [WORD: Hình 3.7] Luồng soát vé QR
@startuml
title Sequence Diagram - Luồng soát vé QR
autonumber
actor Staff
participant "React StaffTicketScannerPage" as FE
participant TicketController as CTRL
participant BookingServiceImpl as SERVICE
database PostgreSQL as DB
Staff -> FE : select assigned cinema and open showtime
FE -> CTRL : GET /tickets/check-in/showtimes?cinemaId
CTRL --> FE : open check-in showtimes in scope
Staff -> FE : scan camera/file QR
FE -> CTRL : POST /tickets/check-in\nqrCode, cinemaId, showtimeId
CTRL -> SERVICE : checkInTicket(...)
SERVICE -> DB : lock/load ticket, booking, showtime, cinema
SERVICE -> SERVICE : verify signature, ACTIVE, SUCCESS, scope, context, window
alt all valid
  SERVICE -> DB : ticket ACTIVE -> USED; set checkInTime/checkedInBy
  SERVICE --> CTRL : TicketResponse
else invalid/used/wrong cinema/wrong showtime/window
  SERVICE --> CTRL : domain error; ticket unchanged
end
CTRL --> FE : result message
@enduml

' [WORD: Hình 3.12] Đồng bộ sơ đồ ghế qua WebSocket
@startuml
title Sequence Diagram - Đồng bộ sơ đồ ghế qua WebSocket/STOMP
autonumber
actor "Khách A" as A
actor "Khách B" as B
participant "SeatSelectionPage A" as FEA
participant "SeatSelectionPage B\nuseSeatWebSocket" as FEB
participant "ShowtimeController" as SHOWTIME
participant "BookingServiceImpl" as BOOKING
participant "SeatStatusRepository" as REPO
database PostgreSQL as DB
participant "SeatStatusPublisher\nSTOMP broker" as WS
B -> FEB : mở màn hình chọn ghế
FEB -> SHOWTIME : GET /api/v1/showtimes/{id}/seats
SHOWTIME -> REPO : đọc seat map hiện tại
REPO --> DB : SELECT seat_status
SHOWTIME --> FEB : snapshot AVAILABLE/HOLD/BOOKED
FEB -> WS : CONNECT /ws-native\nSUBSCRIBE /topic/seatmap/{showtimeId}
A -> FEA : chọn ghế
FEA -> BOOKING : POST /api/v1/bookings/hold
BOOKING -> REPO : lock/cập nhật SeatStatus trong transaction
REPO -> DB : AVAILABLE -> HOLD
BOOKING -> WS : đăng ký publish sau commit
WS -> FEB : SeatStatusEvent(HOLD, heldBy, holdUntil)
FEB -> FEB : cập nhật đúng seat trong state UI
note over BOOKING,WS
Nếu transaction rollback, afterCommit không chạy;
khách B không nhận trạng thái HOLD giả.
end note
@enduml

' [WORD: Hình 3.13] Hủy suất chiếu và tạo yêu cầu hoàn tiền
@startuml
title Activity Diagram - Hủy suất chiếu và tạo yêu cầu hoàn tiền
skinparam shadowing false
skinparam linetype ortho
start
:Admin/Staff gửi POST /api/v1/showtimes/{id}/cancel;
:ShowtimeServiceImpl kiểm tra permission\nvà staff cinema scope;
if (Suất đang UPCOMING/ONGOING?) then (Có)
  if (Có ticket USED?) then (Có)
    :Từ chối hủy suất;
    stop
  else (Không)
    :Lấy Booking PENDING và SUCCESS\ncùng payment liên quan;
    :Giải phóng reservation khuyến mãi;
    while (Mỗi booking bị ảnh hưởng?) is (Còn)
      :Trả ghế AVAILABLE và phát event sau commit;
      :Hủy ticket ACTIVE nếu có;
      if (Booking ban đầu là SUCCESS?) then (Có)
        :Booking -> REFUND_PENDING;
        :Payment SUCCESS -> REFUND_PENDING;
        :Tạo Refund PENDING và PaymentEvent;
        :Đăng ký gửi email hủy suất sau commit;
      else (PENDING)
        :Booking -> CANCELLED;
      endif
    endwhile (Hết)
    :Payment PENDING -> FAILED\nvà ghi PaymentEvent;
:Showtime -> CANCELLED;
    :Lưu transaction;
    :Quản trị viên xử lý hoàn trả bên ngoài hệ thống;\nsau đó ghi nhận thành công hoặc thất bại\nqua RefundServiceImpl;
  endif
else (Không)
  :Từ chối vì suất không thể hủy;
endif
stop
@enduml

' [WORD: Hình 3.14] Xử lý yêu cầu hoàn tiền
@startuml
title Activity Diagram - Xử lý yêu cầu hoàn tiền
skinparam shadowing false
skinparam linetype ortho
start
:Người có thẩm quyền mở yêu cầu Refund PENDING;
:Kiểm tra permission, trạng thái Booking/Payment\nvà phạm vi rạp;
if (Yêu cầu còn hợp lệ?) then (Có)
  :Refund PENDING -> PROCESSING;
  :Thực hiện hoàn trả theo quy trình bên ngoài hệ thống;
  if (Người vận hành xác nhận đã hoàn trả?) then (Có)
    :Refund -> SUCCESS;
    :Payment -> REFUNDED;
    :Booking -> REFUNDED;
    :Ghi mã tham chiếu, thời điểm xử lý\nvà PaymentEvent;
  else (Không)
    :Refund -> FAILED;
    :Payment -> REFUND_FAILED;
    :Booking giữ REFUND_PENDING để theo dõi;
    :Ghi lý do thất bại và PaymentEvent;
  endif
else (Không)
  :Từ chối thao tác;\nkhông thay đổi trạng thái giao dịch;
endif
note right
CinemaBooking chưa gọi API hoàn tiền của provider.
SUCCESS/FAILED là kết quả do người vận hành ghi nhận.
end note
stop
@enduml

' [WORD: Hình 3.9] Sơ đồ trạng thái đơn đặt vé
@startuml
title State Diagram - Trạng thái đơn đặt vé
[*] --> PENDING : create after valid holds
PENDING --> SUCCESS : verified payment success
PENDING --> FAILED : payment failed
PENDING --> EXPIRED : payment timeout scheduler
PENDING --> CANCELLED : customer/operator cancel when policy permits
SUCCESS --> REFUND_PENDING : showtime cancellation/refund policy
REFUND_PENDING --> REFUNDED : operator records external refund success
note right of REFUND_PENDING
CinemaBooking không tự chuyển tiền qua provider API.
Nếu xử lý bên ngoài thất bại, Payment chuyển
REFUND_FAILED; Booking vẫn giữ REFUND_PENDING
để nhân viên tiếp tục theo dõi và xử lý.
end note
SUCCESS --> CANCELLED : cancellation policy if applicable
FAILED --> [*]
EXPIRED --> [*]
CANCELLED --> [*]
REFUNDED --> [*]
@enduml

' [WORD: Hình 3.10] Sơ đồ trạng thái giao dịch thanh toán
@startuml
title State Diagram - Trạng thái giao dịch thanh toán
[*] --> PENDING : initiate/reuse payment
PENDING --> SUCCESS : verified callback/webhook
PENDING --> FAILED : provider failure
PENDING --> EXPIRED : booking timeout
SUCCESS --> REFUND_PENDING : refund request created
REFUND_PENDING --> REFUNDED : operator records success
REFUND_PENDING --> REFUND_FAILED : operator records failure
note right of REFUND_PENDING
Các trạng thái phản ánh kết quả được ghi nhận;
không có provider refund API trong phiên bản hiện tại.
end note
FAILED --> [*]
EXPIRED --> [*]
REFUNDED --> [*]
REFUND_FAILED --> [*]
@enduml
'''


def write_plantuml() -> None:
    PLANTUML.write_text(plantuml_text(), encoding="utf-8")


def write_source_analysis(endpoints: list[Endpoint], counts: dict[str, int], calls: dict[str, list[str]]) -> None:
    service_trace = [
        ("Authentication", "AuthenticationController.java → AuthenticationService.java → UserRepository/RefreshTokenRepository/InvalidatedTokenRepository → users/refresh_tokens/invalidated_token/auth_audit_logs", "Password login, Google ID token verification, JWT issue/verify, refresh rotation, logout and session revoke."),
        ("Authorization", "SecurityConfig.java → CustomJwtDecoder.java → @PreAuthorize on Controller → StaffCinemaScopeService.java", "Filter-level authentication, permission-level access, then cinema scope for staff operations."),
        ("Seat hold/booking", "BookingController.java → BookingServiceImpl.java → SeatStatusRepository/BookingRepository → seat_status/bookings/booking_details", "Rate limit, expiry cleanup, lock/availability check, HOLD/PENDING state, transaction and after-commit event."),
        ("VNPay", "PaymentController.java → PaymentServiceImpl.java → VnPayPaymentGateway.java → PaymentRepository/PaymentEventRepository", "Initiate URL, checksum/amount verification, idempotent success/failure handling."),
        ("SePay", "PaymentController.java → PaymentServiceImpl.java → SePayPaymentGateway.java → PaymentEventServiceImpl", "Build bank QR, public webhook with API key/HMAC check, amount/content reconciliation."),
        ("Ticket/check-in", "TicketController.java → BookingServiceImpl.java → TicketRepository/ShowtimeService", "Signed QR, ACTIVE/SUCCESS/context/window/scope check before USED."),
        ("Realtime", "WebSocketConfig.java + SeatStatusPublisher.java → frontend useSeatWebSocket.ts", "STOMP topic /topic/seatmap/{showtimeId}; frontend only updates UI after event."),
        ("Schedulers", "HoldExpireScheduler.java; PendingBookingExpireScheduler.java; ShowtimeStatusSyncScheduler.java", "Release expired holds, expire pending bookings, synchronize showtime status."),
        ("Refund request/audit", "RefundServiceImpl.java; PaymentEventServiceImpl.java; AdminAuditLogInterceptor.java; AuthAuditService.java", "Internal refund-request lifecycle, operator-recorded result, payment-event history and operational/authentication trace; no provider refund API."),
    ]
    lines = [
        "# CinemaBooking Source Analysis",
        "",
        "## 1. Phạm vi và nguyên tắc", "",
        "Tài liệu được reverse-engineer từ source code hiện tại: backend Spring Boot, frontend React, Flyway migration, application.yaml, docker-compose.yml và test source. Không dùng báo cáo khóa luận cũ làm source of truth. Những điều source không thể chứng minh được đều được đánh dấu `[CHƯA XÁC MINH ĐƯỢC TỪ SOURCE CODE]`.", "",
        "## 2. Bản đồ source đã khảo sát", "",
        f"- Controllers: {counts['controllers']}", f"- Service/interface/implementation: {counts['services']}", f"- Repository/projection: {counts['repositories']}", f"- Entity: {counts['entities']}", f"- Flyway migrations: {counts['migrations']}", f"- React pages: {counts['frontend_pages']}", f"- Java test classes: {counts['tests']}", f"- REST endpoints extracted: {counts['endpoints']}", "",
        "## 3. Kiến trúc tổng thể", "",
        "`React Pages/Components → src/api Axios client → Spring Controller → Service → Repository/Projection → PostgreSQL`", "",
        "Các nhánh ngoài luồng REST thông thường: frontend subscribe STOMP `/topic/seatmap/{showtimeId}`; scheduler chạy trong backend; VNPay callback và SePay webhook quay lại PaymentController; EmailService chạy bất đồng bộ cho các thông báo phù hợp.", "",
        "## 4. Frontend React", "",
        "- `src/router/AppRouter.tsx`: lazy route, ba layout Auth/Admin/Public và `ProtectedRoute` theo permission.",
        "- `src/api/axiosClient.ts`: đính Bearer token; gộp các yêu cầu 401 vào một `refreshPromise`; retry đúng một lần sau refresh; logout khi refresh thất bại.",
        "- `src/stores/authStore.ts`: chỉ lưu access token, user, permission trong localStorage; refresh token bị xóa khỏi localStorage. Đây là bằng chứng frontend ưu tiên session refresh phía server/cookie (`withCredentials: true`).",
        "- `src/hooks/useSeatWebSocket.ts`: native WebSocket STOMP `/ws-native`, subscribe theo showtime, reconnect 3 giây, cleanup unmount và giữ callback bằng ref để tránh resubscribe thừa.",
        "- `src/pages/public/*`, `src/pages/user/*`, `src/pages/admin/*`, `src/pages/staff/*`: tách màn hình public, customer, admin và staff.",
        "- `src/components/RegionalShowtimeBrowser.tsx`: dùng React Query cho truy vấn lịch chiếu vùng; `App.tsx` có ErrorBoundary toàn app.", "",
        "## 5. Backend Spring Boot", "",
        "- `configuration/SecurityConfig.java`: permit có chủ đích cho browse public, auth public, callback/webhook và WebSocket handshake; còn lại yêu cầu authentication. `@EnableMethodSecurity` bật `@PreAuthorize`.",
        "- `security/service/AuthenticationService.java`: password BCrypt, Google ID token, claim `token_use`, `auth_version`, token invalidation, refresh token hash/rotation, rate limit và audit.",
        "- `service/impl/BookingServiceImpl.java`: owner check, hold/booking/cancel/promotion, ticket and check-in; là trung tâm lifecycle booking/seat.",
        "- `service/impl/PaymentServiceImpl.java`: điều phối payment methods, gateway validation, idempotency, event recording and post-payment side effects.",
        "- `payment/*PaymentGateway.java`: adapter theo provider. VnPay/SePay được cấu hình; Momo class tồn tại nhưng default disabled.",
        "- `exception/GlobalExceptionHandler.java`: chuẩn hóa lỗi cho client. DTO dùng Bean Validation.",
        "- `configuration/CacheConfig.java` và service annotations: Caffeine cho dữ liệu đọc ít đổi; không cache seat status.",
        "- `security/task/*Scheduler.java`: cleanup and time-derived state synchronization.", "",
        "## 6. Database và migration", "",
        "Flyway được bật với `classpath:db/migration`, `baseline-on-migrate=true`, `validate-on-migrate=true`; JPA `ddl-auto=validate`, do đó schema được kiểm soát bằng migration thay vì JPA tự tạo production schema.",
        "", "### 6.1. Domain entities", "", ", ".join(entity_names()) + ".", "",
        "### 6.2. Tính toàn vẹn và index", "",
        "- `seat_status` có unique `(seat_id, showtime_id)`, CHECK cho HOLD và `version` để bảo vệ cạnh tranh.",
        "- `bookings`, `payments`, `tickets`, `refunds` có trạng thái CHECK/enum, foreign key và index theo truy vấn vận hành.",
        "- V13/V14 bổ sung partial unique pending payment/booking để chống nhân đôi khi double click, callback chậm hoặc retry.",
        "- Index theo thời gian cho scheduler, payment events, audit, showtimes và filters admin nằm trong V1, V2, V3, V4, V5, V9, V10, V11, V13, V14.", "",
        "## 7. Traceability các luồng trọng yếu", "",
        "| Luồng | Chuỗi file/class | Điều được xác nhận |", "| --- | --- | --- |",
    ]
    for name, trace, evidence in service_trace:
        lines.append(f"| {name} | `{trace}` | {evidence} |")
    lines.extend([
        "", "## 8. Concurrency và race condition", "",
        "`SeatStatus` là nguồn trạng thái ghế theo suất. Luồng hold/confirm được bọc transaction trong service, sử dụng repository lock/query và kiểm tra trạng thái; schema có unique seat-showtime, version và partial unique pending booking/payment. Response 'seat not available' khi cạnh tranh là lỗi nghiệp vụ dự kiến, không phải lỗi hệ thống. Scheduler thu hồi HOLD/PENDING hết hạn; publisher chỉ broadcast sau commit để client không thấy dữ liệu rollback.",
        "", "## 9. Payment và idempotency", "",
        "VNPay và SePay có flow riêng. VNPay dựa callback/checksum; SePay dựa webhook/API key-HMAC/amount-content. PaymentEvent lưu event nhằm trace webhook, signature failure, mismatch, reused/processed payment và thay đổi trạng thái yêu cầu refund. RefundServiceImpl chỉ tạo yêu cầu và ghi nhận kết quả do operator cung cấp; source không có lời gọi provider refund API. Không có bằng chứng source về hàng đợi message broker hoặc reconciliation worker bên ngoài process; source hiện dùng repository/service/scheduler/event table.",
        "", "## 10. Phân loại mức độ hoàn thiện", "",
        "| Trạng thái | Hạng mục | Căn cứ |", "| --- | --- | --- |",
        "| Đã hiện thực | Auth/JWT/refresh/session, RBAC, staff scope, catalog, seat hold, booking, promotion, VNPay, SePay, ticket QR/check-in, email, WebSocket, scheduler, quản lý yêu cầu refund, audit, cache, tests | Controller/service/entity/migration/frontend/test tương ứng; refund chỉ ở mức yêu cầu/trạng thái do operator ghi nhận. |",
        "| Có cấu trúc nhưng cần cấu hình/kiểm thử triển khai | Google login, SMTP, VNPay, SePay, MoMo gateway | Có class/config; hoạt động thực tế phụ thuộc secrets/provider. `MOMO_ENABLED=false` mặc định. |",
        "| Hướng phát triển | Distributed cache/broker, queue, provider refund API tự động, CI/CD, production monitoring/metrics, kiểm thử tải JMeter dài hạn và xác định capacity | [CHƯA XÁC MINH ĐƯỢC TỪ SOURCE CODE]. |",
        "", "## 11. API consistency findings", "",
        "- Inventory được trích từ controller, không tự đặt endpoint.",
        "- Frontend API clients dùng cùng base route `/api/v1` và `/auth`; route bảo vệ UI bổ sung cho server authorization, không thay thế server authorization.",
        "- Callback/webhook không có caller frontend là đúng thiết kế: gateway bên ngoài mới gọi các endpoint đó.",
        "- Có hai endpoint đọc sơ đồ ghế cùng gọi `bookingService.getSeatMap`; frontend đang gọi `/api/v1/showtimes/{id}/seats`. Route alias dưới BookingController chưa có bằng chứng deprecation trong source, vì vậy cần chọn route chuẩn và deprecate có kiểm soát trước khi xóa alias.",
        "- Có hai endpoint check-in cùng đi vào `bookingService.checkInTicket`; frontend gọi `/api/v1/tickets/check-in`. Đây là điểm cần hợp nhất API theo ticket domain, giữ alias có thời hạn nếu cần tương thích client cũ.",
        "- `movieApi.getAll` không truyền keyword xuống backend trong file hiện tại; comment cho biết keyword có thể lọc client-side. Đây là quyết định UX/client hiện có, không phải endpoint mismatch, nhưng cần cân nhắc backend search khi catalog lớn.",
        "- [CHƯA XÁC MINH ĐƯỢC TỪ SOURCE CODE] tất cả caller gián tiếp/dynamic của React ngoài `src/api` vì inventory chỉ dò static API client theo yêu cầu.",
        "", "## 12. Danh sách test source", "",
    ])
    lines.extend([f"- `{name}`" for name in test_names()])
    SOURCE_ANALYSIS.write_text("\n".join(lines) + "\n", encoding="utf-8")


def validate_prerequisites() -> None:
    missing = [path for path in (ROOT / "pom.xml", ROOT / "src" / "main" / "resources" / "application.yaml", BACKEND_JAVA / "controller", MIGRATIONS) if not path.exists()]
    if missing:
        raise RuntimeError("Missing required source paths: " + ", ".join(str(path) for path in missing))


def main() -> None:
    validate_prerequisites()
    endpoints = parse_controllers()
    if not endpoints:
        raise RuntimeError("No controller endpoints were extracted; source parser needs review.")
    calls = frontend_calls()
    counts = source_counts()
    write_api_inventory(endpoints, calls)
    write_source_analysis(endpoints, counts, calls)
    write_plantuml()
    create_report(endpoints, counts)
    print(f"Created {REPORT.name}")
    print(f"Created {PLANTUML.name}")
    print(f"Created {SOURCE_ANALYSIS.name}")
    print(f"Created {API_INVENTORY.name}")
    print(f"Extracted endpoints: {len(endpoints)}")


if __name__ == "__main__":
    main()

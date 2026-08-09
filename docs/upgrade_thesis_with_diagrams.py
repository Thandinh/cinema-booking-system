from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DIAGRAMS = DOCS / "diagrams"
OUT = DOCS / "Bao_cao_khoa_luan_CinemaBooking_Final_Submission.docx"
SOURCE = DOCS / "Bao_cao_khoa_luan_CinemaBooking_Full.docx"

DIAGRAMS.mkdir(parents=True, exist_ok=True)

PURPLE = "3B2496"
ORANGE = "F97316"
BLUE = "1D4ED8"
DARK = "111827"
MUTED = "64748B"
LINE = "CBD5E1"
LIGHT = "F8FAFC"
GREEN = "059669"
RED = "DC2626"
YELLOW = "D97706"


def font_path(name: str) -> str:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/Arial.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/times.ttf"),
    ]
    for p in candidates:
        if p.exists():
            return str(p)
    return ""


FONT_REG = font_path("arial.ttf")
FONT_BOLD = font_path("arialbd.ttf")


def pil_font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_BOLD if bold and FONT_BOLD else FONT_REG, size=size)


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.strip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def wrap_text(text: str, width: int) -> list[str]:
    return textwrap.wrap(text, width=width, break_long_words=False) or [""]


def draw_text_center(draw: ImageDraw.ImageDraw, xy, text: str, font, fill="#111827", width_chars=18):
    x1, y1, x2, y2 = xy
    lines = wrap_text(text, width_chars)
    line_h = font.size + 6
    total_h = line_h * len(lines)
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=hex_to_rgb(fill))
        y += line_h


def draw_box(draw, xy, text, fill="FFFFFF", outline=BLUE, text_color=DARK, radius=24, bold=True, width_chars=18):
    draw.rounded_rectangle(xy, radius=radius, fill=hex_to_rgb(fill), outline=hex_to_rgb(outline), width=3)
    draw_text_center(draw, xy, text, pil_font(26 if bold else 23, bold), text_color, width_chars)


def arrow(draw, start, end, color=MUTED, width=4):
    draw.line([start, end], fill=hex_to_rgb(color), width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - sign * 16, ey - 10), (ex - sign * 16, ey + 10)]
    else:
        sign = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 10, ey - sign * 16), (ex + 10, ey - sign * 16)]
    draw.polygon(pts, fill=hex_to_rgb(color))


def draw_note(draw, xy, title: str, lines: list[str], fill="FFFBEB", outline=YELLOW):
    draw.rounded_rectangle(xy, radius=20, fill=hex_to_rgb(fill), outline=hex_to_rgb(outline), width=3)
    x1, y1, x2, _ = xy
    draw.text((x1 + 24, y1 + 18), title, font=pil_font(23, True), fill=hex_to_rgb(DARK))
    y = y1 + 58
    for line in lines:
        draw.text((x1 + 24, y), line, font=pil_font(19, False), fill=hex_to_rgb(MUTED))
        y += 30


def legend(draw, items: list[tuple[str, str]], x: int, y: int):
    for label, color in items:
        draw.rounded_rectangle([x, y, x + 34, y + 18], radius=9, fill=hex_to_rgb(color))
        draw.text((x + 46, y - 5), label, font=pil_font(18, False), fill=hex_to_rgb(MUTED))
        y += 34


def message_color(message: str) -> str:
    lower = message.lower()
    if any(key in lower for key in ["success", "thành công", "confirm", "booked", "used", "active"]):
        return GREEN
    if any(key in lower for key in ["failed", "expired", "hủy", "hoàn", "release", "trả ghế"]):
        return RED
    if any(key in lower for key in ["verify", "kiểm tra", "callback", "webhook", "hash", "hmac"]):
        return PURPLE
    if any(key in lower for key in ["websocket", "topic", "publish", "email", "smtp"]):
        return BLUE
    return ORANGE


def base_canvas(title: str, w=1800, h=1100):
    img = Image.new("RGB", (w, h), hex_to_rgb("FFFFFF"))
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, w, 92], fill=hex_to_rgb("EEF2FF"))
    draw.text((44, 24), title, font=pil_font(34, True), fill=hex_to_rgb(PURPLE))
    draw.text((w - 380, 30), "CinemaBooking.vn", font=pil_font(24, True), fill=hex_to_rgb(ORANGE))
    return img, draw


def save(img: Image.Image, name: str) -> Path:
    path = DIAGRAMS / name
    img.save(path, quality=95)
    return path


def context_diagram() -> Path:
    img, d = base_canvas("Sơ đồ ngữ cảnh hệ thống")
    center_box = [660, 410, 1140, 650]
    draw_box(d, center_box, "CinemaBooking.vn\nReact + Spring Boot", fill="FFF7ED", outline=ORANGE, width_chars=22)
    actors = [
        ([80, 180, 390, 310], "Khách hàng", "DBEAFE"),
        ([80, 760, 390, 890], "Nhân viên rạp", "DCFCE7"),
        ([1410, 180, 1720, 310], "Admin", "FEE2E2"),
        ([1410, 760, 1720, 890], "Database\nPostgreSQL", "E0F2FE"),
        ([720, 145, 1080, 255], "Google OAuth", "F1F5F9"),
        ([430, 800, 650, 920], "VNPay", "EDE9FE"),
        ([720, 800, 1080, 920], "SePay / VietQR", "FEF3C7"),
        ([1160, 800, 1380, 920], "SMTP Mail", "ECFDF5"),
    ]
    centers = []
    for xy, text, fill in actors:
        draw_box(d, xy, text, fill=fill, outline=LINE, bold=True, width_chars=16)
        centers.append((xy, ((xy[0] + xy[2]) // 2, (xy[1] + xy[3]) // 2)))

    def target_for(point):
        px, py = point
        cx1, cy1, cx2, cy2 = center_box
        if py < cy1:
            return ((cx1 + cx2) // 2, cy1)
        if py > cy2:
            return ((cx1 + cx2) // 2, cy2)
        if px < cx1:
            return (cx1, (cy1 + cy2) // 2)
        return (cx2, (cy1 + cy2) // 2)

    def source_for(xy):
        x1, y1, x2, y2 = xy
        px, py = ((x1 + x2) // 2, (y1 + y2) // 2)
        cx1, cy1, cx2, cy2 = center_box
        if py < cy1:
            return (px, y2)
        if py > cy2:
            return (px, y1)
        if px < cx1:
            return (x2, py)
        return (x1, py)

    for xy, p in centers:
        arrow(d, source_for(xy), target_for(p))
    return save(img, "01_so_do_ngu_canh.png")


def bfd_diagram() -> Path:
    img, d = base_canvas("Sơ đồ phân rã chức năng BFD")
    draw_box(d, [620, 135, 1180, 250], "Hệ thống đặt vé xem phim CinemaBooking.vn", fill="FFF7ED", outline=ORANGE, width_chars=30)
    modules = [
        ("Quản lý tài khoản", 100, 350), ("Quản lý phim", 430, 350), ("Rạp / Phòng / Ghế", 760, 350),
        ("Suất chiếu", 1090, 350), ("Đặt vé", 1420, 350), ("Thanh toán", 100, 620),
        ("Gửi email", 430, 620), ("Soát vé QR", 760, 620), ("Phân quyền", 1090, 620), ("Thống kê / Audit", 1420, 620),
    ]
    for text, x, y in modules:
        draw_box(d, [x, y, x + 280, y + 130], text, fill=LIGHT, outline=BLUE, width_chars=16)
        arrow(d, (900, 250), (x + 140, y))
    return save(img, "02_bfd.png")


def use_case_diagram() -> Path:
    img, d = base_canvas("Sơ đồ Use Case tổng quát", w=2200, h=1250)
    system = [470, 150, 1730, 1110]
    d.rounded_rectangle(system, radius=34, fill=hex_to_rgb("FFFFFF"), outline=hex_to_rgb(PURPLE), width=4)
    d.text((890, 180), "Hệ thống CinemaBooking.vn", font=pil_font(31, True), fill=hex_to_rgb(PURPLE))

    def actor(label: str, x: int, y: int, fill: str):
        d.rounded_rectangle([x, y, x + 285, y + 106], radius=24, fill=hex_to_rgb(fill), outline=hex_to_rgb(LINE), width=3)
        d.ellipse([x + 22, y + 28, x + 70, y + 76], outline=hex_to_rgb(DARK), width=4)
        d.text((x + 90, y + 35), label, font=pil_font(26, True), fill=hex_to_rgb(DARK))
        return (x + 285 if x < system[0] else x, y + 53)

    def use_case(label: str, x: int, y: int, fill: str, outline: str):
        d.ellipse([x, y, x + 300, y + 86], fill=hex_to_rgb(fill), outline=hex_to_rgb(outline), width=3)
        draw_text_center(d, [x, y, x + 300, y + 86], label, pil_font(21, True), DARK, 22)
        return (x, y + 43), (x + 300, y + 43)

    actors = {
        "USER": actor("USER", 95, 285, "DBEAFE"),
        "STAFF": actor("STAFF", 95, 735, "DCFCE7"),
        "ADMIN": actor("ADMIN", 1820, 285, "FEE2E2"),
        "EXT": actor("Dịch vụ ngoài", 1820, 735, "F8FAFC"),
    }

    d.text((560, 270), "Khách hàng", font=pil_font(24, True), fill=hex_to_rgb(ORANGE))
    d.text((560, 650), "Nhân viên rạp", font=pil_font(24, True), fill=hex_to_rgb(GREEN))
    d.text((1185, 270), "Quản trị", font=pil_font(24, True), fill=hex_to_rgb(RED))
    d.text((1185, 695), "Tích hợp ngoài", font=pil_font(24, True), fill=hex_to_rgb(PURPLE))

    cases = {
        "browse": use_case("Xem phim, rạp, lịch chiếu", 560, 320, "EFF6FF", BLUE),
        "booking": use_case("Đặt vé và thanh toán", 560, 440, "EFF6FF", BLUE),
        "profile": use_case("Xem vé, hồ sơ cá nhân", 560, 560, "EFF6FF", BLUE),
        "scan": use_case("Soát vé QR", 560, 700, "ECFDF5", GREEN),
        "staff_showtime": use_case("Quản lý suất theo rạp phụ trách", 560, 820, "ECFDF5", GREEN),
        "staff_scope": use_case("Xem dữ liệu trong phạm vi rạp", 560, 940, "ECFDF5", GREEN),
        "admin_data": use_case("Quản lý dữ liệu hệ thống", 1185, 320, "FEF2F2", RED),
        "rbac": use_case("Phân quyền và gán rạp staff", 1185, 440, "FEF2F2", RED),
        "dashboard": use_case("Dashboard, báo cáo, audit", 1185, 560, "FEF2F2", RED),
        "google": use_case("Đăng nhập Google", 1185, 740, "F5F3FF", PURPLE),
        "payment": use_case("Xác nhận VNPay / SePay", 1185, 860, "F5F3FF", PURPLE),
        "mail": use_case("Gửi email xác thực / vé", 1185, 980, "F5F3FF", PURPLE),
    }

    def connect_left(actor_key: str, case_key: str, color: str):
        src = actors[actor_key]
        dst = cases[case_key][0]
        mid_x = system[0] - 58
        d.line([src, (mid_x, src[1]), (mid_x, dst[1]), dst], fill=hex_to_rgb(color), width=3)

    def connect_right(actor_key: str, case_key: str, color: str):
        src = actors[actor_key]
        dst = cases[case_key][1]
        mid_x = system[2] + 58
        d.line([src, (mid_x, src[1]), (mid_x, dst[1]), dst], fill=hex_to_rgb(color), width=3)

    for key in ["browse", "booking", "profile"]:
        connect_left("USER", key, BLUE)
    for key in ["scan", "staff_showtime", "staff_scope"]:
        connect_left("STAFF", key, GREEN)
    for key in ["admin_data", "rbac", "dashboard"]:
        connect_right("ADMIN", key, RED)
    for key in ["google", "payment", "mail"]:
        connect_right("EXT", key, PURPLE)

    draw_note(
        d,
        [720, 1098, 1495, 1238],
        "Ghi chú",
        ["Các actor chỉ nối tới nhóm chức năng liên quan để tránh giao cắt đường nối.", "Staff bị giới hạn theo rạp được admin phân công."],
        fill="FFFBEB",
        outline=YELLOW,
    )
    return save(img, "03_use_case.png")


def dfd0_diagram() -> Path:
    img, d = base_canvas("DFD Level 0 - luồng dữ liệu tổng quát", w=2200, h=1250)
    legend(d, [
        ("Dữ liệu từ actor", ORANGE),
        ("Phản hồi từ hệ thống", BLUE),
        ("Đọc/ghi database", GREEN),
        ("Tích hợp dịch vụ ngoài", PURPLE),
    ], 90, 1005)

    d.text((120, 145), "Actor", font=pil_font(25, True), fill=hex_to_rgb(ORANGE))
    d.text((870, 145), "Process chính", font=pil_font(25, True), fill=hex_to_rgb(BLUE))
    d.text((1630, 145), "Kho dữ liệu / hệ thống ngoài", font=pil_font(25, True), fill=hex_to_rgb(GREEN))

    process = [820, 310, 1380, 835]
    draw_box(d, process, "P0\nCinemaBooking.vn\nĐặt vé - thanh toán - quản trị", fill="FFF7ED", outline=ORANGE, width_chars=25)

    actors = [
        ("USER\nTìm phim, đặt vé,\nthanh toán, xem vé", [90, 250, 455, 390], 320),
        ("STAFF\nVận hành suất chiếu,\nsoát vé QR", [90, 510, 455, 650], 580),
        ("ADMIN\nQuản trị dữ liệu,\nphân quyền, thống kê", [90, 770, 455, 910], 840),
    ]
    for text, xy, y in actors:
        draw_box(d, xy, text, fill="F8FAFC", outline=BLUE, width_chars=22)
        arrow(d, (xy[2], y - 18), (process[0], y - 18), ORANGE, 4)
        arrow(d, (process[0], y + 18), (xy[2], y + 18), BLUE, 4)

    stores = [
        ("D1 Database\nPostgreSQL\nusers, movies, bookings,\npayments, tickets, audit", [1630, 300, 2085, 500], 400, GREEN),
        ("E1 Dịch vụ ngoài\nGoogle OAuth, VNPay,\nSePay/VietQR, SMTP", [1630, 650, 2085, 850], 750, PURPLE),
    ]
    for text, xy, y, color in stores:
        draw_box(d, xy, text, fill="F8FAFC", outline=color, width_chars=25)
        arrow(d, (process[2], y - 18), (xy[0], y - 18), color, 4)
        arrow(d, (xy[0], y + 18), (process[2], y + 18), BLUE, 4)

    d.text((520, 275), "Request / thao tác", font=pil_font(19, True), fill=hex_to_rgb(ORANGE))
    d.text((520, 352), "Kết quả / trạng thái", font=pil_font(19, True), fill=hex_to_rgb(BLUE))
    d.text((1410, 352), "Đọc/ghi dữ liệu", font=pil_font(19, True), fill=hex_to_rgb(GREEN))
    d.text((1410, 702), "Callback / webhook / email", font=pil_font(19, True), fill=hex_to_rgb(PURPLE))
    draw_note(
        d,
        [760, 980, 1560, 1155],
        "Cách đọc sơ đồ",
        [
            "DFD Level 0 chỉ thể hiện hệ thống như một process duy nhất.",
            "Các luồng chi tiết hơn được tách sang DFD Level 1.",
            "Bố cục LR giúp tránh mũi tên giao cắt và dễ trình bày khi bảo vệ.",
        ],
        fill="FFFBEB",
        outline=YELLOW,
    )
    return save(img, "04_dfd_level_0.png")


def dfd1_diagram() -> Path:
    img, d = base_canvas("DFD Level 1 - các luồng nghiệp vụ chính", w=2200, h=1500)
    legend(d, [
        ("Request từ frontend", ORANGE),
        ("Dữ liệu nghiệp vụ", GREEN),
        ("Tích hợp ngoài", PURPLE),
        ("Phản hồi/kết quả", BLUE),
    ], 90, 1060)
    d.text((95, 140), "Actor / Client", font=pil_font(25, True), fill=hex_to_rgb(ORANGE))
    d.text((610, 140), "Process trong backend", font=pil_font(25, True), fill=hex_to_rgb(BLUE))
    d.text((1300, 140), "Data store / External", font=pil_font(25, True), fill=hex_to_rgb(GREEN))

    actor_boxes = [
        ("USER\nĐăng ký, đặt vé,\nthanh toán, xem vé", [80, 210, 390, 370]),
        ("STAFF\nQuản lý suất trong rạp,\nsoát vé QR", [80, 505, 390, 665]),
        ("ADMIN\nQuản trị dữ liệu,\nphân quyền, audit", [80, 800, 390, 960]),
    ]
    for text, box in actor_boxes:
        draw_box(d, box, text, fill="FFF7ED", outline=ORANGE, width_chars=22)

    proc_boxes = [
        ("P1 Auth\nĐăng ký, xác thực email,\nlogin, refresh, logout", [560, 190, 970, 340]),
        ("P2 Booking\nGiữ ghế, tạo đơn,\náp mã giảm giá", [560, 410, 970, 560]),
        ("P3 Payment\nVNPay callback,\nSePay webhook", [560, 630, 970, 780]),
        ("P4 Ticket\nSinh QR, gửi email,\nsoát vé", [560, 850, 970, 1000]),
        ("P5 Admin/Staff\nCRUD, staff scope,\nthống kê, audit", [560, 1070, 970, 1220]),
    ]
    for text, box in proc_boxes:
        draw_box(d, box, text, fill=LIGHT, outline=BLUE, width_chars=25)

    data_boxes = [
        ("D1 Auth DB\nusers, roles,\npermissions, tokens", [1240, 190, 1610, 340], "DBEAFE"),
        ("D2 Booking DB\nshowtimes, seats,\nseat_status, bookings", [1240, 410, 1610, 560], "DCFCE7"),
        ("D3 Payment DB\npayments,\npayment_events", [1240, 630, 1610, 780], "EDE9FE"),
        ("D4 Ticket DB\ntickets,\ncheck-in data", [1240, 850, 1610, 1000], "DCFCE7"),
        ("D5 Operation DB\nstaff_cinemas,\naudit logs", [1240, 1070, 1610, 1220], "FEF3C7"),
        ("Google OAuth\nđăng nhập bằng Google", [1760, 190, 2080, 340], "F8FAFC"),
        ("VNPay / SePay\ncallback, webhook", [1760, 630, 2080, 780], "F8FAFC"),
        ("SMTP Mail\ngửi xác thực và vé", [1760, 850, 2080, 1000], "F8FAFC"),
    ]
    for text, box, fill in data_boxes:
        draw_box(d, box, text, fill=fill, outline=GREEN if fill != "F8FAFC" else PURPLE, width_chars=24)

    def right_mid(box): return (box[2], (box[1] + box[3]) // 2)
    def left_mid(box): return (box[0], (box[1] + box[3]) // 2)
    def bottom_mid(box): return ((box[0] + box[2]) // 2, box[3])
    def top_mid(box): return ((box[0] + box[2]) // 2, box[1])

    user, staff, admin = [b for _, b in actor_boxes]
    p1, p2, p3, p4, p5 = [b for _, b in proc_boxes]
    d1, d2, d3, d4, d5, google, pay_gateways, smtp = [b for _, b, _ in data_boxes]

    for start, end, color in [
        (right_mid(user), left_mid(p1), ORANGE),
        (right_mid(user), left_mid(p2), ORANGE),
        (right_mid(user), left_mid(p3), ORANGE),
        (right_mid(user), left_mid(p4), ORANGE),
        (right_mid(staff), left_mid(p4), ORANGE),
        (right_mid(staff), left_mid(p5), ORANGE),
        (right_mid(admin), left_mid(p5), ORANGE),
        (right_mid(p1), left_mid(d1), GREEN),
        (right_mid(p2), left_mid(d2), GREEN),
        (right_mid(p3), left_mid(d3), GREEN),
        (right_mid(p4), left_mid(d4), GREEN),
        (right_mid(p5), left_mid(d5), GREEN),
        (right_mid(d1), left_mid(google), PURPLE),
        (right_mid(d3), left_mid(pay_gateways), PURPLE),
        (right_mid(d4), left_mid(smtp), PURPLE),
        (bottom_mid(p2), top_mid(p3), BLUE),
        (bottom_mid(p3), top_mid(p4), BLUE),
    ]:
        arrow(d, start, end, color, 4)

    draw_note(d, [1240, 1260, 2080, 1440], "Cách đọc sơ đồ", [
        "Các actor gửi request vào process.",
        "Process cập nhật data store tương ứng.",
        "Dịch vụ ngoài chỉ đi qua backend.",
        "Sơ đồ tách tầng để tránh vạch đè chữ."
    ], fill="FFFBEB", outline=YELLOW)
    return save(img, "05_dfd_level_1.png")


def erd_diagram() -> Path:
    img, d = base_canvas("ERD tổng quan cơ sở dữ liệu", w=2300, h=1500)

    def entity(name: str, fields: list[str], x: int, y: int, fill: str, w: int = 320, h: int = 210):
        d.rounded_rectangle([x, y, x+w, y+h], radius=22, fill=hex_to_rgb(fill), outline=hex_to_rgb(LINE), width=3)
        d.rectangle([x, y, x+w, y+48], fill=hex_to_rgb(PURPLE))
        d.text((x+18, y+10), name, font=pil_font(22, True), fill=hex_to_rgb("FFFFFF"))
        fy = y + 62
        for f in fields:
            d.text((x+18, fy), f, font=pil_font(20), fill=hex_to_rgb(DARK))
            fy += 28
        return {
            "name": name,
            "bounds": (x, y, x + w, y + h),
            "left": (x, y+h//2),
            "right": (x+w, y+h//2),
            "top": (x+w//2, y),
            "bottom": (x+w//2, y+h),
        }

    def group(x1, y1, x2, y2, title):
        d.rounded_rectangle([x1, y1, x2, y2], radius=28, fill=hex_to_rgb("FFFFFF"), outline=hex_to_rgb("E2E8F0"), width=3)
        d.text((x1+24, y1+18), title, font=pil_font(28, True), fill=hex_to_rgb(PURPLE))

    group(40, 120, 2260, 430, "Cụm xác thực và phân quyền")
    group(40, 470, 2260, 790, "Cụm dữ liệu rạp phim")
    group(40, 830, 2260, 1300, "Cụm đặt vé, thanh toán và vận hành")

    e = {}
    # Auth/RBAC cluster. Junction tables are displayed explicitly to avoid ambiguous N-N lines.
    e["users"] = entity("users", ["PK id", "username, email", "password_hash", "email_verified"], 80, 190, "DBEAFE", 270, 190)
    e["users_roles"] = entity("users_roles", ["FK user_id", "FK role_id"], 390, 190, "F8FAFC", 240, 190)
    e["roles"] = entity("roles", ["PK id", "name"], 670, 190, "F1F5F9", 240, 190)
    e["roles_permissions"] = entity("roles_permissions", ["FK role_id", "FK permission_id"], 950, 190, "F8FAFC", 270, 190)
    e["permissions"] = entity("permissions", ["PK id", "name"], 1260, 190, "F1F5F9", 270, 190)
    e["refresh"] = entity("refresh_tokens", ["PK id", "FK user_id", "token_hash", "revoked"], 1570, 190, "F8FAFC", 270, 190)
    e["staff_cinemas"] = entity("staff_cinemas", ["FK user_id", "FK cinema_id", "assigned_at"], 1880, 190, "DCFCE7", 300, 190)

    # Cinema master cluster.
    e["cinemas"] = entity("cinemas", ["PK id", "name, city", "address", "lat/lng"], 80, 550, "DBEAFE", 300, 190)
    e["rooms"] = entity("rooms", ["PK id", "FK cinema_id", "name"], 440, 550, "DBEAFE", 300, 190)
    e["seats"] = entity("seats", ["PK id", "FK room_id", "row/number", "seat_type"], 800, 550, "DBEAFE", 300, 190)
    e["movies"] = entity("movies", ["PK id", "title", "duration", "status"], 1160, 550, "FEF3C7", 300, 190)
    e["showtimes"] = entity("showtimes", ["PK id", "FK movie_id", "FK room_id", "start/end/status"], 1520, 550, "FEF3C7", 330, 190)
    e["seat_status"] = entity("seat_status", ["PK id", "FK seat_id", "FK showtime_id", "status, hold_until"], 1900, 550, "FEE2E2", 300, 190)

    # Transaction cluster.
    e["bookings"] = entity("bookings", ["PK id", "FK user_id", "FK showtime_id", "status, total"], 80, 940, "FFF7ED", 300, 190)
    e["booking_details"] = entity("booking_details", ["PK id", "FK booking_id", "FK seat_id", "price_at_booking"], 440, 940, "FFF7ED", 330, 190)
    e["tickets"] = entity("tickets", ["PK id", "FK booking_detail_id", "qr_code", "status, check_in"], 830, 940, "DCFCE7", 310, 190)
    e["payments"] = entity("payments", ["PK id", "FK booking_id", "method/status", "transaction_no"], 1200, 940, "EDE9FE", 310, 190)
    e["payment_events"] = entity("payment_events", ["PK id", "FK payment_id", "provider", "payload/status"], 1570, 940, "EDE9FE", 310, 190)
    e["audit"] = entity("audit logs", ["admin_audit_logs", "auth_audit_logs", "actor/action/time"], 1940, 940, "F8FAFC", 260, 190)

    def connect(a, point_a, b, point_b, label=""):
        ax1, ay1, ax2, ay2 = e[a]["bounds"]
        bx1, by1, bx2, by2 = e[b]["bounds"]
        route_y = max(ay2, by2) + 26
        x1 = ax2
        x2 = bx1
        if x2 < x1:
            x1 = ax1
            x2 = bx2
        arrow(d, (x1, route_y), (x2, route_y), MUTED, 3)
        if label:
            mx = (x1 + x2) // 2
            d.rectangle([mx-48, route_y-18, mx+48, route_y+14], fill=hex_to_rgb("FFFFFF"))
            d.text((mx-42, route_y-17), label, font=pil_font(17, True), fill=hex_to_rgb(MUTED))

    for a, b, label in [
        ("users", "users_roles", "1-N"), ("users_roles", "roles", "N-1"), ("roles", "roles_permissions", "1-N"),
        ("roles_permissions", "permissions", "N-1"), ("users", "refresh", "1-N"), ("users", "staff_cinemas", "1-N"),
        ("cinemas", "rooms", "1-N"), ("rooms", "seats", "1-N"), ("movies", "showtimes", "1-N"),
        ("showtimes", "seat_status", "1-N"), ("seats", "seat_status", "1-N"),
        ("bookings", "booking_details", "1-N"), ("booking_details", "tickets", "1-1"),
        ("bookings", "payments", "1-N"), ("payments", "payment_events", "1-N"), ("payments", "audit", "log"),
    ]:
        connect(a, "right", b, "left", label)
    d.text((70, 1335), "Ghi chú: các FK liên cụm như bookings.user_id, bookings.showtime_id, staff_cinemas.cinema_id được ghi trong từng entity để sơ đồ không bị rối khi in trong báo cáo.", font=pil_font(22), fill=hex_to_rgb(MUTED))
    return save(img, "06_erd.png")


def activity_diagram(name: str, title: str, steps: list[tuple[str, str]]) -> Path:
    img, d = base_canvas(title, h=1200)
    x = 560
    y = 145
    d.ellipse([830, y, 970, y+70], fill=hex_to_rgb(GREEN), outline=hex_to_rgb(GREEN))
    draw_text_center(d, [830, y, 970, y+70], "Bắt đầu", pil_font(20, True), "FFFFFF", 10)
    prev = (900, y+70)
    y += 125
    for label, color in steps:
        if color == "decision":
            points = [(900, y), (1130, y+90), (900, y+180), (670, y+90)]
            d.polygon(points, fill=hex_to_rgb("FEF3C7"), outline=hex_to_rgb(YELLOW))
            draw_text_center(d, [700, y+25, 1100, y+155], label, pil_font(21, True), DARK, 28)
            mid = (900, y)
            bottom = (900, y+180)
            arrow(d, prev, mid)
            prev = bottom
            y += 250
        else:
            fill = {"normal":"F8FAFC","success":"DCFCE7","danger":"FEE2E2","external":"EDE9FE"}.get(color, "F8FAFC")
            outline = {"success":GREEN,"danger":RED,"external":PURPLE}.get(color, BLUE)
            draw_box(d, [560, y, 1240, y+105], label, fill=fill, outline=outline, width_chars=45)
            arrow(d, prev, (900, y))
            prev = (900, y+105)
            y += 170
    d.ellipse([830, y, 970, y+70], fill=hex_to_rgb(DARK), outline=hex_to_rgb(DARK))
    draw_text_center(d, [830, y, 970, y+70], "Kết thúc", pil_font(20, True), "FFFFFF", 10)
    arrow(d, prev, (900, y))
    return save(img, name)


def sequence_diagram(name: str, title: str, actors: list[str], messages: list[tuple[int, int, str]]) -> Path:
    w, h = 2200, max(1050, 390 + len(messages)*125)
    img, d = base_canvas(title, w=w, h=h)
    legend(d, [
        ("Yêu cầu từ client/người dùng", ORANGE),
        ("Xác thực, callback, webhook", PURPLE),
        ("Cập nhật thành công/realtime", GREEN),
        ("Hủy, hết hạn, hoàn tiền", RED),
    ], w - 610, 110)
    top = 150
    left = 190
    span = (w - 2 * left) / (len(actors)-1)
    xs = []
    for i, a in enumerate(actors):
        x = int(left + i*span)
        xs.append(x)
        draw_box(d, [x-125, top, x+125, top+72], a, fill=LIGHT, outline=BLUE, width_chars=14)
        d.line([x, top+72, x, h-80], fill=hex_to_rgb(LINE), width=3)
    y = top + 160
    for src, dst, msg in messages:
        x1, x2 = xs[src], xs[dst]
        color = message_color(msg)
        if src == dst:
            if x1 + 520 <= w - 90:
                box = [x1 + 28, y - 44, x1 + 520, y + 28]
            else:
                box = [max(90, x1 - 520), y - 44, x1 - 28, y + 28]
            d.rounded_rectangle(box, radius=16, fill=hex_to_rgb("FFFFFF"), outline=hex_to_rgb(color), width=3)
            for i, line in enumerate(wrap_text(msg, 34)[:2]):
                d.text((box[0] + 16, box[1] + 12 + i * 27), line, font=pil_font(19, True), fill=hex_to_rgb(DARK))
        else:
            arrow(d, (x1, y), (x2, y), color, 4)
            label_w = max(330, min(620, abs(x2 - x1) - 44))
            tx = min(x1, x2) + max(20, (abs(x2 - x1) - label_w) // 2)
            lines = wrap_text(msg, 42)[:2]
            label_h = 28 * len(lines) + 18
            d.rounded_rectangle([tx, y - label_h - 10, tx + label_w, y - 10], radius=12, fill=hex_to_rgb("FFFFFF"), outline=hex_to_rgb(LINE), width=2)
            for i, line in enumerate(lines):
                d.text((tx + 14, y - label_h + 2 + i * 28), line, font=pil_font(18, True), fill=hex_to_rgb(DARK))
        y += 122
    return save(img, name)


def state_diagram() -> Path:
    img, d = base_canvas("State diagram: Booking, Payment, Seat, Ticket", w=2200, h=1300)
    groups = [
        ("Booking", [("PENDING", "SUCCESS"), ("PENDING", "FAILED"), ("PENDING", "CANCELLED"), ("PENDING", "EXPIRED")], 100, 170),
        ("Payment", [("PENDING", "SUCCESS"), ("PENDING", "FAILED"), ("PENDING", "EXPIRED")], 1150, 170),
        ("SeatStatus", [("AVAILABLE", "HOLD"), ("HOLD", "BOOKED"), ("HOLD", "AVAILABLE")], 100, 720),
        ("Ticket", [("ACTIVE", "USED"), ("ACTIVE", "CANCELLED")], 1150, 720),
    ]
    for title, edges, x0, y0 in groups:
        d.text((x0, y0-60), title, font=pil_font(34, True), fill=hex_to_rgb(PURPLE))
        states = sorted({s for e in edges for s in e})
        coords = {}
        for i, st in enumerate(states):
            x = x0 + (i%3)*300
            y = y0 + (i//3)*190
            fill = "DCFCE7" if st in ("SUCCESS","BOOKED","USED","AVAILABLE") else "FEF3C7" if st in ("PENDING","HOLD","ACTIVE") else "FEE2E2"
            draw_box(d, [x, y, x+220, y+86], st, fill=fill, outline=LINE, width_chars=12)
            coords[st] = (x+110, y+43)
        for a,b in edges:
            arrow(d, coords[a], coords[b], MUTED)
    return save(img, "14_state_diagram.png")


def class_diagram() -> Path:
    img, d = base_canvas("Class diagram nghiệp vụ chính", w=2200, h=1350)
    classes = [
        ("Booking\n- id\n- status\n- totalPrice\n- paymentExpiresAt\n+ applyPromotion()\n+ expire()", 100, 170),
        ("BookingDetail\n- seat\n- priceAtBooking", 520, 170),
        ("Ticket\n- qrCode\n- status\n- checkInTime\n+ markUsed()", 940, 170),
        ("Payment\n- method\n- status\n- transactionNo\n+ markSuccess()", 1360, 170),
        ("PaymentEvent\n- provider\n- eventType\n- payload", 1780, 170),
        ("Showtime\n- startTime\n- endTime\n- status\n+ isBookable()", 100, 650),
        ("SeatStatus\n- status\n- holdBy\n- holdUntil\n- version\n+ hold()\n+ release()", 520, 650),
        ("Seat\n- rowLabel\n- seatNumber\n- seatType", 940, 650),
        ("User\n- username\n- email\n- roles", 1360, 650),
        ("Cinema/Room\n- city\n- address\n- roomName", 1780, 650),
    ]
    centers = {}
    for text,x,y in classes:
        key = text.split("\n")[0]
        centers[key] = (x+160, y+110)
        draw_box(d, [x, y, x+320, y+220], text, fill=LIGHT, outline=BLUE, bold=False, width_chars=24)
    rels = [("Booking","BookingDetail"),("BookingDetail","Ticket"),("Booking","Payment"),("Payment","PaymentEvent"),
            ("Booking","Showtime"),("BookingDetail","Seat"),("Showtime","SeatStatus"),("SeatStatus","Seat"),
            ("Booking","User"),("Showtime","Cinema/Room")]
    for a,b in rels:
        d.line([centers[a], centers[b]], fill=hex_to_rgb(MUTED), width=3)
    return save(img, "13_class_business.png")


def architecture_diagram() -> Path:
    img, d = base_canvas("Sơ đồ kiến trúc tổng thể", w=1900, h=1100)
    layers = [
        ("Client React\nVite, Router, Zustand, Axios, React Query", 110, 250, "DBEAFE"),
        ("Spring Boot API\nController - Service - Repository\nSecurity - WebSocket - Scheduler", 610, 250, "FFF7ED"),
        ("PostgreSQL\nFlyway, Index, Constraint, Audit", 1160, 250, "DCFCE7"),
        ("External Services\nGoogle OAuth, VNPay, SePay, SMTP", 610, 650, "EDE9FE"),
    ]
    for text,x,y,fill in layers:
        draw_box(d, [x, y, x+430, y+170], text, fill=fill, outline=LINE, width_chars=28)
    arrow(d, (540,335), (610,335), ORANGE)
    arrow(d, (1040,335), (1160,335), ORANGE)
    arrow(d, (825,420), (825,650), PURPLE)
    arrow(d, (825,650), (825,420), PURPLE)
    draw_box(d, [110, 650, 540, 820], "Trình duyệt người dùng\nKhách hàng / Staff / Admin", fill=LIGHT, outline=BLUE, width_chars=26)
    arrow(d, (325,650), (325,420), BLUE)
    return save(img, "15_architecture.png")


def deployment_diagram() -> Path:
    img, d = base_canvas("Sơ đồ triển khai đề xuất", w=1900, h=1100)
    nodes = [
        ("Browser / Mobile Web\nhttps://cinemabooking.vn", 90, 250, "DBEAFE"),
        ("Nginx / Static Hosting\nReact build", 520, 250, "F1F5F9"),
        ("Spring Boot Server\nREST API + WebSocket", 950, 250, "FFF7ED"),
        ("PostgreSQL Server\nDatabase", 1380, 250, "DCFCE7"),
        ("Google OAuth", 520, 680, "EDE9FE"),
        ("VNPay Sandbox/Production", 850, 680, "EDE9FE"),
        ("SePay Webhook", 1180, 680, "FEF3C7"),
        ("SMTP Mail Provider", 1510, 680, "ECFDF5"),
    ]
    for text,x,y,fill in nodes:
        draw_box(d, [x,y,x+300,y+140], text, fill=fill, outline=LINE, width_chars=20)
    arrows = [((390,320),(520,320)),((820,320),(950,320)),((1250,320),(1380,320)),
              ((1100,390),(670,680)),((1100,390),(1000,680)),((1100,390),(1330,680)),((1100,390),(1660,680))]
    for s,e in arrows:
        arrow(d,s,e,ORANGE)
    return save(img, "16_deployment.png")


def rbac_diagram() -> Path:
    img, d = base_canvas("Sơ đồ phân quyền RBAC và Staff Scope", w=1900, h=1100)
    boxes = [
        ("User", 140, 260, "DBEAFE"), ("Role", 500, 260, "F1F5F9"), ("Permission", 860, 260, "F1F5F9"),
        ("StaffCinema\nuser_id + cinema_id", 500, 620, "DCFCE7"), ("Cinema", 860, 620, "DBEAFE"),
        ("@PreAuthorize\nkiểm tra permission", 1220, 260, "FFF7ED"),
        ("Service scope check\nkiểm tra rạp phụ trách", 1220, 620, "FEF3C7"),
    ]
    for text,x,y,fill in boxes:
        draw_box(d, [x,y,x+280,y+130], text, fill=fill, outline=BLUE, width_chars=20)
    for s,e in [((420,325),(500,325)),((780,325),(860,325)),((1140,325),(1220,325)),
                ((280,390),(500,620)),((780,685),(860,685)),((1360,390),(1360,620))]:
        arrow(d,s,e,ORANGE)
    return save(img, "17_rbac_scope.png")


def backend_component_diagram() -> Path:
    img, d = base_canvas("Component diagram - Backend Spring Boot", w=2200, h=1350)
    legend(d, [
        ("REST/API boundary", BLUE),
        ("Nghiệp vụ lõi", ORANGE),
        ("Tích hợp ngoài", PURPLE),
        ("Tác vụ nền và realtime", GREEN),
    ], 1580, 112)
    draw_box(d, [90, 210, 420, 350], "Security Layer\nJWT Filter\n@PreAuthorize\nCORS", fill="DBEAFE", outline=BLUE, width_chars=22)
    draw_box(d, [90, 470, 420, 610], "Controller Layer\nAuth, Movie, Booking,\nPayment, Admin, Ticket", fill="DBEAFE", outline=BLUE, width_chars=24)
    draw_box(d, [90, 730, 420, 870], "DTO + Mapper\nRequest/Response\nValidation", fill="F8FAFC", outline=BLUE, width_chars=22)
    draw_box(d, [520, 430, 610, 650], "API\nentry", fill="EEF2FF", outline=BLUE, width_chars=8)
    arrow(d, (420, 280), (520, 490), BLUE)
    arrow(d, (420, 540), (520, 540), BLUE)
    arrow(d, (420, 800), (520, 590), BLUE)

    modules = [
        ("AuthService\nlogin, refresh, logout,\nemail verification", 700, 170, "FFF7ED"),
        ("BookingService\nhold seats, create booking,\nexpire, cancel", 1120, 170, "FFF7ED"),
        ("PaymentService\nVNPay, SePay,\npayment events", 1540, 170, "FFF7ED"),
        ("TicketService\nQR signing,\ncheck-in, ticket email", 700, 470, "FFF7ED"),
        ("Admin Services\nmovie/cinema/room/seat,\nshowtime, user, audit", 1120, 470, "FFF7ED"),
        ("Scope Services\nstaff-cinema assignment,\npermission checks", 1540, 470, "FFF7ED"),
        ("WebSocket Publisher\nseatmap topic events", 700, 770, "DCFCE7"),
        ("Schedulers\nexpired holds/bookings,\nshowtime status", 1120, 770, "DCFCE7"),
        ("Email Service\nSMTP templates", 1540, 770, "EDE9FE"),
    ]
    for text, x, y, fill in modules:
        draw_box(d, [x, y, x + 340, y + 150], text, fill=fill, outline=ORANGE if fill == "FFF7ED" else GREEN if fill == "DCFCE7" else PURPLE, width_chars=24)

    draw_box(d, [700, 1080, 1880, 1230], "Repository Layer + PostgreSQL\nEntity, JPQL/native query, index, transaction, Flyway migration", fill="ECFDF5", outline=GREEN, width_chars=62)
    for start, end, color in [
        ((610, 540), (700, 245), ORANGE),
        ((610, 540), (700, 545), ORANGE),
        ((1040, 245), (1120, 245), ORANGE),
        ((1460, 245), (1540, 245), ORANGE),
        ((1290, 320), (1290, 470), ORANGE),
        ((870, 620), (870, 770), GREEN),
        ((1290, 620), (1290, 770), GREEN),
        ((1710, 620), (1710, 770), PURPLE),
        ((1290, 920), (1290, 1080), GREEN),
    ]:
        arrow(d, start, end, color)
    draw_note(d, [90, 955, 575, 1195], "Ghi chú thiết kế", [
        "Controller không chứa nghiệp vụ phức tạp.",
        "Service là nơi giữ rule quan trọng.",
        "Repository có query rõ ràng để tránh N+1.",
        "Scheduler/WebSocket tách khỏi request chính."
    ])
    return save(img, "20_component_backend.png")


def frontend_architecture_diagram() -> Path:
    img, d = base_canvas("Component diagram - Frontend React", w=2200, h=1450)
    legend(d, [
        ("Điều hướng và bảo vệ route", BLUE),
        ("UI component", ORANGE),
        ("State/API", GREEN),
        ("Dịch vụ trình duyệt", PURPLE),
    ], 1530, 1160)
    cols = [
        ("Điều hướng", [
            "AppRouter\nPublic/User/Admin routes",
            "Route Guards\nRequireAuth/Role/Permission",
            "Layouts\nCustomer/Admin/Staff"
        ], 100, BLUE, "DBEAFE"),
        ("Màn hình và UI", [
            "Pages\nHome, MovieDetail,\nSeatSelection, Payment",
            "Admin/Staff Pages\nDashboard, Showtime,\nTicket QR, Staff Cinemas",
            "Reusable Components\nMovieCard, SeatMap,\nTable, Modal, QR"
        ], 600, ORANGE, "FFF7ED"),
        ("State và API", [
            "Hooks + Stores\nAuth, booking timer,\nseat realtime, profile",
            "API Client\nAxios interceptor,\nrefresh retry, normalize error",
            "React Query/Cache\nDữ liệu phim, rạp,\nsuất chiếu ít đổi"
        ], 1100, GREEN, "DCFCE7"),
        ("Tích hợp trình duyệt", [
            "Camera/File QR\nQuét vé bằng camera hoặc ảnh",
            "Local Storage\nLưu phiên, theme,\ntrạng thái UI",
            "Download Ticket\nLưu ảnh vé/QR về máy"
        ], 1600, PURPLE, "EDE9FE"),
    ]
    column_boxes = []
    for title, items, x, color, fill in cols:
        d.text((x, 150), title, font=pil_font(26, True), fill=hex_to_rgb(color))
        boxes = []
        y = 205
        for item in items:
            box = [x, y, x + 380, y + 150]
            draw_box(d, box, item, fill=fill, outline=color, width_chars=26)
            boxes.append(box)
            y += 220
        column_boxes.append(boxes)

    for col in column_boxes:
        for i in range(len(col) - 1):
            top_box, bottom_box = col[i], col[i + 1]
            arrow(d, ((top_box[0] + top_box[2]) // 2, top_box[3]), ((bottom_box[0] + bottom_box[2]) // 2, bottom_box[1]), MUTED)
    for i, color in [(0, ORANGE), (1, GREEN), (2, PURPLE)]:
        for row in range(3):
            left_box, right_box = column_boxes[i][row], column_boxes[i + 1][row]
            arrow(d, (left_box[2], (left_box[1] + left_box[3]) // 2), (right_box[0], (right_box[1] + right_box[3]) // 2), color)

    draw_box(d, [760, 970, 1640, 1115], "Backend API + WebSocket\nREST /api/**, STOMP /ws, payment callback/webhook", fill="ECFDF5", outline=GREEN, width_chars=48)
    cache_box = column_boxes[2][2]
    arrow(d, ((cache_box[0] + cache_box[2]) // 2, cache_box[3]), (1200, 970), GREEN)
    draw_note(d, [90, 1130, 760, 1340], "Điểm cần trình bày khi bảo vệ", [
        "Frontend không gọi database trực tiếp.",
        "Axios interceptor giúp refresh access token tự động.",
        "SeatMap nhận snapshot ban đầu và cập nhật tiếp bằng WebSocket.",
        "QR scanner hỗ trợ camera và file ảnh để phù hợp thiết bị thực tế."
    ])
    return save(img, "21_component_frontend.png")


def class_rbac_diagram() -> Path:
    img, d = base_canvas("Class diagram - RBAC và phạm vi nhân viên", w=2100, h=1150)
    classes = [
        ("User\n- id\n- username\n- email\n- emailVerified\n- isActive\n- roles", 120, 210),
        ("Role\n- id\n- name\n- description\n- permissions", 540, 210),
        ("Permission\n- id\n- name\n- description", 960, 210),
        ("StaffCinema\n- userId\n- cinemaId\n- assignedAt", 540, 610),
        ("Cinema\n- id\n- name\n- city\n- address\n- isActive", 960, 610),
        ("SecurityContext\n- currentUser\n+ hasPermission()", 1380, 210),
        ("StaffScopeService\n+ assertCanAccessCinema()\n+ getAssignedCinemaIds()", 1380, 610),
    ]
    centers = {}
    for text, x, y in classes:
        key = text.split("\n")[0]
        centers[key] = (x + 170, y + 105)
        draw_box(d, [x, y, x + 340, y + 210], text, fill=LIGHT, outline=BLUE, bold=False, width_chars=26)
    rels = [
        ("User", "Role", "n-n", ORANGE),
        ("Role", "Permission", "n-n", ORANGE),
        ("User", "StaffCinema", "1-n nếu là STAFF", GREEN),
        ("StaffCinema", "Cinema", "n-1", GREEN),
        ("SecurityContext", "Role", "đọc quyền", PURPLE),
        ("StaffScopeService", "StaffCinema", "kiểm tra rạp", GREEN),
        ("StaffScopeService", "Cinema", "giới hạn dữ liệu", GREEN),
    ]
    for a, b, label, color in rels:
        arrow(d, centers[a], centers[b], color)
        mx = (centers[a][0] + centers[b][0]) // 2
        my = (centers[a][1] + centers[b][1]) // 2
        d.text((mx + 8, my - 28), label, font=pil_font(18, True), fill=hex_to_rgb(color))
    draw_note(d, [120, 900, 1190, 1050], "Ý nghĩa", [
        "RBAC trả lời câu hỏi: người dùng có được làm hành động này không?",
        "Staff scope trả lời câu hỏi tiếp theo: nhân viên có được thao tác trên rạp này không?"
    ], fill="ECFDF5", outline=GREEN)
    return save(img, "22_class_rbac.png")


def class_backend_architecture_diagram() -> Path:
    img, d = base_canvas("Class diagram - Kiến trúc backend theo lớp", w=2100, h=1150)
    boxes = [
        ("Controller\n@PostMapping\n@GetMapping\n@PreAuthorize\n@Valid", 120, 250, "DBEAFE", BLUE),
        ("Request DTO\nValidation annotation\nKhông expose entity", 500, 250, "F8FAFC", BLUE),
        ("Service Interface\nKhai báo nghiệp vụ", 880, 250, "FFF7ED", ORANGE),
        ("Service Impl\n@Transactional\nBusiness rule\nScope check", 1260, 250, "FFF7ED", ORANGE),
        ("Repository\nJpaRepository\n@EntityGraph/JPQL\nPagination", 1640, 250, "ECFDF5", GREEN),
        ("Entity\nJPA mapping\nConstraint\nRelation", 500, 650, "ECFDF5", GREEN),
        ("Mapper\nEntity -> Response DTO", 880, 650, "F8FAFC", BLUE),
        ("Response DTO\nDữ liệu trả client", 1260, 650, "F8FAFC", BLUE),
        ("GlobalExceptionHandler\nChuẩn hóa lỗi API", 1640, 650, "FEE2E2", RED),
    ]
    centers = {}
    for text, x, y, fill, outline in boxes:
        key = text.split("\n")[0]
        centers[key] = (x + 160, y + 100)
        draw_box(d, [x, y, x + 320, y + 200], text, fill=fill, outline=outline, bold=False, width_chars=23)
    for a, b, color in [
        ("Controller", "Request DTO", BLUE),
        ("Request DTO", "Service Interface", ORANGE),
        ("Service Interface", "Service Impl", ORANGE),
        ("Service Impl", "Repository", GREEN),
        ("Repository", "Entity", GREEN),
        ("Entity", "Mapper", BLUE),
        ("Mapper", "Response DTO", BLUE),
        ("Controller", "GlobalExceptionHandler", RED),
    ]:
        arrow(d, centers[a], centers[b], color)
    draw_note(d, [120, 900, 1040, 1050], "Quy ước clean code", [
        "Controller chỉ điều phối request/response.",
        "Service giữ nghiệp vụ và transaction.",
        "DTO giúp không rò rỉ entity nội bộ ra client."
    ])
    return save(img, "23_class_backend_architecture.png")


def auth_sequence_diagram() -> Path:
    return sequence_diagram(
        "24_sequence_auth_refresh.png",
        "Sequence - Đăng nhập JWT, refresh token và logout",
        ["React", "Auth API", "AuthService", "Database", "Security Filter"],
        [
            (0, 1, "POST /auth/login username/password"),
            (1, 2, "Validate request và gọi nghiệp vụ đăng nhập"),
            (2, 3, "Load user, role, permission; kiểm tra password BCrypt"),
            (2, 3, "Tạo refresh token mới và lưu hash/expiry"),
            (2, 0, "Trả access token ngắn hạn + refresh token"),
            (0, 4, "Gọi API kèm Authorization Bearer access token"),
            (4, 4, "Verify chữ ký, hạn dùng, blacklist token"),
            (0, 1, "Access token hết hạn -> POST /auth/refresh"),
            (1, 2, "Rotate refresh token, thu hồi token cũ"),
            (0, 1, "Logout -> revoke refresh token + invalidate access token"),
        ],
    )


def scheduler_sequence_diagram() -> Path:
    return sequence_diagram(
        "25_sequence_scheduler_expire.png",
        "Sequence - Scheduler trả ghế và booking hết hạn",
        ["Scheduler", "BookingService", "Database", "PaymentService", "WebSocket", "React Clients"],
        [
            (0, 1, "Chạy theo chu kỳ cấu hình trong env"),
            (1, 2, "Tìm seat_status HOLD có hold_until <= now"),
            (1, 2, "Set AVAILABLE, clear hold_by/hold_until"),
            (1, 2, "Tìm booking PENDING có payment_expires_at <= now"),
            (1, 3, "Expire payment pending nếu chưa ghi nhận thanh toán"),
            (1, 2, "Set booking EXPIRED và giải phóng ghế liên quan"),
            (1, 4, "Publish SeatStatusEvent sau khi commit"),
            (4, 5, "Client tự đổi màu ghế realtime"),
        ],
    )


def cancel_showtime_sequence_diagram() -> Path:
    return sequence_diagram(
        "26_sequence_cancel_showtime_refund.png",
        "Sequence - Admin/Staff hủy suất chiếu và xử lý hoàn tiền",
        ["Admin/Staff", "React Admin", "ShowtimeService", "Database", "PaymentService", "Email/Audit"],
        [
            (0, 1, "Chọn hủy suất chiếu và nhập lý do"),
            (1, 2, "PATCH /showtimes/{id}/cancel"),
            (2, 2, "Kiểm tra quyền ADMIN hoặc STAFF thuộc rạp"),
            (2, 3, "Khóa showtime, booking, ticket liên quan"),
            (2, 3, "Set showtime CANCELLED, ticket CANCELLED"),
            (2, 4, "Tạo yêu cầu hoàn tiền cho payment SUCCESS"),
            (2, 3, "Release ghế chưa check-in nếu phù hợp policy"),
            (2, 5, "Ghi audit log và gửi email thông báo khách hàng"),
            (2, 1, "Trả kết quả, UI hiện trạng thái đang xử lý hoàn tiền"),
        ],
    )


def activity_login_refresh_diagram() -> Path:
    return activity_diagram("27_activity_login_refresh.png", "Activity - Đăng nhập, refresh token và logout", [
        ("Người dùng nhập thông tin đăng nhập hoặc chọn Google", "normal"),
        ("Backend xác thực tài khoản, trạng thái active và email", "decision"),
        ("Sinh access token ngắn hạn và refresh token dài hơn", "success"),
        ("Frontend lưu phiên đăng nhập và gọi API bằng access token", "normal"),
        ("Access token hết hạn: Axios interceptor gọi refresh", "decision"),
        ("Refresh hợp lệ: rotate token và gọi lại request cũ", "success"),
        ("Logout: thu hồi refresh token và invalidate access token", "danger"),
    ])


def activity_cancel_refund_diagram() -> Path:
    return activity_diagram("28_activity_cancel_refund.png", "Activity - Hủy suất chiếu và hoàn tiền", [
        ("Admin hoặc staff rạp phụ trách chọn suất cần hủy", "normal"),
        ("Backend kiểm tra permission và staff scope", "decision"),
        ("Cập nhật showtime = CANCELLED", "danger"),
        ("Tìm booking SUCCESS/PENDING liên quan", "normal"),
        ("PENDING: hủy booking và trả ghế", "danger"),
        ("SUCCESS: hủy ticket và tạo trạng thái xử lý hoàn tiền", "external"),
        ("Gửi email thông báo khách và ghi audit log", "success"),
    ])


def websocket_diagram() -> Path:
    return sequence_diagram(
        "12_sequence_websocket.png",
        "Sequence diagram - WebSocket realtime giữ ghế",
        ["User A", "React A", "Backend", "Database", "Publisher", "React B"],
        [
            (0,1,"Chọn ghế"), (1,2,"POST /bookings/hold-seats"), (2,3,"Lock seat_status và set HOLD"),
            (3,2,"Commit transaction"), (2,4,"publish after commit"), (4,5,"/topic/seatmap/{showtimeId}"),
            (5,5,"Cập nhật màu ghế realtime"),
        ],
    )


def generate_diagrams() -> list[tuple[str, Path]]:
    items = [
        ("Hình 1. Sơ đồ ngữ cảnh hệ thống", context_diagram()),
        ("Hình 2. Sơ đồ phân rã chức năng BFD", bfd_diagram()),
        ("Hình 3. Sơ đồ Use Case tổng quát", use_case_diagram()),
        ("Hình 4. DFD Level 0", dfd0_diagram()),
        ("Hình 5. DFD Level 1", dfd1_diagram()),
        ("Hình 6. ERD tổng quan cơ sở dữ liệu", erd_diagram()),
        ("Hình 7. Activity diagram - Đăng ký và xác thực email", activity_diagram("07_activity_register_verify.png", "Activity - Đăng ký và xác thực email", [
            ("Người dùng nhập username, email, mật khẩu", "normal"),
            ("Backend validate dữ liệu và kiểm tra trùng email/username", "decision"),
            ("Tạo user chưa xác thực và token hash", "normal"),
            ("Gửi email xác thực qua SMTP", "external"),
            ("Người dùng bấm link xác thực", "normal"),
            ("Cập nhật email_verified = true", "success"),
        ])),
        ("Hình 8. Activity diagram - Đặt vé và giữ ghế", activity_diagram("08_activity_booking.png", "Activity - Đặt vé và giữ ghế", [
            ("Chọn phim, rạp, ngày, suất chiếu", "normal"),
            ("Mở sơ đồ ghế và nhận snapshot seat_status", "normal"),
            ("Chọn ghế còn AVAILABLE", "decision"),
            ("Backend lock ghế và set HOLD theo user", "normal"),
            ("Tạo booking PENDING và payment_expires_at", "success"),
            ("Chuyển sang màn hình thanh toán", "normal"),
        ])),
        ("Hình 9. Activity diagram - Thanh toán thành công/thất bại", activity_diagram("09_activity_payment.png", "Activity - Thanh toán VNPay/Quét QR ngân hàng", [
            ("User chọn phương thức thanh toán", "normal"),
            ("VNPay redirect hoặc SePay tạo QR ngân hàng", "external"),
            ("Cổng thanh toán/cổng webhook trả kết quả", "decision"),
            ("SUCCESS: booking SUCCESS, seat BOOKED, sinh ticket", "success"),
            ("FAILED/EXPIRED: booking FAILED/EXPIRED, trả ghế", "danger"),
            ("Gửi email vé hoặc thông báo kết quả", "external"),
        ])),
        ("Hình 10. Activity diagram - Soát vé QR", activity_diagram("10_activity_checkin.png", "Activity - Soát vé QR", [
            ("Staff chọn thành phố, rạp và suất đang mở check-in", "normal"),
            ("Quét QR bằng camera hoặc file ảnh", "normal"),
            ("Backend kiểm tra chữ ký QR và ticket ACTIVE", "decision"),
            ("Kiểm tra booking SUCCESS, đúng rạp, đúng suất, đúng cửa sổ", "decision"),
            ("Set ticket USED, lưu checked_in_by và check_in_time", "success"),
            ("Trả kết quả xác thực cho nhân viên", "normal"),
        ])),
        ("Hình 11. Sequence diagram - VNPay payment callback", sequence_diagram("11_sequence_vnpay.png", "Sequence - Thanh toán VNPay", ["React", "Backend", "VNPay", "Database", "Email"], [
            (0,1,"Tạo payment VNPay"), (1,3,"Lưu payment PENDING"), (1,2,"Redirect payment URL"),
            (2,0,"User thanh toán"), (2,1,"Return/IPN callback"), (1,2,"Verify secure hash"),
            (1,3,"Payment SUCCESS, Booking SUCCESS, Seat BOOKED, Ticket ACTIVE"), (1,4,"Gửi email vé"), (1,0,"Trả trang kết quả"),
        ])),
        ("Hình 12. Sequence diagram - WebSocket realtime giữ ghế", websocket_diagram()),
        ("Hình 13. Sequence diagram - SePay webhook tự xác nhận", sequence_diagram("18_sequence_sepay.png", "Sequence - SePay / VietQR webhook", ["React", "Backend", "SePay", "Database", "Email"], [
            (0,1,"Tạo QR ngân hàng"), (1,3,"Payment PENDING + nội dung CK"), (1,0,"Trả QR/số tiền/nội dung"),
            (0,2,"User chuyển khoản"), (2,1,"Webhook tiền vào"), (1,1,"Verify API key/HMAC + đối chiếu amount/content"),
            (1,3,"Confirm payment và booking"), (1,4,"Gửi email vé"), (1,0,"Client polling/WebSocket chuyển kết quả"),
        ])),
        ("Hình 14. Sequence diagram - QR check-in đúng rạp/suất", sequence_diagram("19_sequence_checkin.png", "Sequence - Soát vé QR", ["Staff UI", "Backend", "Ticket Service", "Database", "Audit"], [
            (0,1,"POST /tickets/check-in + QR + cinemaId + showtimeId"), (1,2,"Verify signed QR"), (2,3,"Load ticket + booking + showtime"),
            (1,1,"Check ACTIVE, SUCCESS, đúng rạp/suất, cửa sổ check-in"), (1,3,"Set USED + checked_in_by"), (1,4,"Ghi audit"), (1,0,"Trả kết quả xác thực"),
        ])),
        ("Hình 15. Class diagram nghiệp vụ chính", class_diagram()),
        ("Hình 16. State diagram Booking/Payment/Seat/Ticket", state_diagram()),
        ("Hình 17. Sơ đồ kiến trúc tổng thể", architecture_diagram()),
        ("Hình 18. Sơ đồ triển khai đề xuất", deployment_diagram()),
        ("Hình 19. Sơ đồ RBAC và Staff Scope", rbac_diagram()),
        ("Hình 20. Component diagram backend Spring Boot", backend_component_diagram()),
        ("Hình 21. Component diagram frontend React", frontend_architecture_diagram()),
        ("Hình 22. Class diagram RBAC và Staff Scope", class_rbac_diagram()),
        ("Hình 23. Class diagram kiến trúc backend theo lớp", class_backend_architecture_diagram()),
        ("Hình 24. Sequence diagram đăng nhập JWT và refresh token", auth_sequence_diagram()),
        ("Hình 25. Sequence diagram scheduler trả ghế hết hạn", scheduler_sequence_diagram()),
        ("Hình 26. Sequence diagram hủy suất chiếu và xử lý hoàn tiền", cancel_showtime_sequence_diagram()),
        ("Hình 27. Activity diagram đăng nhập và refresh token", activity_login_refresh_diagram()),
        ("Hình 28. Activity diagram hủy suất chiếu và hoàn tiền", activity_cancel_refund_diagram()),
    ]
    return items


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold=False, color=DARK, size=11):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(0.75)

    for name, size, color in [
        ("Heading 1", 16, PURPLE),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 13, DARK),
    ]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.first_line_indent = Cm(0)
    return doc


def para(doc, text: str, align=None):
    p = doc.add_paragraph()
    p.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    return p


def bullets(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.7)
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(13)


def table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_shading(t.rows[0].cells[i], "EDE9FE")
        set_cell_text(t.rows[0].cells[i], h, bold=True, color=PURPLE, size=11)
        if widths:
            t.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return t


def caption(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def figure(doc, cap: str, path: Path, width=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, cap)


def placeholder(doc, text: str):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    set_cell_shading(cell, "F8FAFC")
    set_cell_text(cell, text, bold=True, color=MUTED, size=11)
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(18)
    doc.add_paragraph()


def cover(doc):
    section = doc.sections[0]
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = tbl.rows[0].cells
    left.width = Cm(5.2)
    right.width = Cm(12.0)
    set_cell_shading(left, PURPLE)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(300)
    run = p.add_run("Học để có nghề nghiệp\nHọc để có việc làm\nHọc để có khả năng tự học")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string("FFFFFF")

    for p in right.paragraphs:
        p.paragraph_format.first_line_indent = Cm(0)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(45)
    r = p.add_run("PHU XUAN UNIVERSITY\n")
    r.bold = True
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor.from_string(PURPLE)
    p.add_run("\nKHÓA LUẬN\nTỐT NGHIỆP").bold = True
    for r in p.runs:
        r.font.name = "Times New Roman"
    p.runs[-1].font.size = Pt(31)
    p.runs[-1].font.color.rgb = RGBColor.from_string(PURPLE)

    p2 = right.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(80)
    p2.paragraph_format.first_line_indent = Cm(0)
    r = p2.add_run("ĐỀ TÀI:\nXÂY DỰNG HỆ THỐNG ĐẶT VÉ XEM PHIM TRỰC TUYẾN CINEMABOOKING.VN")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)

    p3 = right.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(80)
    p3.paragraph_format.first_line_indent = Cm(0)
    r = p3.add_run("Họ và tên: ........................................\nNgành: Công nghệ thông tin\nGiảng viên hướng dẫn: ........................................")
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)

    p4 = right.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(150)
    p4.paragraph_format.first_line_indent = Cm(0)
    r = p4.add_run("HUẾ 2026")
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    doc.add_page_break()

    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)


def front_matter(doc, diagrams: list[tuple[str, Path]]):
    doc.add_heading("KHÓA LUẬN TỐT NGHIỆP", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "Hệ đại học chính quy", WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Tên đề tài: Xây dựng hệ thống đặt vé xem phim trực tuyến CinemaBooking.vn", WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Họ và tên: ........................................", WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Mã số sinh viên: ........................................", WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Khoa/Ngành: Công nghệ thông tin", WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Năm tốt nghiệp: 2026", WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    doc.add_heading("LỜI CAM ĐOAN", level=1)
    para(doc, "Em xin cam đoan khóa luận với đề tài “Xây dựng hệ thống đặt vé xem phim trực tuyến CinemaBooking.vn” là kết quả tìm hiểu, phân tích, thiết kế và xây dựng của cá nhân em dựa trên dự án hiện tại. Các nội dung phân tích, mô tả luồng nghiệp vụ, thiết kế cơ sở dữ liệu, thiết kế giao diện và đánh giá hệ thống được trình bày bằng ngôn ngữ riêng, không sao chép nguyên văn từ tài liệu khác.")
    para(doc, "Các công nghệ, thư viện và dịch vụ bên ngoài như Spring Boot, React, PostgreSQL, VNPay, SePay, Google OAuth và SMTP Mail được sử dụng đúng mục đích học thuật. Những phần tham khảo được tổng hợp lại để phục vụ việc giải thích hệ thống, không nhằm thay thế tài liệu chính thức của nhà cung cấp.")
    para(doc, "Huế, ngày ..... tháng ..... năm 2026", WD_ALIGN_PARAGRAPH.RIGHT)
    para(doc, "Sinh viên thực hiện\n\n................................", WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()

    doc.add_heading("LỜI CẢM ƠN", level=1)
    para(doc, "Trong quá trình thực hiện đề tài, em có cơ hội hệ thống hóa kiến thức về lập trình web, thiết kế cơ sở dữ liệu, bảo mật ứng dụng, thanh toán trực tuyến, xử lý realtime và triển khai phần mềm. Em xin gửi lời cảm ơn đến giảng viên hướng dẫn đã định hướng và góp ý để em hoàn thiện đề tài theo hướng thực tế hơn.")
    para(doc, "Em cũng xin cảm ơn gia đình, bạn bè và các nguồn tài liệu kỹ thuật đã hỗ trợ em trong quá trình học tập, thử nghiệm và hoàn thiện hệ thống. Mặc dù đã cố gắng, báo cáo khó tránh khỏi thiếu sót; em mong nhận được góp ý để tiếp tục hoàn thiện sản phẩm.")
    doc.add_page_break()

    doc.add_heading("TÓM TẮT ĐỀ TÀI", level=1)
    para(doc, "CinemaBooking.vn là hệ thống đặt vé xem phim trực tuyến gồm frontend React và backend Spring Boot. Hệ thống hỗ trợ người dùng xem phim, chọn rạp, chọn suất chiếu, chọn ghế, áp mã giảm giá, thanh toán qua VNPay hoặc quét QR ngân hàng, nhận vé QR qua email và xem lịch sử vé. Nhân viên rạp có thể vận hành trong phạm vi rạp được phân công, quản lý suất chiếu, xem dữ liệu liên quan và soát vé QR đúng rạp, đúng suất. Admin có thể quản trị dữ liệu toàn hệ thống, phân quyền, gán rạp cho nhân viên, theo dõi dashboard, booking, payment, audit log.")
    para(doc, "Điểm nổi bật của hệ thống là luồng giữ ghế có timeout, cập nhật trạng thái ghế realtime bằng WebSocket, mô hình RBAC kết hợp staff scope, xác thực JWT với refresh token, email verification, tích hợp payment gateway, ticket QR có chữ ký và các scheduler xử lý hết hạn. Báo cáo trình bày đầy đủ từ tổng quan, cơ sở lý thuyết, phân tích yêu cầu, thiết kế hệ thống, triển khai, kiểm thử đến hướng phát triển.")
    doc.add_page_break()

    doc.add_heading("MỤC LỤC", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc = [
        "LỜI CAM ĐOAN ........................................................................................................ i",
        "LỜI CẢM ƠN ........................................................................................................... ii",
        "TÓM TẮT ĐỀ TÀI .................................................................................................... iii",
        "DANH MỤC THUẬT NGỮ ...................................................................................... iv",
        "DANH MỤC HÌNH ẢNH .......................................................................................... v",
        "DANH MỤC BẢNG ................................................................................................. vii",
        "LỜI NÓI ĐẦU ........................................................................................................ viii",
        "CHƯƠNG 1: TỔNG QUAN ..................................................................................... 1",
        "CHƯƠNG 2: CƠ SỞ LÝ THUYẾT .......................................................................... 8",
        "CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG ........................................ 20",
        "CHƯƠNG 4: TRIỂN KHAI HỆ THỐNG ............................................................... 55",
        "CHƯƠNG 5: KIỂM THỬ HỆ THỐNG .................................................................. 82",
        "CHƯƠNG 6: ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN .......................................... 92",
        "KẾT LUẬN ............................................................................................................. 100",
        "TÀI LIỆU THAM KHẢO ....................................................................................... 102",
        "PHỤ LỤC ............................................................................................................... 104",
    ]
    for item in toc:
        para(doc, item, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

    doc.add_heading("DANH MỤC THUẬT NGỮ", level=1)
    table(doc, ["Thuật ngữ", "Giải thích"], [
        ["API", "Giao diện lập trình ứng dụng, cho phép frontend gọi chức năng backend."],
        ["JWT", "Token dùng để xác thực request API sau khi đăng nhập."],
        ["RBAC", "Mô hình phân quyền dựa trên vai trò và permission."],
        ["Staff Scope", "Giới hạn dữ liệu nhân viên theo rạp được admin phân công."],
        ["WebSocket", "Kết nối realtime dùng để cập nhật trạng thái ghế ngay lập tức."],
        ["VNPay", "Cổng thanh toán redirect/callback."],
        ["SePay/VietQR", "Thanh toán chuyển khoản QR và xác nhận tự động qua webhook."],
        ["Flyway", "Công cụ quản lý migration database theo phiên bản."],
    ], [3.2, 11.5])
    doc.add_page_break()

    doc.add_heading("DANH MỤC HÌNH ẢNH", level=1)
    for cap, _ in diagrams:
        para(doc, f"{cap} ..............................................................................................")
    for idx, name in enumerate([
        "Giao diện trang chủ",
        "Giao diện chi tiết phim",
        "Giao diện chọn rạp và lịch chiếu",
        "Giao diện chọn ghế",
        "Giao diện thanh toán VNPay và quét QR ngân hàng",
        "Giao diện vé của tôi",
        "Giao diện admin dashboard",
        "Giao diện staff soát vé QR",
    ], start=len(diagrams)+1):
        para(doc, f"Hình {idx}. {name} ..............................................................................................")
    doc.add_page_break()

    doc.add_heading("DANH MỤC BẢNG", level=1)
    for item in [
        "Bảng 1. Các tác nhân trong hệ thống",
        "Bảng 2. Yêu cầu chức năng theo nhóm người dùng",
        "Bảng 3. Yêu cầu phi chức năng",
        "Bảng 4. Danh sách bảng cơ sở dữ liệu",
        "Bảng 5. Ma trận phân quyền chính",
        "Bảng 6. Test case tiêu biểu",
    ]:
        para(doc, f"{item} ..............................................................................................")
    doc.add_page_break()

    doc.add_heading("LỜI NÓI ĐẦU", level=1)
    para(doc, "Sự phát triển của thương mại điện tử và thói quen sử dụng dịch vụ trực tuyến đã làm thay đổi cách người dùng tiếp cận các hoạt động giải trí. Đối với lĩnh vực rạp chiếu phim, người dùng ngày càng mong muốn có thể xem lịch chiếu, chọn ghế, thanh toán và nhận vé điện tử mà không cần đến quầy. Điều này đặt ra yêu cầu xây dựng một hệ thống đặt vé ổn định, dễ sử dụng, xử lý trạng thái ghế chính xác và đảm bảo thanh toán an toàn.")
    para(doc, "Từ nhu cầu đó, đề tài CinemaBooking.vn được xây dựng nhằm mô phỏng một hệ thống đặt vé xem phim trực tuyến theo hướng sản phẩm thực tế. Hệ thống không chỉ tập trung vào thao tác CRUD cơ bản mà còn xử lý các luồng phức tạp như giữ ghế có thời hạn, thanh toán qua nhiều phương thức, xác nhận webhook, gửi vé qua email, QR check-in, phân quyền admin/staff/user và cập nhật seat map realtime.")
    doc.add_page_break()


def chapters(doc, diagrams):
    doc.add_heading("CHƯƠNG 1: TỔNG QUAN", level=1)
    doc.add_heading("1.1 Lý do chọn đề tài", level=2)
    para(doc, "Trong các hệ thống bán vé, lỗi bán trùng ghế, hết hạn thanh toán không trả ghế, thanh toán sai số tiền hoặc quét nhầm vé là những vấn đề thường gặp nếu chỉ thiết kế ở mức CRUD. Vì vậy, đề tài đặt trọng tâm vào việc xây dựng một quy trình đặt vé có trạng thái rõ ràng, có kiểm soát timeout, có realtime và có phân quyền theo vai trò.")
    para(doc, "CinemaBooking.vn phù hợp với bối cảnh học tập và thực tế vì hệ thống bao gồm nhiều mảng kiến thức: frontend, backend, database, bảo mật, payment, email, QR code, WebSocket, scheduler, cache, audit và kiểm thử. Đây là nền tảng tốt để sinh viên chứng minh khả năng phân tích, thiết kế và triển khai một ứng dụng web hoàn chỉnh.")
    doc.add_heading("1.2 Mục đích phát triển", level=2)
    bullets(doc, [
        "Xây dựng website cho phép khách hàng xem phim, chọn suất, chọn ghế và thanh toán trực tuyến.",
        "Đảm bảo trạng thái ghế nhất quán khi nhiều người cùng thao tác.",
        "Tích hợp thanh toán VNPay và quét QR ngân hàng thông qua SePay/VietQR.",
        "Sinh vé điện tử có QR và hỗ trợ staff soát vé đúng rạp, đúng suất chiếu.",
        "Xây dựng trang quản trị cho admin và nhân viên rạp có phân quyền rõ ràng.",
        "Thiết kế hệ thống có khả năng mở rộng thêm cổng thanh toán, loyalty, combo và mobile app.",
    ])
    doc.add_heading("1.3 Đối tượng hướng đến và phạm vi đề tài", level=2)
    para(doc, "Đối tượng sử dụng gồm khách hàng đặt vé, nhân viên rạp và quản trị viên hệ thống. Khách hàng tập trung vào trải nghiệm đặt vé nhanh và rõ ràng. Nhân viên rạp tập trung vào vận hành rạp được phân công, soát vé và theo dõi suất chiếu. Admin chịu trách nhiệm quản trị toàn bộ dữ liệu, phân quyền và theo dõi hoạt động hệ thống.")
    table(doc, ["Tác nhân", "Nhu cầu chính"], [
        ["Khách hàng", "Xem phim/rạp/lịch chiếu, chọn ghế, thanh toán, nhận vé QR, xem lịch sử vé."],
        ["Nhân viên rạp", "Xem rạp phụ trách, tạo/sửa/hủy suất chiếu trong phạm vi, soát vé QR, xem booking/payment liên quan."],
        ["Admin", "Quản lý phim, rạp, phòng, ghế, suất chiếu, khuyến mãi, người dùng, phân quyền, dashboard, audit."],
        ["Dịch vụ ngoài", "Google OAuth, VNPay, SePay/VietQR, SMTP Mail, PostgreSQL."],
    ], [3.5, 11.2])
    doc.add_heading("1.4 Phương pháp thực hiện", level=2)
    bullets(doc, [
        "Khảo sát nghiệp vụ từ các hệ thống đặt vé phim thực tế.",
        "Phân tích chức năng theo actor và luồng dữ liệu.",
        "Thiết kế database trước bằng schema có constraint/index rõ ràng.",
        "Triển khai backend Spring Boot theo mô hình Controller - Service - Repository.",
        "Triển khai frontend React theo route, page, component, API client và state store.",
        "Kiểm thử các luồng quan trọng bằng unit test, integration test và manual test.",
    ])

    doc.add_heading("CHƯƠNG 2: CƠ SỞ LÝ THUYẾT", level=1)
    for title, body in [
        ("2.1 Hệ thống đặt vé xem phim trực tuyến", "Hệ thống đặt vé xem phim trực tuyến là hệ thống giao dịch có trạng thái cao. Khác với website giới thiệu thông thường, mỗi thao tác chọn ghế, giữ ghế, thanh toán và sinh vé đều ảnh hưởng trực tiếp đến dữ liệu tồn kho ghế của suất chiếu."),
        ("2.2 Kiến trúc client-server", "Frontend React đóng vai trò client, hiển thị giao diện và gọi API. Backend Spring Boot xử lý xác thực, nghiệp vụ, giao tiếp database và tích hợp dịch vụ ngoài. Kiến trúc này giúp tách biệt trách nhiệm và dễ mở rộng."),
        ("2.3 Spring Boot và mô hình phân lớp", "Spring Boot giúp xây dựng REST API nhanh, kết hợp Spring Security, Data JPA, Validation, Mail, WebSocket và Scheduler. Mô hình phân lớp giúp controller mỏng, service chứa nghiệp vụ, repository truy vấn database."),
        ("2.4 React, TypeScript và Vite", "React phù hợp xây dựng giao diện tương tác cao như chọn ghế, thanh toán, admin dashboard. TypeScript giúp kiểm soát kiểu dữ liệu, Vite giúp trải nghiệm phát triển nhanh."),
        ("2.5 PostgreSQL, JPA và Flyway", "PostgreSQL đảm bảo transaction, constraint, index và dữ liệu quan hệ. JPA giúp mapping entity. Flyway quản lý lịch sử thay đổi schema để triển khai nhất quán."),
        ("2.6 JWT, refresh token và OAuth", "JWT giúp xác thực API không trạng thái. Refresh token rotation tăng bảo mật phiên đăng nhập. Google OAuth giúp người dùng đăng nhập bằng tài khoản Google."),
        ("2.7 RBAC và staff scope", "RBAC xác định người dùng có quyền làm gì. Staff scope xác định nhân viên được thao tác trên rạp nào. Kết hợp hai lớp này giúp phân quyền sát thực tế hơn."),
        ("2.8 WebSocket realtime", "WebSocket cho phép server đẩy thay đổi trạng thái ghế đến client đang xem cùng suất chiếu, giúp giảm thao tác refresh và tránh hiểu nhầm trạng thái."),
        ("2.9 Payment gateway và webhook", "VNPay sử dụng redirect/callback, SePay/VietQR sử dụng QR chuyển khoản và webhook. Backend cần xác thực chữ ký, đối chiếu số tiền và cập nhật trạng thái giao dịch."),
    ]:
        doc.add_heading(title, level=2)
        para(doc, body)

    doc.add_heading("CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", level=1)
    doc.add_heading("3.1 Bài toán đặt ra", level=2)
    para(doc, "Bài toán yêu cầu xây dựng một hệ thống có khả năng phục vụ đồng thời nhiều người dùng, đảm bảo ghế không bị bán trùng, thanh toán đúng số tiền, vé điện tử không bị dùng lại và nhân viên chỉ thao tác trong phạm vi rạp được phân công.")
    doc.add_heading("3.2 Yêu cầu chức năng", level=2)
    table(doc, ["Nhóm chức năng", "Mô tả"], [
        ["Tài khoản", "Đăng ký, xác thực email, đăng nhập, Google login, refresh token, logout, reset mật khẩu."],
        ["Phim và rạp", "Xem danh sách phim, rạp, bản đồ rạp, lọc theo thành phố, xem lịch chiếu."],
        ["Đặt vé", "Chọn suất, giữ ghế, tạo booking, áp/bỏ mã giảm giá, xử lý hết hạn."],
        ["Thanh toán", "VNPay, quét QR ngân hàng, webhook, payment event, kết quả giao dịch."],
        ["Vé QR", "Sinh QR, gửi email vé, xem vé, tải vé, soát vé đúng rạp/suất."],
        ["Admin/Staff", "Dashboard, CRUD dữ liệu, phân quyền, staff scope, audit, quản lý suất chiếu."],
    ], [4.0, 10.7])
    doc.add_heading("3.3 Đặc tả Use Case và truy vết yêu cầu", level=2)
    para(doc, "Dựa trên các tác nhân chính của hệ thống, các use case được gom theo ba nhóm: khách hàng, nhân viên rạp và quản trị viên. Việc đặt mã UC giúp các yêu cầu, sơ đồ, bảng dữ liệu và kiểm thử có thể truy vết với nhau khi trình bày báo cáo.")
    table(doc, ["Mã UC", "Actor chính", "Use case", "Tiền điều kiện", "Kết quả sau khi hoàn tất"], [
        ["UC01", "USER", "Đăng ký, xác thực email và đăng nhập", "Người dùng có email hợp lệ.", "Tài khoản được kích hoạt, hệ thống cấp access token và refresh token."],
        ["UC02", "USER", "Xem phim, rạp và lịch chiếu", "Có dữ liệu phim/rạp/suất còn hiệu lực.", "Người dùng chọn được phim, rạp, ngày và suất chiếu phù hợp."],
        ["UC03", "USER", "Chọn ghế và giữ ghế", "Suất chiếu chưa bắt đầu, ghế còn AVAILABLE.", "seat_status chuyển HOLD, các client khác nhận realtime update."],
        ["UC04", "USER", "Áp mã giảm giá và thanh toán", "Booking còn PENDING và chưa hết hạn.", "Payment được tạo đúng số tiền; booking chuyển SUCCESS nếu gateway xác nhận thành công."],
        ["UC05", "USER", "Xem vé và lịch sử đơn", "Người dùng đã đăng nhập.", "Người dùng thấy vé hợp lệ, đơn chờ thanh toán, đơn hủy/hết hạn và chi tiết QR."],
        ["UC06", "STAFF", "Soát vé QR đúng rạp/suất", "Staff có rạp phụ trách và chọn suất đang mở check-in.", "Ticket hợp lệ chuyển USED; sai rạp/suất/ngoài giờ bị từ chối."],
        ["UC07", "STAFF", "Quản lý suất trong rạp phụ trách", "Staff được admin gán cinema scope.", "Staff chỉ tạo/sửa/hủy suất chiếu trong phạm vi được phân công."],
        ["UC08", "ADMIN", "Quản lý dữ liệu và phân quyền", "Admin đăng nhập hợp lệ.", "Admin quản lý phim/rạp/phòng/ghế/suất/user/role và gán rạp cho staff."],
        ["UC09", "ADMIN", "Theo dõi dashboard, payment và audit", "Có dữ liệu booking/payment/audit.", "Admin xem số liệu vận hành, filter theo ngày/trạng thái và kiểm tra lịch sử thao tác."],
    ], [1.5, 2.2, 4.4, 5.0, 4.8])
    table(doc, ["Yêu cầu", "Use case liên quan", "Bảng/domain chính", "Điểm kiểm thử tương ứng"], [
        ["Không bán trùng ghế", "UC03, UC04", "seat_status, bookings, booking_details", "TC03, TC04"],
        ["Thanh toán đúng số tiền", "UC04", "payments, payment_events, promotions", "TC05, TC06"],
        ["Vé QR không dùng lại", "UC05, UC06", "tickets, booking_details, showtimes", "TC07"],
        ["Staff chỉ thao tác trong rạp được gán", "UC06, UC07", "staff_cinemas, cinemas, showtimes", "TC07, Staff scope"],
        ["Admin quản trị có audit", "UC08, UC09", "admin_audit_logs, users, roles, permissions", "Admin audit"],
        ["Giao diện phản hồi realtime", "UC03", "seat_status + WebSocket event", "Booking realtime"],
    ], [3.2, 4.2, 4.6, 3.9])

    doc.add_heading("3.4 Sơ đồ phân tích", level=2)
    for cap, path in diagrams:
        figure(doc, cap, path)
    doc.add_heading("3.5 Thiết kế cơ sở dữ liệu", level=2)
    para(doc, "Cơ sở dữ liệu được thiết kế theo hướng tách dữ liệu master và dữ liệu giao dịch. Các bảng phim, rạp, phòng, ghế là dữ liệu cấu hình; các bảng booking, payment, ticket, seat_status là dữ liệu giao dịch thay đổi thường xuyên. Bảng seat_status là trung tâm kiểm soát trạng thái ghế theo từng suất chiếu.")
    table(doc, ["Nhóm bảng", "Các bảng chính", "Ý nghĩa"], [
        ["RBAC/Auth", "users, roles, permissions, users_roles, roles_permissions, refresh_tokens, invalidated_token", "Quản lý tài khoản, phiên đăng nhập và phân quyền."],
        ["Cinema master", "movies, cinemas, rooms, seats, showtimes", "Quản lý phim, rạp, phòng, ghế và lịch chiếu."],
        ["Booking", "seat_status, bookings, booking_details, tickets", "Giữ ghế, đặt vé, sinh vé QR."],
        ["Payment", "payments, payment_events", "Theo dõi giao dịch và sự kiện từ cổng thanh toán."],
        ["Operation", "staff_cinemas, admin_audit_logs, auth_audit_logs", "Giới hạn phạm vi nhân viên và lưu dấu thao tác."],
    ], [3.2, 6.2, 5.3])

    doc.add_heading("3.6 Phân tích các luồng nghiệp vụ trọng tâm", level=2)
    para(doc, "Các luồng dưới đây là phần cốt lõi của hệ thống. Khi bảo vệ, đây là những nội dung nên trình bày rõ vì chúng thể hiện hệ thống không chỉ dừng ở thao tác thêm, sửa, xóa dữ liệu mà còn xử lý được các tình huống thực tế của một nền tảng bán vé.")
    table(doc, ["Luồng", "Mục tiêu nghiệp vụ", "Cách hệ thống xử lý"], [
        ["Đăng ký và xác thực email", "Đảm bảo tài khoản dùng email thật trước khi sử dụng đầy đủ chức năng.", "Backend tạo token xác thực dạng hash, gửi email qua SMTP, khi người dùng bấm link thì cập nhật email_verified."],
        ["Đăng nhập và refresh token", "Giữ trải nghiệm đăng nhập mượt nhưng vẫn hạn chế rủi ro lộ token.", "Access token có thời gian sống ngắn; refresh token được lưu và xoay vòng khi cấp token mới; logout thu hồi phiên."],
        ["Giữ ghế", "Không cho nhiều người mua cùng một ghế trong cùng suất chiếu.", "seat_status chuyển AVAILABLE -> HOLD theo user và hold_until; transaction kiểm tra trạng thái trước khi cập nhật."],
        ["Thanh toán", "Xác nhận đơn chỉ khi cổng thanh toán báo thành công và số tiền hợp lệ.", "VNPay dùng return/callback/IPN; SePay dùng webhook, đối chiếu số tiền và nội dung chuyển khoản."],
        ["Sinh vé QR", "Mỗi ghế có một vé điện tử có thể kiểm tra tại rạp.", "Sau booking SUCCESS, hệ thống sinh ticket ACTIVE cho từng booking_detail và QR được ký bằng secret."],
        ["Soát vé", "Tránh quét nhầm rạp, nhầm suất hoặc dùng lại vé.", "Staff chọn cinemaId/showtimeId; backend kiểm tra QR, booking, rạp, suất, cửa sổ check-in rồi mới set USED."],
        ["Hủy suất chiếu", "Khi rạp hủy suất, khách không nên tự chọn lại ghế cho suất đã bị hủy.", "Showtime chuyển CANCELLED, ticket liên quan bị hủy, booking hiển thị trạng thái xử lý hoàn tiền."],
    ], [3.4, 5.5, 6.4])

    doc.add_heading("3.7 Yêu cầu phi chức năng", level=2)
    table(doc, ["Yêu cầu", "Ý nghĩa", "Thiết kế đáp ứng"], [
        ["Bảo mật", "Bảo vệ tài khoản, token, API admin/staff và QR vé.", "JWT, refresh token, RBAC, staff scope, QR signing, không đưa secret vào client."],
        ["Nhất quán dữ liệu", "Tránh bán trùng ghế hoặc xác nhận sai payment.", "Transaction, constraint, version seat_status, idempotency payment event."],
        ["Hiệu năng", "Danh sách admin, lịch chiếu và seat map cần tải nhanh.", "Index, phân trang, filter, cache dữ liệu ít đổi, tránh N+1."],
        ["Khả dụng vận hành", "Scheduler phải trả ghế và cập nhật trạng thái suất chiếu đúng thời gian.", "Job định kỳ cho booking/hold hết hạn và trạng thái showtime."],
        ["Mở rộng", "Có thể thêm cổng thanh toán, combo, loyalty hoặc mobile app.", "Payment gateway interface, module hóa service, DTO rõ ràng, frontend tách page/component/API."],
        ["Trải nghiệm người dùng", "Người dùng cần hiểu rõ trạng thái vé, thanh toán và lỗi.", "Thông báo tiếng Việt, UI khóa mã giảm giá khi QR đã cố định số tiền, tự chuyển trang khi thanh toán thành công."],
    ], [3.0, 5.0, 7.3])

    doc.add_heading("CHƯƠNG 4: TRIỂN KHAI HỆ THỐNG", level=1)
    sections = [
        ("4.1 Backend Spring Boot", "Backend được triển khai theo mô hình Controller - Service - Repository. Controller nhận request và validate DTO. Service xử lý nghiệp vụ như giữ ghế, tạo booking, thanh toán, soát vé. Repository chịu trách nhiệm truy vấn database, có các query tối ưu để tránh N+1 và hỗ trợ phân trang/filter."),
        ("4.2 Xác thực và phân quyền", "Hệ thống sử dụng Spring Security, JWT, refresh token và RBAC. Các API quan trọng được bảo vệ bằng permission. Với nhân viên, service tiếp tục kiểm tra staff scope theo cinema_id để đảm bảo nhân viên chỉ thao tác tại rạp được phân công."),
        ("4.3 Luồng đặt vé và giữ ghế", "Khi user chọn ghế, backend khóa các dòng seat_status tương ứng, kiểm tra trạng thái AVAILABLE, sau đó chuyển sang HOLD với hold_by và hold_until. Khi tạo booking, hệ thống kiểm tra ghế vẫn thuộc user hiện tại và chưa hết hạn. Nếu thanh toán thành công, ghế chuyển BOOKED; nếu thất bại hoặc hết hạn, ghế được trả AVAILABLE."),
        ("4.4 Thanh toán VNPay và quét QR ngân hàng", "VNPay được xử lý bằng redirect URL và callback/IPN. SePay/VietQR tạo QR ngân hàng với nội dung chuyển khoản riêng cho booking, sau đó webhook tiền vào sẽ đối chiếu nội dung và số tiền để xác nhận tự động. Khi số tiền thay đổi do mã giảm giá, payment pending cũ được làm hết hạn để tránh thanh toán sai."),
        ("4.5 Vé QR và soát vé", "Sau khi booking SUCCESS, hệ thống sinh ticket cho từng ghế. QR được ký bằng secret, không chỉ dựa vào ID thuần. Khi staff quét, backend kiểm tra QR hợp lệ, ticket ACTIVE, booking SUCCESS, đúng rạp, đúng suất chiếu và đúng cửa sổ check-in rồi mới set USED."),
        ("4.6 WebSocket realtime", "Frontend subscribe kênh /topic/seatmap/{showtimeId}. Backend publish event sau khi commit thay đổi trạng thái ghế. Nhờ vậy các user đang xem cùng suất chiếu sẽ thấy ghế chuyển màu ngay khi có người giữ, mua, hủy hoặc hết hạn."),
        ("4.7 Scheduler, cache và tối ưu truy vấn", "Scheduler định kỳ quét ghế HOLD và booking PENDING hết hạn. Cache được dùng cho dữ liệu ít thay đổi như phim, rạp, phòng và khuyến mãi. Các truy vấn admin có phân trang và filter; database có index theo các trường thường dùng như status, created_at, showtime_id, user_id."),
        ("4.8 Frontend React", "Frontend tổ chức theo page, component, API client, route guard và store. Người dùng có các trang public, đặt vé, thanh toán, vé của tôi và hồ sơ. Admin/staff có layout riêng, sidebar, dashboard, quản lý dữ liệu, soát vé QR và rạp phụ trách."),
    ]
    for title, body in sections:
        doc.add_heading(title, level=2)
        para(doc, body)
    doc.add_heading("4.9 Các điểm kỹ thuật nâng cao đã triển khai", level=2)
    para(doc, "Ngoài các chức năng CRUD thông thường, hệ thống được phát triển theo hướng mô phỏng sản phẩm đặt vé thực tế. Các điểm kỹ thuật nâng cao được lựa chọn dựa trên rủi ro nghiệp vụ thường gặp: tranh chấp ghế, thanh toán không đồng bộ, phân quyền nhân viên theo rạp, vé QR dùng lại, dữ liệu admin lớn và trải nghiệm realtime.")
    table(doc, ["Nhóm kỹ thuật", "Triển khai trong backend", "Triển khai trong frontend"], [
        ["Bảo mật phiên đăng nhập", "Spring Security, JWT filter, refresh token rotation, invalidated token, email verification.", "Route guard, lưu phiên an toàn, Axios interceptor tự refresh và retry request."],
        ["Phân quyền", "RBAC theo permission và staff scope theo staff_cinemas.", "Sidebar/menu/form chỉ hiển thị chức năng phù hợp role và phạm vi rạp."],
        ["Giữ ghế realtime", "Transaction khi hold seat_status, version, scheduler trả ghế, publisher WebSocket.", "SeatMap nhận snapshot ban đầu, subscribe topic và cập nhật màu ghế không cần refresh."],
        ["Thanh toán", "VNPay redirect/callback, SePay webhook, payment_events, đối chiếu số tiền/nội dung.", "UI khóa mã giảm giá khi QR đã cố định số tiền, tự chuyển trang khi payment SUCCESS."],
        ["Soát vé QR", "Verify chữ ký QR, ticket ACTIVE, booking SUCCESS, đúng cinemaId/showtimeId, lưu staff check-in.", "Scanner hỗ trợ camera và file ảnh, có bộ lọc rạp/suất đang mở check-in."],
        ["Tối ưu dữ liệu", "Index, phân trang, filter ngày/trạng thái, cache dữ liệu ít đổi, tránh N+1 bằng query phù hợp.", "Danh sách có filter/search, layout compact, hạn chế re-render gây rung UI."],
        ["Vận hành", "Audit log, soft delete, cancel showtime policy, trạng thái showtime tự phản ánh UPCOMING/ONGOING/ENDED.", "Admin/staff nhìn được dữ liệu vận hành, trạng thái rõ màu, thao tác có xác nhận."],
    ], [3.4, 6.0, 5.3])

    doc.add_heading("4.10 Giao diện hệ thống", level=2)
    for idx, name in enumerate([
        "Giao diện trang chủ hiển thị phim đang chiếu và mua vé theo rạp",
        "Giao diện chi tiết phim và lịch chiếu",
        "Giao diện chọn rạp theo thành phố",
        "Giao diện chọn ghế realtime",
        "Giao diện thanh toán VNPay",
        "Giao diện thanh toán quét QR ngân hàng",
        "Giao diện kết quả giao dịch",
        "Giao diện vé của tôi",
        "Giao diện admin dashboard",
        "Giao diện quản lý suất chiếu",
        "Giao diện staff rạp phụ trách",
        "Giao diện soát vé QR",
    ], start=1):
        placeholder(doc, f"CHÈN ẢNH GIAO DIỆN {idx}: {name}")

    doc.add_heading("CHƯƠNG 5: KIỂM THỬ HỆ THỐNG", level=1)
    doc.add_heading("5.1 Chiến lược kiểm thử", level=2)
    para(doc, "Hệ thống được kiểm thử theo nhiều mức: unit test cho các logic nhỏ, integration test cho API và database tạm thời, security test cho phân quyền, manual test cho các luồng thanh toán/email/QR/WebSocket. Các test quan trọng tập trung vào booking, payment, auth, staff scope và exception.")
    table(doc, ["Mã", "Tình huống", "Kết quả mong đợi"], [
        ["TC01", "Đăng ký email mới", "Tạo user và gửi email xác thực."],
        ["TC02", "Đăng nhập sai mật khẩu", "Trả lỗi tiếng Việt, không reload trang."],
        ["TC03", "Hai user giữ cùng ghế", "Chỉ một user giữ thành công."],
        ["TC04", "HOLD hết hạn", "Scheduler trả ghế AVAILABLE và publish realtime."],
        ["TC05", "VNPay thành công", "Booking SUCCESS, Payment SUCCESS, Seat BOOKED, sinh vé QR."],
        ["TC06", "SePay webhook đúng", "Xác nhận payment tự động và gửi email vé."],
        ["TC07", "Staff quét sai rạp", "Không set USED, báo vé không thuộc rạp."],
        ["TC08", "Admin hủy suất", "Booking/ticket liên quan bị hủy và ghi nhận xử lý hoàn tiền."],
    ], [1.6, 7.0, 6.0])
    doc.add_heading("5.2 Phạm vi kiểm thử và mức độ bao phủ", level=2)
    para(doc, "Các kịch bản kiểm thử được lựa chọn theo rủi ro nghiệp vụ thay vì chỉ kiểm tra thao tác CRUD. Trọng tâm là các trạng thái có thể gây lỗi thực tế: giữ ghế đồng thời, thanh toán treo, payment callback/webhook lặp, staff quét sai rạp/suất, token hết hạn và dữ liệu admin cần filter/phân trang.")
    table(doc, ["Nhóm kiểm thử", "Mục tiêu bao phủ", "Lý do quan trọng"], [
        ["Unit test", "Các hàm tính tiền, validate mã giảm giá, định dạng QR/payment description.", "Giảm lỗi logic nhỏ nhưng ảnh hưởng trực tiếp đến tiền và vé."],
        ["Integration test", "API booking, payment, auth, repository query và exception handler.", "Đảm bảo các lớp controller-service-repository hoạt động đúng với database test."],
        ["Security test", "RBAC, staff scope, refresh token, logout, route guard.", "Ngăn user/staff truy cập sai dữ liệu hoặc thao tác ngoài quyền."],
        ["Manual E2E", "VNPay sandbox, SePay webhook, email, camera/file QR scanner, WebSocket realtime.", "Các tích hợp ngoài khó mock hoàn toàn nên cần demo thực tế."],
    ], [3.0, 5.6, 6.0])

    doc.add_heading("CHƯƠNG 6: ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN", level=1)
    doc.add_heading("6.1 Đánh giá kết quả", level=2)
    bullets(doc, [
        "Hệ thống đáp ứng đầy đủ luồng khách hàng, nhân viên và admin.",
        "Luồng giữ ghế và thanh toán có xử lý timeout, tránh ảnh hưởng bán vé.",
        "Phân quyền RBAC kết hợp staff scope phù hợp hệ thống rạp thực tế.",
        "Realtime WebSocket, QR ticket và email vé là các điểm mạnh khi bảo vệ.",
        "Database có constraint, index và migration rõ ràng, dễ bảo trì.",
    ])
    doc.add_heading("6.2 Hạn chế", level=2)
    bullets(doc, [
        "Chưa triển khai refund tự động với API production của từng cổng thanh toán.",
        "Chưa có combo bắp nước, loyalty/member tier và mobile app.",
        "Chưa có CI/CD, monitoring, tracing và alerting production đầy đủ.",
        "Một số ảnh giao diện cần chèn thủ công bằng screenshot cuối cùng trước khi nộp.",
    ])
    doc.add_heading("6.3 Hướng phát triển", level=2)
    bullets(doc, [
        "Tích hợp hoàn tiền tự động và reconciliation với cổng thanh toán.",
        "Bổ sung combo, voucher nâng cao, loyalty và recommendation phim.",
        "Triển khai Redis cho cache phân tán/rate limit/pub-sub khi có nhiều backend instance.",
        "Bổ sung mobile app hoặc PWA, push notification và nhắc lịch chiếu.",
        "Triển khai CI/CD, log tập trung, metrics và dashboard vận hành kỹ thuật.",
    ])

    doc.add_heading("KẾT LUẬN", level=1)
    para(doc, "Đề tài đã xây dựng được hệ thống CinemaBooking.vn với các thành phần quan trọng của một sản phẩm đặt vé xem phim trực tuyến: quản lý phim/rạp/suất/ghế, đặt vé, thanh toán, vé QR, soát vé, gửi email, phân quyền và dashboard quản trị. Hệ thống được thiết kế theo hướng có thể mở rộng, có kiểm soát trạng thái và có cơ chế bảo vệ dữ liệu tốt hơn so với một ứng dụng CRUD đơn giản.")
    para(doc, "Thông qua quá trình thực hiện, sinh viên rèn luyện được tư duy phân tích nghiệp vụ, thiết kế cơ sở dữ liệu, tổ chức code backend/frontend, tích hợp dịch vụ ngoài, xử lý realtime, bảo mật và kiểm thử. Đây là nền tảng quan trọng để tiếp tục phát triển hệ thống trong môi trường thực tế.")

    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    refs = [
        "Tài liệu Spring Boot, Spring Security, Spring Data JPA.",
        "Tài liệu React, TypeScript, Vite, React Router.",
        "Tài liệu PostgreSQL và Flyway Migration.",
        "Tài liệu VNPay Sandbox và quy trình callback/IPN.",
        "Tài liệu SePay/VietQR và webhook xác nhận giao dịch.",
        "Tài liệu Google OAuth 2.0.",
        "Tài liệu ZXing QR Code và Spring Mail.",
        "Mã nguồn hệ thống CinemaBooking.vn trong repository dự án.",
        "Rutgers University, Software Engineering Project Report Requirements: tham khảo cách trình bày use case, system sequence diagram và user interface specification.",
        "Fort Hays State University, Software Engineering Project Report: tham khảo cấu trúc self-contained report gồm yêu cầu, phân tích, thiết kế, domain model và interaction diagram.",
        "Mabini Colleges BSCS Research Manual, System Architecture: tham khảo vai trò của sơ đồ kiến trúc, component và data flow trong báo cáo kỹ thuật.",
    ]
    for r in refs:
        para(doc, r, WD_ALIGN_PARAGRAPH.LEFT)

def append_submission_appendices(doc: Document):
    doc.add_page_break()
    doc.add_heading("PHỤ LỤC A: KẾ HOẠCH KIỂM THỬ HỆ THỐNG", level=1)
    para(doc, "Phụ lục này trình bày các kịch bản kiểm thử tiêu biểu phục vụ nghiệm thu và demo bảo vệ. Nội dung tập trung vào những luồng có rủi ro nghiệp vụ cao như giữ ghế đồng thời, thanh toán không đồng bộ, phân quyền nhân viên theo rạp và soát vé QR.")
    table(doc, ["Nhóm kiểm thử", "Kịch bản", "Điều kiện/dữ liệu kiểm thử", "Kết quả mong đợi"], [
        ["Xác thực", "Đăng ký tài khoản, xác thực email, đăng nhập, refresh token và logout.", "Email hợp lệ, mật khẩu đạt chính sách, link xác thực còn hạn.", "Tài khoản chỉ đăng nhập ổn định sau khi xác thực; refresh token được xoay vòng; logout thu hồi phiên."],
        ["Phân quyền", "USER truy cập trang admin; STAFF thao tác ngoài rạp phụ trách; ADMIN gán rạp cho STAFF.", "Ba tài khoản đại diện USER/STAFF/ADMIN.", "Backend trả lỗi quyền rõ ràng; frontend ẩn chức năng không thuộc quyền; STAFF chỉ thấy dữ liệu trong scope."],
        ["Giữ ghế realtime", "Hai trình duyệt mở cùng seat map, một bên giữ ghế.", "Cùng showtime, cùng danh sách ghế đang AVAILABLE.", "Ghế chuyển HOLD realtime ở trình duyệt còn lại, tránh hai người chọn trùng."],
        ["Hết hạn giữ ghế", "Giữ ghế nhưng không thanh toán đến khi quá hạn.", "Cấu hình timeout giữ ghế ngắn để demo.", "Scheduler chuyển booking/payment sang hết hạn, trả seat_status về AVAILABLE và phát WebSocket update."],
        ["Áp mã giảm giá", "Áp mã hợp lệ, mã hết hạn, mã không đủ điều kiện và mã đã khóa sau khi tạo QR.", "Các promotion trong dữ liệu mẫu.", "Tổng tiền tính đúng; thông báo lỗi tiếng Việt; QR thanh toán yêu cầu tạo lại nếu số tiền thay đổi."],
        ["VNPay", "Thanh toán sandbox thành công, thất bại và quay lại khi booking hết hạn.", "Return URL/IPN public, booking PENDING còn hạn.", "Xác thực checksum, cập nhật payment/booking/seat/ticket đúng trạng thái, không xử lý lặp."],
        ["SePay/VietQR", "Tạo QR chuyển khoản, nhận webhook đúng/sai số tiền hoặc sai nội dung.", "Webhook có API key/HMAC; description chứa mã booking.", "Giao dịch đúng được xác nhận tự động; sai amount/content không xác nhận vé."],
        ["Email vé", "Thanh toán thành công và kiểm tra email vé.", "SMTP sandbox hoặc email thật đã cấu hình.", "Email chứa thông tin phim, rạp, phòng, ghế, địa chỉ, tổng tiền và QR từng vé."],
        ["Soát vé QR", "Quét QR đúng rạp/suất, sai rạp, sai suất, ngoài cửa sổ check-in và vé đã dùng.", "Tài khoản STAFF có rạp phụ trách.", "Chỉ QR đúng điều kiện mới chuyển USED; lỗi nghiệp vụ hiển thị rõ cho nhân viên."],
        ["Admin/Staff vận hành", "Tạo, lọc, hủy suất chiếu; xem đơn đặt vé/thanh toán theo ngày; xem audit.", "Có dữ liệu nhiều ngày, nhiều rạp.", "Phân trang/filter ổn định; hủy suất xử lý vé/hoàn tiền theo chính sách; audit ghi lại thao tác quan trọng."],
    ], [3.0, 5.2, 5.0, 5.0])

    doc.add_heading("PHỤ LỤC B: DANH SÁCH GIAO DIỆN MINH HỌA", level=1)
    para(doc, "Các ảnh giao diện nên được chụp từ hệ thống chạy thật để thay vào phần minh họa ở Chương 4. Ảnh cần rõ bố cục, không che nội dung chính, nhưng phải ẩn thông tin bí mật như mật khẩu, secret, access key, token hoặc số tài khoản nhạy cảm nếu không cần thiết.")
    table(doc, ["STT", "Giao diện cần chèn", "Mục đích minh họa"], [
        ["1", "Trang chủ khách hàng", "Thể hiện danh sách phim, mua vé theo phim và mua vé theo rạp."],
        ["2", "Trang chi tiết phim/lịch chiếu", "Thể hiện cách người dùng chọn ngày, rạp và suất chiếu."],
        ["3", "Trang chọn rạp theo thành phố và bản đồ", "Thể hiện UX lọc thành phố, danh sách rạp và bản đồ vị trí."],
        ["4", "Trang chọn ghế realtime", "Thể hiện seat map, màu ghế, trạng thái HOLD/BOOKED và thông tin vé."],
        ["5", "Trang thanh toán VNPay", "Thể hiện luồng chuyển sang cổng thanh toán."],
        ["6", "Trang thanh toán quét QR ngân hàng", "Thể hiện QR, số tiền, nội dung chuyển khoản và trạng thái chờ xác nhận."],
        ["7", "Trang kết quả giao dịch", "Thể hiện kết quả thành công/thất bại và hướng đi tiếp theo cho người dùng."],
        ["8", "Trang Vé của tôi", "Thể hiện vé hợp lệ, đơn chờ thanh toán, đơn hủy và chi tiết QR."],
        ["9", "Admin dashboard", "Thể hiện số liệu tổng quan, biểu đồ và top phim."],
        ["10", "Admin quản lý phim/rạp/phòng/ghế/suất", "Thể hiện CRUD, filter, phân trang và trạng thái dữ liệu."],
        ["11", "Admin quản lý người dùng và gán rạp staff", "Thể hiện RBAC và staff scope."],
        ["12", "Staff rạp phụ trách và soát vé QR", "Thể hiện phạm vi rạp, chọn suất đang check-in, camera/file scan QR."],
    ], [1.2, 5.4, 7.2])

    doc.add_heading("PHỤ LỤC C: MỘT SỐ API/MODULE TIÊU BIỂU", level=1)
    para(doc, "Phụ lục này không thay thế tài liệu API chi tiết. Bảng chỉ chọn những nhóm API/module tiêu biểu để người đọc thấy được mối liên hệ giữa phân tích nghiệp vụ, thiết kế backend và giao diện frontend.")
    table(doc, ["Nhóm API/module", "Vai trò trong hệ thống", "Điểm cần giải thích khi bảo vệ"], [
        ["Auth API", "Đăng ký, xác thực email, đăng nhập, Google login, refresh token, logout.", "Vì sao dùng access token ngắn hạn, refresh token rotation và token blacklist."],
        ["Movie/Cinema/Showtime API", "Cung cấp dữ liệu phim, rạp, phòng, ghế và lịch chiếu cho web khách hàng/admin/staff.", "Cách phân trang, filter ngày, soft delete và tránh hiện suất đã qua cho khách hàng."],
        ["Booking API", "Giữ ghế, tạo booking, áp mã giảm giá, hủy booking và xử lý hết hạn.", "Transaction, optimistic locking, seat_status và WebSocket sau khi thay đổi trạng thái."],
        ["Payment API", "Khởi tạo VNPay/SePay, xử lý callback/webhook và kiểm tra trạng thái thanh toán.", "Checksum/HMAC, idempotency, đối chiếu số tiền/nội dung chuyển khoản và payment event."],
        ["Ticket API", "Trả vé điện tử, sinh QR và soát vé đúng rạp/suất.", "QR ký bằng secret, trạng thái ACTIVE/USED/CANCELLED, lưu giờ và nhân viên check-in."],
        ["Admin/Staff API", "Quản trị dữ liệu và vận hành theo phạm vi vai trò.", "RBAC ở security layer kết hợp staff scope ở service layer."],
        ["Frontend API client", "Gọi backend, gắn token, xử lý lỗi và điều hướng người dùng.", "Axios interceptor, route guard, thông báo lỗi tiếng Việt và UX theo trạng thái booking/payment."],
    ], [3.4, 5.5, 6.0])

    doc.add_heading("PHỤ LỤC D: GHI CHÚ TRIỂN KHAI THỰC NGHIỆM", level=1)
    para(doc, "Phần này chỉ mô tả các nhóm cấu hình ở mức khái quát để phục vụ chạy thử và demo. Báo cáo không đưa giá trị bí mật thật vào tài liệu nhằm đảm bảo an toàn thông tin.")
    table(doc, ["Thành phần", "Vai trò khi triển khai", "Lưu ý"], [
        ["Frontend React", "Giao diện khách hàng, nhân viên và admin.", "Khi chạy local dùng Vite; khi triển khai thật cần build static và cấu hình API base URL phù hợp."],
        ["Backend Spring Boot", "Cung cấp REST API, WebSocket, scheduler, email và tích hợp thanh toán.", "Các secret JWT/payment/mail nên lấy từ biến môi trường hoặc secret manager."],
        ["PostgreSQL", "Lưu dữ liệu nghiệp vụ, RBAC, booking, payment, ticket và audit.", "Dùng Flyway/schema versioning để tránh lệch cấu trúc database giữa môi trường."],
        ["SMTP Mail", "Gửi email xác thực, reset password và vé điện tử.", "Có thể dùng sandbox khi demo; production cần cấu hình domain/mail uy tín."],
        ["VNPay/SePay", "Xác nhận thanh toán trực tuyến và chuyển khoản QR.", "Callback/webhook cần public URL, xác thực chữ ký và xử lý idempotent."],
        ["WebSocket", "Đẩy realtime trạng thái ghế cho các client đang xem cùng suất.", "Frontend subscribe theo showtimeId; backend publish sau khi cập nhật database."],
    ], [3.2, 6.2, 5.8])


def build():
    diagrams = generate_diagrams()
    doc = setup_doc()
    cover(doc)
    front_matter(doc, diagrams)
    chapters(doc, diagrams)
    append_submission_appendices(doc)
    try:
        doc.save(OUT)
        print(OUT)
    except PermissionError:
        fallback = OUT.with_name(f"{OUT.stem}_Updated.docx")
        doc.save(fallback)
        print(fallback)


if __name__ == "__main__":
    build()

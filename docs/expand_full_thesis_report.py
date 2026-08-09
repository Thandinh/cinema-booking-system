from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parent
REPORT_PATH = DOCS_DIR / "Bao_cao_khoa_luan_CinemaBooking_Full.docx"

SOURCE_DOCS = [
    ("PHU LUC F. LUONG AUTH, JWT, REFRESH TOKEN VA GOOGLE LOGIN", DOCS_DIR / "Luong_Auth_JWT_Session_GoogleLogin_CinemaBooking.docx"),
    ("PHU LUC G. RBAC, PERMISSION VA STAFF SCOPE", DOCS_DIR / "Luong_RBAC_Permission_StaffScope_CinemaBooking.docx"),
    ("PHU LUC H. SCHEDULER GIU GHE, HET HAN VA TRANG THAI SUAT CHIEU", DOCS_DIR / "Luong_Scheduler_GiuGhe_HetHan_TrangThaiSuatChieu_CinemaBooking.docx"),
    ("PHU LUC I. EXCEPTION, AUDIT LOG VA SOFT DELETE", DOCS_DIR / "Luong_Exception_Audit_SoftDelete_CinemaBooking.docx"),
    ("PHU LUC J. CACHE, QUERY, INDEX VA FLYWAY", DOCS_DIR / "Luong_Cache_Query_Index_Flyway_CinemaBooking.docx"),
    ("PHU LUC K. FRONTEND REACT, ROUTER, STATE, API VA UX", DOCS_DIR / "Luong_Frontend_React_Router_State_API_UX_CinemaBooking.docx"),
    ("PHU LUC L. EMAIL VERIFICATION, RESET PASSWORD VA TICKET EMAIL", DOCS_DIR / "Luong_Email_Verification_Reset_Ticket_Notification_CinemaBooking.docx"),
    ("PHU LUC M. ADMIN, STAFF VA VAN HANH RAP", DOCS_DIR / "Luong_Admin_Staff_Operation_CinemaBooking.docx"),
    ("PHU LUC N. LUONG THANH TOAN VNPAY VA QUET QR NGAN HANG", DOCS_DIR / "Luong_thanh_toan_VNPay_SePay_CinemaBooking.docx"),
    ("PHU LUC O. WEBSOCKET REALTIME SEAT MAP", DOCS_DIR / "Luong_WebSocket_Realtime_SeatMap_CinemaBooking.docx"),
]


def normalize_text(text: str) -> str:
    return " ".join((text or "").split())


def add_para(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if style is None:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p


def add_note(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(85, 99, 118)


def append_source_doc(target: Document, title: str, source_path: Path):
    if not source_path.exists():
        target.add_heading(title, level=1)
        add_note(target, f"Chua tim thay file nguon: {source_path.name}")
        return

    target.add_heading(title, level=1)
    add_note(
        target,
        "Phu luc nay duoc bien soan tu tai lieu hoc rieng cua chinh du an CinemaBooking.vn. "
        "No giup nguoi doc hieu sau hon ve luong code, class, config va cach van hanh thuc te.",
    )

    source = Document(source_path)
    copied = 0
    for para in source.paragraphs:
        text = normalize_text(para.text)
        if not text:
            continue

        style_name = para.style.name if para.style is not None else ""
        if style_name.startswith("Title"):
            # The appendix title already identifies the source section.
            continue
        if style_name == "Heading 1":
            target.add_heading(text, level=2)
        elif style_name == "Heading 2":
            target.add_heading(text, level=3)
        elif style_name == "Heading 3":
            p = target.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(11, 31, 51)
        else:
            add_para(target, text)
        copied += 1

    for table in source.tables:
        if not table.rows:
            continue
        rows = [[normalize_text(cell.text) for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        new_table = target.add_table(rows=0, cols=len(rows[0]))
        new_table.style = "Table Grid"
        for row in rows:
            cells = new_table.add_row().cells
            for idx, value in enumerate(row[: len(cells)]):
                cells[idx].text = value
        copied += 1

    add_note(target, f"Da chen {copied} muc noi dung tu {source_path.name}.")


def add_final_defense_pack(doc: Document):
    doc.add_heading("PHU LUC P. BO CAU HOI ON TAP BAO VE", level=1)
    add_note(
        doc,
        "Phan nay danh cho sinh vien on tap truoc khi bao ve. Cau tra loi nen duoc dien dat bang loi cua minh, "
        "khong hoc thuoc may moc.",
    )

    groups = [
        (
            "Nhom cau hoi ve kien truc",
            [
                ("Vi sao tach backend va frontend?", "De backend tap trung nghiep vu/API/bao mat, frontend tap trung UX. Cach tach nay de deploy rieng, scale rieng va thay doi giao dien ma it anh huong logic server."),
                ("Vi sao dung Spring Boot?", "Spring Boot ho tro nhanh cho REST API, Security, JPA, Validation, Mail, WebSocket, cache va test. He sinh thai phu hop ung dung doanh nghiep."),
                ("Vi sao dung PostgreSQL?", "PostgreSQL manh ve rang buoc du lieu, transaction, index, JSONB, UUID va phu hop cac luong booking/payment can tinh nhat quan cao."),
                ("Vi sao dung Flyway?", "Flyway giup database co version ro rang, moi moi truong chay dung thu tu migration thay vi sua tay kho kiem soat."),
                ("Open-in-view false co y nghia gi?", "No buoc service phai lay du lieu can thiet trong transaction, giam nguy co lazy query bat ngo va tranh N+1 an trong tang controller."),
            ],
        ),
        (
            "Nhom cau hoi ve bao mat",
            [
                ("Access token va refresh token khac nhau the nao?", "Access token song ngan de goi API; refresh token song dai hon de cap access token moi. Refresh token duoc luu hash va rotation de giam rui ro bi danh cap."),
                ("Logout xu ly the nao?", "Client goi API logout, backend revoke refresh token hien tai va co the invalid access token den het han. Client xoa state/token va chuyen ve trang dang nhap."),
                ("RBAC va staff scope khac nhau gi?", "RBAC tra loi user co quyen lam hanh dong nay khong. Staff scope tra loi neu co quyen thi duoc lam tren rap/suat nao."),
                ("Tai sao staff khong duoc xem toan bo du lieu?", "Trong thuc te nhan vien chi van hanh rap duoc phan cong. Gioi han scope giam lo du lieu va giam thao tac nham."),
                ("QR ve co an toan khong?", "QR khong chi chua ID thuan; no duoc ky bang secret. Backend verify chu ky, trang thai ticket, booking, rap, suat, cua so check-in roi moi set USED."),
            ],
        ),
        (
            "Nhom cau hoi ve booking va thanh toan",
            [
                ("Tai sao can giu ghe?", "Nguoi dung can thoi gian nhap thong tin va thanh toan. HOLD ngan tranh hai nguoi cung mua mot ghe trong cung thoi diem."),
                ("Neu user bo man hinh thanh toan thi sao?", "Booking PENDING co payment_expires_at. Scheduler quet booking/hold het han, set EXPIRED va tra ghe AVAILABLE."),
                ("Tai sao SePay QR bi khoa ma giam gia sau khi tao?", "QR ngan hang co so tien/noi dung co dinh. Neu doi tong tien sau khi tao QR thi co nguy co chuyen sai so tien, nen product that thuong yeu cau tao QR moi."),
                ("VNPay va SePay khac nhau gi?", "VNPay la cong redirect sang sandbox/payment page va callback/IPN. SePay la chuyen khoan QR, backend nhan webhook tien vao va doi chieu noi dung/sotien."),
                ("Sau thanh toan thanh cong he thong lam gi?", "Cap nhat payment SUCCESS, booking SUCCESS, seat BOOKED, sinh ticket QR, gui email ve, publish realtime va cho user xem ket qua."),
            ],
        ),
        (
            "Nhom cau hoi ve realtime va hieu nang",
            [
                ("WebSocket dung de lam gi?", "De cac man hinh dang xem seat map nhan thay doi HOLD/BOOKED/AVAILABLE ngay ma khong can refresh."),
                ("Tai sao publish sau commit?", "Neu publish truoc commit ma DB rollback thi client se thay trang thai sai. Sau commit dam bao thong tin da ben vung."),
                ("Index la gi?", "Index giong muc luc sach, giup DB tim dong nhanh hon theo cot hay truy van hay dung, doi lai ton them dung luong va chi phi khi insert/update."),
                ("Cache nen dung cho bang nao?", "Nen cache du lieu it doi nhu phim, rap, phong, promotion public. Khong nen cache lau seat_status vi no thay doi lien tuc."),
                ("Tranh N+1 nhu the nao?", "Dung fetch join/entity graph/projection/query theo id page roi fetch details, batch fetch size va tat open-in-view de phat hien som."),
            ],
        ),
        (
            "Nhom cau hoi ve UI/UX",
            [
                ("Vi sao trang chu khong bat dang nhap khi xem phim/rap?", "Xem thong tin la hanh vi public. Chi bat dang nhap khi giu ghe/dat ve/thanh toan de giam friction."),
                ("Vi sao lich chieu chi hien cac ngay toi?", "Product thuc te khong nen hien suat da qua. Thuong hien hom nay va vai ngay toi de nguoi dung chon nhanh."),
                ("Vi sao staff scanner can chon rap/suat?", "De tranh quet nham rap/suat. Ve dung QR nhung sai rap hoac sai suat se khong bi danh dau USED."),
                ("Vi sao trang ve can tach ve hop le va don da dat?", "Ve hop le de vao rap can QR ro rang; don da dat gom pending/cancel/failed/expired can theo doi trang thai."),
                ("Vi sao admin can filter theo ngay?", "Booking/payment/showtime phat sinh nhieu theo thoi gian. Loc theo ngay giup van hanh va doi soat nhanh hon."),
            ],
        ),
    ]

    for group_title, rows in groups:
        doc.add_heading(group_title, level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Cau hoi"
        table.rows[0].cells[1].text = "Y tra loi goi y"
        for question, answer in rows:
            cells = table.add_row().cells
            cells[0].text = question
            cells[1].text = answer


def main():
    doc = Document(REPORT_PATH)
    doc.add_page_break()
    intro = doc.add_heading("PHAN PHU LUC CHUYEN SAU DE HOC VA BAO VE", level=1)
    intro.runs[0].font.color.rgb = RGBColor(13, 71, 161)
    add_para(
        doc,
        "Phan sau mo rong bao cao chinh bang cac tai lieu giai thich rieng tung luong nang cao trong he thong. "
        "No khong thay the cac chuong bao cao, ma dong vai tro nhu so tay ky thuat kem theo de sinh vien nam ro code, class, annotation, config va luong xu ly.",
    )

    for title, path in SOURCE_DOCS:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        append_source_doc(doc, title, path)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_final_defense_pack(doc)
    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    main()

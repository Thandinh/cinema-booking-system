from __future__ import annotations

import json
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document


REPORT = (
    Path(sys.argv[1]).resolve()
    if len(sys.argv) > 1
    else Path(__file__).resolve().parent / "Bao_cao_khoa_luan_CinemaBooking_Final.docx"
)


def is_contiguous(sequence: list[int]) -> bool:
    return sequence == list(range(1, max(sequence, default=0) + 1))


def main() -> None:
    document = Document(REPORT)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    headings = [
        (paragraph.style.name, paragraph.text.strip())
        for paragraph in document.paragraphs
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip()
    ]
    captions = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.style.name in {"Table Caption", "Figure Caption"} and paragraph.text.strip()
    ]
    table_texts = [
        "\n".join(cell.text.strip() for row in table.rows for cell in row.cells if cell.text.strip())
        for table in document.tables
    ]
    caption_counts = Counter(captions)
    with zipfile.ZipFile(REPORT) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        invalid_zip_member = archive.testzip()

    table_numbers = [caption for caption in captions if caption.startswith("Bảng ")]
    figure_numbers = [caption for caption in captions if caption.startswith("Hình ")]
    figure_sequences: dict[str, list[int]] = {}
    for chapter in (3, 4):
        pattern = re.compile(rf"^Hình {chapter}\.(\d+)\.")
        figure_sequences[str(chapter)] = [
            int(match.group(1))
            for caption in figure_numbers
            if (match := pattern.match(caption))
        ]
    placeholders = [
        text
        for text in [*paragraphs, *table_texts]
        if "[CHÈN " in text
    ]
    result = {
        "paragraphs": len(paragraphs),
        "tables": len(document.tables),
        "headings": len(headings),
        "media": len(media),
        "placeholders": len(placeholders),
        "diagram_placeholders": sum(text.startswith("[CHÈN FILE:") for text in placeholders),
        "ui_placeholders": sum(text.startswith("[CHÈN ẢNH") for text in placeholders),
        "literal_backticks": sum("`" in text for text in paragraphs),
        "source_terms": [text for text in paragraphs if re.search(r"\bsource\b", text, re.IGNORECASE)],
        "figure_captions": len(figure_numbers),
        "table_captions": len(table_numbers),
        "duplicate_captions": [caption for caption, count in caption_counts.items() if count > 1],
        "figure_sequences": figure_sequences,
        "figure_sequences_valid": {
            "3": is_contiguous(figure_sequences["3"]),
            "4": is_contiguous(figure_sequences["4"]),
        },
        "zip_valid": invalid_zip_member is None,
        "last_headings": headings[-25:],
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()

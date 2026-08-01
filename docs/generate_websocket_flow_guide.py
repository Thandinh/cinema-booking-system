from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/Luong_WebSocket_Realtime_SeatMap_CinemaBooking.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="111827", size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_code(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.style = "CodeBlock"
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    return paragraph


def add_callout(doc, title, body, fill="EFF6FF", title_color="1D4ED8"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""

    title_paragraph = cell.paragraphs[0]
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor.from_string(title_color)

    body_paragraph = cell.add_paragraph()
    body_paragraph.paragraph_format.space_before = Pt(3)
    body_run = body_paragraph.add_run(body)
    body_run.font.name = "Arial"
    body_run.font.size = Pt(9.5)
    body_run.font.color.rgb = RGBColor.from_string("1F2937")
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, True, "FFFFFF", 9)
        shade_cell(header_cells[index], "111827")
        if widths:
            header_cells[index].width = widths[index]

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, index == 0, "111827", 8.8)
            if widths:
                cells[index].width = widths[index]

    doc.add_paragraph()
    return table


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in [
        ("Title", 24, "111827"),
        ("Heading 1", 16, "0F172A"),
        ("Heading 2", 13, "1F2937"),
        ("Heading 3", 11, "374151"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    if "CodeBlock" not in styles:
        code_style = styles.add_style("CodeBlock", 1)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(8.5)
        code_style.font.color.rgb = RGBColor(31, 41, 55)
        code_style.paragraph_format.left_indent = Inches(0.18)
        code_style.paragraph_format.right_indent = Inches(0.18)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(8)


def build_doc():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Luồng WebSocket Realtime Seat Map")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "CinemaBooking.vn - giải thích backend, frontend, class, code và checklist test"
    )
    subtitle_run.italic = True
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(10.5)

    add_callout(
        doc,
        "Mục tiêu tài liệu",
        "Tài liệu này giải thích toàn bộ luồng realtime của sơ đồ ghế: từ lúc backend mở endpoint WebSocket, "
        "frontend subscribe topic theo suất chiếu, user giữ ghế, thanh toán thành công, hủy đơn, thanh toán thất bại, "
        "đến khi scheduler nhả ghế hết hạn. Cách viết đi từ dễ hiểu đến chi tiết code để bạn có thể đọc, debug và trình bày khi bảo vệ.",
        "ECFDF5",
        "047857",
    )

    doc.add_heading("1. WebSocket trong hệ thống dùng để làm gì?", level=1)
    doc.add_paragraph(
        "Trong hệ thống đặt vé xem phim, nhiều người có thể cùng mở một suất chiếu và cùng nhìn vào sơ đồ ghế. "
        "Nếu một người vừa giữ ghế A1, những người còn lại phải thấy ghế A1 đổi trạng thái gần như ngay lập tức. "
        "Nếu chỉ dùng HTTP thông thường, client phải tự refresh liên tục và trải nghiệm sẽ chậm. WebSocket giải quyết phần này bằng cách giữ một kết nối lâu dài giữa trình duyệt và backend."
    )
    add_bullets(
        doc,
        [
            "HTTP API vẫn là nguồn lấy dữ liệu ban đầu: frontend gọi API để tải toàn bộ seat map.",
            "WebSocket là kênh đẩy sự kiện nhỏ: backend chỉ gửi những ghế vừa đổi trạng thái.",
            "Database vẫn là nguồn sự thật cuối cùng. WebSocket chỉ thông báo, không thay thế transaction trong DB.",
            "Mỗi suất chiếu có một kênh riêng: /topic/seatmap/{showtimeId}.",
            "Client nào đang xem đúng suất chiếu đó thì nhận event và tự đổi màu ghế trên giao diện.",
        ],
    )
    add_callout(
        doc,
        "Câu hiểu nhanh",
        "API giống như chụp ảnh toàn bộ sơ đồ ghế tại thời điểm mở trang. WebSocket giống như loa thông báo: mỗi khi một ghế đổi trạng thái, backend phát tin cho tất cả màn hình đang xem suất chiếu đó.",
        "FFF7ED",
        "C2410C",
    )

    doc.add_heading("2. Sơ đồ tổng quan luồng realtime", level=1)
    add_code(
        doc,
        """User A chọn ghế
  -> Frontend gọi POST /api/v1/bookings/hold
  -> BookingServiceImpl khóa ghế bằng DB lock
  -> seat_status chuyển AVAILABLE -> HOLD
  -> Transaction commit
  -> SeatStatusPublisher gửi SeatStatusEvent
  -> Topic /topic/seatmap/{showtimeId}
  -> Mọi client đang xem suất chiếu nhận event
  -> SeatSelectionPage cập nhật state và đổi màu ghế"""
    )
    doc.add_paragraph(
        "Điểm quan trọng nhất là backend luôn cập nhật database trước rồi mới publish WebSocket sau khi transaction commit. "
        "Vì vậy frontend không bị báo nhầm khi DB rollback."
    )

    doc.add_heading("3. Backend mở cổng WebSocket như thế nào?", level=1)
    add_table(
        doc,
        ["Class", "Vai trò", "Chi tiết quan trọng"],
        [
            (
                "WebSocketConfig",
                "Cấu hình STOMP WebSocket cho Spring Boot.",
                "Bật @EnableWebSocketMessageBroker, mở endpoint /ws và /ws-native, bật simple broker /topic.",
            ),
            (
                "SecurityConfig",
                "Cho phép client kết nối WebSocket.",
                "Permit /ws/** và /ws-native/**. Việc xem trạng thái ghế là public, nhưng API giữ ghế/thanh toán vẫn cần JWT.",
            ),
            (
                "SeatStatusEvent",
                "Payload sự kiện ghế gửi xuống frontend.",
                "Chứa showtimeId, seatId, status, heldByUserId, holdUntil, eventTime.",
            ),
            (
                "SeatStatusPublisher",
                "Lớp chuyên trách gửi event ra topic.",
                "Dùng SimpMessagingTemplate.convertAndSend('/topic/seatmap/' + showtimeId, event).",
            ),
        ],
        [Inches(1.55), Inches(2.1), Inches(2.85)],
    )

    doc.add_heading("3.1. WebSocketConfig.java", level=2)
    doc.add_paragraph(
        "Class này là nơi backend khai báo mình hỗ trợ WebSocket theo giao thức STOMP. STOMP giúp client subscribe theo topic, giống như đăng ký nghe một kênh cụ thể."
    )
    add_code(
        doc,
        """@Configuration
@EnableWebSocketMessageBroker
public class WebSocketConfig implements WebSocketMessageBrokerConfigurer {

    @Override
    public void configureMessageBroker(MessageBrokerRegistry registry) {
        registry.enableSimpleBroker("/topic");
        registry.setApplicationDestinationPrefixes("/app");
    }

    @Override
    public void registerStompEndpoints(StompEndpointRegistry registry) {
        registry.addEndpoint("/ws").setAllowedOriginPatterns("*").withSockJS();
        registry.addEndpoint("/ws-native").setAllowedOriginPatterns("*");
    }
}"""
    )
    add_bullets(
        doc,
        [
            "registry.enableSimpleBroker('/topic'): server có thể publish message xuống các topic bắt đầu bằng /topic.",
            "registry.setApplicationDestinationPrefixes('/app'): dự phòng cho trường hợp client gửi message lên server qua STOMP. Hiện luồng seat map chủ yếu là server push xuống client.",
            "/ws: endpoint có SockJS fallback, phù hợp môi trường cũ hoặc proxy chặn WebSocket thuần.",
            "/ws-native: endpoint WebSocket thuần, frontend hiện đang dùng endpoint này.",
            "setAllowedOriginPatterns('*'): dev/test dễ chạy nhiều host. Khi production nên giới hạn origin về domain thật như https://cinemabooking.vn.",
        ],
    )

    doc.add_heading("3.2. SecurityConfig.java", level=2)
    doc.add_paragraph(
        "WebSocket endpoint được public để người dùng chưa đăng nhập vẫn có thể xem trạng thái ghế. Tuy nhiên điều này không có nghĩa là ai cũng giữ ghế được. "
        "Giữ ghế, tạo booking và thanh toán vẫn đi qua REST API có JWT và permission."
    )
    add_code(
        doc,
        """private static final String[] PUBLIC_WS_ENDPOINTS = {
    "/ws/**",
    "/ws/info/**",
    "/ws-native/**"
};

httpSecurity.authorizeHttpRequests(request -> request
    .requestMatchers(PUBLIC_WS_ENDPOINTS).permitAll()
    // các API nghiệp vụ khác vẫn kiểm tra JWT/permission
);"""
    )
    add_callout(
        doc,
        "Vì sao WebSocket public vẫn an toàn?",
        "Client chỉ nhận trạng thái ghế. Client không thể tự gửi event để đổi ghế thành BOOKED. Muốn giữ ghế hoặc thanh toán bắt buộc phải gọi API có kiểm tra đăng nhập, quyền và transaction trong database.",
        "EFF6FF",
        "1D4ED8",
    )

    doc.add_heading("4. Payload realtime: SeatStatusEvent.java", level=1)
    doc.add_paragraph(
        "Mỗi khi ghế đổi trạng thái, backend đóng gói thông tin vào SeatStatusEvent rồi gửi xuống frontend. Payload này nhỏ, dễ parse và chỉ chứa dữ liệu cần thiết để đổi UI."
    )
    add_code(
        doc,
        """@Getter
@Builder
public class SeatStatusEvent {
    UUID showtimeId;
    UUID seatId;
    SeatStatusType status;      // AVAILABLE | HOLD | BOOKED
    UUID heldByUserId;          // chỉ có khi status = HOLD
    LocalDateTime holdUntil;    // chỉ có khi status = HOLD

    @Builder.Default
    LocalDateTime eventTime = LocalDateTime.now();
}"""
    )
    add_table(
        doc,
        ["Field", "Ý nghĩa trên frontend"],
        [
            ("showtimeId", "Xác định event thuộc suất chiếu nào. Topic cũng chứa showtimeId."),
            ("seatId", "Ghế nào vừa đổi trạng thái."),
            ("status", "Trạng thái mới: AVAILABLE, HOLD hoặc BOOKED."),
            ("heldByUserId", "Nếu ghế HOLD, frontend biết ghế do mình giữ hay người khác giữ để tô màu khác nhau."),
            ("holdUntil", "Thời điểm hết hạn giữ ghế, dùng cho countdown hoặc debug."),
            ("eventTime", "Thời điểm server phát event, hữu ích khi debug lệch thời gian hoặc log realtime."),
        ],
        [Inches(1.7), Inches(4.8)],
    )

    doc.add_heading("5. Backend phát event: SeatStatusPublisher.java", level=1)
    doc.add_paragraph(
        "Thay vì để mọi service tự gọi SimpMessagingTemplate, hệ thống gom việc publish vào SeatStatusPublisher. "
        "Cách này làm code sạch hơn: BookingService chỉ cần nói 'ghế này HOLD/BOOKED/AVAILABLE', còn publisher lo gửi đúng topic."
    )
    add_code(
        doc,
        """@Service
@RequiredArgsConstructor
public class SeatStatusPublisher {
    private final SimpMessagingTemplate messagingTemplate;

    public void publishHold(UUID showtimeId, UUID seatId, UUID userId, LocalDateTime holdUntil) {
        SeatStatusEvent event = SeatStatusEvent.builder()
            .showtimeId(showtimeId)
            .seatId(seatId)
            .status(SeatStatusType.HOLD)
            .heldByUserId(userId)
            .holdUntil(holdUntil)
            .build();
        send(showtimeId, event);
    }

    public void publishBulk(UUID showtimeId, List<UUID> seatIds, SeatStatusType status) {
        seatIds.forEach(seatId -> send(showtimeId, SeatStatusEvent.builder()
            .showtimeId(showtimeId)
            .seatId(seatId)
            .status(status)
            .build()));
    }

    private void send(UUID showtimeId, SeatStatusEvent event) {
        messagingTemplate.convertAndSend("/topic/seatmap/" + showtimeId, event);
    }
}"""
    )
    add_bullets(
        doc,
        [
            "publishHold dùng cho HOLD vì cần gửi thêm heldByUserId và holdUntil.",
            "publishBulk dùng cho BOOKED hoặc AVAILABLE vì nhiều ghế có thể đổi trạng thái cùng lúc.",
            "Destination luôn có dạng /topic/seatmap/{showtimeId}.",
            "Frontend không cần hỏi lại toàn bộ sơ đồ ghế khi nhận event; chỉ update đúng seatId.",
        ],
    )

    doc.add_heading("6. Luồng giữ ghế realtime", level=1)
    doc.add_paragraph(
        "Khi user chọn ghế và bấm tiếp tục, frontend gọi API giữ ghế. Đây là đoạn quan trọng nhất để tránh hai người giữ cùng một ghế."
    )
    add_code(
        doc,
        """// BookingServiceImpl.holdSeats()
List<SeatStatus> seatStatuses =
    seatStatusRepository.findForUpdateByShowtimeAndSeats(showtimeId, seatIds);

boolean hasUnavailable = seatStatuses.stream()
    .anyMatch(ss -> ss.getStatus() != SeatStatusType.AVAILABLE);

for (SeatStatus ss : seatStatuses) {
    ss.setStatus(SeatStatusType.HOLD);
    ss.setHoldBy(user);
    ss.setHoldUntil(holdUntil);
}

seatStatusRepository.saveAll(seatStatuses);
publishHoldAfterCommit(showtimeId, seatIds, userId, holdUntil);"""
    )
    add_table(
        doc,
        ["Bước", "Backend làm gì?", "Vì sao cần?"],
        [
            ("1", "Lấy current user từ SecurityUtils.", "Biết ai đang giữ ghế."),
            ("2", "Load showtime và kiểm tra còn được đặt vé không.", "Không cho đặt sát giờ chiếu hoặc suất không hợp lệ."),
            ("3", "findForUpdateByShowtimeAndSeats dùng PESSIMISTIC_WRITE.", "Khóa DB để hai user không giữ cùng ghế cùng lúc."),
            ("4", "Kiểm tra tất cả ghế phải AVAILABLE.", "Nếu ghế đã HOLD/BOOKED thì từ chối."),
            ("5", "Set status = HOLD, holdBy, holdUntil.", "Tạo trạng thái giữ ghế có thời hạn."),
            ("6", "Commit transaction.", "DB là nguồn sự thật."),
            ("7", "Publish HOLD sau commit.", "Các client khác thấy ghế chuyển sang đang giữ."),
        ],
        [Inches(0.7), Inches(2.7), Inches(3.1)],
    )

    doc.add_heading("6.1. Repository chống race condition", level=2)
    doc.add_paragraph(
        "SeatStatusRepository dùng pessimistic lock ở lúc giữ ghế. Đây là điểm bảo vệ chống bán trùng ghế ở backend."
    )
    add_code(
        doc,
        """@Lock(LockModeType.PESSIMISTIC_WRITE)
@Query("SELECT ss FROM SeatStatus ss JOIN FETCH ss.seat s "
     + "WHERE ss.showtime.id = :showtimeId AND s.id IN :seatIds")
List<SeatStatus> findForUpdateByShowtimeAndSeats(UUID showtimeId, List<UUID> seatIds);"""
    )
    add_callout(
        doc,
        "WebSocket không chống bán trùng ghế",
        "WebSocket chỉ giúp các màn hình cập nhật nhanh. Chống bán trùng ghế phải nằm ở database transaction và lock. Nếu hai user bấm cùng lúc, DB lock quyết định ai giữ được trước.",
        "FEF3C7",
        "92400E",
    )

    doc.add_heading("7. Vì sao publish sau commit?", level=1)
    doc.add_paragraph(
        "Trong BookingServiceImpl có hai helper publishBulkAfterCommit và publishHoldAfterCommit. Hai helper này đăng ký TransactionSynchronization.afterCommit. "
        "Điều đó nghĩa là chỉ sau khi transaction thành công, backend mới gửi event realtime."
    )
    add_code(
        doc,
        """private void publishBulkAfterCommit(UUID showtimeId, List<UUID> seatIds, SeatStatusType status) {
    if (!TransactionSynchronizationManager.isSynchronizationActive()) {
        seatStatusPublisher.publishBulk(showtimeId, seatIds, status);
        return;
    }

    TransactionSynchronizationManager.registerSynchronization(new TransactionSynchronization() {
        @Override
        public void afterCommit() {
            seatStatusPublisher.publishBulk(showtimeId, seatIds, status);
        }
    });
}"""
    )
    add_bullets(
        doc,
        [
            "Nếu DB commit thành công: client được thông báo đổi màu ghế.",
            "Nếu DB rollback vì lỗi: không có event giả gửi xuống client.",
            "Điều này giúp giao diện không bị lệch với database.",
            "Đây là cách làm đúng hơn so với publish ngay giữa transaction.",
        ],
    )

    doc.add_heading("8. Các kịch bản backend phát WebSocket", level=1)
    add_table(
        doc,
        ["Kịch bản", "Method liên quan", "DB đổi gì?", "Event gửi xuống"],
        [
            (
                "User giữ ghế",
                "BookingServiceImpl.holdSeats",
                "seat_status AVAILABLE -> HOLD, set hold_by, hold_until",
                "HOLD kèm heldByUserId và holdUntil",
            ),
            (
                "Thanh toán thành công",
                "BookingServiceImpl.handlePaymentSuccess",
                "booking SUCCESS, seat_status -> BOOKED, clear hold",
                "BOOKED cho toàn bộ ghế trong booking",
            ),
            (
                "Thanh toán thất bại",
                "BookingServiceImpl.handlePaymentFailure",
                "booking FAILED, seat_status HOLD -> AVAILABLE",
                "AVAILABLE cho các ghế được trả lại",
            ),
            (
                "User hủy booking",
                "BookingServiceImpl.cancelBooking",
                "booking CANCELLED, seat_status HOLD -> AVAILABLE",
                "AVAILABLE cho các ghế trong booking",
            ),
            (
                "Booking chờ thanh toán hết hạn",
                "BookingServiceImpl.expirePendingBooking",
                "booking EXPIRED, payment PENDING -> EXPIRED, ghế -> AVAILABLE",
                "AVAILABLE cho ghế được nhả",
            ),
            (
                "Ghế HOLD hết hạn",
                "HoldExpireScheduler.releaseExpiredHolds",
                "seat_status HOLD hết hạn -> AVAILABLE",
                "AVAILABLE theo từng showtime",
            ),
            (
                "User mở lại sơ đồ ghế",
                "BookingServiceImpl.getSeatMap",
                "releaseExpiredHoldsForShowtime dọn hold hết hạn trước khi trả snapshot",
                "AVAILABLE nếu có hold hết hạn được dọn",
            ),
        ],
        [Inches(1.35), Inches(1.7), Inches(2.1), Inches(1.35)],
    )

    doc.add_heading("8.1. Thanh toán thành công: HOLD -> BOOKED", level=2)
    add_code(
        doc,
        """// BookingServiceImpl.handlePaymentSuccess()
booking.setStatus(BookingStatus.SUCCESS);

List<UUID> seatIds = booking.getBookingDetails().stream()
    .map(bd -> bd.getSeat().getId())
    .toList();

seatStatusRepository.bulkUpdateStatusAndClearHold(
    booking.getShowtime().getId(), seatIds, SeatStatusType.BOOKED);

publishBulkAfterCommit(booking.getShowtime().getId(), seatIds, SeatStatusType.BOOKED);"""
    )
    doc.add_paragraph(
        "Sau event BOOKED, mọi client đang xem sơ đồ ghế sẽ thấy ghế chuyển sang đã đặt. User khác không còn chọn được các ghế đó."
    )

    doc.add_heading("8.2. Thanh toán thất bại/hủy/hết hạn: HOLD -> AVAILABLE", level=2)
    add_code(
        doc,
        """int releasedSeatCount = seatStatusRepository.releaseHeldSeatsForBooking(
    booking.getShowtime().getId(),
    seatIds,
    booking.getUser().getId(),
    paymentReleaseCutoff(booking),
    SeatStatusType.AVAILABLE);

if (releasedSeatCount == seatIds.size()) {
    publishBulkAfterCommit(booking.getShowtime().getId(), seatIds, SeatStatusType.AVAILABLE);
}"""
    )
    doc.add_paragraph(
        "Backend chỉ publish AVAILABLE khi thật sự nhả đủ số ghế. Cách này tránh thông báo sai trong trường hợp ghế không còn thuộc booking hoặc đã bị xử lý bởi luồng khác."
    )

    doc.add_heading("9. Scheduler nhả ghế hết hạn", level=1)
    doc.add_paragraph(
        "Không thể trông chờ người dùng tự bấm hủy. Nếu user tắt tab hoặc bỏ đi ở màn hình thanh toán, hệ thống phải có scheduler dọn các ghế HOLD quá hạn."
    )
    add_code(
        doc,
        """@Scheduled(fixedDelayString = "${booking.expired-hold-scan-delay-ms:30000}")
@Transactional
public void releaseExpiredHolds() {
    List<ExpiredSeatHoldProjection> expired =
        seatStatusRepository.findExpiredHoldRows(LocalDateTime.now(), expiredHoldScanLimit);

    Map<UUID, List<UUID>> seatIdsByShowtime = expired.stream()
        .collect(groupingBy(getShowtimeId, mapping(getSeatId, toList())));

    int releasedCount = seatStatusRepository.releaseExpiredHoldsByIds(expiredIds);

    TransactionSynchronizationManager.registerSynchronization(new TransactionSynchronization() {
        @Override
        public void afterCommit() {
            seatIdsByShowtime.forEach(seatStatusPublisher::publishAvailable);
        }
    });
}"""
    )
    add_bullets(
        doc,
        [
            "Scheduler chạy theo booking.expired-hold-scan-delay-ms, mặc định 30 giây.",
            "Chỉ lấy một batch theo expiredHoldScanLimit để tránh quét quá nặng khi dữ liệu lớn.",
            "Group theo showtimeId để publish đúng topic cho từng suất chiếu.",
            "Vẫn publish sau commit để UI không lệch DB.",
            "Ngoài scheduler, getSeatMap cũng dọn hold hết hạn của chính showtime đang mở để UI hồi phục nhanh khi user refresh hoặc mở trang.",
        ],
    )

    doc.add_heading("10. Frontend kết nối WebSocket như thế nào?", level=1)
    add_table(
        doc,
        ["File frontend", "Vai trò"],
        [
            ("src/hooks/useSeatWebSocket.ts", "Hook tạo STOMP client, kết nối /ws-native, subscribe topic theo showtimeId, parse SeatStatusEvent và gọi callback."),
            ("src/pages/user/SeatSelectionPage.tsx", "Trang chọn ghế. Load snapshot bằng HTTP, gọi hook WebSocket, nhận event và cập nhật seatMap state."),
            ("src/api/bookingApi.ts", "Gọi HTTP snapshot /api/v1/showtimes/{showtimeId}/seats và API hold/create booking."),
            ("vite.config.ts", "Proxy /ws và /ws-native sang backend, bật ws: true để dev server chuyển tiếp WebSocket."),
        ],
        [Inches(2.35), Inches(4.15)],
    )

    doc.add_heading("10.1. useSeatWebSocket.ts", level=2)
    doc.add_paragraph(
        "Hook này gom toàn bộ logic WebSocket vào một chỗ. Component dùng hook chỉ cần truyền showtimeId và callback onSeatUpdate."
    )
    add_code(
        doc,
        """function buildWsUrl(): string {
  const proto = window.location.protocol === "https:" ? "wss" : "ws";
  return `${proto}://${window.location.host}/ws-native`;
}

export function useSeatWebSocket({ showtimeId, onSeatUpdate, currentUserId }: Options) {
  const clientRef = useRef<Client | null>(null);
  const callbackRef = useRef(onSeatUpdate);
  callbackRef.current = onSeatUpdate;

  const connect = useCallback(() => {
    if (!showtimeId) return;

    const client = new Client({
      brokerURL: buildWsUrl(),
      reconnectDelay: 3000,
      onConnect: () => {
        client.subscribe(`/topic/seatmap/${showtimeId}`, (message) => {
          const event: SeatStatusEvent = JSON.parse(message.body);
          callbackRef.current(event);
        });
      },
    });

    client.activate();
    clientRef.current = client;
  }, [showtimeId]);

  useEffect(() => {
    connect();
    return () => clientRef.current?.deactivate();
  }, [connect]);
}"""
    )
    add_bullets(
        doc,
        [
            "buildWsUrl tự chọn ws:// hoặc wss:// theo trang hiện tại. Nếu chạy qua HTTPS/ngrok thì dùng wss.",
            "brokerURL trỏ tới /ws-native, khớp endpoint backend.",
            "reconnectDelay = 3000 giúp tự kết nối lại sau 3 giây nếu mạng rớt.",
            "callbackRef giữ callback mới nhất nhưng không làm hook resubscribe liên tục.",
            "cleanup deactivate khi component unmount hoặc showtimeId đổi, tránh một tab subscribe nhiều lần.",
        ],
    )

    doc.add_heading("10.2. Vite proxy cho WebSocket", level=2)
    doc.add_paragraph(
        "Khi dev bằng Vite, frontend chạy ở localhost:5173 còn backend ở localhost:8080. Proxy giúp frontend gọi cùng host /ws-native, Vite chuyển tiếp sang backend."
    )
    add_code(
        doc,
        """// vite.config.ts
server: {
  proxy: {
    "/ws": {
      target: backendTarget,
      changeOrigin: true,
      ws: true,
    },
    "/ws-native": {
      target: backendTarget,
      changeOrigin: true,
      ws: true,
    },
  },
}"""
    )
    add_bullets(
        doc,
        [
            "ws: true là phần quan trọng để Vite proxy được WebSocket upgrade request.",
            "BACKEND_PROXY_TARGET trong .env giúp đổi backend target dễ hơn khi deploy/dev.",
            "Frontend dùng window.location.host nên không hard-code localhost:8080.",
        ],
    )

    doc.add_heading("11. SeatSelectionPage nhận event và đổi UI", level=1)
    doc.add_paragraph(
        "Trang chọn ghế có hai lớp dữ liệu: snapshot HTTP và delta WebSocket. Snapshot tải toàn bộ ghế, delta cập nhật từng ghế khi có thay đổi."
    )
    add_code(
        doc,
        """const { data: fetchedSeatMap = [] } = useQuery({
  queryKey: ["seats", showtimeId],
  queryFn: () => bookingApi.getSeatMap(showtimeId!).then(r => r.data.result),
  staleTime: 0,
  refetchOnMount: "always",
  refetchOnWindowFocus: true,
  refetchInterval: 5000,
});

useEffect(() => {
  setSeatMap(fetchedSeatMap ?? []);
}, [fetchedSeatMap]);"""
    )
    doc.add_paragraph(
        "refetchInterval 5 giây là lớp dự phòng. Nếu WebSocket bị rớt vài giây, UI vẫn tự đồng bộ lại bằng API snapshot."
    )
    add_code(
        doc,
        """const handleSeatUpdate = useCallback((event: SeatStatusEvent) => {
  setSeatMap(previous =>
    previous.map(seat => {
      if (seat.seatId !== event.seatId) return seat;

      if (event.status === "HOLD" && event.heldByUserId) {
        holdersRef.current.set(event.seatId, event.heldByUserId);
      } else {
        holdersRef.current.delete(event.seatId);
      }

      return { ...seat, status: event.status };
    })
  );

  if (event.status !== "AVAILABLE") {
    setSelected(previous => previous.filter(id => id !== event.seatId));
  }
}, []);"""
    )
    add_bullets(
        doc,
        [
            "Nếu event không thuộc seatId hiện tại trong vòng lặp, giữ nguyên seat.",
            "Nếu status = HOLD, lưu heldByUserId vào holdersRef.",
            "Nếu status khác HOLD, xóa holder khỏi holdersRef.",
            "Cập nhật seat.status để UI đổi màu.",
            "Nếu ghế không còn AVAILABLE, xóa ghế khỏi selected để user không tiếp tục giữ ghế vừa bị người khác lấy.",
        ],
    )

    doc.add_heading("11.1. Phân biệt ghế mình giữ và người khác giữ", level=2)
    add_code(
        doc,
        """const getSeatStyle = (seat: SeatMapItem): string => {
  if (heldSeats.includes(seat.seatId) || selected.includes(seat.seatId)) {
    return SEAT_STYLES[seat.seatType].SELECTED;
  }
  if (seat.status === "HOLD") {
    const holderId = holdersRef.current.get(seat.seatId);
    if (holderId && user?.id && holderId === user.id) {
      return SEAT_STYLES[seat.seatType].MY_HOLD;
    }
    return SEAT_STYLES[seat.seatType].HOLD;
  }
  return SEAT_STYLES[seat.seatType][seat.status];
};"""
    )
    doc.add_paragraph(
        "Nhờ heldByUserId trong event, frontend có thể hiển thị ghế mình giữ khác với ghế người khác giữ. "
        "Điều này giúp user không bị nhầm: ghế đang HOLD bởi chính mình vẫn nằm trong luồng thanh toán của mình."
    )

    doc.add_heading("12. Flow frontend giữ ghế rồi chuyển thanh toán", level=1)
    add_numbered(
        doc,
        [
            "User mở SeatSelectionPage, frontend gọi getSeatMap để lấy snapshot.",
            "useSeatWebSocket subscribe /topic/seatmap/{showtimeId}.",
            "User chọn ghế AVAILABLE.",
            "Bấm tiếp tục, frontend gọi POST /api/v1/bookings/hold.",
            "Backend giữ ghế thành công và publish HOLD sau commit.",
            "Frontend nhận HOLD event, seatMap đổi trạng thái.",
            "holdMutation onSuccess lưu heldSeats, bật holdActive, set countdown theo holdUntil.",
            "Frontend gọi create booking rồi navigate sang /checkout/{bookingId}.",
        ],
    )
    add_code(
        doc,
        """const holdMutation = useMutation({
  mutationFn: () => axiosClient.post("/api/v1/bookings/hold", {
    showtimeId,
    seatIds: selected,
  }),
  onSuccess: response => {
    const result = response.data.result;
    setHeldSeats(result?.heldSeatIds?.map(String) ?? selected);
    setSelected([]);
    setHoldActive(true);
    setTimeLeft(secondsUntil(result?.holdUntil));
  },
});"""
    )

    doc.add_heading("13. Khi countdown hết thời gian giữ ghế", level=1)
    doc.add_paragraph(
        "Frontend có countdown riêng để user thấy thời gian còn lại. Nhưng countdown không phải nguồn quyết định cuối cùng. "
        "Backend scheduler và payment timeout mới là nguồn xử lý thật."
    )
    add_code(
        doc,
        """useEffect(() => {
  if (!holdActive) return;
  const interval = window.setInterval(() => {
    setTimeLeft(current => {
      if (current <= 1) {
        setHoldActive(false);
        setHeldSeats([]);
        queryClient.invalidateQueries({ queryKey: ["seats", showtimeId] });
        toast.warning("Hết thời gian giữ ghế. Vui lòng chọn lại.");
        return HOLD_SECONDS;
      }
      return current - 1;
    });
  }, 1000);
  return () => window.clearInterval(interval);
}, [holdActive, queryClient, showtimeId]);"""
    )
    add_bullets(
        doc,
        [
            "Frontend hết countdown thì bỏ trạng thái holdActive và refetch seat map.",
            "Backend scheduler vẫn dọn HOLD quá hạn kể cả user đã tắt tab.",
            "Nếu scheduler publish AVAILABLE, các client khác cũng thấy ghế sáng lại.",
            "Nếu client bỏ lỡ event, refetch 5 giây hoặc focus lại tab sẽ tự sửa UI.",
        ],
    )

    doc.add_heading("14. Vì sao cần cả HTTP snapshot và WebSocket delta?", level=1)
    add_table(
        doc,
        ["Cơ chế", "Vai trò", "Nếu thiếu thì sao?"],
        [
            (
                "HTTP getSeatMap",
                "Lấy toàn bộ trạng thái ghế tại thời điểm mở trang.",
                "Frontend không biết trạng thái ban đầu của hàng trăm ghế.",
            ),
            (
                "WebSocket event",
                "Cập nhật tức thời những ghế vừa đổi.",
                "User phải chờ refetch hoặc refresh mới thấy ghế bị giữ/đặt.",
            ),
            (
                "React Query refetch",
                "Tự hồi phục khi mạng rớt, tab ngủ, hoặc event bị bỏ lỡ.",
                "UI có thể lệch trạng thái nếu WebSocket tạm mất kết nối.",
            ),
        ],
        [Inches(1.65), Inches(2.35), Inches(2.5)],
    )
    add_callout(
        doc,
        "Thiết kế chuẩn product",
        "Realtime tốt không có nghĩa là bỏ HTTP. Product thực tế thường dùng snapshot + realtime delta + fallback refetch. Như vậy hệ thống vừa nhanh vừa tự phục hồi.",
        "ECFDF5",
        "047857",
    )

    doc.add_heading("15. Các trạng thái ghế và màu UI", level=1)
    add_table(
        doc,
        ["Trạng thái", "Ý nghĩa nghiệp vụ", "Cách UI xử lý"],
        [
            ("AVAILABLE", "Ghế trống, có thể chọn.", "Hiển thị màu theo loại ghế NORMAL/VIP/COUPLE và cho phép click."),
            ("HOLD", "Ghế đang được giữ tạm thời.", "Không cho click. Nếu heldByUserId là user hiện tại thì dùng màu MY_HOLD."),
            ("BOOKED", "Ghế đã thanh toán thành công.", "Không cho click, hiển thị như đã đặt."),
            ("SELECTED", "Trạng thái frontend tạm thời trước khi gọi hold.", "Chỉ tồn tại trong UI, chưa phải trạng thái DB."),
            ("MY_HOLD", "Ghế HOLD bởi chính user hiện tại.", "Dùng để user biết đây là ghế của mình đang chờ thanh toán."),
        ],
        [Inches(1.4), Inches(2.5), Inches(2.6)],
    )

    doc.add_heading("16. Tính mở rộng và giới hạn hiện tại", level=1)
    add_bullets(
        doc,
        [
            "Simple broker của Spring phù hợp đồ án, demo, hoặc hệ thống vừa và nhỏ.",
            "Nếu scale nhiều instance backend, nên chuyển sang broker ngoài như RabbitMQ hoặc Redis pub/sub để event đi qua nhiều node.",
            "Topic theo showtimeId giúp giảm nhiễu: client chỉ nhận event của suất chiếu đang xem.",
            "Payload nhỏ giúp giảm băng thông và giảm re-render.",
            "Frontend update từng seatId, không replace toàn bộ seat map mỗi event.",
            "Nên giới hạn allowed origins khi lên production.",
            "Nếu cần audit realtime, có thể log SeatStatusEvent hoặc thêm bảng seat_status_events, nhưng hiện tại chưa cần để tránh phức tạp.",
        ],
    )

    doc.add_heading("17. Checklist test WebSocket khi bảo vệ", level=1)
    doc.add_heading("17.1. Test realtime giữ ghế", level=2)
    add_numbered(
        doc,
        [
            "Mở hai trình duyệt hoặc một trình duyệt thường và một ẩn danh.",
            "Đăng nhập hai user khác nhau.",
            "Cùng mở một suất chiếu và cùng nhìn sơ đồ ghế.",
            "User A chọn ghế A1 và bấm giữ ghế.",
            "Quan sát màn hình User B: ghế A1 phải đổi sang HOLD gần như ngay lập tức.",
            "Kiểm tra User B không thể chọn ghế A1.",
        ],
    )

    doc.add_heading("17.2. Test thanh toán thành công", level=2)
    add_numbered(
        doc,
        [
            "User A giữ ghế và tạo booking.",
            "Thanh toán thành công bằng VNPay hoặc Quét QR ngân hàng.",
            "Quay lại màn hình User B đang mở seat map.",
            "Ghế của booking phải đổi từ HOLD sang BOOKED.",
            "Refresh trang để xác nhận snapshot HTTP cũng trả BOOKED.",
        ],
    )

    doc.add_heading("17.3. Test hủy/thất bại/hết hạn", level=2)
    add_numbered(
        doc,
        [
            "User A giữ ghế nhưng hủy booking hoặc để hết thời gian.",
            "Scheduler hoặc API hủy chuyển ghế về AVAILABLE.",
            "Màn hình User B phải thấy ghế chuyển lại AVAILABLE.",
            "User B chọn lại ghế đó và giữ ghế thành công.",
        ],
    )

    doc.add_heading("17.4. Test reconnect và cleanup", level=2)
    add_numbered(
        doc,
        [
            "Mở DevTools Network, lọc WS.",
            "Vào một suất chiếu, xác nhận có kết nối /ws-native.",
            "Chuyển sang suất chiếu khác, xác nhận subscription cũ được cleanup và không nhận event nhầm.",
            "Tắt mạng vài giây rồi bật lại, quan sát hook tự reconnect sau khoảng 3 giây.",
            "Nếu bỏ lỡ event, đợi refetch 5 giây hoặc focus lại tab để seat map đồng bộ.",
        ],
    )

    doc.add_heading("18. Cách trình bày ngắn gọn khi bảo vệ", level=1)
    add_callout(
        doc,
        "Câu trả lời mẫu",
        "Hệ thống của em dùng HTTP API để tải snapshot ban đầu của sơ đồ ghế, sau đó dùng WebSocket STOMP để backend đẩy các thay đổi ghế theo topic /topic/seatmap/{showtimeId}. "
        "Khi user giữ ghế, thanh toán thành công, hủy, thanh toán thất bại hoặc hết hạn giữ ghế, backend cập nhật database trong transaction trước. Sau khi commit, SeatStatusPublisher mới gửi SeatStatusEvent xuống frontend. "
        "Frontend dùng useSeatWebSocket để subscribe topic, nhận event và update đúng seatId trong state. React Query vẫn refetch định kỳ để tự phục hồi nếu WebSocket mất kết nối. Vì vậy hệ thống vừa realtime, vừa không lệch database, vừa chống bán trùng ghế bằng lock ở backend.",
        "ECFDF5",
        "047857",
    )

    doc.add_heading("19. Những lỗi thường gặp và cách hiểu", level=1)
    add_table(
        doc,
        ["Hiện tượng", "Nguyên nhân thường gặp", "Cách kiểm tra"],
        [
            (
                "Không realtime, phải refresh mới đổi ghế.",
                "WebSocket không kết nối, proxy thiếu ws: true, topic sai showtimeId, hoặc backend chưa publish.",
                "Kiểm tra DevTools Network WS, console frontend, log backend SeatStatusPublisher.",
            ),
            (
                "Một event bị nhận nhiều lần.",
                "Component subscribe trùng do cleanup không chạy hoặc showtimeId thay đổi nhưng client cũ chưa deactivate.",
                "Kiểm tra useEffect cleanup và số kết nối /ws-native trong Network.",
            ),
            (
                "Ghế đổi màu nhưng DB không đúng.",
                "Publish trước commit hoặc frontend tự đổi state quá sớm.",
                "Trong hệ thống hiện tại publish sau commit để tránh lỗi này.",
            ),
            (
                "User vẫn chọn được ghế người khác vừa giữ.",
                "Event chưa tới hoặc UI chưa xóa selected.",
                "handleSeatUpdate đã xóa selected nếu event.status không phải AVAILABLE.",
            ),
            (
                "Mất kết nối khi chạy qua ngrok/HTTPS.",
                "Sai ws/wss hoặc Vite allowed host/proxy.",
                "buildWsUrl tự chọn wss nếu trang là https; kiểm tra /ws-native qua Vite proxy.",
            ),
        ],
        [Inches(1.65), Inches(2.45), Inches(2.4)],
    )

    doc.add_heading("20. Kết luận", level=1)
    doc.add_paragraph(
        "Luồng WebSocket hiện tại đi theo hướng đúng cho hệ thống đặt vé xem phim: database transaction quyết định trạng thái thật, WebSocket chỉ phát event sau commit, frontend nhận event để cập nhật nhanh, và HTTP refetch làm lớp tự phục hồi. "
        "Thiết kế này đủ tốt cho đồ án bảo vệ và có đường nâng cấp rõ ràng nếu sau này hệ thống chạy nhiều backend instance hoặc cần broker ngoài."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CinemaBooking.vn - WebSocket Realtime Guide").italic = True

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)

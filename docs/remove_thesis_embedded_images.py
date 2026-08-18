from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


REPORT = Path(__file__).resolve().parent / "Bao_cao_khoa_luan_CinemaBooking_Final.docx"


def next_figure_number(paragraphs, start_index: int) -> str:
    for paragraph in paragraphs[start_index + 1 : start_index + 4]:
        match = re.search(r"Hình\s+(\d+\.\d+)", paragraph.text)
        if match:
            return match.group(1)
    raise RuntimeError(f"Không tìm thấy chú thích hình sau đoạn {start_index}.")


def main() -> None:
    document = Document(REPORT)
    paragraphs = document.paragraphs
    image_paragraphs = [
        (index, paragraph)
        for index, paragraph in enumerate(paragraphs)
        if paragraph._p.xpath(".//w:drawing")
    ]

    if not image_paragraphs:
        print("Không còn ảnh nhúng trong báo cáo.")
        return

    relationship_ids: set[str] = set()
    for index, paragraph in image_paragraphs:
        for blip in paragraph._p.xpath(".//a:blip"):
            relationship_id = blip.get(qn("r:embed"))
            if relationship_id:
                relationship_ids.add(relationship_id)

        figure_number = next_figure_number(paragraphs, index)
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(f"[CHÈN HÌNH {figure_number} TẠI ĐÂY]")
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 116, 139)

    for relationship_id in relationship_ids:
        if relationship_id in document.part.rels:
            document.part.drop_rel(relationship_id)

    temporary_report = REPORT.with_suffix(".without-images.docx")
    document.save(temporary_report)
    temporary_report.replace(REPORT)
    print(f"Đã gỡ {len(image_paragraphs)} ảnh và giữ lại vị trí chèn trong {REPORT.name}.")


if __name__ == "__main__":
    main()

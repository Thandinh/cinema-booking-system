from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/Luong_thanh_toan_VNPay_SePay_CinemaBooking.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="111827"):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_code(doc, code):
    p = doc.add_paragraph()
    p.style = "CodeBlock"
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Thành phần", True, "FFFFFF")
    set_cell_text(hdr[1], "Vai trò trong luồng thanh toán", True, "FFFFFF")
    shade_cell(hdr[0], "111827")
    shade_cell(hdr[1], "111827")
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, fill="FFF7ED"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(146, 64, 14)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(3)
    p2.add_run(body)
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 24, "111827"),
        ("Heading 1", 16, "0F172A"),
        ("Heading 2", 13, "1F2937"),
        ("Heading 3", 11, "374151"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    if "CodeBlock" not in styles:
        code_style = styles.add_style("CodeBlock", 1)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(8.5)
        code_style.font.color.rgb = RGBColor(31, 41, 55)
        code_style.paragraph_format.left_indent = Inches(0.18)
        code_style.paragraph_format.right_indent = Inches(0.18)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(8)


def build_doc():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Luồng thanh toán VNPay và Quét QR ngân hàng")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("CinemaBooking.vn - tài liệu giải thích backend, frontend, bảo mật và checklist test").italic = True

    add_callout(
        doc,
        "Mục tiêu tài liệu",
        "Tài liệu này giúp bạn hiểu 2 luồng thanh toán chính của hệ thống: VNPay redirect và Quét QR ngân hàng/SePay webhook. "
        "Cách giải thích đi từ tổng quan đến từng class, từng đoạn code quan trọng, rồi nối sang frontend để bạn có thể trình bày khi bảo vệ.",
        "ECFDF5",
    )

    doc.add_heading("1. Bức tranh tổng quan", level=1)
    doc.add_paragraph(
        "Trong hệ thống đặt vé xem phim, thanh toán không chỉ là tạo URL hoặc QR. "
        "Nó là một chuỗi trạng thái nối giữa booking, payment, seat_status và ticket. "
        "Nếu xử lý sai, hệ thống có thể bán trùng ghế, xác nhận nhầm vé, hoặc giữ ghế quá lâu làm giảm doanh thu."
    )
    add_bullets(
        doc,
        [
            "Booking là đơn đặt vé của người dùng. Khi mới tạo, booking thường ở trạng thái PENDING.",
            "Payment là giao dịch thanh toán gắn với booking. Một booking có thể có nhiều payment nếu người dùng tạo lại thanh toán.",
            "Seat status là trạng thái ghế theo từng suất chiếu: AVAILABLE, HOLD hoặc BOOKED.",
            "Ticket chỉ được sinh sau khi booking thanh toán thành công.",
            "Payment event là nhật ký/audit trail, giúp admin đối soát và điều tra lỗi callback/webhook.",
        ],
    )
    add_callout(
        doc,
        "Nguyên tắc product thật",
        "Nguồn quyết định cuối cùng không phải giao diện frontend. Backend phải kiểm tra user, trạng thái booking, thời hạn giữ vé, số tiền, chữ ký callback/webhook và trạng thái payment trước khi xác nhận vé.",
    )

    doc.add_heading("2. Các class chính trong backend", level=1)
    add_key_value_table(
        doc,
        [
            ("PaymentController", "Nhận request từ frontend và callback/webhook từ cổng thanh toán. Nó không xử lý nghiệp vụ sâu mà chuyển sang PaymentService."),
            ("PaymentService", "Interface định nghĩa các hành động thanh toán: initiatePayment, handleVNPayCallback, handleSePayWebhook, getAllPayments, reconciliation."),
            ("PaymentServiceImpl", "Trung tâm nghiệp vụ payment. Kiểm tra quyền, trạng thái booking, số tiền, tạo payment, xử lý callback/webhook, cập nhật payment và gọi BookingService."),
            ("PaymentGateway", "Interface chung cho các cổng thanh toán. Mỗi cổng chỉ cần biết cách tạo URL/QR thanh toán."),
            ("VnPayPaymentGateway", "Sinh URL redirect sang VNPay, ký tham số bằng HMAC SHA512, gắn thời gian hết hạn thanh toán."),
            ("SePayPaymentGateway", "Sinh QR VietQR/SePay, lưu cache QR trong providerResponse và trả payload dạng sepay://pay?... cho frontend."),
            ("PaymentRepository", "Truy vấn payment, tìm payment PENDING, khóa payment theo transactionNo để xử lý callback/webhook an toàn."),
            ("BookingServiceImpl", "Khi payment thành công/thất bại/hết hạn, service này đổi trạng thái booking, ghế và sinh vé."),
            ("PaymentEventService", "Ghi nhật ký từng bước: initiated, URL created, callback received, amount mismatch, success, failed, expired."),
            ("SePayConfig, VNPayConfig", "Đọc cấu hình từ application.yaml/.env: merchant info, secret key, QR account, webhook credentials."),
        ],
    )

    doc.add_heading("3. Tầng frontend tham gia thanh toán", level=1)
    add_key_value_table(
        doc,
        [
            ("CheckoutPage.tsx", "Trang người dùng chọn phương thức thanh toán, áp mã giảm giá, tạo VNPay URL hoặc QR ngân hàng."),
            ("paymentApi.ts", "Gọi API /api/v1/payments/initiate với bookingId, method và amount."),
            ("bookingApi.ts", "Lấy booking hiện tại, áp/gỡ mã giảm giá, kiểm tra trạng thái booking sau khi chuyển khoản."),
            ("PaymentResultPage.tsx", "Hiển thị kết quả sau callback VNPay hoặc sau khi SePay tự xác nhận thành công."),
            ("React Query", "Cache booking và tự refetch khi chờ SePay webhook. Giúp giao diện cập nhật gần realtime."),
            ("toastBus", "Hiển thị lỗi/ngữ cảnh cho người dùng: hết hạn giữ vé, chưa ghi nhận thanh toán, đã xác nhận thanh toán."),
        ],
    )

    doc.add_heading("4. Luồng khởi tạo thanh toán chung", level=1)
    doc.add_paragraph(
        "Cả VNPay và Quét QR ngân hàng đều đi qua cùng một cửa: PaymentController.initiatePayment(). "
        "Frontend gửi bookingId, method và amount. Backend không tin tuyệt đối số tiền frontend gửi lên, mà so sánh với tổng tiền trong booking."
    )
    add_numbered(
        doc,
        [
            "Frontend gọi paymentApi.initiatePayment(bookingId, method, booking.totalPrice).",
            "PaymentController nhận request POST /api/v1/payments/initiate.",
            "PaymentServiceImpl.initiatePayment lấy user hiện tại từ JWT bằng SecurityUtils.getCurrentUserId().",
            "Service tìm booking trong DB và kiểm tra booking thuộc về user đang đăng nhập.",
            "Nếu booking không còn PENDING hoặc đã hết thời gian giữ vé, backend từ chối thanh toán.",
            "Backend so sánh amount frontend gửi lên với booking.totalPrice trong DB.",
            "Nếu đã có payment PENDING cùng method và cùng số tiền, backend reuse payment đó.",
            "Nếu payment PENDING cũ khác số tiền, backend đánh dấu EXPIRED và tạo payment mới.",
            "Backend gọi đúng PaymentGateway theo method để tạo URL/QR.",
            "Backend ghi PaymentEvent để admin có nhật ký đối soát.",
        ],
    )
    add_code(
        doc,
        """// PaymentServiceImpl.initiatePayment - ý chính
if (!booking.getUser().getId().equals(userId)) throw UNAUTHORIZED;
if (booking.getStatus() != PENDING) throw BOOKING_ALREADY_PROCESSED;
if (isPaymentWindowExpired(booking)) expirePendingBooking(...);
if (amount == null || amount.compareTo(booking.getTotalPrice()) != 0) throw PAYMENT_AMOUNT_MISMATCH;

// Chỉ reuse payment nếu số tiền còn khớp.
if (pendingPayment.isPresent() && isSameAmount(payment.getAmount(), booking.getTotalPrice())) {
    return createPaymentUrlWithAudit(payment, booking, request);
}""",
    )

    doc.add_heading("5. Vì sao phải kiểm tra amount ở backend?", level=2)
    doc.add_paragraph(
        "Frontend có thể bị lỗi, bị sửa request bằng DevTools, hoặc người dùng áp mã giảm giá sau khi đã tạo QR. "
        "Nếu backend chỉ tin số tiền frontend gửi lên thì có thể xác nhận vé với số tiền sai. Vì vậy backend luôn lấy booking.totalPrice trong DB làm chuẩn."
    )
    add_callout(
        doc,
        "Case vừa fix",
        "Nếu người dùng tạo QR trước, sau đó áp mã giảm giá, payment PENDING cũ sẽ bị EXPIRED. QR cũ không được reuse. Người dùng phải tạo QR mới đúng số tiền.",
        "FEF3C7",
    )

    doc.add_heading("6. Luồng thanh toán VNPay", level=1)
    doc.add_heading("6.1. Tạo URL thanh toán", level=2)
    doc.add_paragraph(
        "VNPay là luồng redirect. Người dùng rời website của mình, sang trang VNPay sandbox/production, thanh toán xong VNPay redirect về callback backend."
    )
    add_numbered(
        doc,
        [
            "CheckoutPage chọn method VNPAY.",
            "Người dùng bấm Thanh toán VNPay.",
            "Frontend gọi /api/v1/payments/initiate?bookingId=...&method=VNPAY&amount=...",
            "PaymentServiceImpl tạo hoặc reuse Payment PENDING.",
            "PaymentServiceImpl gọi VnPayPaymentGateway.createPaymentUrl().",
            "VnPayPaymentGateway tạo bộ tham số vnp_*: mã giao dịch, số tiền, order info, IP, create date, expire date.",
            "Gateway ký hash bằng VNPay hash secret.",
            "Backend trả URL VNPay về frontend.",
            "Frontend dùng window.location.assign(paymentUrl) để chuyển người dùng sang VNPay.",
        ],
    )
    add_code(
        doc,
        """// CheckoutPage.tsx - xử lý kết quả initiate payment
if (paymentUrl && /^https?:\\/\\//i.test(paymentUrl)) {
    window.location.assign(paymentUrl);
    return;
}""",
    )
    add_code(
        doc,
        """// VnPayPaymentGateway - ý chính
vnpParamsMap.put("vnp_TxnRef", payment.getTransactionNo());
vnpParamsMap.put("vnp_OrderInfo", "Thanh toan ve xem phim|" + booking.getSecureToken());
vnpParamsMap.put("vnp_Amount", String.valueOf(payment.getAmount().multiply(100).longValue()));
vnpParamsMap.put("vnp_CreateDate", format(now Asia/Ho_Chi_Minh));
vnpParamsMap.put("vnp_ExpireDate", format(booking.paymentExpiresAt));
String secureHash = hmacSHA512(hashSecret, hashData);
return vnpayUrl + "?" + query + "&vnp_SecureHash=" + secureHash;""",
    )

    doc.add_heading("6.2. Callback từ VNPay về backend", level=2)
    doc.add_paragraph(
        "Callback là điểm quan trọng nhất. Backend phải xác thực chữ ký, tìm payment bằng transactionNo, kiểm tra số tiền, kiểm tra hạn giữ vé, rồi mới cho booking thành công."
    )
    add_numbered(
        doc,
        [
            "VNPay redirect về GET /api/v1/payments/vnpay-callback với các tham số vnp_*.",
            "PaymentController gọi paymentService.handleVNPayCallback(request).",
            "PaymentServiceImpl gom toàn bộ tham số vào callbackPayload để lưu audit.",
            "Service loại vnp_SecureHash khỏi dữ liệu ký, sort tham số và tính lại HMAC SHA512.",
            "Nếu chữ ký không khớp, ghi VNPAY_CALLBACK_INVALID_SIGNATURE và redirect về payment result FAILED.",
            "Nếu chữ ký hợp lệ, lấy vnp_TxnRef để tìm payment trong DB bằng transactionNo.",
            "Service lấy booking từ payment và dùng secureToken của booking trong DB, không tin orderInfo từ callback.",
            "Nếu payment đã xử lý, trả kết quả hiện tại để tránh xử lý trùng.",
            "Nếu booking đã xử lý, đồng bộ payment theo booking.",
            "Kiểm tra vnp_Amount có đúng payment.amount * 100 không.",
            "Nếu quá hạn giữ vé, payment EXPIRED và booking EXPIRED.",
            "Nếu vnp_ResponseCode = 00, payment SUCCESS và gọi bookingService.handlePaymentSuccess().",
            "Nếu responseCode khác 00, payment FAILED và gọi bookingService.handlePaymentFailure().",
        ],
    )
    add_code(
        doc,
        """// PaymentServiceImpl.handleVNPayCallback - phần kiểm tra chữ ký
String signValue = VNPayUtil.hmacSHA512(hashSecret, hashData.toString());
if (!signValue.equalsIgnoreCase(vnp_SecureHash)) {
    recordDetached(VNPAY_CALLBACK_INVALID_SIGNATURE);
    return "redirect:/payment/result?status=FAILED&reason=invalid-signature";
}""",
    )
    add_code(
        doc,
        """// Không dùng token từ vnp_OrderInfo để điều khiển booking
Payment payment = paymentRepository.findLockedByTransactionNo(txnRef).orElseThrow(...);
Booking booking = payment.getBooking();
String secureToken = booking.getSecureToken();""",
    )
    add_callout(
        doc,
        "Tại sao không dùng secureToken trong orderInfo?",
        "Callback là dữ liệu đi từ cổng thanh toán quay về. Dù đã có chữ ký, cách an toàn hơn là dùng vnp_TxnRef để tìm payment nội bộ, rồi lấy booking/secureToken trong DB. Như vậy callback không thể điều khiển nhầm booking khác.",
    )

    doc.add_heading("6.3. Sau khi VNPay thành công", level=2)
    add_bullets(
        doc,
        [
            "payment.status chuyển từ PENDING sang SUCCESS.",
            "payment.paymentTime được set bằng thời gian hiện tại.",
            "BookingServiceImpl.handlePaymentSuccess(secureToken) đổi booking sang SUCCESS.",
            "Ghế trong booking chuyển sang BOOKED.",
            "Ticket được sinh cho từng booking_detail.",
            "Email vé được gửi cho khách nếu email service được cấu hình.",
            "Frontend được redirect về /payment/result?status=SUCCESS&bookingId=...",
        ],
    )

    doc.add_heading("7. Luồng Quét QR ngân hàng / SePay", level=1)
    doc.add_heading("7.1. Tạo QR thanh toán", level=2)
    doc.add_paragraph(
        "Luồng QR ngân hàng khác VNPay ở chỗ người dùng không rời sang cổng thanh toán. Backend trả về thông tin QR, frontend hiển thị QR để người dùng chuyển khoản bằng app ngân hàng hoặc ví điện tử hỗ trợ quét QR."
    )
    add_numbered(
        doc,
        [
            "CheckoutPage mặc định chọn Quét QR ngân hàng.",
            "Người dùng bấm Thanh toán Quét QR ngân hàng.",
            "Backend tạo Payment PENDING với method SEPAY.",
            "SePayPaymentGateway kiểm tra SePayConfig đã sẵn sàng: enabled, bankCode, accountNumber.",
            "Gateway tạo transferCode từ transactionNo, ví dụ CBK1234567890.",
            "Gateway tạo transferContent = transferCode + ' thanh toan ve'.",
            "Gateway tạo qrUrl theo VietQR: bank, account, amount, description.",
            "Gateway lưu payload QR vào payment.providerResponse.sepayQr.",
            "Backend trả chuỗi sepay://pay?... cho frontend.",
            "Frontend parse payload và render SePayQrPanel.",
        ],
    )
    add_code(
        doc,
        """// SePayPaymentGateway - tạo payload QR
long amount = toVndAmount(payment.getAmount());
String transferCode = payment.getTransactionNo();
String transferContent = transferCode + " thanh toan ve";
String qrUrl = buildQrUrl(amount, transferContent);
providerResponse.put("sepayQr", payload);
return "sepay://pay?qrUrl=...&amount=...&transferContent=...";""",
    )
    add_code(
        doc,
        """// CheckoutPage.tsx - nhận payload QR
if (paymentUrl && paymentUrl.startsWith("sepay://pay?")) {
    setSePayQr(parseSePayPayload(paymentUrl));
    queryClient.invalidateQueries({ queryKey: ["booking", bookingId] });
    return;
}""",
    )

    doc.add_heading("7.2. Vì sao QR phải cố định số tiền?", level=2)
    doc.add_paragraph(
        "QR chuyển khoản là hướng dẫn thanh toán cụ thể: chuyển bao nhiêu tiền, vào tài khoản nào, nội dung gì. "
        "Nếu QR đã hiện ra mà người dùng áp mã giảm giá sau đó, QR cũ phải bị bỏ. Nếu không, người dùng có thể chuyển sai số tiền và webhook sẽ báo amount mismatch."
    )
    add_bullets(
        doc,
        [
            "Frontend khóa form mã giảm giá khi QR đã được tạo.",
            "Nếu muốn đổi mã giảm giá, người dùng bấm Đổi mã giảm giá.",
            "Frontend ẩn QR hiện tại bằng setSePayQr(null).",
            "Sau khi áp/gỡ mã, người dùng tạo QR mới đúng số tiền.",
            "Backend cũng bảo vệ bằng cách đánh dấu payment PENDING cũ là EXPIRED nếu booking amount đổi.",
        ],
    )

    doc.add_heading("7.3. Webhook SePay xác nhận tự động", level=2)
    doc.add_paragraph(
        "Sau khi tiền vào tài khoản, SePay gửi webhook server-to-server về backend. Đây là nguồn xác nhận đáng tin hơn việc người dùng bấm 'Tôi đã chuyển khoản'."
    )
    add_numbered(
        doc,
        [
            "SePay gọi POST /api/v1/payments/sepay-webhook.",
            "PaymentController chuyển rawPayload và request sang paymentService.handleSePayWebhook().",
            "Service parse JSON payload.",
            "Service tìm mã giao dịch trong code/content/description theo prefix CBK.",
            "Service xác thực webhook bằng HMAC hoặc API key.",
            "Nếu không có credential khi SePay đang bật thật, webhook bị từ chối.",
            "Service khóa payment bằng transactionNo để tránh webhook xử lý trùng song song.",
            "Service lưu payload webhook vào providerResponse để audit.",
            "Nếu transferType không phải in, bỏ qua.",
            "Nếu transferAmount khác payment.amount, ghi SEPAY_AMOUNT_MISMATCH.",
            "Nếu payment không còn PENDING, trả Already processed để idempotent.",
            "Nếu booking không còn PENDING, đồng bộ payment theo booking.",
            "Nếu booking hết hạn giữ vé, payment EXPIRED và booking EXPIRED.",
            "Nếu hợp lệ, payment SUCCESS và gọi bookingService.handlePaymentSuccess().",
        ],
    )
    add_code(
        doc,
        """// PaymentServiceImpl.handleSePayWebhook - ý chính
if (!isValidSePayWebhookSignature(rawPayload, request)) return failed;
Payment payment = paymentRepository.findLockedByTransactionNo(transactionNo).orElseThrow(...);
if (transferAmount != toVndAmount(payment.getAmount())) return amountMismatch;
if (payment.getStatus() != PENDING) return alreadyProcessed;
if (isPaymentWindowExpired(booking)) expirePendingBooking(...);
payment.setStatus(SUCCESS);
bookingService.handlePaymentSuccess(booking.getSecureToken());""",
    )

    doc.add_heading("7.4. Nút 'Tôi đã chuyển khoản' trong frontend", level=2)
    doc.add_paragraph(
        "Nút này không tự xác nhận vé. Nó chỉ gọi lại API lấy booking mới nhất để kiểm tra webhook đã tới chưa. Nếu webhook chưa tới, frontend báo người dùng chờ thêm vài giây."
    )
    add_code(
        doc,
        """// CheckoutPage.tsx - kiểm tra thủ công
bookingApi.getBookingById(bookingId).then(latestBooking => {
    if (latestBooking.status === "SUCCESS") navigate("/payment/result?status=SUCCESS");
    else toast.info("Chưa ghi nhận thanh toán. Vui lòng kiểm tra lại sau vài giây.");
});""",
    )

    doc.add_heading("8. Frontend checkout hoạt động như thế nào?", level=1)
    add_key_value_table(
        doc,
        [
            ("useQuery booking", "Lấy booking hiện tại theo bookingId. Khi có QR SePay và booking PENDING, tự refetch mỗi 3 giây."),
            ("remainingSeconds", "Đếm ngược thời gian giữ vé từ booking.paymentExpiresAt. Hết giờ thì điều hướng về chọn ghế."),
            ("method", "Lưu phương thức thanh toán hiện tại: SEPAY hoặc VNPAY."),
            ("sePayQr", "Nếu khác null, frontend hiển thị panel QR và khóa form mã giảm giá."),
            ("paymentMutation", "Gọi initiatePayment. Nếu response là sepay://pay thì parse QR; nếu là URL http thì redirect sang VNPay."),
            ("applyPromotionMutation", "Áp mã giảm giá, cập nhật booking cache, ẩn QR cũ nếu có."),
            ("removePromotionMutation", "Gỡ mã giảm giá, cập nhật booking cache, ẩn QR cũ nếu có."),
            ("refreshBookingMutation", "Nút Tôi đã chuyển khoản dùng mutation này để kiểm tra booking mới nhất."),
            ("PaymentResultPage", "Chuẩn hóa status từ query string, hiển thị thành công/thất bại/hết hạn và link xem vé/chọn lại ghế."),
        ],
    )
    add_callout(
        doc,
        "UX quan trọng",
        "Nếu người dùng đổi từ Quét QR ngân hàng sang VNPay, frontend ẩn QR cũ. Điều này tránh nhầm lẫn khi cùng lúc nhìn thấy QR cũ và nút thanh toán VNPay.",
        "EFF6FF",
    )

    doc.add_heading("9. BookingService liên quan gì đến payment?", level=1)
    doc.add_paragraph(
        "PaymentService không trực tiếp đổi ghế hoặc sinh vé. Nó chỉ xác định payment thành công/thất bại/hết hạn, sau đó gọi BookingService để cập nhật nghiệp vụ đặt vé."
    )
    add_bullets(
        doc,
        [
            "handlePaymentSuccess: booking SUCCESS, ghế BOOKED, sinh ticket, gửi email vé, publish realtime seat map.",
            "handlePaymentFailure: booking FAILED, ghế trả về AVAILABLE, payment ghi FAILED.",
            "expirePendingBooking: booking EXPIRED, ghế đang HOLD được nhả lại.",
            "applyPromotion/removePromotion: cập nhật tổng tiền và expire payment PENDING cũ để QR/URL không sai amount.",
        ],
    )

    doc.add_heading("10. Bảo mật và tính đúng đắn", level=1)
    add_key_value_table(
        doc,
        [
            ("JWT + PAYMENT_CREATE", "Chỉ user có quyền thanh toán mới gọi được initiatePayment."),
            ("Kiểm tra chủ booking", "User chỉ thanh toán booking của chính mình."),
            ("Kiểm tra booking PENDING", "Không thanh toán lại đơn đã SUCCESS/FAILED/CANCELLED/EXPIRED."),
            ("Kiểm tra paymentExpiresAt", "Không xác nhận thanh toán sau khi hết thời gian giữ vé."),
            ("Kiểm tra amount", "Frontend gửi amount nhưng backend đối chiếu với booking.totalPrice."),
            ("Pessimistic lock", "findLockedByTransactionNo khóa payment khi callback/webhook đến, tránh xử lý song song."),
            ("Idempotency", "Callback/webhook lặp lại không làm sinh vé nhiều lần, chỉ trả trạng thái đã xử lý."),
            ("VNPay signature", "Callback VNPay phải khớp HMAC SHA512."),
            ("SePay HMAC/API key", "Webhook SePay phải có chữ ký hoặc API key khi hệ thống bật SePay thật."),
            ("PaymentEvent audit", "Mọi bước quan trọng được ghi log để admin đối soát."),
        ],
    )

    doc.add_heading("11. Các tình huống lỗi thường gặp", level=1)
    add_key_value_table(
        doc,
        [
            ("VNPay báo Không tìm thấy website/giao dịch", "Thường do sai TMN code/hash secret, sai return URL, ngrok tắt, hoặc transaction hết hạn ở sandbox."),
            ("Payment amount mismatch", "Số tiền callback/webhook khác payment.amount. Có thể do QR cũ, mã giảm giá đổi sau khi tạo payment, hoặc request bị sửa."),
            ("Webhook SePay invalid authentication", "Sai HMAC secret/API key hoặc chưa cấu hình credential khi SEPAY_ENABLED=true."),
            ("Chưa ghi nhận thanh toán", "Webhook chưa tới, nội dung chuyển khoản sai, số tiền sai, hoặc SePay chưa nhận giao dịch từ ngân hàng."),
            ("Booking expired", "Webhook/callback đến sau thời gian giữ vé. Hệ thống không xác nhận vé để tránh giữ ghế quá lâu."),
            ("Payment already processed", "Cổng thanh toán gửi callback/webhook lại. Đây là bình thường; hệ thống xử lý idempotent."),
        ],
    )

    doc.add_heading("12. Checklist test khi bảo vệ", level=1)
    doc.add_heading("12.1. Test VNPay", level=2)
    add_numbered(
        doc,
        [
            "Đăng nhập user.",
            "Chọn phim, suất chiếu, ghế.",
            "Tạo booking PENDING.",
            "Vào checkout, chọn VNPay.",
            "Bấm Thanh toán VNPay và kiểm tra redirect sang sandbox VNPay.",
            "Thanh toán thành công.",
            "Kiểm tra redirect về PaymentResultPage SUCCESS.",
            "Kiểm tra Vé của tôi có vé mới và QR check-in.",
            "Kiểm tra ghế đã chuyển BOOKED.",
            "Trong admin Payment, kiểm tra payment SUCCESS và có payment event callback received/success.",
        ],
    )
    doc.add_heading("12.2. Test Quét QR ngân hàng/SePay", level=2)
    add_numbered(
        doc,
        [
            "Đăng nhập user và tạo booking PENDING.",
            "Ở checkout, chọn Quét QR ngân hàng.",
            "Bấm Thanh toán Quét QR ngân hàng.",
            "Kiểm tra QR hiện đúng số tiền, ngân hàng, số tài khoản, nội dung chuyển khoản.",
            "Chuyển khoản đúng số tiền và đúng nội dung.",
            "Đợi webhook SePay hoặc bấm Tôi đã chuyển khoản.",
            "Kiểm tra hệ thống tự chuyển sang PaymentResultPage SUCCESS.",
            "Kiểm tra payment SUCCESS, booking SUCCESS, ghế BOOKED, ticket đã sinh.",
            "Kiểm tra Payment Events có SEPAY_WEBHOOK_RECEIVED và PAYMENT_SUCCESS.",
        ],
    )
    doc.add_heading("12.3. Test lỗi mã giảm giá với QR", level=2)
    add_numbered(
        doc,
        [
            "Tạo QR trước.",
            "Quan sát form mã giảm giá bị khóa và có hướng dẫn đổi mã.",
            "Bấm Đổi mã giảm giá.",
            "Áp mã giảm giá.",
            "Tạo QR mới.",
            "Kiểm tra số tiền QR mới bằng tổng tiền sau giảm.",
            "Kiểm tra payment PENDING cũ trong DB/admin bị EXPIRED.",
        ],
    )

    doc.add_heading("13. Cách trình bày ngắn gọn khi bảo vệ", level=1)
    add_callout(
        doc,
        "Câu trả lời mẫu",
        "Hệ thống của em tách phần tạo thanh toán thành PaymentGateway để dễ mở rộng cổng mới. VNPay là luồng redirect có callback ký HMAC. SePay là luồng QR ngân hàng, xác nhận bằng webhook server-to-server. Backend luôn kiểm tra quyền user, trạng thái booking, thời gian giữ vé, số tiền, chữ ký callback/webhook và idempotency trước khi chuyển booking sang SUCCESS. Khi thành công, BookingService mới đổi ghế sang BOOKED và sinh ticket QR.",
        "ECFDF5",
    )

    doc.add_heading("14. Ghi chú mở rộng sau này", level=1)
    add_bullets(
        doc,
        [
            "Có thể thêm cổng mới bằng cách tạo class implements PaymentGateway, thêm enum PaymentMethod và config tương ứng.",
            "Nên có job reconciliation định kỳ so sánh payment/booking/ticket để phát hiện lệch trạng thái.",
            "Nên giữ PaymentEvent lâu dài để phục vụ đối soát và hỗ trợ khách hàng.",
            "Khi triển khai thật, bắt buộc dùng HTTPS public URL cho callback/webhook.",
            "Không commit secret key, webhook secret, API key lên GitHub.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CinemaBooking.vn - Payment Flow Guide").italic = True

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)

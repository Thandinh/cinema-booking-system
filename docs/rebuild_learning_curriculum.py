"""Rebuild the final CinemaBooking.vn Word learning curriculum.

The source guides in docs/ keep their original pedagogical formatting. This
builder copies each focused guide and appends a current-source learning section
with an end-to-end flow, class map, code patterns, failure cases, defense
questions, and a hands-on checklist.
"""

from __future__ import annotations

import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from enhance_flow_guides import (
    add_bullets,
    add_callout,
    add_code,
    add_numbered,
    add_table,
    setup_learning_styles,
)
from generate_advanced_learning_guides import add_title, apply_styles


DOCS_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = DOCS_DIR / "learning-rebuild"
CURRENT_UPDATE_MARKER = "PHẦN CẬP NHẬT THEO SOURCE HIỆN TẠI"
DEFENSE_FRAME_MARKER = "KHUNG 7 PHẦN: HỌC SÂU - BẢO VỆ - PHỎNG VẤN"


def ensure_compatibility_styles(document: Document) -> None:
    """Supply the few styles helper functions expect in legacy/localized DOCX files."""
    if "Table Grid" not in document.styles:
        document.styles.add_style("Table Grid", WD_STYLE_TYPE.TABLE)


def add_heading(document: Document, text: str, level: int = 1):
    """Add a heading even when an older DOCX lacks Word's built-in heading styles."""
    style_name = f"Heading {level}"
    if style_name in document.styles:
        return document.add_heading(text, level=level)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    paragraph_pr = paragraph._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(max(0, level - 1)))
    paragraph_pr.append(outline)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16 if level == 1 else 12)
    run.font.color.rgb = RGBColor.from_string("1F4E78" if level == 1 else "243B53")
    return paragraph


def add_bullets(document: Document, items):
    """Use native list styles when available and a stable visual fallback otherwise."""
    has_style = "List Bullet" in document.styles
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet" if has_style else None)
        if not has_style:
            paragraph.paragraph_format.left_indent = Inches(0.24)
            paragraph.paragraph_format.first_line_indent = Inches(-0.16)
            paragraph.add_run("• ")
        paragraph.add_run(item)


def add_numbered(document: Document, items):
    """Keep numbered learning steps readable in localized/custom Word templates."""
    has_style = "List Number" in document.styles
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph(style="List Number" if has_style else None)
        if not has_style:
            paragraph.paragraph_format.left_indent = Inches(0.24)
            paragraph.paragraph_format.first_line_indent = Inches(-0.16)
            paragraph.add_run(f"{index}. ").bold = True
        paragraph.add_run(item)


@dataclass(frozen=True)
class Topic:
    source: str
    output: str
    title: str
    purpose: str
    quick_steps: tuple[str, ...]
    core_idea: str
    flow: str
    classes: tuple[tuple[str, str, str], ...]
    code_sections: tuple[tuple[str, str, str], ...]
    failures: tuple[tuple[str, str, str], ...]
    questions: tuple[tuple[str, str], ...]
    checklist: tuple[str, ...]


@dataclass(frozen=True)
class DefenseLab:
    tldr: str
    sequence_steps: tuple[str, ...]
    request_payload: str
    database_state: str
    ide_order: tuple[str, ...]
    arrow_flow: str
    comparisons: tuple[tuple[str, str, str, str], ...]
    test_steps: tuple[str, ...]
    edge_cases: tuple[str, ...]
    senior_findings: tuple[tuple[str, str, str], ...]
    scale_summary: str
    hard_questions: tuple[tuple[str, str], ...]


def add_update_marker(document: Document, title: str, purpose: str) -> None:
    document.add_page_break()
    marker = document.add_paragraph()
    marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    marker_run = marker.add_run(CURRENT_UPDATE_MARKER)
    marker_run.bold = True
    marker_run.font.name = "Arial"
    marker_run.font.size = Pt(18)
    marker_run.font.color.rgb = RGBColor.from_string("0F172A")

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(8)
    title_run = heading.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(15)
    title_run.font.color.rgb = RGBColor.from_string("1D4ED8")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(purpose)
    subtitle_run.italic = True
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor.from_string("475569")


def append_topic_update(document: Document, topic: Topic) -> None:
    setup_learning_styles(document)
    add_update_marker(document, topic.title, topic.purpose)

    add_heading(document, "A. Cách học chuyên đề này", level=1)
    add_numbered(document, topic.quick_steps)
    add_callout(
        document,
        "Ý chính phải nhớ",
        topic.core_idea,
        "ECFDF5",
        "047857",
    )

    add_heading(document, "B. Sơ đồ luồng end-to-end", level=1)
    add_code(document, topic.flow)

    add_heading(document, "C. Bản đồ class và file cần mở trong IDE", level=1)
    add_table(
        document,
        ["Class/file", "Trách nhiệm", "Điểm cần quan sát"],
        topic.classes,
    )

    add_heading(document, "D. Những đoạn code/pattern cần hiểu", level=1)
    for heading, code_text, explanation in topic.code_sections:
        add_heading(document, heading, level=2)
        add_code(document, code_text)
        document.add_paragraph(explanation)

    add_heading(document, "E. Tình huống lỗi và cách suy luận", level=1)
    add_table(
        document,
        ["Hiện tượng", "Nguyên nhân thường gặp", "Cách kiểm tra/xử lý"],
        topic.failures,
    )

    add_heading(document, "F. Câu hỏi bảo vệ và trả lời trọng tâm", level=1)
    add_table(document, ["Câu hỏi", "Trả lời trọng tâm"], topic.questions)

    add_heading(document, "G. Checklist thực hành", level=1)
    add_bullets(document, topic.checklist)

    deep_sections = DEEP_DIVES.get(topic.output, ())
    if deep_sections:
        add_heading(document, "H. Đọc code chuyên sâu và liên hệ production", level=1)
        for heading, paragraphs in deep_sections:
            add_heading(document, heading, level=2)
            for paragraph_text in paragraphs:
                document.add_paragraph(paragraph_text)

    defense_lab = DEFENSE_LABS.get(topic.output)
    if defense_lab:
        append_defense_lab(document, defense_lab)


def append_defense_lab(document: Document, lab: DefenseLab) -> None:
    document.add_page_break()
    marker = document.add_paragraph()
    marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    marker_run = marker.add_run(DEFENSE_FRAME_MARKER)
    marker_run.bold = True
    marker_run.font.name = "Arial"
    marker_run.font.size = Pt(17)
    marker_run.font.color.rgb = RGBColor.from_string("9A3412")

    add_heading(document, "1. Bức tranh tổng quan (TL;DR)", level=1)
    add_callout(document, "Nói trong 20 giây", lab.tldr, "FFF7ED", "C2410C")

    add_heading(document, "2. Phân tích luồng chạy chi tiết (Sequence Flow)", level=1)
    add_numbered(document, lab.sequence_steps)
    add_callout(document, "Payload/request quan trọng", lab.request_payload, "EFF6FF", "1D4ED8")
    add_callout(document, "Database thay đổi thế nào?", lab.database_state, "ECFDF5", "047857")

    add_heading(document, "3. Bí kíp học nhanh và dò code trong 15 phút", level=1)
    add_numbered(document, lab.ide_order)
    add_code(document, lab.arrow_flow)

    add_heading(document, "4. Bảng so sánh và lựa chọn công nghệ", level=1)
    add_table(
        document,
        ["Quyết định", "Giải pháp hiện tại", "Phương án thay thế", "Vì sao chọn/đánh đổi"],
        lab.comparisons,
    )

    add_heading(document, "5. Kịch bản test thực tế", level=1)
    add_numbered(document, lab.test_steps)
    add_heading(document, "Edge cases phải thử", level=2)
    add_bullets(document, lab.edge_cases)

    add_heading(document, "6. Soi code smell và đề xuất tối ưu ở mức senior", level=1)
    add_table(
        document,
        ["Điểm cần soi", "Rủi ro", "Cách đo và tối ưu"],
        lab.senior_findings,
    )
    add_callout(document, "Nếu có 10.000 người dùng đồng thời", lab.scale_summary, "FEF2F2", "B91C1C")

    add_heading(document, "7. Top 3 câu hỏi xoáy từ hội đồng/interviewer", level=1)
    add_table(document, ["Câu hỏi", "Cách trả lời ghi điểm"], lab.hard_questions)


TOPICS = (
    Topic(
        source="Luong_Auth_JWT_Session_GoogleLogin_CinemaBooking.docx",
        output="02_Auth_JWT_Session_GoogleLogin_CinemaBooking.docx",
        title="Auth, JWT, refresh token, session và Google Login",
        purpose="Đọc từ màn hình đăng nhập đến SecurityContext, refresh rotation, logout và audit bảo mật.",
        quick_steps=(
            "Đọc AuthenticationController để biết contract HTTP của login, refresh, logout và session.",
            "Đặt breakpoint trong AuthenticationService để thấy password login và Google login hội tụ về cùng bước phát token.",
            "Mở JwtProperties và CustomJwtDecoder để hiểu issuer, key, thời hạn và cách Spring kiểm tra access token.",
            "Mở axiosClient.ts để quan sát một request 401 chỉ tạo đúng một refresh request rồi retry các request đang chờ.",
            "Cuối cùng kiểm tra RefreshToken, InvalidatedToken và AuthAuditLog trong database.",
        ),
        core_idea=(
            "Access token dùng ngắn hạn để xác thực từng API; refresh token đại diện cho phiên đăng nhập và được hash trong DB. "
            "Refresh rotation cấp token mới rồi thu hồi token cũ, còn logout phải vô hiệu cả phiên chứ không chỉ xóa state trên React."
        ),
        flow="""LoginPage
  -> authApi.login(username, password)
  -> POST /auth/token
  -> AuthenticationController
  -> AuthenticationService xác minh user + BCrypt
  -> tạo access JWT
  -> tạo refresh token ngẫu nhiên, lưu hash trong refresh_tokens
  -> trả AuthenticationResponse
  -> authStore giữ user/access token

Access token hết hạn
  -> axiosClient nhận 401
  -> dùng một refreshPromise dùng chung
  -> POST /auth/refresh
  -> kiểm tra hash + expiry + revoked
  -> rotate refresh token
  -> retry request cũ với access token mới""",
        classes=(
            ("AuthenticationController", "Nhận login/Google/refresh/logout/session request.", "Mapping, @Valid và response contract."),
            ("AuthenticationService", "Business logic xác thực và quản lý phiên.", "Password check, token issue, rotation, revoke, audit."),
            ("JwtProperties", "Bind cấu hình JWT từ application.yaml/.env.", "Issuer, access TTL, refresh TTL và secret."),
            ("CustomJwtDecoder", "Xác minh chữ ký và điều kiện token.", "Decoder được Spring Resource Server gọi trước controller."),
            ("RefreshTokenRepository", "Lưu/tìm/khóa refresh session.", "Không lưu raw refresh token."),
            ("AuthRateLimitService", "Giới hạn spam login/refresh.", "Key theo IP/user và cửa sổ thời gian."),
            ("AuthAuditService", "Ghi dấu vết login, refresh, revoke.", "Không ghi password/token thô."),
            ("axiosClient.ts", "Gắn Bearer token và tự refresh.", "Single-flight refresh và retry đúng một lần."),
            ("authStore.ts", "State đăng nhập phía React.", "User, permission helper và clear session."),
        ),
        code_sections=(
            (
                "D.1. Password không bao giờ so sánh dạng plain text",
                """if (!passwordEncoder.matches(request.password(), user.getPassword())) {
    auditFailure(...);
    throw new AuthenticationException("INVALID_CREDENTIALS");
}""",
                "BCrypt tự lấy salt/cost từ chuỗi hash. Backend trả thông báo chung để không tiết lộ username hay email nào tồn tại.",
            ),
            (
                "D.2. Refresh rotation",
                """lock refresh session
verify token hash + expiry + revokedAt
mark old token revoked/rotated
create new refresh token hash
issue new access token
commit transaction""",
                "Khóa và transaction ngăn hai request refresh đồng thời cùng sử dụng một token cũ để tạo hai phiên hợp lệ.",
            ),
            (
                "D.3. Frontend single-flight refresh",
                """if (!refreshPromise) {
  refreshPromise = refreshAccessToken().finally(() => {
    refreshPromise = null;
  });
}
await refreshPromise;
return axiosClient(originalRequest);""",
                "Nhiều request cùng nhận 401 sẽ chờ một Promise thay vì bắn nhiều refresh request và tự làm token rotation xung đột.",
            ),
        ),
        failures=(
            ("Đăng nhập đúng nhưng nhận 401", "Password hash/user active/email verification hoặc issuer JWT sai.", "Kiểm tra audit auth, user flags và JwtProperties."),
            ("Nhiều request refresh liên tục", "Interceptor thiếu single-flight hoặc retry loop.", "Kiểm tra cờ _retry và refreshPromise trong axiosClient."),
            ("Logout rồi token vẫn gọi API", "Chỉ xóa frontend state, chưa invalidate jti/session.", "Kiểm tra InvalidatedToken và refresh token revokedAt."),
            ("Google báo invalid_client", "Client ID/redirect origin sai môi trường.", "Đối chiếu Google Console và VITE_GOOGLE_CLIENT_ID."),
            ("403 sau khi login", "Token hợp lệ nhưng thiếu permission.", "Phân biệt JwtAuthenticationEntryPoint 401 với JwtAccessDeniedHandler 403."),
        ),
        questions=(
            ("Vì sao cần cả access và refresh token?", "Access sống ngắn giảm rủi ro; refresh quản lý phiên và có thể revoke/rotate."),
            ("Vì sao lưu hash refresh token?", "Nếu DB lộ, kẻ tấn công không dùng ngay token thô để lấy access token."),
            ("JWT có thật sự stateless không?", "Access verification gần stateless, nhưng logout/session/reuse protection cần state refresh và blacklist."),
            ("Frontend có phải lớp bảo mật không?", "Không. Route guard chỉ là UX; backend @PreAuthorize và data scope mới là lớp quyết định."),
            ("Google login tin dữ liệu frontend gửi không?", "Không. Backend phải xác minh Google credential/token và các claim bắt buộc."),
        ),
        checklist=(
            "Test login đúng, sai password, user bị khóa và email chưa xác thực.",
            "Mở hai tab, làm access token hết hạn và xác nhận chỉ một refresh request.",
            "Logout một phiên và kiểm tra access/refresh token cũ không còn sử dụng được.",
            "Kiểm tra log/audit không chứa password, raw refresh token hoặc JWT đầy đủ.",
            "Dùng user/staff/admin gọi cùng API để phân biệt 401 và 403.",
        ),
    ),
    Topic(
        source="Luong_RBAC_Permission_StaffScope_CinemaBooking.docx",
        output="03_RBAC_Permission_StaffScope_CinemaBooking.docx",
        title="RBAC, permission và staff scope theo rạp",
        purpose="Hiểu quyền chức năng và quyền dữ liệu là hai lớp khác nhau.",
        quick_steps=(
            "Đọc User, Role và Permission để hiểu quan hệ nhiều-nhiều.",
            "Xem ApplicationInitConfig để biết role/permission mặc định được seed thế nào.",
            "Tìm @PreAuthorize trong controller/service để thấy permission chặn chức năng.",
            "Đọc StaffCinema và StaffCinemaScopeService để thấy staff chỉ thao tác dữ liệu rạp được gán.",
            "Đối chiếu ProtectedRoute phía React nhưng luôn nhớ backend mới là lớp bảo mật cuối cùng.",
        ),
        core_idea=(
            "RBAC trả lời ‘người này được làm loại hành động nào’; staff scope trả lời ‘được làm trên dữ liệu của rạp nào’. "
            "Chỉ có @PreAuthorize mà không kiểm tra cinema scope vẫn có thể tạo lỗ hổng IDOR."
        ),
        flow="""JWT hợp lệ
  -> Spring chuyển scope claim thành GrantedAuthority
  -> @PreAuthorize kiểm tra permission
  -> Controller gọi Service
  -> StaffCinemaScopeService kiểm tra cinemaId thuộc assignment
  -> Repository chỉ query dữ liệu trong scope
  -> thao tác được audit

ADMIN
  -> có toàn bộ permission + global scope
STAFF
  -> permission vận hành + assigned cinema scope
USER
  -> public/own-resource permission + ownership check""",
        classes=(
            ("User / Role / Permission", "Mô hình RBAC.", "Quan hệ users_roles và roles_permissions được map trực tiếp."),
            ("ApplicationInitConfig", "Seed role, permission và tài khoản mặc định.", "Idempotent, không thay Flyway."),
            ("SecurityConfig", "Bật method security và cấu hình endpoint public.", "Request đã public không có nghĩa dữ liệu write được public."),
            ("StaffCinema / StaffCinemaId", "Bảng gán staff-rạp.", "Composite key và unique assignment."),
            ("StaffCinemaScopeService", "Kiểm tra/quy đổi phạm vi rạp.", "assert access trước thao tác dữ liệu."),
            ("ProtectedRoute.tsx", "Ẩn/chặn màn hình theo auth/permission.", "Chỉ hỗ trợ UX, không thay backend."),
        ),
        code_sections=(
            (
                "D.1. Kiểm tra permission",
                """@PreAuthorize("hasAuthority('SHOWTIME_CREATE')")
public ApiResponse<ShowtimeResponse> create(...) { ... }""",
                "Permission đặt tại biên API giúp policy dễ đọc. Service vẫn cần kiểm tra cinema scope vì permission không chứa ID rạp.",
            ),
            (
                "D.2. Kiểm tra data scope",
                """if (currentUserIsStaff()) {
    staffCinemaScopeService.assertAssigned(currentUserId, cinemaId);
}""",
                "Không được tin cinemaId do frontend gửi. Backend truy assignment trong DB rồi mới cho create/update/search.",
            ),
            (
                "D.3. Ownership của USER",
                """booking = bookingRepository.findById(id)
if (!booking.user.id.equals(currentUserId)) {
    throw new ForbiddenException(...)
}""",
                "BOOKING_VIEW_OWN cần kiểm tra chủ sở hữu resource; role USER không được dùng để đọc booking của user khác.",
            ),
        ),
        failures=(
            ("Staff xem được dữ liệu toàn hệ thống", "Query không áp scope hoặc controller dùng endpoint admin chung.", "Kiểm tra StaffCinemaScopeService và filter cinema IDs."),
            ("Có menu nhưng API trả 403", "Frontend permission stale hoặc role vừa đổi.", "Refresh phiên/token và đối chiếu @PreAuthorize."),
            ("Đổi bookingId trên URL xem được đơn khác", "Thiếu ownership check.", "Thêm find-owned query hoặc assert owner trong service."),
            ("Staff mới không thấy rạp", "Chưa có staff_cinemas assignment.", "Admin gán rạp và user tải lại scope."),
        ),
        questions=(
            ("Role khác permission thế nào?", "Role là nhóm quyền; permission là hành động cụ thể được kiểm tra ở API."),
            ("Vì sao STAFF không chỉ cần role STAFF?", "Role quá rộng; permission giúp policy rõ và mở rộng vai trò mới mà ít sửa code."),
            ("Staff scope có phải RBAC không?", "Nó bổ sung data authorization cho RBAC, giới hạn theo resource/cinema."),
            ("Vì sao frontend vẫn cần permission?", "Để ẩn thao tác không dùng được và cải thiện UX, nhưng backend vẫn xác minh lại."),
        ),
        checklist=(
            "Đăng nhập staff được gán một rạp và kiểm tra chỉ thấy dữ liệu rạp đó.",
            "Thay cinemaId bằng rạp không được gán và xác nhận backend trả 403.",
            "Dùng USER truy booking/ticket của user khác và xác nhận bị chặn.",
            "Đổi permission của role rồi đăng nhập lại để quan sát scope trong token.",
            "Kiểm tra thao tác admin/staff quan trọng có audit actor và resource ID.",
        ),
    ),
    Topic(
        source="Luong_Scheduler_GiuGhe_HetHan_TrangThaiSuatChieu_CinemaBooking.docx",
        output="06_Scheduler_GiuGhe_HetHan_TrangThaiSuatChieu_CinemaBooking.docx",
        title="Scheduler giữ ghế, booking hết hạn và trạng thái suất chiếu",
        purpose="Hiểu công việc nền giúp hệ thống tự phục hồi khi người dùng đóng tab hoặc bỏ thanh toán.",
        quick_steps=(
            "Phân biệt hold_until của seat_status với payment_expires_at của booking.",
            "Đọc HoldExpireScheduler và service/repository mà nó gọi.",
            "Đọc PendingBookingExpireScheduler để hiểu vì sao đơn PENDING không tồn tại mãi.",
            "Đọc ShowtimeStatusSyncScheduler để hiểu UPCOMING/ONGOING/ENDED được đồng bộ theo đồng hồ.",
            "Theo dõi event WebSocket được phát sau khi scheduler commit thay đổi.",
        ),
        core_idea=(
            "Countdown trên trình duyệt chỉ là UX. Database timestamp và scheduler backend mới là nguồn quyết định ghế/booking đã hết hạn. "
            "Job phải idempotent, xử lý theo batch và không làm sống lại dữ liệu đã CANCELLED/SUCCESS."
        ),
        flow="""@Scheduled fixedDelay
  -> lấy một batch ID đã hết hạn
  -> transaction cập nhật có điều kiện
  -> bỏ qua row đã được request khác xử lý
  -> commit
  -> publish AVAILABLE event sau commit
  -> vòng chạy sau xử lý batch tiếp theo

HoldExpireScheduler: HOLD quá hold_until -> AVAILABLE
PendingBookingExpireScheduler: PENDING quá payment_expires_at -> EXPIRED
ShowtimeStatusSyncScheduler: thời gian thực -> UPCOMING/ONGOING/ENDED""",
        classes=(
            ("HoldExpireScheduler", "Kích hoạt dọn seat hold.", "fixed delay từ env và batch giới hạn."),
            ("PendingBookingExpireScheduler", "Expire booking chưa thanh toán.", "Gọi business service, không tự sửa rải rác."),
            ("ShowtimeStatusSyncScheduler", "Đồng bộ trạng thái suất chiếu.", "Không ghi đè CANCELLED."),
            ("TokenCleanupTask", "Dọn token hết hạn.", "Job bảo trì bảo mật riêng."),
            ("SeatStatusRepository", "Query/conditional update hold hết hạn.", "Index hold_until và batch ID."),
            ("BookingServiceImpl", "Rule expire booking và trả ghế.", "State transition + promotion + realtime."),
        ),
        code_sections=(
            (
                "D.1. Fixed delay lấy từ cấu hình",
                """@Scheduled(fixedDelayString = "${app.booking.expired-hold-scan-delay-ms}")
public void releaseExpiredHolds() { ... }""",
                "Thời gian quét là cấu hình vận hành, không hard-code. fixedDelay tính từ lúc lần chạy trước hoàn tất.",
            ),
            (
                "D.2. Conditional batch update",
                """UPDATE seat_status
SET status = 'AVAILABLE', hold_by = NULL, hold_until = NULL
WHERE id IN (:ids)
  AND status = 'HOLD'
  AND hold_until <= :now""",
                "Điều kiện được kiểm tra lại lúc UPDATE để job không ghi đè ghế vừa được tiến trình khác xử lý.",
            ),
            (
                "D.3. Trạng thái suất chiếu suy ra từ thời gian",
                """startTime > now             -> UPCOMING
startTime <= now < endTime -> ONGOING
endTime <= now              -> ENDED
status == CANCELLED         -> giữ nguyên""",
                "Frontend có thể tính trạng thái để hiển thị nhanh, nhưng backend và DB vẫn cần đồng bộ cho filter/report.",
            ),
        ),
        failures=(
            ("Countdown hết nhưng ghế chưa đổi", "Scheduler chưa đến chu kỳ hoặc WebSocket mất kết nối.", "Refetch seat map; kiểm tra job log và hold_until DB."),
            ("Job xử lý cùng row hai lần", "Thiếu điều kiện status/timestamp khi update.", "Dùng conditional update và kiểm tra affected rows."),
            ("Job chạy lâu dần", "Quét toàn bảng, thiếu index hoặc không chia batch.", "EXPLAIN query, index hold_until/payment_expires_at và batch."),
            ("Suất CANCELLED thành ENDED", "Sync job chỉ dựa thời gian.", "Loại CANCELLED khỏi update."),
        ),
        questions=(
            ("Vì sao cần scheduler nếu frontend có countdown?", "Frontend có thể đóng/mất mạng; backend phải tự giải phóng tài nguyên."),
            ("Scheduler có chống race condition không?", "Có bằng transaction và update có điều kiện; job không được giả định dữ liệu vẫn như lúc SELECT."),
            ("fixedDelay khác fixedRate?", "fixedDelay chờ lần chạy hoàn tất; phù hợp job DB tránh chồng lượt."),
            ("Scale nhiều instance thì sao?", "Cần distributed lock/job ownership hoặc query SKIP LOCKED để tránh nhiều node xử lý trùng."),
        ),
        checklist=(
            "Giảm thời gian giữ ghế ở môi trường test và quan sát HOLD -> AVAILABLE.",
            "Đóng tab checkout rồi xác nhận backend vẫn expire booking.",
            "Chạy hai scheduler instance/test đồng thời và kiểm tra state không sai.",
            "Tạo showtime qua start/end và xác nhận UPCOMING -> ONGOING -> ENDED.",
            "Kiểm tra event realtime chỉ phát sau khi transaction thành công.",
        ),
    ),
)


TOPICS += (
    Topic(
        source="Luong_Cache_Query_Index_Flyway_CinemaBooking.docx",
        output="08_Database_Flyway_Query_Index_Cache_NPlusOne_CinemaBooking.docx",
        title="Database, Flyway, query, index, cache và N+1",
        purpose="Đọc dữ liệu từ mô hình quan hệ đến execution plan và chiến lược cache an toàn.",
        quick_steps=(
            "Đọc migration V1 trước để hiểu schema nền, sau đó lần V2 đến V14 để biết hệ thống tiến hóa thế nào.",
            "Mở entity và repository cùng lúc để nối tên field Java với cột/index PostgreSQL.",
            "Bật SQL profiling có kiểm soát, chạy một use case và đếm số query thay vì đoán N+1.",
            "Dùng EXPLAIN (ANALYZE, BUFFERS) trên query nóng với dữ liệu đủ lớn.",
            "Cuối cùng đọc CacheConfig và @Cacheable/@CacheEvict để biết dữ liệu nào được cache và khi nào bị làm mới.",
        ),
        core_idea=(
            "Database là nguồn sự thật cho booking/payment/seat. Index chỉ hữu ích khi khớp WHERE/JOIN/ORDER BY; "
            "cache chỉ nên dùng cho dữ liệu đọc nhiều, ít đổi và không mang tính cạnh tranh như phim/rạp/phòng."
        ),
        flow="""HTTP request
  -> Service chọn use case
  -> Repository query projection/entity cần thiết
  -> PostgreSQL planner chọn index/scan/join
  -> Hibernate map kết quả
  -> Mapper tạo response DTO

Master data read
  -> kiểm tra Caffeine cache
  -> cache hit: trả ngay
  -> cache miss: query DB rồi cache
Admin update
  -> transaction commit
  -> evict cache liên quan""",
        classes=(
            ("db/migration/V1...V14", "Lịch sử schema có version.", "Constraint/index mới phải đi bằng migration mới."),
            ("BaseEntity và entity package", "Ánh xạ bảng/quan hệ/audit time.", "Fetch type, nullable, enum và soft delete."),
            ("SeatStatusRepository", "Query nóng cho seat map/hold/expire.", "Lock, native batch query và index."),
            ("BookingRepository", "Ownership/page/expired pending.", "Hai bước page IDs rồi fetch graph khi cần."),
            ("PaymentRepository", "Lookup transaction/pending/filter.", "Pessimistic lock và unique pending gateway."),
            ("AnalyticsServiceImpl", "Read model dashboard.", "Projection/aggregate, tránh load entity graph."),
            ("CacheConfig", "Caffeine cache names/policy.", "Không cache seat map, booking hoặc payment."),
        ),
        code_sections=(
            (
                "D.1. N+1 query",
                """// Không tốt: 1 query bookings + N query showtime/movie
for (Booking booking : bookings) {
    booking.getShowtime().getMovie().getTitle();
}

// Hướng xử lý: projection/entity graph/fetch query đúng use case""",
                "Không bật EAGER toàn cục để chữa N+1. Hãy fetch đúng quan hệ tại query cần nó, đặc biệt khi phân trang collection.",
            ),
            (
                "D.2. Index ghép",
                """CREATE INDEX idx_seat_status_showtime_status
ON seat_status(showtime_id, status);

SELECT ... FROM seat_status
WHERE showtime_id = ? AND status = 'HOLD';""",
                "Thứ tự cột index dựa trên cách query lọc. Index làm INSERT/UPDATE tốn thêm nên không tạo cho mọi cột.",
            ),
            (
                "D.3. Cache master data",
                """@Cacheable(cacheNames = "cinemas", key = "#id")
public CinemaResponse getById(UUID id) { ... }

@CacheEvict(cacheNames = {"cinemas", "cinema-map"}, allEntries = true)
public CinemaResponse update(...) { ... }""",
                "Đọc và ghi phải thiết kế cùng nhau. Cache không có eviction đúng sẽ trả dữ liệu cũ dù DB đã đổi.",
            ),
            (
                "D.4. Đọc EXPLAIN ANALYZE",
                """EXPLAIN (ANALYZE, BUFFERS)
SELECT ...
FROM payments
WHERE transaction_no = ?;""",
                "Quan sát actual rows, loops, execution time, shared hit/read và scan type. Luôn test gần với dữ liệu production.",
            ),
        ),
        failures=(
            ("API list chậm khi dữ liệu tăng", "N+1, fetch graph quá lớn hoặc thiếu page.", "Bật SQL log/profiling, đếm query và xem execution plan."),
            ("Có index nhưng PostgreSQL không dùng", "Bảng nhỏ, selectivity thấp, cast/function sai hoặc statistics cũ.", "ANALYZE bảng và đọc planner cost/actual rows."),
            ("Admin sửa nhưng user vẫn thấy dữ liệu cũ", "Thiếu @CacheEvict hoặc key cache không nhất quán.", "Kiểm tra cache names/key và đường update."),
            ("Phân trang bị duplicate/mất row", "Fetch join collection trực tiếp với Pageable.", "Page ID trước, fetch detail theo IDs sau."),
            ("Flyway checksum mismatch", "Đã sửa migration từng chạy.", "Khôi phục file cũ và tạo migration version mới."),
        ),
        questions=(
            ("UUID v4 có phải vấn đề lớn nhất không?", "Chưa; query/index/N+1 và transaction thường quan trọng hơn ở quy mô đồ án."),
            ("Vì sao không cache seat map?", "Trạng thái ghế đổi liên tục và liên quan cạnh tranh; cache stale có thể gây UX sai."),
            ("Index có luôn tăng tốc không?", "Không. Nó đổi chi phí write/storage và chỉ giúp query có pattern phù hợp."),
            ("Flyway khác database.sql thế nào?", "Flyway là lịch sử schema chạy theo version; database.sql là bản tham khảo/bootstrap tổng hợp."),
            ("Projection giúp gì?", "Chỉ lấy cột cần dùng, giảm entity hydration và tránh kéo graph không cần thiết."),
        ),
        checklist=(
            "Vẽ ERD các bảng Movie/Cinema/Room/Seat/Showtime và Booking/Payment/Ticket.",
            "Chạy EXPLAIN ANALYZE seat map, expired hold, payment lookup và dashboard query.",
            "Bật statistics SQL để xác nhận list API không còn N+1.",
            "Sửa một master data và kiểm tra cache được evict.",
            "Đọc lần lượt V1-V14 và giải thích lý do của từng migration.",
        ),
    ),
    Topic(
        source="Luong_Frontend_React_Router_State_API_UX_CinemaBooking.docx",
        output="09_Frontend_React_Router_State_API_Realtime_UX_CinemaBooking.docx",
        title="Frontend React, router, state, API, realtime và UX",
        purpose="Theo dõi một thao tác UI từ component đến API, cache, WebSocket và trạng thái hiển thị.",
        quick_steps=(
            "Bắt đầu từ main.tsx/App.tsx để hiểu provider tree trước khi đọc từng page.",
            "Đọc AppRouter.tsx và ProtectedRoute để thấy public/user/staff/admin route.",
            "Phân biệt server state trong React Query, global client state trong Zustand và state cục bộ trong component.",
            "Đọc api/axiosClient.ts trước các file *Api.ts để hiểu token, refresh và error normalization.",
            "Dùng SeatSelectionPage và CheckoutPage làm hai ví dụ end-to-end đầy đủ nhất.",
        ),
        core_idea=(
            "Frontend không phải nguồn sự thật của ghế, giá hay thanh toán. React tối ưu trải nghiệm và đồng bộ server state; "
            "mọi quyết định nghiệp vụ quan trọng vẫn được backend xác minh lại."
        ),
        flow="""User tương tác Page/Component
  -> local form/UI state
  -> useMutation gọi domain API
  -> axiosClient gắn access token
  -> backend response/error chuẩn hóa
  -> React Query cập nhật/invalidate cache
  -> component render lại

SeatSelectionPage
  -> useQuery lấy HTTP snapshot
  -> useSeatWebSocket nhận delta
  -> merge event vào query cache
  -> periodic/focus refetch tự phục hồi nếu mất event""",
        classes=(
            ("main.tsx / App.tsx", "Mount app và provider tree.", "QueryClient, Router, error/toast/theme provider."),
            ("AppRouter.tsx", "Khai báo lazy route.", "Public/user/admin/staff layout."),
            ("ProtectedRoute.tsx", "Route guard UX.", "Loading auth, role/permission và redirect."),
            ("axiosClient.ts", "HTTP transport chung.", "Bearer token, refresh single-flight, retry và error mapping."),
            ("authStore.ts", "Global auth state.", "User, access token, permissions và clear."),
            ("React Query hooks/pages", "Server state.", "queryKey, staleTime, refetch, mutation invalidation."),
            ("useSeatWebSocket.ts", "Realtime subscription lifecycle.", "connect/subscribe/reconnect/cleanup."),
            ("SeatSelectionPage.tsx", "Seat UX phức tạp nhất.", "Selection, hold, rate limit, countdown và realtime."),
            ("CheckoutPage.tsx", "Promotion và payment UI.", "QR amount locking, polling và result transition."),
        ),
        code_sections=(
            (
                "D.1. Query key phải mô tả dữ liệu",
                """useQuery({
  queryKey: ['seat-map', showtimeId],
  queryFn: () => bookingApi.getSeatMap(showtimeId),
  enabled: Boolean(showtimeId),
});""",
                "showtimeId nằm trong key để cache hai suất chiếu không trộn dữ liệu. enabled ngăn request khi tham số chưa sẵn sàng.",
            ),
            (
                "D.2. Mutation phải đồng bộ cache",
                """useMutation({
  mutationFn: holdSeats,
  onSuccess: () => queryClient.invalidateQueries({
    queryKey: ['seat-map', showtimeId]
  })
});""",
                "Mutation đổi server state; UI cần cập nhật cache hoặc invalidate để không tiếp tục hiển thị snapshot cũ.",
            ),
            (
                "D.3. Cleanup realtime",
                """useEffect(() => {
  const client = connect(showtimeId, onSeatEvent);
  return () => client.deactivate();
}, [showtimeId, onSeatEvent]);""",
                "Cleanup tránh hai subscription cùng nhận một event, rò camera/WebSocket và lỗi removeChild khi component unmount.",
            ),
            (
                "D.4. Responsive theo nội dung",
                """grid-cols-1 md:grid-cols-2 xl:grid-cols-4
overflow-x-auto
min-w-0 break-words""",
                "Responsive không chỉ thu nhỏ font. Phải đổi grid, cho vùng dữ liệu cuộn có chủ đích và ngăn text làm vỡ layout.",
            ),
        ),
        failures=(
            ("Maximum update depth exceeded", "Effect setState phụ thuộc object/function đổi mỗi render.", "Ổn định dependency bằng useMemo/useCallback và tránh derived state dư thừa."),
            ("Click card không hoạt động khi kéo", "Drag handler nuốt click hoặc pointer capture sai.", "Dùng threshold phân biệt click/drag và cleanup pointer state."),
            ("Nhiều request refresh", "Mỗi 401 tự refresh riêng.", "Dùng refreshPromise single-flight."),
            ("Realtime nhận hai lần", "Subscribe trùng/cleanup thiếu.", "Kiểm tra StrictMode, dependency và deactivate."),
            ("Mobile bị tràn ngang", "Fixed width/min-width hoặc text dài.", "Rà overflow, min-w-0, break-words và viewport thật."),
        ),
        questions=(
            ("Zustand và React Query khác nhau thế nào?", "Zustand cho client/global state; React Query quản lý dữ liệu từ server và cache/refetch."),
            ("Vì sao không gọi Axios trực tiếp trong page?", "Domain API tập trung endpoint/type/error handling và giúp page chỉ lo UX."),
            ("ProtectedRoute có bảo mật API không?", "Không; người dùng có thể gọi API ngoài UI, nên backend vẫn kiểm tra JWT/permission."),
            ("Vì sao WebSocket vẫn cần refetch?", "WebSocket delta có thể mất khi disconnect; HTTP snapshot là cơ chế phục hồi."),
            ("Optimistic UI dùng ở giữ ghế được không?", "Có thể hiển thị loading/selected tạm, nhưng phải chờ backend xác nhận quyền giữ ghế."),
        ),
        checklist=(
            "Dùng React DevTools và Network lần theo Home -> Movie -> Showtime -> Seat -> Checkout.",
            "Làm access token hết hạn và xác nhận refresh không tạo vòng lặp.",
            "Mở hai tab seat map, kiểm tra realtime và reconnect/refetch.",
            "Kiểm tra trang public/user/admin/staff ở mobile, tablet và desktop.",
            "Thử loading, empty, error, disabled và success state của các form chính.",
        ),
    ),
    Topic(
        source="Luong_Email_Verification_Reset_Ticket_Notification_CinemaBooking.docx",
        output="10_Email_Verification_Reset_Ticket_Notification_CinemaBooking.docx",
        title="Email verification, reset password, ticket và notification",
        purpose="Hiểu token một lần, template email, gửi bất đồng bộ và xử lý lỗi nhà cung cấp mail.",
        quick_steps=(
            "Đọc luồng đăng ký trước: tạo user chưa verified, sinh token và gửi verification link.",
            "Đọc luồng forgot/reset password, chú ý response không tiết lộ email có tồn tại.",
            "Đọc EmailServiceImpl và template ticket để thấy QR và thông tin rạp được dựng thế nào.",
            "Kiểm tra AsyncConfig để hiểu request không phải chờ SMTP hoàn tất.",
            "Phân biệt lỗi nghiệp vụ chính với lỗi notification hậu xử lý.",
        ),
        core_idea=(
            "Verification/reset token là credential dùng một lần: chỉ lưu hash, có hạn sử dụng và bị vô hiệu sau khi dùng. "
            "Email là side effect; không được để SMTP chậm làm giữ transaction booking/payment quá lâu."
        ),
        flow="""RegisterPage -> POST register
  -> tạo user emailVerified=false
  -> sinh random token, lưu hash + expiresAt
  -> @Async EmailService gửi link
  -> VerifyEmailPage gọi token
  -> backend hash token, đối chiếu DB, kiểm tra expiry
  -> emailVerified=true, xóa token hash

Payment SUCCESS
  -> commit booking/ticket
  -> gửi email vé với QR từng ghế
  -> lỗi email được log/retry riêng, không rollback vé đã mua""",
        classes=(
            ("UserServiceImpl", "Register/verify/reset business rules.", "Token hash, expiry, active/verified flags."),
            ("EmailServiceImpl", "Dựng và gửi email.", "Template variables, QR attachment/inline và exception."),
            ("AsyncConfig", "Executor riêng cho mail.", "Pool size, queue, uncaught exception."),
            ("email-verification.html", "Template xác thực.", "Không chèn dữ liệu chưa escape."),
            ("ticket-email.html", "Template vé điện tử.", "Cinema/address/city/room/time/seat/QR."),
            ("VerifyEmailPage / ResetPasswordPage", "Nhận token từ URL.", "Loading/success/expired/error UX."),
        ),
        code_sections=(
            (
                "D.1. Token một lần phải lưu hash",
                """rawToken = secureRandomToken()
user.verificationTokenHash = sha256(rawToken)
user.verificationExpiresAt = now + ttl
emailLink = frontendUrl + '/verify-email?token=' + rawToken""",
                "DB không lưu token thô. Khi client gửi token, backend hash lại rồi so sánh và kiểm tra thời hạn.",
            ),
            (
                "D.2. Async không đồng nghĩa fire-and-forget vô trách nhiệm",
                """@Async("mailExecutor")
public void sendTicketEmail(...) {
    try { mailSender.send(message); }
    catch (MailException ex) { log + metric + retry policy; }
}""",
                "Cần log có correlation/booking ID và không đưa secret/QR payload đầy đủ vào log.",
            ),
        ),
        failures=(
            ("Mailtrap không thấy thư", "Sai SMTP host/port/auth/TLS hoặc app chưa nạp .env.", "Kiểm tra startup config và mail exception log."),
            ("Link luôn hết hạn", "Timezone/TTL/hash không khớp.", "So expiresAt DB và clock ứng dụng."),
            ("Thanh toán thành công nhưng không có mail", "Async mail thất bại sau commit.", "Vé vẫn xem trong My Tickets; retry mail độc lập."),
            ("Reset password lộ email tồn tại", "Response khác nhau cho email có/không.", "Luôn trả thông báo trung tính và audit/rate limit."),
        ),
        questions=(
            ("Vì sao không lưu verification token thô?", "Token là credential; hash giảm hậu quả khi DB bị lộ."),
            ("Email lỗi có rollback payment không?", "Không. Payment/ticket là nghiệp vụ chính; email được xử lý hậu commit."),
            ("Vì sao dùng @Async?", "Giảm latency request, nhưng vẫn cần executor, logging và retry phù hợp."),
            ("Một booking nhiều ghế có mấy QR?", "Mỗi booking detail/ghế có ticket QR riêng để check-in độc lập."),
        ),
        checklist=(
            "Đăng ký email mới, kiểm tra token chỉ dùng được một lần.",
            "Thử token verify/reset hết hạn và token đã dùng.",
            "Tắt SMTP rồi thanh toán để xác nhận booking không rollback.",
            "Kiểm tra email vé có phim, rạp, địa chỉ, thành phố, phòng, giờ, ghế và QR.",
            "Kiểm tra log không chứa raw verification/reset token." ,
        ),
    ),
)


BOOKING_TOPIC = Topic(
    source="",
    output="04_Booking_GiuGhe_RaceCondition_CinemaBooking.docx",
    title="Booking, giữ ghế và race condition",
    purpose="Chuyên đề viết mới về use case quan trọng nhất của hệ thống đặt vé.",
    quick_steps=(
        "Vẽ mô hình Seat khác SeatStatus: Seat là ghế vật lý, SeatStatus là trạng thái ghế theo từng showtime.",
        "Đọc HoldSeatRequest/HoldSeatResponse rồi lần vào BookingServiceImpl.holdSeats.",
        "Đọc SeatStatusRepository.findForUpdateByShowtimeAndSeats để hiểu khóa hàng và thứ tự lock.",
        "Đọc createBooking để thấy backend xác minh ghế vẫn do đúng user giữ và tính giá từ DB.",
        "Đọc handlePaymentSuccess/failure/cancel/expire để hoàn thiện state machine.",
        "Cuối cùng đọc V13/V14 để thấy database bảo vệ invariant khi application có lỗi hoặc request đồng thời.",
    ),
    core_idea=(
        "Không bán trùng ghế nhờ nhiều lớp phối hợp: transaction, pessimistic row lock theo thứ tự ổn định, kiểm tra owner/expiry, "
        "state transition có điều kiện và unique constraint. WebSocket chỉ giúp UI biết kết quả nhanh, không thay các lớp bảo vệ này."
    ),
    flow="""1. User chọn ghế trên SeatSelectionPage
2. POST /api/v1/bookings/hold {showtimeId, seatIds[]}
3. Auth/rate limit/DTO validation
4. BookingServiceImpl.holdSeats @Transactional
5. Khóa các SeatStatus theo showtime + seat IDs
6. Kiểm tra đủ số row, đúng phòng, trạng thái/hold expiry
7. AVAILABLE hoặc hold cũ của chính user -> HOLD
8. Ghi holdBy + holdUntil + version
9. Commit
10. Publish HOLD event sau commit

Tạo booking
  -> lock/xác minh toàn bộ hold của user còn hạn
  -> kiểm tra một pending booking tương đương
  -> snapshot priceAtBooking từng ghế
  -> tính subtotal/promotion/total ở backend
  -> booking PENDING + paymentExpiresAt

Thanh toán SUCCESS
  -> lock booking/payment
  -> PENDING -> SUCCESS
  -> HOLD -> BOOKED
  -> tạo Ticket cho từng BookingDetail
  -> after commit: WebSocket + email""",
    classes=(
        ("BookingController", "API seat map/hold/create/cancel/promotion/ticket.", "@Valid, permission và current user context."),
        ("BookingServiceImpl", "State machine và transaction booking.", "holdSeats/createBooking/success/failure/cancel/expire/check-in."),
        ("SeatStatusRepository", "Lock và batch transition ghế.", "PESSIMISTIC_WRITE, ordered IDs, conditional release/update."),
        ("BookingRepository", "Pending/ownership/lock booking.", "find locked for payment và unique-equivalent pending policy."),
        ("Seat / SeatStatus", "Ghế vật lý và trạng thái theo suất.", "unique(seat_id, showtime_id), holdBy, holdUntil, version."),
        ("Booking / BookingDetail", "Đơn và snapshot từng ghế.", "status, total, secureToken, paymentExpiresAt, priceAtBooking."),
        ("SeatHoldRateLimitService", "Chống spam endpoint hold.", "Không thay lock DB; chỉ bảo vệ tài nguyên/UX."),
        ("SeatStatusPublisher", "Thông báo sau commit.", "HOLD/BOOKED/AVAILABLE theo showtime."),
        ("SeatSelectionPage.tsx", "UI chọn/hold/countdown.", "Rate limit message, stale selection cleanup và realtime merge."),
    ),
    code_sections=(
        (
            "D.1. Khóa hàng để tuần tự hóa hai người tranh cùng ghế",
            """@Lock(LockModeType.PESSIMISTIC_WRITE)
@Query("select ss from SeatStatus ss " +
       "where ss.showtime.id = :showtimeId " +
       "and ss.seat.id in :seatIds order by ss.seat.id")
List<SeatStatus> findForUpdateByShowtimeAndSeats(...);""",
            "Transaction A giữ lock trước sẽ kiểm tra và cập nhật. Transaction B chờ, sau đó đọc trạng thái HOLD mới và bị từ chối. ORDER BY ổn định giảm nguy cơ deadlock khi lock nhiều ghế.",
        ),
        (
            "D.2. Validation sau khi lock",
            """locked = repository.findForUpdateByShowtimeAndSeats(showtimeId, sortedSeatIds)
require locked.size == requestedUniqueSeatIds.size
for each seatStatus:
  require seat belongs to showtime.room
  require AVAILABLE
      or expired HOLD
      or active HOLD owned by current user
  set HOLD, holdBy=currentUser, holdUntil=deadline""",
            "Kiểm tra trước lock có thể bị thay đổi ngay sau khi đọc. Vì vậy điều kiện quyết định phải được đánh giá trong transaction sau khi sở hữu lock.",
        ),
        (
            "D.3. Giá được snapshot",
            """seatPrice = showtime.basePrice * seat.priceMultiplier
bookingDetail.priceAtBooking = seatPrice
subtotal = sum(detail.priceAtBooking)
discount = promotionPolicy.calculate(subtotal)
booking.totalPrice = subtotal - discount""",
            "Không đọc lại giá hiện tại khi xem lịch sử. priceAtBooking giữ đúng giá lúc mua dù admin đổi giá ghế/suất sau đó.",
        ),
        (
            "D.4. Transition idempotent khi payment thành công",
            """lock booking
if booking.status == SUCCESS:
    return existing response
require booking.status == PENDING
require payment not expired
booking.status = SUCCESS
mark held seats BOOKED
create missing tickets only
commit""",
            "Callback lặp không được sinh thêm ticket hoặc tăng usedCount promotion lần nữa. Terminal state được nhận diện trước khi side effect.",
        ),
        (
            "D.5. Database invariant",
            """unique (seat_id, showtime_id)
unique (booking_id, seat_id)
check booking/payment/seat/ticket status
partial unique index cho pending payment theo policy
indexes cho hold_until và payment_expires_at""",
            "Validation Java cho message đẹp; constraint DB là lớp cuối chống dữ liệu sai khi có concurrency, bug hoặc script ngoài ứng dụng.",
        ),
    ),
    failures=(
        ("One or more seats are not available", "Người khác giữ/mua trước hoặc snapshot UI đã cũ.", "Trả thông báo tiếng Việt, refetch seat map và bỏ selected không còn hợp lệ."),
        ("Hai booking PENDING cùng bộ ghế", "createBooking thiếu lock/invariant hoặc retry request không idempotent.", "Lock hold, kiểm tra pending tương đương và constraint V14."),
        ("Hết countdown nhưng ghế vẫn HOLD", "Countdown chỉ chạy client; scheduler/event chậm.", "Backend timestamp là chuẩn, refetch và kiểm tra scheduler."),
        ("Hủy trả nhầm ghế người khác", "Release chỉ lọc showtime/seat, thiếu holdBy/booking.", "Conditional release theo owner/status/expiry."),
        ("Deadlock khi giữ nhiều ghế", "Hai transaction lock cùng tập ghế theo thứ tự khác.", "Sort seat IDs trước query/lock và giữ transaction ngắn."),
        ("Promotion count sai", "Apply/success/retry tăng usedCount nhiều lần.", "Chỉ consume promotion tại transition được chọn và idempotent."),
        ("Thanh toán thành công nhưng booking expired", "Callback đến trễ.", "Ghi nhận tiền và tạo refund request; không lấy lại ghế đã bán."),
    ),
    questions=(
        ("Pessimistic lock khác optimistic lock?", "Pessimistic khóa row ngay; optimistic phát hiện version conflict khi update. Luồng tranh ghế dùng row lock để quyết định tức thời."),
        ("Vì sao cần SeatStatus riêng?", "Một Seat vật lý có trạng thái khác nhau ở mỗi Showtime; không thể lưu AVAILABLE/BOOKED trực tiếp trên Seat."),
        ("WebSocket có chống race condition?", "Không. Nó thông báo kết quả; transaction và DB mới quyết định người thắng."),
        ("Vì sao có cả holdUntil và paymentExpiresAt?", "Hold quản lý tài nguyên ghế; payment expiry quản lý vòng đời booking/checkout. Chúng liên quan nhưng không đồng nhất."),
        ("Vì sao transaction phải ngắn?", "Lock giữ trong transaction; gọi network/email lâu sẽ giảm throughput và tăng timeout/deadlock."),
        ("Nếu backend có nhiều instance?", "DB lock/constraint vẫn chung; rate limit/scheduler/WebSocket broker in-memory cần nâng cấp thành shared infrastructure."),
    ),
    checklist=(
        "Mở hai user và gửi hold cùng ghế gần như đồng thời; chỉ một request thành công.",
        "Gửi seatIds trùng, rỗng, sai showtime/phòng và ghế không tồn tại.",
        "User A giữ ghế rồi user B cố create booking bằng seat ID đó.",
        "Double-click create booking và initiate payment.",
        "Để hold/booking hết hạn rồi đặt lại mà không refresh trang.",
        "Gửi callback thành công hai lần và kiểm tra chỉ có một bộ ticket.",
        "Kiểm tra rollback giữa chừng không phát WebSocket event và không để state nửa vời.",
        "Đối chiếu seat_status, bookings, booking_details, payments, tickets sau từng test.",
    ),
)


TEST_TOPIC = Topic(
    source="",
    output="13_Kiem_thu_Trien_khai_Bao_ve_CinemaBooking.docx",
    title="Kiểm thử, triển khai và kịch bản bảo vệ",
    purpose="Từ test tự động đến checklist demo và cách chẩn đoán lỗi trong ngày bảo vệ.",
    quick_steps=(
        "Chạy backend unit/integration test trước, sau đó frontend lint/typecheck/Vitest/build.",
        "Chuẩn bị database riêng, migration sạch và mock data có đủ case success/failure/expired/cancelled/refund.",
        "Test luồng quan trọng theo trạng thái DB, không chỉ nhìn thông báo UI.",
        "Diễn tập callback/webhook/ngrok/camera/SMTP trước ngày demo.",
        "Chuẩn bị câu trả lời về race condition, idempotency, RBAC, N+1 và giới hạn khi scale nhiều instance.",
    ),
    core_idea=(
        "Không thể chứng minh hệ thống không có bug, nhưng có thể chứng minh các invariant quan trọng bằng test nhiều lớp, "
        "quan sát được failure và có đường phục hồi. Ngày bảo vệ cần một môi trường có thể tái lập, không phụ thuộc dữ liệu ngẫu nhiên."
    ),
    flow="""Static checks
  -> compile/typecheck/lint
Unit tests
  -> pure rule/mapper/validator
Service tests
  -> state transition + mocked collaborator
Repository/integration tests
  -> PostgreSQL lock/query/constraint/Flyway
API tests
  -> security/validation/response contract
Manual scenario
  -> browser + two users + gateway/webhook + WebSocket
Observability
  -> log/audit/payment event/DB assertion""",
    classes=(
        ("src/test/java", "Backend unit/integration tests.", "Service, security, repository, payment callbacks."),
        ("PostgresIntegrationTest", "DB thật/giống production.", "Migration, native query, lock và constraint."),
        ("Vitest tests", "Frontend logic/component tests.", "API error mapping, hooks/util/state."),
        ("mvn test", "Chạy backend suite.", "Phải chạy tại thư mục có pom.xml."),
        ("npm run check", "Lint + test + TypeScript + build.", "Chạy tại cinema-client."),
        ("mock-data.sql", "Dữ liệu demo.", "Không thay migration; phải khớp secret/config test."),
        (".env.example", "Danh mục cấu hình triển khai.", "Không chứa secret thật."),
    ),
    code_sections=(
        (
            "D.1. Given-When-Then",
            """Given: booking PENDING, hold còn hạn, payment hợp lệ
When: xử lý callback SUCCESS hai lần
Then: booking SUCCESS, seat BOOKED,
      một ticket mỗi ghế, không side effect trùng""",
            "Test phải mô tả invariant và trạng thái cuối, không chỉ verify một method được gọi.",
        ),
        (
            "D.2. Concurrency test",
            """start transaction A and B together
both request same showtime/seat
wait for both futures
assert exactly one success
assert DB has one active hold/booking owner""",
            "Race condition không thể chứng minh chỉ bằng unit test mock repository; cần integration test với PostgreSQL.",
        ),
        (
            "D.3. Contract test lỗi",
            """POST /bookings/hold invalid body
expect 400
expect stable code/message/path/fieldErrors
expect no stack trace or SQL""",
            "Client phụ thuộc error contract; thay đổi message/framework không được làm vỡ UX.",
        ),
    ),
    failures=(
        ("mvn báo không có project", "Chạy ở thư mục cha không có pom.xml.", "cd cinema-booking-system rồi mvn test."),
        ("Test DB không chạy khi Docker tắt", "Integration profile dùng Testcontainers.", "Bật Docker hoặc cấu hình PostgreSQL test riêng; không dùng DB production."),
        ("VNPay Không tìm thấy website/GD", "Sai terminal config hoặc callback/ngrok không hoạt động.", "Kiểm tra sandbox credentials, return/IPN URL và tunnel."),
        ("Camera không mở trên điện thoại", "HTTP không secure hoặc permission/device conflict.", "Dùng HTTPS, cấp quyền và dừng stream cũ."),
        ("Demo phụ thuộc thời gian", "Mock showtime/hold đã hết hạn.", "Seed relative NOW và chạy sanity query trước demo."),
    ),
    questions=(
        ("Unit test khác integration test?", "Unit cô lập logic; integration xác minh wiring/DB/query/constraint thật."),
        ("Test race condition thế nào?", "Hai transaction đồng thời trên PostgreSQL và assert chỉ một state transition thành công."),
        ("Vì sao không test trên DB production?", "Test có thể phá dữ liệu, không tái lập và vi phạm an toàn."),
        ("Ưu tiên test gì trước bảo vệ?", "Auth, hold race, booking expiry, payment idempotency, QR check-in scope và cancel/refund."),
        ("Có đảm bảo không bug không?", "Không tuyệt đối; trình bày invariant, coverage, monitoring và recovery strategy."),
    ),
    checklist=(
        "Chạy mvn test và npm run check từ đúng thư mục.",
        "Khởi tạo DB sạch bằng Flyway, seed mock data và đăng nhập admin/staff/user.",
        "Demo public browse, giữ ghế hai user, promotion, VNPay/SePay, ticket email và QR check-in.",
        "Demo hủy suất/refund, staff scope, audit log và dashboard.",
        "Chuẩn bị phương án khi ngrok, SMTP hoặc gateway sandbox lỗi.",
        "Không chiếu .env/secret/API key trên màn hình bảo vệ.",
    ),
)


def build_topic(topic: Topic) -> Path:
    source = DOCS_DIR / topic.source
    if not source.exists():
        raise FileNotFoundError(source)
    output = OUTPUT_DIR / topic.output
    shutil.copy2(source, output)
    document = Document(output)
    ensure_compatibility_styles(document)
    setup_learning_styles(document)
    if any(CURRENT_UPDATE_MARKER in paragraph.text for paragraph in document.paragraphs):
        raise RuntimeError(f"Current update marker already exists in {output.name}")
    append_topic_update(document, topic)
    document.save(output)
    return output


def build_standalone(topic: Topic) -> Path:
    output = OUTPUT_DIR / topic.output
    document = Document()
    ensure_compatibility_styles(document)
    apply_styles(document)
    setup_learning_styles(document)
    add_title(document, topic.title, topic.purpose)
    append_topic_update(document, topic)
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CinemaBooking.vn - Learning Curriculum").italic = True
    document.save(output)
    return output


def build_roadmap() -> Path:
    source = DOCS_DIR / "Lo_trinh_hoc_toan_bo_he_thong_CinemaBooking.docx"
    output = OUTPUT_DIR / "00_Lo_trinh_hoc_toan_bo_he_thong_CinemaBooking.docx"
    shutil.copy2(source, output)
    document = Document(output)
    ensure_compatibility_styles(document)
    setup_learning_styles(document)
    add_update_marker(
        document,
        "Lộ trình học bộ tài liệu cuối",
        "Học theo tầng, sau đó quay lại source và tự chạy từng kịch bản.",
    )
    add_heading(document, "A. Bộ tài liệu theo thứ tự", level=1)
    add_table(
        document,
        ["Thứ tự", "Tài liệu", "Mục tiêu"],
        (
            ("01", "Kiến trúc và bản đồ code", "Nhìn toàn hệ thống trước khi học chi tiết."),
            ("02", "Auth/JWT/Session/Google", "Hiểu đăng nhập và vòng đời phiên."),
            ("03", "RBAC/Permission/Staff Scope", "Hiểu quyền chức năng và quyền dữ liệu."),
            ("04", "Booking/Giữ ghế/Race condition", "Hiểu transaction cạnh tranh cốt lõi."),
            ("05", "WebSocket realtime seat map", "Hiểu snapshot, delta và after commit."),
            ("06", "Scheduler hết hạn", "Hiểu tự nhả ghế/expire/sync showtime."),
            ("07", "VNPay/SePay/Refund/Reconciliation", "Hiểu thanh toán đáng tin cậy."),
            ("08", "Database/Flyway/Query/Cache", "Hiểu dữ liệu và hiệu năng."),
            ("09", "Frontend React", "Hiểu router/state/API/realtime/UX."),
            ("10", "Email/Notification", "Hiểu token một lần và side effect async."),
            ("11", "Admin/Staff/QR/Map", "Hiểu vận hành rạp."),
            ("12", "Exception/Audit/Soft delete", "Hiểu lỗi, lịch sử và monitoring."),
            ("13", "Kiểm thử/Triển khai/Bảo vệ", "Biết chứng minh hệ thống hoạt động."),
            ("14", "Tra cứu class/file/API", "Dùng khi mở IDE và debug."),
        ),
    )

    add_heading(document, "B. Lộ trình học bốn tuần", level=1)
    add_table(
        document,
        ["Tuần", "Trọng tâm", "Bài thực hành"],
        (
            ("1", "01-03: kiến trúc, Auth, RBAC", "Trace login; test 401/403; vẽ security flow."),
            ("2", "04-07: booking, realtime, scheduler, payment", "Hai user tranh ghế; callback lặp; expire/late payment."),
            ("3", "08-12: DB, frontend, email, operation, exception", "EXPLAIN query; trace React; QR/email/audit."),
            ("4", "13-14 và source", "Chạy test, diễn tập demo, tự trả lời câu hỏi bảo vệ."),
        ),
    )

    add_heading(document, "C. Cách học một chuyên đề", level=1)
    add_numbered(
        document,
        (
            "Đọc phần học nhanh và tự vẽ lại luồng bằng giấy.",
            "Mở các class trong bảng bản đồ file; không đọc tuần tự cả repository.",
            "Đặt breakpoint ở controller, service và callback/scheduler quan trọng.",
            "Chạy một case thành công và ít nhất hai case lỗi.",
            "Kiểm tra database/log/event thay vì chỉ nhìn UI.",
            "Tự trả lời câu hỏi bảo vệ không nhìn tài liệu rồi mới chuyển chương.",
        ),
    )
    add_callout(
        document,
        "Nguyên tắc nguồn sự thật",
        "Tài liệu giúp hiểu ý tưởng và đường đi. Khi tên method, endpoint hoặc policy có khác biệt, source code hiện tại và Flyway migration là nguồn đúng cuối cùng.",
        "FFF7ED",
        "C2410C",
    )
    document.save(output)
    return output


def build_reference() -> Path:
    source = DOCS_DIR / "learning" / "14_Tra_cuu_Class_File_API_CinemaBooking.docx"
    output = OUTPUT_DIR / "14_Tra_cuu_Class_File_API_CinemaBooking.docx"
    if not source.exists():
        raise FileNotFoundError(source)
    shutil.copy2(source, output)
    document = Document(output)
    ensure_compatibility_styles(document)
    setup_learning_styles(document)
    if not any(CURRENT_UPDATE_MARKER in paragraph.text for paragraph in document.paragraphs):
        add_update_marker(
            document,
            "Cách dùng tài liệu tra cứu",
            "Đây là index mở IDE, không phải chương cần học thuộc.",
        )
        add_numbered(
            document,
            (
                "Bắt đầu từ page/component khi lỗi xuất hiện trên UI.",
                "Mở file domain API để xác định endpoint/request/response.",
                "Mở controller để xem validation và permission.",
                "Mở service để tìm business rule và transaction.",
                "Mở repository/migration để xác minh query, lock và constraint.",
                "Mở test liên quan để xem invariant đã được chứng minh thế nào.",
            ),
        )
    document.save(output)
    return output


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    generated = [build_roadmap()]
    for topic in TOPICS:
        generated.append(build_topic(topic))
    generated.append(build_standalone(BOOKING_TOPIC))
    generated.append(build_standalone(TEST_TOPIC))
    generated.append(build_reference())
    generated.sort(key=lambda path: path.name)

    print("Generated focused learning curriculum:")
    for path in generated:
        print(f"- {path.name}")


TOPICS += (
    Topic(
        source="Huong_dan_kien_truc_va_luong_hoat_dong_CinemaBookingSystem.docx",
        output="01_Kien_truc_va_ban_do_code_CinemaBooking.docx",
        title="Kiến trúc và bản đồ code theo source hiện tại",
        purpose="Nối kiến trúc tổng thể với package, class và file frontend đang tồn tại trong repository.",
        quick_steps=(
            "Đọc sơ đồ React -> Security -> Controller -> Service -> Repository -> PostgreSQL.",
            "Chọn một use case như giữ ghế và lần đúng đường đi trước khi học từng annotation.",
            "Phân biệt dữ liệu đồng bộ trong transaction với side effect sau commit như WebSocket/email.",
            "Đọc configuration/security trước, sau đó mới đi sâu từng domain service.",
            "Dùng tài liệu tra cứu class cuối bộ khi cần tìm file, không học thuộc tên toàn bộ class.",
        ),
        core_idea=(
            "Controller giữ contract HTTP mỏng; Service giữ business rule và transaction; Repository giữ truy vấn; "
            "DTO/Mapper giữ API không phụ thuộc entity. Frontend chia route, domain API, server state và UI state."
        ),
        flow="""React Page/Component
  -> domainApi.ts
  -> axiosClient.ts
  -> Spring Security JWT filter
  -> Controller (@Valid, @PreAuthorize)
  -> Service (@Transactional, business policy)
  -> Repository (JPA/JPQL/native/projection)
  -> PostgreSQL constraints/indexes
  -> Mapper/Response DTO
  -> React Query cache

Sau commit
  -> WebSocket event / email / audit / payment notification""",
        classes=(
            ("CinemaBookingSystemApplication", "Entry point Spring Boot.", "@SpringBootApplication, scheduling/async config được nạp."),
            ("SecurityConfig", "Security filter chain.", "Public endpoints, JWT resource server, 401/403."),
            ("controller package", "HTTP adapter.", "Request mapping, validation, permission, response."),
            ("service/impl package", "Application/domain rules.", "Transaction boundary, state transition, orchestration."),
            ("repository package", "Persistence adapter.", "Query, lock, projection, page."),
            ("entity/enums", "Mô hình dữ liệu và trạng thái.", "Relationship, constraint và lifecycle."),
            ("dto/mapper", "API contract.", "Không trả entity/password/internal fields."),
            ("payment/websocket/security", "Adapter/cross-cutting nâng cao.", "Gateway, realtime, JWT/RBAC/rate limit."),
            ("frontend src/pages/components/api", "Presentation và client orchestration.", "Page -> API -> query/store/hook."),
        ),
        code_sections=(
            (
                "D.1. Controller mỏng",
                """@PostMapping("/hold")
@PreAuthorize("hasAuthority('BOOKING_CREATE')")
public ApiResponse<HoldSeatResponse> hold(@Valid @RequestBody HoldSeatRequest request) {
    return ApiResponse.success(bookingService.holdSeats(request));
}""",
                "Controller không tự lock ghế hoặc tính giá. Nó kiểm tra contract/quyền rồi giao use case cho service.",
            ),
            (
                "D.2. Transaction boundary ở service",
                """@Transactional
public HoldSeatResponse holdSeats(HoldSeatRequest request) {
    // load/lock -> validate -> transition -> save
}""",
                "Transaction bao quanh một thay đổi nghiệp vụ hoàn chỉnh. Gọi gateway/email trong transaction dài có thể giữ lock không cần thiết.",
            ),
            (
                "D.3. DTO bảo vệ contract",
                """Request DTO -> validation -> service
Entity      -> mapper     -> Response DTO""",
                "DTO giúp API ổn định, tránh vòng lặp JSON và không làm lộ password/hash/internal relation.",
            ),
        ),
        failures=(
            ("Controller chứa nhiều logic", "Business rule bị đặt ở HTTP layer.", "Chuyển orchestration/transaction về service và viết test service."),
            ("LazyInitializationException", "Mapper truy lazy relation sau transaction.", "Fetch đúng use case hoặc map trong transaction/projection."),
            ("Circular JSON/N+1", "Trả entity trực tiếp.", "Response DTO + mapper + query plan phù hợp."),
            ("WebSocket báo trạng thái DB rollback", "Publish trước commit.", "Đăng ký afterCommit rồi mới broadcast."),
            ("Frontend page quá lớn", "Trộn API, state, format và UI.", "Tách domain API/hook/component theo trách nhiệm."),
        ),
        questions=(
            ("Vì sao cần Service interface/implementation?", "Tách contract, dễ test/thay implementation và giữ controller không phụ thuộc chi tiết."),
            ("Entity khác DTO?", "Entity map persistence; DTO là contract request/response có validation và giới hạn field."),
            ("Business rule nên nằm đâu?", "Trong service/domain policy, không ở frontend hay controller."),
            ("Cross-cutting concern là gì?", "Security, exception, audit, cache, logging áp dụng nhiều domain và được cấu hình tập trung."),
            ("Hệ thống mở rộng nhiều instance cần gì?", "Redis/distributed rate limit, shared broker, distributed job coordination và externalized session/cache."),
        ),
        checklist=(
            "Lần HomePage -> movie API -> controller/service/repository.",
            "Lần SeatSelectionPage -> hold -> DB lock -> WebSocket after commit.",
            "Lần CheckoutPage -> payment gateway -> callback/webhook -> ticket/email.",
            "Chỉ ra transaction boundary và source of truth của ba luồng trên.",
            "Giải thích vai trò từng package mà không đọc tài liệu.",
        ),
    ),
    Topic(
        source="Luong_WebSocket_Realtime_SeatMap_CinemaBooking_hoc_nhanh.docx",
        output="05_WebSocket_Realtime_SeatMap_CinemaBooking.docx",
        title="Đối chiếu WebSocket với source hiện tại",
        purpose="Giữ nguyên bản học nhanh được trình bày tốt và bổ sung checklist source/reliability hiện tại.",
        quick_steps=(
            "Học phần Học nhanh có sẵn trước, sau đó quay lại phần chuyên sâu từng class.",
            "Mở WebSocketConfig, SeatStatusEvent và SeatStatusPublisher song song với sơ đồ.",
            "Mở useSeatWebSocket.ts và SeatSelectionPage.tsx để thấy connect/subscribe/cache update/cleanup.",
            "Đặt breakpoint tại BookingServiceImpl và scheduler để xác nhận publish sau commit.",
            "Test bằng hai trình duyệt, sau đó chủ động tắt WebSocket để quan sát HTTP refetch tự phục hồi.",
        ),
        core_idea=(
            "HTTP snapshot + WebSocket delta + fallback refetch tạo thành một hệ thống realtime có khả năng tự phục hồi. "
            "Database transaction mới chống bán trùng; WebSocket chỉ phát trạng thái đã commit để các client cập nhật nhanh."
        ),
        flow="""Mount SeatSelectionPage
  -> GET seat map snapshot
  -> connect /ws-native
  -> subscribe /topic/seatmap/{showtimeId}

Backend transition thành công
  -> transaction commit
  -> SeatStatusPublisher.publishBulk/publishHold
  -> SeatStatusEvent
  -> useSeatWebSocket callback
  -> queryClient.setQueryData/invalidate
  -> UI đổi màu

Disconnect/mất event
  -> reconnect + HTTP refetch
Unmount/đổi showtime
  -> unsubscribe/deactivate""",
        classes=(
            ("WebSocketConfig", "STOMP endpoint và broker prefix.", "Endpoint /ws-native, /topic và allowed origin."),
            ("SeatStatusEvent", "Realtime contract.", "showtimeId, seatId/status, hold metadata."),
            ("SeatStatusPublisher", "Publish theo showtime topic.", "after-commit và bulk event."),
            ("BookingServiceImpl", "Nguồn state transition.", "Hold/book/release rồi mới publish."),
            ("HoldExpireScheduler", "Nguồn event hết hạn.", "Release batch rồi publish AVAILABLE."),
            ("useSeatWebSocket.ts", "STOMP client lifecycle.", "URL, reconnectDelay, subscribe, JSON parse, cleanup."),
            ("SeatSelectionPage.tsx", "Merge event vào UI.", "MY_HOLD/other HOLD/BOOKED và selected cleanup."),
        ),
        code_sections=(
            (
                "D.1. Topic cô lập theo suất chiếu",
                """/topic/seatmap/{showtimeId}""",
                "Client chỉ nhận thay đổi của suất đang xem; payload nhỏ hơn và không phải lọc sự kiện toàn hệ thống.",
            ),
            (
                "D.2. Publish sau commit",
                """afterCommit(() ->
  messagingTemplate.convertAndSend(topic, event)
);""",
                "Nếu transaction rollback, không có event giả khiến UI lệch database.",
            ),
            (
                "D.3. Reconnect không được subscribe trùng",
                """connect -> subscribe once
on reconnect -> restore subscription
on cleanup -> unsubscribe + deactivate""",
                "Lifecycle phải gắn với showtimeId; callback nên ổn định để effect không tạo client mới mỗi render.",
            ),
        ),
        failures=(
            ("Hai khung camera/WebSocket", "Component mount/subscription hai lần và cleanup thiếu.", "Theo dõi active client và cleanup trong effect."),
            ("Ghế chỉ đổi sau refresh", "Không nhận event hoặc handler không cập nhật cache.", "Kiểm tra Network WS frames, topic và query key."),
            ("Event tới nhưng UI sai owner", "Không phân biệt holdBy với current user.", "Derive MY_HOLD ở client từ user/hold metadata."),
            ("Qua HTTPS không kết nối", "Dùng ws thay vì wss/proxy host sai.", "Build URL theo protocol và kiểm tra Vite/nginx websocket upgrade."),
            ("Message parse làm crash", "Payload khác contract hoặc không validate.", "Guard JSON/type và refetch khi event không hợp lệ."),
        ),
        questions=(
            ("WebSocket có thay REST không?", "Không. REST lấy snapshot/thực hiện command; WebSocket phát delta."),
            ("WebSocket có cần JWT không?", "Seat status có thể public; command hold/payment vẫn bắt buộc JWT. Private topic thì cần auth handshake."),
            ("Mất event xử lý thế nào?", "Reconnect và refetch snapshot để hội tụ lại state DB."),
            ("Scale nhiều backend instance?", "Cần shared broker/Redis/Kafka relay thay simple in-memory broker."),
            ("Vì sao không gửi toàn bộ seat map mỗi lần?", "Delta nhỏ hơn; snapshot chỉ dùng lúc mở/refetch."),
        ),
        checklist=(
            "Mở hai browser cùng showtime và giữ/hủy/thanh toán cùng ghế.",
            "Kiểm tra một transition tạo đúng số event mong đợi.",
            "Tắt mạng rồi bật lại, xác nhận reconnect và snapshot hội tụ.",
            "Rời trang/đổi showtime, kiểm tra subscription cũ đã cleanup.",
            "Buộc transaction rollback và xác nhận không phát event sai.",
        ),
    ),
    Topic(
        source="Luong_thanh_toan_VNPay_SePay_CinemaBooking.docx",
        output="07_Thanh_toan_VNPay_SePay_Refund_Reconciliation_CinemaBooking.docx",
        title="Cập nhật payment reliability theo source hiện tại",
        purpose="Giữ nguyên bản VNPay/SePay dễ học và bổ sung lock, unique pending, event, late payment, refund và reconciliation.",
        quick_steps=(
            "Học phần VNPay redirect và SePay webhook có sẵn trước.",
            "Đọc PaymentServiceImpl.initiatePayment để hiểu lock booking và tái sử dụng pending payment.",
            "Đọc từng callback/webhook với thứ tự signature -> transaction lookup -> amount -> state -> booking transition.",
            "Đọc PaymentEventServiceImpl và RefundServiceImpl để hiểu vận hành sau lỗi/late payment.",
            "Mở admin Payment page để thấy giao dịch, hoàn tiền, đối soát và nhật ký là bốn read model khác nhau.",
        ),
        core_idea=(
            "Gateway không quyết định trực tiếp vé. Backend xác minh callback/webhook, khóa payment/booking, áp dụng state transition idempotent "
            "và chỉ phát hành ticket khi booking còn hợp lệ. Tiền đến trễ được ghi nhận rồi đưa vào refund workflow, không cưỡng ép mở lại booking."
        ),
        flow="""initiatePayment
  -> lock booking
  -> require owner + PENDING + not expired
  -> amount phải bằng booking.totalPrice
  -> lock pending payment
  -> reuse nếu cùng method/amount còn hợp lệ
  -> expire stale pending nếu amount đổi
  -> tạo payment/transactionNo duy nhất
  -> gateway tạo redirect URL hoặc QR data

callback/webhook
  -> verify authenticity
  -> lock payment by transactionNo
  -> validate amount + provider status
  -> idempotency check
  -> booking success hoặc late-success refund request
  -> PaymentEvent audit
  -> commit -> ticket/email/realtime""",
        classes=(
            ("PaymentController", "Initiate/callback/webhook/admin APIs.", "Public callback nhưng luôn verify signature/key."),
            ("PaymentServiceImpl", "Payment orchestration.", "Lock, amount, idempotency, late payment."),
            ("PaymentGateway", "Abstraction theo cổng.", "VNPay/SePay/MoMo adapter không chứa booking policy."),
            ("PaymentRepository", "Locked transaction/pending lookup.", "Pessimistic lock và unique partial index."),
            ("PaymentEventServiceImpl", "Append-only event/audit.", "Callback outcome, signature fail, retry, status transition."),
            ("RefundServiceImpl", "Refund workflow.", "REQUESTED -> PROCESSING/SUCCESS/FAILED theo schema hiện tại."),
            ("BookingServiceImpl", "Chốt booking/seat/ticket.", "Idempotent SUCCESS/failure/expired transition."),
            ("CheckoutPage.tsx", "Chọn method/promotion/QR polling UX.", "Khóa mã giảm giá sau tạo QR và tự chuyển kết quả."),
        ),
        code_sections=(
            (
                "D.1. Lock khi khởi tạo",
                """bookingRepository.findLockedForPaymentInitiation(bookingId)
paymentRepository.findLockedPendingByBookingId(bookingId)""",
                "Hai click đồng thời được tuần tự hóa; unique constraint trong DB là lớp chặn cuối cho một pending payment theo gateway/policy.",
            ),
            (
                "D.2. Idempotent callback",
                """if (payment.status == SUCCESS) {
    recordDuplicateCallbackEvent();
    return currentResult;
}""",
                "Gateway có thể retry. Callback lặp phải trả kết quả ổn định mà không sinh vé, tăng promotion count hoặc gửi email lần nữa.",
            ),
            (
                "D.3. Late successful payment",
                """money received
booking no longer payable
  -> retain SUCCESS receipt/event
  -> create refund request
  -> do not restore expired/cancelled booking""",
                "Đây là cách bảo toàn tài chính: không bỏ qua tiền đã nhận nhưng cũng không bán lại ghế đã được người khác mua.",
            ),
        ),
        failures=(
            ("Bấm thanh toán nhiều lần tạo nhiều giao dịch", "Thiếu booking lock/pending reuse/unique DB rule.", "Kiểm tra V13/V14 và locked repositories."),
            ("Đổi mã nhưng QR còn số tiền cũ", "Payment đã cố định amount.", "Bỏ QR/pending cũ rồi tạo payment mới sau khi cập nhật booking total."),
            ("Đã chuyển khoản nhưng chưa SUCCESS", "Webhook không tới/auth sai/content không match.", "Kiểm tra ngrok, SePay webhook log và PaymentEvent."),
            ("Callback đến sau expiry", "Ngân hàng/gateway xử lý chậm.", "Tạo refund request thay vì hồi sinh booking."),
            ("Callback xử lý hai lần", "Không lock/idempotency guard.", "Lock theo transactionNo và kiểm tra terminal state."),
        ),
        questions=(
            ("Vì sao amount từ frontend không đáng tin?", "Frontend có thể stale/tampered; DB booking.totalPrice là chuẩn."),
            ("Webhook khác return URL?", "Return phụ thuộc browser; webhook là server-to-server và mới phù hợp xác nhận tự động."),
            ("Idempotency được bảo đảm ở đâu?", "Service state check + DB lock/constraint + payment event."),
            ("Reconciliation là gì?", "Đối chiếu booking/payment/ticket/refund và dữ liệu gateway để tìm trạng thái lệch."),
            ("Refund thật có thể mở rộng thế nào?", "RefundService gọi provider adapter, lưu provider refund ID và xử lý callback/retry."),
        ),
        checklist=(
            "Double-click initiate và gửi hai request song song.",
            "Gửi callback/webhook thành công lặp nhiều lần.",
            "Sửa signature, amount, transactionNo và transfer content để xác nhận bị từ chối.",
            "Thanh toán sau expiry/cancel để kiểm tra refund request.",
            "Đối chiếu payment, booking, seat, ticket, payment_events và refunds sau từng case.",
        ),
    ),
)


TOPICS += (
    Topic(
        source="Luong_Admin_Staff_Operation_CinemaBooking.docx",
        output="11_Admin_Staff_QR_Map_Operation_CinemaBooking.docx",
        title="Admin, staff, QR check-in, bản đồ và vận hành rạp",
        purpose="Theo dõi các use case quản trị từ master data đến suất chiếu, hoàn tiền và soát vé.",
        quick_steps=(
            "Phân biệt master data (phim/rạp/phòng/ghế) với transaction data (booking/payment/ticket).",
            "Đọc flow tạo suất chiếu: thành phố -> rạp -> phòng -> phim -> thời gian -> giá và kiểm tra trùng lịch.",
            "Đọc staff assignment trước các trang staff để hiểu vì sao dữ liệu được giới hạn theo rạp.",
            "Đọc QR check-in theo thứ tự validate; chỉ cập nhật USED ở bước cuối.",
            "Cuối cùng đọc dashboard projection, audit log, refund queue và map/location UX.",
        ),
        core_idea=(
            "Admin quản trị toàn hệ thống; staff vận hành các rạp được phân công. Mọi thao tác nhạy cảm phải có permission, "
            "data scope, business policy và audit. Xóa/hủy dữ liệu giao dịch không được làm mất lịch sử tài chính."
        ),
        flow="""ADMIN
  -> quản lý movie/cinema/room/seat/promotion/user/assignment
  -> quản lý showtime toàn hệ thống
  -> xem dashboard/payment/refund/reconciliation/audit

STAFF
  -> tải danh sách assigned cinemas
  -> tạo/sửa/hủy showtime trong scope
  -> xem booking/payment trong scope
  -> chọn cinema + showtime đang mở check-in
  -> quét QR
  -> backend xác minh QR, ACTIVE, booking SUCCESS,
     đúng rạp, đúng suất, đúng cửa sổ
  -> ticket USED + checkInTime + audit""",
        classes=(
            ("Movie/Cinema/Room/Seat services", "Master data CRUD.", "Validation, soft delete, cache eviction."),
            ("ShowtimeServiceImpl", "Create/update/cancel showtime.", "Room overlap, duration, status và staff scope."),
            ("StaffCinemaScopeService", "Giới hạn staff theo rạp.", "Assert scope cả read lẫn write."),
            ("BookingServiceImpl.checkInTicket", "QR validation/check-in.", "Thứ tự kiểm tra trước khi ACTIVE -> USED."),
            ("RefundServiceImpl", "Workflow hoàn tiền.", "Request/complete/fail và permission admin."),
            ("AnalyticsServiceImpl", "Dashboard read model.", "Projection doanh thu/vé/top phim."),
            ("AdminAuditLogInterceptor/Service", "Ghi thao tác vận hành.", "Actor/action/resource/result."),
            ("CinemaMapPage/location.ts", "Map và rạp gần tôi.", "Geolocation, Haversine, radius và fallback city."),
        ),
        code_sections=(
            (
                "D.1. Kiểm tra trùng lịch phòng",
                """newStart < existingEnd && newEnd > existingStart
  -> khoảng thời gian giao nhau
  -> từ chối tạo/cập nhật showtime""",
                "Cần loại chính showtime đang update và bỏ qua bản đã CANCELLED/deleted theo policy hiện tại.",
            ),
            (
                "D.2. Thứ tự check-in an toàn",
                """verify QR signature
find ticket
require ticket ACTIVE
require booking SUCCESS
require selected cinema matches
require selected showtime matches
require within check-in window
atomic ACTIVE -> USED
record checkInTime/operator""",
                "Không đánh dấu USED trước khi kiểm tra rạp/suất; nếu khách đến nhầm, vé vẫn dùng được tại đúng nơi.",
            ),
            (
                "D.3. Hủy suất chiếu",
                """showtime -> CANCELLED
active/pending bookings -> CANCELLED
tickets -> CANCELLED
held/booked seats -> AVAILABLE theo policy
successful payments -> create refund requests
send notification after commit""",
                "Hủy suất là workflow nhiều bảng, không phải một câu UPDATE showtimes đơn lẻ.",
            ),
        ),
        failures=(
            ("Staff thấy rạp không phụ trách", "List/filter endpoint không áp assignment.", "Scope query bằng assigned cinema IDs và test IDOR."),
            ("Tạo showtime báo end before start", "Date/time parse hoặc user chọn giờ kết thúc cùng ngày sai.", "Chuẩn hóa local datetime và validate backend."),
            ("Vé đúng nhưng quét nhầm rạp bị USED", "Update status trước scope validation.", "Đảm bảo toàn bộ validate chạy trước atomic transition."),
            ("Xóa phim làm mất lịch sử vé", "Hard delete/cascade sai policy.", "Soft delete master data đã được dùng."),
            ("Map có rạp dưới biển", "Tọa độ seed/admin nhập sai.", "Validate range, preview map và kiểm tra địa chỉ thực."),
        ),
        questions=(
            ("Staff có được hoàn tiền không?", "Mặc định không; staff theo dõi/hỗ trợ, admin hoặc quy trình có quyền mới complete refund."),
            ("Vì sao hủy suất không xóa suất?", "Cần giữ lịch sử booking, payment, ticket, audit và đối soát."),
            ("Trạng thái showtime có ONGOING không?", "Có; scheduler đồng bộ UPCOMING/ONGOING/ENDED và giữ CANCELLED."),
            ("Gần tôi được tính thế nào?", "Lấy geolocation rồi tính khoảng cách Haversine và lọc/sắp xếp theo bán kính."),
            ("Dashboard tránh N+1 thế nào?", "Dùng aggregate/projection/read model thay vì duyệt entity graph."),
        ),
        checklist=(
            "Tạo phòng/ghế và kiểm tra layout, unique seat cùng cache eviction.",
            "Tạo hai showtime trùng phòng/thời gian và xác nhận bị chặn.",
            "Đăng nhập staff ở một rạp, thử truy rạp ngoài scope.",
            "Quét vé đúng/sai rạp, sai suất, quá sớm, đã dùng và từ file/camera.",
            "Hủy suất có payment SUCCESS và quan sát refund request, ticket/booking status, email/audit.",
        ),
    ),
    Topic(
        source="Luong_Exception_Audit_SoftDelete_CinemaBooking.docx",
        output="12_Exception_Audit_SoftDelete_Monitoring_CinemaBooking.docx",
        title="Exception, audit, soft delete và monitoring",
        purpose="Hiểu cách hệ thống xử lý lỗi nhất quán, bảo toàn lịch sử và hỗ trợ điều tra vận hành.",
        quick_steps=(
            "Bắt đầu từ ApiResponse/error contract để biết frontend luôn nhận cấu trúc gì.",
            "Đọc GlobalExceptionHandler theo nhóm: validation, business, auth, access denied, database và unknown.",
            "Lần một lỗi business từ service exception đến toast/message trên frontend.",
            "Đọc AdminAuditLog và AuthAuditLog để phân biệt audit nghiệp vụ với log kỹ thuật.",
            "Đọc soft-delete policy và monitoring payment để hiểu hệ thống xử lý dữ liệu sau sự cố.",
        ),
        core_idea=(
            "Exception handler không che giấu bug; nó biến lỗi đã hiểu thành contract ổn định và giữ lỗi chưa hiểu trong log có correlation. "
            "Audit trả lời ai làm gì, còn application log/metric trả lời hệ thống đã vận hành ra sao."
        ),
        flow="""Controller @Valid / Service business rule / Repository DB
  -> exception cụ thể
  -> GlobalExceptionHandler
  -> HTTP status + code + message + timestamp + path
  -> axiosClient normalize
  -> page hiển thị thông báo tiếng Việt phù hợp

Admin/Auth/Payment action
  -> business transaction
  -> audit/event record có actor/resource/result
  -> admin page lọc và điều tra""",
        classes=(
            ("GlobalExceptionHandler", "Ánh xạ exception thành API error.", "Không trả stack trace/SQL ra client."),
            ("JwtAuthenticationEntryPoint", "Xử lý chưa xác thực.", "401 token thiếu/sai/hết hạn."),
            ("JwtAccessDeniedHandler", "Xử lý thiếu quyền.", "403 có login nhưng không đủ authority."),
            ("AdminAuditLogInterceptor", "Chặn/ghi thao tác admin.", "Actor, action, endpoint/resource và kết quả."),
            ("AuthAuditService", "Audit login/session security.", "Success/failure/IP/user agent an toàn."),
            ("PaymentEventServiceImpl", "Audit vòng đời payment callback.", "Payload đã sanitize, event type và status."),
            ("Soft-delete services", "Ẩn dữ liệu nhưng giữ quan hệ lịch sử.", "isDeleted và policy dữ liệu đã được tham chiếu."),
        ),
        code_sections=(
            (
                "D.1. Exception có ngữ nghĩa",
                """throw new SeatUnavailableException(showtimeId, seatIds);
throw new BookingExpiredException(bookingId);
throw new ForbiddenException("STAFF_CINEMA_SCOPE_DENIED");""",
                "Tên exception và error code giúp handler/client xử lý đúng, thay vì mọi lỗi đều trở thành 500 Uncategorized error.",
            ),
            (
                "D.2. Response lỗi ổn định",
                """{
  "code": 1027,
  "message": "Dữ liệu gửi lên không hợp lệ",
  "timestamp": "...",
  "path": "/api/v1/bookings/hold",
  "fieldErrors": { "seatIds": "..." }
}""",
                "Client không phụ thuộc message tiếng Anh từ framework; field error cho phép đặt thông báo cạnh input.",
            ),
            (
                "D.3. Soft delete",
                """UPDATE movies SET is_deleted = true WHERE id = ?;

SELECT ... FROM movies
WHERE is_deleted = false;""",
                "Không phải bảng nào cũng cần soft delete. Booking/payment/ticket thường giữ nguyên lịch sử và chuyển trạng thái, không xóa.",
            ),
        ),
        failures=(
            ("Client hiện Uncategorized error", "Exception chưa có handler hoặc handler lỗi khi serialize.", "Xem stack trace gốc và bổ sung mapping có code ổn định."),
            ("Staff nhận You do not have permission", "Frontend dùng message thô từ backend/security.", "Map 403 thành thông báo theo hành động và scope."),
            ("Audit làm request thất bại", "Audit nằm trong transaction chính và insert lỗi.", "Xác định policy; audit quan trọng nhưng không nên che lỗi nghiệp vụ gốc."),
            ("Dữ liệu đã xóa vẫn xuất hiện", "Query thiếu is_deleted=false hoặc cache stale.", "Specification/repository filter và evict cache."),
            ("Log lộ secret", "Ghi raw token/webhook payload/header.", "Mask credential và chỉ log ID/correlation cần thiết."),
        ),
        questions=(
            ("401 khác 403 thế nào?", "401 chưa xác thực/token không hợp lệ; 403 đã xác thực nhưng thiếu quyền/scope."),
            ("Vì sao không trả exception message thẳng?", "Có thể lộ SQL/internal detail và contract không ổn định."),
            ("Audit log khác application log?", "Audit là hồ sơ nghiệp vụ có cấu trúc; application log phục vụ chẩn đoán kỹ thuật."),
            ("Khi nào soft delete?", "Khi dữ liệu cần ẩn nhưng vẫn được lịch sử/quan hệ tham chiếu; cần policy và filter nhất quán."),
            ("Monitoring payment cần gì?", "Tỷ lệ callback lỗi, signature fail, pending quá lâu, late success, refund backlog và reconciliation mismatch."),
        ),
        checklist=(
            "Gửi request validation sai và kiểm tra fieldErrors tiếng Việt.",
            "Test 401, 403, 404, conflict/race, expired booking và gateway callback lỗi.",
            "Xác nhận response không có stack trace, SQL hoặc secret.",
            "Xóa mềm movie/cinema đã dùng và kiểm tra lịch sử booking vẫn đọc được.",
            "Lọc audit theo actor/action/resource và đối chiếu với thao tác vừa thực hiện.",
        ),
    ),
)


DEEP_DIVES: dict[str, tuple[tuple[str, tuple[str, ...]], ...]] = {
    "01_Kien_truc_va_ban_do_code_CinemaBooking.docx": (
        (
            "1. Đọc một request theo chiều dọc",
            (
                "Khi học một tính năng, đừng đọc toàn bộ package theo thứ tự alphabet. Hãy bắt đầu từ hành động trên React, tìm hàm trong domain API, ghi lại method và URL, rồi mở controller có mapping tương ứng. Từ controller lần xuống service interface, service implementation, repository và migration tạo bảng. Sau đó đi ngược từ response DTO, mapper, React Query hook đến component hiển thị. Cách đọc theo chiều dọc giúp thấy một use case hoàn chỉnh và tránh nhầm class cùng tên nhưng khác trách nhiệm.",
                "Ở mỗi tầng hãy trả lời một câu hỏi khác nhau: UI quyết định tương tác nào; API client gửi contract gì; Security xác thực ai; Controller kiểm tra hình dạng request và permission nào; Service giữ quy tắc nghiệp vụ và transaction nào; Repository bảo đảm truy vấn/lock nào; PostgreSQL giữ constraint/index nào. Nếu một class đang trả lời quá nhiều câu hỏi, đó là dấu hiệu coupling hoặc trách nhiệm bị đặt sai tầng.",
            ),
        ),
        (
            "2. Transaction boundary và side effect sau commit",
            (
                "`@Transactional` tạo ranh giới nguyên tử cho thay đổi booking, payment, ticket và seat status. Nếu một bước ném exception runtime, JPA đánh dấu rollback và dữ liệu không được để ở trạng thái nửa thành công. Tuy nhiên email, WebSocket và gọi cổng thanh toán là side effect bên ngoài database; phát chúng trước commit có thể khiến client nhận sự kiện dù transaction sau đó rollback.",
                "Vì vậy hệ thống dùng event/publisher sau commit ở các điểm cần thiết. Khi bảo vệ, hãy phân biệt consistency trong database với delivery ra ngoài. Database transaction bảo đảm tính đúng của state; event after-commit bảo đảm chỉ thông báo state đã lưu; retry/reconciliation xử lý trường hợp thông báo hoặc callback bên ngoài thất bại. Đây là tư duy quan trọng hơn việc chỉ nhớ tên annotation.",
            ),
        ),
        (
            "3. DTO, mapper và entity không phải cùng một thứ",
            (
                "Entity phản ánh mô hình persistence và quan hệ JPA, còn request/response DTO là contract công khai với frontend. Trả entity trực tiếp dễ làm lộ password, token hash, trường audit, gây vòng lặp JSON và vô tình kích hoạt lazy loading. Mapper là điểm kiểm soát dữ liệu được phép rời backend, đồng thời giúp API ổn định khi schema nội bộ thay đổi.",
                "Một response tốt chỉ mang dữ liệu màn hình cần và đã được gom đúng cách. Ví dụ ticket response cần phim, rạp, phòng, ghế và QR; frontend không nên gọi thêm năm endpoint để tự ghép. Ngược lại, danh sách quản trị nên dùng projection/page DTO nhẹ thay vì tải toàn bộ graph entity. Đây là cách kiến trúc vừa sạch vừa tránh N+1.",
            ),
        ),
        (
            "4. Cách nhận biết coupling và refactor an toàn",
            (
                "Coupling đáng chú ý xuất hiện khi controller tự sửa nhiều entity, component tự biết chi tiết token storage, hoặc service nghiệp vụ tự xây HTML/email/gateway payload. Refactor an toàn là giữ nguyên contract và test, sau đó tách một trách nhiệm có biên rõ: gateway adapter, mapper, publisher, scope service hoặc policy validator. Không nên tạo abstraction chỉ để đổi tên một dòng code.",
                "Khi source đã gần ngày bảo vệ, ưu tiên refactor có bằng chứng: query lặp, logic trạng thái trùng, dependency vòng, file quá lớn hoặc test khó viết. Mọi thay đổi phải được kiểm chứng bằng build/test và một kịch bản UI liên quan. Clean code trong product là code dễ thay đổi mà vẫn giữ hành vi, không phải số lượng class càng nhiều càng tốt.",
            ),
        ),
    ),
    "02_Auth_JWT_Session_GoogleLogin_CinemaBooking.docx": (
        (
            "1. Password login từ HTTP đến SecurityContext",
            (
                "Frontend gửi username/password qua HTTPS tới endpoint đăng nhập. Backend không tự so chuỗi password mà dùng `PasswordEncoder.matches` với BCrypt hash trong bảng users. Sau khi kiểm tra tài khoản active, chưa deleted và các điều kiện xác thực email phù hợp, service nạp role/permission để tạo access token. Password thô chỉ tồn tại trong request ngắn hạn và tuyệt đối không được ghi log.",
                "Các request kế tiếp mang `Authorization: Bearer <access-token>`. Spring Resource Server gọi `CustomJwtDecoder` để kiểm tra chữ ký, issuer/thời hạn và trạng thái session; converter biến claim scope/authority thành `GrantedAuthority`. Khi đó `@PreAuthorize` mới có dữ liệu để quyết định. JWT hợp lệ chỉ chứng minh danh tính và claim, chưa tự động chứng minh user được thao tác trên cinema cụ thể; data scope vẫn phải kiểm tra trong service.",
            ),
        ),
        (
            "2. Access token ngắn hạn và refresh token rotation",
            (
                "Access token nên sống ngắn vì nó thường là bearer token: ai cầm được có thể gọi API tới khi hết hạn. Refresh token sống dài hơn nhưng được quản lý như một session trong database. Khi refresh, backend khóa/tìm record hiện tại, kiểm tra chưa revoke/chưa hết hạn, thu hồi token cũ và phát cặp token mới. Rotation thu hẹp cửa sổ replay và cho phép phát hiện việc một refresh token cũ bị dùng lại.",
                "Frontend dùng cơ chế single-flight trong Axios interceptor: nhiều request cùng gặp 401 chỉ tạo một lời gọi refresh; các request còn lại chờ cùng Promise rồi retry đúng một lần. Nếu refresh thất bại, client xóa session và đưa người dùng về trạng thái chưa đăng nhập. Không làm single-flight dễ tạo một loạt refresh đồng thời, các token quay vòng chồng chéo và người dùng bị logout ngẫu nhiên.",
            ),
        ),
        (
            "3. Logout và giới hạn tự nhiên của JWT",
            (
                "Logout không thể xóa access token đã phát khỏi mọi nơi như session server truyền thống. Hệ thống thu hồi refresh session và có thể đưa `jti` access token vào invalidated-token store tới khi token hết hạn. Decoder kiểm tra blacklist để chặn token đã logout. Đổi lại, mỗi request có thêm một lần kiểm tra trạng thái; đây là lựa chọn bảo mật có chủ đích.",
                "Cần phân biệt logout một thiết bị với logout tất cả thiết bị. Nếu refresh session có định danh thiết bị/family, backend có thể revoke một session hoặc toàn bộ session của user. Các thao tác nhạy cảm như đổi password, khóa user hoặc phát hiện reuse nên revoke toàn bộ refresh token liên quan. Scheduler dọn record hết hạn để bảng không tăng vô hạn.",
            ),
        ),
        (
            "4. Google Login và trust boundary",
            (
                "Google Login không có nghĩa frontend gửi email bất kỳ rồi backend tin. Frontend nhận Google credential, backend phải xác minh chữ ký, audience/client ID, issuer và expiry bằng thư viện phù hợp. Sau khi token Google hợp lệ, backend mới ánh xạ email/sub về user nội bộ, cập nhật avatar/name theo policy và phát token của chính CinemaBooking.",
                "Email từ Google có thể dùng để đánh dấu verified vì nguồn đã xác minh, nhưng không nên ghi đè phone, dob hoặc dữ liệu người dùng đã chỉnh nếu Google không cung cấp hay policy không cho phép. Tài khoản password và Google cùng email cần quy tắc liên kết rõ để tránh tạo hai user hoặc account takeover. Client secret và key cấu hình phải ở environment, không commit vào repository.",
            ),
        ),
        (
            "5. Checklist đe dọa bảo mật cần biết",
            (
                "Các rủi ro chính gồm brute force login, credential stuffing, refresh token theft, token replay, XSS lấy token, CSRF nếu dùng cookie, IDOR khi truy cập tài nguyên người khác và rò secret qua log. Hệ thống giảm rủi ro bằng rate limit, BCrypt, token expiry/rotation/revocation, validation, ownership/scope checks và response lỗi không lộ chi tiết nội bộ.",
                "Không có cấu hình nào tự biến ứng dụng thành tuyệt đối an toàn. Khi triển khai thật cần HTTPS, cookie flags nếu token nằm trong cookie, CSP/XSS hardening, secret manager, audit cảnh báo và kế hoạch rotate key. Khi demo, hãy giải thích threat model và trade-off thay vì chỉ nói 'dùng JWT nên an toàn'.",
            ),
        ),
    ),
    "03_RBAC_Permission_StaffScope_CinemaBooking.docx": (
        (
            "1. Ba lớp kiểm soát: authentication, permission, data scope",
            (
                "Authentication trả lời người gọi là ai. Permission trả lời vai trò đó được thực hiện loại hành động nào, ví dụ `SHOWTIME_CREATE`. Data scope trả lời hành động đó được thực hiện trên bản ghi nào, ví dụ staff chỉ tạo suất chiếu trong cinema đã được gán. Bỏ lớp thứ ba sẽ tạo lỗ hổng IDOR dù `@PreAuthorize` vẫn chạy đúng.",
                "Frontend có thể ẩn nút để UX rõ ràng, nhưng đó không phải hàng rào bảo mật. Người dùng vẫn có thể gọi API bằng DevTools/Postman. Controller dùng permission để chặn coarse-grained access; service dùng `StaffCinemaScopeService` hoặc ownership query để chặn theo dữ liệu. Repository nên đưa cinema/showtime vào query khi cần để tránh tải dữ liệu rồi mới lọc trong bộ nhớ.",
            ),
        ),
        (
            "2. Authority đi từ database vào annotation",
            (
                "Quan hệ users_roles và roles_permissions là nguồn sự thật. Khi đăng nhập hoặc refresh, backend tổng hợp permission thành claim/authority. Spring chuyển chúng thành `GrantedAuthority`, sau đó biểu thức `hasAuthority(...)` trong `@PreAuthorize` so khớp. Tên permission là contract giữa initializer, security annotation và frontend capability map nên phải nhất quán tuyệt đối.",
                "Nếu thay permission trong database mà access token cũ còn sống, claim cũ có thể còn hiệu lực tới khi token hết hạn. Với thay đổi khẩn cấp như khóa user, decoder/session check phải có khả năng vô hiệu hóa ngay. Với thay đổi role thông thường, access token ngắn hạn giúp quyền hội tụ nhanh mà không phải query toàn bộ permissions mỗi request.",
            ),
        ),
        (
            "3. StaffCinema là assignment nghiệp vụ",
            (
                "StaffCinema không thay thế role STAFF; nó bổ sung phạm vi vận hành. Admin tạo hoặc thu hồi assignment. Các endpoint danh sách rạp, suất chiếu, booking, payment, ticket và dashboard dành cho staff phải dùng cùng một scope policy để tránh trang này lọc đúng nhưng trang khác lộ dữ liệu toàn hệ thống.",
                "Khi staff tạo suất chiếu, backend lấy room rồi suy ra cinema, sau đó xác nhận assignment trước khi ghi. Khi check-in QR, backend kiểm tra context cinema/showtime scanner đang chọn. Khi xem payment, query phải join booking-showtime-room-cinema và giới hạn assignment. Đây là lý do scope nên được đóng gói trong service/specification thay vì copy điều kiện ở từng controller.",
            ),
        ),
        (
            "4. Thiết kế quyền để không phình vô hạn",
            (
                "Permission nên mô tả capability ổn định như VIEW, CREATE, UPDATE, CANCEL, CHECKIN, REFUND. Không nên tạo permission riêng cho từng cinema hoặc từng bản ghi; phạm vi đó thuộc assignment/data scope. Cũng không nên gộp mọi thứ vào role ADMIN/STAFF check vì sẽ khó mở rộng vai trò quản lý vùng, kế toán hay CSKH.",
                "Ma trận quyền cần có owner, lý do và test. Mỗi endpoint mới phải trả lời: public hay authenticated; cần authority gì; cần ownership hay cinema scope gì; response có trường nhạy cảm không. Test tối thiểu gồm đúng quyền đúng scope, đúng quyền sai scope, thiếu quyền và admin override. Đây là cách chứng minh RBAC hoạt động chứ không chỉ có bảng role/permission.",
            ),
        ),
    ),
    "04_Booking_GiuGhe_RaceCondition_CinemaBooking.docx": (
        (
            "1. Vì sao kiểm tra rồi update thông thường vẫn bị tranh chấp",
            (
                "Nếu hai request cùng đọc ghế AVAILABLE trước khi request nào commit, cả hai đều có thể vượt qua câu lệnh `if` và tạo booking. Đây là race condition kiểu check-then-act. `synchronized` chỉ bảo vệ trong một JVM và không đủ khi deploy nhiều instance. Giải pháp phải dựa vào transaction và khóa/constraint ở database, nơi mọi instance cùng gặp nhau.",
                "Repository khóa các dòng seat_status mục tiêu bằng `PESSIMISTIC_WRITE` hoặc `SELECT ... FOR UPDATE`. Transaction thứ hai chờ transaction thứ nhất; khi được chạy tiếp nó đọc state mới là HOLD/BOOKED và thất bại có kiểm soát. Việc sort seatIds trước khi khóa giúp các transaction khóa cùng thứ tự, giảm nguy cơ deadlock khi người dùng chọn các bộ ghế giao nhau.",
            ),
        ),
        (
            "2. Invariant của hold, booking và payment",
            (
                "Invariant quan trọng: tại một thời điểm một seat_status của một showtime chỉ thuộc một trạng thái hợp lệ; HOLD phải có holdBy/holdUntil phù hợp; booking PENDING chỉ được tạo từ hold còn hiệu lực của chính user; booking SUCCESS phải có payment SUCCESS và các ghế BOOKED; ticket chỉ sinh cho booking thành công. Constraint unique `(seat_id, showtime_id)` chặn hai dòng trạng thái cho cùng ghế/suất.",
                "Service không tin giá từ client. Nó khóa và đọc ghế/showtime/promotion hiện tại, tính `price_at_booking`, tổng tiền và discount ở server. Booking detail giữ snapshot giá để lịch sử không đổi khi base price hoặc multiplier thay đổi. Sau commit, WebSocket mới phát trạng thái để các client khác cập nhật.",
            ),
        ),
        (
            "3. Hai lần bấm tạo booking và tính idempotent",
            (
                "Reload, double click hoặc retry mạng có thể gửi create booking nhiều lần. Backend phải tìm booking PENDING hiện có tương ứng user/showtime/hold hoặc dùng idempotency key, thay vì tạo vô hạn. Ở payment initiation, repository khóa booking/payment và tái sử dụng giao dịch pending hợp lệ. Frontend disable nút khi mutation đang chạy chỉ cải thiện UX, không thay thế bảo vệ server.",
                "Khi request timeout, client không được suy đoán là thất bại. Nó nên query lại booking/payment bằng identifier an toàn. Đây là nguyên tắc at-least-once delivery: cùng một ý định có thể tới server nhiều lần, nhưng state transition chỉ xảy ra một lần. Những thao tác tài chính càng cần idempotency và audit rõ hơn thao tác đọc thông thường.",
            ),
        ),
        (
            "4. Hết hạn và late payment",
            (
                "Hold hết hạn làm ghế AVAILABLE; booking PENDING hết payment_expires_at chuyển EXPIRED. Scheduler là cơ chế dọn nền, nhưng request đọc/ghi vẫn phải tự kiểm tra thời gian vì scheduler có thể chạy trễ. Không nên coi việc hàng đợi chưa quét là quyền tiếp tục thanh toán.",
                "Late payment là tình huống cổng báo thành công sau khi booking đã hết hạn và ghế có thể được người khác giữ. Backend không được tự chuyển ghế sang BOOKED bằng mọi giá. Nó ghi payment event, đưa vào trạng thái cần reconciliation/refund theo policy và không phá invariant ghế. Đây là phần chứng minh hệ thống xử lý thực tế chứ không chỉ happy path.",
            ),
        ),
        (
            "5. Cách test concurrency có ý nghĩa",
            (
                "Unit test mock repository không chứng minh row lock hoạt động. Cần integration test trên PostgreSQL thật hoặc Testcontainers: tạo một showtime/seat, dùng hai thread với transaction riêng cùng gọi hold, đồng bộ thời điểm bắt đầu rồi assert chỉ một request thành công và database có đúng một owner. Test tương tự cho create booking và payment callback trùng.",
                "Quan sát cả state cuối và side effect: không có hai booking hợp lệ cùng ghế, promotion used_count không tăng hai lần, ticket không trùng, email/WebSocket không phát sai. Nếu test chỉ assert HTTP status mà không kiểm tra invariant database, race condition vẫn có thể lọt qua.",
            ),
        ),
    ),
    "05_WebSocket_Realtime_SeatMap_CinemaBooking.docx": (
        (
            "1. STOMP/SockJS giải quyết phần nào",
            (
                "WebSocket tạo kết nối hai chiều lâu dài; STOMP bổ sung mô hình destination/subscription; SockJS là fallback transport khi môi trường không hỗ trợ WebSocket thuần. Client kết nối endpoint `/ws`, sau đó subscribe topic theo showtime. Server không gửi toàn bộ seat map mỗi giây mà chỉ phát event những ghế thay đổi, giảm payload và độ trễ.",
                "Realtime không thay thế REST. REST vẫn tải snapshot ban đầu và là nguồn khôi phục sau reconnect; WebSocket truyền delta để UI phản ứng nhanh. Nếu client nghi ngờ bỏ lỡ event, nó invalidate/refetch seat map. Thiết kế snapshot cộng delta giúp hệ thống bền hơn việc coi socket là nguồn dữ liệu duy nhất.",
            ),
        ),
        (
            "2. Chỉ publish sau khi transaction commit",
            (
                "Nếu publisher gửi BOOKED trước khi transaction database commit rồi transaction rollback, mọi client sẽ hiển thị sai. Vì vậy event phải gắn với after-commit hoặc được phát sau khi service hoàn tất thay đổi. Payload chứa showtimeId, seatIds, status, owner/expiry tối thiểu cần thiết và không lộ dữ liệu cá nhân.",
                "Trong hệ thống lớn hơn, outbox pattern giúp đảm bảo event không mất giữa DB commit và broker publish. Với kiến trúc monolith hiện tại, after-commit listener là trade-off hợp lý. Khi bảo vệ, nên nêu rõ giới hạn này và hướng mở rộng thay vì khẳng định WebSocket tự bảo đảm exactly-once.",
            ),
        ),
        (
            "3. Lifecycle phía React",
            (
                "Hook chỉ tạo một client cho showtime hiện tại, subscribe khi component mount/showtimeId đổi và unsubscribe/deactivate khi unmount. Callback tránh giữ closure cũ; reconnect có backoff; listener không được đăng ký lặp sau mỗi render. Hai camera/socket hoặc hai subscription thường bắt nguồn từ effect cleanup không đối xứng.",
                "Khi event đến, client cập nhật React Query cache bất biến hoặc invalidate query. Trạng thái HOLD của chính user có thể hiển thị khác HOLD của người khác, nhưng quyền đặt vẫn do backend quyết định. Đồng hồ hết hạn ở UI chỉ là chỉ báo; server time và database mới là nguồn quyết định cuối.",
            ),
        ),
        (
            "4. Test mất kết nối và thứ tự event",
            (
                "Kịch bản cần thử gồm hai tab cùng showtime, giữ/trả ghế, thanh toán thành công, scheduler nhả hold, đổi showtime, mất mạng rồi reconnect và unmount/remount. Sau reconnect phải lấy snapshot mới để bù event bị bỏ lỡ. Không nên chỉ kiểm tra một tab trong điều kiện mạng hoàn hảo.",
                "Event có thể tới gần nhau hoặc trễ. Payload nên đủ để client áp state mới idempotently; nếu cần thứ tự mạnh hơn có thể thêm version. Entity đã có version phục vụ optimistic locking, nhưng việc đưa version vào protocol phải được thiết kế đồng bộ. Hiện tại refetch khi nghi ngờ là cách đơn giản và an toàn.",
            ),
        ),
    ),
}


DEEP_DIVES.update({
    "06_Scheduler_GiuGhe_HetHan_TrangThaiSuatChieu_CinemaBooking.docx": (
        (
            "1. Fixed delay, fixed rate và ý nghĩa vận hành",
            (
                "Scheduler dùng fixed delay chờ một khoảng sau khi lần chạy trước kết thúc rồi mới chạy tiếp. Điều này tránh hai vòng quét chồng lên nhau khi database chậm. Fixed rate bám theo nhịp thời gian tuyệt đối và có thể dồn lần chạy nếu công việc kéo dài. Với cleanup hold/booking, fixed delay thường dễ kiểm soát tải hơn.",
                "Các biến `BOOKING_EXPIRED_HOLD_SCAN_DELAY_MS` và `BOOKING_EXPIRED_BOOKING_SCAN_DELAY_MS` điều chỉnh tần suất quét, còn thời gian nghiệp vụ như `BOOKING_SEAT_HOLD_MINUTES` quyết định hold hợp lệ bao lâu. Không được nhầm scan mỗi 30 giây với hold 30 giây. Request path vẫn kiểm tra timestamp để không phụ thuộc scheduler chạy đúng khoảnh khắc.",
            ),
        ),
        (
            "2. Idempotency của cleanup job",
            (
                "Một job tốt có thể chạy lại mà không làm hỏng state. Query chỉ chọn HOLD đã quá hạn hoặc booking PENDING đã quá payment_expires_at; update có điều kiện trạng thái hiện tại. Nếu hai lượt quét hoặc hai instance cùng chạm một bản ghi, chỉ lượt đầu chuyển trạng thái, lượt sau cập nhật 0 dòng và kết thúc an toàn.",
                "Khi nhả ghế, job phải xóa holdBy/holdUntil nhất quán, cập nhật booking/payment liên quan theo policy và publish AVAILABLE sau commit. Không dùng vòng lặp entity gây N+1 cho hàng nghìn bản ghi; ưu tiên bulk query hoặc phân batch có giới hạn để transaction không quá lớn.",
            ),
        ),
        (
            "3. Multi-instance và khóa job",
            (
                "Trong một instance, `@Scheduled` là đủ để kích hoạt. Khi scale nhiều backend, mỗi instance đều chạy scheduler. Nếu query/update idempotent và có điều kiện, dữ liệu vẫn đúng nhưng tạo tải thừa. Có thể dùng ShedLock, PostgreSQL advisory lock hoặc chuyển job sang queue/worker chuyên trách khi quy mô tăng.",
                "Không nên vội thêm distributed scheduler cho đồ án monolith nếu chưa cần. Hãy trình bày đường nâng cấp: hiện tại database condition bảo vệ correctness; production nhiều instance bổ sung leader/job lock để tối ưu efficiency. Đây là trade-off có căn cứ.",
            ),
        ),
        (
            "4. Thời gian, timezone và trạng thái suất chiếu",
            (
                "Thời gian được so sánh ở backend theo clock thống nhất. UI định dạng theo timezone người dùng/rạp và không hiển thị giây nếu nghiệp vụ không cần. Trạng thái hiệu lực của showtime được suy ra: trước start là UPCOMING, giữa start/end là ONGOING, sau end là ENDED; CANCELLED là trạng thái nghiệp vụ ưu tiên hơn trạng thái suy ra.",
                "Test nên đóng băng Clock hoặc truyền thời điểm cố định thay vì phụ thuộc `now()` thật. Các ca biên gồm đúng thời điểm hold_until, showtime qua nửa đêm, DST nếu triển khai quốc tế và callback tới sát expiry. Dữ liệu demo dùng NOW cần được hiểu là thay đổi theo lúc seed.",
            ),
        ),
        (
            "5. Quan sát scheduler",
            (
                "Job cần log có cấu trúc: bắt đầu/kết thúc, số bản ghi quét, số bản ghi đổi trạng thái, thời lượng và lỗi. Không log từng ghế ở INFO vì gây nhiễu. Metric nên theo dõi backlog hold hết hạn, booking pending quá lâu và thời gian job chạy để phát hiện cleanup bị kẹt.",
                "Exception của một batch không nên làm scheduler im lặng vĩnh viễn. Log phải có correlation/job name và nguyên nhân gốc; lần chạy sau vẫn tiếp tục. Với lỗi lặp, alert giúp vận hành biết trước khi người dùng thấy ghế bị giữ quá lâu.",
            ),
        ),
    ),
    "07_Thanh_toan_VNPay_SePay_Refund_Reconciliation_CinemaBooking.docx": (
        (
            "1. Payment state machine và nguồn sự thật",
            (
                "Booking và Payment là hai state machine liên quan nhưng không đồng nhất. Booking PENDING thể hiện đơn đang chờ trả tiền; Payment PENDING là một lần thử thanh toán. Thành công hợp lệ chuyển payment và booking sang SUCCESS, ghế sang BOOKED rồi sinh ticket. FAILED/EXPIRED giải phóng theo policy. REFUND_PENDING/REFUNDED mô tả tiền trả lại và không được giả bằng việc chỉ đổi booking thành CANCELLED.",
                "Frontend không quyết định thanh toán thành công dựa trên redirect hoặc nút 'đã chuyển khoản'. Nguồn sự thật là callback/webhook đã xác minh và state trong database. Trang kết quả poll/query backend để hiển thị trạng thái. Redirect chỉ là trải nghiệm điều hướng; IPN/webhook mới có thể tới ngay cả khi người dùng đóng trình duyệt.",
            ),
        ),
        (
            "2. PaymentGateway giúp mở rộng cổng",
            (
                "Service nghiệp vụ gọi interface gateway thay vì chứa điều kiện VNPay/SePay ở mọi nơi. Adapter VNPay biết cách ký tham số và tạo redirect URL; adapter SePay biết cách tạo VietQR/reference; adapter mới sau này triển khai cùng contract. Gateway registry/factory chọn adapter theo method cấu hình và từ chối cổng disabled.",
                "Phần chung vẫn nằm trong PaymentService: kiểm tra booking owner/state/expiry, khóa bản ghi, tạo hoặc tái sử dụng payment pending, ghi event/audit và gọi BookingService khi kết quả cuối. Tách đúng biên giúp thêm cổng không sao chép business rule và giảm nguy cơ mỗi cổng xử lý trạng thái khác nhau.",
            ),
        ),
        (
            "3. Idempotency, row lock và callback trùng",
            (
                "Gateway có thể retry webhook; người dùng có thể double click; proxy có thể gửi lại request. Payment repository dùng truy vấn khóa theo booking/transaction number, còn service kiểm tra trạng thái trước khi transition. Callback SUCCESS thứ hai phải trả acknowledgment phù hợp nhưng không tăng promotion, tạo ticket hoặc gửi mail lần nữa.",
                "Transaction number và event identifier cần unique khi provider bảo đảm duy nhất. Payload callback được lưu ở payment event/audit để truy vết nhưng phải che secret. So sánh amount, reference, booking, gateway và signature trước khi chấp nhận. Không dùng dữ liệu query/body chưa xác minh để cập nhật tiền.",
            ),
        ),
        (
            "4. Khác biệt VNPay và SePay",
            (
                "VNPay là redirect gateway: backend tạo URL ký HMAC, trình duyệt sang trang VNPay, sau đó return URL phục vụ UX và IPN phục vụ xác nhận server-to-server. SePay/VietQR tạo QR có số tiền và nội dung chuyển khoản duy nhất; webhook giao dịch vào được đối chiếu reference và amount. Vì QR cố định số tiền, promotion phải khóa sau khi payment QR đã tạo hoặc payment cũ phải được hủy và tạo QR mới.",
                "Ngrok chỉ cần khi provider ngoài Internet gọi callback vào máy local. Redirect từ trình duyệt tới localhost có thể vẫn hoạt động trên cùng máy, nhưng IPN/webhook không thể gọi localhost của bạn. Production dùng domain HTTPS ổn định, endpoint public đúng gateway, secret riêng theo môi trường và allowlist/rate limit phù hợp.",
            ),
        ),
        (
            "5. Refund, late payment và reconciliation",
            (
                "Hủy suất có booking đã trả tiền tạo refund request/event ở trạng thái chờ xử lý. Staff có thể xem nhưng không tự hoàn nếu policy chỉ cho admin/kế toán. Refund thật cần API provider hoặc quy trình chuyển khoản, idempotency key, amount không vượt số đã thu, audit actor và kết quả cuối. Không xóa payment gốc vì lịch sử tài chính phải bất biến.",
                "Reconciliation so sánh booking, payment, ticket và dữ liệu gateway để phát hiện lệch: tiền SUCCESS nhưng booking chưa SUCCESS, booking SUCCESS thiếu ticket, webhook signature fail, late success sau expiry hoặc refund tồn đọng. Job/report không tự sửa mù; nó phân loại, tự chữa ca an toàn và đưa ca rủi ro vào hàng đợi xử lý. Đây là lớp bảo vệ khi callback không hoàn hảo.",
            ),
        ),
        (
            "6. Các bài test thanh toán bắt buộc",
            (
                "Test initiation đúng owner, booking hết hạn, method disabled, reuse pending và double click. Test callback chữ ký đúng/sai, amount sai, transaction lạ, callback duplicate, failure sau success và late success. Test SePay webhook reference không khớp, số tiền thiếu/thừa theo policy. Test cancel showtime tạo đúng refund request.",
                "Integration test phải assert cả payment, booking, seat status, promotion count, ticket và event log trong cùng kịch bản. Mock unit test tốt cho signature/parser và policy, nhưng transaction/lock cần PostgreSQL integration. Khi demo không nên chuyển tiền thật nhiều lần; chuẩn bị sandbox và dữ liệu có thể tái lập.",
            ),
        ),
    ),
    "08_Database_Flyway_Query_Index_Cache_NPlusOne_CinemaBooking.docx": (
        (
            "1. Flyway là lịch sử tiến hóa schema",
            (
                "`database.sql` mô tả snapshot tạo mới thuận tiện để đọc hoặc bootstrap thủ công, còn `db/migration/V*.sql` là chuỗi thay đổi có thứ tự để môi trường thật nâng cấp an toàn. Flyway lưu checksum và version đã chạy; không sửa migration đã áp dụng ở môi trường dùng chung. Thay đổi mới phải tạo migration mới, có forward path và dữ liệu backfill nếu thêm cột NOT NULL.",
                "Mock data chỉ phục vụ demo/test thủ công và không phải migration production. Entity `ddl-auto=validate` giúp phát hiện Java/schema lệch mà không tự thay đổi database. Khi bảo vệ, hãy nêu quy trình: database trống -> Flyway V1..Vn -> ApplicationInitConfig seed RBAC tối thiểu -> tùy chọn chạy mock-data ở môi trường local.",
            ),
        ),
        (
            "2. Constraint là hàng rào cuối",
            (
                "Validation DTO cho thông báo đẹp nhưng có thể bị bỏ qua bởi job, script hoặc race. Database giữ invariant cuối bằng PK, FK, UNIQUE, NOT NULL và CHECK. Ví dụ unique seat/showtime, unique QR/secure token, status check, end_time > start_time và amount dương. Service vẫn kiểm tra trước để trả lỗi nghiệp vụ rõ; constraint xử lý trường hợp cạnh tranh còn lại.",
                "Xóa mềm cần query mặc định loại `is_deleted=true`, trong khi lịch sử booking vẫn tham chiếu được dữ liệu đã ẩn. FK delete action phải phản ánh nghiệp vụ: dữ liệu giao dịch thường không cascade theo master; bảng trung gian/chi tiết thuộc aggregate có thể cascade. Mọi hard delete cần xem graph tham chiếu và audit trước.",
            ),
        ),
        (
            "3. N+1 và cách nhận biết",
            (
                "N+1 xảy ra khi lấy một page N entity rồi truy cập lazy relation khiến Hibernate gửi thêm query cho từng dòng. Dấu hiệu là log có một SELECT danh sách và hàng chục SELECT gần giống. Sửa bằng projection DTO, `join fetch` có kiểm soát, `@EntityGraph` hoặc query batch; không chuyển mọi relation sang EAGER vì dễ tạo graph khổng lồ và Cartesian product.",
                "Pagination với collection fetch join cần cẩn thận vì count/duplicate row. Cách ổn định là page IDs/summary projection rồi tải detail theo batch, hoặc xây read model riêng cho dashboard. Mapper không nên âm thầm chạm relation chưa fetch trong vòng lặp. Test query count hoặc profiling log giúp ngăn N+1 quay lại.",
            ),
        ),
        (
            "4. Index và EXPLAIN ANALYZE",
            (
                "Index là cấu trúc phụ giúp PostgreSQL tìm dòng mà không quét toàn bảng, đổi lại tốn dung lượng và chi phí insert/update. Thứ tự cột composite dựa trên điều kiện lọc/sort thực tế, ví dụ showtime_id + status hoặc user_id + created_at. Index trên cột ít chọn lọc đơn lẻ như boolean có thể không hữu ích; partial index đôi khi phù hợp hơn.",
                "`EXPLAIN ANALYZE` chạy query và cho biết Seq Scan/Index Scan, estimated rows, actual rows, loops, sort và thời gian. Phải dùng dữ liệu đủ lớn và tham số gần production. Không kết luận chỉ từ việc 'có index'; planner có thể chọn seq scan vì bảng nhỏ. Sau tối ưu so sánh plan trước/sau và bảo đảm index không trùng lặp.",
            ),
        ),
        (
            "5. Cache và consistency",
            (
                "Caffeine phù hợp dữ liệu đọc nhiều, đổi ít như danh sách phim/rạp/phòng. Cache key phải bao gồm bộ lọc ảnh hưởng kết quả; TTL giới hạn độ cũ; mutation create/update/delete phải evict đúng cache. Không cache seat status, booking/payment đang thay đổi nhanh nếu chưa thiết kế consistency chặt.",
                "Cache local nằm riêng trên từng backend instance. Khi scale nhiều instance, eviction của instance A không tự tới B; khi đó cần TTL ngắn, event invalidation hoặc Redis. Cache không sửa query sai và không thay index. Thứ tự tối ưu đúng là đo query, sửa N+1/index/read model rồi mới cache phần ổn định.",
            ),
        ),
    ),
    "09_Frontend_React_Router_State_API_Realtime_UX_CinemaBooking.docx": (
        (
            "1. Phân loại state trước khi viết component",
            (
                "Server state như movies, cinemas, bookings và seat map thuộc React Query: có loading/error/stale/refetch/invalidation. Session/global UI state nhỏ thuộc Zustand/provider. State cục bộ như modal mở, input và tab thuộc component. Nhét tất cả vào một store làm stale data và side effect khó kiểm soát; gọi fetch thủ công khắp component làm trùng request.",
                "Query key phải chứa mọi tham số ảnh hưởng response như city, cinema, date, page và status. Mutation thành công invalidate hoặc cập nhật chính xác key liên quan. Dữ liệu realtime là delta trên server state, không nên tạo một bản seat map song song không đồng bộ với React Query.",
            ),
        ),
        (
            "2. Router, public route và protected route",
            (
                "Các trang xem phim/rạp/lịch chiếu là public; chọn ghế hoặc thanh toán chuyển sang yêu cầu login ở thời điểm cần bảo vệ. Protected route kiểm tra session đã hydrate, tránh redirect sớm trong lúc refresh token đang khôi phục. Admin/staff route kiểm tra capability để điều hướng UX, còn backend vẫn kiểm tra permission thật.",
                "Deep link như `/cinemas/:id` phải tự tải dữ liệu theo route param, không phụ thuộc user đã đi qua trang danh sách. Khi ID sai, hiển thị not-found có đường quay lại. Sau login nên trở lại intent trước đó nếu hợp lệ. Đây là những chi tiết làm app giống product hơn demo tuyến tính.",
            ),
        ),
        (
            "3. Axios client và contract lỗi",
            (
                "Một Axios instance tập trung base URL, Authorization, timeout, refresh single-flight và normalize error. Domain API chỉ mô tả endpoint/type, không tự lặp token logic. Frontend ánh xạ error code/message có kiểm soát sang tiếng Việt; không hiển thị raw exception/JDBC message. Validation field error gắn ngay input, lỗi nghiệp vụ hiển thị callout/toast phù hợp.",
                "Abort/cancel request khi component unmount hoặc filter đổi nhanh giúp tránh response cũ ghi đè. Nút mutation có pending state và kích thước ổn định để không rung. Retry tự động chỉ dùng cho read hoặc thao tác idempotent; không retry mù create booking/payment nếu backend chưa có idempotency.",
            ),
        ),
        (
            "4. Effect lifecycle và lỗi render vô hạn",
            (
                "`useEffect` chỉ dùng để đồng bộ với hệ thống bên ngoài. Dependency là object/function tạo mới mỗi render có thể khiến effect chạy lại và setState vô hạn. Dữ liệu suy ra nên tính trong render/useMemo thay vì effect setState; callback cần useCallback khi identity thực sự quan trọng; subscription/camera/socket phải cleanup đối xứng.",
                "StrictMode development có thể mount-cleanup-mount để phát hiện side effect không an toàn. Nếu xuất hiện hai camera hoặc hai socket, không nên tắt StrictMode ngay; hãy bảo đảm effect tạo đúng một resource và cleanup đúng resource đã tạo. Đây cũng là nguyên nhân của `removeChild` khi thư viện tự quản DOM không tương thích lifecycle.",
            ),
        ),
        (
            "5. Responsive, accessibility và cảm giác product",
            (
                "Mobile ưu tiên luồng tác vụ, touch target đủ lớn, không phụ thuộc hover và không để bảng tràn ngang vô nghĩa. Desktop có thể dùng grid nhiều cột; tablet cần breakpoint trung gian. Container scroll riêng phải có kích thước ổn định và dấu hiệu còn nội dung; drag carousel chỉ kích hoạt sau ngưỡng để click card vẫn hoạt động.",
                "Label gắn với input, focus visible, button có disabled/aria state, icon có accessible name và màu đạt tương phản. Skeleton giữ layout; empty state đưa hành động tiếp theo; error state cho retry. Typography dùng hierarchy vừa phải, tránh chữ quá đậm/to trong dashboard. Những điều này không đổi business flow nhưng giảm lỗi sử dụng thực tế.",
            ),
        ),
    ),
})


DEEP_DIVES.update({
    "10_Email_Verification_Reset_Ticket_Notification_CinemaBooking.docx": (
        (
            "1. Token email phải dùng một lần và không lưu thô",
            (
                "Verification/reset token được sinh bằng random đủ mạnh, gửi bản thô qua link email nhưng database chỉ lưu hash cùng expiry. Khi người dùng bấm link, backend hash token nhận được rồi so sánh. Nếu database bị lộ, attacker không thể dùng trực tiếp hash làm link. Sau thành công phải xóa hash/expiry hoặc đánh dấu consumed để token không dùng lại.",
                "Reset password không tiết lộ email có tồn tại hay không; endpoint có thể trả thông báo trung tính. Khi đổi password thành công, nên revoke refresh sessions để thiết bị cũ không tiếp tục truy cập. Verification email có thể resend nhưng rate limit và vô hiệu token cũ để tránh spam và nhiều link còn hiệu lực song song.",
            ),
        ),
        (
            "2. Gửi email và transaction",
            (
                "Không gửi SMTP ở giữa transaction tạo user/payment vì mạng chậm có thể giữ connection DB và email có thể được gửi dù transaction rollback. Cách hiện tại phù hợp là phát event sau commit rồi xử lý async. Nếu gửi thất bại, dữ liệu chính vẫn đúng và log/retry xử lý notification riêng.",
                "Production quy mô lớn có thể dùng outbox + queue để bảo đảm delivery. Với monolith, async executor cần giới hạn queue/thread và có handler exception. Không dùng common pool không kiểm soát. Log chỉ ghi recipient đã mask, template/event id và trạng thái; không log token thô hoặc toàn bộ nội dung nhạy cảm.",
            ),
        ),
        (
            "3. Ticket email và QR",
            (
                "Sau payment SUCCESS và transaction commit, hệ thống dựng email vé từ booking snapshot: phim, giờ, rạp, địa chỉ/city, phòng, ghế, tổng tiền và từng QR. Mỗi booking detail tương ứng một ticket/QR để check-in từng ghế. QR chứa opaque signed payload, không chứa thông tin cá nhân dễ đọc và không phụ thuộc ảnh QR để xác minh.",
                "Logo trong email cần URL public HTTPS hoặc inline CID; đường dẫn local/static React không hiển thị trong Gmail. HTML email dùng layout/table và CSS tương thích client email, có plain-text fallback, escape dữ liệu người dùng và giới hạn kích thước ảnh. Tải vé ở frontend và email phải cùng format thời gian/địa điểm để tránh mâu thuẫn.",
            ),
        ),
        (
            "4. Deliverability và vận hành",
            (
                "Mailtrap Sandbox chỉ bắt email test, không gửi tới inbox Gmail thật. Production cần provider/domain đã cấu hình SPF, DKIM, DMARC và sender hợp lệ. SMTP timeout, authentication và TLS lấy từ environment; `.env.example` chỉ có tên biến, không chứa credential thật.",
                "Theo dõi tỷ lệ gửi thành công, bounce, complaint và retry. Không retry vô hạn lỗi permanent như địa chỉ sai. Notification có idempotency key theo event/template/recipient để callback payment trùng không gửi nhiều vé. User vẫn phải xem vé trong app ngay cả khi email chậm hoặc thất bại.",
            ),
        ),
    ),
    "11_Admin_Staff_QR_Map_Operation_CinemaBooking.docx": (
        (
            "1. Admin portal là công cụ vận hành, không phải landing page",
            (
                "Màn hình quản trị ưu tiên scan, lọc, phân trang và thao tác lặp. Danh sách dùng summary DTO, filter ở server, sort ổn định và tổng số rõ. Form create/update có validation client để phản hồi nhanh nhưng backend vẫn là nguồn rule. Action nguy hiểm dùng confirm mô tả hậu quả và audit actor.",
                "Dashboard chỉ hiển thị chỉ số có định nghĩa nhất quán: doanh thu từ payment SUCCESS trong kỳ, vé bán từ ticket/booking detail hợp lệ, showtime theo scope. Chart dùng read model/projection thay vì tải entity rồi tính ở Java. Staff dashboard phải lọc cinema assignment, admin mới thấy toàn hệ thống.",
            ),
        ),
        (
            "2. Tạo suất chiếu và kiểm tra xung đột",
            (
                "Form staff chỉ hiển thị city/cinema/room được phân công; nhưng backend vẫn kiểm tra scope từ roomId. End time có thể được gợi ý từ duration phim cộng thời gian dọn phòng, nhưng phải sau start và không overlap suất khác trong cùng room. Client gửi ISO timestamp thống nhất, tránh parse locale AM/PM gây lỗi end trước start.",
                "Status UPCOMING/ONGOING/ENDED nên suy ra theo thời gian hoặc được job đồng bộ; CANCELLED là quyết định nghiệp vụ. Chỉ showtime còn mở bán mới xuất hiện ở client. Update/hủy phải khóa hoặc kiểm tra booking liên quan và dùng policy rõ, không hard delete lịch sử đã bán vé.",
            ),
        ),
        (
            "3. Hủy suất và hoàn tiền",
            (
                "`cancelShowtimeWithPolicy` kiểm tra quyền/scope, trạng thái thời gian và booking. Booking chưa trả được cancel/expire và nhả ghế; booking đã trả chuyển sang trạng thái hủy do suất, tạo refund pending, hủy ticket và thông báo khách. Thao tác chạy transaction để state không dở dang, còn gọi refund provider/email chạy theo event/job phù hợp.",
                "UI khách không còn nút chọn lại ghế cho suất bị hủy mà hiển thị đang xử lý/đã hoàn tiền. Admin/payment ops thấy queue hoàn tiền và lịch sử. Staff thông thường không tự xác nhận tiền đã hoàn nếu không có permission/quy trình kế toán. Mọi transition cần audit lý do, actor và timestamp.",
            ),
        ),
        (
            "4. QR check-in theo context",
            (
                "Scanner yêu cầu chọn city/cinema/showtime đang mở check-in trước khi quét. API gửi QR cùng cinemaId/showtimeId. Backend xác minh chữ ký QR, ticket ACTIVE, booking SUCCESS, đúng cinema, đúng showtime và đúng cửa sổ check-in rồi mới atomic update ACTIVE -> USED. Thứ tự này tránh vé nhầm rạp bị tiêu thụ.",
                "Nếu vé USED, trả giờ và nhân viên đã quét để xử lý tại quầy. Camera và upload file dùng cùng hàm decode/submit, có debounce chống quét liên tục và cleanup camera khi rời trang. Lỗi không đọc thấy QR là trạng thái bình thường trong stream camera, không nên hiện toast đỏ mỗi frame.",
            ),
        ),
        (
            "5. Map và dữ liệu tọa độ",
            (
                "Cinema lưu latitude/longitude; frontend tạo marker trực tiếp từ dữ liệu API. Nút Gần tôi xin geolocation có consent, tính khoảng cách Haversine và sắp xếp client hoặc server tùy quy mô. City gần nhất là gợi ý, người dùng vẫn đổi được. Các rạp thiếu tọa độ không được làm hỏng toàn bản đồ.",
                "Tọa độ mock phải nằm đúng vị trí thực, tránh dưới biển/sông. GeoJSON Hoàng Sa/Trường Sa là lớp hiển thị bản đồ và không ảnh hưởng marker/interactions. Chỉ đường mở URL map đã encode. Production nên có bước geocoding/preview marker khi admin tạo rạp để giảm nhập sai.",
            ),
        ),
    ),
    "12_Exception_Audit_SoftDelete_Monitoring_CinemaBooking.docx": (
        (
            "1. Taxonomy lỗi và HTTP status",
            (
                "Lỗi validation request trả 400 với fieldErrors; chưa đăng nhập/token lỗi trả 401; đã đăng nhập nhưng thiếu quyền/scope trả 403; resource không tồn tại trả 404; xung đột trạng thái như ghế vừa bị giữ trả 409; rate limit trả 429; lỗi bất ngờ trả 500 với message trung tính. Error code ổn định giúp frontend dịch thông báo mà không parse câu tiếng Anh.",
                "Business exception nên mang code/context cần thiết, không mang SQL hay stack trace ra response. `GlobalExceptionHandler` là adapter chuyển exception thành contract lỗi thống nhất gồm code, message, timestamp, path và fieldErrors khi có. Log server giữ correlation id và nguyên nhân gốc theo mức phù hợp.",
            ),
        ),
        (
            "2. 401 và 403 không được trộn",
            (
                "AuthenticationEntryPoint xử lý request chưa xác thực hoặc token không hợp lệ và trả 401. AccessDeniedHandler xử lý principal hợp lệ nhưng không đủ authority/scope và trả 403. Nếu mọi lỗi đều thành 'You do not have permission', người dùng không biết cần đăng nhập lại hay liên hệ admin.",
                "Frontend với 401 có thể thử refresh đúng một lần; nếu thất bại thì kết thúc session. Với 403 không refresh vì token mới không tạo thêm quyền; hiển thị thông báo phù hợp thao tác, ví dụ staff không có quyền xác nhận hoàn tiền. Đây là UX và bảo mật nhất quán.",
            ),
        ),
        (
            "3. Exception và rollback",
            (
                "Exception runtime trong `@Transactional` làm rollback mặc định. Nếu catch rồi nuốt exception trong service, transaction có thể commit state dở dang; nếu cần chuyển exception, hãy ném business exception tiếp và giữ cause trong log. Side effect ngoài DB phải sau commit hoặc có compensation.",
                "Optimistic lock, unique violation và deadlock cần ánh xạ cẩn thận. Conflict ghế là tình huống nghiệp vụ dự kiến nên trả 409 thân thiện; deadlock/transient DB có thể retry giới hạn ở lớp phù hợp; lỗi constraint bất ngờ phải log và sửa invariant, không biến tất cả thành 400.",
            ),
        ),
        (
            "4. Audit log khác application log",
            (
                "Application log phục vụ chẩn đoán kỹ thuật và có thể rotate. Audit log là hồ sơ ai làm gì với tài nguyên nào, trước/sau ra sao, kết quả và thời điểm; nó cần truy vấn trong portal và retention rõ. Hành động quan trọng gồm assignment staff, thay role, hủy suất, refund, check-in, thay trạng thái booking/payment và soft delete.",
                "Audit không lưu secret, password, token, full payment credential. Metadata cần sanitize và giới hạn kích thước. Với lỗi callback, payment event giữ provider reference, signature status, payload đã lọc và retry count. Correlation id nối HTTP log, audit và payment event để điều tra một giao dịch end-to-end.",
            ),
        ),
        (
            "5. Soft delete và khả năng khôi phục",
            (
                "Movie/cinema/room/seat đã tham gia lịch sử không nên hard delete tùy tiện. Soft delete ẩn khỏi danh sách active nhưng giữ FK cho booking/ticket báo cáo. Service phải chặn tạo mới dựa trên master đã deleted và query public/admin phải có semantics rõ. Unique key với soft delete cần policy tái sử dụng tên/code.",
                "Dữ liệu giao dịch tài chính thường append/state transition, không xóa. Hard delete chỉ hợp lý cho token hết hạn, dữ liệu tạm hoặc record chưa từng được tham chiếu theo policy. Backup/restore và retention vẫn cần ngay cả khi có soft delete; cột cờ không phải bản sao lưu.",
            ),
        ),
        (
            "6. Monitoring hướng theo triệu chứng người dùng",
            (
                "Theo dõi error rate theo endpoint/code, p95 latency, connection pool, scheduler backlog, hold conflict/rate limit, WebSocket connection, email failure, payment signature fail, pending quá lâu, refund backlog và reconciliation mismatch. Alert dựa trên tỷ lệ/ngưỡng trong cửa sổ để tránh báo động vì một lỗi đơn lẻ.",
                "Log phải đủ điều tra nhưng không ồn: INFO cho state transition quan trọng, WARN cho tình huống cần chú ý, ERROR cho thất bại ngoài dự kiến; camera không đọc QR mỗi frame không phải ERROR server. Dashboard monitoring khác dashboard doanh thu. Khi có incident, runbook nêu cách tìm correlation id, kiểm tra DB/event và hành động an toàn.",
            ),
        ),
    ),
    "13_Kiem_thu_Trien_khai_Bao_ve_CinemaBooking.docx": (
        (
            "1. Kim tự tháp kiểm thử áp dụng cho dự án",
            (
                "Unit test nhiều và nhanh cho policy, mapper, signature, price/promotion, token và state transition thuần. Slice test kiểm tra controller/security hoặc repository riêng. Integration test trên PostgreSQL xác minh Flyway, JPA mapping, lock, constraint và transaction. Một số smoke test API/manual UI bao phủ hành trình chính; không cần E2E trình duyệt nếu bạn đã quyết định không duy trì Playwright.",
                "Mục tiêu không phải 100% coverage mà là bảo vệ invariant có rủi ro cao. Mỗi bug từng gặp nên trở thành regression test ở tầng thấp nhất có thể tái hiện: giữ ghế đồng thời, callback duplicate, QR sai scope, refresh đồng thời, showtime overlap và filter staff scope.",
            ),
        ),
        (
            "2. Testcontainers và database thật trong test",
            (
                "H2 không mô phỏng đầy đủ PostgreSQL lock, JSONB, UUID, native query và constraint. Testcontainers khởi động PostgreSQL tạm trong Docker, chạy Flyway rồi test repository/service trên engine thật. Container/test database tách biệt hoàn toàn database dev và bị hủy sau test.",
                "Nếu máy không mở Docker, unit test vẫn chạy nhưng integration test Testcontainers sẽ skip/fail tùy cấu hình; điều này phải được ghi rõ. Trước bảo vệ nên chạy một lần với Docker và lưu kết quả. Không trỏ test vào database thật có dữ liệu cá nhân hoặc demo quan trọng.",
            ),
        ),
        (
            "3. Ma trận kịch bản trọng yếu",
            (
                "Auth: login sai/đúng, account disabled, access expiry, refresh rotation/reuse và logout. Booking: hai user cùng ghế, hold expiry, create duplicate, promotion snapshot. Payment: double click, callback duplicate/sai chữ ký/sai amount, late success, refund. QR: đúng, đã dùng, sai rạp, sai suất, ngoài giờ. RBAC: user/admin/staff đúng và sai cinema scope.",
                "Frontend/manual: public browse không login, deep link, mobile/tablet, dark mode, loading/empty/error, refresh trang thanh toán, WebSocket reconnect, camera permission denied và upload QR. Mỗi test case ghi precondition, action, expected API/state/UI và dữ liệu cleanup.",
            ),
        ),
        (
            "4. Build và cấu hình môi trường",
            (
                "Backend chạy `mvn test` rồi `mvn package`; frontend chạy lint/typecheck/build theo scripts trong package.json. `.env.example` phải liệt kê đủ biến nhưng không có secret. CI/production inject environment, Flyway chạy trước app hoặc khi startup theo policy. CORS, backend URL, redirect URL và webhook URL khác nhau giữa local/ngrok/production.",
                "Smoke test sau deploy gồm health/startup, login, public movies/showtimes, một booking sandbox, callback/webhook và WebSocket. Có rollback plan cho application và migration. Migration phá hủy dữ liệu cần tránh; thay đổi schema theo expand-migrate-contract khi production lớn.",
            ),
        ),
        (
            "5. Kịch bản bảo vệ 15 phút",
            (
                "Chuẩn bị dữ liệu cố định và ba tài khoản admin/staff/user. Mở bằng user: chọn phim -> suất -> hai tab tranh ghế -> áp mã -> thanh toán sandbox/QR -> vé. Mở staff: thấy rạp phụ trách -> scanner đúng/sai context. Mở admin: dashboard, tạo/hủy suất và refund/audit. Không phụ thuộc hoàn toàn Internet; có ảnh/log hoặc dữ liệu backup cho email/gateway.",
                "Trong phần hỏi đáp, trả lời theo cấu trúc vấn đề -> rủi ro -> giải pháp hiện tại -> bằng chứng code/test -> giới hạn -> hướng scale. Không nói 'chuẩn enterprise tuyệt đối'. Sự trung thực về trade-off và khả năng chỉ đúng class/query thường thuyết phục hơn liệt kê công nghệ.",
            ),
        ),
    ),
})


DEFENSE_LABS: dict[str, DefenseLab] = {
    "01_Kien_truc_va_ban_do_code_CinemaBooking.docx": DefenseLab(
        tldr="CinemaBooking.vn là monolith phân lớp: React gọi REST/WebSocket; Spring Security, controller và service thực thi quyền cùng nghiệp vụ; JPA/PostgreSQL giữ dữ liệu và invariant. Điểm thiết kế quan trọng là controller mỏng, transaction nằm ở service, DTO tách API khỏi entity và side effect chỉ phát sau commit.",
        sequence_steps=(
            "Client tạo hành động từ page/component, domain API chuẩn hóa URL, query/body và type response.",
            "Axios client gắn access token, xử lý timeout và chỉ refresh một lần nếu nhận 401.",
            "Spring Security xác minh JWT, dựng SecurityContext và kiểm tra permission trước khi vào controller.",
            "Controller bind request DTO, chạy Bean Validation và chuyển use case sang service.",
            "Service kiểm tra business rule/data scope trong transaction rồi gọi repository/mapper/publisher phù hợp.",
            "Repository phát SQL; PostgreSQL áp FK, UNIQUE, CHECK, index và row lock nếu use case cạnh tranh.",
            "Service map entity/projection thành response DTO; sau commit mới gửi WebSocket, email hoặc event ngoài DB.",
            "React Query lưu server state, render loading/success/error và invalidate/refetch sau mutation.",
        ),
        request_payload="Payload thay đổi theo use case nhưng luôn là DTO có kiểu, ví dụ HoldSeatsRequest {showtimeId, seatIds[]} hoặc PaymentInitiationRequest {bookingId, method}. Không gửi entity JPA hay tin totalPrice/role/userId nhạy cảm do client tự quyết định.",
        database_state="Mỗi transaction chỉ commit khi toàn bộ invariant đúng. Master data thường soft delete; booking/payment/ticket đổi trạng thái thay vì bị xóa; audit/event lưu dấu vết. Side effect ngoài database không được làm DB rơi vào trạng thái nửa thành công.",
        ide_order=(
            "Mở AppRouter và page phát sinh hành động để xác định route và UI state.",
            "Mở file `*Api.ts` rồi `axiosClient.ts` để lấy HTTP contract.",
            "Tìm controller bằng endpoint với `rg @RequestMapping`/`@PostMapping`.",
            "Đi vào service implementation, đánh dấu `@Transactional`, policy và state transition.",
            "Mở repository query/entity/migration liên quan; cuối cùng đọc mapper và test.",
        ),
        arrow_flow="React Page -> domainApi.ts -> axiosClient -> SecurityFilterChain\n-> Controller(@Valid/@PreAuthorize) -> Service(@Transactional)\n-> Repository -> PostgreSQL constraint/index/lock\n-> DTO -> React Query -> UI\nAfter commit -> WebSocket / Email / Audit / Gateway event",
        comparisons=(
            ("Kiến trúc", "Modular monolith phân lớp", "Microservices", "Monolith dễ transaction/debug/deploy cho quy mô đồ án; có thể tách payment/email khi tải và đội ngũ đủ lớn."),
            ("API model", "Request/Response DTO + mapper", "Trả entity trực tiếp", "DTO tránh lộ trường nhạy cảm, vòng lặp JSON và lazy loading ngoài ý muốn."),
            ("Business rule", "Service layer", "Đặt trong controller/repository", "Service là ranh giới use case, transaction và test; controller/repository giữ trách nhiệm kỹ thuật rõ."),
            ("Side effect", "After-commit event", "Gửi email/socket trong transaction", "Không phát thông báo cho dữ liệu sẽ rollback; đánh đổi là cần retry/outbox khi yêu cầu delivery mạnh hơn."),
        ),
        test_steps=(
            "Chọn một use case như giữ ghế và ghi lại request/response từ DevTools Network.",
            "Đặt breakpoint lần lượt controller, service và mapper; xác nhận Security chạy trước controller.",
            "Bật SQL profiling ở local, thực hiện một request và đối chiếu query với repository.",
            "Cố tình vi phạm validation, permission và DB constraint; kiểm tra rollback cùng error contract.",
            "Thực hiện mutation thành công và kiểm tra WebSocket/email chỉ xuất hiện sau dữ liệu commit.",
        ),
        edge_cases=(
            "Controller trả entity làm lộ password/token hash hoặc kích hoạt N+1.",
            "Service bắt exception rồi không ném lại khiến transaction commit dở dang.",
            "Frontend retry mutation không idempotent và tạo dữ liệu trùng.",
            "Event phát trước commit khiến UI hiển thị state chưa tồn tại.",
        ),
        senior_findings=(
            ("Service quá lớn", "Coupling cao, khó test và thay đổi", "Đo complexity/dependency; tách policy, mapper, gateway hoặc publisher theo biên nghiệp vụ thật."),
            ("DTO tải graph rộng", "N+1, JSON lớn, latency tăng", "Dùng projection/read model, query count profiling và pagination."),
            ("Side effect đồng bộ", "Giữ DB connection khi SMTP/gateway chậm", "After-commit async; production lớn dùng outbox/queue."),
            ("Cấu hình phân tán", "Sai URL/timeout/secret khi deploy", "Typed properties, `.env.example`, validation lúc startup và secret manager."),
        ),
        scale_summary="Nút nghẽn đầu tiên thường là PostgreSQL connection pool, query N+1, các transaction giữ lock lâu và SMTP/gateway đồng bộ. Scale theo thứ tự: đo p95/query count, thêm index/read model/cache dữ liệu tĩnh, rút ngắn transaction, tách side effect sang queue, rồi scale nhiều app instance; WebSocket cần broker relay và cache local cần cơ chế invalidation chung.",
        hard_questions=(
            ("Vì sao chưa dùng microservices nếu gọi là chuẩn product?", "Nêu đúng trade-off: domain và đội ngũ hiện tại cần transaction nhất quán, deploy/debug đơn giản; module/gateway/event đã tạo biên để tách khi có bằng chứng tải."),
            ("Làm sao chứng minh controller mỏng chứ không chỉ nói?", "Mở một endpoint, chỉ ra bind/validation/permission; business state transition và transaction nằm ở service, SQL ở repository, response ở mapper."),
            ("Nếu email gửi thất bại sau DB commit thì dữ liệu có sai không?", "Dữ liệu chính vẫn đúng; notification là side effect sau commit, được log/retry. Cần outbox nếu production yêu cầu không mất sự kiện."),
        ),
    ),
    "02_Auth_JWT_Session_GoogleLogin_CinemaBooking.docx": DefenseLab(
        tldr="Luồng Auth xác minh password hoặc Google credential, phát access token ngắn hạn và refresh session có rotation. Mỗi API được CustomJwtDecoder kiểm tra chữ ký, expiry và trạng thái thu hồi; logout/revoke xử lý phía server chứ không chỉ xóa token trên React.",
        sequence_steps=(
            "LoginForm gửi POST `/auth/token` với username và password; Google button gửi POST `/auth/google` với credential Google.",
            "AuthenticationController dùng `@Valid` rồi chuyển request tới AuthenticationService.",
            "Password login nạp user, kiểm tra BCrypt, active/deleted/email policy; Google login xác minh issuer, audience, signature và expiry.",
            "Service tổng hợp roles/permissions, tạo access JWT có jti/issuer/expiry và tạo refresh token ngẫu nhiên.",
            "Database lưu refresh token dưới dạng hash/session metadata; audit ghi login success/failure mà không ghi password/token thô.",
            "Client lưu session theo cơ chế hiện tại, gắn Bearer token cho request và hiển thị route theo capability.",
            "Khi 401 do access token hết hạn, Axios single-flight gọi POST `/auth/refresh`; backend rotate và revoke refresh cũ.",
            "POST `/auth/logout` hoặc xóa session thu hồi refresh session và invalidated jti; request sau bị decoder từ chối.",
        ),
        request_payload="Login: {username, password}. Google: {credential}. Refresh: refresh token/session credential theo contract hiện tại. Response gồm authenticated, accessToken, refreshToken/session data và thông tin user/permissions cần cho UI; password và token hash không bao giờ trả về.",
        database_state="users giữ BCrypt hash và trạng thái tài khoản; refresh_tokens giữ token hash, family/session, expiry/revoked; invalidated_token giữ jti tới khi access token hết hạn; auth_audit_log ghi hành động/kết quả/IP/user-agent đã lọc. Refresh rotation tạo record mới và vô hiệu record cũ trong transaction.",
        ide_order=(
            "AuthenticationController: `/auth/token`, `/google`, `/refresh`, `/logout`, `/sessions`.",
            "AuthenticationService: password login, Google mapping, issue/rotate/revoke token.",
            "JwtProperties, JwtTokenService và CustomJwtDecoder: key, issuer, claims, blacklist/session check.",
            "SecurityConfig và authority converter: public endpoint, 401/403, claim -> GrantedAuthority.",
            "RefreshToken/InvalidatedToken/AuthAuditLog repository và migration.",
            "Frontend Login/Register, AuthProvider/store và `axiosClient.ts` single-flight refresh.",
        ),
        arrow_flow="LoginPage -> POST /auth/token -> AuthenticationController\n-> AuthenticationService -> UserRepository + PasswordEncoder\n-> JwtTokenService + RefreshTokenRepository -> token response\n-> Auth store -> Axios Bearer\n401 -> one refresh Promise -> rotate token -> retry waiting requests",
        comparisons=(
            ("Xác thực API", "JWT access token", "Server session cookie", "JWT thuận lợi nhiều client/API và scale stateless phần lớn; session thu hồi tức thời đơn giản hơn nhưng cần shared store."),
            ("Mật khẩu", "BCrypt", "SHA-256 thuần", "BCrypt có salt và cost chống brute force; hash nhanh không phù hợp password."),
            ("Refresh", "Rotation + hash trong DB", "Refresh token cố định", "Rotation phát hiện reuse và giới hạn replay; phức tạp hơn vì quản lý session/family."),
            ("Social login", "Backend verify Google credential", "Tin email client gửi", "Xác minh server bảo vệ trust boundary; tin email trực tiếp dẫn tới account takeover."),
        ),
        test_steps=(
            "POST `/auth/token` bằng user1/đúng password; lưu access và gọi GET `/api/v1/users/me`.",
            "Dùng password sai nhiều lần để quan sát error code/rate limit và audit.",
            "Đợi hoặc dùng token expiry test, gửi nhiều request song song và xác nhận chỉ một `/auth/refresh`.",
            "Dùng lại refresh token cũ sau rotation; kỳ vọng bị từ chối và session xử lý theo policy.",
            "Logout, gọi API bằng access token cũ và kiểm tra 401/invalidated token.",
            "Đăng nhập Google đúng/sai audience; kiểm tra profile không ghi đè phone/dob ngoài policy.",
        ),
        edge_cases=(
            "Hai tab refresh đồng thời gây token reuse nếu không single-flight/phối hợp session.",
            "User bị block sau khi JWT đã phát nhưng token vẫn còn hạn.",
            "Refresh token bị đánh cắp và dùng trước chủ tài khoản.",
            "Clock skew làm token vừa phát bị coi là hết hạn.",
            "Google email trùng tài khoản password tạo duplicate hoặc liên kết sai.",
        ),
        senior_findings=(
            ("Query quyền mỗi request", "DB tăng tải ở endpoint nóng", "Claim access ngắn hạn + session/revocation check tối thiểu; cache có TTL nếu đo thấy cần."),
            ("Token trong nơi dễ bị JS đọc", "XSS đánh cắp bearer token", "CSP/sanitize; cân nhắc HttpOnly cookie với CSRF protection theo deployment."),
            ("Login brute force", "CPU BCrypt và tài khoản bị dò", "Rate limit theo IP+identity, backoff, metric và cảnh báo."),
            ("Bảng token tăng vô hạn", "Index lớn, lookup chậm", "Index hash/expiry, scheduler cleanup và retention."),
        ),
        scale_summary="10.000 người login đồng thời làm CPU BCrypt và connection pool nghẽn trước. Tách rate limiter/shared cache, giới hạn thread, scale auth instances, dùng index token hash, access token ngắn để giảm DB lookup, và tuyệt đối không giảm BCrypt cost tùy tiện. Refresh storm cần jitter/single-flight phía client và capacity/rate limit phía server.",
        hard_questions=(
            ("JWT stateless nhưng sao vẫn có bảng refresh/invalidated token?", "Access verification chủ yếu tự chứa claim; session/revocation là state có chủ đích để logout, rotation và khóa tài khoản an toàn. Đây là hybrid thực tế."),
            ("Kẻ xấu dùng refresh token cũ thì phát hiện thế nào?", "Token cũ đã revoked sau rotation; reuse cho thấy khả năng bị đánh cắp, backend từ chối và có thể revoke cả family, ghi audit."),
            ("Frontend ẩn nút admin có đủ bảo mật?", "Không. UI chỉ tối ưu UX; `@PreAuthorize` và data scope phía service mới là hàng rào chống gọi API trực tiếp/IDOR."),
        ),
    ),
    "03_RBAC_Permission_StaffScope_CinemaBooking.docx": DefenseLab(
        tldr="RBAC quyết định người dùng được làm loại hành động nào, còn staff-cinema assignment giới hạn hành động đó trên dữ liệu rạp nào. Backend kiểm tra cả permission và data scope; frontend chỉ dùng capability để trình bày đúng chức năng.",
        sequence_steps=(
            "Admin gán role/permission và assignment staff-cinema; database lưu quan hệ many-to-many có khóa duy nhất.",
            "Khi login/refresh, AuthenticationService nạp permission vào authority claim của access token.",
            "Request staff đi qua JWT decoder và converter để tạo SecurityContext.",
            "`@PreAuthorize` chặn coarse-grained permission như SHOWTIME_CREATE hoặc PAYMENT_VIEW_ALL.",
            "Service lấy cinema từ room/showtime/booking và gọi StaffCinemaScopeService để xác minh assignment.",
            "Repository query danh sách có điều kiện cinema scope, pagination và filter; không tải toàn bộ rồi lọc trên Java.",
            "Nếu thiếu permission trả 403; nếu có permission nhưng ngoài phạm vi cũng trả 403 có message nghiệp vụ phù hợp.",
        ),
        request_payload="Request mang resource id như roomId, showtimeId, cinemaId hoặc bookingId. Backend không tin cinemaId đứng riêng mà truy theo quan hệ thật từ resource để chống giả phạm vi. Response staff chỉ chứa dữ liệu thuộc assignment; admin có scope toàn hệ thống theo policy.",
        database_state="roles, permissions, users_roles, roles_permissions giữ capability; staff_cinemas giữ assignment. Unique(user_id, cinema_id) tránh gán trùng. Thay đổi quyền/assignment được audit; access token cũ hội tụ theo expiry hoặc bị revoke nếu cần hiệu lực ngay.",
        ide_order=(
            "ApplicationInitConfig và tài liệu permission matrix để biết nguồn role/permission.",
            "SecurityConfig + method có `@PreAuthorize` để thấy coarse-grained authorization.",
            "StaffCinemaScopeService và StaffCinemaRepository để thấy data scope.",
            "Một use case thật: ShowtimeService, PaymentService hoặc TicketService kiểm tra phạm vi.",
            "Repository pageable/specification có filter cinema; frontend admin/staff route và capability guard.",
        ),
        arrow_flow="JWT authorities -> @PreAuthorize(permission)\n-> load resource -> derive actual cinemaId\n-> StaffCinemaScopeService(userId, cinemaId)\n-> scoped repository query -> DTO\nFrontend guard = UX only; backend checks = security",
        comparisons=(
            ("Mô hình quyền", "Permission-based RBAC", "Chỉ kiểm tra role", "Permission chi tiết, mở rộng vai trò mới không sửa mọi endpoint; cần quản trị ma trận nhất quán."),
            ("Phạm vi staff", "Assignment theo cinema", "Permission riêng từng rạp", "Tách capability khỏi data scope, tránh permission bùng nổ."),
            ("Chặn quyền", "Method security + service scope", "Chỉ route guard frontend", "Backend chống API trực tiếp/IDOR; frontend không phải security boundary."),
            ("Policy nâng cao", "RBAC + scope", "ABAC/policy engine", "Đủ rõ cho domain hiện tại; ABAC phù hợp khi rule phụ thuộc nhiều thuộc tính/thời gian/địa bàn."),
        ),
        test_steps=(
            "Login ADMIN, STAFF được gán rạp A, STAFF chỉ rạp B và USER; lưu bốn token.",
            "Gọi endpoint tạo suất ở rạp A bằng từng token và ghi status mong đợi.",
            "Thay roomId sang phòng rạp B nhưng giữ cinemaId A; backend phải suy ra scope thật và từ chối.",
            "Kiểm tra danh sách booking/payment/dashboard của staff không chứa rạp ngoài assignment.",
            "Thu hồi assignment rồi refresh/login lại; kiểm tra quyền dữ liệu thay đổi và audit có record.",
        ),
        edge_cases=(
            "Có SHOWTIME_CREATE nhưng dùng roomId thuộc rạp không được gán.",
            "Endpoint detail lọc scope nhưng endpoint export/count lại lộ toàn hệ thống.",
            "Staff được bỏ assignment trong khi access token cũ còn authority.",
            "Admin tự xóa role/permission đang được hệ thống dùng.",
        ),
        senior_findings=(
            ("Scope check copy-paste", "Một endpoint mới dễ quên kiểm tra", "Đóng gói scope service/specification, test ma trận endpoint."),
            ("Lọc sau khi query", "Lộ dữ liệu và tốn memory", "Đưa scope vào SQL/projection/page query."),
            ("Authority claim quá lớn", "JWT/header phình khi permission tăng", "Giữ permission coarse ổn định; cân nhắc role/version/cache nếu quy mô rất lớn."),
            ("Thay quyền không tức thời", "Token cũ còn capability", "Access TTL ngắn, revoke session/user security version cho thay đổi khẩn cấp."),
        ),
        scale_summary="Với 10.000 staff/user request, query `exists assignment` lặp có thể nóng. Index staff_cinemas(user_id, cinema_id), đưa scope vào query chính hoặc cache assignment TTL ngắn có invalidation. Không cache permission vô hạn; thay đổi quyền cần hội tụ. Báo cáo staff dùng projection và index cinema/date thay vì lọc entity.",
        hard_questions=(
            ("RBAC đã đủ, tại sao cần StaffCinemaScopeService?", "RBAC chỉ nói staff được tạo suất; scope mới nói được tạo ở rạp nào. Thiếu scope là IDOR theo resource id."),
            ("Tại sao không tạo ROLE_STAFF_CINEMA_1, ROLE_STAFF_CINEMA_2?", "Role sẽ bùng nổ theo dữ liệu và khó bảo trì. Role/permission mô tả capability; assignment mô tả tập tài nguyên."),
            ("Nếu frontend không hiển thị rạp B thì staff có truy cập được không?", "Họ vẫn sửa request được, nên backend luôn suy ra cinema từ resource và kiểm tra assignment. UI không phải bằng chứng bảo mật."),
        ),
    ),
    "04_Booking_GiuGhe_RaceCondition_CinemaBooking.docx": DefenseLab(
        tldr="Luồng booking dùng row lock trong transaction để chỉ một người có thể giữ cùng ghế của cùng suất chiếu. Hold có thời hạn, booking/payment chuyển trạng thái theo state machine và mọi giá tiền được tính lại ở backend trước khi commit.",
        sequence_steps=(
            "Seat page tải GET `/api/v1/bookings/showtimes/{showtimeId}/seats` và subscribe topic seat map.",
            "Client gửi POST `/api/v1/bookings/hold` với showtimeId và mảng seatIds.",
            "BookingController xác thực user/validation/rate limit rồi gọi `BookingServiceImpl.holdSeats`.",
            "Service sort seatIds và repository `findForUpdateByShowtimeAndSeats` khóa các dòng seat_status bằng pessimistic write.",
            "Trong transaction, service kiểm tra đủ số ghế, đúng room, AVAILABLE hoặc hold hợp lệ của chính user; sau đó set HOLD, holdBy, holdUntil.",
            "Commit thành công rồi publisher phát delta HOLD; client khác đổi màu ghế. Nếu transaction rollback thì không phát state giả.",
            "Client gửi POST `/api/v1/bookings`; service khóa/kiểm tra hold còn hạn, tính giá server-side và tạo booking PENDING, booking_details cùng payment_expires_at.",
            "Thanh toán SUCCESS chuyển booking SUCCESS, seat_status BOOKED và tạo ticket; thất bại/hủy/hết hạn trả ghế AVAILABLE theo policy.",
        ),
        request_payload="HoldSeatsRequest: {showtimeId: UUID, seatIds: UUID[]}; seatIds bắt buộc là mảng, không phải chuỗi. CreateBookingRequest dùng showtime/seat hoặc hold context theo contract hiện tại; promotion được áp qua endpoint booking riêng. Client không được gửi trạng thái ghế hay tổng tiền làm nguồn sự thật.",
        database_state="seat_status: AVAILABLE -> HOLD(hold_by, hold_until) -> BOOKED hoặc AVAILABLE. bookings: PENDING(payment_expires_at) -> SUCCESS/FAILED/CANCELLED/EXPIRED/REFUND state theo policy. booking_details lưu seat_id và price_at_booking snapshot; unique(seat_id, showtime_id) cùng lock bảo vệ invariant.",
        ide_order=(
            "SeatSelection page/hook và booking API để thấy payload cùng UX rate-limit/conflict.",
            "BookingController `/showtimes/{id}/seats`, `/hold`, POST booking, cancel/promotion.",
            "BookingServiceImpl: `holdSeats`, `createBooking`, success/failure/cancel/expire.",
            "SeatStatusRepository: `findForUpdateByShowtimeAndSeats` và conditional bulk update.",
            "SeatStatus/Booking/BookingDetail entity + Flyway constraints/indexes.",
            "SeatStatusPublisher/WebSocket hook và HoldExpireScheduler.",
            "Concurrency/integration tests trên PostgreSQL.",
        ),
        arrow_flow="SeatPage -> POST /bookings/hold -> row lock seat_status\n-> validate state/owner/expiry -> HOLD + commit -> WebSocket\n-> POST /bookings -> re-lock + price snapshot -> PENDING\n-> payment SUCCESS -> BOOKED + tickets\n-> fail/cancel/expire -> AVAILABLE",
        comparisons=(
            ("Khóa ghế", "Pessimistic row lock", "Optimistic version retry", "Tranh chấp ghế cần một người thắng rõ; row lock đơn giản và đúng. Optimistic phù hợp xung đột hiếm nhưng retry UX phức tạp."),
            ("Khóa JVM", "Database lock", "`synchronized`", "DB lock hoạt động qua nhiều instance; synchronized chỉ trong một process."),
            ("Giữ ghế", "seat_status hold_by/hold_until", "Bảng seat_hold riêng", "Mô hình hiện tại gọn và đủ; bảng riêng hữu ích cho lịch sử/nhiều hold attempt ở quy mô lớn."),
            ("Giá", "Server snapshot", "Tin total từ client", "Chống sửa giá và giữ lịch sử đúng khi bảng giá thay đổi."),
        ),
        test_steps=(
            "Lấy showtimeId và hai seatId AVAILABLE; login user1 và user2.",
            "Gửi cùng lúc hai POST `/bookings/hold` cho cùng seat bằng runner/thread; chỉ một request thành công.",
            "Kiểm tra seat_status có đúng một hold_by, hold_until và WebSocket hai tab đồng bộ.",
            "User thắng tạo booking; user thua thử tạo booking trực tiếp và phải bị từ chối.",
            "Đợi expiry hoặc chỉnh timeout test; xác nhận scheduler/request path trả ghế và UI cập nhật không cần refresh.",
            "Double click tạo booking/thanh toán; kiểm tra không có hai booking/ticket/promotion count.",
        ),
        edge_cases=(
            "seatIds gửi dạng chuỗi thay vì mảng gây 400 Request body is invalid.",
            "Mảng chứa seat khác room/showtime hoặc trùng seatId.",
            "Hold vừa hết hạn đúng lúc create booking chạy.",
            "Hai bộ ghế giao nhau bị deadlock nếu thứ tự khóa khác nhau.",
            "Payment thành công muộn sau khi ghế đã được người khác giữ.",
        ),
        senior_findings=(
            ("Transaction giữ lock lâu", "Throughput giảm, timeout/deadlock", "Chỉ query/validate/update DB trong lock; không gọi email/gateway; sort ID và đo lock wait."),
            ("Quét expiry toàn bảng", "Scheduler gây IO cao", "Index status/hold_until, batch update và metric backlog."),
            ("Seat map polling", "Read spike ở suất hot", "WebSocket delta + snapshot/refetch có kiểm soát; cache chỉ dữ liệu tĩnh."),
            ("Tạo booking trùng", "Double click/retry", "Unique/idempotency key hoặc reuse PENDING dưới row lock."),
        ),
        scale_summary="10.000 user tranh suất hot tạo hotspot trên cùng các dòng seat_status. Row lock vẫn bảo vệ đúng nhưng latency tăng; cần giới hạn hold/rate limit, transaction cực ngắn, index đúng, connection pool cân bằng và phân vùng tải theo showtime. Có thể đưa admission queue cho sự kiện cực lớn, nhưng không bỏ DB invariant. WebSocket broadcast dùng broker relay khi nhiều instance.",
        hard_questions=(
            ("Pessimistic lock có làm hệ thống chậm không?", "Có chi phí chờ khi tranh chấp, nhưng ghế là tài nguyên contention cao cần correctness. Giảm bằng lock đúng dòng, thứ tự nhất quán và transaction ngắn; đo lock wait."),
            ("Nếu backend chạy hai instance thì synchronized có đủ?", "Không. Database row lock/constraint là điểm đồng bộ chung nên thiết kế hiện tại không phụ thuộc một JVM."),
            ("WebSocket có ngăn hai người mua cùng ghế không?", "Không. WebSocket chỉ cập nhật UX; transaction + lock + constraint mới quyết định người thắng."),
        ),
    ),
}


DEFENSE_LABS.update({
    "05_WebSocket_Realtime_SeatMap_CinemaBooking.docx": DefenseLab(
        tldr="WebSocket/STOMP đẩy thay đổi ghế theo showtime ngay sau khi transaction commit, giúp mọi tab thấy HOLD, BOOKED hoặc AVAILABLE gần như tức thời. REST vẫn tải snapshot ban đầu và là nguồn phục hồi sau reconnect; socket chỉ truyền delta chứ không quyết định quyền sở hữu ghế.",
        sequence_steps=(
            "Seat page gọi REST lấy snapshot seat map của showtime và render trạng thái ban đầu.",
            "Hook tạo STOMP client kết nối endpoint `/ws`, sau đó subscribe `/topic/seatmap/{showtimeId}`.",
            "Một user gọi hold/cancel/payment; BookingService cập nhật database trong transaction có lock.",
            "Sau commit, SeatStatusPublisher tạo SeatStatusEvent gồm showtimeId, seatIds, status và metadata expiry cần thiết.",
            "SimpMessagingTemplate phát event tới topic đúng showtime, không broadcast mọi rạp.",
            "STOMP client nhận delta, cập nhật/invalidate React Query cache và UI đổi màu mà không reload.",
            "Khi mất mạng, client reconnect có backoff rồi refetch snapshot để bù event bị bỏ lỡ.",
            "Khi đổi showtime/unmount, hook unsubscribe và deactivate đúng client để không rò socket/listener.",
        ),
        request_payload="Client không gửi payload nghiệp vụ qua socket trong luồng này; mutation vẫn qua REST có JWT. Event server -> client có dạng {showtimeId, seatIds[], status, holdBy?, holdUntil?, occurredAt?}; chỉ chứa dữ liệu tối thiểu, không chứa user profile/token.",
        database_state="WebSocket không lưu state riêng. Nguồn sự thật vẫn là seat_status/booking/payment đã commit. HOLD/BOOKED/AVAILABLE được ghi trước; event chỉ phản ánh kết quả. Nếu publish thất bại, client khôi phục bằng REST refetch; production yêu cầu delivery mạnh có thể dùng outbox/broker.",
        ide_order=(
            "WebSocketConfig: endpoint, broker prefix và allowed origin.",
            "SeatStatusEvent: contract delta giữa backend và frontend.",
            "SeatStatusPublisher: destination và publish single/bulk.",
            "BookingServiceImpl/HoldExpireScheduler: tìm điểm gọi publisher sau state transition.",
            "Frontend seat-map WebSocket hook: connect, subscribe, reconnect, cleanup.",
            "SeatSelection page và React Query cache update/refetch.",
        ),
        arrow_flow="REST snapshot -> render seat map -> STOMP connect /ws\n-> subscribe /topic/seatmap/{showtimeId}\nDB transaction commit -> SeatStatusPublisher -> broker topic\n-> React hook -> patch/invalidate cache -> UI\nReconnect -> REST snapshot reconciliation",
        comparisons=(
            ("Realtime transport", "WebSocket + STOMP", "Short/long polling", "Socket giảm request lặp và latency; polling đơn giản hơn nhưng tốn tải, cập nhật chậm."),
            ("Fallback", "SockJS", "WebSocket thuần", "SockJS hỗ trợ môi trường hạn chế; thêm dependency/protocol overhead."),
            ("Dữ liệu", "REST snapshot + socket delta", "Chỉ WebSocket", "Có đường phục hồi khi reconnect/mất event và dễ cache/read model."),
            ("Broker", "Simple broker trong app", "RabbitMQ broker relay", "Phù hợp một instance/demo; relay cần khi scale nhiều instance/kết nối lớn."),
        ),
        test_steps=(
            "Mở cùng một showtime ở hai trình duyệt/tài khoản và quan sát một kết nối/subscription mỗi tab.",
            "User A giữ ghế; user B phải thấy HOLD không refresh và vẫn không giữ được qua API.",
            "Thanh toán thành công; cả hai tab nhận BOOKED. Hủy/thất bại/expiry phải nhận AVAILABLE.",
            "Tắt mạng của B, thay đổi ghế ở A, bật mạng và xác nhận B reconnect rồi refetch đúng snapshot.",
            "Đổi showtime liên tục; kiểm tra tab cũ không còn nhận event và số socket không tăng vô hạn.",
        ),
        edge_cases=(
            "Event tới trước khi DB commit hoặc transaction rollback.",
            "Subscribe trùng do useEffect chạy lại/StrictMode tạo hai kết nối.",
            "Event bị mất trong lúc reconnect làm cache lệch.",
            "Event của showtime A cập nhật nhầm seat map B.",
            "Client nhận event cũ sau event mới do mạng/reconnect.",
        ),
        senior_findings=(
            ("Một socket mỗi render", "Memory/socket leak và event nhân đôi", "Ổn định dependency, cleanup đối xứng, quan sát WebSocket stats."),
            ("Broadcast payload lớn", "CPU/network tăng theo fan-out", "Topic theo showtime, bulk delta nhỏ, không gửi full seat map."),
            ("Simple broker nhiều instance", "Client ở instance khác không nhận event", "Broker relay/shared messaging hoặc sticky + cross-node pub/sub."),
            ("Không version/reconcile", "Event trễ làm cache sai", "Refetch sau reconnect; cân nhắc event version/sequence khi cần."),
        ),
        scale_summary="10.000 kết nối đồng thời làm file descriptor, heartbeat, outbound executor và fan-out thành điểm nghẽn. Cần reverse proxy hỗ trợ upgrade, giới hạn message, tune thread/heartbeat, broker relay như RabbitMQ, scale connection nodes và metric active sessions/queue lag. DB lock vẫn quyết định ghế; socket chỉ giảm polling và phát delta.",
        hard_questions=(
            ("Nếu WebSocket mất event thì ghế có bị bán trùng không?", "Không, DB lock/constraint bảo vệ correctness. UI có thể tạm stale; reconnect/refetch snapshot khôi phục."),
            ("Tại sao không gửi toàn bộ seat map sau mỗi thay đổi?", "Delta nhỏ giảm bandwidth/fan-out; snapshot REST chỉ dùng lúc đầu hoặc reconciliation."),
            ("Simple broker có chạy khi deploy nhiều instance?", "Không bảo đảm cross-instance. Hiện tại phù hợp monolith một node; hướng production là broker relay/shared pub-sub và observability."),
        ),
    ),
    "06_Scheduler_GiuGhe_HetHan_TrangThaiSuatChieu_CinemaBooking.docx": DefenseLab(
        tldr="Scheduler định kỳ giải phóng hold/booking quá hạn, dọn token và đồng bộ trạng thái suất chiếu. Job được thiết kế idempotent bằng điều kiện trạng thái/thời gian; request path vẫn tự kiểm tra expiry nên correctness không phụ thuộc job chạy đúng từng giây.",
        sequence_steps=(
            "Spring khởi tạo bean scheduler và đọc delay/timeout từ typed configuration/environment.",
            "Theo fixed delay, HoldExpireScheduler truy vấn seat_status HOLD có hold_until <= now theo batch/index.",
            "Job update có điều kiện HOLD -> AVAILABLE, xóa hold_by/hold_until trong transaction.",
            "Sau commit, publisher gửi AVAILABLE theo showtime để UI cập nhật realtime.",
            "Booking expiry job tìm booking PENDING có payment_expires_at <= now, khóa/chuyển EXPIRED và xử lý payment/ghế liên quan idempotently.",
            "Showtime status job suy ra UPCOMING/ONGOING/ENDED theo start/end, không ghi đè CANCELLED.",
            "Token cleanup xóa refresh/invalidated/verification record đã quá retention.",
            "Log/metric ghi số quét, số cập nhật, thời lượng và lỗi; vòng chạy sau tiếp tục nếu một lần thất bại.",
        ),
        request_payload="Scheduler không nhận HTTP payload. Input là Clock/now, các config như BOOKING_SEAT_HOLD_MINUTES và scan delay, cùng tập bản ghi thỏa status + expiry. Output là số record chuyển trạng thái và event realtime/audit phù hợp.",
        database_state="seat_status HOLD quá hạn -> AVAILABLE; bookings PENDING quá hạn -> EXPIRED; payment pending liên quan được đánh dấu theo policy; showtimes chuyển UPCOMING -> ONGOING -> ENDED trừ CANCELLED; token hết hạn bị xóa. Conditional update giúp chạy lặp không chuyển lại state đã xử lý.",
        ide_order=(
            "BookingProperties/application.yaml/.env.example để phân biệt timeout và scan delay.",
            "HoldExpireScheduler và các method `@Scheduled`.",
            "BookingServiceImpl.expirePendingBooking/release logic.",
            "Repository query theo status + timestamp và index Flyway.",
            "SeatStatusPublisher sau commit và frontend countdown/realtime.",
            "Scheduler/integration tests có Clock cố định.",
        ),
        arrow_flow="@Scheduled(fixedDelayString) -> query expired batch\n-> conditional update under transaction\n-> commit -> publish AVAILABLE / audit\nRequest path also checks expiry\n=> scheduler latency affects cleanup speed, not correctness",
        comparisons=(
            ("Nhịp chạy", "Fixed delay", "Fixed rate", "Không chồng job khi lần trước chạy lâu; thời điểm thực tế có thể trễ một khoảng."),
            ("Cleanup", "DB conditional batch", "Load entity rồi loop", "Ít query/memory và idempotent; bulk update cần chủ động event/audit."),
            ("Nhiều instance", "Idempotent DB condition", "Distributed lock", "Correctness vẫn giữ; distributed lock tối ưu tải khi scale."),
            ("Thời gian test", "Inject/fixed Clock", "Gọi `now()` khắp code", "Test deterministic và kiểm tra ca biên chính xác."),
        ),
        test_steps=(
            "Đặt hold timeout ngắn ở môi trường test và scan delay phù hợp; restart backend để config có hiệu lực.",
            "Giữ một ghế, kiểm tra hold_by/hold_until và countdown frontend.",
            "Đợi quá hạn; xác nhận seat AVAILABLE, booking EXPIRED nếu có và WebSocket đổi UI không refresh.",
            "Chạy method/job hai lần trên cùng dữ liệu; lần hai không tạo transition/event trùng.",
            "Tạo showtime trước/trong/sau khoảng thời gian và kiểm tra status, đặc biệt suất qua nửa đêm và CANCELLED.",
        ),
        edge_cases=(
            "Scheduler chạy trễ trong khi user bấm thanh toán đúng thời điểm hết hạn.",
            "Hai instance cùng quét một batch.",
            "Một batch quá lớn giữ transaction/lock lâu.",
            "Timezone/Clock lệch giữa app và database.",
            "Job lỗi một record làm bỏ cả batch hoặc ngừng các lần sau.",
        ),
        senior_findings=(
            ("Full table scan", "IO tăng theo dữ liệu", "Index partial/composite status+expiry, EXPLAIN ANALYZE và batch size."),
            ("Transaction batch quá lớn", "Lock/bloat/rollback tốn kém", "Chunk theo ID/time, commit từng batch có metric."),
            ("Nhiều node chạy cùng job", "Tải duplicate", "ShedLock/advisory lock hoặc worker riêng khi scale."),
            ("Log từng bản ghi", "Log volume/IO lớn", "Aggregate count, sample error và structured metrics."),
        ),
        scale_summary="10.000 hold hết hạn gần nhau có thể tạo burst update và WebSocket fan-out. Dùng index status/hold_until, batch giới hạn, bulk conditional update, group event theo showtime và backpressure. Tách scheduler worker/distributed lock khi nhiều app node; theo dõi backlog thay vì giảm delay mù làm DB bị quét liên tục.",
        hard_questions=(
            ("Scheduler chạy mỗi 30 giây thì hold có thể dài hơn timeout không?", "Cleanup vật lý có thể trễ tối đa gần scan delay, nhưng service kiểm tra hold_until nên người dùng không có quyền thanh toán sau expiry."),
            ("Hai backend cùng chạy scheduler có làm trả ghế hai lần?", "Conditional state update/idempotency giữ correctness; có thể tốn tải. Production nhiều node bổ sung distributed lock."),
            ("Tại sao không dùng Redis TTL cho hold?", "Redis TTL nhanh nhưng vẫn cần DB invariant/khôi phục và xử lý event; mô hình DB hiện tại đơn giản, nhất quán. Redis là hướng scale sau khi đo hotspot."),
        ),
    ),
    "07_Thanh_toan_VNPay_SePay_Refund_Reconciliation_CinemaBooking.docx": DefenseLab(
        tldr="PaymentService điều phối state machine chung, còn PaymentGateway adapter xử lý chi tiết VNPay redirect hoặc SePay/VietQR webhook. Backend chỉ công nhận tiền sau callback đã xác minh, xử lý idempotent dưới row lock rồi cập nhật booking, ghế, ticket, refund và audit/reconciliation.",
        sequence_steps=(
            "Payment page gửi POST `/api/v1/payments/initiate` với bookingId và method; backend kiểm tra owner, PENDING và chưa expiry.",
            "PaymentService khóa booking/payment, tái sử dụng pending hợp lệ hoặc tạo payment PENDING với transaction number duy nhất.",
            "Gateway registry chọn VNPay hoặc SePay theo method/config. VNPay tạo signed redirect URL; SePay tạo QR/reference cố định amount.",
            "Client chuyển tới VNPay hoặc hiển thị QR. Nút/redirect không tự đổi SUCCESS.",
            "VNPay callback/IPN hoặc POST `/api/v1/payments/sepay-webhook` đi vào controller public chuyên biệt.",
            "Adapter/service xác minh chữ ký/API key-HMAC, reference, amount, method và transaction; repository khóa payment theo transaction number.",
            "Nếu callback hợp lệ và chưa xử lý, payment -> SUCCESS, booking -> SUCCESS, seats -> BOOKED, tickets được tạo trong transaction.",
            "Sau commit gửi WebSocket/email/audit. Callback trùng trả acknowledgment nhưng không lặp side effect.",
            "Late success/refund/cancel showtime được ghi payment event và đưa reconciliation/refund policy thay vì sửa state mù.",
        ),
        request_payload="Initiate: {bookingId: UUID, method: 'VNPAY'|'SEPAY'}. VNPay callback là map `vnp_*` có SecureHash; SePay webhook chứa id/gateway/transactionDate/accountNumber/content/transferAmount/reference fields theo provider. Backend dùng amount/reference đã ký/xác thực, không dùng số tiền UI gửi.",
        database_state="payments PENDING -> SUCCESS/FAILED/EXPIRED và refund state theo migration; bookings PENDING -> SUCCESS hoặc terminal phù hợp; seat_status HOLD -> BOOKED/AVAILABLE; tickets ACTIVE sinh một lần. payment_events/audit lưu callback đã lọc, signature result, retry/reconciliation; unique transaction/reference hỗ trợ idempotency.",
        ide_order=(
            "Frontend PaymentPage, paymentApi và result polling/auto-scroll UX.",
            "PaymentController `/initiate`, VNPay callback, SePay webhook, events/refunds/reconciliation.",
            "PaymentServiceImpl: lock, reuse pending, validate amount/state và apply result.",
            "PaymentGateway interface/registry, VNPayPaymentGateway, SepayPaymentGateway.",
            "Payment/PaymentEvent/Refund entity + repository locked queries + Flyway indexes/constraints.",
            "BookingServiceImpl payment success/failure và email/ticket/WebSocket after commit.",
            "Payment integration/idempotency/security tests.",
        ),
        arrow_flow="PaymentPage -> POST /payments/initiate -> lock booking/payment\n-> PaymentGateway(VNPay redirect | SePay QR)\nProvider -> callback/webhook -> verify signature/reference/amount\n-> lock transaction -> idempotent state transition\n-> booking SUCCESS + seats BOOKED + tickets\n-> after commit email/WebSocket/audit\n-> reconciliation/refund handles mismatch/late result",
        comparisons=(
            ("Tích hợp cổng", "Strategy/adapter PaymentGateway", "if/else trong service", "Mở rộng cổng và test riêng, business state machine dùng chung; thêm abstraction có chi phí nhưng loại duplication thực."),
            ("Xác nhận", "Server callback/webhook", "Tin redirect/nút client", "Hoạt động khi browser đóng và chống giả kết quả phía client."),
            ("Chống trùng", "Row lock + unique + state guard", "Disable nút frontend", "Bảo vệ cả retry/provider/multi-instance; UI disable chỉ giảm thao tác nhầm."),
            ("Độ tin cậy", "Audit + reconciliation", "Chỉ callback một lần", "Phát hiện callback mất/late/mismatch; thêm job và vận hành."),
        ),
        test_steps=(
            "Tạo booking PENDING còn hạn; initiate SePay và kiểm tra QR amount/reference khớp DB.",
            "Nhấn initiate hai lần nhanh; xác nhận reuse pending hoặc chỉ một transaction hợp lệ.",
            "Gửi webhook/callback đúng; kiểm tra payment/booking/seats/tickets và trang kết quả tự chuyển.",
            "Gửi lại callback y hệt; kiểm tra không sinh ticket/email/promotion lần hai.",
            "Sửa signature, amount hoặc reference; backend từ chối, ghi event nhưng không đổi SUCCESS.",
            "Để booking hết hạn rồi gửi success muộn; kiểm tra reconciliation/refund policy.",
            "Hủy suất đã trả tiền; kiểm tra refund pending, permission staff/admin và audit.",
        ),
        edge_cases=(
            "Người dùng đóng tab trước callback hoặc callback tới trước redirect.",
            "SePay chuyển sai nội dung/sai số tiền hoặc cùng nội dung được dùng lại.",
            "Promotion đổi sau khi QR đã cố định amount.",
            "Provider retry callback nhiều lần hoặc callback đảo thứ tự FAILED/SUCCESS.",
            "DB commit thành công nhưng email/WebSocket thất bại.",
        ),
        senior_findings=(
            ("Gọi provider trong DB lock", "Giữ connection/lock khi mạng chậm", "Chuẩn bị state ngắn; gọi ngoài lock phù hợp rồi xác nhận bằng transaction/idempotency."),
            ("Search join graph rộng", "N+1/Cartesian/latency admin", "Projection pageable, index status/method/created_at/transaction."),
            ("Webhook không rate limit/audit", "Spam và khó điều tra", "Signature trước business, rate limit hợp lý, event log/correlation."),
            ("Retry không backoff", "Thundering herd provider/DB", "Exponential backoff+jitter, max attempt và dead-letter/manual queue."),
        ),
        scale_summary="10.000 payment cùng lúc làm gateway latency, callback burst, DB lock theo booking và email fan-out nghẽn. Payment initiation/callback phải idempotent và transaction ngắn; index transaction/reference; queue email/refund/reconciliation; rate limit initiation; autoscale stateless adapters. Tiền không được xử lý chỉ bằng cache. Metric pending age, callback failure, signature failure và refund backlog là bắt buộc.",
        hard_questions=(
            ("Nếu callback thành công nhưng response về provider bị mất?", "Provider retry; idempotency lock/state guard nhận lại và trả acknowledgment mà không lặp side effect."),
            ("Tại sao redirect SUCCESS chưa đủ?", "Redirect đi qua browser và có thể giả/mất; chỉ callback server đã xác minh signature, amount, reference mới đổi dữ liệu."),
            ("Tiền tới sau khi ghế đã trả thì xử lý sao?", "Không chiếm lại ghế. Ghi late payment event, reconciliation và refund/manual policy để bảo toàn invariant ghế lẫn tài chính."),
        ),
    ),
    "08_Database_Flyway_Query_Index_Cache_NPlusOne_CinemaBooking.docx": DefenseLab(
        tldr="PostgreSQL là nguồn sự thật cho quan hệ, trạng thái và concurrency; Flyway quản lý lịch sử schema, repository/projection kiểm soát truy vấn, index phục vụ access pattern đã đo và Caffeine chỉ cache dữ liệu ít đổi. Tối ưu bắt đầu bằng SQL profiling/EXPLAIN chứ không thêm cache hoặc EAGER theo cảm tính.",
        sequence_steps=(
            "Application startup kết nối datasource; Flyway đọc schema history và chạy migration V1..Vn còn thiếu theo thứ tự.",
            "Hibernate `ddl-auto=validate` so entity mapping với schema, không tự tạo/sửa production database.",
            "Controller/service gọi repository với page/filter/scope; Hibernate sinh SQL hoặc dùng JPQL/native/projection định nghĩa sẵn.",
            "PostgreSQL planner chọn Seq Scan/Index Scan/Join dựa statistics; constraint/lock bảo vệ invariant khi ghi.",
            "Mapper chuyển projection/entity đã fetch thành DTO mà không chạm lazy relation ngoài kế hoạch.",
            "Cacheable read kiểm tra Caffeine trước DB; mutation evict key liên quan sau thay đổi.",
            "Profiling log và EXPLAIN ANALYZE đo actual rows/loops/time; migration mới thêm/sửa index nếu có bằng chứng.",
        ),
        request_payload="Database không nhận payload HTTP trực tiếp; service chuyển request filter/page/sort thành tham số query đã bind. Tuyệt đối không nối chuỗi SQL từ input. Pageable có giới hạn size; search/filter được normalize và scope cinema/user đi vào WHERE.",
        database_state="Flyway schema history ghi version/checksum. Transaction commit entity/state; constraint PK/FK/UNIQUE/CHECK từ chối invariant sai. Cache không phải state chính và phải evict khi master data đổi. Mock-data chỉ dành local/demo, không thay migration production.",
        ide_order=(
            "`src/main/resources/db/migration/V*.sql` theo version và `database/database.sql` để xem snapshot.",
            "Entity quan hệ/fetch type và index/constraint tương ứng.",
            "Repository query/projection/EntityGraph/Pageable của một màn hình.",
            "Service/mapper xem có chạm lazy collection trong loop không.",
            "application cache/JPA/Flyway config và cache eviction.",
            "Bật SQL statistics, chạy endpoint, lấy SQL sang EXPLAIN ANALYZE.",
        ),
        arrow_flow="HTTP filters/page -> Service -> Repository bind parameters\n-> PostgreSQL planner -> index/join/lock/constraint\n-> projection/DTO -> cache/response\nMigration change -> new Flyway Vn (never edit applied Vn)\nMeasure SQL -> EXPLAIN ANALYZE -> optimize -> re-measure",
        comparisons=(
            ("Schema evolution", "Flyway versioned migration", "Hibernate ddl-auto=update", "Có lịch sử/review/repeatable deploy; cần viết migration cẩn thận."),
            ("ID", "UUID v4 hiện tại", "BIGINT/UUIDv7", "UUID thuận tiện phân tán và không lộ sequence; index lớn/random hơn. Chưa cần migration rủi ro nếu tải chưa chứng minh."),
            ("Fetch danh sách", "Projection/EntityGraph có mục tiêu", "EAGER toàn bộ", "Giảm N+1 mà không tải graph khổng lồ."),
            ("Cache", "Caffeine local cho master data", "Redis/distributed cache", "Đơn giản, nhanh một node; Redis cần khi nhiều node và invalidation chung."),
        ),
        test_steps=(
            "Tạo database trống, chạy ứng dụng/test và kiểm tra Flyway áp đủ migration không lỗi checksum.",
            "Gọi endpoint list 1, 20 và 100 records; đếm số SQL để phát hiện N+1.",
            "Copy query chậm cùng tham số vào `EXPLAIN (ANALYZE, BUFFERS)` trên dữ liệu test đủ lớn.",
            "So sánh plan/thời gian trước sau index hoặc projection; kiểm tra write cost và index trùng.",
            "Gọi endpoint cache hai lần, quan sát hit; update dữ liệu rồi kiểm tra cache đã evict.",
            "Cố insert duplicate/invalid status/FK và xác nhận DB constraint cùng exception mapping.",
        ),
        edge_cases=(
            "Sửa migration đã chạy làm checksum mismatch.",
            "Fetch join collection cùng pagination tạo duplicate/count sai.",
            "Index có nhưng planner không dùng vì bảng nhỏ/selectivity thấp.",
            "Cache key thiếu city/status/page trả nhầm dữ liệu.",
            "Soft delete record nhưng query/report vẫn tính hoặc không còn đọc lịch sử.",
        ),
        senior_findings=(
            ("N+1", "Query tăng theo số dòng", "Hibernate statistics/query-count test, projection/join fetch/batch đúng chỗ."),
            ("Index thừa/thiếu", "Read chậm hoặc write/bloat tăng", "pg_stat/EXPLAIN, composite theo WHERE+ORDER, remove duplicate có bằng chứng."),
            ("Page size không giới hạn", "Memory/DB spike", "Max page size, keyset pagination cho bảng giao dịch lớn."),
            ("Cache stale đa node", "Mỗi instance trả dữ liệu khác", "TTL+event invalidation hoặc Redis; không cache transactional state."),
        ),
        scale_summary="10.000 request đồng thời thường nghẽn ở connection pool và query nóng, không phải CPU Java trước. Giới hạn pool theo DB capacity, page size, index access pattern, projection/read model dashboard, cache master data và replica đọc khi cần. Partition/archival bảng event/payment lớn sau khi đo. EXPLAIN trên dữ liệu gần production và load test là bằng chứng, không dựa vào số lượng index.",
        hard_questions=(
            ("Có index rồi sao query vẫn chậm?", "Planner có thể chọn seq scan do selectivity/bảng nhỏ; join/sort/rows estimate sai. Dùng EXPLAIN ANALYZE BUFFERS và statistics, không chỉ nhìn tên index."),
            ("Tại sao không đổi toàn bộ UUID sang BIGINT?", "Có trade-off index/storage, nhưng migration xuyên hệ thống rủi ro. Chỉ đổi khi benchmark chứng minh bottleneck; UUID vẫn phù hợp tính phân tán/opaque."),
            ("Cache có giải quyết N+1 không?", "Không. Cache có thể che tạm và tạo stale state. Sửa fetch/query trước, cache dữ liệu ổn định sau."),
        ),
    ),
})


DEFENSE_LABS.update({
    "09_Frontend_React_Router_State_API_Realtime_UX_CinemaBooking.docx": DefenseLab(
        tldr="Frontend React tổ chức theo route và domain API; React Query quản lý server state, store/provider quản lý session và state UI nhỏ, Axios tập trung token/refresh/error. Realtime chỉ vá hoặc invalidate cache, còn responsive/accessibility giữ cùng business flow trên desktop, tablet và mobile.",
        sequence_steps=(
            "AppRouter chọn public/protected/admin/staff route và chờ session hydrate trước khi redirect.",
            "Page đọc route param/filter local, gọi domain hook/API có React Query key đầy đủ.",
            "Axios client ghép base URL/proxy, Authorization, timeout và normalize error; 401 đi qua single-flight refresh.",
            "Backend trả response DTO/page ổn định; React Query cache theo key và component render loading/empty/error/success.",
            "Mutation disable thao tác đang chạy, gửi DTO và khi thành công update/invalidate đúng query.",
            "Seat map subscribe WebSocket; event delta cập nhật cache hoặc refetch khi reconnect.",
            "Responsive CSS thay layout chứ không nhân đôi logic; focus, label và touch target hỗ trợ keyboard/mobile.",
            "Cleanup effect hủy request/subscription/camera khi unmount hoặc route param đổi.",
        ),
        request_payload="TypeScript request type phải khớp backend DTO, ví dụ seatIds là UUID[], date là ISO/local contract thống nhất và page/filter không gửi undefined vô nghĩa. Response/error được type/normalize; UI không parse entity nội bộ hoặc raw exception string.",
        database_state="Frontend không trực tiếp lưu DB. Mutation thành công chỉ được coi là hoàn tất theo response backend; React Query cache là bản sao tạm có thể stale và được invalidation/refetch. Session storage không phải nguồn permission cuối; backend database và JWT/session policy quyết định.",
        ide_order=(
            "AppRouter và layout để hiểu route/public/protected/role UX.",
            "Page cần học, sau đó component con và hook được gọi trực tiếp.",
            "Domain `*Api.ts`/types để lấy endpoint và payload.",
            "`axiosClient.ts` để hiểu token, refresh, proxy/base URL và error.",
            "React Query key/mutation invalidation và Zustand/Auth provider.",
            "WebSocket/camera/map hook có lifecycle; cuối cùng đọc responsive styles/tests/build config.",
        ),
        arrow_flow="Route -> Page -> useQuery(domainApi) -> axiosClient\n-> REST DTO -> React Query cache -> Component\nMutation -> pending UI -> API -> invalidate/patch cache\nWebSocket delta -> same cache\n401 -> one refresh promise -> retry\nUnmount -> abort/unsubscribe/deactivate",
        comparisons=(
            ("Server state", "React Query", "useEffect + useState thủ công", "Cache/deduplicate/loading/retry/invalidate chuẩn; cần thiết kế query key đúng."),
            ("Global state", "Zustand/provider nhỏ", "Redux cho mọi dữ liệu", "Ít boilerplate cho session/UI; Redux phù hợp state machine client rất lớn."),
            ("Routing", "React Router SPA", "Server-side routing/Next.js", "Phù hợp app hiện tại; SSR tốt hơn SEO/first paint public khi cần."),
            ("Realtime cache", "Patch/invalidate React Query", "Seat map state song song", "Một nguồn server state, giảm lệch và closure cũ."),
        ),
        test_steps=(
            "Mở public movie/cinema ở cửa sổ ẩn danh; protected action phải dẫn login rồi quay lại intent.",
            "Dùng DevTools Network kiểm tra query key/filter tạo request đúng và không gọi lặp vô hạn.",
            "Làm access token hết hạn, bắn nhiều request và kiểm tra một refresh rồi các request retry.",
            "Test loading/empty/400/401/403/409/429/500; message tiếng Việt và layout không rung/vỡ.",
            "Mở hai tab seat map, mất mạng/reconnect và đổi route để kiểm tra cleanup realtime.",
            "Kiểm tra 360px, tablet và desktop; keyboard tab/focus, dark mode, input/password và text dài.",
        ),
        edge_cases=(
            "Effect dependency tạo object mới dẫn Maximum update depth exceeded.",
            "Hai camera/socket vì cleanup không đối xứng trong StrictMode.",
            "Response cũ ghi đè filter mới do request không abort/query key thiếu.",
            "Double click mutation và retry tự động tạo booking/payment trùng.",
            "Text dài hoặc countdown đổi width làm card rung/vỡ.",
        ),
        senior_findings=(
            ("Bundle/page lớn", "First load và mobile chậm", "Route lazy loading, bundle analyze, tối ưu ảnh/poster và library nặng."),
            ("Render toàn list", "Main thread lag", "Memo đúng chỗ, pagination/windowing khi dữ liệu lớn, tránh state thay đổi toàn cây."),
            ("Query key sai", "Cache nhầm hoặc refetch thừa", "Query key factory theo domain/filter; React Query Devtools/profile."),
            ("Listener leak", "Memory/CPU tăng theo điều hướng", "AbortController, unsubscribe/deactivate cleanup và test mount/unmount."),
        ),
        scale_summary="10.000 user không nằm trong một browser, nhưng backend/API/CDN chịu tải. Frontend giảm tải bằng cache HTTP/React Query, debounce search, pagination, WebSocket delta thay polling và ảnh qua CDN. Một client vẫn phải tránh render hàng nghìn node/listener. Khi traffic public lớn, cân nhắc SSR/CDN cho catalog; booking/payment giữ API động và idempotent.",
        hard_questions=(
            ("Tại sao vừa Zustand vừa React Query?", "Hai loại state khác nhau: React Query quản lý server state/staleness; Zustand/provider giữ session hoặc UI state nhỏ. Không nhân bản cùng dữ liệu ở cả hai."),
            ("Ẩn nút theo permission có bảo mật không?", "Không, chỉ UX. Backend `@PreAuthorize` và data scope mới quyết định; frontend xử lý 403 thân thiện."),
            ("Vì sao StrictMode làm hai socket và sửa thế nào?", "StrictMode kiểm tra side effect bằng mount-cleanup-mount. Effect phải tạo resource một lần theo dependency và cleanup đúng instance, không tắt StrictMode để che lỗi."),
        ),
    ),
    "10_Email_Verification_Reset_Ticket_Notification_CinemaBooking.docx": DefenseLab(
        tldr="Email flow tạo token ngẫu nhiên dùng một lần, chỉ lưu hash và gửi link sau khi transaction chính commit. Vé điện tử được dựng từ booking SUCCESS cùng từng QR; lỗi SMTP không được làm sai booking/payment và có log/retry riêng.",
        sequence_steps=(
            "Register/ForgotPassword page gửi email/username tới UserController endpoint tương ứng.",
            "Service validate request, tạo user hoặc tìm user theo response trung tính để không lộ account tồn tại.",
            "Token random được sinh; database lưu token hash + expiry, không lưu bản thô.",
            "Transaction commit rồi application event gọi async email service/template renderer.",
            "Mail sender gửi link chứa raw token tới SMTP/provider; log không ghi token.",
            "Verify/reset page gửi raw token về backend; service hash, so sánh expiry/consumed và cập nhật email/password.",
            "Reset thành công revoke refresh sessions; verify xóa token hash/expiry.",
            "Payment SUCCESS phát ticket email: mapper lấy movie/cinema/room/seats/tickets, renderer tạo QR và gửi sau commit.",
        ),
        request_payload="Register {username,email,password}; verify {token}; resend/forgot {email}; reset {token,newPassword}; ticket notification dùng internal event chứa bookingId chứ không tin dữ liệu vé từ client. Link frontend/base URL và SMTP config lấy từ environment.",
        database_state="users lưu email_verified, verification_token_hash/expiry, reset_token_hash/expiry và BCrypt password. Raw token chỉ ở link email. Booking/payment/ticket đã SUCCESS trước khi ticket email phát; notification failure được log/retry, không rollback tiền hoặc xóa vé.",
        ide_order=(
            "UserController register/verify/resend/forgot/reset endpoint và request DTO validation.",
            "UserService/Auth service tạo/hash/consume token và revoke session.",
            "Email event/listener/AsyncConfig để thấy after-commit và executor.",
            "EmailService/template builder và ticket QR rendering.",
            "application mail properties + `.env.example`; HTML templates/resources nếu có.",
            "Frontend Register/Verify/Reset pages và URL token parsing.",
        ),
        arrow_flow="Register/reset request -> service transaction\n-> store token HASH + expiry -> commit\n-> async event -> SMTP sends RAW token link\nUser clicks -> frontend -> verify/reset API\n-> hash received token -> compare -> consume once\nPayment commit -> ticket event -> QR email",
        comparisons=(
            ("Token storage", "Hash token trong DB", "Lưu raw token", "Giảm tác hại khi DB lộ; không thể resend đúng token cũ, phải sinh token mới."),
            ("Gửi mail", "After-commit async", "SMTP trong transaction", "Không giữ DB connection/rollback vì mạng; cần retry/observability riêng."),
            ("Template", "HTML + plain fallback", "Text thuần", "Trình bày vé/QR tốt hơn; HTML email có giới hạn CSS/client."),
            ("Đảm bảo delivery", "Async log/retry", "Outbox + queue", "Đủ monolith hiện tại; outbox mạnh hơn khi email bắt buộc không mất."),
        ),
        test_steps=(
            "Đăng ký email mới, kiểm tra DB chỉ có hash và Mailtrap nhận đúng link.",
            "Bấm verify một lần thành công, lần hai bị từ chối/idempotent theo contract.",
            "Yêu cầu resend, kiểm tra token/link cũ hết hiệu lực và rate limit.",
            "Forgot password với email có/không tồn tại; response không tiết lộ khác biệt đáng kể.",
            "Reset đúng/sai/hết hạn; password mới login được và refresh session cũ bị revoke.",
            "Thanh toán thành công nhiều ghế; email có địa chỉ, city, phòng, giờ và QR riêng từng vé.",
            "Tắt SMTP rồi thanh toán; booking/ticket vẫn đúng, notification có log/retry.",
        ),
        edge_cases=(
            "Hai yêu cầu resend/reset gần nhau tạo nhiều token còn hiệu lực.",
            "Email listener chạy trước commit hoặc callback duplicate gửi nhiều vé.",
            "Logo dùng localhost nên Gmail không tải được.",
            "Tên/phim chứa HTML gây injection template nếu không escape.",
            "Async executor queue đầy làm mất/đẩy chậm notification.",
        ),
        senior_findings=(
            ("SMTP đồng bộ", "Latency/thread exhaustion", "After-commit executor giới hạn; queue/provider API ở quy mô lớn."),
            ("Không idempotency", "Email vé trùng", "Notification key event+template+recipient/booking."),
            ("Token query thiếu index", "Verify/reset chậm khi bảng lớn", "Index hash/expiry, cleanup retention."),
            ("Template dựng N+1", "Nhiều query theo ghế/ticket", "Fetch/projection booking email một lần và map ngoài transaction."),
        ),
        scale_summary="10.000 email cùng lúc không nên tạo 10.000 SMTP call trên request threads. Commit dữ liệu, ghi event/outbox, worker tiêu thụ có rate limit theo provider, batch/ retry backoff và dead-letter. QR generation có thể tốn CPU/memory nên giới hạn concurrency. Theo dõi queue age, failure/bounce và bảo đảm user vẫn lấy vé trong app.",
        hard_questions=(
            ("Tại sao DB chỉ lưu hash nhưng email cần raw token?", "Raw token chỉ tồn tại lúc sinh và trong link; khi callback, backend hash token nhận được để so. DB leak không cung cấp token dùng được."),
            ("SMTP lỗi sau thanh toán có rollback booking không?", "Không. Booking/payment/ticket là transaction chính đã commit; email là side effect retry được, user vẫn xem vé trong app."),
            ("Làm sao callback duplicate không gửi hai email?", "Payment transition idempotent và notification có key/audit; chỉ event của transition đầu được xử lý thành công."),
        ),
    ),
    "11_Admin_Staff_QR_Map_Operation_CinemaBooking.docx": DefenseLab(
        tldr="Portal Admin/Staff vận hành phim, rạp, phòng/ghế, suất chiếu, booking, payment/refund và check-in; mọi thao tác staff bị giới hạn theo rạp được gán. QR chỉ chuyển ACTIVE sang USED sau khi đúng chữ ký, booking, rạp, suất và cửa sổ check-in; thao tác nguy hiểm có policy/audit.",
        sequence_steps=(
            "Admin/staff route tải session/capability và danh sách cinema assignment; UI chỉ hiện module được phép.",
            "List page gửi page/filter city/cinema/date/status/search tới endpoint server-side; response là summary DTO.",
            "Create/update form validate client rồi controller `@Valid` và `@PreAuthorize`; service kiểm tra staff scope cùng invariant domain.",
            "Tạo suất kiểm tra room thuộc cinema, assignment, end > start và không overlap; transaction ghi showtime/seat status cần thiết.",
            "Hủy suất gọi policy: chưa trả nhả ghế; đã trả tạo refund pending/hủy ticket/thông báo và audit.",
            "Scanner chọn cinema/showtime trong scope, camera/file decode QR và POST check-in cùng context.",
            "Backend verify signed QR, ticket ACTIVE, booking SUCCESS, đúng cinema/showtime/time window rồi atomic update USED/checkInTime/staff.",
            "Map page tải `/api/v1/cinemas/map` hoặc nearest, tính/sắp khoảng cách và marker từ tọa độ API.",
        ),
        request_payload="Danh sách dùng query page,size,search,status,city,cinemaId,date. Showtime create có movieId, roomId, startTime, endTime, basePrice; backend suy cinema từ roomId. Check-in có qrCode, cinemaId, showtimeId. Cancel/refund action có reason khi policy yêu cầu; actor lấy từ JWT, không từ body.",
        database_state="CRUD master chủ yếu soft delete; showtime trạng thái UPCOMING/ONGOING/ENDED/CANCELLED; cancel có thể tạo refund request và ticket CANCELLED. tickets ACTIVE -> USED lưu check_in_time/check_in_by. staff_cinemas giới hạn dữ liệu; admin_audit_logs lưu actor/action/resource/reason/result.",
        ide_order=(
            "AdminLayout/AppRouter và capability navigation để thấy module theo role.",
            "Page list/form tương ứng và domain API/filter/pagination.",
            "Controller + `@PreAuthorize` của Showtimes, Payments, Tickets, Cinemas/Rooms/Seats.",
            "Service policy: StaffCinemaScopeService, showtime overlap/cancel, refund và ticket check-in.",
            "Repository scoped projection/locked ticket query + migrations/audit.",
            "QR scanner camera/file lifecycle và cinema map/nearest hook.",
        ),
        arrow_flow="Admin/Staff UI -> filtered pageable API\n-> permission -> derive actual cinema -> staff scope\n-> business policy -> transaction -> audit\nShowtime cancel -> booking/refund/ticket workflow\nScanner context + QR -> verify -> atomic ACTIVE->USED\nMap coordinates -> marker + nearest sorting",
        comparisons=(
            ("Admin list", "Server filter/page/projection", "Load all rồi lọc client", "Ổn định khi dữ liệu lớn và bảo đảm staff scope ngay trong query."),
            ("Xóa", "Soft delete/policy cancel", "Hard delete", "Giữ lịch sử booking/tài chính; cần query active nhất quán."),
            ("Check-in", "QR + cinema/showtime context", "Chỉ QR", "Ngăn quét nhầm rạp/suất làm vé bị USED; thêm một bước chọn context."),
            ("Tọa độ", "Lưu lat/lng và preview", "Geocode mỗi lần tải", "Nhanh/ổn định; admin phải nhập/kiểm tra dữ liệu đúng."),
        ),
        test_steps=(
            "Login staff chỉ rạp A; kiểm tra list/filter/form không thấy rạp B và sửa request roomId B bị 403.",
            "Tạo suất hợp lệ; thử end trước start và overlap cùng phòng, kiểm tra lỗi tiếng Việt.",
            "Hủy suất không booking, pending booking và paid booking; kiểm tra ghế/refund/ticket/audit từng nhánh.",
            "Scanner chọn đúng rạp/suất, quét vé ACTIVE; quét lại phải báo USED kèm giờ/nhân viên.",
            "Quét vé đúng QR nhưng sai rạp, sai suất, ngoài cửa sổ và file ảnh mờ; ticket không được dùng.",
            "Map Gần tôi: cho phép/từ chối geolocation, giới hạn km, dữ liệu thiếu/sai tọa độ và mobile layout.",
        ),
        edge_cases=(
            "Staff thấy nút nhưng assignment vừa bị thu hồi.",
            "Hai admin cùng sửa/hủy suất hoặc check-in cùng QR.",
            "Hủy suất đang chiếu/đã kết thúc hoặc có payment late.",
            "Camera StrictMode tạo hai preview hay removeChild lỗi.",
            "Export/dashboard quên staff scope dù list đã lọc đúng.",
        ),
        senior_findings=(
            ("List join rộng", "Admin latency/N+1", "Projection pageable, index filter+sort, read model dashboard."),
            ("Policy rải ở UI/controller", "Bypass và hành vi không nhất quán", "Service policy dùng chung + integration test ma trận role/scope."),
            ("Check-in read-then-update", "Hai scanner cùng dùng vé", "Conditional/locked ACTIVE -> USED atomic."),
            ("Map render nhiều marker", "UI lag", "Bounding-box query, clustering/virtual list khi số rạp lớn."),
        ),
        scale_summary="10.000 request vận hành cần phân biệt traffic customer với staff. List/report dùng replica/read model, pagination và index; mutation giữ transaction ngắn/audit async phù hợp. QR check-in có hotspot theo showtime nhưng conditional update nhỏ. Map dùng bounding box/clustering. Hủy hàng nghìn booking/refund nên tạo batch job/queue thay vì một HTTP transaction khổng lồ.",
        hard_questions=(
            ("Vì sao staff không được hoàn tiền trực tiếp?", "Tách nhiệm vụ và rủi ro tài chính; staff tạo/theo dõi yêu cầu trong scope, admin/kế toán có permission xác nhận và audit."),
            ("QR đã ký rồi sao còn cần chọn rạp/suất?", "Chữ ký chỉ chứng minh QR hợp lệ; context ngăn nhân viên quét nhầm địa điểm/suất và tiêu thụ vé sai."),
            ("Hủy suất có thể chạy một transaction cho 100.000 vé không?", "Không nên. Policy tạo trạng thái/công việc idempotent rồi worker batch refund/notify, có progress, retry và reconciliation."),
        ),
    ),
    "12_Exception_Audit_SoftDelete_Monitoring_CinemaBooking.docx": DefenseLab(
        tldr="Exception layer chuyển lỗi validation, authentication, authorization, business conflict và lỗi bất ngờ thành contract ổn định cho client mà không lộ SQL/stack trace. Audit ghi hành động nghiệp vụ có trách nhiệm; monitoring đo triệu chứng hệ thống; soft delete giữ lịch sử trong khi ẩn dữ liệu khỏi luồng active.",
        sequence_steps=(
            "Request sai Bean Validation bị bắt trước service; handler tạo 400 cùng fieldErrors.",
            "JWT thiếu/sai đi AuthenticationEntryPoint -> 401; principal đủ nhưng thiếu authority/scope đi AccessDeniedHandler -> 403.",
            "Service ném NotFound/Conflict/Business exception có error code khi rule không thỏa.",
            "GlobalExceptionHandler map exception dự kiến sang status/code/message/timestamp/path/correlationId.",
            "Runtime/JDBC exception ngoài dự kiến được log cause phía server và trả 500 message trung tính.",
            "Transaction rollback nếu exception phù hợp; event ngoài DB không phát trước commit.",
            "Audit interceptor/service ghi actor, action, resource, result và metadata đã sanitize cho thao tác quan trọng.",
            "Metrics/log dashboard tổng hợp error rate, latency, payment failure, scheduler backlog và alert theo ngưỡng.",
        ),
        request_payload="ErrorResponse chuẩn gồm code số/chuỗi ổn định, message thân thiện, timestamp, path, correlationId và fieldErrors tùy validation. Client dựa code/status để xử lý, không parse raw Java message. Audit metadata không chứa password/token/secret/full sensitive payload.",
        database_state="Business transaction rollback khi lỗi; audit thất bại/thành công được lưu theo policy độc lập an toàn. Master data soft delete set is_deleted và updated_at; history booking/payment/ticket vẫn giữ FK. Payment events/refunds không bị xóa mà chuyển trạng thái/retention theo quy định.",
        ide_order=(
            "Các custom exception/error code enum và ErrorResponse DTO.",
            "GlobalExceptionHandler: mapping validation/not found/conflict/rate limit/unexpected.",
            "Security AuthenticationEntryPoint và AccessDeniedHandler để phân biệt 401/403.",
            "Một service ném exception trong `@Transactional` và test rollback.",
            "Admin/Auth/Payment audit entity/service/controller và sanitization.",
            "Soft-delete policy/query và monitoring endpoint/log configuration.",
        ),
        arrow_flow="Request -> validation/security/business\n-> expected exception -> stable HTTP status + error code\n-> unexpected exception -> server log cause + generic 500\nTransaction exception -> rollback\nImportant action -> sanitized audit\nLogs + metrics -> alert/runbook\nSoft delete -> hidden active data, preserved history",
        comparisons=(
            ("Error contract", "Global handler + code", "Trả exception message", "Ổn định, localize và không lộ nội bộ; cần quản lý catalog code."),
            ("Lịch sử", "Soft delete/state transition", "Hard delete", "Giữ tham chiếu/audit; query active phải nhất quán và retention rõ."),
            ("Theo dõi", "Structured logs + metrics", "Chỉ println/log text", "Tìm correlation và alert được; cần cấu hình volume/retention."),
            ("Audit", "Bảng nghiệp vụ riêng", "Dùng application log", "Truy vấn/retention/quy trách nhiệm rõ; tăng storage và cần sanitize."),
        ),
        test_steps=(
            "Gửi DTO thiếu/sai field; kiểm tra 400, fieldErrors và không có stack/SQL.",
            "Gọi protected API không token, token hết hạn, thiếu permission và sai cinema scope; phân biệt 401/403.",
            "Tạo conflict ghế/duplicate/check-in USED; kiểm tra 409/error code thân thiện.",
            "Gây exception giữa transaction test; assert mọi state rollback và không có WebSocket/email giả.",
            "Soft delete movie/cinema đã được tham chiếu; public ẩn nhưng history vẫn đọc được.",
            "Thực hiện hủy suất/refund/role assignment; lọc audit theo actor/action/resource.",
            "Mô phỏng payment callback lỗi lặp; kiểm tra metric/monitoring summary và log correlation.",
        ),
        edge_cases=(
            "Catch exception rồi trả success khiến transaction commit sai.",
            "Mọi security error thành 403 và client refresh token vô ích.",
            "Log raw Authorization, QR secret, password hoặc webhook secret.",
            "Soft-deleted data vẫn xuất hiện ở export/cache hoặc unique key không tái sử dụng rõ.",
            "Audit ghi thất bại làm thao tác chính rollback ngoài ý muốn.",
        ),
        senior_findings=(
            ("Exception handler catch-all sớm", "Che business status/cause", "Handler cụ thể trước, test mapping và preserve cause trong log."),
            ("Log volume cao", "IO/cost và khó tìm tín hiệu", "Structured level, sampling, correlation và retention."),
            ("Audit sync mọi read", "Write amplification", "Audit thao tác nhạy cảm; async/outbox phù hợp nhưng bảo đảm ordering cần thiết."),
            ("Soft delete filter rải rác", "Rò dữ liệu hoặc báo cáo sai", "Repository policy/specification, test public/admin/history và cache eviction."),
        ),
        scale_summary="Ở 10.000 user, log đồng bộ/stack trace và audit mỗi request có thể thành bottleneck. Dùng structured async appender/collector, metric aggregation, sampling cho success, giữ đầy đủ security/payment failure, index audit theo time/actor/action và partition/retention. Error storm phải có rate limit/circuit breaker và alert theo tỷ lệ, không làm logger kéo sập app.",
        hard_questions=(
            ("Tại sao 401 khác 403 quan trọng?", "401 là chưa xác thực/token không hợp lệ nên có thể refresh/login; 403 là đã xác thực nhưng thiếu quyền/scope, refresh không giúp. Cả UX và security flow khác nhau."),
            ("Soft delete có phải backup không?", "Không. Nó chỉ giữ record trong DB và vẫn có thể bị sửa/xóa/lỗi hệ thống; production cần backup/restore và retention riêng."),
            ("Audit log có được rollback cùng business transaction?", "Tùy mục tiêu: success audit nên phản ánh commit; failure/security audit có thể cần transaction độc lập/event. Phải tránh audit failure phá nghiệp vụ và giữ correlation."),
        ),
    ),
    "13_Kiem_thu_Trien_khai_Bao_ve_CinemaBooking.docx": DefenseLab(
        tldr="Chiến lược test ưu tiên invariant rủi ro cao bằng unit test nhanh và PostgreSQL integration test cho query, lock, transaction; smoke/manual checklist xác nhận hành trình UI. Quy trình bảo vệ có dữ liệu cố định, phương án offline và trả lời theo vấn đề - giải pháp - bằng chứng - giới hạn - hướng mở rộng.",
        sequence_steps=(
            "Maven test khởi tạo Spring context/slice/unit tùy lớp; Testcontainers bật PostgreSQL tạm cho test cần engine thật.",
            "Flyway tạo schema test; fixture/SQL factory tạo dữ liệu tối thiểu độc lập cho từng case.",
            "Test gọi controller/service/repository, thu response và quan sát database transaction/state.",
            "Concurrency test dùng hai thread/transaction đồng bộ thời điểm để tranh cùng seat/payment/ticket.",
            "Security test tạo principal/role/scope khác nhau và assert 401/403/ownership.",
            "Frontend build/type/lint bắt contract và lỗi compile; manual responsive/realtime/camera/gateway kiểm tra phần khó tự động hiện tại.",
            "Sau deploy, smoke test health/login/public browse/sandbox payment/WebSocket và kiểm tra migration/log.",
            "Kết quả test và checklist demo trở thành bằng chứng khi bảo vệ, không chỉ lời khẳng định.",
        ),
        request_payload="Test dùng DTO giống API thật và fixture có UUID/state/time xác định. Integration test không gọi database dev/prod; Testcontainers tạo DB tạm. Gateway/SMTP bên ngoài dùng mock/sandbox và callback fixture đã ký; secret không hard-code trong source test công khai.",
        database_state="Mỗi test tạo/rollback hoặc dọn dữ liệu độc lập; Flyway schema phải khớp production. Assert state cuối ở nhiều bảng, không chỉ HTTP: booking, payment, seat_status, ticket, promotion count, refund và audit/event. Concurrency test xác nhận chỉ một transition thắng.",
        ide_order=(
            "pom.xml test dependencies/profiles và application-test config.",
            "Unit tests policy/mapper/security/gateway signature.",
            "Repository/PostgresIntegrationTest/Testcontainers setup và Flyway.",
            "Booking/payment concurrency + idempotency integration tests.",
            "Controller/security/exception tests.",
            "Frontend package scripts và manual defense checklist/README.",
        ),
        arrow_flow="mvn test -> unit + Spring slices + PostgreSQL container\n-> Flyway -> fixture -> action -> assert response + DB invariant\nConcurrent threads -> one winner\nFrontend check/build -> manual responsive/realtime/gateway smoke\nDeploy -> migration -> health -> critical journey",
        comparisons=(
            ("DB test", "PostgreSQL Testcontainers", "H2/mocked repository", "Xác minh lock/JSONB/native SQL/constraint thật; cần Docker và chạy chậm hơn."),
            ("Phạm vi", "Test pyramid theo rủi ro", "100% E2E", "Nhanh, chẩn đoán dễ; manual/E2E ít cho tích hợp quan trọng."),
            ("Gateway", "Sandbox + signed fixtures", "Thanh toán thật mỗi test", "Tái lập/an toàn; vẫn cần smoke production có kiểm soát."),
            ("Migration", "Flyway trên DB trống trong test", "Schema thủ công", "Bắt migration lỗi và drift sớm."),
        ),
        test_steps=(
            "Chạy `mvn test` trong backend; nếu integration cần Docker, mở Docker và kiểm tra container PostgreSQL khởi động.",
            "Chạy frontend lint/typecheck/build theo package scripts từ đúng thư mục client.",
            "Seed mock data và login ba role; kiểm tra public browse, auth refresh/logout, staff scope.",
            "Demo hai user tranh ghế, promotion, VNPay/SePay callback duplicate và vé email.",
            "Demo QR đúng/sai rạp/suất/USED; hủy suất và refund/audit/reconciliation.",
            "Kiểm tra mobile/tablet/dark mode/loading/error và realtime reconnect.",
            "Lưu log/test report ngắn, chuẩn bị fallback ảnh/video nếu ngrok/SMTP/gateway sandbox lỗi.",
        ),
        edge_cases=(
            "Test phụ thuộc thứ tự/dữ liệu cũ nên chạy riêng khác chạy cả suite.",
            "Test race dùng chung transaction nên không tạo cạnh tranh thật.",
            "Mock quá nhiều khiến query/constraint production chưa từng được chạy.",
            "Clock/port/ngrok/env khác ngày bảo vệ làm demo thất bại.",
            "Test chỉ happy path, không có duplicate/late/authorization/rollback.",
        ),
        senior_findings=(
            ("Flaky time/concurrency test", "CI không tin cậy", "Inject Clock, latch/barrier, timeout rõ và assert DB state."),
            ("Suite integration quá chậm", "Developer bỏ chạy test", "Phân profile/tag, reuse container hợp lý, fixture tối thiểu; unit test phần thuần."),
            ("Không test query count", "N+1 quay lại", "Hibernate statistics/query-count regression ở endpoint quan trọng."),
            ("Không load test", "Không biết capacity", "k6/JMeter profile hold/list/callback với dữ liệu gần thật, theo p95/error/DB pool."),
        ),
        scale_summary="Để chứng minh 10.000 user, không chạy mù trên laptop. Xây workload thực tế: phần lớn browse, một tỷ lệ chọn suất/giữ ghế, ít payment callback; tăng tải theo bậc, đo p95/p99, error, DB pool, lock wait, CPU/memory và WebSocket sessions. Tìm saturation point, tối ưu rồi chạy lại. Báo cáo giới hạn trung thực ghi điểm hơn tuyên bố chưa đo.",
        hard_questions=(
            ("Bạn chứng minh không double-book bằng test nào?", "Integration test PostgreSQL với hai transaction/thread giữ cùng ghế, assert một winner và invariant DB; mock unit test không đủ chứng minh lock."),
            ("Coverage bao nhiêu mới gọi là tốt?", "Không có số tuyệt đối; ưu tiên branch/invariant rủi ro. Trình bày ma trận Auth, booking race, payment idempotency, QR scope, cancel/refund và residual risk."),
            ("Nếu ngày bảo vệ cổng thanh toán/ngrok lỗi thì sao?", "Có sandbox/config đã kiểm tra, fixture callback ký đúng và dữ liệu/log/video fallback; vẫn demo state machine/idempotency local mà không giả dữ liệu production."),
        ),
    ),
})


if __name__ == "__main__":
    main()

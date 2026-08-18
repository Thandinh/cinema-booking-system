"""Academic expansion sections for the source-driven CinemaBooking thesis.

The module contains prose only. It does not inspect or depend on any previous
thesis document. Technical claims are tied to source files already inspected by
the main generator.
"""

from __future__ import annotations

from collections.abc import Iterable

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


NAVY = "07172E"
BLUE = "1D4ED8"
LIGHT_BLUE = "EAF2FF"
BORDER = "CBD5E1"
def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), BORDER)
        borders.append(element)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _add_inline_text(paragraph, text: str) -> None:
    for index, part in enumerate(text.split("`")):
        if not part:
            continue
        run = paragraph.add_run(part)
        if index % 2 == 1:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10.5)


def body(doc, text: str, *, indent: bool = True) -> None:
    paragraph = doc.add_paragraph()
    _add_inline_text(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        paragraph.paragraph_format.first_line_indent = Pt(24)


def body_many(doc, paragraphs: Iterable[str]) -> None:
    for text in paragraphs:
        body(doc, text)


def bullet_list(doc, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        _add_inline_text(paragraph, item)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def numbered_list(doc, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        _add_inline_text(paragraph, item)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def academic_table(doc, caption: str, headers: list[str], rows: list[list[str]], font_size: float = 9.0) -> None:
    caption_paragraph = doc.add_paragraph(caption, style="Table Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _borders(table)
    _repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _shade(cell, LIGHT_BLUE)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value.replace("`", ""))
            run.font.size = Pt(font_size)
    doc.add_paragraph()


def note(doc, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _borders(table)
    cell = table.cell(0, 0)
    _shade(cell, "F8FAFC")
    paragraph = cell.paragraphs[0]
    heading = paragraph.add_run(title + "\n")
    heading.bold = True
    heading.font.color.rgb = RGBColor.from_string(BLUE)
    paragraph.add_run(text)
    doc.add_paragraph()


def diagram_placeholder(doc, marker: str, explanation: str, caption: str) -> None:
    """Create a consistent insertion point for a rendered PlantUML diagram."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _borders(table)
    cell = table.cell(0, 0)
    _shade(cell, "FFF8E1")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = paragraph.add_run(marker + "\n")
    title.bold = True
    title.font.color.rgb = RGBColor.from_string("9A5B00")
    paragraph.add_run(explanation)
    caption_paragraph = doc.add_paragraph(caption, style="Figure Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def add_preliminary_pages(doc) -> None:
    doc.add_heading("LỜI CAM ĐOAN", level=1)
    body_many(doc, [
        "Tôi cam đoan nội dung khóa luận là kết quả của quá trình nghiên cứu, phân tích, thiết kế và hiện thực hệ thống CinemaBooking dưới sự hướng dẫn của giảng viên. Các mô tả về kiến trúc, dữ liệu, API và luồng nghiệp vụ được đối chiếu với mã nguồn, cấu hình, migration và test của dự án. Những nội dung sử dụng tài liệu bên ngoài đều được trích dẫn trong danh mục tài liệu tham khảo.",
        "Các kết quả kiểm thử và nhận định kỹ thuật trong báo cáo được trình bày trên cơ sở bằng chứng thu thập được. Những nội dung chưa có số liệu thực nghiệm, chẳng hạn kiểm thử tải hoặc cấu hình triển khai trong môi trường vận hành thực tế, được xác định là kịch bản dự kiến hoặc hướng phát triển. Tôi chịu trách nhiệm về tính trung thực của nội dung báo cáo theo quy định của cơ sở đào tạo.",
    ])
    body(doc, "[SINH VIÊN TỰ BỔ SUNG HỌ TÊN, MÃ SỐ SINH VIÊN, NGÀY VÀ CHỮ KÝ]", indent=False)
    doc.add_page_break()

    doc.add_heading("LỜI CẢM ƠN", level=1)
    body_many(doc, [
        "Trong quá trình thực hiện đề tài, sinh viên đã vận dụng kiến thức về phân tích hệ thống, lập trình web, cơ sở dữ liệu, an toàn thông tin và kiểm thử phần mềm để xây dựng hệ thống có các luồng nghiệp vụ liên kết. Những góp ý về phương pháp nghiên cứu, cách xác định phạm vi, cách trình bày kết quả và yêu cầu kiểm chứng đã tạo điều kiện để đề tài được hoàn thiện.",
        "Sinh viên xin trân trọng cảm ơn giảng viên hướng dẫn, quý thầy cô trong khoa và các cá nhân đã hỗ trợ về chuyên môn, góp ý giao diện, kiểm thử các tình huống đặt vé, thanh toán và vận hành rạp. Sự hỗ trợ này là nguồn động viên quan trọng trong quá trình hoàn thành khóa luận.",
    ])
    doc.add_page_break()

    doc.add_heading("TÓM TẮT", level=1)
    body_many(doc, [
        "Khóa luận trình bày quá trình phân tích, thiết kế và hiện thực hệ thống đặt vé xem phim trực tuyến CinemaBooking. Hệ thống được tổ chức theo mô hình client-server: giao diện người dùng phát triển bằng ReactJS và TypeScript; back-end sử dụng Java Spring Boot; dữ liệu được quản lý trong PostgreSQL thông qua Spring Data JPA, Hibernate và các migration Flyway. Giao tiếp chính giữa hai phía sử dụng RESTful API, trong khi WebSocket/STOMP được dùng để đồng bộ thay đổi trạng thái ghế theo thời gian thực.",
        "Bài toán trọng tâm không chỉ là hiển thị phim và lịch chiếu mà còn là bảo đảm tính nhất quán khi nhiều người cùng chọn một ghế, quản lý thời gian giữ ghế, xử lý đơn chờ thanh toán, tiếp nhận callback hoặc webhook từ cổng thanh toán, phát hành vé QR và kiểm soát check-in tại đúng rạp, đúng suất chiếu. Để giải quyết yêu cầu này, hệ thống kết hợp transaction, khóa bi quan trên các bản ghi trạng thái ghế, trường phiên bản, ràng buộc duy nhất ghế-suất chiếu và các chỉ mục một phần nhằm hạn chế đơn đặt vé hoặc giao dịch chờ bị tạo trùng. Các tác vụ định kỳ thu hồi trạng thái hết hạn, còn sự kiện thay đổi ghế chỉ được phát sau khi transaction hoàn tất.",
        "Về an toàn truy cập, hệ thống sử dụng Spring Security, JWT access token, refresh token được băm trước khi lưu, cơ chế rotation và thu hồi phiên. Phân quyền được xây dựng theo RBAC ở mức permission; nhân viên còn bị giới hạn theo danh sách rạp được quản trị viên phân công. Hai phương thức thanh toán chính được phân tích tách biệt: VNPay dùng URL chuyển hướng và callback có checksum; SePay/VietQR dùng QR chuyển khoản và webhook được đối chiếu thông tin giao dịch. Payment event, reconciliation và dữ liệu yêu cầu refund hỗ trợ vận hành, đối soát và truy vết sự cố.",
        "Báo cáo được xây dựng trên cơ sở khảo sát và đối chiếu các thành phần hiện thực của hệ thống. Các endpoint được xác định từ lớp Controller; quan hệ dữ liệu được kiểm tra giữa Entity và migration Flyway; các luồng phía React được đối chiếu với API client, định tuyến và kho trạng thái. Kết quả của khóa luận là một hệ thống đáp ứng các quy trình chính dành cho khách hàng, nhân viên và quản trị viên. Bên cạnh đó, báo cáo cũng xác định những giới hạn của phiên bản hiện tại, gồm bộ môi giới WebSocket và bộ nhớ đệm còn hoạt động trong phạm vi từng tiến trình, chưa có kết quả kiểm thử tải thực nghiệm và chưa triển khai trong môi trường vận hành chính thức.",
    ])
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Từ khóa: ")
    run.bold = True
    paragraph.add_run("đặt vé xem phim, ReactJS, Spring Boot, PostgreSQL, JWT, WebSocket, concurrency, VNPay, SePay, QR ticket.")
    doc.add_page_break()

    doc.add_heading("ABSTRACT", level=1)
    body_many(doc, [
        "This thesis presents the analysis, design, and implementation of CinemaBooking, an online cinema ticket booking system. The application separates a React and TypeScript client from a Java Spring Boot back-end. PostgreSQL stores persistent data, Spring Data JPA and Hibernate provide data access, and Flyway controls schema evolution. RESTful APIs are used for request-response communication, while STOMP over WebSocket distributes seat-status changes to connected clients.",
        "The main engineering concern is data consistency during seat selection and payment. The implemented source combines transactional service methods, pessimistic row locking, optimistic versioning, database constraints, expiration schedulers, and post-commit notifications. Authentication uses signed JWT access tokens and server-managed refresh-token rotation. Authorization combines permission-based RBAC with cinema assignments for staff users. VNPay and SePay are modeled as separate provider flows because their callback, signature, and confirmation mechanisms differ.",
        "The report was reverse-engineered from the current source tree. Controller annotations were used to build the API inventory; Entity mappings were checked against Flyway migrations; React API clients and protected routes were compared with server authorization. Unsupported claims, fabricated performance measurements, and unverified production deployment details are intentionally excluded.",
    ])
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Keywords: ")
    run.bold = True
    paragraph.add_run("cinema booking, React, Spring Boot, PostgreSQL, JWT, WebSocket, transaction, payment gateway.")
    doc.add_page_break()

    doc.add_heading("DANH MỤC TỪ VIẾT TẮT", level=1)
    academic_table(doc, "Bảng 0.1. Danh mục từ viết tắt", ["Từ viết tắt", "Tiếng Anh", "Ý nghĩa trong đề tài"], [
        ["API", "Application Programming Interface", "Giao diện lập trình giữa React và Spring Boot hoặc giữa hệ thống với cổng thanh toán."],
        ["DTO", "Data Transfer Object", "Đối tượng request/response tách khỏi Entity JPA."],
        ["ERD", "Entity Relationship Diagram", "Sơ đồ quan hệ các bảng PostgreSQL."],
        ["FK", "Foreign Key", "Khóa ngoại bảo đảm quan hệ tham chiếu."],
        ["HMAC", "Hash-based Message Authentication Code", "Cơ chế xác minh tính toàn vẹn của một số payload thanh toán."],
        ["HTTP", "Hypertext Transfer Protocol", "Giao thức nền của REST API."],
        ["JWT", "JSON Web Token", "Định dạng token ký số dùng cho access/refresh trong hệ thống."],
        ["JPA", "Jakarta Persistence API", "Chuẩn ánh xạ đối tượng-quan hệ được Hibernate hiện thực."],
        ["PK", "Primary Key", "Khóa chính định danh duy nhất bản ghi."],
        ["QR", "Quick Response", "Mã vé điện tử và mã chuyển khoản ngân hàng."],
        ["RBAC", "Role-Based Access Control", "Mô hình gán permission thông qua role."],
        ["REST", "Representational State Transfer", "Phong cách thiết kế API tài nguyên qua HTTP."],
        ["STOMP", "Simple Text Oriented Messaging Protocol", "Giao thức nhắn tin dùng trên WebSocket."],
        ["TPS", "Transactions Per Second", "Chỉ số tải chỉ được báo cáo khi có đo thực tế."],
        ["UML", "Unified Modeling Language", "Ngôn ngữ mô hình hóa dùng cho các sơ đồ PlantUML."],
    ])
    doc.add_page_break()


def add_formal_pre_toc_pages(doc) -> None:
    """Front matter placed before the table of contents in the formal thesis."""
    doc.add_heading("LỜI CAM ĐOAN", level=1)
    body_many(doc, [
        "Tôi cam đoan khóa luận với đề tài “Xây dựng hệ thống đặt vé xem phim trực tuyến CinemaBooking” là kết quả của quá trình tìm hiểu, phân tích, thiết kế và hiện thực của cá nhân dưới sự hướng dẫn của giảng viên. Những mô tả về kiến trúc, dữ liệu, API và nghiệp vụ trong báo cáo được đối chiếu với mã nguồn, cấu hình, migration và test của dự án. Các kiến thức kế thừa từ tài liệu bên ngoài đều được trích dẫn trong phần tài liệu tham khảo.",
        "Các kết quả kiểm thử và nhận định kỹ thuật trong khóa luận không được cố ý làm sai lệch. Những nội dung chưa có bằng chứng thực nghiệm, chẳng hạn số liệu kiểm thử tải hoặc cấu hình triển khai production, được trình bày dưới dạng kịch bản dự kiến hay hướng phát triển. Tôi chịu trách nhiệm về tính trung thực của nội dung báo cáo theo quy định của cơ sở đào tạo.",
    ])
    body(doc, "[ĐỊA DANH], ngày ..... tháng ..... năm 2026", indent=False)
    body(doc, "Sinh viên thực hiện", indent=False)
    body(doc, "[HỌ VÀ TÊN SINH VIÊN]", indent=False)
    doc.add_page_break()

    doc.add_heading("LỜI CẢM ƠN", level=1)
    body_many(doc, [
        "Tôi xin trân trọng cảm ơn [HỌ VÀ TÊN GIẢNG VIÊN HƯỚNG DẪN] đã dành thời gian hướng dẫn, nhận xét và giúp tôi điều chỉnh phạm vi của đề tài. Những góp ý của giảng viên là cơ sở để tôi xem xét lại cách tổ chức hệ thống, kiểm tra các trường hợp cạnh tranh khi giữ ghế và trình bày kết quả theo bằng chứng từ chương trình đã xây dựng.",
        "Tôi cảm ơn các thầy cô của [TÊN KHOA/BỘ MÔN] đã trang bị kiến thức về lập trình, cơ sở dữ liệu, phân tích thiết kế hệ thống và an toàn thông tin trong suốt quá trình học tập. Tôi cũng biết ơn gia đình và bạn bè đã động viên, góp ý cho giao diện và hỗ trợ kiểm tra một số hành trình sử dụng. Tôi mong tiếp tục nhận được nhận xét của hội đồng để nhận ra những điểm còn hạn chế và hoàn thiện đề tài tốt hơn.",
    ])
    body(doc, "[SINH VIÊN ĐIỀU CHỈNH NỘI DUNG CẢM ƠN THEO THỰC TẾ TRƯỚC KHI NỘP]", indent=False)
    doc.add_page_break()


def add_formal_abbreviations(doc) -> None:
    doc.add_heading("DANH MỤC TỪ VIẾT TẮT", level=1)
    academic_table(doc, "Bảng 0.1. Danh mục từ viết tắt", ["Từ viết tắt", "Tiếng Anh", "Ý nghĩa trong đề tài"], [
        ["API", "Application Programming Interface", "Giao diện lập trình giữa React, Spring Boot và các hệ thống ngoài."],
        ["DTO", "Data Transfer Object", "Đối tượng request/response tách khỏi Entity JPA."],
        ["ERD", "Entity Relationship Diagram", "Sơ đồ quan hệ các bảng PostgreSQL."],
        ["FK", "Foreign Key", "Khóa ngoại bảo đảm quan hệ tham chiếu."],
        ["HMAC", "Hash-based Message Authentication Code", "Cơ chế kiểm tra tính toàn vẹn và nguồn gốc payload."],
        ["HTTP", "Hypertext Transfer Protocol", "Giao thức nền của REST API."],
        ["JWT", "JSON Web Token", "Định dạng token ký số dùng cho xác thực."],
        ["JPA", "Jakarta Persistence API", "Chuẩn ánh xạ đối tượng - quan hệ được Hibernate hiện thực."],
        ["PK", "Primary Key", "Khóa chính định danh duy nhất bản ghi."],
        ["QR", "Quick Response", "Mã vé điện tử và mã chuyển khoản ngân hàng."],
        ["RBAC", "Role-Based Access Control", "Mô hình phân quyền dựa trên vai trò và permission."],
        ["REST", "Representational State Transfer", "Phong cách thiết kế API tài nguyên qua HTTP."],
        ["STOMP", "Simple Text Oriented Messaging Protocol", "Giao thức nhắn tin dùng trên WebSocket."],
        ["TPS", "Transactions Per Second", "Chỉ số tải chỉ được báo cáo khi có đo thực tế."],
        ["UML", "Unified Modeling Language", "Ngôn ngữ mô hình hóa dùng cho sơ đồ hệ thống."],
    ])
    doc.add_page_break()


def add_formal_summary_and_preface(doc) -> None:
    doc.add_heading("TÓM TẮT KHÓA LUẬN", level=1)
    body_many(doc, [
        "CinemaBooking được xây dựng nhằm hỗ trợ tra cứu lịch chiếu, lựa chọn ghế, thanh toán và sử dụng vé điện tử trên một ứng dụng web thống nhất. Giao diện hệ thống được phát triển bằng ReactJS và TypeScript; các dịch vụ nghiệp vụ được hiện thực bằng Java Spring Boot; PostgreSQL đảm nhiệm lưu trữ dữ liệu. Các yêu cầu thông thường được trao đổi qua RESTful API, trong khi thay đổi của sơ đồ ghế được cập nhật đến trình duyệt bằng WebSocket/STOMP.",
        "Khóa luận tập trung phân tích tính nhất quán của trạng thái ghế từ thời điểm khách hàng lựa chọn cho đến khi giao dịch thanh toán được xác nhận. Thao tác giữ ghế và xác nhận ghế được xử lý trong transaction cơ sở dữ liệu, kết hợp cơ chế khóa bản ghi, trường phiên bản và các ràng buộc duy nhất. Các tác vụ định kỳ có nhiệm vụ thu hồi ghế hoặc đơn đặt vé đã hết hạn. VNPay và SePay được tổ chức thành hai quy trình riêng vì khác nhau về cách khởi tạo giao dịch, xác minh chữ ký và tiếp nhận kết quả thanh toán.",
        "Hệ thống phục vụ ba nhóm người dùng gồm khách hàng, nhân viên rạp và quản trị viên. Bên cạnh quy trình mua vé, hệ thống còn có xác thực JWT, quản lý phiên làm mới, phân quyền theo vai trò và phạm vi rạp, phát hành vé QR, soát vé, gửi thư điện tử, quản lý yêu cầu hoàn tiền và ghi nhận nhật ký hoạt động. Chức năng hoàn tiền được giới hạn ở việc tạo, theo dõi và ghi nhận kết quả do người vận hành cung cấp; hệ thống chưa tự động chuyển tiền qua API nhà cung cấp. Kết quả phân tích được đối chiếu với các thành phần hiện thực ở cả phía giao diện, phía máy chủ, cấu trúc dữ liệu và mã kiểm thử. Những nội dung chưa có số liệu thực nghiệm, đặc biệt là khả năng xử lý tải đồng thời ở quy mô lớn và triển khai nhiều máy chủ, được xác định là giới hạn của đề tài.",
    ])
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run("Từ khóa: ")
    run.bold = True
    paragraph.add_run("đặt vé xem phim, ReactJS, Spring Boot, PostgreSQL, JWT, WebSocket, xử lý đồng thời, VNPay, SePay, vé QR.")
    doc.add_page_break()

    doc.add_heading("LỜI NÓI ĐẦU", level=1)
    body_many(doc, [
        "Một lần mua vé trực tuyến bắt đầu từ việc xem lịch nhưng chỉ hoàn tất khi ghế đã được xác nhận và khách hàng nhận được vé có thể sử dụng tại rạp. Giữa hai thời điểm đó, hệ thống phải xử lý nhiều thay đổi: ghế có thể vừa được người khác giữ, mã khuyến mãi làm thay đổi số tiền, cổng thanh toán phản hồi chậm hoặc người dùng đóng trình duyệt trước khi giao dịch kết thúc. Vì thế, bài toán đặt vé không chỉ là hiển thị danh sách phim và lưu một đơn hàng.",
        "CinemaBooking được lựa chọn làm đề tài để tìm hiểu cách tổ chức một quy trình bán vé có trạng thái và có nhiều tác nhân tham gia. Khách hàng cần thao tác ngắn gọn; nhân viên phải soát vé đúng rạp, đúng suất; quản trị viên cần theo dõi dữ liệu vận hành mà không làm mất tính nhất quán của đơn đã thanh toán. Việc hiện thực đồng thời giao diện, dịch vụ, cơ sở dữ liệu, kết nối thanh toán và kênh cập nhật ghế tạo điều kiện vận dụng các kiến thức đã học vào một sản phẩm có luồng nghiệp vụ liên tục.",
        "Báo cáo được trình bày trong năm chương. Chương 1 xác định vấn đề, mục tiêu và phạm vi. Chương 2 nêu các cơ sở lý thuyết được sử dụng. Chương 3 chuyển yêu cầu nghiệp vụ thành mô hình dữ liệu, API và các sơ đồ thiết kế. Chương 4 trình bày cách hiện thực những thành phần chính và kế hoạch kiểm thử. Chương 5 đánh giá kết quả, những giới hạn còn tồn tại và hướng phát triển tiếp theo. Hai phụ lục cuối báo cáo cung cấp ma trận truy vết yêu cầu và các kịch bản kiểm thử trọng yếu.",
    ])


def expand_chapter_one(doc) -> None:
    doc.add_heading("1.7. Ý nghĩa khoa học và thực tiễn", level=2)
    body_many(doc, [
        "Về mặt học thuật, đề tài kết nối nhiều nội dung thường được học tách rời: mô hình dữ liệu quan hệ, kiến trúc phân lớp, bảo mật web, xử lý giao dịch, lập trình giao diện, tích hợp hệ thống ngoài và kiểm thử. Một thao tác tưởng như đơn giản là bấm chọn ghế tạo ra chuỗi vấn đề liên quan đến đồng thời: client cần phản hồi nhanh, server phải xác định quyền sở hữu hold, database phải ngăn hai transaction bán cùng một tài nguyên, scheduler phải thu hồi hold hết hạn và các client khác phải nhận được thay đổi sau commit. Vì vậy, CinemaBooking là ngữ cảnh phù hợp để chứng minh khả năng vận dụng kiến thức hơn là chỉ trình diễn CRUD.",
        "Về mặt thực tiễn, hệ thống mô hình hóa ba nhóm người dùng có nhu cầu khác nhau. Khách hàng cần hành trình mua vé rõ ràng và trạng thái thanh toán đáng tin cậy. Nhân viên cần vận hành trong đúng rạp được phân công và kiểm tra vé theo ngữ cảnh rạp-suất chiếu. Quản trị viên cần quản lý dữ liệu nền, người dùng, khuyến mãi, doanh thu, payment event, yêu cầu refund và audit. Việc tách permission khỏi phạm vi rạp phản ánh yêu cầu kiểm soát truy cập theo cả chức năng lẫn dữ liệu.",
        "Đề tài còn có giá trị ở khả năng duy trì mối liên hệ giữa yêu cầu, thiết kế và thành phần hiện thực. Các luồng nghiệp vụ được mô tả thống nhất với kiến trúc phân lớp và lược đồ dữ liệu, nhờ đó những thay đổi trong chương trình có thể được phản ánh có hệ thống vào tài liệu thiết kế và kiểm thử.",
    ])
    doc.add_heading("1.8. Tiểu kết chương 1", level=2)
    body(doc, "Chương 1 đã trình bày lý do lựa chọn đề tài, mục tiêu, đối tượng nghiên cứu, phạm vi và phương pháp thực hiện hệ thống CinemaBooking. Bài toán không chỉ yêu cầu cung cấp chức năng tra cứu và đặt vé mà còn phải bảo đảm tính nhất quán của ghế, an toàn trong thanh toán và kiểm soát quyền truy cập theo từng nhóm người sử dụng. Những nội dung này là cơ sở để lựa chọn công nghệ và xây dựng phương án thiết kế được trình bày trong các chương tiếp theo.")


def expand_chapter_two(doc) -> None:
    doc.add_heading("2.10. Đánh giá và lựa chọn công nghệ", level=2)
    body_many(doc, [
        "Việc lựa chọn công nghệ cho CinemaBooking được xem xét theo đặc điểm của bài toán thay vì chỉ dựa trên mức độ phổ biến của công cụ. Hệ thống cần giao diện có nhiều trạng thái tương tác, dữ liệu giao dịch có quan hệ chặt chẽ, khả năng kiểm soát đồng thời khi giữ ghế, cơ chế phân quyền theo vai trò và kênh cập nhật trạng thái theo thời gian thực. Bên cạnh yêu cầu chức năng, phạm vi khóa luận cũng đòi hỏi giải pháp có thể triển khai, kiểm thử và bảo trì với nguồn lực phù hợp.",
        "Bảng 2.5 đối chiếu các công nghệ được lựa chọn với một số phương án thay thế. Kết quả cho thấy không có công nghệ tối ưu trong mọi hoàn cảnh; mỗi lựa chọn gắn với một nhóm yêu cầu cụ thể. Các giải pháp phân tán như Redis, message broker hoặc distributed lock có thể cần thiết khi hệ thống vận hành trên nhiều instance, nhưng chưa tạo ra lợi ích tương xứng với độ phức tạp trong phạm vi triển khai hiện tại.",
    ])
    academic_table(doc, "Bảng 2.5. Đánh giá lựa chọn công nghệ và phương án thay thế", ["Bài toán", "Giải pháp lựa chọn", "Phương án thay thế", "Cơ sở lựa chọn"], [
        ["Xây dựng front-end", "React, TypeScript và Vite", "Angular, Vue hoặc server-rendered template", "React phù hợp với giao diện SPA nhiều trạng thái; TypeScript hỗ trợ kiểm soát contract và Vite rút ngắn chu trình phát triển."],
        ["Xây dựng back-end", "Spring Boot theo kiến trúc phân lớp", "Node.js/NestJS hoặc .NET", "Hệ sinh thái Spring tích hợp transaction, security, validation, JPA và WebSocket phù hợp với nghiệp vụ đặt vé."],
        ["Lưu trữ dữ liệu", "PostgreSQL", "MySQL hoặc cơ sở dữ liệu tài liệu", "Mô hình quan hệ, constraint, transaction, khóa bản ghi và partial index phù hợp với booking và payment."],
        ["Quản lý phiên đăng nhập", "JWT access token kết hợp refresh token có luân chuyển", "HTTP session thuần túy", "Access token giảm phụ thuộc vào session lookup; refresh token phía server vẫn hỗ trợ quản lý và thu hồi phiên."],
        ["Kiểm soát đồng thời", "Khóa cơ sở dữ liệu, cập nhật nguyên tử và constraint", "Chỉ dùng optimistic retry hoặc distributed lock", "PostgreSQL là nguồn dữ liệu có thẩm quyền; khóa và constraint bảo vệ trực tiếp tài nguyên ghế trong kiến trúc hiện tại."],
        ["Đồng bộ trạng thái ghế", "STOMP trên WebSocket", "Polling hoặc Server-Sent Events", "Topic theo suất chiếu giảm request lặp và chuyển thay đổi đến nhiều client; REST vẫn cung cấp dữ liệu ban đầu."],
        ["Quản lý lược đồ", "Flyway migration", "Hibernate auto-DDL hoặc Liquibase", "Migration SQL cho phép kiểm soát rõ constraint, index và thứ tự nâng cấp PostgreSQL."],
        ["Bộ nhớ đệm", "Caffeine cục bộ", "Redis phân tán", "Cache trong tiến trình đơn giản và phù hợp triển khai một instance; Redis là hướng mở rộng khi có nhiều instance."],
    ])
    doc.add_heading("2.11. Tiểu kết chương 2", level=2)
    body(doc, "Chương 2 đã trình bày các cơ sở lý thuyết và công nghệ được sử dụng để xây dựng CinemaBooking, từ kiến trúc client-server, React, Spring Boot, PostgreSQL và RESTful API đến xác thực, phân quyền, xử lý giao dịch và WebSocket. Việc phân tích không tách rời đặc điểm của bài toán: dữ liệu ghế cần nhất quán khi có nhiều yêu cầu đồng thời, phiên đăng nhập phải có khả năng thu hồi, còn trạng thái giao diện cần được cập nhật sau khi dữ liệu đã được ghi nhận thành công. Trên nền tảng đó, Chương 3 chuyển sang phân tích yêu cầu, mô hình hóa dữ liệu và thiết kế các luồng nghiệp vụ của hệ thống.")
def expand_chapter_three(doc) -> None:
    doc.add_heading("3.9. Thiết kế các luồng nghiệp vụ trọng yếu", level=2)
    body(doc, "Các luồng dưới đây được trình bày theo thứ tự từ hành động của tác nhân, điều kiện xử lý, thay đổi trạng thái đến kết quả trả về. Sơ đồ tuần tự làm rõ sự phối hợp giữa client, back-end, cơ sở dữ liệu và hệ thống bên ngoài; sơ đồ hoạt động được sử dụng tại những quy trình có nhiều nhánh quyết định. Chi tiết tổ chức lớp và mã nguồn được trình bày tại Chương 4.")

    doc.add_heading("3.9.1. Xác thực và quản lý phiên đăng nhập", level=3)
    body_many(doc, [
        "Khi đăng ký, người dùng cung cấp tên đăng nhập, email và mật khẩu đáp ứng chính sách của hệ thống. Sau khi kiểm tra tính duy nhất, mật khẩu được băm và tài khoản được gán vai trò khách hàng. Token xác thực email có thời hạn được gửi qua thư điện tử, trong khi cơ sở dữ liệu chỉ lưu giá trị băm. Quy trình đặt lại mật khẩu áp dụng nguyên tắc tương tự và sử dụng phản hồi thống nhất để hạn chế việc dò tìm tài khoản qua địa chỉ email.",
        "Đăng nhập có thể được thực hiện bằng mật khẩu hoặc Google ID Token. Phía server xác minh thông tin, trạng thái tài khoản và giới hạn tần suất trước khi phát access token cùng một phiên refresh token. Access token có thời hạn ngắn và được gửi trong Bearer header; refresh token được lưu trong cookie HttpOnly, đồng thời bản băm của nó được quản lý tại server để hỗ trợ thu hồi phiên.",
        "Khi access token hết hạn, client phối hợp các request đang chờ để chỉ thực hiện một lần làm mới tại cùng thời điểm. Server khóa phiên refresh token, kiểm tra thời hạn, trạng thái thu hồi và phiên bản xác thực, sau đó vô hiệu hóa token cũ trước khi phát token thay thế. Đăng xuất hoặc thu hồi phiên làm mất hiệu lực refresh token tương ứng và ngăn token đó tiếp tục được sử dụng.",
    ])
    diagram_placeholder(doc, "[CHÈN FILE: Sequence Diagram - Access Token hết hạn và Refresh Token Rotation.svg]", "Sơ đồ thể hiện access token hết hạn, cơ chế phối hợp refresh phía client, xác minh phiên và luân chuyển refresh token tại server.", "Hình 3.7. Sơ đồ tuần tự của quy trình làm mới access token")

    doc.add_heading("3.9.2. Giữ ghế và tạo đơn đặt vé", level=3)
    body_many(doc, [
        "Khách hàng chọn một suất chiếu và mở sơ đồ ghế. Client tải trạng thái hiện tại qua RESTful API, sau đó đăng ký kênh sự kiện của suất chiếu để tiếp nhận các thay đổi phát sinh. Thông tin hiển thị hỗ trợ lựa chọn, nhưng quyền giữ ghế chỉ được xác lập khi server chấp nhận yêu cầu.",
        "Yêu cầu giữ ghế gồm mã suất chiếu và danh sách ghế, không chứa giá vé do client tự tính. Server kiểm tra tài khoản, thời gian mở bán, giới hạn tần suất và sự tồn tại của các ghế, sau đó khóa các dòng trạng thái theo thứ tự ổn định. Chỉ ghế còn khả dụng mới được chuyển sang HOLD cùng người giữ và thời điểm hết hạn. Nếu có yêu cầu cạnh tranh, transaction xử lý sau sẽ đọc được trạng thái mới và trả về lỗi nghiệp vụ thay vì ghi đè kết quả trước.",
        "Khi tạo đơn, hệ thống kiểm tra lại lượt giữ còn hiệu lực và thuộc đúng khách hàng. Tổng tiền được tính từ giá cơ bản của suất chiếu, hệ số loại ghế và khuyến mãi hợp lệ. Đơn được tạo ở trạng thái PENDING, kèm các chi tiết ghế và thời hạn thanh toán. Cơ sở dữ liệu giới hạn số đơn chờ đang hoạt động cho cùng khách hàng và suất chiếu, nhờ đó thao tác gửi lặp không tạo nhiều đơn cạnh tranh.",
    ])
    diagram_placeholder(doc, "[CHÈN FILE: 3.3 Sequence Diagram - Giữ ghế tạm thời.svg]", "Sơ đồ thể hiện yêu cầu giữ ghế, transaction khóa trạng thái, cập nhật người giữ và thời hạn, sau đó phát sự kiện khi commit thành công.", "Hình 3.8. Sơ đồ tuần tự của quy trình giữ ghế tạm thời")
    diagram_placeholder(doc, "[CHÈN FILE: 3.4 Activity Diagram - Giữ ghế và tạo booking PENDING.svg]", "Sơ đồ thể hiện các điểm quyết định từ kiểm tra ghế, giữ ghế đến tạo đơn PENDING, bao gồm nhánh ghế không khả dụng hoặc lượt giữ đã hết hạn.", "Hình 3.9. Sơ đồ hoạt động của quy trình giữ ghế và tạo đơn đặt vé")

    doc.add_heading("3.9.3. Áp dụng khuyến mãi và xác định số tiền", level=3)
    body_many(doc, [
        "Khuyến mãi có thể giảm theo tỷ lệ phần trăm hoặc số tiền cố định và đi kèm các điều kiện về thời gian, giá trị đơn tối thiểu, mức giảm tối đa và giới hạn sử dụng. Những điều kiện này được kiểm tra tại server; giao diện chỉ hiển thị kết quả tạm tính, mức giảm và tổng thanh toán do server trả về.",
        "Lượt sử dụng được tạm giữ trong thời gian đơn chờ thanh toán và được giải phóng nếu đơn bị hủy, thất bại hoặc hết hạn. Khi một mã thanh toán đã được tạo với số tiền cố định, khuyến mãi không thể thay đổi trực tiếp. Người dùng cần kết thúc giao dịch chờ và tạo lại mã thanh toán để số tiền trên đơn thống nhất với số tiền gửi đến cổng thanh toán hoặc ngân hàng.",
    ])

    doc.add_heading("3.9.4. Thanh toán qua VNPay", level=3)
    body_many(doc, [
        "Khách hàng chọn VNPay để khởi tạo giao dịch cho đơn đang chờ. Server kiểm tra quyền sở hữu, trạng thái, thời hạn và số tiền, sau đó tạo mã giao dịch cùng URL thanh toán có chữ ký. Bản ghi payment PENDING được lưu trước khi trình duyệt chuyển sang trang của VNPay.",
        "Khi giao dịch kết thúc, VNPay chuyển kết quả về callback của hệ thống. Back-end xác minh secure hash, mã phản hồi, số tiền, mã giao dịch và trạng thái hiện tại của booking, payment. Nếu kết quả hợp lệ, payment và booking được xác nhận, ghế chuyển sang BOOKED và vé được phát hành. Callback lặp chỉ trả lại kết quả đã xử lý, không tạo thêm vé hoặc lặp lại thay đổi trạng thái.",
    ])
    diagram_placeholder(doc, "[CHÈN FILE: 3.5 Sequence Diagram - Thanh toán VNPay.svg]", "Sơ đồ thể hiện quá trình khởi tạo, chuyển hướng, callback, xác minh chữ ký, đối chiếu số tiền và phát hành vé theo nguyên tắc lũy đẳng.", "Hình 3.10. Sơ đồ tuần tự của quy trình thanh toán qua VNPay")

    doc.add_heading("3.9.5. Thanh toán chuyển khoản qua SePay/VietQR", level=3)
    body_many(doc, [
        "Với SePay/VietQR, hệ thống tạo mã QR chứa ngân hàng, số tài khoản, số tiền và nội dung định danh giao dịch. Khách hàng chuyển khoản bằng ứng dụng ngân hàng. Việc client kiểm tra định kỳ chỉ giúp quan sát trạng thái; hành động xác nhận từ phía người dùng không thể tự chuyển đơn sang thành công.",
        "Sau khi ngân hàng ghi nhận tiền vào, SePay gửi webhook đến back-end. Dữ liệu được xác minh bằng API key hoặc HMAC theo cấu hình, sau đó đối chiếu nội dung chuyển khoản, mã giao dịch, số tiền và trạng thái hiện tại. Chỉ webhook hợp lệ mới đi vào quy trình xác nhận dùng chung để cập nhật payment, booking, ghế và phát hành vé. Sự kiện không hợp lệ được lưu phục vụ theo dõi mà không làm thay đổi kết quả bán vé.",
    ])
    diagram_placeholder(doc, "[CHÈN FILE: 3.6 Sequence Diagram - Luồng thanh toán SePay VietQR.svg]", "Sơ đồ phân biệt bước tạo QR, thao tác chuyển khoản, webhook từ SePay và quá trình đối chiếu trước khi xác nhận đơn.", "Hình 3.11. Sơ đồ tuần tự của quy trình thanh toán qua SePay/VietQR")

    doc.add_heading("3.9.6. Đồng bộ trạng thái ghế theo thời gian thực", level=3)
    body_many(doc, [
        "RESTful API cung cấp ảnh chụp trạng thái ghế tại thời điểm người dùng mở màn hình. Sau đó, client đăng ký topic STOMP theo mã suất chiếu để nhận các sự kiện HOLD, AVAILABLE hoặc BOOKED. Sự kết hợp này giúp giao diện có dữ liệu ban đầu đầy đủ và tiếp tục phản ánh các thay đổi mà không phải liên tục gửi request thăm dò.",
        "Sự kiện chỉ được phát sau khi transaction cập nhật dữ liệu đã commit. Nếu transaction bị rollback, không có thông điệp thay đổi được gửi đến client. PostgreSQL vì vậy vẫn là nguồn dữ liệu có thẩm quyền; WebSocket chỉ làm nhiệm vụ phân phối thay đổi đã được xác nhận đến các trình duyệt đang kết nối.",
    ])
    diagram_placeholder(doc, "[CHÈN FILE: 3.14 Sequence Diagram - Đồng bộ sơ đồ ghế qua WebSocket STOMP.svg]", "Sơ đồ thể hiện REST snapshot, kết nối STOMP, sự kiện thay đổi trạng thái và cập nhật giao diện sau commit.", "Hình 3.12. Sơ đồ tuần tự của quy trình đồng bộ sơ đồ ghế qua WebSocket/STOMP")

    doc.add_heading("3.9.7. Phát hành và soát vé bằng mã QR", level=3)
    body_many(doc, [
        "Sau khi đơn được thanh toán thành công, mỗi chi tiết ghế có tối đa một vé điện tử. Mã QR chứa payload có phiên bản, định danh và chữ ký HMAC; thông tin phim, rạp, phòng, thời gian, ghế và địa chỉ được lấy từ dữ liệu booking để hiển thị trên vé, không được dùng thay thế cho việc xác minh tại server.",
        "Trước khi mở camera hoặc đọc mã từ tệp ảnh, nhân viên phải chọn rạp và suất chiếu đang soát. Với tài khoản nhân viên, danh sách rạp chỉ gồm những địa điểm được quản trị viên phân công; tài khoản quản trị có thể lựa chọn trong toàn hệ thống. Danh sách suất chiếu tiếp tục được thu hẹp theo rạp và cửa sổ check-in. Theo cấu hình của phiên bản được đánh giá, một suất được mở soát từ 60 phút trước giờ bắt đầu đến 30 phút sau giờ bắt đầu. Khi không có suất phù hợp, giao diện thông báo rõ trạng thái thay vì cho phép quét trong một ngữ cảnh không xác định.",
        "Yêu cầu soát vé gửi đồng thời mã QR, mã rạp và mã suất chiếu. Phía máy chủ lần lượt kiểm tra tính hợp lệ của mã, rạp và suất được chọn, phạm vi phụ trách của nhân viên, trạng thái vé, trạng thái đơn đặt vé và thời gian check-in. Chỉ khi toàn bộ điều kiện được thỏa mãn, vé ACTIVE mới chuyển sang USED và được lưu thời điểm cùng người thực hiện. Vé đã sử dụng, sai rạp, sai suất, chưa đến giờ hoặc quá thời gian soát đều nhận kết quả riêng và không bị thay đổi ngoài ý muốn. Khi hai thiết bị quét đồng thời cùng một vé, khóa dữ liệu bảo đảm chỉ một thao tác có thể hoàn tất việc sử dụng vé.",
    ])
    diagram_placeholder(doc, "[CHÈN FILE: 3.7 2 Sequence Diagram - Xác thực và check-in vé QR.svg]", "Sơ đồ thể hiện việc kiểm tra chữ ký, trạng thái vé, booking, phạm vi rạp, suất chiếu và cửa sổ check-in trước khi chuyển vé sang USED.", "Hình 3.13. Sơ đồ tuần tự của quy trình xác thực và soát vé QR")

    doc.add_heading("3.9.8. Hủy suất chiếu và quản lý yêu cầu hoàn tiền", level=3)
    body_many(doc, [
        "Người thao tác chỉ có thể hủy suất chiếu khi có quyền phù hợp và, đối với nhân viên, suất thuộc rạp được phân công. Hệ thống từ chối hủy khi điều kiện vận hành không cho phép, chẳng hạn đã có vé được sử dụng. Sau khi hủy, suất không còn được mở bán; đơn PENDING được kết thúc và các ghế đang giữ được giải phóng.",
        "Đối với đơn đã thanh toán, hệ thống chuyển booking và payment sang trạng thái chờ xử lý, đồng thời tạo một yêu cầu refund nếu chưa có yêu cầu đang hoạt động. Yêu cầu này được theo dõi độc lập với thao tác hủy suất. Chỉ chủ thể có permission PAYMENT_REFUND mới được ghi nhận kết quả thành công hoặc thất bại; STAFF hiện không có permission này. Phiên bản hiện tại không gọi API nhà cung cấp để tự động hoàn tiền. Sau khi việc hoàn trả được xử lý theo quy trình bên ngoài hệ thống, quản trị viên ghi nhận kết quả; khi đó payment và booking mới chuyển sang REFUNDED, hoặc trạng thái thất bại được lưu để tiếp tục đối soát.",
    ])
    diagram_placeholder(doc, "[CHÈN FILE: 3.15 Activity Diagram - Hủy suất chiếu và tạo yêu cầu hoàn tiền.svg]", "Sơ đồ phân biệt nhánh đơn PENDING, đơn SUCCESS, vé đã sử dụng và điều kiện tạo yêu cầu hoàn tiền.", "Hình 3.14. Sơ đồ hoạt động của quy trình hủy suất chiếu và tạo yêu cầu hoàn tiền")
    diagram_placeholder(doc, "[CHÈN FILE: 3.15 Activity Diagram - Xử lý hoàn tiền.svg]", "Sơ đồ thể hiện việc quản trị viên ghi nhận kết quả xử lý bên ngoài hệ thống qua các trạng thái PENDING, PROCESSING, SUCCESS hoặc FAILED; sơ đồ không biểu diễn một API hoàn tiền tự động.", "Hình 3.15. Sơ đồ hoạt động của quy trình quản lý yêu cầu hoàn tiền")

    doc.add_heading("3.9.9. Quản lý lịch chiếu và ngăn xung đột phòng chiếu", level=3)
    body_many(doc, [
        "Mỗi phòng chiếu chỉ có thể phục vụ một suất tại một thời điểm. Khi tạo hoặc cập nhật lịch, hệ thống yêu cầu thời gian kết thúc phải sau thời gian bắt đầu và kiểm tra các suất còn hiệu lực trong cùng phòng. Khoảng kiểm tra được mở rộng thêm 15 phút trước giờ bắt đầu và 15 phút sau giờ kết thúc của suất mới để dành thời gian ổn định phòng, vệ sinh và chuẩn bị cho lượt chiếu tiếp theo. Vì vậy, hai suất dù không đè trực tiếp lên nhau vẫn bị từ chối nếu khoảng cách giữa chúng ngắn hơn thời gian chuyển phòng đã quy định.",
        "Lịch dành cho khách hàng được tách khỏi danh sách vận hành. Các màn hình công khai chỉ nhận suất ở trạng thái UPCOMING, thuộc phim, phòng và rạp còn hoạt động, bắt đầu sau thời điểm ngừng bán vé và nằm trong khoảng ngày được công bố. Cấu hình hiện tại ngừng nhận đặt vé trước giờ chiếu 15 phút và công bố lịch trong bảy ngày. Suất đã qua, đang chiếu, đã kết thúc, bị hủy hoặc nằm ngoài khoảng mở bán không xuất hiện trong hành trình mua vé, nhưng vẫn được giữ trong dữ liệu quản trị để phục vụ theo dõi và đối chiếu.",
        "Trạng thái của suất được cập nhật theo thời gian từ UPCOMING sang ONGOING và ENDED bởi tác vụ định kỳ. Cách tổ chức này giúp giao diện quản lý phản ánh đúng vòng đời suất chiếu mà không phụ thuộc vào việc có người mở trang tại thời điểm chuyển trạng thái. Suất CANCELLED được bảo toàn như một trạng thái kết thúc riêng để lịch sử đơn đặt vé, giao dịch và yêu cầu xử lý liên quan vẫn có thể được truy vết.",
    ])

    doc.add_heading("3.9.10. Tìm rạp gần vị trí hiện tại", level=3)
    body_many(doc, [
        "Trường hợp sử dụng bắt đầu khi khách hàng chủ động chọn Gần tôi tại trang chủ hoặc trang bản đồ rạp. Client chỉ yêu cầu Browser Geolocation API tại thời điểm này; hệ thống không tự động xin quyền ngay khi mở trang [17]. Nếu trình duyệt trả về vị trí hợp lệ, client dùng danh sách rạp công khai từ GET `/api/v1/cinemas/map`. Mỗi phần tử bản đồ gồm mã rạp, tên, địa chỉ, thành phố, vĩ độ, kinh độ và trạng thái hoạt động. Back-end chỉ lấy rạp đang hoạt động, chưa bị xóa và có đủ tọa độ; client tiếp tục loại tọa độ không hữu hạn hoặc nằm ngoài miền vĩ độ [-90, 90], kinh độ [-180, 180].",
        "Khoảng cách từ người dùng đến từng rạp được tính tại trình duyệt bằng công thức Haversine [18]. Utility sử dụng hằng số bán kính 6.371 km; danh sách được sắp tăng dần theo khoảng cách, sau đó theo tên rạp nếu kết quả bằng nhau. Khu vực lịch chiếu chọn rạp gần nhất và ưu tiên thành phố tương ứng; trang bản đồ di chuyển khung nhìn đến vị trí người dùng, đánh dấu điểm hiện tại, mở rạp gần nhất và hiển thị khoảng cách trên từng mục. Hằng số 10 km được dùng để đếm, gắn nhãn các rạp gần, không phải điều kiện loại bỏ mọi rạp ở xa. Khi không có rạp trong bán kính, danh sách gần nhất vẫn được giữ để khách hàng có thể tiếp tục tra cứu.",
        "Luồng thay thế được duy trì cho mọi trường hợp không thể định vị. Nếu người dùng từ chối quyền, yêu cầu hết thời gian chờ 10 giây, trình duyệt không hỗ trợ hoặc không xác định được vị trí, client hiển thị thông báo phù hợp và vẫn cho phép chọn thành phố, rạp thủ công. Bản ghi rạp có tọa độ sai không được tạo marker hoặc ưu tiên trong phép sắp xếp. Endpoint GET `/api/v1/cinemas/nearest` có sẵn ở Back-end để trả một số rạp gần theo truy vấn Haversine, nhưng luồng React hiện tại không gọi endpoint này; việc tính ở client giúp dùng chung dữ liệu bản đồ đã tải và tránh gửi vị trí chính xác đến Back-end của CinemaBooking.",
        "Về quyền riêng tư, latitude và longitude của người dùng chỉ tồn tại trong state của component trong phiên hiển thị. Client không ghi hai giá trị này vào localStorage, không có API lưu chúng vào cơ sở dữ liệu và không đưa chúng vào hồ sơ tài khoản. Chỉ tên thành phố ưu tiên cùng nguồn lựa chọn `manual` hoặc `location` được lưu để duy trì trải nghiệm. Quy tắc này vừa đủ cho mục tiêu gợi ý rạp mà không tạo lịch sử vị trí chính xác không cần thiết.",
    ])

    doc.add_heading("3.10. Thiết kế xử lý đồng thời và bảo đảm toàn vẹn dữ liệu", level=2)
    body_many(doc, [
        "Trạng thái đang hiển thị trên trình duyệt không được dùng làm căn cứ duy nhất để xác nhận ghế hoặc giao dịch. Mọi thao tác có khả năng cạnh tranh đều được kiểm tra lại tại server trong transaction. Mỗi cặp ghế–suất chiếu có một bản ghi trạng thái; các bản ghi được khóa theo thứ tự ổn định khi giữ hoặc xác nhận ghế; trường phiên bản và ràng buộc duy nhất tạo thêm lớp phát hiện xung đột.",
        "Khi hai khách hàng cùng giữ một ghế, PostgreSQL chỉ cho một transaction nắm khóa ghi tại một thời điểm. Transaction hoàn tất trước chuyển ghế sang HOLD. Transaction còn lại tiếp tục sau đó, đọc trạng thái mới và nhận kết quả ghế không khả dụng. WebSocket giúp thông báo thay đổi sớm nhưng không thay thế vai trò kiểm soát của transaction và constraint.",
        "Các thao tác gửi lặp khi tạo booking hoặc khởi tạo payment được kiểm soát bằng cách sử dụng lại bản ghi PENDING phù hợp và chỉ mục duy nhất một phần. Callback, webhook và thao tác check-in áp dụng nguyên tắc lũy đẳng: dữ liệu mục tiêu được khóa, trạng thái kết thúc được kiểm tra trước khi tạo hiệu ứng nghiệp vụ, nhờ đó cùng một thông điệp không phát hành nhiều vé hoặc sử dụng vé nhiều lần.",
        "Scheduler sử dụng điều kiện về trạng thái, người giữ và thời hạn khi giải phóng ghế hoặc kết thúc đơn chờ. Chỉ transaction còn thỏa điều kiện mới cập nhật được dữ liệu. Retry chỉ được xem xét đối với lỗi kỹ thuật tạm thời; lỗi nghiệp vụ như ghế đã có người giữ được trả về ngay để người dùng chọn phương án khác.",
    ])
    academic_table(doc, "Bảng 3.10. Các tình huống cạnh tranh dữ liệu và cơ chế phòng vệ", ["Tình huống", "Nguy cơ", "Cơ chế phòng vệ", "Kết quả mong đợi"], [
        ["Hai khách cùng giữ một ghế", "Cùng sở hữu hoặc bán trùng ghế.", "Khóa bản ghi, kiểm tra trạng thái và ràng buộc duy nhất ghế–suất.", "Một yêu cầu thành công; yêu cầu còn lại nhận thông báo ghế không khả dụng."],
        ["Nhấn tạo đơn nhiều lần", "Sinh nhiều booking PENDING.", "Tái sử dụng đơn đang hoạt động và chỉ mục duy nhất một phần.", "Tối đa một đơn chờ thanh toán cho mỗi khách và suất chiếu."],
        ["Khởi tạo nhiều phương thức thanh toán", "Nhiều payment cùng chờ trên một đơn.", "Khóa booking và giới hạn payment PENDING.", "Chỉ một giao dịch chờ thanh toán còn hiệu lực."],
        ["Callback hoặc webhook lặp", "Tạo vé hoặc gửi thông báo nhiều lần.", "Khóa dữ liệu, kiểm tra trạng thái kết thúc, ràng buộc vé và payment event.", "Hiệu ứng nghiệp vụ chỉ được thực hiện một lần."],
        ["Hai thiết bị cùng soát vé", "Một vé được sử dụng hai lần.", "Khóa ticket và kiểm tra trạng thái ACTIVE.", "Một thao tác chuyển vé sang USED; thao tác còn lại nhận trạng thái đã sử dụng."],
        ["Scheduler chạy cùng thao tác người dùng", "Giải phóng hoặc xác nhận sai trạng thái.", "Cập nhật có điều kiện theo trạng thái, chủ sở hữu và thời hạn.", "Chỉ nhánh còn thỏa điều kiện được phép cập nhật."],
    ])

    doc.add_heading("3.11. Tiểu kết chương 3", level=2)
    body(doc, "Chương 3 đã phân tích bài toán, xác định tác nhân và yêu cầu, sau đó chuyển các yêu cầu thành kiến trúc mô-đun, mô hình lớp, trạng thái nghiệp vụ, cơ sở dữ liệu và RESTful API. Các quy trình xác thực, tìm rạp gần, quản lý lịch chiếu, giữ ghế, thanh toán, đồng bộ thời gian thực, soát vé và quản lý yêu cầu hoàn tiền được mô tả cùng những điều kiện nghiệp vụ tương ứng. Thiết kế xử lý đồng thời sử dụng nhiều lớp bảo vệ từ transaction, khóa bản ghi và cập nhật có điều kiện đến ràng buộc cơ sở dữ liệu; thiết kế định vị đồng thời duy trì luồng chọn thủ công và giới hạn dữ liệu vị trí được lưu. Chương 4 trình bày cách các thiết kế này được hiện thực trong front-end, back-end và được kiểm tra bằng các kịch bản phù hợp.")
def expand_chapter_four(doc, test_files: list[str]) -> None:
    doc.add_heading("4.2.5. Cấu hình khởi động và môi trường", level=3)
    body_many(doc, [
        "`CinemaBookingSystemApplication.java` là điểm khởi động của ứng dụng Spring Boot. `application.yaml` khai báo kết nối PostgreSQL, Flyway, JPA, thư điện tử, cache, cổng thanh toán, thời hạn đặt vé, scheduler, JWT, Google Identity và QR. Các biến môi trường tách thông tin bí mật khỏi mã nguồn; `.env.example` mô tả những tham số cấu hình cần thiết nhưng không chứa thông tin xác thực thật. Docker Compose định nghĩa PostgreSQL 15 và vùng lưu trữ dữ liệu cho môi trường phát triển cục bộ.",
        "JPA sử dụng `ddl-auto=validate`; vì vậy quá trình khởi động sẽ báo lỗi khi ánh xạ Entity không phù hợp với lược đồ thay vì tự động thay đổi bảng. Flyway thực thi các migration theo phiên bản trong `src/main/resources/db/migration`. Cấu hình cache tạo Caffeine CacheManager với danh sách vùng nhớ đệm xác định. `AsyncConfig` cung cấp executor riêng cho các tác vụ như gửi thư; `WebMvcConfig` cấu hình CORS và interceptor; `OpenApiConfig` cung cấp tài liệu Swagger phục vụ kiểm tra API.",
    ])

    doc.add_heading("4.2.6. Chuỗi xử lý bảo mật", level=3)
    body_many(doc, [
        "`SecurityConfig` xây dựng SecurityFilterChain theo mô hình stateless resource server. Các endpoint đăng ký, xác thực email, đặt lại mật khẩu, đăng nhập, dữ liệu công khai, callback hoặc webhook thanh toán và WebSocket handshake được cho phép theo HTTP method và đường dẫn cụ thể; phần còn lại yêu cầu người dùng đã xác thực. `JwtAuthenticationEntryPoint` và `JwtAccessDeniedHandler` chuẩn hóa phản hồi HTTP 401 và 403, trong khi `CustomJwtDecoder` xác minh token trước khi request được chuyển đến Controller.",
        "Method security được bật để Controller khai báo permission gần endpoint. Ví dụ BOOKING_CREATE bảo vệ hold/create, PAYMENT_CREATE bảo vệ initiation và TICKET_CHECKIN bảo vệ scanner. ApplicationInitConfig đồng bộ PermissionName và role mặc định; ADMIN nhận toàn bộ permission, STAFF và USER nhận tập giới hạn. StaffCinemaScopeService thực hiện data-level authorization ở service, bổ sung cho `@PreAuthorize`.",
        "AuthRateLimitService và SeatHoldRateLimitService dựa trên FixedWindowRateLimitService để hạn chế request nhạy cảm. Đây là rate limit trong memory của instance; khi scale nhiều instance cần store phân tán để quota nhất quán. AuthAuditService ghi event login/refresh/logout thất bại hoặc thành công, hỗ trợ phát hiện hành vi bất thường và giải thích sự cố.",
    ])

    doc.add_heading("4.2.7. Hiện thực nghiệp vụ đặt vé", level=3)
    body_many(doc, [
        "BookingServiceImpl là một trong các service trung tâm. Các method write được đánh dấu `@Transactional`; read method dùng readOnly khi phù hợp. Service không nhận userId từ client cho nghiệp vụ own mà lấy identity qua SecurityUtils. Giá ghế được tính từ ShowTime và Seat, không dùng amount do client tự cung cấp. Mapper tạo BookingResponse sau khi quan hệ cần thiết đã được fetch.",
        "SeatStatusRepository có query `JOIN FETCH` để tải sơ đồ ghế không phát sinh lazy query theo từng seat. `findForUpdateByShowtimeAndSeats` khóa pessimistic và sắp xếp row. `confirmHeldSeatsForBooking` là bulk conditional update yêu cầu status HOLD, đúng holdBy và holdUntil còn hiệu lực; service kiểm tra số dòng update bằng số ghế. Các query release expired dùng native SQL theo batch, tránh tải toàn bộ entity vào memory.",
        "BookingRepository áp dụng chiến lược phân trang hai bước cho màn hình phức tạp: query ID theo bộ lọc/count riêng, sau đó fetch toàn bộ details cho tập ID. Cách này tránh pagination sai khi JOIN FETCH collection và giảm N+1. Query admin/staff có nhánh cinemaIds để thực thi scope tại database thay vì tải tất cả rồi lọc trong Java.",
    ])

    doc.add_heading("4.2.8. Hiện thực nghiệp vụ thanh toán", level=3)
    body_many(doc, [
        "PaymentServiceImpl nhận danh sách PaymentGateway và lập bản đồ theo PaymentMethod. Pattern này giúp Controller không chứa `if/else` provider. Initiation khóa booking, kiểm tra owner/state/expiry, xử lý payment PENDING hiện có và gọi gateway tạo dữ liệu redirect hoặc QR. Repository và partial index V14 là lớp chống double initiation.",
        "Callback/webhook được chuẩn hóa thành quá trình xác minh, tìm payment/booking, khóa dữ liệu, kiểm tra idempotency, cập nhật trạng thái và kích hoạt side effect sau thành công. PaymentEventServiceImpl ghi event với propagation phù hợp để duy trì audit trong các tình huống lỗi đã được xử lý. Monitoring summary và reconciliation query trả read model cho AdminPaymentPage thay vì để frontend tự so sánh trạng thái nhiều bảng.",
        "RefundServiceImpl quản lý vòng đời của yêu cầu hoàn tiền. Việc tách thực thể Refund giúp lưu lịch sử yêu cầu, trạng thái xử lý, mã tham chiếu và nguyên nhân thất bại. Các thao tác hoàn tất hoặc đánh dấu thất bại chỉ ghi nhận kết quả do người vận hành cung cấp, không gọi API hoàn tiền của nhà cung cấp. Cách tổ chức này tránh chuyển Payment trực tiếp sang REFUNDED mà không xác định được chủ thể yêu cầu và thời điểm xử lý; tích hợp hoàn tiền tự động được xếp vào hướng phát triển.",
    ])

    doc.add_heading("4.2.9. Email, mã QR, cập nhật thời gian thực và tác vụ định kỳ", level=3)
    body_many(doc, [
        "EmailServiceImpl dựng email xác thực, reset password và vé điện tử. Spring Mail dùng SMTP cấu hình qua env; template được tạo trong Java/service sau khi các file HTML static thừa đã được loại bỏ khỏi client flow. Gửi email là side effect không nên quyết định commit payment: nếu email lỗi, booking đã thanh toán không được rollback về thất bại; cần log/retry theo chính sách vận hành.",
        "TicketQrCodeService ký payload bằng secret riêng hoặc fallback cấu hình được xác định, còn QrCodeImageService dùng ZXing tạo bitmap. QR không chứa toàn bộ dữ liệu vé tin cậy; server giải mã/xác minh rồi truy vấn Ticket hiện tại. Nhờ đó ticket đã USED/CANCELLED không thể được sử dụng chỉ bằng ảnh QR cũ.",
        "SeatStatusPublisher phát event STOMP. HoldExpireScheduler quét projection theo batch, update conditional và publish sau commit. PendingBookingExpireScheduler terminalize đơn/payment quá hạn và release seat/promotion reservation. ShowtimeStatusSyncScheduler chuyển UPCOMING/ONGOING/ENDED theo thời gian. TokenCleanupTask dọn refresh/invalidated token hết hạn. Các task đều cần idempotent vì lịch chạy có thể lặp hoặc ứng dụng restart.",
    ])

    doc.add_heading("4.3.4. Hành trình tra cứu và mua vé của khách hàng", level=3)
    body_many(doc, [
        "HomePage tải phim nổi bật và khu vực lịch chiếu; MovieCarousel hỗ trợ điều hướng ngang trên máy tính và thao tác cảm ứng trên thiết bị di động. RegionalShowtimeBrowser kết hợp lựa chọn thành phố, rạp, ngày và danh sách phim - suất chiếu trong một quy trình thống nhất. Khi người dùng nhấn Gần tôi, component mới xin quyền Browser Geolocation [17], chọn rạp gần nhất và sắp các thành phố, rạp theo khoảng cách. CinemaMapPage dùng cùng utility Haversine [18], đồng thời đặt dấu vị trí, di chuyển bản đồ Leaflet và đếm rạp trong bán kính 10 km.",
        "Nếu bị từ chối quyền, timeout, trình duyệt không hỗ trợ hoặc không có rạp trong bán kính 10 km, giao diện không khóa hành trình: thông báo lỗi hoặc trạng thái rỗng theo ngữ cảnh được hiển thị, còn bộ chọn thủ công và danh sách rạp gần nhất vẫn hoạt động. Tọa độ rạp sai bị bỏ qua khi tạo marker; tọa độ chính xác của người dùng chỉ nằm trong state, không được lưu vào localStorage hay cơ sở dữ liệu. Thành phố ưu tiên được lưu riêng để hỗ trợ lần truy cập sau.",
        "MovieDetailPage hiển thị metadata, ngôn ngữ, phụ đề, quốc gia và lịch theo thành phố. CinemaDetailPage nhóm suất theo phim tại một rạp. SeatSelectionPage là màn hình giao dịch: tải seat map, hiển thị loại/giá/trạng thái, giữ ghế, lắng nghe WebSocket và điều hướng checkout. Countdown phía client phục vụ UX nhưng expiry phía server mới có giá trị quyết định.",
        "CheckoutPage cho phép chọn promotion trước khi initiation, ưu tiên QR ngân hàng và giữ VNPay là lựa chọn redirect. Khi QR đã được tạo, thay đổi amount yêu cầu tạo payment mới để tránh mismatch. PaymentResultPage đặt viewport/scroll phù hợp và truy vấn trạng thái server. MyBookingsPage phân loại vé hợp lệ và đơn đã đặt; hành động chọn lại ghế bị ẩn với suất đã kết thúc hoặc bị hủy theo trạng thái response. TicketDetailPage hiển thị mỗi ghế một QR và hỗ trợ lưu ảnh vé có thông tin rạp/phòng/địa chỉ/thành phố.",
    ])

    doc.add_heading("4.3.5. Giao diện quản trị và vận hành rạp", level=3)
    body_many(doc, [
        "AdminLayout tạo sidebar theo permission nên STAFF chỉ thấy menu được phép. Dashboard dùng Recharts để biểu diễn doanh thu và cơ cấu đơn/vé; analytics API cung cấp summary, daily/monthly revenue, top movie, showtime stats và CSV. Các trang quản trị phim, rạp, phòng-ghế, suất, người dùng, khuyến mãi, booking, payment và audit có filter, pagination và modal/form tương ứng.",
        "AdminRoomSeatPage hỗ trợ chọn thành phố-rạp-phòng, tạo phòng và bulk generate sơ đồ ghế theo template. AdminShowtimePage tự giới hạn rạp đối với STAFF, kiểm tra thời gian và hủy suất theo policy. AdminUserPage tách tab người dùng/nhân viên/admin; assignment rạp được quản lý cho staff. AdminPaymentPage có giao dịch, yêu cầu refund, reconciliation và event log, đồng thời lọc theo gateway, trạng thái, thành phố, rạp và thời gian.",
        "StaffAssignedCinemasPage cho nhân viên biết nơi mình phụ trách. StaffTicketScannerPage chọn cinema và showtime trước khi mở camera hoặc đọc file, tránh tiêu thụ vé sai ngữ cảnh. Component scanner phải khởi tạo/cleanup camera đúng vòng đời để không tạo hai video element hoặc lỗi removeChild. Thông báo lỗi dịch business code thành tiếng Việt dễ hiểu thay vì hiển thị `You do not have permission` hoặc exception kỹ thuật.",
    ])

    doc.add_heading("4.3.6. Khả năng đáp ứng, tiếp cận và phục hồi giao diện", level=3)
    body_many(doc, [
        "Layout sử dụng breakpoint, grid/flex, overflow container và kích thước ổn định để hoạt động trên mobile, tablet và desktop. Các danh sách ngang cho phép scrollbar hoặc carousel có chỉ dấu để người dùng nhận biết còn nội dung. Form admin dùng modal có vùng cuộn riêng, tránh vượt chiều cao màn hình. QR và seat grid có kích thước responsive nhưng không làm nội dung chồng lấn.",
        "Lucide icon được dùng thống nhất; button có trạng thái disabled/loading; input hiển thị lỗi và màu dark mode có độ tương phản. Các thao tác nguy hiểm như hủy suất, hủy booking hoặc ghi nhận kết quả yêu cầu refund cần xác nhận và giải thích hậu quả. ErrorBoundary toàn app bắt lỗi render nghiêm trọng, trong khi Toast xử lý feedback nghiệp vụ. API errors được chuẩn hóa để page không phụ thuộc message tiếng Anh thô từ server.",
        "Khả năng truy cập vẫn cần được kiểm tra bằng bàn phím, screen reader và công cụ contrast trước khi phát hành. Source có semantic button/input ở nhiều nơi nhưng báo cáo không khẳng định đạt chuẩn WCAG nếu chưa có audit. Đây là ví dụ phân biệt thiết kế responsive đã hiện thực với chứng nhận accessibility chưa được đo.",
    ])

    doc.add_heading("4.6.1. Phạm vi kiểm thử phía máy chủ", level=3)
    body(doc, "Các lớp kiểm thử phía máy chủ tập trung vào những khu vực có rủi ro nghiệp vụ cao như vòng đời phiên đăng nhập, quyền sở hữu đơn đặt vé, cạnh tranh giữ ghế, lịch chiếu, callback thanh toán, giới hạn tần suất và chữ ký mã QR. Trong lần thực thi ngày 17/08/2026, Maven đã chạy 53 trường hợp kiểm thử và toàn bộ đều đạt; không có trường hợp thất bại, lỗi hoặc bị bỏ qua.")
    academic_table(doc, "Bảng 4.8. Kiểm thử phía máy chủ", ["Tệp kiểm thử", "Loại/đối tượng kiểm tra"], [
        [name, _test_purpose(name)] for name in test_files
    ], font_size=8.5)
    body_many(doc, [
        "Integration test phù hợp với booking/payment vì transaction, lock, unique index và security filter khó được chứng minh bằng mock đơn lẻ. Testcontainers tạo PostgreSQL tạm thời theo container; máy không chạy Docker có thể không chạy được nhóm test phụ thuộc container. Unit test vẫn chạy không cần DB nếu không khởi tạo container. Báo cáo phải ghi rõ điều kiện này khi trình diễn.",
        "Mockito hữu ích để cô lập service nhỏ hoặc gateway, nhưng mock Repository không thay thế test constraint/lock thật. Do đó chiến lược hợp lý là unit test cho thuật toán ký QR, rate limit, mapping payload; integration test cho auth rotation, ownership, booking race và callback idempotency; controller test cho validation/exception/security response.",
    ])

    doc.add_heading("4.6.2. Kịch bản kiểm thử API bằng Postman", level=3)
    academic_table(doc, "Bảng 4.9. Bộ test API đề xuất", ["Mã", "Request", "Dữ liệu", "Kết quả kỳ vọng"], [
        ["API-01", "POST /auth/token", "Sai password nhiều lần", "401/business code; sau ngưỡng nhận rate limit, không lộ user tồn tại."],
        ["API-02", "POST /auth/refresh", "Refresh token hợp lệ rồi dùng lại token cũ", "Lần đầu rotate; lần dùng lại bị từ chối/revoke theo policy."],
        ["API-03", "POST /bookings/hold", "seatIds không phải mảng hoặc rỗng", "400 validation, ApiResponse ổn định."],
        ["API-04", "POST /bookings/hold", "Hai user cùng showtimeId/seatId", "Một thành công, một lỗi seat unavailable; DB chỉ có một HOLD."],
        ["API-05", "POST /bookings", "Hold hết hạn", "Không tạo booking SUCCESS/PENDING sai; ghế được release theo flow."],
        ["API-06", "PATCH booking promotion", "Mã hết hạn/quota hết/đơn dưới tối thiểu", "Business error; total không đổi."],
        ["API-07", "POST payment initiate", "Double request/cổng khác nhau cùng booking", "Tối đa một PENDING payment."],
        ["API-08", "GET VNPay callback", "Checksum sai hoặc amount mismatch", "Không success; event ghi failure."],
        ["API-09", "POST SePay webhook", "Secret sai, amount/content sai", "Bị từ chối; không tạo ticket."],
        ["API-10", "POST ticket check-in", "Sai cinema/showtime", "Ticket vẫn ACTIVE, message đúng ngữ cảnh."],
        ["API-11", "POST ticket check-in", "Cùng QR gửi hai lần", "Lần đầu USED; lần sau báo đã dùng kèm thông tin."],
        ["API-12", "GET admin payment", "STAFF thay cinemaId ngoài scope", "403/business forbidden hoặc dữ liệu rỗng theo service policy; không rò dữ liệu."],
    ], font_size=8.3)

    doc.add_heading("4.6.3. Kịch bản kiểm thử giao diện", level=3)
    academic_table(doc, "Bảng 4.10. Các trường hợp kiểm thử giao diện theo hành trình", ["Màn hình", "Tình huống", "Điều cần quan sát"], [
        ["Đăng ký", "Username/email trùng; password yếu; confirm khác", "Lỗi tiếng Việt tại trường, không reload trang, không gửi request sai."],
        ["Đăng nhập", "Sai credential; dark mode; refresh session", "Text nhìn rõ, loading chống double click, route sau login đúng role."],
        ["Trang chủ", "Mobile/tablet/desktop; poster lỗi; danh sách dài", "Card đều chiều cao, nút không lệch, fallback ảnh, scroll/carousel mượt."],
        ["Bản đồ - định vị thành công", "Cho phép vị trí; rạp trong/ngoài bán kính 10 km; tọa độ rạp sai", "Haversine và thứ tự khoảng cách đúng; đánh dấu rạp gần; bỏ qua marker sai nhưng không làm hỏng danh sách."],
        ["Bản đồ - luồng thay thế", "Từ chối quyền; timeout; không hỗ trợ; không có rạp trong 10 km", "Thông báo đúng ngữ cảnh, vẫn chọn thủ công và vẫn hiển thị các rạp gần nhất; không lưu tọa độ chính xác."],
        ["Chọn ghế", "Hai browser giữ cùng ghế; hold hết hạn", "WebSocket đổi màu không cần refresh; countdown không rung; lỗi cạnh tranh dễ hiểu."],
        ["Checkout", "Áp/bỏ mã trước và sau khi tạo QR", "Amount server và QR đồng bộ; UI yêu cầu tạo QR mới khi cần."],
        ["Kết quả", "Callback chậm, reload, back/forward", "Trang lấy trạng thái server và scroll vị trí phù hợp; không tạo payment mới."],
        ["Vé của tôi", "Success, cancelled showtime, expired, refund pending", "CTA đúng policy; hiển thị yêu cầu đang chờ xử lý và không hiện chọn lại ghế khi không hợp lệ."],
        ["Scanner", "Camera/file, QR mờ, camera denied, quét lặp", "Một preview, cleanup đúng, lỗi ổn định, ticket không dùng sai ngữ cảnh."],
        ["Admin", "Filter/pagination/modal dài/dark mode", "Không mất filter, form cuộn được, permission menu đúng, không vỡ layout."],
    ], font_size=8.3)

    doc.add_heading("4.7.1. Kiểm thử tải đối với API tra cứu", level=3)
    body_many(doc, [
        "Thread Group có thể tăng dần 10, 25, 50 và 100 virtual user theo tài nguyên máy kiểm thử. Sampler gọi home feed, movie list, cinema map và showtime query với dữ liệu phân bố. HTTP Header Manager đặt Accept và token khi cần; CSV Data Set Config cung cấp thành phố/movieId khác nhau để tránh chỉ đo một cache key. Listener nặng như View Results Tree chỉ dùng lúc debug, không dùng khi chạy tải chính.",
        "Chỉ số cần thu thập gồm throughput, median/p90/p95/p99 response time, error rate, số connection DB, CPU/memory và slow query. Báo cáo hiện không điền số vì chưa có kết quả .jmx thực tế. Acceptance threshold phải được xác định theo môi trường bảo vệ hoặc yêu cầu giảng viên trước khi chạy.",
    ])
    doc.add_heading("4.7.2. Kiểm thử đồng thời khi giữ ghế", level=3)
    numbered_list(doc, [
        "Chuẩn bị một showtime sắp chiếu và một seat đang AVAILABLE; lấy token của 50 user test khác nhau.",
        "CSV Data Set Config phân phối token nhưng giữ nguyên showtimeId và seatId cho toàn bộ thread.",
        "Synchronizing Timer giữ các thread rồi giải phóng gần như cùng thời điểm.",
        "HTTP Request gửi POST `/api/v1/bookings/hold` với `seatIds` là mảng JSON.",
        "Response Assertion phân loại đúng một response thành công; các response còn lại phải là lỗi nghiệp vụ seat unavailable hoặc rate limit nếu cấu hình giới hạn can thiệp.",
        "JDBC/SQL kiểm tra cuối: đúng một seat_status HOLD với hold_by xác định; không có nhiều booking/payment PENDING ngoài invariant.",
        "Lặp kịch bản nhiều vòng sau khi release hold để quan sát deadlock, timeout và tính ổn định, đồng thời thu log PostgreSQL/Spring.",
    ])
    body(doc, "Kết quả kỳ vọng là chỉ một người dùng có thể giữ ghế trong mỗi thời điểm kiểm thử. Báo cáo không đưa ra số liệu về thông lượng hoặc thời gian đáp ứng khi chưa thực hiện phép đo. Nếu cơ chế giới hạn tần suất chặn phần lớn yêu cầu trước khi đến bước khóa dữ liệu, cần tách cấu hình kiểm thử hoặc điều chỉnh ngưỡng phù hợp để đánh giá đúng cơ chế xử lý đồng thời, sau đó khôi phục cấu hình vận hành.")

    doc.add_heading("4.8.1. Đánh giá theo các thuộc tính chất lượng", level=3)
    academic_table(doc, "Bảng 4.11. Đánh giá theo thuộc tính chất lượng", ["Thuộc tính", "Điểm đã có", "Khoảng trống cần tiếp tục"], [
        ["An toàn thông tin", "JWT rotation, RBAC, phạm vi rạp của nhân viên, validation, audit, rate limit, signed QR.", "Secret manager, TLS, reverse proxy, cookie trong môi trường vận hành và security scan chưa được kiểm chứng."],
        ["Độ tin cậy", "Transaction, kiểm tra trạng thái, scheduler, event log, callback lũy đẳng.", "Retry email/nhà cung cấp và khôi phục sau thảm họa cần kế hoạch vận hành."],
        ["Performance", "Index, projection, pagination, cache catalog, batch update.", "Cần SQL profile/JMeter thật và capacity planning."],
        ["Maintainability", "Layered package, DTO/mapper, gateway adapter, Flyway, tests.", "Hai cặp route alias cần deprecate; tiếp tục giảm class/service quá lớn."],
        ["Scalability", "Stateless access API và cấu hình externalized.", "Caffeine/simple broker/rate limit local cần giải pháp phân tán khi nhiều instance."],
        ["Observability", "PaymentEvent, auth/admin audit, logging và monitoring summary.", "Metrics/tracing/alerting tập trung chưa được xác minh."],
        ["Usability", "Responsive flows, location, QR file/camera, thông báo tiếng Việt.", "Cần usability test và accessibility audit có người dùng thật."],
    ])

    doc.add_heading("4.9. Tiểu kết chương 4", level=2)
    body(doc, "Qua các nội dung của chương, có thể thấy luồng đặt vé được hình thành từ nhiều thành phần phối hợp thay vì một thao tác lưu đơn đơn lẻ. Phía máy chủ chịu trách nhiệm xác thực, giữ ghế, xử lý kết quả thanh toán và phát hành vé; phía giao diện tổ chức hành trình sử dụng và phản ánh trạng thái do máy chủ xác nhận. Các kiểm thử hiện có được trình bày tách biệt với kịch bản Postman và JMeter chưa thực hiện, nhờ đó kết quả của đề tài không bị lẫn với số liệu dự kiến. Thông tin triển khai thực tế vẫn cần được bổ sung sau khi sinh viên hoàn tất môi trường công bố hệ thống.")


def _test_purpose(path: str) -> str:
    name = path.rsplit("/", 1)[-1]
    mapping = {
        "AuthenticationServiceIntegrationTest.java": "Integration: login, refresh token/session và các quy tắc xác thực.",
        "BookingPaymentSecurityIntegrationTest.java": "Integration: ownership và authorization booking/payment.",
        "BookingWorkflowIntegrationTest.java": "Integration: hold, booking, expiry và lifecycle ghế.",
        "PaymentCallbackIntegrationTest.java": "Integration: callback/webhook, idempotency và trạng thái payment/booking.",
        "GlobalExceptionHandlerIntegrationTest.java": "Web integration: định dạng response exception/validation/security.",
        "UserManagementIntegrationTest.java": "Integration: quản lý user, role, block/delete/scope.",
        "HomeShowtimeFeedIntegrationTest.java": "Integration: feed phim-suất public và bộ lọc thời gian/trạng thái.",
        "TicketQrCodeServiceTest.java": "Unit: tạo, ký và xác minh payload QR ticket.",
        "QrCodeImageServiceTest.java": "Unit: tạo ảnh QR từ payload.",
        "FixedWindowRateLimitServiceTest.java": "Unit: cửa sổ rate limit và thời điểm reset.",
        "SePayPaymentGatewayTest.java": "Unit: dữ liệu QR/webhook/signature của SePay gateway.",
        "CinemaBookingSystemApplicationTests.java": "Smoke test: Spring application context.",
    }
    return mapping.get(name, "Test source tồn tại; đọc method test để xác định từng assertion cụ thể.")


def expand_chapter_five(doc) -> None:
    doc.add_heading("5.4. Nhận xét từ quá trình thực hiện", level=2)
    body_many(doc, [
        "Thứ nhất, giao diện không thể được xem là nguồn quyết định trạng thái trong nghiệp vụ giao dịch. Màu ghế, bộ đếm thời gian và thao tác xác nhận chuyển khoản chỉ hỗ trợ tương tác; quyền sở hữu ghế, thời hạn giữ chỗ, số tiền và trạng thái vé phải được phía máy chủ xác minh trong giao dịch cơ sở dữ liệu. WebSocket hỗ trợ cập nhật kịp thời nhưng không thay thế ràng buộc dữ liệu; tương tự, endpoint callback công khai vẫn phải thực hiện xác minh chữ ký và dữ liệu giao dịch.",
        "Thứ hai, lược đồ cơ sở dữ liệu là một bộ phận của thiết kế hệ thống. Ràng buộc duy nhất và CHECK constraint không chỉ mô tả dữ liệu mà còn góp phần ngăn ngừa các tình huống tranh chấp mà kiểm tra tuần tự ở tầng ứng dụng không thể loại bỏ hoàn toàn. Các chỉ mục cần được xây dựng từ nhu cầu truy vấn và tác vụ định kỳ; migration cần được quản lý theo phiên bản thay vì chỉnh sửa thủ công trên từng môi trường.",
        "Thứ ba, tích hợp cổng thanh toán cần bảo đảm tính lũy đẳng (idempotency) và khả năng truy vết. Trình duyệt có thể tải lại trang, nhà cung cấp có thể gửi lại thông báo và webhook có thể đến muộn. Vì vậy, PaymentEvent, lịch sử chuyển trạng thái, đối soát và quy trình quản lý yêu cầu hoàn tiền cung cấp căn cứ cần thiết để xác định nguyên nhân khi một giao dịch chưa thể phát hành vé.",
        "Các nội dung trong báo cáo được đối chiếu với thành phần hiện thực của hệ thống nhằm hạn chế mô tả không còn phù hợp. Trước khi nộp, sinh viên cần hoàn thiện các thông tin cá nhân, hình minh họa, số liệu kiểm thử đã thực hiện và nội dung triển khai thực tế theo yêu cầu của cơ sở đào tạo.",
    ])
    doc.add_heading("5.5. Kết luận chung", level=2)
    body(doc, "CinemaBooking đáp ứng mục tiêu xây dựng hệ thống đặt vé xem phim có các chức năng phục vụ khách hàng, nhân viên rạp và quản trị viên; đồng thời hỗ trợ xử lý đồng thời khi chọn ghế, tích hợp nhiều phương thức thanh toán, phát hành vé QR, cập nhật thời gian thực và ghi nhận nhật ký hoạt động. Kiến trúc hiện tại phù hợp với mô hình modular monolith (nguyên khối theo mô-đun), sử dụng chung PostgreSQL và có khả năng mở rộng theo nhu cầu. Tuy nhiên, việc triển khai ở môi trường vận hành thực tế, giám sát tập trung, kiểm thử tải, rà soát an toàn thông tin và bổ sung các thành phần phân tán vẫn là những nội dung cần tiếp tục thực hiện khi quy mô hệ thống tăng. Trong phạm vi khóa luận, kết quả thể hiện quá trình phân tích bài toán, thiết kế dữ liệu, xây dựng ứng dụng web toàn diện và xử lý các tình huống đồng thời đặc trưng của hệ thống giao dịch.")
